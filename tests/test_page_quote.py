"""Объём в страницах и смета: счёт знаков, норма, цена.

Проверяется то, чем эта функция может СОВРАТЬ в деньгах:

  1. знаки считаются как считает их Word — с пробелами, но знак абзаца
     не знак и наш разделитель кусков тоже (иначе книга в 2670 абзацев
     дорожала бы на полторы страницы из-за склейки);
  2. норма страницы берётся у языка ИСХОДНИКА, и откуда взята — сказано
     словом (`source`): догадка по умолчанию не должна выглядеть расчётом;
  3. в .docx не считается то, за что платить нельзя, — скрытый текст,
     вычисляемые поля (номер страницы в оглавлении) и вложенные надписи,
     чьи абзацы уже посчитаны отдельно;
  4. цена не задана — это `null` и отказ считать сумму, а НЕ ноль:
     бесплатный перевод в смете хуже отсутствующей сметы;
  5. цены правит владелец организации, переводчику остаётся чтение;
  6. смета по файлу и смета по сегментам проекта — РАЗНЫЕ величины,
     и каждая называет себя (`basis`);
  7. история смет ЗАМОРОЖЕНА: смена прайса старую оферту не пересчитывает,
     повторный расчёт того же файла не плодит заказ, а судьбу сметы
     (выставлена/оплачена) решает владелец.

Ни одного вызова модели, состояние на диск не пишется.
"""
import os
import sys

os.environ["APP_PASSWORD"] = "boot-password-1"
os.environ["AUTHORITY_CORPUS"] = "0"
os.environ["OPENAI_API_KEY"] = "test-key"
sys.path.insert(0, "backend")
import io                                     # noqa: E402
import json                                   # noqa: E402
import main                                   # noqa: E402
import textcount                              # noqa: E402
from starlette.testclient import TestClient   # noqa: E402

main.save_state = lambda *a, **k: None
main.STATE["users"], main.STATE["tenants"] = [], []
main._SESSIONS.clear()
main._LOGIN_FAILS.clear()
fail = []


def check(cond, label):
    print(("  OK   " if cond else "  FAIL ") + label)
    if not cond:
        fail.append(label)


print("=== 1. Счёт знаков ===")
c1 = textcount.count_blocks(["Привет мир.", "", "Вторая  строка тут."])
check(c1["chars"] == 29, "знаки как в Word: 11 + 18, знак абзаца не считается (%d)" % c1["chars"])
check(c1["charsNoSpaces"] == 26, "без пробелов — 26 (%d)" % c1["charsNoSpaces"])
check(c1["words"] == 5, "слов 5 (%d)" % c1["words"])
one_line = textcount.count_blocks(["Привет мир. Вторая строка тут."])
check(one_line["chars"] == c1["chars"] + 1,
      "одним абзацем ровно на 1 знак больше — на настоящий пробел, а не на наш разделитель")
rep = textcount.count_blocks(["Заголовок", "Заголовок", "Текст"])
check(rep["repeatBlocks"] == 1 and rep["chars"] == len("ЗаголовокЗаголовокТекст"),
      "повтор посчитан отдельной строкой, но из объёма НЕ вычтен")

print("=== 2. Норма страницы ===")
check(textcount.norm_for("RU")["chars"] == 1800, "RU — 1800 знаков (канон СНГ)")
check(textcount.norm_for("EN")["chars"] == 1500, "EN — 1500 знаков")
check(textcount.norm_for("ZH")["chars"] == 400 and textcount.norm_for("ZH")["spaceless"],
      "ZH — 400 знаков, письмо без пробелов")
check(textcount.norm_for("RU")["source"] == "table", "источник нормы назван: table")
check(textcount.norm_for("RU", {"RU": 1667})["chars"] == 1667
      and textcount.norm_for("RU", {"RU": 1667})["source"] == "tenant",
      "переопределение организации сильнее таблицы и помечено tenant")
unknown = textcount.norm_for("QQ")
check(unknown["source"] == "default" and unknown["basis"] == "assumed",
      "языка нет в таблице — число по умолчанию помечено как догадка")
check(all(r.get("chars") for r in textcount.norms()["rows"].values()),
      "у каждой строки таблицы есть норма")

print("=== 3. Страницы и округление ===")
p = textcount.pages_of(1800, 1800, min_pages=1, round_to=0.1)
check(p["exact"] == 1.0 and p["billed"] == 1.0, "ровно норма — одна страница, а не 1.1")
p = textcount.pages_of(1801, 1800, min_pages=1, round_to=0.1)
check(p["billed"] == 1.1, "знак сверх нормы — округление вверх до шага (%s)" % p["billed"])
p = textcount.pages_of(200, 1800, min_pages=1, round_to=0.1)
check(p["billed"] == 1.0, "минимальный заказ поднимает 0.11 стр. до 1")
p = textcount.pages_of(200, 1800, min_pages=0, round_to=0)
check(p["billed"] == p["exact"], "без минимума и без округления — как есть")

print("=== 4. Форматы: чего не считаем и почему ===")
try:
    textcount.extract("scan.tiff", b"II*\x00garbage")
    check(False, "неизвестное расширение обязано быть отказом")
except textcount.Unsupported as e:
    check("tiff" in str(e) or "Формат" in str(e), "неизвестный формат — Unsupported: %s" % e)
try:
    textcount.measure("empty.txt", b"   \n\n  ", "RU")
    check(False, "пустой текст обязан быть отказом")
except textcount.Unsupported:
    check(True, "ноль знаков — отказ, а не смета на ноль")
try:
    textcount.extract("big.txt", b"x" * (textcount.MAX_BYTES + 1))
    check(False, "файл сверх потолка обязан быть отказом")
except textcount.TooBig:
    check(True, "файл больше потолка — TooBig (413), воркер один")
try:
    textcount.extract("old.doc", b"\xd0\xcf\x11\xe0")
    check(False, ".doc обязан быть отказом")
except textcount.Unsupported:
    check(True, "старый .doc — отказ с советом пересохранить")
html = textcount.measure("page.html", "<p>Раз два</p><script>var x=1;</script>".encode("utf-8"), "RU")
check(html["counts"]["chars"] == len("Раз два"), "теги и скрипты в оплату не идут (%d)" % html["counts"]["chars"])

print("=== 5. .docx: считается ровно то, куда встанет перевод ===")
try:
    from docx import Document
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls, qn

    doc = Document()
    p = doc.add_paragraph("Видимый текст")
    r = p.add_run("СКРЫТОЕ")                       # скрытый текст: его никто не увидит
    r.font.hidden = True
    p2 = doc.add_paragraph()
    p2._p.append(parse_xml('<w:fldSimple %s w:instr=" PAGE "><w:r><w:t>42</w:t></w:r></w:fldSimple>'
                           % nsdecls("w")))        # номер страницы считает Word
    p2.add_run("Хвост")
    p3 = doc.add_paragraph("Абзац")
    p3._p.append(parse_xml('<w:p %s><w:r><w:t>Надпись</w:t></w:r></w:p>' % nsdecls("w")))
    doc.sections[0].header.paragraphs[0].text = "Колонтитул"
    buf = io.BytesIO()
    doc.save(buf)
    raw = buf.getvalue()

    blocks = main._docx_bill_paragraphs(raw)
    joined = " ".join(blocks)
    check("СКРЫТОЕ" not in joined, "скрытый текст в счёт не идёт")
    check("42" not in joined, "номер страницы (вычисляемое поле) в счёт не идёт")
    check("Колонтитул" in joined, "колонтитул считается — его переводят")
    check(joined.count("Надпись") == 1, "вложенный абзац посчитан ОДИН раз, а не дважды")
    m = textcount.measure("t.docx", raw, "RU", docx_paragraphs=main._docx_bill_paragraphs)
    check(m["kind"] == "docx" and m["counts"]["chars"] > 0, "docx считается через разбор main.py")
    zip_blocks = textcount.extract("t.docx", raw)["blocks"]     # запасной разбор по XML
    check(any("Видимый текст" in b for b in zip_blocks), "запасной разбор .docx без python-docx работает")
except ImportError:
    check(True, "python-docx не установлен — проверка docx пропущена")

print("=== 6. Цена: не задана — это отказ, а не ноль ===")
card = {**main.PRICING_DEFAULTS}
check(main._rate_for(card, "RU", "EN")["price"] is None, "пустая карточка не даёт цены")
check(main._quote_of({"chars": 1800, "charsNoSpaces": 1500, "words": 250, "blocks": 1,
                      "repeatBlocks": 0, "repeatChars": 0}, "RU", "EN", card, "file")["total"] is None,
      "без цены суммы нет (не ноль)")
card = {**card, "default": 10.0, "rates": [{"src": "RU", "tgt": "EN", "price": 12.5}]}
check(main._rate_for(card, "RU", "EN") == {"price": 12.5, "source": "pair"}, "строка пары сильнее общей цены")
check(main._rate_for(card, "DE", "EN") == {"price": 10.0, "source": "default"}, "пары нет — общая цена")
check(main._money(0.1, 12.30) == 1.23, "деньги считаются десятичными: 0.1 × 12.30 = 1.23 (%s)"
      % main._money(0.1, 12.30))

print("=== 7. Эндпоинты и права ===")
c = TestClient(main.app)
H = lambda t: {"Authorization": "Bearer " + t}
A = c.post("/api/auth/login", json={"login": "admin", "password": "boot-password-1"}).json()["token"]
c.post("/api/admin/users", headers=H(A),
       json={"login": "tr", "password": "tr-pass-12345", "role": "translator"})
T = c.post("/api/auth/login", json={"login": "tr", "password": "tr-pass-12345"}).json()["token"]

check(main._owner_only("POST", "/api/pricing") and not main._owner_only("GET", "/api/pricing"),
      "цены правит владелец, читают все")
check(not main._is_paid("POST", "/api/quote") and not main._is_paid("GET", "/api/pricing"),
      "смета бесплатна: вызовов модели в ней нет")
check(c.post("/api/pricing", headers=H(T), json={"default": 5}).status_code == 403,
      "переводчику прайс править нельзя")
r = c.get("/api/pricing", headers=H(T))
check(r.status_code == 200 and len(r.json()["norms"]["rows"]) > 50,
      "переводчик читает карточку и таблицу норм")
r = c.post("/api/pricing", headers=H(A), json={
    "currency": "USD", "default": 10, "minPages": 1, "roundTo": 0.1,
    "rates": [{"src": "RU", "tgt": "EN", "price": 12.5}], "norms": {"RU": 1800}})
check(r.status_code == 200 and r.json()["pricing"]["rates"][0]["price"] == 12.5, "владелец записал прайс")
check(c.post("/api/pricing", headers=H(A), json={"currency": "доллар"}).status_code == 400,
      "валюта не по ISO — 400")
check(c.post("/api/pricing", headers=H(A), json={"norms": {"RU": 5}}).status_code == 400,
      "норма в 5 знаков — 400 (опечатка дороже отказа)")
check(c.post("/api/pricing", headers=H(A), json={
    "rates": [{"src": "RU", "tgt": "EN", "price": 1}, {"src": "ru", "tgt": "en", "price": 2}]
}).status_code == 400, "дубль пары в прайсе — 400: две цены на одну работу")

body = ("Текст на восемнадцать сотен знаков. " * 60).encode("utf-8")
r = c.post("/api/quote", headers=H(A), files={"file": ("doc.txt", body, "text/plain")},
           data={"src": "RU", "tgt": "EN"})
q = r.json()
check(r.status_code == 200 and q["counts"]["chars"] == len(("Текст на восемнадцать сотен знаков. " * 60).strip()),
      "смета по файлу: знаки посчитаны (%s)" % q["counts"]["chars"])
check(q["basis"] == "file" and q["norm"]["chars"] == 1800 and q["rate"]["price"] == 12.5,
      "норма языка исходника и цена пары применены")
check(q["total"] == main._money(q["pages"]["billed"], 12.5) and "÷" in q["formula"],
      "сумма сходится с расчётом, и расчёт показан строкой")
r = c.post("/api/quote", headers=H(A), files={"file": ("x.tiff", b"II*\x00", "image/tiff")},
           data={"src": "RU", "tgt": "EN"})
check(r.status_code == 415, "нечитаемый формат — 415, а не смета на ноль (%d)" % r.status_code)

pid = c.post("/api/projects", headers=H(A), json={"title": "Q", "src": "RU", "tgt": "EN"}).json()["id"]
prj = main.get_project(pid)
prj["segments"] = [{"id": 1, "source": "Первый сегмент.", "target": "", "status": "new"},
                   {"id": 2, "source": "Второй сегмент.", "target": "", "status": "new"}]
r = c.get("/api/projects/%d/quote" % pid, headers=H(A))
pq = r.json()
check(r.status_code == 200 and pq["basis"] == "segments"
      and pq["counts"]["chars"] == len("Первый сегмент.Второй сегмент."),
      "смета по проекту считает исходники сегментов и называет себя segments")
check(any("Объём файла бывает больше" in n for n in pq["notes"]),
      "разница с объёмом файла названа вслух, а не спрятана")


print("=== 8. Отказы приходят кодом, а не пятисоткой ===")
import zipfile                                  # noqa: E402
buf = io.BytesIO()
z = zipfile.ZipFile(buf, "w"); z.writestr("readme.txt", "hi"); z.close()
notdoc = buf.getvalue()
r = c.post("/api/quote", headers=H(A), files={"file": ("fake.odt", notdoc, "application/vnd.oasis")},
           data={"src": "RU", "tgt": "EN"})
check(r.status_code == 415, "zip без content.xml — 415 с причиной, а не 500 (%d)" % r.status_code)
r = c.post("/api/quote", headers=H(A), files={"file": ("fake.docx", notdoc, "application/msword")},
           data={"src": "RU", "tgt": "EN"})
check(r.status_code == 415, "docx без word/document.xml — 415 (%d)" % r.status_code)
r = c.post("/api/quote", headers=H(A), files={"file": ("huge.txt", b"x" * (textcount.MAX_BYTES + 10), "text/plain")},
           data={"src": "RU", "tgt": "EN"})
check(r.status_code == 413, "файл сверх потолка — 413 ДО разбора (%d)" % r.status_code)
try:
    import pypdf                                # noqa: F401
    check(True, "pypdf установлен — путь 503 проверять нечем")
except ImportError:
    r = c.post("/api/quote", headers=H(A), files={"file": ("s.pdf", b"%PDF-1.4 fake", "application/pdf")},
               data={"src": "RU", "tgt": "EN"})
    check(r.status_code == 503, "нет pypdf — 503 «нечем прочитать», а не 415 (%d)" % r.status_code)

print("=== 9. Ноль — не цена ===")
check(c.post("/api/pricing", headers=H(A), json={"default": 0}).status_code == 400,
      "общая цена 0 — 400: пустое поле в браузере приходит нулём")
check(c.post("/api/pricing", headers=H(A), json={
    "rates": [{"src": "RU", "tgt": "EN", "price": 0}]}).status_code == 400,
      "строка прайса с нулём — 400, бесплатной работы в прайсе не бывает")
r = c.post("/api/pricing", headers=H(A), json={"clearDefault": True})
check(r.status_code == 200 and r.json()["pricing"]["default"] is None,
      "снять общую цену можно — и это null, а не 0")

print("=== 10. Прайс не течёт ===")
c.post("/api/admin/tenants", headers=H(A),
       json={"id": "beta", "name": "Beta", "ownerLogin": "beta", "ownerPassword": "beta-pass-123"})
B = c.post("/api/auth/login", json={"login": "beta", "password": "beta-pass-123"}).json()["token"]
c.post("/api/pricing", headers=H(A), json={"default": 99, "currency": "UZS"})
check(c.get("/api/pricing", headers=H(B)).json()["pricing"]["default"] is None,
      "чужой прайс не виден: у B своя пустая карточка")
c.post("/api/pricing", headers=H(B), json={"default": 5})
check(c.get("/api/pricing", headers=H(A)).json()["pricing"]["default"] == 99,
      "B правит только своё — у A цена не изменилась")
me = c.get("/api/auth/me", headers=H(A)).json()
check("pricing" not in (me.get("tenant") or {}) and set(me["tenant"]) <= {"id", "name", "active"},
      "/api/auth/me отдаёт организацию белым списком полей, без прайса")
adm = json.dumps(c.get("/api/admin/tenants", headers=H(A)).json(), ensure_ascii=False)
ovw = json.dumps(c.get("/api/admin/overview", headers=H(A)).json(), ensure_ascii=False)
check("pricing" not in adm and "pricing" not in ovw,
      "суперпользователю прайсы агентств не показываются")
seed = json.dumps(c.get("/api/seed", headers=H(B)).json(), ensure_ascii=False)
check("pricing" not in seed, "/api/seed прайса не несёт")
r = c.get("/api/projects/%d/quote" % pid, headers=H(B))
check(r.status_code == 404, "смета по чужому проекту — 404 (%d)" % r.status_code)

print("=== 11. Округление на разных шагах ===")
check(textcount.pages_of(1801, 1800, 0, 0.25)["billed"] == 1.25, "шаг 0.25: 1.001 → 1.25")
check(textcount.pages_of(450, 1800, 0, 0.25)["billed"] == 0.25, "шаг 0.25: ровно четверть остаётся четвертью")
check(textcount.pages_of(1802, 1800, 0, 1.0)["billed"] == 2.0, "шаг 1: знак сверх страницы — вторая страница")
check(textcount.pages_of(1_000_000, 1800, 0, 0.1)["billed"] == 555.6, "большие числа не плывут")
check(textcount.pages_of(5000, 1800, 10, 0.1)["billed"] == 10.0, "минимальный заказ сильнее расчёта")
check(main._money(555.6, 12.5) == 6945.0, "деньги на большом объёме точны (%s)" % main._money(555.6, 12.5))

print("=== 12. withFile: объём файла рядом с объёмом сегментов ===")
try:
    from docx import Document as _D
    d2 = _D()
    d2.add_paragraph("Первый сегмент.")
    d2.add_paragraph("Второй сегмент.")
    d2.add_paragraph("7")                       # номер страницы: сегментом не станет
    b2 = io.BytesIO(); d2.save(b2)
    r = c.post("/api/projects/%d/source" % pid, headers=H(A),
               files={"file": ("src.docx", b2.getvalue(),
                               "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
               data={"force": "true"})
    q2 = c.get("/api/projects/%d/quote?withFile=true" % pid, headers=H(A)).json()
    check("file" in q2 and (q2["file"].get("counts") or q2["file"].get("error")),
          "withFile отвечает: либо объём файла, либо причина, почему его нет")
    if q2["file"].get("counts"):
        check(q2["file"]["counts"]["chars"] >= q2["counts"]["chars"],
              "объём файла не меньше объёма сегментов — разница законна и названа")
except ImportError:
    check(True, "python-docx нет — проверка withFile пропущена")


print("=== 13. История смет: к ней возвращаются платить ===")
c.post("/api/pricing", headers=H(A), json={"default": 10, "currency": "USD",
                                           "rates": [{"src": "RU", "tgt": "EN", "price": 12.5}]})
doc = ("Текст для сметы. " * 200).encode("utf-8")
r1 = c.post("/api/quote", headers=H(A), files={"file": ("dogovor.docx.txt", doc, "text/plain")},
            data={"src": "RU", "tgt": "EN"}).json()
check(r1.get("saved") and r1["saved"]["total"] == r1["total"],
      "расчёт сам лёг в историю с той же суммой")
qid = r1["saved"]["id"]
hist = c.get("/api/quotes", headers=H(A)).json()["quotes"]
was = len(hist)
check(hist[0]["id"] == qid, "новая смета встала первой в истории")
r2 = c.post("/api/quote", headers=H(A), files={"file": ("dogovor.docx.txt", doc, "text/plain")},
            data={"src": "RU", "tgt": "EN"}).json()
hist = c.get("/api/quotes", headers=H(A)).json()["quotes"]
check(len(hist) == was and hist[0]["count"] == 2,
      "повторный расчёт того же файла по той же цене — не новый заказ, а второе обращение")

# Смена прайса не трогает выданную оферту: клиент считал вчера, платит сегодня.
frozen_total, frozen_price = hist[0]["total"], hist[0]["pricePerPage"]
c.post("/api/pricing", headers=H(A), json={"rates": [{"src": "RU", "tgt": "EN", "price": 40}]})
hist = c.get("/api/quotes", headers=H(A)).json()["quotes"]
check(hist[0]["total"] == frozen_total and hist[0]["pricePerPage"] == frozen_price,
      "смена прайса СТАРУЮ смету не пересчитала (%s)" % hist[0]["total"])
r3 = c.post("/api/quote", headers=H(A), files={"file": ("dogovor.docx.txt", doc, "text/plain")},
            data={"src": "RU", "tgt": "EN"}).json()
check(r3["saved"]["id"] != qid and r3["total"] != frozen_total,
      "по новой цене — НОВАЯ запись: это другая оферта")

print("=== 14. Судьба сметы: право владельца, след человека ===")
check(c.post("/api/quotes/%d" % qid, headers=H(T), json={"status": "paid"}).status_code == 403,
      "переводчик смету оплаченной не объявляет")
r = c.post("/api/quotes/%d" % qid, headers=H(A), json={"status": "paid", "note": "счёт 14"})
check(r.status_code == 200 and r.json()["quote"]["status"] == "paid"
      and r.json()["quote"].get("paidAt"), "владелец пометил оплаченной, время записано")
check(r.json()["quote"]["total"] == frozen_total, "отметка об оплате чисел сметы не тронула")
r = c.post("/api/quotes/%d" % qid, headers=H(A), json={"status": "new"})
check(len(r.json()["quote"].get("log") or []) == 2,
      "возврат в черновик не стирает след: кто и когда что объявил — вопрос денег")
check(c.post("/api/quotes/%d" % qid, headers=H(A), json={"status": "оплачено"}).status_code == 400,
      "выдуманное состояние — 400")
check(c.get("/api/quotes", headers=H(B)).json()["quotes"] == [],
      "чужая история смет не видна")
check(c.post("/api/quotes/%d" % qid, headers=H(B), json={"status": "paid"}).status_code == 404,
      "чужую смету не пометить — 404, а не 403")
seed = json.dumps(c.get("/api/seed", headers=H(B)).json(), ensure_ascii=False)
check("dogovor" not in seed, "/api/seed чужих смет не несёт")
check(c.delete("/api/quotes/%d" % qid, headers=H(T)).status_code == 403, "удаляет сметы владелец")
check(c.delete("/api/quotes/%d" % qid, headers=H(A)).status_code == 200
      and all(q["id"] != qid for q in c.get("/api/quotes", headers=H(A)).json()["quotes"]),
      "удалённая смета из истории ушла")
r = c.post("/api/quote", headers=H(A), files={"file": ("nosave.txt", doc, "text/plain")},
           data={"src": "RU", "tgt": "EN", "save": "false"})
check(r.json().get("saved") is None
      and all(q["file"] != "nosave.txt" for q in c.get("/api/quotes", headers=H(A)).json()["quotes"]),
      "save=false считает, но в историю не пишет")

print()
if fail:
    print("ПРОВАЛЕНО %d:" % len(fail))
    for f in fail:
        print("  -", f)
    sys.exit(1)
print("ВСЁ ПРОШЛО")
