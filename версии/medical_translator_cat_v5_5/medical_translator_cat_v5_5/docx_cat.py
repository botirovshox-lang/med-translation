from pathlib import Path
from docx import Document
from db import create_project,add_segment,get_project,get_segments
def iter_docx_text_blocks(doc):
    idx=0
    for p in doc.paragraphs:
        text=p.text.strip()
        if text: yield ('paragraph',idx,p,text); idx+=1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text=p.text.strip()
                    if text: yield ('table_cell',idx,p,text); idx+=1
def import_docx_project(uploaded_file,title,source_language='Russian',target_language='English'):
    Path('data/source_docs').mkdir(parents=True,exist_ok=True); path=Path('data/source_docs')/uploaded_file.name.replace('/','_').replace('\\','_')
    with open(path,'wb') as f: f.write(uploaded_file.getbuffer())
    pid=create_project(title or uploaded_file.name,str(path),source_language,target_language); doc=Document(str(path)); order=1
    for bt,bi,p,text in iter_docx_text_blocks(doc): add_segment(pid,order,bt,bi,text); order+=1
    return pid,order-1
def replace_paragraph_text(p,new_text):
    if p.runs:
        p.runs[0].text=new_text
        for r in p.runs[1:]: r.text=''
    else: p.add_run(new_text)
def export_translated_docx(pid):
    project=get_project(pid); doc=Document(project['source_docx_path']); blocks=list(iter_docx_text_blocks(doc)); segs=get_segments(pid)
    for seg,block in zip(segs,blocks): replace_paragraph_text(block[2],seg.get('target_text') or seg.get('source_text') or '')
    Path('exports').mkdir(exist_ok=True); out=Path('exports')/(project['title']+'_translated.docx'); doc.save(str(out)); return str(out)
