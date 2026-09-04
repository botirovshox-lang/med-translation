"""
FastAPI backend for CAT Translator v5.6 (перевод документов с проверками).

Сервис НЕ медицинский: предметная область — параметр проекта (`DOMAINS`),
языковая пара — любая из каталога `languages.json`. Медицина осталась одной
из встроенных областей и данными первого клиента, а не свойством системы.
Serves the React design at /, exposes REST API at /api/*

Авторизация: POST /api/auth/login отдаёт токен сессии; его нужно присылать
в заголовке `Authorization: Bearer <token>` во ВСЕ остальные /api/* запросы.
Публичны только /api/auth/login, /api/auth/logout и /api/health.

Endpoints:
  GET  /api/seed                          → all initial data (projects, glossary, tm, etc.)
  POST /api/auth/login                    → password check → token
  POST /api/auth/logout                   → invalidate token
  GET  /api/projects                      → list projects
  POST /api/projects                      → create project (from DOCX or empty)
  GET  /api/projects/{pid}                → project detail
  POST /api/segments/{pid}/{sid}/translate → translate via Google or GPT
  POST /api/segments/{pid}/{sid}/qa       → run QA
  POST /api/segments/{pid}/{sid}/confirm  → confirm + add to TM
  POST /api/segments/{pid}/{sid}/revert   → revert confirmed/failed
  POST /api/segments/{pid}/{sid}/update   → update target/comment
  POST /api/projects/{pid}/batch          → batch translate (выбранной моделью)
  POST /api/projects/{pid}/preflight      → run preflight analysis
  POST /api/projects/{pid}/export         → trigger export (docx|pdf|xlsx)
  POST /api/glossary                      → add/update term
  DELETE /api/glossary/{src}              → delete term
  DELETE /api/tm/{src}                    → delete TM entry

Run:
  cd backend
  uvicorn main:app --reload --port 8000
  → open http://localhost:8000
"""
import os
import re
import sys
import json
import time
import hmac
import hashlib
import secrets
import asyncio
import threading
from pathlib import Path
from typing import Optional, Any, List
import io
import html as _html_mod
import unicodedata
from datetime import datetime

from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Request
from fastapi import Response
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from starlette.concurrency import run_in_threadpool
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

# Add med_translation to path so we can import existing modules
ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "backend"))
sys.path.insert(0, str(ROOT / "med_translation"))

# Try to import existing modules (graceful fallback if missing deps)
_BACKEND_MODULES = {}

def _safe_import(name: str):
    try:
        mod = __import__(name)
        _BACKEND_MODULES[name] = mod
        return mod
    except Exception as e:
        print(f"[backend] WARN: could not import {name}: {e}", file=sys.stderr)
        return None

db = _safe_import("db")
pipeline = _safe_import("pipeline")
tm_mod = _safe_import("tm")
# Модуль детерминированных проверок. Назывался medical_qa, пока сервис был
# медицинским; правила в нём — числа, единицы, отрицания, соответствие
# глоссарию — предметной области не касаются вовсе, а те, что касаются,
# лежат таблицей DOMAIN_RULES по ключу «область + пара языков».
checks_mod = _safe_import("checks")
# Внешние источники приказов: отраслевые справочники и корпуса целевого языка.
# Без них термин может заверить только человек — а он целевого языка может
# и не знать. См. шапку authorities.py.
authorities_mod = _safe_import("authorities")

# Google Translate убран из системы. Перевод делает ТОЛЬКО выбранная модель:
# бесплатный движок не знает предметной области, игнорирует глоссарий (кроме
# грубой подстановки плейсхолдерами) и давал текст, который потом всё равно
# приходилось чинить платными прогонами. Экономия была мнимой.
#
# Отсюда следует: нет ключа OpenAI — нет перевода. Не заглушка, не «черновик
# бесплатным движком», а честная ошибка. Исторические сегменты с
# provider="google" и route="GOOGLE_SAFE" остаются как есть: это факт о том,
# чем их перевели, и переписывать его нельзя.

import re as _re

_PATTERN_CACHE: dict = {}
_PATTERN_CACHE_MAX = 20000


# Границы слова для поиска термина. Класс намеренно шире одной письменности:
# он отвечает на вопрос «не влезли ли мы в середину чужого токена», а соседним
# словом в русском тексте законно бывает латинское («Mycobacterium bovis»).
# ЦИФРЫ входят в границу: без них «CD4» ловил «CD40», «HIV» — «HIV1»,
# «2HRE» — «2HREZ», а это разные рецепторы, вирусы и схемы лечения. Дефис
# и апостроф в границу НЕ входят намеренно: «БЦЖ» обязан находиться
# в «БЦЖ-вакцина», а «patient» — в «patient's».
_LETTERS = "а-яёА-ЯЁa-zA-ZÀ-ÖØ-öø-ÿ"
_DIGITS = "0-9"


def _bound_l(term: str) -> str:
    """Левая граница термина. Цифра запрещена перед ним, только если сам термин
    с цифры и НАЧИНАЕТСЯ: иначе «2HRE» нашёлся бы внутри «12HRE»."""
    return ("(?<![" + _LETTERS + (_DIGITS if (term or "")[:1].isdigit() else "") + "])")


def _bound_r(term: str, cls: str = "") -> str:
    """Правая граница. Цифра запрещена после термина, только если он цифрой
    КОНЧАЕТСЯ: «CD4» не должен ловить «CD40», а «Т1» — «Т12», это разные
    рецепторы и стадии.

    Запрещать цифру всегда нельзя: импорт этого учебника приклеивал к словам
    номера страниц («ТУБЕРКУЛЕЗА ОРГАНОВ ДЫХАНИЯ16»), и глухой запрет отнимал
    у приказной записи законное совпадение ради защиты от случая, которого
    в тексте нет. Дефис и апостроф в границу не входят никогда: «БЦЖ» обязан
    находиться в «БЦЖ-вакцина», а «patient» — в «patient's»."""
    last = (term or "")[-1:]
    return "(?![" + (cls or _LETTERS) + (_DIGITS if last.isdigit() else "") + "])"


# ── Словоизменение: закрытые списки окончаний по языкам ──────────────
#
# Списки нужны ровно для одного вопроса: найденная форма — форма ТОГО ЖЕ
# слова или другого слова с общим началом. Словообразовательных суффиксов
# (-ома, -оз, -ит, -ация, -ary, -itis) здесь нет намеренно: ими понятия
# и различаются.
#
# Главный закон таблиц — ОНИ ОШИБАЮТСЯ В БЕЗОПАСНУЮ СТОРОНУ. Слишком короткий
# список означает, что часть законных форм не найдётся: запись просто не
# сработает на этом сегменте. Слишком длинный означает, что запись сядет
# на чужое слово и прикажет модели неверный перевод. Первое — потеря, второе —
# в медицинском переводчике подмена понятия. Поэтому списки минимальны,
# а расширять их можно только по замеру на реальных данных.
#
# Нет списка для языка — стем-поиск НЕ РАБОТАЕТ ВООБЩЕ, остаётся только точное
# совпадение по границам слова. Это тот же закон, что у DOMAIN_RULES и правил
# регистра: нет правил для этой пары — молчим. Раньше отсутствие правил
# означало обратное, максимальную нестрогость: класс букв в стем-шаблоне был
# кириллическим, поэтому на латинице он совпадал с нулём символов и находкой
# становился ОБРУБОК («Mycobacterium» → «Mycobacteri», «infiltrate» →
# «infiltra»), а `_same_lexeme` для некириллицы молчала — то есть защита
# от подмены была выключена целиком: «test» ловил «testosterone», «gene»
# ловил «generalized».
# Окончания, при которых найденная форма ОБЯЗАНА сохранить «и»: существительное
# на «-ия» держит его во всех формах (терапия, терапии, терапию, терапией,
# терапий). Форма без него — другое слово с тем же началом: дисциплина против
# врача. «Фтизиатрия» ловила «фтизиатр», «фтизиатра», «фтизиатром», и приказ
# «Фтизиатрия → Phthisiology» шёл на сегменты про врача (то же «хирургия →
# хирурга», «педиатрия → педиатром»).
_KEEP_I_ENDINGS = frozenset(("ия",))
_RU_ENDINGS = frozenset((
    "", "а", "е", "и", "й", "о", "у", "ы", "ь", "ю", "я",
    "ам", "ах", "ая", "ев", "ей", "ем", "ею", "ие", "ий", "им", "их", "ия",
    "ов", "ое", "ой", "ом", "ою", "ую", "ые", "ый", "ым", "ых", "ья", "ье",
    "ью", "юю", "яя", "ям", "ях",
    # «ее» — средний род мягких прилагательных и вся сравнительная степень:
    # «верхний → верхнее», «средний → среднее», «ранний → раннее».
    "ее",
    "ами", "его", "ему", "ими", "ого", "ому", "ыми", "ями", "ией", "иям",
    "иях", "ьев", "ьям", "ьях", "иями", "ьями",
))
# Английский: число, включая латинские множественные, без которых медицинский
# текст не живёт («bronchus → bronchi», «Mycobacterium → Mycobacteria»,
# «plica → plicae»). Чего здесь НЕТ и почему (замер по всему глоссарию против
# всех 2711 переводов):
#   • «ed» и «ing» — они путают существительное с глаголом и причастием:
#     «conditioned → condition» (83 срабатывания), «forms → formed»,
#     «injured → injury», «diseases → diseased». Законные их срабатывания
#     все до одного на записях, которые сами не словарная форма
#     («increased», «stained»), — то есть на том, что `_looks_like_term`
#     и сверка смысла и так обязаны отклонить;
#   • «on»/«a» («criterion → criteria») — давало «plasmon → plasma»,
#     а это разные вещи;
#   • производных прилагательных («cavity» / «cavitary») здесь нет никогда.
_EN_ENDINGS = frozenset(("", "s", "es", "y", "ies", "um", "a", "us", "i", "ae"))
# Немецкий, испанский, французский: ТОЛЬКО число. Род и падеж сюда не входят,
# и это не лень, а тот же закон безопасной ошибки: на минимальных парах
# «puerto/puerta», «grupo/grupa», «rein/reine», «sein/seine» одиночная гласная
# в окончании превращает одно слово в другое. Законное «medico → medica»
# уходит вместе с ними — потерянная подсказка дешевле подмены понятия.
# На данных эти три таблицы НЕ проверялись: пар DE/ES/FR в работе не было.
# Композиты («Lungenentzündung» внутри «Lungenentzündungsherd») и умлаут
# в основе («Buch → Bücher») суффиксом не берутся и браться не будут.
_DE_ENDINGS = frozenset(("", "e", "en", "n", "s"))
_ES_ENDINGS = frozenset(("", "s", "es"))
_FR_ENDINGS = frozenset(("", "s", "es", "x"))

_LANG_ENDINGS = {
    "RU": _RU_ENDINGS,
    "EN": _EN_ENDINGS,
    "DE": _DE_ENDINGS,
    "ES": _ES_ENDINGS,
    "FR": _FR_ENDINGS,
}

# Окончания, которые ПРИПИСЫВАЮТСЯ к записи целиком, а не заменяют её хвост.
# Разводить эти два случая обязательно: латинское множественное в английском
# именно ЗАМЕНЯЕТ окончание («bronchus → bronchi»), и разрешив его припиской,
# мы получаем «ARV → ARVI» — антиретровирусные препараты против острой
# респираторной вирусной инфекции. Заодно отпадают «echocardiograph →
# Echocardiography» и «mammograph → Mammography»: там тоже приписка.
#
# Список задан для КАЖДОГО языка, а запасной вариант — пустой. «Нет правил —
# молчим»: раньше запасным вариантом был полный набор окончаний языка, и это
# давало ровно ту беду, от которой список заведён, — «sal → sala», «rein →
# reine», «Rat → Rate», «Not → Note».
_LANG_APPEND = {
    "RU": _RU_ENDINGS,          # в русском приписывается почти всё: «рак → рака»
    "EN": frozenset(("", "s", "es")),
    "DE": frozenset(("", "n", "en", "s")),
    "ES": frozenset(("", "s", "es")),
    "FR": frozenset(("", "s", "x")),
}
_APPEND_NONE = frozenset(("",))

# Послабление «найденное начинается с записи целиком» действует от пяти букв...
_WHOLE_WORD_MIN = 5
# ...и только в языках, где оно измерено безвредным. В русском без него
# перестало бы находиться «туберкулёз» в «туберкулёзного» — 96 законных
# совпадений на боевом проекте. В английском оно, наоборот, вредно: слова
# там наращиваются суффиксом без соединительной морфемы, и запись садится
# на чужое слово целиком. Замер на английской стороне боевого проекта:
# «creatin → creating», «derma → dermatologists», «magnet → Magnetic»,
# «asbestos → asbestosis», «keloid → keloidal», «bipolar → bipolarity»,
# «opportunist → Opportunistic», «broncholith → broncholithiasis».
# Для остальных языков послабление выключено: неизмеренное правило,
# ошибающееся в опасную сторону, включать нельзя.
_WHOLE_WORD_LANGS = frozenset(("RU",))

# Письменность берётся у САМОГО СЛОВА, а не у языка проекта: класс букв
# в стем-шаблоне должен покрывать то, из чего слово состоит, иначе шаблон
# обрывается на середине слова. Латинский класс с диакритикой — иначе
# «Lungenentzündung» и «pneumonía» рвались бы на умлауте и ударении.
_SCRIPTS = (
    ("cyr", "а-яёА-ЯЁ"),
    ("lat", "A-Za-zÀ-ÖØ-öø-ÿ"),
)
_SCRIPT_CLASS = dict(_SCRIPTS)
_SCRIPT_RE = {name: _re.compile("[" + cls + "]") for name, cls in _SCRIPTS}
_SCRIPT_ONLY_RE = {name: _re.compile("[" + cls + "]+") for name, cls in _SCRIPTS}
# «Слово целиком» для склонения: буквы одной письменности, между которыми
# допустим дефис или апостроф. Дефис обязателен: русская медицинская
# терминология на нём стоит — «фиброзно-кавернозный», «клинико-
# рентгенологический», «черепно-мозговая». Запрети их — и они перестанут
# склоняться, то есть перестанут находиться в косвенных падежах, которых
# в тексте большинство. А вот цифра или скобка внутри («ТУБЕРКУЛЁЗА16»,
# «(CO2)») склонению не подлежит: там мы не умеем прочитать слово как слово,
# и такие записи требуются точно.
_WORDISH_RE = {name: _re.compile("[" + cls + "]+(?:[-’'][" + cls + "]+)*")
               for name, cls in _SCRIPTS}


def _src_lang(entry: dict) -> str:
    """Язык ОРИГИНАЛА записи глоссария («RU→EN» → «RU»).

    Берётся из самой записи, а не из проекта: так язык доезжает до каждого
    места, где ищут термин, без протаскивания параметра через полдюжины
    вызовов. Записи без поля читаются как DEFAULT_GLOSS_LANG — весь
    исторический импорт именно такой (см. CLAUDE.md про область записи).

    Пара, записанная не латиницей («Русский→English»), даёт пустую строку,
    то есть стем-поиска не будет. Это осознанно: не разобрав язык, безопаснее
    искать только точное совпадение, чем угадать таблицу окончаний."""
    pair = (entry or {}).get("lang") or DEFAULT_GLOSS_LANG
    head = _re.split(r"[^A-Za-z]+", str(pair).strip())
    return (head[0] if head else "").upper()


def _word_script(term: str):
    """Письменность слова — по большинству букв. Ни одной знакомой буквы
    (иероглифы, арабица, греческий, одни цифры) — None, и тогда стем-поиска
    не будет: судить о словоизменении письма, для которого у нас нет ни класса
    букв, ни окончаний, значит гадать."""
    best, best_n = None, 0
    for name, _cls in _SCRIPTS:
        n = len(_SCRIPT_RE[name].findall(term or ""))
        if n > best_n:
            best, best_n = name, n
    return best


def _same_lexeme(term_word: str, found_word: str, lang: str = "") -> bool:
    """Найденная форма — форма ТОГО ЖЕ слова, а не другого с общим началом.

    Зачем: стем режет слово до 85% букв, и на длинном слове это срезает ровно
    тот хвост, которым в медицине различаются понятия. «Туберкулема» (11 букв)
    даёт стем «туберкуле» — и ловит «туберкулез» со всеми его формами. На боевом
    учебнике так вышло 1006 раз у одной этой записи, ещё 136 у «Туберкулёма
    лёгких», плюс «аллергия → аллергена», «инфильтрат → инфильтрации»,
    «алкоголизм → алкоголики». Итого 1185 ложных совпадений на 875 сегментах —
    и в промпт перевода уходил приказ «Туберкулема → Tuberculoma» на сегменты
    про туберкулёз. Комментарий к порогу 85% предупреждал ровно об этом на паре
    «циклоз/циклит» — там шести букв хватило, на одиннадцати перестало.

    Правило: у записи и у найденной формы должен найтись ОБЩИЙ стем, после
    которого у обеих остаётся окончание из закрытого списка ЭТОГО ЯЗЫКА.
    «первичный» и «первичного» дают стем «первичн» + «ый»/«ого» — одно слово;
    «туберкулема» и «туберкулеза» не дают такого стема ни при какой длине:
    «ма»/«за», «ема»/«еза», «лема»/«леза» окончаниями не бывают.

    Хвост записи пуст — значит найденное ПРИПИСАЛО что-то к записи целиком,
    а не изменило её окончание. Это другой случай, и список у него свой
    (`_LANG_APPEND`), потому что приписка и замена ведут себя по-разному:
    «bronchus → bronchi» законно, «ARV → ARVI» нет.

    Два намеренных исключения, оба ради того, чтобы не менять лишнего:
      • найденное НАЧИНАЕТСЯ с записи целиком, запись не короче
        _WHOLE_WORD_MIN, а язык — из _WHOLE_WORD_LANGS («туберкулёз» →
        «туберкулёзного», «инфильтрат» → «инфильтративный») — так работало
        и так работает. Список языков, а не всеобщее правило, потому что
        в английском оно ошибается в опасную сторону, см. комментарий там же.
        Это НЕ значит, что производные прилагательные не судятся вообще:
        когда запись кончается гласной или мягким знаком, послабление
        не срабатывает, и «полость» → «полостной», «бронхит» →
        «бронхиальный» отвергаются;
      • слово написано НЕ ОДНОЙ письменностью — МОЛЧИМ. Под это попадает
        не только чужой алфавит: «HIV/TB» и всё с цифрой или знаком внутри
        проверку минуют. Судить о словоизменении такого слова нечем.

    Нет таблицы окончаний для языка — совпадением считается только полное
    равенство. Стем-поиск в этом случае и не запускается (см. _term_patterns),
    так что сюда мы попадаем разве что прямым вызовом.

    Чего правило НЕ умеет, и это названо намеренно:
      • беглая гласная: «лимфоузел» → «лимфоузлы», «ребёнок» → «ребёнка»,
        «палец» → «пальца». Основа меняется, общего стема с окончаниями
        не находится, и законная форма отвергается. Примеры настоящие —
        обе записи в боевом глоссарии есть;
      • разносклоняемые на -мя («время» → «времени») и на -анин
        («гражданин» → «граждане»): та же беда с основой;
      • глагольные формы («разрушать» → «разрушается»): список окончаний
        именной, а глагол словарной записью глоссария почти не бывает;
      • немецкие композиты и умлаут в основе — суффиксом их не взять.
    На боевом глоссарии весь этот класс — около 4% снятых пар и НИ ОДНОЙ
    приказной записи: цена ошибки здесь потерянная подсказка, а не приказ."""
    a = (term_word or "").lower().replace("ё", "е")
    b = (found_word or "").lower().replace("ё", "е")
    lang = (lang or "").upper()
    endings = _LANG_ENDINGS.get(lang)
    if endings is None:
        return a == b
    script = _word_script(a)
    if script is None or _word_script(b) != script:
        return True
    only = _WORDISH_RE[script]
    if not only.fullmatch(a) or not only.fullmatch(b):
        return True
    n = 0
    for x, y in zip(a, b):
        if x != y:
            break
        n += 1
    if (n >= len(a) and len(a) >= _WHOLE_WORD_MIN
            and lang in _WHOLE_WORD_LANGS):
        return True
    # Потолок считается по ТАБЛИЦЕ ЭТОГО ЯЗЫКА, а не по всем сразу: иначе
    # длинное окончание, добавленное в турецкую таблицу, расширило бы цикл
    # и русскому правилу — то есть ослабило бы защиту «туберкулема ≠
    # туберкулез» правкой, к русскому отношения не имеющей.
    ending_max = max(len(e) for e in endings)
    append = _LANG_APPEND.get(lang, _APPEND_NONE)
    for k in range(n, max(0, n - ending_max) - 1, -1):
        if a[k:] not in endings:
            continue
        allowed = append if not a[k:] else endings
        if a[k:] in _KEEP_I_ENDINGS and not b[k:].startswith("и"):
            continue
        if b[k:] in allowed:
            return True
    return False


def _term_forms(term: str, text: str, lang: str = "") -> list:
    """ВСЕ вхождения термина в текст, прошедшие проверку «то же слово».

    Один расчёт на всех: по нему `_term_match` отвечает «есть ли термин»,
    а `_agreed_form` смотрит, одинаково ли он написан во всех местах.
    Разойдись они — отчёт показывал бы одно, а правка делала другое: до этого
    `_agreed_form` считала начертание в том числе по словам, которые
    `_term_match` отвергал, и на боевых данных из-за этого молча теряла
    проверку начертания на двух сегментах («Инфильтративный» вперемешку
    с «инфильтрация», «Туберкулема» с «туберкулеза»)."""
    exact, stem = _term_patterns(term, lang)
    if exact is None:
        return []
    if stem is None:
        return [m.group(0) for m in exact.finditer(text)]
    tw = (term or "").lower().split()
    out = []
    for m in stem.finditer(text):
        found = m.group(0)
        fw = found.lower().split()
        # Число слов у шаблона и у термина совпадает по построению (шаблон
        # склеен ровно len(parts)-1 разделителями, а кусок пробела не
        # содержит). Условие оставлено сторожем: если оно однажды сработает,
        # термин пропадёт молча, и знать об этом надо.
        if len(tw) != len(fw):
            print("[backend] _term_match: %r нашлось %d словами вместо %d (%r)"
                  % (term, len(fw), len(tw), found), file=sys.stderr)
            continue
        if all(_same_lexeme(a, b, lang) for a, b in zip(tw, fw)):
            out.append(found)
    return out


def _term_match(term: str, text: str, lang: str = "") -> "str | None":
    """Ищет термин (в любой грамматической форме) в тексте.
    Возвращает фактически найденную подстроку или None.

    Алгоритм:
    1. Точное совпадение по границам слова (быстро, работает на любом языке).
    2. Стем-поиск: первые 85% символов каждого слова термина (мин. 4) + любое
       окончание ТОЙ ЖЕ письменности, и найденное слово обязано оказаться
       формой того же слова (`_same_lexeme`). Без таблицы окончаний для языка
       шаг 2 не выполняется вовсе.

    Порог 85%, а не 75%: на 75% «циклоз» (стем «цикл») ловил «циклит»
    и подсовывал модели «cyclosis» вместо «cyclitis». Медицинские термины
    различаются как раз хвостом (-ит / -оз / -ома), срезать его нельзя.
    Одного порога, впрочем, мало — см. `_same_lexeme`.

    Перебираются ВСЕ вхождения, а не первое: стем совпасть мог и с другим
    словом, а настоящая форма стоять следующим («Очаговый туберкулез легких.
    Туберкулемы выявлены…»). По первому хиту запись «Туберкулема» была бы
    объявлена отсутствующей, и приказ не ушёл бы ни в промпт, ни в проверку
    соответствия. В учебнике фтизиатрии эти слова стоят в одном абзаце
    постоянно, так что правило гасило бы само себя, и молча.
    """
    exact, _stem = _term_patterns(term, lang)
    if exact is None:
        return None
    hit = exact.search(text)
    if hit:
        return hit.group(0)
    got = _term_forms(term, text, lang)
    return got[0] if got else None


def _term_patterns(term: str, lang: str = ""):
    """Скомпилированные шаблоны термина: точный (по границам слова) и стем-поиск.

    Раньше обе регулярки собирались и компилировались на КАЖДОЕ сравнение.
    На словаре в 10 000 записей внутренний кэш re переполнялся и компилировал
    заново — отбор кандидатов для одного сегмента стоил 130 мс вместо единиц.
    Ключ кэша включает язык: от него зависит, собирается ли стем-шаблон вообще.

    Склоняется КАЖДОЕ СЛОВО ОТДЕЛЬНО и только если оно целиком состоит из букв
    одной письменности. Остальные слова требуются точно. Иначе:
      • «рак in situ» получал бы латинский класс на всё (латинских букв
        больше) и переставал находить «рака in situ»;
      • «вирус HIV» находил бы себя внутри «вирус HIVB»;
      • «повторное вдыхание (CO2)» возвращало бы «повторного вдыхания (CO2»
        без закрывающей скобки — стем резал слово «(co2)» до «(co2», а класс
        букв дописать её не мог.
    Склонять можно то, что мы умеем прочитать как слово; всё прочее — точно."""
    key = (term, (lang or "").upper())
    cached = _PATTERN_CACHE.get(key)
    if cached is not None:
        return cached
    tl = (term or "").lower()
    parts = tl.split()
    if not parts:
        pair = (None, None)
    else:
        # 1. Точное совпадение — обязательно по границам слова. Без этой проверки
        #    трёхбуквенные записи глоссария лезли внутрь чужих слов: «жалобы» →
        #    «лоб», «профилактика» → «лак», «диагностики» → «нос».
        exact = _re.compile(_bound_l(tl) + _re.escape(tl) + _bound_r(tl),
                            _re.IGNORECASE)
        stem = None
        if (lang or "").upper() in _LANG_ENDINGS:
            chunks, tail, any_stem = [], _bound_r(tl), False
            for w in parts:
                sc = _word_script(w)
                cls = _SCRIPT_CLASS.get(sc) if sc else None
                if cls and _WORDISH_RE[sc].fullmatch(w):
                    chunks.append(_re.escape(w[:max(4, int(len(w) * 0.85))])
                                  + "[" + cls + "]*")
                    # Слово склоняется, значит кончится буквой своей
                    # письменности: граница по её классу.
                    tail = "(?![" + cls + "])"
                    any_stem = True
                else:
                    chunks.append(_re.escape(w))
                    tail = _bound_r(w)
            # Ни одно слово не склоняется — шаблон совпал бы с точным,
            # и второй проход по тексту был бы работой впустую.
            if any_stem:
                stem = _re.compile(_bound_l(tl) + r"\s+".join(chunks) + tail,
                                   _re.IGNORECASE)
        pair = (exact, stem)
    if len(_PATTERN_CACHE) >= _PATTERN_CACHE_MAX:
        _PATTERN_CACHE.clear()
    _PATTERN_CACHE[key] = pair
    return pair


# ─── Уровни доверия глоссария ────────────────────────────────────────
# 9132 записи из 10022 пришли массовым автоимпортом (Sources=baldwin_*), и там
# лежит, например, «задний → rear». Такие записи модель получает как подсказку,
# а не как приказ: "use these exact translations" на них давало кальки вроде
# "rear cyclitis". Проверенные (отраслевые списки + одобренное человеком)
# остаются жёстким правилом. Записи не удаляем — понижаем в правах.
GLOSSARY_TIER_HARD = "verified"
GLOSSARY_TIER_SOFT = "auto"


def _tier_from_origin(origin: str) -> str:
    return GLOSSARY_TIER_SOFT if "baldwin" in (origin or "").lower() else GLOSSARY_TIER_HARD


# ─── Область действия записи: языковая пара + тематика ───────────────
# Глоссарий и TM одни на весь сервис, а проекты бывают разные. Запись,
# одобренная в RU→DE юридическом проекте, не должна лезть в промпт RU→EN
# медицинского: это ровно тот путь, которым «задний → rear» переезжает из
# одного текста в другой. У старых записей полей нет — они пришли из
# медицинского RU→EN импорта, поэтому ОТСУТСТВИЕ поля читается как этот
# дефолт, а не как «годится везде». Массовую миграцию не делаем: 10 022
# записи × два поля — лишние сотни килобайт в state.json на каждое сохранение.
DEFAULT_GLOSS_LANG = "RU→EN"
LEGACY_DOMAIN = "medical"        # как читается запись/проект БЕЗ поля domain
DEFAULT_GLOSS_DOMAIN = LEGACY_DOMAIN


def _lang_pair(project: Optional[dict]) -> str:
    return f"{(project or {}).get('src', 'RU')}→{(project or {}).get('tgt', 'EN')}"


# ─── Организация (арендатор) — ТРЕТЬЕ измерение области ──────────────
# Изоляция клиентов держится не на проверке в каждом из 72 эндпоинтов,
# а на конструкции: организация — такое же поле области, как пара и
# тематика. Всё, что сравнивает области (глоссарий, TM, очередь терминов,
# обходы «по всем проектам»), закрывается этим одним кортежем. Организация
# стоит ПОСЛЕДНЕЙ, чтобы `scope[0]`/`scope[1]` во всех прежних местах
# остались парой и тематикой. Записи без поля читаются как организация
# по умолчанию — тот же закон миграции, что у `lang`/`domain`: боевой файл
# не переписывается. НО каждая НОВАЯ запись обязана нести `tenant`: без него
# она уедет в организацию по умолчанию, то есть к чужому клиенту.
DEFAULT_TENANT = "default"
_JOB_TENANT = threading.local()      # организация фонового прогона (в его потоках)
_JOB_LANG = threading.local()        # язык объяснений прогона (там же и по той же причине)


def _current_tenant() -> str:
    """Организация текущего запроса (из сессии) или текущего прогона
    (из его потока). Ни того ни другого — организация по умолчанию:
    так ходят миграции при старте и тесты."""
    sess = CURRENT_SESSION.get() if "CURRENT_SESSION" in globals() else None
    if sess and sess.get("tenant"):
        return sess["tenant"]
    return getattr(_JOB_TENANT, "id", None) or DEFAULT_TENANT


# Язык, на котором модель пишет ЧЕЛОВЕКУ (`why` у проверки терминов, `comment`
# судьи, довод арбитра). Вопрос модели при этом не меняется — меняется только
# язык объяснения, поэтому версии вердиктов НЕ поднимаются: они сторожат
# ПРИГОДНОСТЬ ответа, а пригодность от языка пояснительной фразы не зависит.
# Подъём версии перекупил бы тысячи оплаченных вердиктов ради одной фразы.
#
# Обратная сторона названа честно: уже написанное объяснение остаётся на своём
# языке до ближайшей перепроверки этого сегмента. Переписать его задним числом
# нечем — это текст модели, а не наш шаблон.
#
# Язык берётся у ТОГО, КТО ЗАПУСТИЛ работу: у запроса — из сессии, у прогона —
# из задачи (ContextVar в рабочие потоки не доезжает, ровно как организация).
EXPLAIN_LANG_NAME = {"ru": "Russian", "uz": "Uzbek (Latin script)"}


def _explain_lang() -> str:
    sess = CURRENT_SESSION.get() if "CURRENT_SESSION" in globals() else None
    if sess and sess.get("uiLang"):
        return sess["uiLang"]
    return getattr(_JOB_LANG, "code", None) or DEFAULT_UI_LANG


def _explain_lang_name() -> str:
    """Как назвать язык в промпте. Незнакомый код — русский: пустое имя
    языка модель истолкует по-своему, и объяснение придёт неизвестно на чём."""
    return EXPLAIN_LANG_NAME.get(_explain_lang()) or EXPLAIN_LANG_NAME["ru"]


def _tenant_of(obj: Optional[dict]) -> str:
    return (obj or {}).get("tenant") or DEFAULT_TENANT


def _tenant_projects() -> list:
    """Проекты ТЕКУЩЕЙ организации — единственный законный обход списка
    в эндпоинтах. Прямой `STATE["projects"]` остаётся миграциям и id."""
    t = _current_tenant()
    return [p for p in STATE["projects"] if _tenant_of(p) == t]


def _scope(lang: Optional[str], domain: Optional[str]) -> tuple:
    """Область из полей запроса — с организацией текущей сессии."""
    return (lang or DEFAULT_GLOSS_LANG, domain or DEFAULT_GLOSS_DOMAIN, _current_tenant())


def _scope_of(entry: dict) -> tuple:
    """(языковая пара, тематика, организация) записи глоссария или кандидата."""
    return (entry.get("lang") or DEFAULT_GLOSS_LANG,
            entry.get("domain") or DEFAULT_GLOSS_DOMAIN,
            _tenant_of(entry))


def _project_scope(project: Optional[dict]) -> tuple:
    if project is None:
        return (DEFAULT_GLOSS_LANG, DEFAULT_GLOSS_DOMAIN, _current_tenant())
    return (_lang_pair(project), _resolve_domain(project.get("domain"))["id"],
            _tenant_of(project))


def _hit_tier(h: dict) -> str:
    """Уровень доверия записи. Отсутствие поля читается как ПОДСКАЗКА, а не
    приказ: запись неизвестного происхождения не должна принуждать модель и
    подставляться в перевод плейсхолдером мимо всех проверок. Исторические
    записи уровень получают миграцией при загрузке состояния."""
    return h.get("tier") or GLOSSARY_TIER_SOFT


def _hit_rank(h: dict) -> tuple:
    """Кого оставить, когда на одну форму претендуют несколько записей:
    сначала проверенные, потом более длинный (более специфичный) термин, при
    равенстве — с более коротким переводом. Последнее отсекает обрезанные
    записи вроде «периферическая → peripheral nervous system», которые тянут
    в перевод лишнее понятие."""
    return (1 if _hit_tier(h) == GLOSSARY_TIER_HARD else 0,
            len(h.get("src", "")), -len(h.get("tgt", "")))


# ─── Индекс глоссария ────────────────────────────────────────────────
# Без него _get_context гонял ~10 000 регулярок на КАЖДЫЙ сегмент: 10-17 секунд
# чистого CPU, которые прятались за временем ответа модели (и превращали пакет
# из 2000 сегментов в лишние часы). Теперь по тексту собираются 4-символьные
# ключи, и полную проверку проходят только записи из совпавших корзин.
_GLOSS_INDEX = None          # {ключ: [записи глоссария]}
_GLOSS_INDEX_LOCK = threading.Lock()


def _index_key(word: str) -> str:
    return word[:4]


def _entry_keys(src: str) -> set:
    """Ключи записи — по первому слову термина: именно с него начинается и
    точное совпадение, и стем-поиск."""
    first = (src or "").lower().split()
    if not first:
        return set()
    w = first[0]
    return {_index_key(w)} if len(w) >= 4 else {w}


def _text_keys(text: str) -> set:
    """Ключи текста — только НАЧАЛА слов. С тех пор как точное совпадение требует
    границы слова, термин не может начаться в середине чужого слова, и окна по
    всей длине слова давали лишь лишних кандидатов: на длинном сегменте отбор
    вырождался в перебор всего глоссария (134 мс против 2 мс на сегмент)."""
    keys = set()
    # Класс тот же, что у _LETTERS: без диакритики индекс разрезал бы «Ärzte»
    # на «rzte», а ключ записи — «ärzt», и записи немецкого, французского
    # и испанского не находились бы ВООБЩЕ, сколько ни настраивай окончания.
    for w in _re.findall(r"[а-яёa-zà-öø-ÿ0-9]+", (text or "").lower()):
        # Короткие записи («ЭКГ», «КТ») лежат в корзине целого слова
        keys.update(w[:n] for n in (1, 2, 3, 4) if len(w) >= n)
    return keys


def _gloss_index() -> dict:
    global _GLOSS_INDEX
    idx = _GLOSS_INDEX
    if idx is not None:
        return idx
    with _GLOSS_INDEX_LOCK:
        if _GLOSS_INDEX is None:
            idx = {}
            for g in STATE.get("glossary", []):
                for k in _entry_keys(g.get("src", "")):
                    idx.setdefault(k, []).append(g)
            _GLOSS_INDEX = idx
            print(f"[backend] glossary index: {len(idx)} keys / "
                  f"{len(STATE.get('glossary', []))} terms", file=sys.stderr)
        return _GLOSS_INDEX          # локальная ссылка: см. _gloss_by_src


_GLOSS_BY_SRC: Optional[dict] = None


def _invalidate_gloss_index():
    """Любая правка глоссария роняет индекс — соберётся заново при следующем поиске.
    Заодно двигает поколение: отчёт о расхождениях считается от глоссария."""
    global _GLOSS_INDEX, _GLOSS_BY_SRC
    _GLOSS_INDEX = None
    _GLOSS_BY_SRC = None
    _GLOSS_EPOCH[0] += 1


def _get_context(text: str, with_tm: bool = True, project: Optional[dict] = None):
    """Возвращает (gloss_hits, tm_hit) для исходного текста.

    gloss_hits — список dict {src, tgt, ..., _form} где _form — фактическая
                 форма термина найденная в тексте (нужна для замены).
    tm_hit     — точное совпадение в TM или None. with_tm=False пропускает поиск
                 по памяти переводов: отчёту о расхождениях он не нужен, а это
                 линейный проход по всей TM на каждый сегмент.
    """
    hits = []
    idx = _gloss_index()
    keys = _text_keys(text)
    seen_ids = set()
    candidates = []
    for k in keys:
        for g in idx.get(k, ()):
            if id(g) not in seen_ids:
                seen_ids.add(id(g))
                candidates.append(g)
    scope = _project_scope(project)
    for g in candidates:
        src = g.get("src", "")
        if not src:
            continue
        # Чужая языковая пара или тематика — мимо: в промпт уходят только
        # записи, заведённые для таких же проектов.
        if _scope_of(g) != scope:
            continue
        # Язык берём у САМОЙ записи: от него зависит таблица окончаний,
        # а записи чужой языковой пары сюда и не доходят (проверка выше).
        form = _term_match(src, text, _src_lang(g))
        if form:
            hits.append({**g, "_form": form})
    # Одна форма — один перевод. «периферические» ловило сразу три записи
    # (peripheral / periphery / peripheral nervous system), и модель выбирала
    # сама. Оставляем лучшую по _hit_rank, остальные отбрасываем.
    best: dict = {}
    for h in hits:
        key = h["_form"].lower()
        if key not in best or _hit_rank(h) > _hit_rank(best[key]):
            best[key] = h
    # Длинные термины первыми: меньше риск перекрытия плейсхолдеров
    hits = sorted(best.values(), key=lambda x: len(x["src"]), reverse=True)[:15]
    if not with_tm:
        return hits, None
    # TM хранит языковую пару с самого начала — и обязана её учитывать: без
    # этого RU→DE проект получил бы английский перевод как точное совпадение,
    # да ещё со статусом confirmed.
    tm_hit = next(
        (t for t in STATE.get("tm", [])
         if t.get("src", "").strip().lower() == text.strip().lower()
         and (t.get("lang") or DEFAULT_GLOSS_LANG) == scope[0]
         and _tenant_of(t) == scope[2]),
        None,
    )
    return hits, tm_hit


def _replace_target(seg: dict, text: str, provider: str, route: str):
    """Записать в сегмент новый перевод. Единственное место, где машина имеет
    право заменить текст, и правило тут одно на всех: если перевод заверял
    человек, прежний текст сохраняется в prevTarget, статус становится
    «требует проверки», а отметка «подтвердил человек» снимается — она
    относилась к тексту, которого больше нет."""
    if seg.get("status") == "confirmed":
        seg["prevTarget"] = seg.get("target", "")
        seg.pop("confirmedBy", None)
        seg.pop("confirmedAt", None)
        seg.pop("confirmedRole", None)
        seg["status"] = "review"
    else:
        seg["status"] = "translated"
    seg["target"] = text
    seg["provider"] = provider
    seg["route"] = route


def _tm_trusted(t: Optional[dict]) -> bool:
    """Право подменить перевод МИМО модели есть только у записей, рождённых
    подтверждением человека в этой системе (`quality: verified`).

    Память переводов — единственное место, где чужая ошибка попадала в перевод,
    минуя и глоссарий, и все проверки: совпадение отдавалось как есть. Записи
    неизвестного происхождения (импорт, `draft`) остаются в промпте справкой,
    а перевод делает модель — и его есть чем проверить."""
    return bool(t) and (t.get("quality") or "draft") == GLOSSARY_TIER_HARD


# ─── Предметные области ──────────────────────────────────────────────
# Сервис не привязан к медицине: область — параметр проекта. Это ЕДИНСТВЕННОЕ
# место, где живёт доменная специфика промптов (перевод и проверка терминов).
# Добавили направление — добавили строку, больше править нечего.
#   expert      — кем модель себя считает при переводе;
#   terminology — что считать эталоном терминологии;
#   examples    — типичные кальки этой области (можно пусто).
DOMAINS = [
    {"id": "medical", "label": "Медицина", "en": "medical",
     "cats": ["Anatomy", "Cardiology", "Disease", "Dosage", "Symptom", "Lab", "Procedure", "Device", "Document"],
     "extract": "Domain terminology only: diseases, anatomy, procedures, drugs, lab tests, devices.",
     "expert": "medical translator specializing in biomedical and clinical texts",
     "terminology": "standard medical terminology as used in peer-reviewed clinical literature",
     "examples": "BAD: 'oxide nitrogena', 'leukocidin', 'rear cyclitis'. "
                 "GOOD: 'nitric oxide', 'leukocytes', 'posterior cyclitis'."},
    {"id": "pharma", "label": "Фармацевтика", "en": "pharmaceutical",
     "cats": ["Substance", "Dosage", "Form", "Route", "Regulatory", "Trial", "Document"],
     "extract": "Domain terminology only: substances, dosage forms, routes, regulatory and trial terms.",
     "expert": "pharmaceutical translator working on drug labels, SmPCs and clinical trial documents",
     "terminology": "standard pharmaceutical and regulatory terminology (INN names, dosage forms, routes)",
     "examples": ""},
    {"id": "legal", "label": "Юриспруденция", "en": "legal",
     "cats": ["Contract", "Court", "Corporate", "Property", "Obligation", "Party", "Document"],
     "extract": "Domain terminology only: legal concepts, instruments, procedural and corporate terms.",
     "expert": "legal translator working on contracts, court documents and corporate filings",
     "terminology": "standard legal terminology of the target language, keeping the legal effect intact",
     "examples": ""},
    {"id": "technical", "label": "Техника", "en": "technical",
     "cats": ["Part", "Material", "Process", "Measurement", "Safety", "Equipment", "Document"],
     "extract": "Domain terminology only: parts, materials, processes, measurements, equipment.",
     "expert": "technical translator working on engineering documentation and manuals",
     "terminology": "standard engineering terminology and unit conventions",
     "examples": ""},
    {"id": "finance", "label": "Финансы", "en": "financial",
     "cats": ["Accounting", "Reporting", "Tax", "Instrument", "Metric", "Audit", "Document"],
     "extract": "Domain terminology only: accounting, reporting, tax and instrument terms.",
     "expert": "financial translator working on reports, statements and audit documents",
     "terminology": "standard accounting and financial terminology",
     "examples": ""},
    {"id": "it", "label": "IT", "en": "software",
     "cats": ["UI", "API", "Data", "Security", "Infrastructure", "Process", "Document"],
     "extract": "Domain terminology only: interface, API, data, security and infrastructure terms.",
     "expert": "software localization specialist",
     "terminology": "established terminology of the platform and the target locale",
     "examples": ""},
    {"id": "general", "label": "Общая тематика", "en": "general-purpose",
     "cats": ["Term", "Name", "Document"],
     "extract": "Only stable multi-word terms and named entities; skip ordinary vocabulary.",
     "expert": "professional translator",
     "terminology": "standard contemporary usage",
     "examples": ""},
]
# Область НОВОГО проекта — общая: сервис не про медицину, и подставлять
# её человеку, который пришёл с договором, нельзя. А вот запись и проект
# БЕЗ поля `domain` читаются как медицинские (LEGACY_DOMAIN) — это закон
# миграции, тот же, что у `lang`: у первого клиента поля нет ни у проекта,
# ни у 1307 записей глоссария, и смена этого чтения оторвала бы глоссарий
# от его проекта. Разводить эти два смысла обязательно: «чем заполнить
# пустое поле сейчас» и «как понимать пустое поле в старых данных».
DEFAULT_DOMAIN = "general"
_DOMAINS_BY_ID = {d["id"]: d for d in DOMAINS}


# ─── Каталог языков ──────────────────────────────────────────────────
# Пара проекта может быть ЛЮБОЙ: список живёт файлом, а не в .jsx (там было
# пять языков с флагами стран). Код — ISO 639-1 в верхнем регистре, тот же,
# что ключует таблицы окончаний и маркеры отрицания; то, чего в каталоге
# нет, проектом стать не может — иначе «язык» станет произвольной строкой,
# и ни одна таблица к нему не подойдёт. Флагов нет намеренно: флаг — это
# страна, а не язык.
def _load_languages() -> list:
    try:
        with open(ROOT / "backend" / "languages.json", encoding="utf-8") as f:
            langs = json.load(f)["languages"]
    except Exception as e:
        print(f"[backend] WARN: languages.json не прочитан: {e}", file=sys.stderr)
        langs = [{"code": "RU", "ru": "Русский", "en": "Russian", "native": "Русский", "script": "CYRILLIC"},
                 {"code": "EN", "ru": "Английский", "en": "English", "native": "English", "script": "LATIN"}]
    return sorted(langs, key=lambda l: l["ru"])


LANGUAGES = _load_languages()
_LANG_BY_CODE = {l["code"]: l for l in LANGUAGES}


def _check_lang_pair(src: str, tgt: str) -> tuple:
    """Коды языков проекта — из каталога, и пара не вырождена. Ответ 400,
    а не молчаливое «RU» по умолчанию: проект с языком «Русский→English»
    не найдёт ни одной таблицы, и все проверки на нём просто умолкнут."""
    s, t = (src or "").strip().upper(), (tgt or "").strip().upper()
    if s not in _LANG_BY_CODE or t not in _LANG_BY_CODE:
        raise HTTPException(400, f"Неизвестный код языка: {src!r} → {tgt!r}. "
                                 "Ожидается код ISO 639-1 из каталога (/api/models → languages).")
    if s == t:
        raise HTTPException(400, "Язык оригинала и язык перевода совпадают.")
    return s, t


# Свои области организации (`STATE["domains"]`): те же поля, что у встроенных,
# потому что промпты читают ровно их. Встроенные — неудаляемый шаблон.
# Ищутся ПЕРВЫМИ и только в своей организации: одноимённая область другого
# клиента не видна. Правило области для автоодобрения — поле `strict`
# (приказ только от человека), для самодельной по умолчанию включено:
# незнакомая область не должна получать право самоодобрения молча.
def _tenant_domains(tenant: Optional[str] = None) -> list:
    t = tenant or _current_tenant()
    return [d for d in STATE.get("domains") or [] if _tenant_of(d) == t]


def _resolve_domain(domain_id: Optional[str]) -> dict:
    """Своя область организации, иначе встроенная, иначе LEGACY_DOMAIN.
    Пустое поле — это СТАРАЯ запись (см. DEFAULT_DOMAIN), а не «общая
    тематика»: у новых проектов область проставлена всегда."""
    if domain_id:
        for d in _tenant_domains():
            if d.get("id") == domain_id:
                return d
    return _DOMAINS_BY_ID.get(domain_id or "") or _DOMAINS_BY_ID[LEGACY_DOMAIN]


# ─── Каталог моделей OpenAI ──────────────────────────────────────────
# Цены — USD за 1M токенов, сверено с developers.openai.com/api/docs/pricing 15.08.2026.
# ЕДИНСТВЕННОЕ место, где правятся модели и цены: добавили модель — добавили строку.
# "api": "modern" — семейство GPT-5.x: НЕ принимает max_tokens/temperature, нужен
# max_completion_tokens (проверено вызовами к API 15.08.2026). "classic" — GPT-4.x.
OPENAI_MODELS = [
    {"id": "gpt-5.6-sol",   "label": "GPT-5.6 Sol",   "in": 5.00, "out": 30.00, "api": "modern",  "note": "Флагман, максимальное качество"},
    {"id": "gpt-5.6-terra", "label": "GPT-5.6 Terra", "in": 2.00, "out": 12.00, "api": "modern",  "note": "Баланс качества и цены"},
    {"id": "gpt-5.6-luna",  "label": "GPT-5.6 Luna",  "in": 0.20, "out": 1.20,  "api": "modern",  "note": "Быстрая и недорогая"},
    {"id": "gpt-5.5",       "label": "GPT-5.5",       "in": 5.00, "out": 30.00, "api": "modern",  "note": "Предыдущий флагман"},
    {"id": "gpt-5.4",       "label": "GPT-5.4",       "in": 2.50, "out": 15.00, "api": "modern",  "note": ""},
    {"id": "gpt-5.4-mini",  "label": "GPT-5.4 mini",  "in": 0.75, "out": 4.50,  "api": "modern",  "note": ""},
    {"id": "gpt-4.1",       "label": "GPT-4.1",       "in": 2.00, "out": 8.00,  "api": "classic", "note": ""},
    {"id": "gpt-4o",        "label": "GPT-4o",        "in": 2.50, "out": 10.00, "api": "classic", "note": "По умолчанию"},
    {"id": "gpt-4o-mini",   "label": "GPT-4o mini",   "in": 0.15, "out": 0.60,  "api": "classic", "note": "Самая дешёвая"},
]
DEFAULT_OPENAI_MODEL = "gpt-4o"

# Название сервиса — из окружения: «Medical» в шапке у клиента-юриста
# выглядит как чужой продукт. Отдаётся через /api/models.
APP_BRAND = os.environ.get("APP_BRAND", "").strip() or "CAT Translator"
_MODELS_BY_ID = {m["id"]: m for m in OPENAI_MODELS}

# ── Справочник силы моделей ──────────────────────────────────────────────────
# Нужен для одного решения: вправе ли проверка одной моделью перезаписать
# готовую проверку, сделанную другой. Более слабая не вправе — иначе вердикт
# Sol на тысяче сегментов молча заменяется вердиктом Terra, и человек платит
# за понижение качества.
#
# Ранг лежит ФАЙЛОМ, а не в коде: список моделей меняется чаще, чем выходят
# релизы сервиса, и обновление ранга не должно требовать деплоя. Копия в
# data/ сильнее репозиторной: её правят на сервере, и git pull её не затрёт.
# Из цены ранг не выводится (GPT-5.5 стоит как Sol и слабее его), из даты
# тоже (у одного поколения три разных уровня) — только руками.
MODEL_RANK_FILES = [ROOT / "backend" / "model_ranks.json",
                    ROOT / "backend" / "data" / "model_ranks.json"]
_MODEL_RANKS: dict = {}
_MODEL_RANKS_STAMP: tuple = ()
_MODEL_RANKS_CHECKED: float = 0.0
# Как часто ходим на диск за временем правки. Ранг спрашивают на каждый сегмент
# каждого шага: на проекте в 2670 строк это 37 000 вызовов stat() на один разбор
# прогона — половина его времени уходила в файловую систему. Две секунды — это
# по-прежнему «правка подхватывается без рестарта», но уже не по разу на сегмент.
MODEL_RANKS_RECHECK_SEC = 2.0


def _model_ranks() -> dict:
    """Справочник с диска. Перечитывается, когда файл менялся: смысл выносить
    его из кода теряется, если для правки нужен рестарт сервиса."""
    global _MODEL_RANKS, _MODEL_RANKS_STAMP, _MODEL_RANKS_CHECKED
    now = time.monotonic()
    # Пустой _MODEL_RANKS_STAMP — просьба перечитать немедленно (так это делает
    # тест). Обычный путь: не чаще раза в MODEL_RANKS_RECHECK_SEC.
    if (_MODEL_RANKS and _MODEL_RANKS_STAMP
            and (now - _MODEL_RANKS_CHECKED) < MODEL_RANKS_RECHECK_SEC):
        return _MODEL_RANKS
    _MODEL_RANKS_CHECKED = now
    stamp = []
    for p in MODEL_RANK_FILES:
        try:
            stamp.append(p.stat().st_mtime_ns)
        except OSError:
            stamp.append(0)
    stamp = tuple(stamp)
    if stamp == _MODEL_RANKS_STAMP and _MODEL_RANKS:
        return _MODEL_RANKS
    ranks: dict = {}
    for p in MODEL_RANK_FILES:          # порядок важен: data/ идёт вторым и побеждает
        try:
            data = json.loads(p.read_text(encoding="utf-8"))
        except FileNotFoundError:
            continue
        except Exception as e:
            # Битый справочник не роняет сервис и не подменяется пустым молча:
            # без рангов защита от понижения выключится, и об этом надо знать.
            print(f"[backend] справочник рангов {p} не прочитан: {e}", file=sys.stderr)
            continue
        for mid, rank in (data.get("ranks") or {}).items():
            if isinstance(rank, int) and not isinstance(rank, bool):
                ranks[str(mid)] = rank
    _MODEL_RANKS, _MODEL_RANKS_STAMP = ranks, stamp
    return ranks


def model_rank(mid: str) -> Optional[int]:
    """None — «не знаю», а не «слабая». Неизвестность не даёт права ни
    перезаписать чужую проверку, ни назвать её действительной."""
    return _model_ranks().get(mid or "")


def _rank_not_weaker(have: str, want: str) -> bool:
    """Проверка моделью have не слабее той, что дала бы want.

    Неизвестный ранг с любой стороны — False, то есть «проверить заново».
    Обратный выбор был бы хуже: сегмент остался бы с вердиктом модели, про
    которую мы ничего не знаем, а человек читал бы его как вердикт выбранной.
    Цена ошибки — один вызов и строка в разборе прогона, после которой модель
    дописывают в справочник."""
    if have and have == want:
        return True
    rh, rw = model_rank(have), model_rank(want)
    if rh is None or rw is None:
        return False
    return rh >= rw

# seg["provider"] — чем сегмент переведён по факту: id модели OpenAI, либо эти константы.
# Поле проставляется в момент перевода; у сегментов, переведённых до его появления,
# его нет, и фронтенд показывает приблизительное значение по seg["route"].
# Осталось только для чтения истории: старые сегменты помечены этим
# провайдером и маршрутом GOOGLE_SAFE. Новый перевод так не помечается
# никогда — движок один, выбранная модель.
PROVIDER_GOOGLE = "google"
PROVIDER_TM = "tm"

# Для обратного перевода нужна не лучшая модель, а самая буквальная и дешёвая:
# её задача — зеркалить текст, а не переводить его хорошо. Чем «умнее» модель,
# тем охотнее она чинит ошибки на лету и прячет их от проверки.
BACKCHECK_DEFAULT_MODEL = "gpt-5.6-luna"
# Запасная — когда запрошенная модель совпала с автором текста (см.
# _backcheck_model). Требования те же: буквальная и дешёвая.
BACKCHECK_FALLBACK_MODEL = "gpt-4o-mini"

# Семантическая близость оригинала и обратного перевода. Лексическая база не умеет
# в синонимы: «больному назначен» против «пациенту назначили» — это один смысл и
# разные основы. Эмбеддинги закрывают ровно этот разрыв.
EMBED_MODEL = "text-embedding-3-small"
# Судья вызывается только в средней зоне: наверху и внизу шкалы решение уже принято
# детерминированными проверками, и платить за подтверждение очевидного незачем.
# Полоса баллов, в которой вызов судьи имеет смысл. НИЗ: ниже него решение
# уже принято детерминированной находкой. ВЕРХ: выше него спорить не о чем —
# так считалось, пока не выяснилось, что «бугорка → cusps» и «Prevalence
# (prevalence)» имеют балл 100 и к судье не попадают никогда. Потолок поэтому
# вынесен в окружение: поднять его — решение с прямой ценой (каждый сегмент
# зоны стоит вызова), и принимать его должен человек, а не автор кода.
JUDGE_ZONE = (int(os.environ.get("JUDGE_ZONE_LOW", "50")),
              int(os.environ.get("JUDGE_ZONE_TOP", "97")))


def _lex_blind(source: str) -> bool:
    """«Оригинал слишком короток, чтобы лексическая мера что-то значила».

    Через getattr, а не прямым вызовом: checks подключается через
    _safe_import и может отсутствовать или оказаться старее кода (то же
    правило, что у BACKCHECK_BANDS). Молчаливый ответ здесь — False, то есть
    «ничего не меняем»: не зная длины, безопаснее оставить прежнюю зону судьи
    и прежнюю подпись корзины, чем открыть низ шкалы наугад и платить за
    судью по всему проекту."""
    fn = getattr(checks_mod, "lexically_blind", None) if checks_mod else None
    return bool(fn(source or "")) if fn else False


def _bc_version() -> int:
    """Версия ПРАВИЛ подсчёта back-check (checks.BACKCHECK_VERSION).

    Через getattr по той же причине, что и _lex_blind: checks подключается
    через _safe_import и может отсутствовать. Молчаливый ответ — 0, то есть
    «версия неизвестна»: записи с таким числом никто не клеймит (у прежних
    ключа `v` нет вовсе), и ни одна проверка на равенство не сработает
    случайно."""
    v = getattr(checks_mod, "BACKCHECK_VERSION", 0) if checks_mod else 0
    try:
        return int(v or 0)
    except (TypeError, ValueError):                          # pragma: no cover
        return 0


def _judge_zone(source: str, judge_all: bool = False) -> tuple:
    """Зона вызова судьи для КОНКРЕТНОГО сегмента.

    Низ шкалы закрыт потому, что там решение уже принято детерминированной
    проверкой. Но на коротком оригинале никакого решения принято не было:
    доля выживших основ при одном-двух содержательных словах даёт только 0 или
    1, и любой синоним в обратном переводе роняет балл в ноль при верном
    переводе (см. lexically_blind). Ноль в этом случае значит «нечем измерить»,
    а не «перевод не тот» — и это ровно тот вопрос, на который отвечает судья.
    Поэтому для таких сегментов низ зоны открыт.

    Верх не двигаем: выше 97 спорить не о чем при любой длине.
    Жёсткая находка (числа, единицы, отрицание) судью по-прежнему отменяет —
    её он не вправе отменить, и проверяется она отдельно от зоны.

    `judge_all` — разовое разрешение прогона поднять ВЕРХ до предела: балл
    выше потолка означает лишь, что довольны детерминированные проверки,
    а смысл при этом не читал никто — там и живёт «беглое неверное слово»
    (monostable, sanguiferous), у которого другой меры нет. Низ разрешение
    не открывает: ниже низа решение уже вынесено, и спорить не о чем.

    Потолок разрешения живёт ЗДЕСЬ и только здесь. Раньше `hi = 100` стояло
    отдельными строками у каждого спрашивающего, а зона — это правило, у
    которого обязано быть одно место: разойдись копии (например, когда
    потолок разрешения станет ниже 100 или переедет в окружение), разбор
    обещал бы один состав, а прогон делал другой."""
    lo, hi = JUDGE_ZONE
    if judge_all:
        hi = 100
    return (0, hi) if _lex_blind(source) else (lo, hi)

# У судьи СВОЯ модель, отдельная от модели обратного перевода, и это намеренно:
# задачи прямо противоположные. Обратному переводу нужна максимально буквальная
# и тупая модель, которая не чинит ошибки; судье — наоборот, сильная, способная
# отличить подмену понятия от синонима. Одна модель на обе роли работала бы плохо
# в одной из них.
JUDGE_DEFAULT_MODEL = "gpt-5.6-terra"
# Проверке терминологии нужна сильная модель: слабая либо пропускает кальки,
# либо начинает придираться к нормальным синонимам. Дефолт перевода (gpt-4o)
# для этой роли слабоват, поэтому берём ту же модель, что и судья.
TERMCHECK_DEFAULT_MODEL = os.environ.get("TERMCHECK_MODEL", JUDGE_DEFAULT_MODEL)
# Ремонт переписывает текст по списку претензий — это работа для сильной модели.
REPAIR_DEFAULT_MODEL = os.environ.get("REPAIR_MODEL", JUDGE_DEFAULT_MODEL)


def _openai_embed(texts: list) -> list:
    import openai
    client = openai.OpenAI(api_key=os.environ.get("OPENAI_API_KEY"), timeout=60, max_retries=2)
    resp = client.embeddings.create(model=EMBED_MODEL, input=texts)
    _note_usage("embed", EMBED_MODEL, resp)
    return [d.embedding for d in resp.data]


def _cosine(a: list, b: list) -> float:
    dot = sum(x * y for x, y in zip(a, b))
    na = sum(x * x for x in a) ** 0.5
    nb = sum(x * x for x in b) ** 0.5
    return (dot / (na * nb)) if na and nb else 0.0


# Тексты повторяются: один и тот же оригинал эмбеддится на каждом перепрогоне
# back-check и в каждом дубле сегмента. Кэш живёт в памяти процесса.
_EMBED_CACHE: dict = {}
_EMBED_CACHE_MAX = 4000


def _embed_cached(texts: list) -> list:
    keys = [_text_hash(t) for t in texts]
    missing = [t for t, k in zip(texts, keys) if k not in _EMBED_CACHE]
    if missing:
        uniq = list(dict.fromkeys(missing))
        for t, vec in zip(uniq, _openai_embed(uniq)):
            _EMBED_CACHE[_text_hash(t)] = vec
        if len(_EMBED_CACHE) > _EMBED_CACHE_MAX:
            for k in list(_EMBED_CACHE)[:len(_EMBED_CACHE) - _EMBED_CACHE_MAX]:
                _EMBED_CACHE.pop(k, None)
    return [_EMBED_CACHE[k] for k in keys]


def _semantic_similarity(source_ru: str, back_ru: str):
    """Косинус между оригиналом и обратным переводом. None — если посчитать не вышло."""
    if not os.environ.get("OPENAI_API_KEY"):
        return None
    try:
        vecs = _embed_cached([source_ru, back_ru])
        return _cosine(vecs[0], vecs[1])
    except Exception as e:
        print(f"[backend] embed failed: {e}", file=sys.stderr)
        return None


def _judge_system(domain: dict, src_lang: str) -> str:
    """Промпт судьи. Раньше он был зашито медицинским и русским: «Ты —
    медицинский редактор», примеры про лимфаденит, «текст на русском».
    Для юридического проекта на немецкий это мешало, а не помогало."""
    return (
        "Ты — редактор перевода, специализация: " + domain["label"].lower() + ". "
        "Тебе дают исходный текст (язык: " + src_lang + ") и его ОБРАТНЫЙ перевод "
        "(текст перевели на другой язык, затем обратно). Твоя задача — понять, "
        "сохранился ли смысл.\n\n"
        "Считай расхождением: подмену понятия или термина на другое (даже похожее по "
        "звучанию — это разные вещи), изменение числа, количества, единицы, отрицания, "
        "стороны, направления, степени уверенности утверждения.\n"
        "НЕ считай расхождением: синонимы, изменённый порядок слов, стилистические "
        "различия, разные грамматические формы, если смысл тот же.\n\n"
        "Верни ТОЛЬКО JSON без пояснений:\n"
        '{"same_meaning": true|false, "severity": "none"|"minor"|"major"|"critical", '
        '"divergences": ["one short sentence in ' + _explain_lang_name() + '"], '
        '"comment": "one sentence in ' + _explain_lang_name() + '"}\n'
        "severity: none — смысл идентичен; minor — стилистика; major — заметное смысловое "
        "расхождение; critical — подмена понятия, числа, отрицания или стороны."
    )


# ─── Проверка терминологии перевода ──────────────────────────────────
# Back-check спрашивает «пережил ли смысл круг» и на кальке всегда отвечает
# «да»: «rear cyclitis» дословно возвращается как «задний циклит» и совпадает
# с оригиналом. Здесь задан противоположный вопрос — «нормальный ли это термин
# целевого языка», и смотрим мы ТОЛЬКО на перевод, оригинал нужен лишь для
# привязки термина. Область берётся из проекта: медицина ничем не выделена.
TERMCHECK_SEVERITY = ["critical", "major", "minor"]


def _termcheck_system(domain: dict, src_lang: str, tgt_lang: str) -> str:
    return (
        "You are a terminology reviewer for " + domain["en"] + " translations from "
        + src_lang + " into " + tgt_lang + ".\n"
        "You get SOURCE and TRANSLATION. Judge ONLY the terminology of the TRANSLATION.\n\n"
        "Flag a term when it is:\n"
        "  - a calque or word-by-word rendering that is not a real term in " + tgt_lang + ";\n"
        "  - a transliteration of a source-language or Latin word instead of the accepted term;\n"
        "  - a different concept than the source term (substitution);\n"
        "  - garbled, truncated or fused with digits/other words;\n"
        "  - wrong register for " + domain["en"] + " documents (everyday word instead of the professional term).\n\n"
        "DO NOT flag: style preferences, synonyms that are both standard, "
        "British vs American spelling, sentence structure, punctuation, anything in the SOURCE.\n"
        "Be conservative: if unsure whether a term is standard, do NOT flag it.\n"
        "The suggestion must be a term actually used in " + tgt_lang + " " + domain["en"]
        + " literature, never your own invention.\n\n"
        'Return ONLY JSON, no prose:\n'
        '{"findings": [{"src_term": "<the matching fragment of SOURCE, or empty>", '
        '"tgt_term": "<the exact fragment of TRANSLATION that is wrong>", '
        '"suggestion": "<the correct term>", "severity": "critical|major|minor", '
        '"why": "<one short sentence in ' + _explain_lang_name() + '>"}]}\n'
        "severity: critical — a different concept or an unreadable fragment; "
        "major — not a real term of the target language; minor — understandable but non-standard.\n"
        'If the terminology is fine, return {"findings": []}.'
    )


def _openai_termcheck(source: str, target: str, src_lang: str, tgt_lang: str,
                      domain_id: Optional[str] = None, model: str = None) -> Optional[dict]:
    """Разбор перевода моделью. None — вызов не удался (сегмент не трогаем)."""
    import json as _json
    import openai
    dom = _resolve_domain(domain_id)
    mdl = _resolve_model(model or TERMCHECK_DEFAULT_MODEL)
    client = openai.OpenAI(api_key=os.environ.get("OPENAI_API_KEY"), timeout=90, max_retries=1)
    extra = ({"max_completion_tokens": 2048} if mdl["api"] == "modern"
             else {"max_tokens": 700, "temperature": 0})
    try:
        resp = client.chat.completions.create(
            model=mdl["id"],
            messages=[
                {"role": "system", "content": _termcheck_system(dom, src_lang, tgt_lang)},
                {"role": "user", "content": "SOURCE:\n" + source + "\n\nTRANSLATION:\n" + target},
            ],
            **extra,
        )
        _note_usage("termcheck", mdl["id"], resp)
        raw = (resp.choices[0].message.content or "").strip()
        m = re.search(r"\{.*\}", raw, re.S)
        if not m:
            return None
        data = _json.loads(m.group(0))
        out = []
        for f in (data.get("findings") or []):
            if not isinstance(f, dict):
                continue
            tgt_term = (f.get("tgt_term") or "").strip()
            if not tgt_term:
                continue
            sev = (f.get("severity") or "major").lower()
            out.append({
                "src_term": (f.get("src_term") or "").strip(),
                "tgt_term": tgt_term,
                "suggestion": (f.get("suggestion") or "").strip(),
                "severity": sev if sev in TERMCHECK_SEVERITY else "major",
                "why": (f.get("why") or "").strip(),
            })
        return {"findings": out, "model": mdl["id"]}
    except Exception as e:
        print(f"[backend] termcheck failed: {e}", file=sys.stderr)
        return None


def _openai_judge(source_ru: str, back_ru: str, model: str = None,
                  domain_id: Optional[str] = None, src_lang: str = "RU") -> Optional[dict]:
    """Вердикт модели по паре «оригинал / обратный перевод»."""
    import json as _json
    import openai
    dom = _resolve_domain(domain_id)
    mdl = _resolve_model(model or JUDGE_DEFAULT_MODEL)
    client = openai.OpenAI(api_key=os.environ.get("OPENAI_API_KEY"), timeout=90, max_retries=1)
    extra = ({"max_completion_tokens": 2048} if mdl["api"] == "modern"
             else {"max_tokens": 500, "temperature": 0})
    try:
        resp = client.chat.completions.create(
            model=mdl["id"],
            messages=[
                {"role": "system", "content": _judge_system(dom, src_lang)},
                {"role": "user", "content": "ОРИГИНАЛ:\n" + source_ru + "\n\nОБРАТНЫЙ ПЕРЕВОД:\n" + back_ru},
            ],
            **extra,
        )
        _note_usage("judge", mdl["id"], resp)
        raw = (resp.choices[0].message.content or "").strip()
        # Модель иногда оборачивает JSON в ```json ... ``` — вырезаем тело
        m = re.search(r"\{.*\}", raw, re.S)
        if not m:
            return None
        verdict = _json.loads(m.group(0))
        verdict["model"] = mdl["id"]
        return verdict
    except Exception as e:
        print(f"[backend] judge failed: {e}", file=sys.stderr)
        return None


def _resolve_model(model_id: Optional[str]) -> dict:
    """Неизвестная/пустая модель → дефолт. Клиент не может подсунуть произвольную строку."""
    return _MODELS_BY_ID.get(model_id or "") or _MODELS_BY_ID[DEFAULT_OPENAI_MODEL]


# ─── Учёт фактического расхода ───────────────────────────────────────────────
# Смета до прогона считается по объёму текста и обязана ошибаться: сколько
# модель ответит, знает только сама модель. Беда была не в этом, а в том, что
# фактический расход НИГДЕ не записывался: поправить смету было не по чему —
# сравнивать не с чем. Так и жила ошибка termcheck: смета клала 450 выходных
# токенов на сегмент, а ответ занимал 37.
#
# Поэтому usage снимается с КАЖДОГО ответа модели. Это не оценка и не наш
# пересчёт объёма текста — это то, за что выставит счёт провайдер.
#
# Куда писать — одна переменная процесса, а не контекст вызова: прогон
# в системе один (тот же инвариант, что и один воркер uvicorn), а
# _run_parallel раздаёт работу в ThreadPoolExecutor, который contextvars
# не наследует, так что contextvars сюда просто не доедут.
#
# Учёт не имеет права ломать работу: всё, что он делает, обёрнуто в try.
# Перевод, не состоявшийся из-за бухгалтерии, — цена, несоизмеримая
# с точностью сметы.

# Модели, которые человек не выбирает, но которые стоят денег. В OPENAI_MODELS
# им не место: оттуда строится выпадающий список, и эмбеддингом никто не
# переводит. Цена — та же единица измерения, USD за 1M токенов.
AUX_MODEL_PRICES = {EMBED_MODEL: {"in": 0.02, "out": 0.0}}

RUN_COST_HISTORY = 100          # сколько прогонов помним в state.json

_USAGE_LOCK = threading.Lock()
_USAGE_SINK: Optional[dict] = None      # счётчик идущего прогона; None — вне прогона


def _usage_zero() -> dict:
    return {"calls": 0, "in": 0, "cached_in": 0, "out": 0, "reasoning": 0,
            "cost": 0.0, "unpriced": 0, "steps": {}, "models": {}}


def _usage_leaf() -> dict:
    return {"calls": 0, "in": 0, "cached_in": 0, "out": 0, "reasoning": 0,
            "cost": 0.0, "unpriced": 0}


# Расход процесса с момента старта. Прогоны живут в памяти и теряются при
# рестарте, а одиночные вызовы (перевод сегмента по кнопке) не принадлежат
# ни одному прогону — но деньги стоят.
_USAGE_TOTAL: dict = _usage_zero()


def _model_price(mid: Optional[str]) -> Optional[dict]:
    m = _MODELS_BY_ID.get(mid or "")
    return {"in": m["in"], "out": m["out"]} if m else AUX_MODEL_PRICES.get(mid or "")


def _usage_cost(mid: Optional[str], tin: int, tout: int) -> Optional[float]:
    """None — «цена неизвестна», а не ноль. Считать неизвестное нулём значит
    показать расход меньше настоящего — ровно то враньё, ради которого учёт
    и заводится. Такие вызовы считаются отдельно (unpriced).

    Скидка на кэшированный вход не применяется: её цены в каталоге нет,
    а выдумывать цену нельзя. Значит, цифра завышена ровно на неё — насколько,
    видно по cached_in рядом."""
    p = _model_price(mid)
    if not p:
        return None
    return tin / 1e6 * p["in"] + tout / 1e6 * p["out"]


def _usage_field(obj, name: str) -> int:
    """usage приходит объектом SDK, но в разных версиях бывает и словарём."""
    if obj is None:
        return 0
    v = obj.get(name) if isinstance(obj, dict) else getattr(obj, name, 0)
    try:
        return int(v or 0)
    except (TypeError, ValueError):
        return 0


def _usage_part(obj, name: str):
    if obj is None:
        return None
    return obj.get(name) if isinstance(obj, dict) else getattr(obj, name, None)


def _usage_add(bucket: dict, step: str, mid: str, tin: int, cached: int,
               tout: int, think: int, cost: Optional[float]) -> None:
    for d in (bucket,
              bucket["steps"].setdefault(step or "?", _usage_leaf()),
              bucket["models"].setdefault(mid or "?", _usage_leaf())):
        d["calls"] += 1
        d["in"] += tin
        d["cached_in"] += cached
        d["out"] += tout
        d["reasoning"] += think
        if cost is None:
            d["unpriced"] += 1
        else:
            # Округление на каждом шаге, а не в конце: доли цента в сумме
            # из тысяч вызовов накапливают мусор в младших разрядах float,
            # и число перестаёт совпадать само с собой при пересчёте.
            d["cost"] = round(d["cost"] + cost, 6)


# ─── Расход и лимит по организации ───────────────────────────────────
# Факт расхода уже снимается с каждого ответа модели; здесь он ещё и
# складывается по организации и месяцу (`STATE["spend"][tenant][YYYY-MM]`).
# Лимит (`tenant["limitUsd"]`, ставит суперпользователь) режет ДЕНЬГИ, а не
# работу: платные команды отвечают 402 с остатком и датой сброса, а всё
# бесплатное — правка начертания, откат правок ремонта, пересчёт back-check,
# принятие кандидатов, разбор состава, экспорт — работает и на исчерпанном
# лимите. Иначе лимит превращается в «сервис сломался», и человек не может
# забрать то, за что уже заплатил. Цена неизвестна — в расход не идёт
# (считается `unpriced`), потому что неизвестное, посчитанное нулём, — это
# расход меньше настоящего, а посчитанное наугад — отказ по выдуманному числу.
def _next_batch_seq() -> int:
    """Номер пачки автоодобрения. В базе — атомарный счётчик: пачку может
    завести и API, и воркер (apply_terms), и два процесса не должны выдать
    один номер. floor — прежний счётчик из state.json, ниже не опускаемся."""
    if STORE.kind == "pg":
        try:
            return STORE.next_counter("autoBatchSeq", int(STATE.get("autoBatchSeq") or 0))
        except Exception as e:
            print(f"[backend] счётчик пачек из базы не взялся: {e}", file=sys.stderr)
    batch = STATE.get("autoBatchSeq", 0) + 1
    STATE["autoBatchSeq"] = batch
    return batch


def _month_key() -> str:
    return datetime.now().strftime("%Y-%m")


def _spend_add(tenant: str, cost: Optional[float]) -> None:
    # В базе расход — счётчик с прямым инкрементом (store.add_spend):
    # снимок словаря из двух процессов терял бы приращения друг друга.
    if STORE.kind == "pg":
        try:
            STORE.add_spend(tenant or DEFAULT_TENANT, _month_key(), cost)
            return
        except Exception as e:
            print(f"[backend] расход не записан в базу: {e}", file=sys.stderr)
    sp = STATE.setdefault("spend", {}).setdefault(tenant or DEFAULT_TENANT, {})
    m = sp.setdefault(_month_key(), {"usd": 0.0, "calls": 0, "unpriced": 0})
    m["calls"] += 1
    if cost is None:
        m["unpriced"] += 1
    else:
        m["usd"] = round(m["usd"] + cost, 6)


def _tenant_rec(tid: str) -> Optional[dict]:
    return next((t for t in _tenants() if t.get("id") == tid), None)


def _spend_status(tenant: Optional[str] = None) -> dict:
    t = tenant or _current_tenant()
    if STORE.kind == "pg":
        try:
            m = STORE.get_spend(t, _month_key())
        except Exception as e:
            print(f"[backend] расход не прочитан из базы: {e}", file=sys.stderr)
            m = {"usd": 0.0, "calls": 0, "unpriced": 0}
    else:
        m = (STATE.get("spend") or {}).get(t, {}).get(_month_key()) or {"usd": 0.0, "calls": 0, "unpriced": 0}
    rec = _tenant_rec(t) or {}
    limit = rec.get("limitUsd")
    return {"tenant": t, "month": _month_key(), "spentUsd": round(m["usd"], 4),
            "calls": m["calls"], "unpriced": m["unpriced"], "limitUsd": limit,
            "over": bool(limit is not None and m["usd"] >= float(limit))}


# Что стоит денег: платные команды по путям. Таблица здесь, а не флаг
# на каждом обработчике — как `_OWNER_ONLY`: одна точка, забыть строку видно.
_PAID = [
    ("POST", re.compile(r"/api/projects/\d+/jobs$")),
    # `review` — сам шаг; откат `/review/{stamp}/undo` под эту строку не
    # попадает ($ на конце) и остаётся бесплатным, как и все откаты: модель
    # он не зовёт, а запирать возврат к прежнему тексту на исчерпанном лимите
    # значило бы держать клиента в заложниках у его же счёта.
    ("POST", re.compile(r"/api/projects/\d+/(batch|extract-terms|term-context|review|termcheck/batch|backcheck/batch|images/scan)$")),
    ("POST", re.compile(r"/api/segments/\d+/\d+/(translate|backcheck|termcheck|repair|medical-qa|checks)$")),
    ("POST", re.compile(r"/api/term-queue/\d+/explain$")),
    ("POST", re.compile(r"/api/glossary/audit$")),
]


def _is_paid(method: str, path: str) -> bool:
    return any(m == method and rx.match(path) for m, rx in _PAID)


def _note_usage(step: str, model_id: str, resp) -> None:
    """Записать то, что посчитал провайдер, а не то, что мы предполагали."""
    try:
        u = getattr(resp, "usage", None) if not isinstance(resp, dict) else resp.get("usage")
        if u is None:
            return
        tin = _usage_field(u, "prompt_tokens")
        tout = _usage_field(u, "completion_tokens")
        cached = _usage_field(_usage_part(u, "prompt_tokens_details"), "cached_tokens")
        think = _usage_field(_usage_part(u, "completion_tokens_details"), "reasoning_tokens")
        if not (tin or tout):
            return
        cost = _usage_cost(model_id, tin, tout)
        with _USAGE_LOCK:
            for bucket in (_USAGE_TOTAL, _USAGE_SINK):
                if bucket is not None:
                    _usage_add(bucket, step, model_id, tin, cached, tout, think, cost)
            _spend_add(_current_tenant(), cost)
    except Exception as e:
        print(f"[backend] учёт расхода не сработал ({step}/{model_id}): {e}", file=sys.stderr)


def _usage_begin(job: dict) -> None:
    global _USAGE_SINK
    with _USAGE_LOCK:
        _USAGE_SINK = job.setdefault("usage", _usage_zero())


def _usage_end(job: dict) -> None:
    """Закрыть счётчик прогона и оставить след, который переживёт рестарт.

    Сами прогоны живут в памяти процесса и теряются при рестарте — это давно
    так и осознанно. Но расход терять нельзя: смету калибруют по нему, а для
    этого нужны десятки прогонов, а не один. Запись компактная: цена по шагам,
    без списков сегментов."""
    global _USAGE_SINK
    with _USAGE_LOCK:
        _USAGE_SINK = None
        u = job.get("usage")
    if not u or not u.get("calls"):
        return
    try:
        rec = {"job": job.get("id"), "kind": job.get("kind"), "project": job.get("project"),
               "tenant": job.get("tenant") or DEFAULT_TENANT,
               "status": job.get("status"), "finished": job.get("finished"),
               "segments": job.get("done"),
               # Смету кладём рядом с фактом: врозь они не сравниваются, а
               # ради сравнения всё и затевалось. Прислал её тот, кто её
               # человеку и показал, — иначе рядом стояли бы два разных числа.
               "est": (job.get("params") or {}).get("est_cost"),
               "cost": u["cost"], "calls": u["calls"], "unpriced": u["unpriced"],
               "in": u["in"], "cached_in": u["cached_in"],
               "out": u["out"], "reasoning": u["reasoning"],
               "steps": {k: {"calls": v["calls"], "cost": v["cost"],
                             "in": v["in"], "out": v["out"], "reasoning": v["reasoning"]}
                         for k, v in u["steps"].items()}}
        hist = STATE.setdefault("runCosts", [])
        hist.append(rec)
        del hist[:-RUN_COST_HISTORY]
    except Exception as e:
        print(f"[backend] расход прогона не записан: {e}", file=sys.stderr)


# Direct OpenAI GPT translation
# Сколько символов соседнего сегмента показывать переводчику. Соседи нужны
# как обстановка, а не как текст для перевода: длинный абзац рядом отвлекает
# и стоит денег, а первых строк хватает, чтобы понять, о чём речь и в каком
# ряду стоит сегмент.
NEIGHBOUR_CHARS = 320


def _neighbours(project: Optional[dict], seg: Optional[dict]) -> tuple:
    """Исходники сегментов ДО и ПОСЛЕ. Пустые строки, если соседей нет.

    Обход списком, а не по id: id не обязаны идти подряд (сегменты удаляют),
    а «сосед» — это про порядок в документе."""
    if not project or not seg:
        return ("", "")
    segs = project.get("segments") or []
    try:
        i = next(k for k, x in enumerate(segs) if x["id"] == seg["id"])
    except StopIteration:
        return ("", "")
    def txt(k):
        return ((segs[k].get("source") or "").strip()[:NEIGHBOUR_CHARS]
                if 0 <= k < len(segs) else "")
    return (txt(i - 1), txt(i + 1))


def _translate_system(src: str, tgt: str, gloss_hits: list, tm_context: dict,
                      literal: bool, domain, mdl: dict,
                      prev_src: str = "", next_src: str = "",
                      style: str = "") -> str:
    """Системный промпт перевода. Вынесен отдельно, чтобы его можно было
    проверить тестом без обращения к модели: от того, каким уровнем уходит
    запись глоссария — приказом или подсказкой, — зависит, повторит ли модель
    чужую ошибку, а такое нельзя оставлять без проверки."""
    if literal:
        system = (
            f"Translate the following text from {src} to {tgt} as literally as possible. "
            "This is a back-translation used for quality control.\n\n"
            "STRICT RULES:\n"
            "1. Return ONLY the translation — no explanations, no comments, no quotes.\n"
            "2. Translate exactly what is written, word for word wherever the language allows.\n"
            "3. Do NOT correct, improve, normalize or standardise anything. If the text contains\n"
            "   an odd, wrong or non-existent term, translate it literally and preserve the oddity.\n"
            "   Your job is to mirror the text, NOT to make it sound correct.\n"
            "4. Preserve every number, unit and negation exactly.\n"
            "5. Keep the word order of the original as close as the target language permits.\n"
        )
        gloss_hits = None
        tm_context = None
    else:
        dom = _resolve_domain(domain)
        system = (
        f"You are a senior {dom['expert']}. "
        f"Translate the following text from {src} to {tgt}.\n\n"
        "STRICT RULES:\n"
        "1. Return ONLY the translated text — no explanations, no comments, no quotes.\n"
        f"2. Use {dom['terminology']}. NEVER transliterate source-language or Latin word forms,\n"
        f"   and never invent word-by-word calques that are not real terms in {tgt}.\n"
        + (f"   {dom['examples']}\n" if dom.get("examples") else "") +
        f"3. NEVER mix languages. Output must be 100% {tgt}. Not a single letter of the\n"
        "   source script may survive — not inside numbers, units, formulas or abbreviations,\n"
        "   and not as a look-alike character.\n"
        "4. NEVER use parenthetical alternatives: NOT 'biologic(al)', NOT 'cell(s)'. Choose ONE correct form.\n"
        "5. NEVER list multiple synonyms separated by semicolons for the same concept.\n"
        "6. Preserve all numbers, abbreviations, and punctuation exactly as in the source.\n"
        f"7. Abbreviations that are identical in {tgt} may be kept as they are.\n"
        # Регистр букв. Без этого правила перевод молча терял заглавную в начале
        # заголовка и подписи, а КАПС заголовка становился обычной строкой: на
        # боевом учебнике так вышло у 36 сегментов. Правило сказано с ДВУХ сторон
        # намеренно — портится регистр в обе (см. оговорку у глоссария ниже):
        # «6. Кавернозный туберкулёз» → «6. cavitary tuberculosis», а
        # «туберкулёза органов дыхания» → «RESPIRATORY TUBERCULOSIS» посреди фразы.
        "8. Follow the capitalisation of the source. A sentence, heading or caption that\n"
        "   starts with a capital letter must start with a capital letter in the translation;\n"
        "   an ALL-CAPS heading stays ALL-CAPS; a word written in lower case inside a sentence\n"
        f"   stays in lower case. Apply {tgt} rules on top of that (proper nouns, nationalities\n"
        "   and months are capitalised even where the source writes them small), but NEVER open\n"
        "   a sentence with a lower-case letter and never shout a word the source does not shout.\n"
        )
        # Стайл-шит документа (`_style_block`) — только в обычный перевод,
        # никогда в обратный: тот обязан ОТРАЖАТЬ текст, а не причёсывать его.
        # Пусто — промпт байт в байт прежний, версии вердиктов не трогаются.
        if style:
            system += "\n" + style
    hard = [h for h in (gloss_hits or []) if _hit_tier(h) == GLOSSARY_TIER_HARD]
    # Подсказки автоимпорта в промпт НЕ уходят — см. блок ниже, где раньше
    # стоял их список. Отбор оставлен: по нему считается строка журнала.
    soft = [h for h in (gloss_hits or []) if _hit_tier(h) == GLOSSARY_TIER_SOFT]

    def _gloss_line(h) -> str:
        """Строка глоссария для промпта: слева — форма, стоящая В ЭТОМ сегменте,
        справа — перевод, уже подогнанный под её начертание.

        Регистр не оставляем на усмотрение модели: запись хранит одно
        начертание, а мест, куда она встаёт, много, и «use these exact
        translations» модель понимает буквально — вместе с заглавной или капсом
        записи. Подгонка детерминированная (`_case_like`), поэтому в промпт
        уходит ровно та строка, которую надо поставить."""
        form = h.get("_form") or h.get("src") or ""
        return "  " + form + " → " + _case_like(form, (h.get("tgt") or ""),
                                                h.get("src") or "")

    if hard:
        terms = "\n".join(_gloss_line(h) for h in hard)
        # Слева — форма ИЗ ЭТОГО сегмента, справа — перевод в её начертании.
        # Раньше сюда уходила словарная запись как есть, и модель копировала её
        # регистр: «Туберкулема → Tuberculoma» ставило заглавную посреди фразы,
        # «ТУБЕРКУЛЕЗ ОРГАНОВ ДЫХАНИЯ → RESPIRATORY TUBERCULOSIS» — капс посреди
        # фразы, «Фиброзно-кавернозный туберкулёз → fibrocavitary tuberculosis»
        # роняло строчную в начало заголовка. Слово при этом верное — испорчено
        # ровно начертание, и правкой записей это не чинится: у записи один
        # регистр, а мест, куда она встаёт, много.
        system += (f"\nApproved glossary — use these exact translations:\n{terms}\n"
                   "On the left is the fragment exactly as it stands in THIS source segment;\n"
                   "on the right the translation is already written in the letter case that\n"
                   "fragment calls for. Copy the right-hand side letter for letter — do not\n"
                   "re-capitalise and do not lower it. Only the grammatical form still follows\n"
                   "the sentence.\n")
    # Терм-лист документа (`tier: "doc"`, см. фазу 0): просьба или правило
    # в зависимости от строгости области; приказ глоссария выше сильнее всегда.
    doc = [] if literal else [h for h in (gloss_hits or []) if h.get("tier") == "doc"]
    if doc:
        system += _termsheet_block("\n".join(_gloss_line(h) for h in doc),
                                   not _auto_policy(domain).get("allow_verified", True))
    # ── Подсказки автоимпорта в промпт не уходят ─────────────────────
    # Здесь стоял блок «Unverified glossary hints (bulk-imported, NOT reviewed
    # — some are wrong)». Его убрали, и вот на каких числах.
    #
    # ПОЛЬЗУ доказать не удалось. Замер на боевом проекте: подсказки попадали
    # в 2316 сегментов из 2711, в среднем 2.6 записи на промпт, и в 72%
    # случаев их вариант оказывался в переводе. Но это не доказательство
    # пользы: «больной → patient» и «мокрота → sputum» модель напишет и без
    # подсказки, а отличить «подсказка помогла» от «подсказка совпала»
    # можно только платным сравнением двух прогонов.
    #
    # ВРЕД доказан и конкретен. Строго: подсказка стоит в переводе И termcheck
    # забраковал ровно это слово — 15 случаев из 11 414 вставок. Каждый
    # из них медицинская ошибка, которой модель послушалась: «лимфаденит →
    # adenolymphitis» (это разные вещи), «пунктат → punctate» (точечный вместо
    # аспирата), «воспалительное → IBD», «ПТК → PTC», «микобактерии
    # туберкулёза → Mycobacteria tuberculosis».
    #
    # Ноль целых тринадцать сотых процента — мало, но чаша весов здесь
    # несимметрична: с одной стороны недоказуемая польза, с другой доказанная
    # подмена понятия в МЕДИЦИНСКОМ переводе. Плюс промпт становится короче
    # на 130 тысяч входных токенов на книгу.
    #
    # Сами записи при этом остаются и работают: они растут в приказ через
    # согласие независимых чистых сегментов, ловят расхождения conflict-
    # кандидатами и разбираются сверкой смысла. Убрано ровно одно — право
    # неверенной записи влиять на перевод напрямую.
    if tm_context:
        system += (
            f"\nTranslation Memory (similar segment, for reference):\n"
            f"  Source: {tm_context.get('src', '')}\n"
            f"  Translation: {tm_context.get('tgt', '')}\n"
        )
    # Соседние сегменты — обстановка, а не задание. До этого сегмент переводился
    # в одиночку, и обрывки списка («Выделяют:», «плевры,», «формирование каверн
    # и распространение процесса.») уходили в модель без всякого признака того,
    # к чему они относятся. Запрет переводить их сказан прямо и дважды: модель
    # добросовестно переписывает всё, что ей дали, — на импорте это уже
    # случалось с соседними абзацами.
    # В literal-режиме соседей НЕТ и быть не должно: обратный перевод обязан
    # ОТРАЖАТЬ текст, а не понимать его. Дай ему обстановку — он починит кривой
    # английский по смыслу соседей и спрячет ровно ту ошибку, которую ищет
    # back-check. Поэтому проверка literal стоит здесь явно.
    if not literal and (prev_src or next_src):
        system += "\nSurrounding context (for disambiguation ONLY — do NOT translate it):\n"
        if prev_src:
            system += f"  [previous segment] {prev_src}\n"
        if next_src:
            system += f"  [next segment] {next_src}\n"
        system += ("Use it to resolve ellipsis, list items, headings and pronouns. "
                   "Translate ONLY the user message.\n")
    if gloss_hits or tm_context:
        print(f"[backend] GPT+context: {len(gloss_hits or [])} gloss, TM={'yes' if tm_context else 'no'}"
              f", model={mdl['id']}", file=sys.stderr)
    # GPT-5.x отвергает max_tokens/temperature; лимит выставлен с запасом, потому что
    # у этого семейства в completion_tokens входят ещё и reasoning-токены.
    return system


def _openai_translate(text: str, src: str, tgt: str,
                      gloss_hits: list = None, tm_context: dict = None,
                      model: str = None, literal: bool = False,
                      domain: Optional[str] = None, step: Optional[str] = None,
                      prev_src: str = "", next_src: str = "",
                      style: str = "") -> str:
    """GPT-перевод с инъекцией глоссария (базовые формы — GPT знает склонения).

    literal=True — режим для обратного перевода. Обычный промпт тут вреден:
    сильная модель видит кривой английский, «чинит» его на лету и возвращает
    правильный русский, маскируя ровно ту ошибку, которую back-check ищет.
    Поэтому в этом режиме требуем дословности и запрещаем править термины,
    а глоссарий и TM не подсовываем вовсе — иначе модель подгонит ответ под них."""
    import openai
    mdl = _resolve_model(model)
    # timeout + retries: зависший вызов не должен блокировать поток бесконечно
    client = openai.OpenAI(api_key=os.environ.get("OPENAI_API_KEY"), timeout=90, max_retries=2)
    system = _translate_system(src, tgt, gloss_hits, tm_context, literal, domain, mdl,
                               prev_src, next_src, "" if literal else style)
    extra = ({"max_completion_tokens": 4096} if mdl["api"] == "modern"
             else {"max_tokens": 1024, "temperature": 0.1})
    resp = client.chat.completions.create(
        model=mdl["id"],
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": text},
        ],
        **extra,
    )
    # Шаг называет вызывающий: буквальный режим — это обратный перевод, но
    # заказывают его двое (back-check и Medical QA), и складывать их расход
    # в одну корзину значит потерять, кто из них сколько стоит.
    _note_usage(step or ("backcheck" if literal else "translate"), mdl["id"], resp)
    return (resp.choices[0].message.content or "").strip()

# ─────────────────────────────────────────────────────────────────────
# Аутентификация
#
# Пользователи и организации лежат в STATE («users», «tenants»). Каждый вход
# выдаёт свой токен, сессия помнит пользователя, организацию и роль, и БЕЗ
# токена не работает ни один /api/* эндпоинт. Токены живут только в памяти
# процесса: рестарт = всем перелогиниться.
#
# Ролей две, и третья появится по первой просьбе, а не заранее:
#   owner       — пользователи, вынос глоссария, удаление проектов, лимиты;
#   translator  — всё остальное, включая платные прогоны.
# Флаг `super` у пользователя — право заводить организации (первый владелец).
#
# Право показать кнопку и право сделать — разные права: роль проверяется
# СЕРВЕРОМ (в `require_token` по таблице `_OWNER_ONLY`), гашение кнопки
# в браузере — удобство. `APP_PASSWORD` остался ТОЛЬКО паролем первого
# владельца при пустой базе пользователей (`_ensure_users`).
# ─────────────────────────────────────────────────────────────────────
import contextvars
import mail_texts                      # тексты писем на языке получателя

_RAW_PASSWORD = os.environ.get("APP_PASSWORD", "").strip()
if not _RAW_PASSWORD:
    # Зашитого дефолта быть не должно: публичный сервис оказался бы открыт
    # каждому, кто читал репозиторий. Без пароля в env — одноразовый случайный.
    _RAW_PASSWORD = secrets.token_urlsafe(9)
    print(f"[backend] WARN: APP_PASSWORD не задан. Пароль на этот запуск: {_RAW_PASSWORD}",
          file=sys.stderr)

# ─── Самостоятельная регистрация по почте ────────────────────────────
# Человек заводит организацию сам: почта + пароль, письмо с кодом,
# подтверждение. Два предохранителя, без которых открытая регистрация —
# это открытый кран к нашему ключу OpenAI и к чужим документам:
#   1) НОВАЯ организация получает лимит расхода SIGNUP_TRIAL_USD (по
#      умолчанию 0 — платное недоступно, пока администратор не поставит
#      лимит). Бесплатные команды и экспорт работают всегда;
#   2) регистрация выключается одной переменной (SIGNUP_ENABLED=0) и
#      ограничена по IP (SIGNUP_MAX_PER_HOUR).
# Вход по почте ИЛИ по логину — обе двери ведут к одному пользователю.
SIGNUP_ENABLED = os.environ.get("SIGNUP_ENABLED", "1").strip().lower() not in ("0", "false", "no", "off")
SIGNUP_TRIAL_USD = float(os.environ.get("SIGNUP_TRIAL_USD", "0") or 0)
SIGNUP_MAX_PER_HOUR = int(os.environ.get("SIGNUP_MAX_PER_HOUR", "5") or 5)
CODE_TTL = int(os.environ.get("AUTH_CODE_TTL_MIN", "30")) * 60
CODE_MAX_TRIES = 5
_EMAIL_RE = re.compile(r"^[^@\s]{1,64}@[^@\s.]+(\.[^@\s.]+)+$")
_SIGNUP_FAILS: dict = {}        # ip -> [время каждой регистрации за час]

BOOTSTRAP_LOGIN = "admin"
# Роли — три. Право читается РАНГОМ («не ниже»): владельцу — всё, включая
# необратимое (удаление записей и проектов, деньги, пользователи); редактор
# и переводчик сегодня в правах РАВНЫ — оба заверяют сегменты, решают
# по терминам, понижают и выносят записи, а роль идёт в СЛЕД ответственного
# («подтвердил: Ева · переводчик»). Это решение
# владельца сервиса, а не забывчивость: первая версия закрывала переводчику
# заверение, и её отменили по прямой просьбе. Что закрыто рангом —
# `_OWNER_ONLY` и (пустая) `_EDITOR_ONLY`.
ROLES = ("owner", "editor", "translator")
ROLE_RANK = {"translator": 0, "editor": 1, "owner": 2}


def _role_at_least(role: Optional[str], need: str) -> bool:
    return ROLE_RANK.get(role or "", -1) >= ROLE_RANK[need]

# Языки интерфейса. Список ЗДЕСЬ, а не в браузере: язык лежит на пользователе
# (`uiLang`), и сервер обязан отказать в том, чего у него нет, — иначе
# в записи окажется код, для которого каталога перевода не существует,
# и человек увидит пустые надписи вместо русских.
# По умолчанию — узбекская латиница: сервис продаётся в Узбекистане.
# Кто выбрал русский, тот его и видит: `uiLang` у него записан явно.
UI_LANGS = ("uz", "ru")
DEFAULT_UI_LANG = os.environ.get("DEFAULT_UI_LANG", "uz").strip().lower() or "uz"
if DEFAULT_UI_LANG not in UI_LANGS:
    DEFAULT_UI_LANG = "uz"
PBKDF2_ITERS = 200_000

# Адрес входа в админку: ADMIN_PATH из окружения, иначе выводится из APP_PASSWORD
# (стабилен для установки, не угадывается). Обфускация входа, не защита.
ADMIN_PATH = re.sub(r"[^A-Za-z0-9._-]", "", os.environ.get("ADMIN_PATH", "").strip()) or (
    "console-" + hashlib.sha256(("admin-console:" + _RAW_PASSWORD).encode()).hexdigest()[:12])
print(f"[backend] админка: /{ADMIN_PATH}", file=sys.stderr)

# Текущая сессия запроса — для мест, куда Request не доезжает (get_project
# и фильтры по организации). В рабочие потоки прогона ContextVar НЕ
# наследуется (`_run_parallel` — ThreadPoolExecutor): задача обязана нести
# организацию в себе и выставлять её сама.
CURRENT_SESSION: "contextvars.ContextVar[Optional[dict]]" = contextvars.ContextVar(
    "session", default=None)


def _hash_password(password: str, salt: Optional[str] = None) -> tuple:
    """pbkdf2, не голый sha256: база паролей клиентов ломается словарём за вечер."""
    salt = salt or secrets.token_hex(16)
    h = hashlib.pbkdf2_hmac("sha256", password.encode(), bytes.fromhex(salt), PBKDF2_ITERS).hex()
    return h, salt


def _verify_password(user: dict, password: str) -> bool:
    if not user or not user.get("active", True):
        return False
    h, _ = _hash_password(password, user.get("salt") or "")
    return hmac.compare_digest(h, user.get("hash") or "")


def _initials(name: str) -> str:
    parts = [p for p in (name or "").split() if p]
    return ("".join(p[0] for p in parts[:2]) or (name or "?")[:2]).upper()


_USER_COLORS = ("#2c7be5", "#22b07d", "#f1a040", "#cc4a4a", "#7a5af8", "#0aa2c0")


def _user_public(u: dict) -> dict:
    return {"id": u["id"], "login": u["login"], "email": u.get("email") or "",
            "emailVerified": bool(u.get("emailVerified")),
            "name": u.get("name") or u["login"],
            "initials": u.get("initials") or _initials(u.get("name") or u["login"]),
            "color": u.get("color") or _USER_COLORS[u["id"] % len(_USER_COLORS)],
            "role": u.get("role", "translator"), "tenant": u.get("tenant", DEFAULT_TENANT),
            "super": bool(u.get("super")), "active": u.get("active", True),
            "uiLang": u.get("uiLang") or DEFAULT_UI_LANG, "created": u.get("created")}


def _users() -> list:
    return STATE.setdefault("users", [])


# ─── Авторство и журнал действий ─────────────────────────────────────
# Кто перевёл, кто подтвердил, кто одобрил термин — медицинский и юридический
# перевод продаются вместе с ответственностью. Отметка «подтвердил человек»
# теперь несёт идентификатор пользователя; прежнее значение "human" остаётся
# ДЕЙСТВИТЕЛЬНЫМ (это факт о том, что заверение было — потерян только автор),
# и читать поле надо ТОЛЬКО через `_confirmed_by_human`: буквальное сравнение
# со строкой молча выключило бы защиту заверений — прогоны начали бы
# переписывать подтверждённый текст, а донор глоссария потерял бы сильнейший
# голос. Журнал — кольцо в STATE (как runCosts): AUDIT_MAX последних записей.
AUDIT_MAX = max(500, int(os.environ.get("AUDIT_MAX", "5000")))


def _actor() -> Optional[dict]:
    sess = CURRENT_SESSION.get()
    if not sess:
        return None
    for u in _users():
        if u["id"] == sess.get("user"):
            return u
    return None


def _actor_id():
    """Что писать в confirmedBy: идентификатор пользователя, а без сессии
    (тесты, миграции) — прежнее "human"."""
    u = _actor()
    return u["id"] if u else "human"


def _actor_role() -> Optional[str]:
    """Роль, с которой человек действует СЕЙЧАС, — из сессии (роль в активной
    команде), а не с записи пользователя: та про домашнюю организацию."""
    sess = CURRENT_SESSION.get() or {}
    return sess.get("role") or None


def _user_label(uid) -> Optional[str]:
    """Имя для показа по идентификатору. "human" и None — прежние отметки
    без автора, имени у них нет. Удалённый пользователь — «#N»: след остаётся
    следом и когда учётной записи больше нет."""
    if not isinstance(uid, int):
        return None
    for u in _users():
        if u["id"] == uid:
            return u.get("name") or u.get("login") or ("#%d" % uid)
    return "#%d" % uid


def _signed(action: str) -> Optional[dict]:
    """След ответственного на записи глоссария: кто, в какой роли, когда и что
    сделал (approve | add | edit | demote | import). Без сессии (миграции,
    тесты) следа нет — выдумывать автора нельзя."""
    u = _actor()
    if not u:
        return None
    return {"user": u["id"], "name": u.get("name") or u["login"],
            "role": _actor_role() or u.get("role", "translator"),
            "at": datetime.now().strftime("%Y-%m-%d %H:%M"), "action": action}


def _signed_field(action: str) -> dict:
    """Для `**` в литерале записи: `{"signedBy": …}` либо ничего."""
    sig = _signed(action)
    return {"signedBy": sig} if sig else {}


def _withdraw_confirmation(seg: dict, how: str) -> None:
    """Снять заверение человека с сегмента — при снятии руками (`revert`)
    и при ПРАВКЕ заверенного текста (`edit`). Подпись относилась к тексту
    и решению, которых больше нет: оставь `confirmedBy` — и `_confirmed_by_human`
    считал бы снятый сегмент заверенным, а на чужой формулировке стояла бы
    подпись человека, который её не видел. Кто снял — на сегменте
    (`unconfirmed`), не только в кольцевом журнале."""
    if seg.get("confirmedBy") is not None:
        seg["unconfirmed"] = {"by": seg.get("confirmedBy"), "at": seg.get("confirmedAt"),
                              "role": seg.get("confirmedRole"), "how": how,
                              "withdrawnBy": _actor_id(),
                              "withdrawnAt": datetime.now().strftime("%Y-%m-%d %H:%M")}
    for k in ("confirmedBy", "confirmedAt", "confirmedRole"):
        seg.pop(k, None)
    seg["status"] = "translated"


def _confirmed_by_human(seg: dict) -> bool:
    by = (seg or {}).get("confirmedBy")
    return by == "human" or isinstance(by, int)


def _audit(action: str, **fields) -> None:
    try:
        u = _actor()
        sess = CURRENT_SESSION.get() or {}
        rec = {"at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
               "tenant": sess.get("tenant") or _current_tenant(),
               "user": u["id"] if u else None, "login": u["login"] if u else None,
               "role": sess.get("role"),
               "action": action}
        rec.update({k: v for k, v in fields.items() if v is not None})
        log = STATE.setdefault("audit", [])
        log.append(rec)
        if len(log) > AUDIT_MAX:
            del log[:-AUDIT_MAX]
    except Exception as e:      # журнал не вправе ронять действие
        print(f"[backend] audit failed: {e}", file=sys.stderr)


def _tenants() -> list:
    return STATE.setdefault("tenants", [])


def _user_by_login(login: str) -> Optional[dict]:
    """Пользователь по логину ИЛИ по почте: человек помнит что-то одно,
    и заставлять его гадать, чем он регистрировался, — плохая дверь."""
    key = (login or "").strip().lower()
    if not key:
        return None
    for u in _users():
        if u.get("login", "").lower() == key or (u.get("email") or "").lower() == key:
            return u
    return None


def _user_by_email(email: str) -> Optional[dict]:
    key = (email or "").strip().lower()
    for u in _users():
        if (u.get("email") or "").lower() == key:
            return u
    return None


def _check_email(email: str) -> str:
    e = (email or "").strip().lower()
    if not _EMAIL_RE.match(e) or len(e) > 190:
        raise HTTPException(400, "Неверный адрес почты")
    return e


def _issue_code(user: dict, kind: str) -> str:
    """Код подтверждения: шесть цифр, живёт CODE_TTL, хранится ХЭШОМ
    (в базу утечь может, а код — нет) с потолком попыток."""
    code = "%06d" % secrets.randbelow(1000000)
    h, salt = _hash_password(code)
    user["authCode"] = {"hash": h, "salt": salt, "kind": kind,
                        "exp": time.time() + CODE_TTL, "tries": 0}
    return code


def _check_code(user: dict, code: str, kind: str) -> None:
    rec = user.get("authCode") or {}
    if rec.get("kind") != kind or rec.get("exp", 0) < time.time():
        raise HTTPException(400, "Код устарел — запросите новый")
    if rec.get("tries", 0) >= CODE_MAX_TRIES:
        raise HTTPException(429, "Слишком много попыток — запросите новый код")
    h, _ = _hash_password((code or "").strip(), rec.get("salt") or "")
    if not hmac.compare_digest(h, rec.get("hash") or ""):
        rec["tries"] = rec.get("tries", 0) + 1
        save_state(STATE)
        raise HTTPException(400, "Неверный код")
    user.pop("authCode", None)


def _signup_blocked(ip: str) -> bool:
    now = time.time()
    hits = [t for t in _SIGNUP_FAILS.get(ip, []) if now - t < 3600]
    _SIGNUP_FAILS[ip] = hits
    return len(hits) >= SIGNUP_MAX_PER_HOUR


def _new_tenant_id(base: str) -> str:
    slug = re.sub(r"[^a-z0-9-]+", "-", (base or "").lower()).strip("-")[:24] or "org"
    if not re.match(r"^[a-z]", slug):
        slug = "org-" + slug
    taken = {t.get("id") for t in _tenants()}
    tid, n = slug, 1
    while tid in taken:
        n += 1
        tid = f"{slug}-{n}"
    return tid


def _ensure_users() -> None:
    """Пустая база пользователей → организация по умолчанию и её владелец
    с паролем APP_PASSWORD. Зовётся лениво (на входе), а не при импорте:
    тесты импортируют модуль с боевой копией state.json, и писать в неё
    пользователя ради импорта нельзя."""
    if _users():
        return
    if not any(t.get("id") == DEFAULT_TENANT for t in _tenants()):
        _tenants().append({"id": DEFAULT_TENANT, "name": "Организация",
                           "created": datetime.now().strftime("%Y-%m-%d"), "active": True})
    h, salt = _hash_password(_RAW_PASSWORD)
    _users().append({"id": 1, "tenant": DEFAULT_TENANT, "login": BOOTSTRAP_LOGIN,
                     "hash": h, "salt": salt, "role": "owner", "super": True,
                     "name": "Администратор", "active": True, "uiLang": DEFAULT_UI_LANG,
                     "created": datetime.now().strftime("%Y-%m-%d")})
    print(f"[backend] пользователей не было — заведён владелец «{BOOTSTRAP_LOGIN}» "
          f"организации «{DEFAULT_TENANT}» с паролем из APP_PASSWORD", file=sys.stderr)
    save_state(STATE)

SESSION_TTL = max(300, int(os.environ.get("SESSION_TTL_HOURS", "12")) * 3600)
LOGIN_MAX_FAILS = 10           # неудачных попыток с одного IP
LOGIN_FAIL_WINDOW = 15 * 60    # за это окно; потом счётчик обнуляется

_SESSIONS: dict = {}           # token -> {"exp", "user", "tenant", "role", "super"}
_LOGIN_FAILS: dict = {}        # ip -> (fail_count, window_started_at)
_AUTH_LOCK = threading.Lock()

# Единственный список исключений. Всё прочее под /api/ требует токен.
# Двери самостоятельной регистрации публичны по своей природе: их зовёт
# человек, у которого ещё нет ни токена, ни учётной записи. Каждая из них
# ограничена по IP и по числу попыток, а «забыли пароль» отвечает одинаково
# при любом адресе — иначе она превратилась бы в проверку «есть ли такой
# клиент». Всё остальное под /api/ по-прежнему требует токен.
PUBLIC_API_PATHS = {"/api/auth/login", "/api/auth/logout", "/api/health",
                    "/api/auth/signup-info", "/api/auth/register",
                    "/api/auth/verify", "/api/auth/resend",
                    "/api/auth/forgot", "/api/auth/reset"}


def _new_session(user: dict) -> str:
    token = secrets.token_urlsafe(32)
    now = time.time()
    with _AUTH_LOCK:
        for dead in [t for t, s in _SESSIONS.items() if s["exp"] <= now]:
            _SESSIONS.pop(dead, None)
        _SESSIONS[token] = {"exp": now + SESSION_TTL, "user": user["id"],
                            "tenant": user.get("tenant", DEFAULT_TENANT),
                            "role": user.get("role", "translator"),
                            # Язык объяснений модели: на нём она пишет `why`
                            # и `comment`. Держим в сессии, чтобы не ходить
                            # в базу пользователей на каждый вызов.
                            "uiLang": user.get("uiLang") or DEFAULT_UI_LANG,
                            "super": bool(user.get("super"))}
    return token


def _session_of(token: Optional[str]) -> Optional[dict]:
    if not token:
        return None
    with _AUTH_LOCK:
        s = _SESSIONS.get(token)
        if s is None:
            return None
        if s["exp"] <= time.time():
            _SESSIONS.pop(token, None)
            return None
    return s


def _session_valid(token: Optional[str]) -> bool:
    return _session_of(token) is not None


# Что вправе делать только владелец. Таблица здесь, а не декоратор на каждом
# обработчике: проверка в одной точке (как и сам токен), и забыть её на новом
# эндпоинте нельзя — забыть можно только строку в таблице, и это видно.
_OWNER_ONLY = [
    ("DELETE", re.compile(r"/api/projects/\d+$")),
    ("DELETE", re.compile(r"/api/glossary$")),
    ("DELETE", re.compile(r"/api/tm$")),
    # Понижения записи и выноса здесь НЕТ по решению владельца сервиса: у обоих
    # есть откат (`prevTier`, копия в data/backups), а кто нажал — в подписи
    # записи (`signedBy`) и в журнале.
    ("POST",   re.compile(r"/api/pricing(/.*)?$")),
    ("POST",   re.compile(r"/api/style$")),
    ("POST",   re.compile(r"/api/quotes/\d+$")),
    ("DELETE", re.compile(r"/api/quotes/\d+$")),
    ("*",      re.compile(r"/api/admin/")),
    # Команд здесь НЕТ намеренно: таблица берёт роль из сессии — роль
    # в АКТИВНОЙ команде, — а путь /api/teams/{tid}/… называет другую.
    # Право владельца команды проверяет `_team_owner_or_403` в обработчике.
]


def _owner_only(method: str, path: str) -> bool:
    return any((m == "*" or m == method) and rx.match(path) for m, rx in _OWNER_ONLY)


# Права редактора и переводчика сегодня РАВНЫ, и таблица пуста намеренно:
# владелец сервиса решил, что заверяет каждый, а роль — подпись
# ответственности в следе («подтвердил: Ева · переводчик»), а не замок.
# Первая версия закрывала переводчику заверение сегментов и решения по
# терминам; отменено по прямой просьбе. Таблица оставлена как место, куда
# лягут права, когда роли разойдутся, — проверка в `require_token` уже стоит.
_EDITOR_ONLY: list = []
_ROLE_DENIED = {
    "owner": "Это действие доступно только владельцу организации",
    "editor": "Это действие доступно редактору или владельцу: заверение и решения "
              "по терминам оставляют подпись ответственного",
}


def _editor_only(method: str, path: str) -> bool:
    return any((m == "*" or m == method) and rx.match(path) for m, rx in _EDITOR_ONLY)


def _drop_session(token: Optional[str]) -> None:
    if token:
        with _AUTH_LOCK:
            _SESSIONS.pop(token, None)


def _client_ip(request: Request) -> str:
    # За nginx реальный адрес приходит в X-Forwarded-For (см. deploy/nginx.conf).
    fwd = request.headers.get("x-forwarded-for", "")
    if fwd:
        return fwd.split(",")[0].strip()
    return request.client.host if request.client else "?"


def _login_blocked(ip: str) -> bool:
    with _AUTH_LOCK:
        rec = _LOGIN_FAILS.get(ip)
        if not rec:
            return False
        count, started = rec
        if time.time() - started > LOGIN_FAIL_WINDOW:
            _LOGIN_FAILS.pop(ip, None)
            return False
        return count >= LOGIN_MAX_FAILS


def _note_login_fail(ip: str) -> None:
    now = time.time()
    with _AUTH_LOCK:
        count, started = _LOGIN_FAILS.get(ip, (0, now))
        if now - started > LOGIN_FAIL_WINDOW:
            count, started = 0, now
        _LOGIN_FAILS[ip] = (count + 1, started)


def _token_from_request(request: Request) -> Optional[str]:
    auth = request.headers.get("authorization", "")
    if auth.lower().startswith("bearer "):
        return auth[7:].strip() or None
    header = request.headers.get("x-auth-token", "").strip()
    if header:
        return header
    # Скачивание файла идёт обычной ссылкой <a href>, заголовок туда не подставить,
    # поэтому только для этого одного пути токен допускается в query-строке.
    if request.url.path.endswith("/export/download"):
        return request.query_params.get("token")
    return None


# ─────────────────────────────────────────────────────────────────────
# App setup
# ─────────────────────────────────────────────────────────────────────
app = FastAPI(title=APP_BRAND + " API", version="5.6.0")


@app.middleware("http")
async def require_token(request: Request, call_next):
    """Одна точка проверки: новый /api/* эндпоинт защищён автоматически."""
    path = request.url.path
    if (request.method != "OPTIONS"              # preflight обслуживает CORSMiddleware
            and path.startswith("/api/")
            and path not in PUBLIC_API_PATHS):
        sess = _session_of(_token_from_request(request))
        if sess is None:
            return JSONResponse({"ok": False, "error": "Требуется вход в систему"}, status_code=401)
        if _owner_only(request.method, path) and sess.get("role") != "owner":
            return JSONResponse({"ok": False, "error": _ROLE_DENIED["owner"]}, status_code=403)
        if _editor_only(request.method, path) and not _role_at_least(sess.get("role"), "editor"):
            return JSONResponse({"ok": False, "error": _ROLE_DENIED["editor"]}, status_code=403)
        if _is_paid(request.method, path):
            st = _spend_status(sess.get("tenant"))
            if st["over"]:
                return JSONResponse({"ok": False, "error":
                    "Месячный лимит расхода организации исчерпан: $%.2f из $%.2f. Бесплатные команды "
                    "(правка начертания, откаты, пересчёт, экспорт) работают; лимит сбрасывается "
                    "1-го числа." % (st["spentUsd"], float(st["limitUsd"])), "spend": st},
                    status_code=402)
        request.state.session = sess
        tok = CURRENT_SESSION.set(sess)
        try:
            if STORE.kind == "pg":
                await run_in_threadpool(_sync_shared)
            return await call_next(request)
        finally:
            CURRENT_SESSION.reset(tok)
    return await call_next(request)


# CORS добавляется ПОСЛЕ require_token: последняя добавленная мидлварь —
# внешняя, поэтому preflight и ответы 401 тоже получают CORS-заголовки.
# Список origin'ов вместо прежнего "*": со звёздочкой и allow_credentials
# любой сторонний сайт мог дёргать API из браузера пользователя.
_DEFAULT_ORIGINS = [
    "https://simpletranslate.me",
    "https://www.simpletranslate.me",
    # Прежний адрес: на него ведут закладки и ссылки из старых писем.
    "https://trasnlateuz.duckdns.org",
    "http://localhost:8000",
    "http://127.0.0.1:8000",
]
ALLOWED_ORIGINS = [o.strip() for o in os.environ.get(
    "ALLOWED_ORIGINS", ",".join(_DEFAULT_ORIGINS)).split(",") if o.strip()]

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["GET", "POST", "DELETE", "OPTIONS"],
    allow_headers=["Content-Type", "Authorization", "X-Auth-Token"],
)

FRONTEND_DIR = ROOT / "frontend"
DATA_DIR = ROOT / "backend" / "data"
DATA_DIR.mkdir(exist_ok=True)
STATE_FILE = DATA_DIR / "state.json"

# Хранилище: файл или PostgreSQL (DATABASE_URL). STATE остаётся моделью
# в памяти, меняется только то, куда пишет save_state и откуда читает
# load_state — см. backend/store.py. Отказ соединения при заданном URL —
# громкий, а не тихий откат на файл.
try:
    import store as _store_mod
except ImportError:                       # запуск как backend.main:app
    from backend import store as _store_mod
try:
    import textcount
except ImportError:                       # запуск как backend.main:app
    from backend import textcount
mailer_mod = _safe_import("mailer")
legal_mod = _safe_import("legal")
STORE = _store_mod.open_store(os.environ.get("DATABASE_URL"), STATE_FILE)

# Прогоны отдельным процессом (systemd-юнит medcat-worker, backend/worker.py):
# API только ставит задачу в таблицу jobs, воркер забирает её claim_job
# (SKIP LOCKED). Включается ТОЛЬКО с базой: файлу второй процесс запрещён
# (инвариант 1). Роль процесса — MEDCAT_ROLE=worker у самого воркера.
EXTERNAL_WORKER = (STORE.kind == "pg" and os.environ.get(
    "MEDCAT_EXTERNAL_WORKER", "").strip().lower() in ("1", "true", "yes", "on"))
IS_WORKER = os.environ.get("MEDCAT_ROLE", "").strip() == "worker"

def checks_enabled() -> bool:
    """Детерминированные проверки включены. Переменная называется
    CHECKS_ENABLED; прежняя MEDICAL_TRANSLATION_QA_ENABLED читается ещё
    релиз — выкат не должен молча менять поведение боевого сервера."""
    raw = (os.environ.get("CHECKS_ENABLED")
           or os.environ.get("MEDICAL_TRANSLATION_QA_ENABLED") or "1")
    if checks_mod and hasattr(checks_mod, "enabled_from_env"):
        return checks_mod.enabled_from_env(raw)
    return raw.strip().lower() not in {"0", "false", "no", "off"}

# ─────────────────────────────────────────────────────────────────────
# Стартовые данные
# ─────────────────────────────────────────────────────────────────────
# Пустой старт — норма: словарь клиента приходит импортом, а не из кода,
# и демонстрационный «Эпикриз — кардиология» с выдуманными сотрудниками
# не должен подниматься у клиента, потерявшего файл состояния. Демо лежит
# файлом (`backend/demo_seed.json`) и включается ТОЛЬКО `DEMO_SEED=1`.
DEMO_SEED = os.environ.get("DEMO_SEED", "").strip().lower() in ("1", "true", "yes", "on")
_EMPTY_SEED = {"projects": [], "glossary": [], "tm": [], "exportHistory": [], "team": []}


def _demo_seed() -> dict:
    if not DEMO_SEED:
        return dict(_EMPTY_SEED)
    try:
        with open(ROOT / "backend" / "demo_seed.json", encoding="utf-8") as f:
            return {**_EMPTY_SEED, **json.load(f)}
    except Exception as e:
        print(f"[backend] WARN: demo_seed.json не прочитан: {e}", file=sys.stderr)
        return dict(_EMPTY_SEED)


def _load_glossary_from_tsv() -> list:
    """Стартовый словарь первого клиента из TSV. Читается ЛЕНИВО и только
    миграцией уровней доверия: у новой установки глоссарий пустой."""
    import csv
    _CAT_MAP = {
        "diagnosis": "Disease", "anatomy": "Anatomy", "symptom": "Symptom",
        "medication": "Dosage", "procedure": "Procedure", "other_medical": "Disease",
        "test": "Lab", "": "Disease",
    }
    tsv = ROOT / "med_translation" / "assets" / "glossary" / "approved_glossary_FINAL.tsv"
    if not tsv.exists():
        tsv = ROOT / "med_translation" / "data" / "approved_glossary_FINAL.tsv"
    if not tsv.exists():
        return []
    terms, seen = [], set()
    try:
        with open(tsv, encoding="utf-8-sig") as f:
            reader = csv.DictReader(f, delimiter="\t")
            for row in reader:
                ru = (row.get("Russian") or "").strip().strip('"')
                en = (row.get("English") or "").strip().strip('"')
                if not ru or not en or ru in seen:
                    continue
                if len(ru) < 3 or len(en) < 3:
                    continue
                # skip strings starting with special chars or digits
                if not ru[0].isalpha() or not en[0].isalpha():
                    continue
                seen.add(ru)
                cat_raw = (row.get("Category") or "").strip()
                origin = (row.get("Sources") or "").strip()
                terms.append({
                    "src": ru, "tgt": en,
                    "cat": _CAT_MAP.get(cat_raw, "Disease"),
                    "freq": 1, "conf": "high", "note": "",
                    # tier решает, приказ это для модели или подсказка (см. _tier_from_origin)
                    "tier": _tier_from_origin(origin), "origin": origin[:60],
                })
    except Exception as e:
        print(f"[backend] WARN: could not load glossary TSV: {e}", file=sys.stderr)
    return terms


BACKUP_DIR = DATA_DIR / "backups"
# Сериализует записи состояния (эндпоинты идут в threadpool). ИМЕННО RLock:
# save_state под этим локом при DocConflict зовёт _apply_doc, который берёт
# его же, — с обычным Lock это самоблокировка НАВСЕГДА. На боевом сервере
# так замер старт API: рестарт вместе с воркером дал конфликт документа
# на модульном save_state, и порт 8000 не поднялся вовсе.
_SAVE_LOCK = threading.RLock()
_BACKUP_KEEP = 48                      # почасовые бэкапы, ~2 суток


# Отказ ремонта, к КАЧЕСТВУ правки отношения не имеющий: перепроверка не
# ответила. Константой, потому что об этой строке говорят двое — writer
# в `_run_segment_repair` и миграция, разжимающая записи прежнего кода.
# В самих данных лежит этот же литерал: переформулировать его задним числом
# нельзя, поэтому миграция сверяется с ним, а не с кодом причины.
REPAIR_RECHECK_FAILED = "перепроверка терминов не выполнилась"

def _apply_migrations(state: dict) -> dict:
    # Migrate: fix TM quality field (verified bool → quality string)
    for t in state.get("tm", []):
        if "quality" not in t:
            t["quality"] = "verified" if t.get("verified") else "draft"
    # Migrate: уровни доверия появились позже самого глоссария. Проставляем их
    # по эталонному TSV; чего в массовом импорте нет — добавлено руками, значит
    # проверено. Иначе весь автоимпорт так и остался бы приказом для модели.
    if any("tier" not in t for t in state.get("glossary", [])):
        # TSV читается лениво и только здесь: при обычном старте он не нужен.
        tiers = {t["src"]: t.get("tier", GLOSSARY_TIER_HARD) for t in _load_glossary_from_tsv()}
        for t in state.get("glossary", []):
            if "tier" not in t:
                t["tier"] = tiers.get(t.get("src"), GLOSSARY_TIER_HARD)
    # Migrate: Medical QA раньше писала свою оценку в seg["risk"], где живёт
    # длина сегмента — а по ней выбирается движок перевода. Возвращаем длину
    # тем же расчётом, что при импорте: миграция идемпотентна и полей не
    # добавляет. Трогаем только сегменты, где проверка действительно была
    # (есть risk_color) или где остался её след — «critical» больше никто
    # не ставит.
    for _p in state.get("projects", []):
        for _s in _p.get("segments", []):
            if _s.get("risk_color") or _s.get("risk") == "critical":
                _w = len((_s.get("source") or "").split())
                _s["risk"] = "high" if _w > 30 else "medium" if _w > 8 else "low"
    state.setdefault("termQueue", [])
    # Migrate: до появления origTgt одобрение конфликта затирало пустой перевод
    # решением человека, и дедупликация теряла кандидата — термин возвращался
    # в очередь неодобренным. Конфликт всегда рождается с пустым tgt, поэтому
    # прежнюю пару восстанавливаем точно. Правленые пары других видов не
    # угадываем: там текущий tgt и есть исходный, если человек его не менял.
    for c in state["termQueue"]:
        if (c.get("kind") == "conflict" and c.get("status") == "approved"
                and c.get("tgt") and "origTgt" not in c):
            c["origTgt"] = ""
    # Migrate: прежний код засчитывал заход ремонта даже тогда, когда правку
    # отвергла не оценка, а ОБОРВАННЫЙ вызов перепроверки. `repair.source_hash`
    # при этом ставился, `_repair_tried` возвращал True — и сегмент оказывался
    # закрыт от ремонта навсегда из-за чужой сетевой ошибки. На боевом проекте
    # так заклеймлены 9 сегментов (#645: «accidental» → «casual», балл не падал
    # вовсе). Снимаем клеймо: причина отказа сохраняется, теряется только запись
    # о попытке — то есть ровно то, чего прежний код не должен был писать.
    #
    # Сверяемся с ПОЛНЫМ совпадением строки: «;» в причине означает, что рядом
    # стояла претензия по существу, и тогда вердикт вынесен — повторный заход
    # даст то же самое за те же деньги. Миграция идемпотентна: снимать второй
    # раз нечего. Хешей тут нет намеренно — `_apply_migrations` зовётся при
    # импорте раньше, чем определён `_text_hash`.
    for _p in state.get("projects", []):
        for _s in _p.get("segments", []):
            _r = _s.get("repair") or {}
            # Причина целиком состоит из сбоя перепроверки и, возможно,
            # упавшего балла — то есть по существу ничего сказано не было.
            # Проверять только полное совпадение мало: заход почти всегда
            # смешанный, и такие записи остались бы заклеймлёнными навсегда.
            _parts = [x.strip() for x in (_r.get("reason") or "").split(";") if x.strip()]
            # «жёсткая находка» в строке про балл — отказ ПО СУЩЕСТВУ, и такую
            # запись разжимать нельзя: её заклеймил нынешний код и намеренно.
            _only_infra = bool(_parts) and REPAIR_RECHECK_FAILED in _parts and all(
                x == REPAIR_RECHECK_FAILED
                or (x.startswith("балл back-check упал") and "жёсткая находка" not in x)
                for x in _parts)
            # Правило отмены ИЗМЕНИЛОСЬ: серьёзные замечания сверяются теперь
            # поимённо, а не счётом. Вердикты, вынесенные прежним правилом,
            # описывают правила, которых больше нет, — и держат сегмент
            # закрытым навсегда. Тот же закон, что у `BACKCHECK_VERSION`:
            # поменялась формула — прежняя оценка перестала быть свежей.
            # Денег это само по себе не тратит: сегмент лишь становится виден
            # разбору прогона, а тот называет число и цену до запуска.
            # Тот же случай у вето по БАЛЛУ: «стало чище» теперь включает
            # снятое заказанное замечание (`ordered_fixed`), поэтому правка
            # с неизменившимся счётчиком («1 → 1» — сняли одно, пришло другое)
            # прежде не проходила ни в одну ветку и откатывалась баллом.
            # Берём ровно этот признак: отказ ТОЛЬКО по баллу и счётчик
            # серьёзных замечаний не изменился. Где счётчик вырос или неизвестен,
            # прежний вердикт правилам не противоречит и остаётся в силе.
            _b, _a = _r.get("before") or {}, _r.get("after") or {}
            _score_only = ((_r.get("reason") or "").startswith("балл back-check упал")
                           and ";" not in (_r.get("reason") or "")
                           and "жёсткая находка" not in (_r.get("reason") or ""))
            _same_terms = (_b.get("terms") is not None and _a.get("terms") is not None
                           and _b["terms"] == _a["terms"])
            if (_r.get("applied") is False and "source_hash" in _r
                    and not _r.get("hardAfter")
                    and (("замечаний по терминам стало больше" in (_r.get("reason") or ""))
                         or (_score_only and _same_terms))):
                _r["attemptHash"] = _r.pop("source_hash")
                _r["retryable"] = True
                _r["retryReason"] = "rules"
                continue
            if (_r.get("applied") is False and _only_infra
                    and not _r.get("hardAfter") and "source_hash" in _r):
                # Снятый клеймящий хеш становится информационным: он и есть
                # хеш текста, на котором заход не состоялся. Пересчитывать
                # нечего — `_text_hash` здесь ещё не определён.
                _r["attemptHash"] = _r.pop("source_hash")
                _r["retryable"] = True
    return state


def load_state() -> dict:
    """Из базы, если она включена и не пуста; иначе из файла (и это же —
    первичный перенос: пустая база + state.json на диске = первое
    сохранение уложит всё в базу документами)."""
    if STORE.kind == "pg":
        st = STORE.load()
        if st is not None:
            return _apply_migrations(st)
        print("[backend] база пуста — состояние поднимается из state.json и уйдёт "
              "в базу при первом сохранении", file=sys.stderr)
    return _load_state_file()


def _load_state_file() -> dict:
    """Загрузка из файла. Если state.json повреждён — НЕ теряем данные молча:
    битый файл сохраняется как state.corrupt-*, затем пробуем свежайший бэкап."""
    candidates = [STATE_FILE] + sorted(BACKUP_DIR.glob("state-*.json"), reverse=True)
    for path in candidates:
        if not path.exists():
            continue
        try:
            state = json.loads(path.read_text(encoding="utf-8"))
            if path != STATE_FILE:
                print(f"[backend] CRITICAL: state.json повреждён — восстановлено из бэкапа {path.name}",
                      file=sys.stderr)
            return _apply_migrations(state)
        except Exception as e:
            print(f"[backend] ERROR: не удалось прочитать {path.name}: {e}", file=sys.stderr)
            if path == STATE_FILE:
                try:
                    corrupt = STATE_FILE.with_name(
                        f"state.corrupt-{datetime.now().strftime('%Y%m%d-%H%M%S')}.json")
                    STATE_FILE.rename(corrupt)
                    print(f"[backend] повреждённый файл сохранён как {corrupt.name} "
                          f"(для ручного восстановления)", file=sys.stderr)
                except Exception:
                    pass
    print("[backend] CRITICAL: рабочее состояние не найдено — "
          + ("старт с демо-данными (DEMO_SEED=1)" if DEMO_SEED else "пустой старт"),
          file=sys.stderr)
    return {**json.loads(json.dumps(_demo_seed())), "termQueue": []}


def _hourly_backup(state: dict):
    """Раз в час откладывает копию состояния в data/backups/ (хранится ~2 суток).
    JSON собирается только когда файл этого часа ещё не записан — при базе
    это единственное место, где состояние сериализуется целиком."""
    try:
        BACKUP_DIR.mkdir(exist_ok=True)
        bak = BACKUP_DIR / f"state-{datetime.now().strftime('%Y%m%d-%H')}.json"
        if not bak.exists():
            bak.write_text(json.dumps(state, ensure_ascii=False), encoding="utf-8")
            for old in sorted(BACKUP_DIR.glob("state-*.json"))[:-_BACKUP_KEEP]:
                old.unlink()
    except Exception as e:
        print(f"[backend] WARN: hourly backup failed: {e}", file=sys.stderr)


def save_state(state: dict):
    """Атомарная запись: tmp-файл + fsync + os.replace под глобальным локом.
    Раньше файл писался напрямую — параллельные запросы могли оставить битый JSON,
    а load_state() молча сбрасывал всё на демо-данные (потеря проектов)."""
    try:
        with _SAVE_LOCK:
            try:
                STORE.save(state)
            except _store_mod.DocConflict as e:
                # Документ переписал другой процесс. Наши правки ЭТОГО
                # документа теряются — громко; остальное сохраняется повтором.
                # Штатно сюда не попадаем: ручные правки проекта закрыты 409
                # на время прогона, пакетные команды такие проекты пропускают.
                print(f"[backend] CRITICAL: конфликт документа {e.key} — "
                      f"перечитан из базы, локальные правки этого документа потеряны",
                      file=sys.stderr)
                _apply_doc(e.key, STORE.load_doc(e.key))
                STORE.save(state)
            _hourly_backup(state)
    except Exception as e:
        print(f"[backend] WARN: could not save state: {e}", file=sys.stderr)


STATE = load_state()

# ── Синхронизация разделяемых коллекций между процессами ─────────────
# Глоссарий и очередь кандидатов лежат в базе по строке на запись с эпохой
# поколения (store.ROW_COLLECTIONS). Пока процесс один, эпоху поднимаем
# только мы сами и перечитывания не случаются; появится воркер отдельным
# процессом — его правки станут видны здесь тем же механизмом, без гонки
# «кто последний сохранил документ». Проверка эпох — один SELECT, но и он
# не бесплатен на каждый запрос, поэтому не чаще раза в SYNC_EVERY секунд.
SYNC_EVERY = float(os.environ.get("SYNC_EVERY", "3"))
_SYNC_LAST = {"t": 0.0}


def _sync_shared(force: bool = False) -> None:
    if STORE.kind != "pg":
        return
    now = time.time()
    if not force and now - _SYNC_LAST["t"] < SYNC_EVERY:
        return
    _SYNC_LAST["t"] = now
    try:
        stale = STORE.stale_collections()
    except Exception as e:
        print(f"[backend] сверка эпох не удалась: {e}", file=sys.stderr)
        return
    for coll in stale:
        try:
            if coll.startswith("doc:"):
                key = coll[4:]
                _apply_doc(key, STORE.load_doc(key))
                print(f"[backend] {key}: перечитан после чужого прогона", file=sys.stderr)
                continue
            items = STORE.load_rows(coll)
        except Exception as e:
            print(f"[backend] коллекция {coll} не перечитана: {e}", file=sys.stderr)
            continue
        with _SAVE_LOCK:
            STATE[coll] = items
        if coll == "glossary":
            _invalidate_gloss_index()
        print(f"[backend] {coll}: подтянуто чужое изменение ({len(items)} записей)",
              file=sys.stderr)


def _apply_doc(key: str, doc) -> None:
    """Подменить документ в STATE тем, что лежит в базе (None — убрать)."""
    with _SAVE_LOCK:
        if key.startswith("projects:"):
            pid = int(key.split(":", 1)[1])
            lst = STATE["projects"]
            for i, pr in enumerate(lst):
                if pr["id"] == pid:
                    if doc is None:
                        lst.pop(i)
                    else:
                        lst[i] = doc
                    return
            if doc is not None:
                lst.insert(0, doc)
        elif doc is None:
            STATE.pop(key, None)
        else:
            STATE[key] = doc


def _active_job_for(pid: int) -> Optional[int]:
    """Идущий или ждущий прогон по проекту — в зеркале и в базе."""
    for j in _JOBS.values():
        if j.get("project") == pid and j.get("status") in ("queued", "running"):
            return j["id"]
    if EXTERNAL_WORKER and not IS_WORKER:
        try:
            return STORE.active_job_for(pid)
        except Exception as e:
            print(f"[backend] проверка прогона по проекту не удалась: {e}", file=sys.stderr)
    return None


def _guard_project_write(pid: int) -> None:
    """Пока прогон идёт в ОТДЕЛЬНОМ процессе, проект принадлежит ему: правка
    из API писала бы тот же документ и затирала работу прогона (или он — её).
    В одном процессе, как раньше, правки и прогон сериализует сам процесс."""
    if not EXTERNAL_WORKER or IS_WORKER:
        return
    jid = _active_job_for(pid)
    if jid:
        raise HTTPException(409, f"По проекту идёт прогон №{jid}: правка подождёт его конца "
                                 f"(или остановите прогон)")


def get_project(pid: int) -> dict:
    """Единственное горло к проекту по номеру. Чужой организации — 404,
    а не 403: 403 подтверждал бы, что проект с таким номером существует."""
    t = _current_tenant()
    for p in STATE["projects"]:
        if p["id"] == pid:
            if _tenant_of(p) != t:
                break
            return p
    raise HTTPException(404, f"Project {pid} not found")


def get_segment(pid: int, sid: int) -> dict:
    project = get_project(pid)
    for s in project["segments"]:
        if s["id"] == sid:
            return s
    raise HTTPException(404, f"Segment {sid} not found in project {pid}")


# ─────────────────────────────────────────────────────────────────────
# API endpoints
# ─────────────────────────────────────────────────────────────────────

class LoginRequest(BaseModel):
    login: str = ""
    password: str

@app.post("/api/auth/login")
def login(req: LoginRequest, request: Request):
    ip = _client_ip(request)
    if _login_blocked(ip):
        raise HTTPException(429, "Слишком много попыток входа. Повторите через 15 минут.")
    _ensure_users()
    # Прежний формат {password} без логина — ещё один релиз: он маппится
    # на первого владельца, иначе выкат запирает нынешнего пользователя.
    login_name = (req.login or "").strip() or BOOTSTRAP_LOGIN
    if not req.login:
        print("[backend] вход без логина — принят как «%s» (прежний формат)" % BOOTSTRAP_LOGIN,
              file=sys.stderr)
    user = _user_by_login(login_name)
    if not user or not _verify_password(user, req.password):
        _note_login_fail(ip)
        raise HTTPException(401, "Неверный логин или пароль")
    # Незавершённая регистрация — не «неверный пароль»: человек ввёл всё
    # правильно, ему нужен код из письма, и сказать об этом надо прямо.
    if user.get("email") and not user.get("emailVerified"):
        raise HTTPException(403, "Почта не подтверждена: введите код из письма "
                                 "или запросите новый")
    token = _new_session(user)
    tok = CURRENT_SESSION.set(_SESSIONS[token])
    try:
        _audit("login", ip=ip)
    finally:
        CURRENT_SESSION.reset(tok)
    return {"ok": True, "token": token, "expiresIn": SESSION_TTL,
            "me": _user_public(user)}


class RegisterRequest(BaseModel):
    email: str
    password: str
    org: str = ""
    name: str = ""
    accept: bool = False        # согласие с офертой и политикой ПДн


class CodeRequest(BaseModel):
    email: str
    code: str = ""
    password: str = ""


def _mail_lang(user: dict) -> str:
    """Язык письма — язык ПОЛУЧАТЕЛЯ, а не отправителя и не сервера.
    Письмо уходит мимо браузера, и подставить перевод на границе показа,
    как везде, здесь некому."""
    return (user or {}).get("uiLang") or DEFAULT_UI_LANG


def _mail_code(user: dict, kind: str, code: str) -> bool:
    lang = _mail_lang(user)
    key = "verify" if kind == "verify" else "reset"
    subject = mail_texts.text(lang, key + ".subject", brand=APP_BRAND, code=code)
    body = mail_texts.text(lang, key + ".body", brand=APP_BRAND, code=code,
                           minutes=CODE_TTL // 60)
    return mailer_mod.send(user["email"], subject, body) if mailer_mod else False


@app.get("/api/auth/signup-info")
def signup_info():
    """Что показывать на экране входа: открыта ли регистрация и уходят ли
    письма. Без второго признака человек упёрся бы в «проверьте почту»
    при ненастроенном SMTP."""
    return {"ok": True, "signup": SIGNUP_ENABLED,
            "mail": bool(mailer_mod and mailer_mod.configured()),
            "brand": APP_BRAND, "trialUsd": SIGNUP_TRIAL_USD,
            "legal": {"version": (legal_mod.VERSION if legal_mod else ""),
                      "terms": "/terms", "privacy": "/privacy",
                      # Реквизиты не заполнены — документ ещё не работает
                      # как договор, и владелец сервиса обязан это видеть.
                      "ready": bool(legal_mod and legal_mod.complete())}}


@app.post("/api/auth/register")
def register(req: RegisterRequest, request: Request):
    """Своя организация одним шагом: почта, пароль, название. Человек
    становится её владельцем; лимит расхода — SIGNUP_TRIAL_USD."""
    if not SIGNUP_ENABLED:
        raise HTTPException(403, "Самостоятельная регистрация выключена — "
                                 "обратитесь к администратору сервиса")
    ip = _client_ip(request)
    if _signup_blocked(ip):
        raise HTTPException(429, "Слишком много регистраций с этого адреса. Повторите через час.")
    _ensure_users()
    email = _check_email(req.email)
    _check_user_fields(None, req.password, None)
    # Согласие — условие заключения договора, а не галочка для красоты:
    # без него регистрации нет. Фиксируем РЕДАКЦИЮ документа, дату и адрес —
    # иначе через год не ответить, с чем именно человек согласился.
    if not req.accept:
        raise HTTPException(400, "Без согласия с офертой и политикой обработки "
                                 "персональных данных регистрация невозможна")
    if _user_by_email(email) or _user_by_login(email):
        raise HTTPException(409, "Такая почта уже зарегистрирована — войдите или "
                                 "восстановите пароль")
    tid = _new_tenant_id(req.org or email.split("@")[0])
    today = datetime.now().strftime("%Y-%m-%d")
    tenant = {"id": tid, "name": (req.org or "").strip() or email.split("@")[0],
              "created": today, "active": True, "signup": True}
    if SIGNUP_TRIAL_USD >= 0:
        # Ноль — тоже решение: платное закрыто до тех пор, пока лимит
        # не поставит администратор. Открытый кран к ключу дороже неудобства.
        tenant["limitUsd"] = SIGNUP_TRIAL_USD
    _tenants().append(tenant)
    h, salt = _hash_password(req.password)
    user = {"id": max((x["id"] for x in _users()), default=0) + 1, "tenant": tid,
            "login": email, "email": email, "emailVerified": False,
            "hash": h, "salt": salt, "role": "owner", "name": (req.name or "").strip() or email.split("@")[0],
            "active": True, "uiLang": DEFAULT_UI_LANG, "created": today,
            "acceptedTerms": {"version": (legal_mod.VERSION if legal_mod else ""),
                              "at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"), "ip": ip}}
    _users().append(user)
    code = _issue_code(user, "verify")
    _SIGNUP_FAILS.setdefault(ip, []).append(time.time())
    tok = CURRENT_SESSION.set({"tenant": tid, "user": user["id"], "role": "owner"})
    try:
        _audit("signup", email=email, tenant_new=tid, ip=ip)
    finally:
        CURRENT_SESSION.reset(tok)
    save_state(STATE)
    sent = _mail_code(user, "verify", code)
    return {"ok": True, "email": email, "tenant": tid, "mailSent": sent,
            "next": "verify",
            "note": ("Код отправлен на почту." if sent else
                     "Почта на сервере не настроена — код записан в журнал сервиса, "
                     "запросите его у администратора.")}


@app.post("/api/auth/verify")
def verify_email(req: CodeRequest, request: Request):
    ip = _client_ip(request)
    if _login_blocked(ip):
        raise HTTPException(429, "Слишком много попыток. Повторите через 15 минут.")
    user = _user_by_email(_check_email(req.email))
    if not user:
        _note_login_fail(ip)
        raise HTTPException(400, "Неверный код")
    if user.get("emailVerified"):
        return {"ok": True, "already": True}
    _check_code(user, req.code, "verify")
    user["emailVerified"] = True
    token = _new_session(user)
    tok = CURRENT_SESSION.set(_SESSIONS[token])
    try:
        _audit("email.verify", email=user["email"])
    finally:
        CURRENT_SESSION.reset(tok)
    save_state(STATE)
    return {"ok": True, "token": token, "expiresIn": SESSION_TTL, "me": _user_public(user)}


@app.post("/api/auth/resend")
def resend_code(req: CodeRequest, request: Request):
    """Повторный код подтверждения. Ответ одинаков при любом адресе:
    иначе эта дверь отвечала бы на вопрос «а есть ли у вас такой клиент»."""
    ip = _client_ip(request)
    if _signup_blocked(ip):
        raise HTTPException(429, "Слишком много запросов. Повторите через час.")
    user = _user_by_email(_check_email(req.email))
    sent = False
    if user and not user.get("emailVerified"):
        code = _issue_code(user, "verify")
        save_state(STATE)
        sent = _mail_code(user, "verify", code)
        _SIGNUP_FAILS.setdefault(ip, []).append(time.time())
    return {"ok": True, "mailSent": sent}


@app.post("/api/auth/forgot")
def forgot_password(req: CodeRequest, request: Request):
    ip = _client_ip(request)
    if _signup_blocked(ip):
        raise HTTPException(429, "Слишком много запросов. Повторите через час.")
    user = _user_by_email(_check_email(req.email))
    sent = False
    if user and user.get("active", True):
        code = _issue_code(user, "reset")
        save_state(STATE)
        sent = _mail_code(user, "reset", code)
        _SIGNUP_FAILS.setdefault(ip, []).append(time.time())
    # Про существование адреса не говорим — ответ один на все случаи.
    return {"ok": True, "mailSent": sent}


@app.post("/api/auth/reset")
def reset_password(req: CodeRequest, request: Request):
    ip = _client_ip(request)
    if _login_blocked(ip):
        raise HTTPException(429, "Слишком много попыток. Повторите через 15 минут.")
    user = _user_by_email(_check_email(req.email))
    if not user:
        _note_login_fail(ip)
        raise HTTPException(400, "Неверный код")
    _check_user_fields(None, req.password, None)
    _check_code(user, req.code, "reset")
    user["hash"], user["salt"] = _hash_password(req.password)
    user["emailVerified"] = True        # код пришёл на эту почту — она рабочая
    with _AUTH_LOCK:                    # прежние сессии закрываем
        for t in [t for t, sess in _SESSIONS.items() if sess["user"] == user["id"]]:
            _SESSIONS.pop(t, None)
    token = _new_session(user)
    tok = CURRENT_SESSION.set(_SESSIONS[token])
    try:
        _audit("password.reset", email=user["email"])
    finally:
        CURRENT_SESSION.reset(tok)
    save_state(STATE)
    return {"ok": True, "token": token, "expiresIn": SESSION_TTL, "me": _user_public(user)}


def _current_user(request: Request) -> dict:
    sess = getattr(request.state, "session", None) or CURRENT_SESSION.get()
    if not sess:
        raise HTTPException(401, "Требуется вход в систему")
    for u in _users():
        if u["id"] == sess["user"]:
            return u
    raise HTTPException(401, "Пользователь удалён — войдите заново")


@app.get("/api/auth/me")
def auth_me(request: Request):
    """Кто я: пользователь, организация, роль. Из этого браузер берёт аватар
    и решает, какие кнопки показывать; право сделать проверяет сервер."""
    u = _current_user(request)
    sess = getattr(request.state, "session", None) or CURRENT_SESSION.get() or {}
    # Организация и роль берутся из СЕССИИ, а не с записи пользователя:
    # человек может состоять в нескольких командах, и работает он в той,
    # которую выбрал. `u["tenant"]`/`u["role"]` — про ДОМАШНЮЮ организацию,
    # и показывать их как текущие значило бы врать про то, где он находится
    # и что ему разрешено.
    tid = sess.get("tenant") or u.get("tenant", DEFAULT_TENANT)
    role = sess.get("role") or u.get("role", "translator")
    tenant = next((t for t in _tenants() if t.get("id") == tid), None)
    # Поля организации перечислены ЯВНО: запись растёт (лимит, прайс, что
    # появится дальше), и отдача её целиком означает, что следующее поле уедет
    # браузеру само, без единого решения.
    tpub = {k: (tenant or {}).get(k) for k in ("id", "name", "active")}
    tpub["id"] = tpub["id"] or tid
    tpub["name"] = tpub["name"] or ""
    return {"ok": True, "me": _user_public(u), "tenant": tpub,
            "can": {"owner": role == "owner", "super": bool(u.get("super")), "role": role},
            "teams": _teams_of(u), "invites": _my_invites(u),
            "spend": _spend_status(tid),
            **({"adminPath": "/" + ADMIN_PATH} if u.get("super") else {})}


class UserCreate(BaseModel):
    login: str
    password: str
    role: str = "translator"
    name: str = ""
    email: str = ""
    tenant: Optional[str] = None        # только для суперпользователя


class UserPatch(BaseModel):
    password: Optional[str] = None
    role: Optional[str] = None
    name: Optional[str] = None
    active: Optional[bool] = None
    uiLang: Optional[str] = None


def _check_user_fields(login: str = None, password: str = None, role: str = None):
    if login is not None and not re.fullmatch(r"[A-Za-z0-9._@-]{2,64}", login):
        raise HTTPException(400, "Логин: 2–64 символа, латиница, цифры, . _ @ -")
    if password is not None and len(password) < 8:
        raise HTTPException(400, "Пароль короче 8 символов")
    if role is not None and role not in ROLES:
        raise HTTPException(400, "Роль: " + " | ".join(ROLES))


def _is_super(request: Request) -> bool:
    return bool((getattr(request.state, "session", None) or CURRENT_SESSION.get() or {}).get("super"))


@app.get("/api/admin/users")
def admin_users(request: Request, all: bool = False):
    """Владелец — своих; суперпользователь с `all=1` — всех, с организацией."""
    me = _current_user(request)
    if all and not _is_super(request):
        raise HTTPException(403, "Все пользователи — только суперпользователю")
    return {"ok": True, "users": [_user_public(u) for u in _users()
                                  if all or u.get("tenant") == me.get("tenant")]}


@app.post("/api/admin/users")
def admin_user_create(req: UserCreate, request: Request):
    _audit("user.create", login=req.login, role=req.role)
    me = _current_user(request)
    _check_user_fields(req.login, req.password, req.role)
    if _user_by_login(req.login):
        raise HTTPException(409, "Такой логин уже есть")
    tenant = me.get("tenant")
    if req.tenant and req.tenant != tenant:
        if not _is_super(request):
            raise HTTPException(403, "Пользователей в другой организации заводит только суперпользователь")
        if not _tenant_rec(req.tenant):
            raise HTTPException(404, "Организация не найдена")
        tenant = req.tenant
    h, salt = _hash_password(req.password)
    email = _check_email(req.email) if req.email else ""
    if email and _user_by_email(email):
        raise HTTPException(409, "Такая почта уже занята")
    u = {"id": max((x["id"] for x in _users()), default=0) + 1, "tenant": tenant,
         "login": req.login.strip(), "hash": h, "salt": salt, "role": req.role,
         "email": email, "emailVerified": bool(email),   # завёл владелец — подтверждать нечего
         "name": req.name.strip() or req.login.strip(), "active": True, "uiLang": DEFAULT_UI_LANG,
         "created": datetime.now().strftime("%Y-%m-%d")}
    _users().append(u)
    save_state(STATE)
    return {"ok": True, "user": _user_public(u)}


@app.post("/api/admin/users/{uid}")
def admin_user_update(uid: int, req: UserPatch, request: Request):
    _audit("user.update", target=uid)
    me = _current_user(request)
    u = next((x for x in _users() if x["id"] == uid
              and (x.get("tenant") == me.get("tenant") or _is_super(request))), None)
    if not u:
        raise HTTPException(404, "Пользователь не найден")
    _check_user_fields(None, req.password, req.role)
    if req.role is not None:
        if u["id"] == me["id"] and req.role != "owner":
            raise HTTPException(400, "Нельзя снять роль владельца с самого себя")
        u["role"] = req.role
    if req.password is not None:
        u["hash"], u["salt"] = _hash_password(req.password)
        # Сменили пароль — чужие сессии этого пользователя закрываются.
        with _AUTH_LOCK:
            for t in [t for t, s in _SESSIONS.items() if s["user"] == u["id"]]:
                _SESSIONS.pop(t, None)
    if req.name is not None:
        u["name"] = req.name.strip() or u["login"]
    if req.uiLang is not None:
        if req.uiLang not in UI_LANGS:
            raise HTTPException(400, "Язык интерфейса: " + " | ".join(UI_LANGS))
        u["uiLang"] = req.uiLang
        u["uiLangSet"] = True
    if req.active is not None:
        if u["id"] == me["id"] and not req.active:
            raise HTTPException(400, "Нельзя отключить самого себя")
        u["active"] = bool(req.active)
    save_state(STATE)
    return {"ok": True, "user": _user_public(u)}


# ─────────────────────────────────────────────────────────────────────
# Профиль пользователя и КОМАНДЫ
# ─────────────────────────────────────────────────────────────────────
# Команда — это рабочее пространство, то есть АРЕНДАТОР (`tenants`), а не
# новая сущность внутри организации. Иначе рушится инвариант 11: область
# записи уже трёхмерна (пара языков, тематика, организация), и «команда
# внутри организации» потребовала бы ЧЕТВЁРТОГО измерения у глоссария,
# памяти переводов, очереди кандидатов, расхода и задач — то есть правки
# каждого места, где сегодня стоит `tenant`.
#
# Один человек может состоять в НЕСКОЛЬКИХ командах: `user["memberships"]` —
# список `{tenant, role, since}`. Домашняя организация (`user["tenant"]`)
# читается членством ВСЕГДА и в список не переписывается — тот же закон
# миграции, что у `lang`/`domain`: боевые записи не трогаем.
#
# АКТИВНАЯ команда живёт в СЕССИИ (`sess["tenant"]`, `sess["role"]`) — ровно
# то, что уже читает `_current_tenant()`. Поэтому изоляция работает без
# единой правки в проектах, глоссарии и прогонах: у запроса по-прежнему
# РОВНО ОДНА организация, просто теперь человек выбирает, какая.
#
# Роль тоже переключается вместе с командой: владелец своей команды может
# быть переводчиком в чужой, и оставить ему права владельца значило бы
# отдать чужой глоссарий на вынос.
TEAM_MAX_PER_USER = int(os.environ.get("TEAM_MAX_PER_USER", "5") or 5)
INVITE_TTL_DAYS = int(os.environ.get("INVITE_TTL_DAYS", "30") or 30)


def _invites() -> list:
    return STATE.setdefault("invites", [])


def _memberships(u: dict) -> list:
    """Все команды человека: домашняя организация ПЛЮС принятые приглашения.
    Домашняя всегда первая и удалению не подлежит — иначе человек остался бы
    без единого рабочего пространства."""
    home = u.get("tenant") or DEFAULT_TENANT
    out = [{"tenant": home, "role": u.get("role", "translator"),
            "since": u.get("created"), "home": True}]
    seen = {home}
    for m in (u.get("memberships") or []):
        tid = m.get("tenant")
        if not tid or tid in seen:
            continue
        seen.add(tid)
        out.append({"tenant": tid, "role": m.get("role", "translator"),
                    "since": m.get("since"), "home": False})
    return out


def _member_role(u: dict, tid: str) -> Optional[str]:
    """Роль человека в КОМАНДЕ tid или None, если он в ней не состоит."""
    for m in _memberships(u):
        if m["tenant"] == tid:
            return m["role"]
    return None


def _team_public(u: dict, m: dict) -> dict:
    rec = _tenant_rec(m["tenant"]) or {}
    return {"id": m["tenant"], "name": rec.get("name") or m["tenant"],
            "role": m["role"], "home": m["home"], "since": m.get("since"),
            "active": rec.get("active", True),
            "members": sum(1 for x in _users() if _member_role(x, m["tenant"]))}


def _teams_of(u: dict) -> list:
    return [_team_public(u, m) for m in _memberships(u)]


def _drop_user_sessions(uid: int, tenant: Optional[str] = None) -> None:
    """Закрыть сессии пользователя (все или только в одной команде).
    Исключённый из команды обязан потерять доступ СРАЗУ, а не через
    SESSION_TTL: до истечения токена он работал бы в чужих проектах."""
    with _AUTH_LOCK:
        for t in [t for t, s in _SESSIONS.items()
                  if s.get("user") == uid and (tenant is None or s.get("tenant") == tenant)]:
            _SESSIONS.pop(t, None)


def _team_owner_or_403(u: dict, tid: str) -> dict:
    """Право распоряжаться КОМАНДОЙ проверяется здесь, а не таблицей
    `_OWNER_ONLY` (инвариант 12), и это не забывчивость: таблица берёт роль
    из сессии — роль в АКТИВНОЙ команде, — а путь называет ДРУГУЮ. Владелец
    своей команды получил бы права владельца в чужой."""
    rec = _tenant_rec(tid)
    if not rec:
        raise HTTPException(404, "Команда не найдена")
    if _member_role(u, tid) != "owner":
        raise HTTPException(403, "Это действие доступно только владельцу команды")
    return rec


def _invite_public(inv: dict) -> dict:
    rec = _tenant_rec(inv.get("tenant")) or {}
    return {"id": inv.get("id"), "tenant": inv.get("tenant"),
            "teamName": rec.get("name") or inv.get("tenant"),
            "email": inv.get("email"), "role": inv.get("role", "translator"),
            "by": inv.get("byName") or inv.get("byLogin") or "",
            "at": inv.get("at"), "status": inv.get("status", "pending"),
            "expired": _invite_expired(inv)}


def _invite_expired(inv: dict) -> bool:
    try:
        at = datetime.strptime((inv.get("at") or "")[:10], "%Y-%m-%d")
    except Exception:
        return False
    return (datetime.now() - at).days > INVITE_TTL_DAYS


def _my_invites(u: dict) -> list:
    """Приглашения, ждущие решения ЭТОГО человека. Ключ — почта: приглашают
    по ней, а не по номеру, и до принятия человек с командой не связан ничем."""
    key = (u.get("email") or "").strip().lower()
    if not key:
        return []
    return [_invite_public(i) for i in _invites()
            if (i.get("email") or "").lower() == key
            and i.get("status") == "pending"
            and not _invite_expired(i)
            and _member_role(u, i.get("tenant")) is None]


class ProfilePatch(BaseModel):
    name: Optional[str] = None
    uiLang: Optional[str] = None
    password: Optional[str] = None
    currentPassword: Optional[str] = None


@app.get("/api/profile")
def profile_get(request: Request):
    """Профиль: кто я, в каких командах состою, что ждёт моего решения.
    Доступен КАЖДОМУ вошедшему — в отличие от `/api/admin/users`, куда
    переводчику хода нет, а язык интерфейса менять ему надо."""
    u = _current_user(request)
    sess = getattr(request.state, "session", None) or CURRENT_SESSION.get() or {}
    return {"ok": True, "me": _user_public(u),
            "activeTeam": sess.get("tenant") or u.get("tenant"),
            "activeRole": sess.get("role") or u.get("role"),
            "teams": _teams_of(u), "invites": _my_invites(u),
            "canCreateTeam": len(_memberships(u)) < TEAM_MAX_PER_USER,
            "teamLimit": TEAM_MAX_PER_USER,
            "spend": _spend_status(sess.get("tenant") or u.get("tenant"))}


@app.post("/api/profile")
def profile_update(req: ProfilePatch, request: Request):
    """Свои имя, язык интерфейса и пароль человек меняет сам. Смена пароля
    требует НЫНЕШНЕГО: украденный токен иначе означал бы украденную
    учётную запись навсегда."""
    u = _current_user(request)
    if req.name is not None:
        u["name"] = req.name.strip() or u["login"]
        u["initials"] = _initials(u["name"])
    if req.uiLang is not None:
        lang = (req.uiLang or "").strip().lower()[:5]
        if lang not in UI_LANGS:
            raise HTTPException(400, "Язык интерфейса: " + " | ".join(UI_LANGS))
        u["uiLang"] = lang
        sess = CURRENT_SESSION.get()
        if sess is not None:
            sess["uiLang"] = lang        # иначе модель до перелогина пишет на прежнем
        # След решения ЧЕЛОВЕКА — тот же приём, что у `_human_touched`
        # в глоссарии: без него не отличить выбранный язык от языка,
        # доставшегося записи по умолчанию кода (см. `_migrate_ui_lang`).
        u["uiLangSet"] = True
    if req.password is not None:
        _check_user_fields(None, req.password, None)
        if not _verify_password(u, req.currentPassword or ""):
            raise HTTPException(403, "Нынешний пароль указан неверно")
        u["hash"], u["salt"] = _hash_password(req.password)
        token = _token_from_request(request)
        _drop_user_sessions(u["id"])            # чужие сессии закрываются
        if token:                               # свою оставляем — иначе выкинет себя же
            with _AUTH_LOCK:
                _SESSIONS[token] = {"exp": time.time() + SESSION_TTL, "user": u["id"],
                                    "tenant": (CURRENT_SESSION.get() or {}).get("tenant", u.get("tenant")),
                                    "role": (CURRENT_SESSION.get() or {}).get("role", u.get("role")),
                                    "super": bool(u.get("super"))}
        _audit("password.change")
    save_state(STATE)
    return {"ok": True, "me": _user_public(u)}


class TeamSwitch(BaseModel):
    tenant: str


@app.post("/api/profile/team")
def profile_switch_team(req: TeamSwitch, request: Request):
    """Переключение активной команды. Меняется И роль: владелец своей
    команды в чужой может быть переводчиком, и `_OWNER_ONLY` обязан
    увидеть ту роль, которая действует ЗДЕСЬ."""
    u = _current_user(request)
    role = _member_role(u, req.tenant)
    if role is None:
        raise HTTPException(404, "Вы не состоите в этой команде")
    rec = _tenant_rec(req.tenant) or {}
    if not rec.get("active", True):
        raise HTTPException(403, "Команда отключена администратором сервиса")
    token = _token_from_request(request)
    with _AUTH_LOCK:
        s = _SESSIONS.get(token)
        if s is None:
            raise HTTPException(401, "Требуется вход в систему")
        s["tenant"], s["role"] = req.tenant, role
    sess = _SESSIONS.get(token) or {}
    tok = CURRENT_SESSION.set(sess)
    try:
        _audit("team.switch", team=req.tenant)
    finally:
        CURRENT_SESSION.reset(tok)
    return {"ok": True, "activeTeam": req.tenant, "activeRole": role,
            "can": {"owner": role == "owner", "super": bool(u.get("super")), "role": role},
            "teams": _teams_of(u)}


class InviteDecision(BaseModel):
    action: str                     # accept | decline


@app.post("/api/profile/invites/{iid}")
def profile_invite_decide(iid: str, req: InviteDecision, request: Request):
    """Решение по приглашению принимает САМ приглашённый, и оба исхода
    записываются: молча исчезнувшее приглашение неотличимо от потерянного."""
    u = _current_user(request)
    key = (u.get("email") or "").strip().lower()
    inv = next((i for i in _invites() if i.get("id") == iid
                and (i.get("email") or "").lower() == key), None)
    if not inv:
        raise HTTPException(404, "Приглашение не найдено")
    if inv.get("status") != "pending":
        raise HTTPException(409, "Решение по этому приглашению уже принято")
    if _invite_expired(inv):
        raise HTTPException(410, "Срок приглашения истёк — попросите новое")
    if req.action not in ("accept", "decline"):
        raise HTTPException(400, "Решение: accept | decline")
    inv["status"] = "accepted" if req.action == "accept" else "declined"
    inv["decidedAt"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    if req.action == "accept":
        tid = inv.get("tenant")
        if not _tenant_rec(tid):
            raise HTTPException(404, "Команда удалена")
        if _member_role(u, tid) is None:
            u.setdefault("memberships", []).append(
                {"tenant": tid, "role": inv.get("role", "translator"),
                 "since": datetime.now().strftime("%Y-%m-%d")})
    _audit("invite." + req.action, team=inv.get("tenant"))
    save_state(STATE)
    return {"ok": True, "teams": _teams_of(u), "invites": _my_invites(u)}


class TeamCreate(BaseModel):
    name: str


@app.get("/api/teams")
def teams_list(request: Request):
    u = _current_user(request)
    return {"ok": True, "teams": _teams_of(u), "invites": _my_invites(u),
            "canCreateTeam": len(_memberships(u)) < TEAM_MAX_PER_USER}


@app.post("/api/teams")
def team_create(req: TeamCreate, request: Request):
    """Новая команда — новое рабочее пространство со своим глоссарием,
    памятью переводов и РАСХОДОМ. Лимит ей ставится тот же, что при
    регистрации (SIGNUP_TRIAL_USD, по умолчанию 0): иначе платный ключ
    открывался бы кнопкой «создать команду» столько раз, сколько нужно.
    Потолок числа команд — по той же причине."""
    u = _current_user(request)
    name = (req.name or "").strip()
    if not (2 <= len(name) <= 64):
        raise HTTPException(400, "Название команды: 2–64 символа")
    if len(_memberships(u)) >= TEAM_MAX_PER_USER:
        raise HTTPException(409, "Больше %d команд на одного человека нельзя" % TEAM_MAX_PER_USER)
    tid = _new_tenant_id(name)
    _tenants().append({"id": tid, "name": name, "created": datetime.now().strftime("%Y-%m-%d"),
                       "active": True, "team": True, "limitUsd": SIGNUP_TRIAL_USD,
                       "createdBy": u["id"]})
    u.setdefault("memberships", []).append(
        {"tenant": tid, "role": "owner", "since": datetime.now().strftime("%Y-%m-%d")})
    _audit("team.create", team=tid, name=name)
    save_state(STATE)
    return {"ok": True, "team": _team_public(u, {"tenant": tid, "role": "owner", "home": False,
                                                 "since": datetime.now().strftime("%Y-%m-%d")}),
            "teams": _teams_of(u)}


@app.get("/api/teams/{tid}")
def team_detail(tid: str, request: Request):
    """Состав команды видит её участник; приглашения — только владелец:
    список чужих почт не дело переводчика."""
    u = _current_user(request)
    role = _member_role(u, tid)
    if role is None:
        raise HTTPException(404, "Команда не найдена")
    rec = _tenant_rec(tid) or {}
    members = []
    for x in _users():
        r = _member_role(x, tid)
        if r is None:
            continue
        members.append({"id": x["id"], "login": x["login"], "name": x.get("name") or x["login"],
                        "email": x.get("email") or "", "role": r,
                        "home": (x.get("tenant") or DEFAULT_TENANT) == tid,
                        "active": x.get("active", True),
                        "initials": x.get("initials") or _initials(x.get("name") or x["login"]),
                        "color": x.get("color") or _USER_COLORS[x["id"] % len(_USER_COLORS)]})
    out = {"ok": True, "team": {"id": tid, "name": rec.get("name") or tid,
                                "active": rec.get("active", True), "created": rec.get("created")},
           "myRole": role, "members": members}
    if role == "owner":
        out["invites"] = [_invite_public(i) for i in _invites() if i.get("tenant") == tid]
    return out


class TeamInvite(BaseModel):
    email: str
    role: str = "translator"


@app.post("/api/teams/{tid}/invite")
def team_invite(tid: str, req: TeamInvite, request: Request):
    """Приглашается ЗАРЕГИСТРИРОВАННЫЙ человек по почте, и отказ «такого нет»
    здесь честный: приглашение — не дверь регистрации, а сообщение внутри
    сервиса. Заводить учётную запись за человека нельзя — пароль знает он.

    Про существование адреса эта дверь говорит намеренно, в отличие от
    `/api/auth/forgot`: там спрашивал кто угодно из интернета, здесь —
    вошедший владелец команды, и без ответа он не поймёт, почему приглашение
    не дошло."""
    u = _current_user(request)
    _team_owner_or_403(u, tid)
    email = _check_email(req.email)
    if req.role not in ROLES:
        raise HTTPException(400, "Роль: " + " | ".join(ROLES))
    target = _user_by_email(email)
    if not target:
        raise HTTPException(404, "Такой почты в сервисе нет — человек должен "
                                 "сначала зарегистрироваться")
    if not target.get("active", True):
        raise HTTPException(409, "Учётная запись отключена")
    if _member_role(target, tid) is not None:
        raise HTTPException(409, "Этот человек уже в команде")
    if any(i.get("tenant") == tid and (i.get("email") or "").lower() == email
           and i.get("status") == "pending" and not _invite_expired(i) for i in _invites()):
        raise HTTPException(409, "Приглашение уже отправлено и ждёт решения")
    inv = {"id": secrets.token_urlsafe(9), "tenant": tid, "email": email,
           "role": req.role, "byUser": u["id"], "byName": u.get("name") or u["login"],
           "byLogin": u["login"], "at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
           "status": "pending"}
    _invites().append(inv)
    _audit("invite.send", team=tid, email=email, role=req.role)
    save_state(STATE)
    sent = _mail_invite(target, _tenant_rec(tid) or {}, u)
    return {"ok": True, "invite": _invite_public(inv), "mailSent": sent}


def _mail_invite(target: dict, team: dict, by: dict) -> bool:
    """Письмо — уведомление, а НЕ дверь: решение принимается внутри сервиса,
    в профиле. Ссылки-приглашения по почте здесь нет намеренно — она была бы
    вторым входом мимо пароля."""
    if not mailer_mod:
        return False
    name = team.get("name") or team.get("id")
    who = by.get("name") or by.get("login")
    lang = _mail_lang(target)
    subject = mail_texts.text(lang, "invite.subject", brand=APP_BRAND, team=name)
    body = mail_texts.text(lang, "invite.body", brand=APP_BRAND, team=name, who=who)
    try:
        return mailer_mod.send(target["email"], subject, body)
    except Exception as e:
        print(f"[backend] приглашение не отправлено: {e}", file=sys.stderr)
        return False


@app.post("/api/teams/{tid}/invites/{iid}/revoke")
def team_invite_revoke(tid: str, iid: str, request: Request):
    u = _current_user(request)
    _team_owner_or_403(u, tid)
    inv = next((i for i in _invites() if i.get("id") == iid and i.get("tenant") == tid), None)
    if not inv:
        raise HTTPException(404, "Приглашение не найдено")
    if inv.get("status") != "pending":
        raise HTTPException(409, "Решение по этому приглашению уже принято")
    inv["status"] = "revoked"
    inv["decidedAt"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    _audit("invite.revoke", team=tid, email=inv.get("email"))
    save_state(STATE)
    return {"ok": True}


class MemberPatch(BaseModel):
    role: Optional[str] = None
    remove: bool = False


@app.post("/api/teams/{tid}/members/{uid}")
def team_member_update(tid: str, uid: int, req: MemberPatch, request: Request):
    """Роль участника и исключение из команды. Три запрета, и все три —
    про необратимое: нельзя снять последнего владельца (команда осталась бы
    без хозяина), нельзя исключить человека из его ДОМАШНЕЙ организации
    (он остался бы без рабочего пространства вовсе) и нельзя исключить
    себя — выход из команды это отдельное решение, не правка состава."""
    u = _current_user(request)
    _team_owner_or_403(u, tid)
    target = next((x for x in _users() if x["id"] == uid), None)
    if not target or _member_role(target, tid) is None:
        raise HTTPException(404, "Участник не найден")
    if (target.get("tenant") or DEFAULT_TENANT) == tid:
        raise HTTPException(400, "Это домашняя организация человека — "
                                 "её состав правится на экране «Организация»")
    owners = [x for x in _users() if _member_role(x, tid) == "owner"]
    if req.remove:
        if target["id"] == u["id"]:
            raise HTTPException(400, "Себя из команды исключить нельзя — "
                                     "передайте владение или удалите команду")
        target["memberships"] = [m for m in (target.get("memberships") or [])
                                 if m.get("tenant") != tid]
        _drop_user_sessions(target["id"], tid)
        _audit("team.member.remove", team=tid, target=uid)
    elif req.role is not None:
        if req.role not in ROLES:
            raise HTTPException(400, "Роль: " + " | ".join(ROLES))
        if req.role != "owner" and len(owners) <= 1 and _member_role(target, tid) == "owner":
            raise HTTPException(400, "В команде должен остаться хотя бы один владелец")
        for m in (target.get("memberships") or []):
            if m.get("tenant") == tid:
                m["role"] = req.role
        _drop_user_sessions(target["id"], tid)   # роль в сессии протухла
        _audit("team.member.role", team=tid, target=uid, role=req.role)
    save_state(STATE)
    return {"ok": True}


@app.post("/api/teams/{tid}/leave")
def team_leave(tid: str, request: Request):
    """Выход из команды по своей воле. Домашнюю организацию покинуть нельзя:
    человек остался бы без рабочего пространства."""
    u = _current_user(request)
    if (u.get("tenant") or DEFAULT_TENANT) == tid:
        raise HTTPException(400, "Домашнюю организацию покинуть нельзя")
    if _member_role(u, tid) is None:
        raise HTTPException(404, "Вы не состоите в этой команде")
    owners = [x for x in _users() if _member_role(x, tid) == "owner"]
    if _member_role(u, tid) == "owner" and len(owners) <= 1:
        raise HTTPException(400, "Вы единственный владелец команды — "
                                 "назначьте другого или удалите команду")
    u["memberships"] = [m for m in (u.get("memberships") or []) if m.get("tenant") != tid]
    _drop_user_sessions(u["id"], tid)
    _audit("team.leave", team=tid)
    save_state(STATE)
    return {"ok": True, "teams": _teams_of(u)}


@app.get("/api/admin/audit")
def admin_audit(request: Request, limit: int = 200, action: str = "", all: bool = False):
    _current_user(request)
    if all and not _is_super(request):
        raise HTTPException(403, "Журнал всех организаций — только суперпользователю")
    t = _current_tenant()
    rows = [r for r in reversed(STATE.get("audit") or [])
            if (all or r.get("tenant") == t) and (not action or r.get("action", "").startswith(action))]
    return {"ok": True, "items": rows[:max(1, min(limit, 1000))]}


_DOMAIN_FIELDS = ("label", "en", "expert", "terminology", "extract", "examples")


class DomainBody(BaseModel):
    id: Optional[str] = None
    base: Optional[str] = None          # встроенная область-шаблон
    label: Optional[str] = None
    en: Optional[str] = None
    expert: Optional[str] = None
    terminology: Optional[str] = None
    extract: Optional[str] = None
    examples: Optional[str] = None
    cats: Optional[List[str]] = None
    strict: Optional[bool] = None


def _domain_public(d: dict) -> dict:
    return {k: d.get(k) for k in ("id", "base", "cats", "strict", "created", "updated") + _DOMAIN_FIELDS} | {"custom": True}


@app.get("/api/admin/domains")
def admin_domains(request: Request):
    _current_user(request)
    return {"ok": True, "builtin": [{"id": d["id"], "label": d["label"], "cats": d["cats"],
                                    "expert": d["expert"], "terminology": d["terminology"],
                                    "extract": d["extract"], "examples": d.get("examples", ""),
                                    "strict": d["id"] in AUTO_APPROVE_BY_DOMAIN} for d in DOMAINS],
            "domains": [_domain_public(d) for d in _tenant_domains()]}


@app.post("/api/admin/domains")
def admin_domain_create(req: DomainBody, request: Request):
    """Своя область — копия встроенного шаблона с правками. Поля ровно те,
    что читают промпты; ничего нового изобретать не надо."""
    _current_user(request)
    base = _DOMAINS_BY_ID.get(req.base or "") or _DOMAINS_BY_ID["general"]
    did = (req.id or "").strip().lower() or re.sub(r"[^a-z0-9]+", "-", (req.label or "").lower()).strip("-")
    if not did:                        # название не латиницей — номерной идентификатор
        did = "area-%d" % (len(STATE.get("domains") or []) + 1)
        while did in _DOMAINS_BY_ID or any(d["id"] == did for d in _tenant_domains()):
            did = did + "x"
    if not re.fullmatch(r"[a-z0-9-]{2,32}", did or ""):
        raise HTTPException(400, "Идентификатор области: 2–32 символа, a-z, 0-9, дефис")
    if did in _DOMAINS_BY_ID or any(d["id"] == did for d in _tenant_domains()):
        raise HTTPException(409, "Область с таким идентификатором уже есть")
    today = datetime.now().strftime("%Y-%m-%d")
    d = {"id": did, "tenant": _current_tenant(), "base": base["id"], "custom": True,
         "cats": [c.strip() for c in (req.cats or base["cats"]) if c.strip()] or ["Term", "Document"],
         "strict": True if req.strict is None else bool(req.strict),
         "created": today, "updated": today}
    for k in _DOMAIN_FIELDS:
        v = getattr(req, k)
        d[k] = (v.strip() if isinstance(v, str) and v.strip() else base.get(k, ""))
    STATE.setdefault("domains", []).append(d)
    _audit("domain.create", domain=did)
    save_state(STATE)
    return {"ok": True, "domain": _domain_public(d)}


@app.post("/api/admin/domains/{did}")
def admin_domain_update(did: str, req: DomainBody, request: Request):
    _current_user(request)
    d = next((x for x in _tenant_domains() if x["id"] == did), None)
    if not d:
        raise HTTPException(404, "Область не найдена")
    for k in _DOMAIN_FIELDS:
        v = getattr(req, k)
        if v is not None:
            d[k] = v.strip()
    if req.cats is not None:
        d["cats"] = [c.strip() for c in req.cats if c.strip()] or d["cats"]
    if req.strict is not None:
        d["strict"] = bool(req.strict)
    d["updated"] = datetime.now().strftime("%Y-%m-%d")
    _audit("domain.update", domain=did)
    save_state(STATE)
    return {"ok": True, "domain": _domain_public(d)}


@app.delete("/api/admin/domains/{did}")
def admin_domain_delete(did: str, request: Request):
    _current_user(request)
    d = next((x for x in _tenant_domains() if x["id"] == did), None)
    if not d:
        raise HTTPException(404, "Область не найдена")
    used = [p["id"] for p in _tenant_projects() if p.get("domain") == did]
    if used:
        raise HTTPException(409, "Область используют проекты: %s" % used[:10])
    STATE["domains"] = [x for x in STATE.get("domains", []) if x is not d]
    _audit("domain.delete", domain=did)
    save_state(STATE)
    return {"ok": True}


class ProjectDomainRequest(BaseModel):
    domain: str


@app.post("/api/projects/{pid}/domain")
def set_project_domain(pid: int, req: ProjectDomainRequest):
    """Смена области у готового проекта. Меняется область глоссария — то есть
    состав приказных терминов, — поэтому оценки back-check с претензией
    «потерян термин» устаревают; пересчёт бесплатный (`/backcheck/rescore`,
    `force: true`), и ответ его называет."""
    project = get_project(pid)
    dom = _resolve_domain(req.domain)
    if dom["id"] != (req.domain or "").strip():
        raise HTTPException(400, "Неизвестная область: %r" % req.domain)
    prev = project.get("domain")
    project["domain"] = dom["id"]
    n_hard = sum(1 for g in STATE["glossary"]
                 if _scope_of(g) == _project_scope(project) and _hit_tier(g) == GLOSSARY_TIER_HARD)
    _audit("project.domain", project=pid, domain=dom["id"], prev=prev)
    save_state(STATE)
    return {"ok": True, "domain": dom["id"], "prev": prev, "verifiedTerms": n_hard,
            "note": "Состав приказных терминов изменился: оценки с претензией «потерян термин» "
                    "устарели — пересчитайте back-check (бесплатно, /backcheck/rescore с force)."}


@app.delete("/api/admin/users/{uid}")
def admin_user_delete(uid: int, request: Request):
    """Убрать учётную запись. Отключение (`active: false`) закрывает вход,
    но оставляет строку — этого мало, когда в базе копится мусор
    самостоятельных регистраций. Два запрета: себя и последнего владельца
    организации (иначе у неё не останется никого, кто вправе её вести)."""
    me = _current_user(request)
    u = next((x for x in _users() if x["id"] == uid
              and (x.get("tenant") == me.get("tenant") or _is_super(request))), None)
    if not u:
        raise HTTPException(404, "Пользователь не найден")
    if u["id"] == me["id"]:
        raise HTTPException(400, "Нельзя удалить самого себя")
    if u.get("role") == "owner":
        others = [x for x in _users() if x.get("tenant") == u.get("tenant")
                  and x["id"] != uid and x.get("role") == "owner" and x.get("active", True)]
        if not others:
            raise HTTPException(409, "Это последний владелец организации: сначала назначьте другого")
    _users().remove(u)
    _drop_user_sessions(uid)
    _audit("user.delete", target=uid, login=u.get("login"))
    save_state(STATE)
    return {"ok": True}


@app.delete("/api/admin/tenants/{tid}")
def admin_tenant_delete(tid: str, request: Request):
    """Снести организацию целиком — вместе с её людьми. Только суперпользователь
    и только пустую: проекты (а с ними переводы и оплаченная работа) молча
    не удаляются НИКОГДА. Глоссарий, память и очередь организации уходят
    вместе с ней — они и так видны только ей."""
    me = _current_user(request)
    if not me.get("super"):
        raise HTTPException(403, "Удалять организации вправе только суперпользователь")
    if tid == DEFAULT_TENANT:
        raise HTTPException(400, "Организацию по умолчанию удалить нельзя")
    rec = _tenant_rec(tid)
    if not rec:
        raise HTTPException(404, "Организация не найдена")
    if me.get("tenant") == tid:
        raise HTTPException(400, "Нельзя удалить организацию, в которой вы сами состоите")
    projects = [p["id"] for p in STATE["projects"] if _tenant_of(p) == tid]
    if projects:
        raise HTTPException(409, "В организации есть проекты: %s — удалите их сначала" % projects[:10])
    users = [u for u in _users() if u.get("tenant") == tid]
    for u in users:
        _drop_user_sessions(u["id"])
        _users().remove(u)
    # Люди из ДРУГИХ организаций, состоявшие в этой командой: членство
    # снимается вместе с самой командой. Иначе оно остаётся висеть, в списке
    # команд человека появляется строка без организации, а переключение на
    # неё сажает его в рабочее пространство, которого больше нет.
    # Приглашения этой команды закрываются по той же причине.
    orphans = 0
    for u in _users():
        rest = [m for m in (u.get("memberships") or []) if m.get("tenant") != tid]
        if len(rest) != len(u.get("memberships") or []):
            u["memberships"] = rest
            _drop_user_sessions(u["id"], tid)
            orphans += 1
    STATE["invites"] = [i for i in _invites() if i.get("tenant") != tid]
    for coll in ("glossary", "tm", "termQueue", "autoBatches", "exportHistory",
                 "domains", "quotes"):
        rows = STATE.get(coll)
        if isinstance(rows, list):
            STATE[coll] = [r for r in rows if _tenant_of(r) != tid]
    _tenants().remove(rec)
    _invalidate_gloss_index()
    _audit("tenant.delete", tenant_target=tid, users=len(users), members=orphans)
    save_state(STATE)
    return {"ok": True, "usersRemoved": len(users), "membershipsRemoved": orphans}


class TenantPatch(BaseModel):
    name: Optional[str] = None
    active: Optional[bool] = None
    limitUsd: Optional[float] = None      # 0 — запретить платное; None — оставить как есть
    clearLimit: bool = False


@app.post("/api/admin/tenants/{tid}")
def admin_tenant_update(tid: str, req: TenantPatch, request: Request):
    me = _current_user(request)
    if not me.get("super"):
        raise HTTPException(403, "Лимиты и организации правит только суперпользователь")
    rec = _tenant_rec(tid)
    if not rec:
        raise HTTPException(404, "Организация не найдена")
    if req.name is not None:
        rec["name"] = req.name.strip() or rec["name"]
    if req.active is not None:
        rec["active"] = bool(req.active)
    if req.clearLimit:
        rec.pop("limitUsd", None)
    elif req.limitUsd is not None:
        if req.limitUsd < 0:
            raise HTTPException(400, "Лимит не может быть отрицательным")
        rec["limitUsd"] = float(req.limitUsd)
    _audit("tenant.update", tenant_target=tid, limitUsd=rec.get("limitUsd"))
    save_state(STATE)
    return {"ok": True, "tenant": rec, "spend": _spend_status(tid)}


@app.get("/api/admin/overview")
def admin_overview(request: Request):
    """Одним экраном для администратора сервиса: организации с людьми,
    проектами и расходом; прогоны всех организаций; расход процесса;
    здоровье. Ни одного вызова модели; STATE не обходится по сегментам —
    только счётчики."""
    me = _current_user(request)
    if not me.get("super"):
        raise HTTPException(403, "Сводка — только суперпользователю")
    users, projects = _users(), STATE["projects"]
    tenants = []
    for t in _tenants():
        tid = t["id"]
        tp = [p for p in projects if _tenant_of(p) == tid]
        # Прайс агентства суперпользователю не показываем: лимит расхода —
        # наше дело, цена страницы — то, что агентство продаёт своим клиентам.
        tenants.append({**{k: v for k, v in t.items() if k != "pricing"}, "users": sum(1 for u in users if u.get("tenant") == tid),
                        "activeUsers": sum(1 for u in users if u.get("tenant") == tid and u.get("active", True)),
                        "projects": len(tp), "segments": sum(len(p["segments"]) for p in tp),
                        "glossary": sum(1 for g in STATE["glossary"] if _tenant_of(g) == tid),
                        "domains": len(_tenant_domains(tid)),
                        "spend": _spend_status(tid)})
    jobs = sorted(_JOBS.values(), key=lambda j: j["id"], reverse=True)
    with _USAGE_LOCK:
        proc = json.loads(json.dumps(_USAGE_TOTAL))
    try:
        state_bytes = STATE_FILE.stat().st_size
    except Exception:
        state_bytes = None
    worker_alive = bool(_JOB_WORKER and _JOB_WORKER.is_alive())
    return {"ok": True,
            "tenants": tenants,
            "jobs": {"active": [dict(_job_public(j), tenant=_tenant_of(j))
                                for j in jobs if j["status"] in ("queued", "running")],
                     "recent": [dict(_job_public(j), tenant=_tenant_of(j)) for j in jobs[:15]],
                     "queued": _JOB_QUEUE.qsize(), "workerAlive": worker_alive},
            "process": {"uptimeSec": int(time.time() - _SERVER_STARTED), "usage": proc,
                        "stateBytes": state_bytes, "sessions": len(_SESSIONS),
                        "openaiKey": bool(os.environ.get("OPENAI_API_KEY")),
                        "version": "5.6.0", "termQueue": len(_term_queue()),
                        "auditRows": len(STATE.get("audit") or [])},
            "month": _month_key()}


@app.get("/api/admin/tenants")
def admin_tenants(request: Request):
    me = _current_user(request)
    if not me.get("super"):
        raise HTTPException(403, "Список организаций — только суперпользователю")
    return {"ok": True, "tenants": [dict({k: v for k, v in t.items() if k != "pricing"},
                                         spend=_spend_status(t["id"])) for t in _tenants()]}


class TenantCreate(BaseModel):
    id: str
    name: str
    ownerLogin: str
    ownerPassword: str


@app.post("/api/admin/tenants")
def admin_tenant_create(req: TenantCreate, request: Request):
    """Только с флагом super: организация и её первый владелец одним шагом."""
    _audit("tenant.create", tenant_new=req.id)
    me = _current_user(request)
    if not me.get("super"):
        raise HTTPException(403, "Заводить организации вправе только суперпользователь")
    tid = req.id.strip().lower()
    if not re.fullmatch(r"[a-z0-9-]{2,32}", tid):
        raise HTTPException(400, "Идентификатор организации: 2–32 символа, a-z, 0-9, дефис")
    if any(t.get("id") == tid for t in _tenants()):
        raise HTTPException(409, "Организация с таким идентификатором уже есть")
    _check_user_fields(req.ownerLogin, req.ownerPassword, "owner")
    if _user_by_login(req.ownerLogin):
        raise HTTPException(409, "Такой логин уже есть")
    _tenants().append({"id": tid, "name": req.name.strip() or tid,
                       "created": datetime.now().strftime("%Y-%m-%d"), "active": True})
    h, salt = _hash_password(req.ownerPassword)
    u = {"id": max((x["id"] for x in _users()), default=0) + 1, "tenant": tid,
         "login": req.ownerLogin.strip(), "hash": h, "salt": salt, "role": "owner",
         "name": req.ownerLogin.strip(), "active": True, "uiLang": DEFAULT_UI_LANG,
         "created": datetime.now().strftime("%Y-%m-%d")}
    _users().append(u)
    save_state(STATE)
    return {"ok": True, "tenant": _tenants()[-1], "owner": _user_public(u)}


# ═══ Стоимость перевода: страницы и цена ════════════════════════════
# Перевод продают СТРАНИЦАМИ исходника, а не токенами. Поэтому счёт знаков
# и цена лежат отдельно от расхода на модели (`_spend_*`, `limitUsd`): то —
# наши затраты на прогоны, это — выручка организации, и путать их нельзя.
# Ни один эндпоинт этого раздела не зовёт модель, значит в `_PAID` им не место:
# на исчерпанном лимите смету посчитать по-прежнему можно.
#
# Ценовая карточка живёт В ЗАПИСИ ОРГАНИЗАЦИИ (`tenant["pricing"]`), а НЕ
# отдельным верхним ключом STATE, и довод ровно один: `/api/seed` отдаёт ВСЕ
# верхние ключи, кроме `users`/`tenants`/`audit`/`spend`, — новый ключ уехал бы
# каждому вошедшему вместе с ценами чужих агентств, и держалась бы эта тайна
# на строчке фильтра, которую забудут при следующем ключе. Записи организаций
# из выдачи исключены уже сегодня, и лимит расхода лежит там же.
# Про ЗАПИСЬ в хранилище это ничего не меняет, и обещать обратное нечестно:
# `tenants` — такой же ОДИН документ на всех арендаторов (`store._docs_of`),
# каким был бы и `pricing`, и два владельца, правящих цены одновременно,
# дерутся за него ровно так же, как сегодня дерутся за лимиты. Сеть
# безопасности та же — версия документа (`DocConflict`).
from decimal import Decimal, ROUND_HALF_UP

PRICING_DEFAULTS = {"currency": "USD", "default": None, "rates": [], "norms": {},
                    "minPages": 1.0, "roundTo": 0.1}
PRICE_MAX = 100000.0          # цена страницы: защита от опечатки в ноль лишний
NORM_MIN, NORM_MAX = 100, 20000


def _pricing_of(tid: Optional[str] = None) -> dict:
    """Карточка организации, дополненная умолчаниями. Отсутствие карточки —
    это «цена не задана», а не «бесплатно»: `default: None` доезжает до ответа
    и превращается там в честный отказ считать сумму."""
    rec = _tenant_rec(tid or _current_tenant()) or {}
    return {**PRICING_DEFAULTS, **(rec.get("pricing") or {})}


def _rate_for(card: dict, src: str, tgt: str) -> dict:
    """Цена за страницу для ПАРЫ языков: сначала строка прайса, затем общая
    цена карточки. Не нашлось — `price: None`, и сумма не считается вовсе.
    Ноль вместо неизвестной цены показал бы клиенту бесплатный перевод."""
    s, t = (src or "").strip().upper(), (tgt or "").strip().upper()
    for r in card.get("rates") or []:
        if str(r.get("src", "")).upper() == s and str(r.get("tgt", "")).upper() == t:
            return {"price": float(r.get("price") or 0), "source": "pair"}
    if card.get("default") is not None:
        return {"price": float(card["default"]), "source": "default"}
    return {"price": None, "source": None}


def _money(pages: float, price: Optional[float]) -> Optional[float]:
    """Деньги считаются десятичными, а не двоичными: 0.1 страницы по $12.30
    во float даёт неповторяемый хвост, а смета — это оферта, и она обязана
    воспроизводиться до копейки."""
    if price is None:
        return None
    # Две цифры после запятой — при любой валюте: в UZS и JPY копеек нет,
    # и лишние нули там просто ничего не значат. Округлять до целых по коду
    # валюты нельзя — округление это условие договора, а не свойство денег.
    return float((Decimal(str(pages)) * Decimal(str(price))).quantize(
        Decimal("0.01"), rounding=ROUND_HALF_UP))


def _quote_of(counts: dict, src: str, tgt: str, card: dict, basis: str,
              notes: Optional[list] = None) -> dict:
    """Смета по готовому счёту знаков. `basis` называет, ЧТО посчитано:
    `file` — весь текст файла, `segments` — сегменты проекта. Две величины
    на одном документе расходятся (импорт выбрасывает чисто цифровые абзацы
    и склеивает соседние повторы), и подать их под одним именем значило бы
    показать человеку две разные суммы за одну работу без объяснения."""
    norm = textcount.norm_for(src, card.get("norms"))
    notes = list(notes or [])
    if norm["source"] == "default":
        # Оговорка живёт ЗДЕСЬ, а не только в textcount.measure: смету
        # по сегментам проекта считают мимо measure, и там догадка о норме
        # молчала бы — то есть в деньгах.
        notes.append("Нормы для языка %s в таблице нет — взята норма по умолчанию (%d знаков). "
                     "Задайте свою в ценовой карточке." % (norm["lang"] or "?", norm["chars"]))
    pages = textcount.pages_of(counts["chars"], norm["chars"],
                               card.get("minPages", 1.0), card.get("roundTo", 0.1))
    rate = _rate_for(card, src, tgt)
    total = _money(pages["billed"], rate["price"])
    return {"basis": basis, "src": src, "tgt": tgt, "counts": counts, "norm": norm,
            "pages": pages, "rate": rate, "currency": card.get("currency") or "USD",
            "total": total, "pricingUpdated": card.get("updated"), "notes": notes,
            # Расчёт строкой — чтобы человек мог проверить сумму глазами,
            # а не верить ей на слово.
            "formula": "%d знаков с пробелами ÷ %d = %s стр.; к оплате %s стр. × %s %s = %s"
                       % (counts["chars"], norm["chars"], pages["exact"], pages["billed"],
                          rate["price"] if rate["price"] is not None else "—",
                          card.get("currency") or "USD",
                          total if total is not None else "цена не задана")}


def _docx_bill_paragraphs(content: bytes) -> list:
    """Текст .docx для СЧЁТА — ровно тот, куда экспорт «как в оригинале»
    пишет перевод (`_para_slots`).

    Почему не `_docx_paragraphs`: тот склеивает `.//w:t` всего абзаца, то есть
    берёт и вложенные надписи (их абзацы идут в списке ещё раз — двойной счёт),
    и скрытый текст, и результат вычисляемых полей — номера страниц оглавления.
    Клиент платил бы за то, что переводить не станут. И не `_docx_units`: тот
    режет короткие абзацы и схлопывает соседние повторы — это правила ИМПОРТА,
    а объём файла от них не уменьшается."""
    from docx import Document
    from docx.oxml.ns import qn as _qn
    doc = Document(io.BytesIO(content))
    out = []
    for p in _docx_flat_paragraphs(doc):
        slots, _full, _dropped = _para_slots(p, _qn)
        out.append(_docx_clean("".join((t.text or "") for t, _sig in slots)))
    return out


def _measure_kwargs(card: dict) -> dict:
    return {"overrides": card.get("norms"), "min_pages": card.get("minPages", 1.0),
            "round_to": card.get("roundTo", 0.1),
            "docx_paragraphs": _docx_bill_paragraphs}


def _count_error(e: Exception) -> HTTPException:
    """Три разных отказа и три разных кода: 413 — файл больше потолка,
    503 — формат поддержан, но библиотеки нет, 415 — не разбираем такое.
    Один общий код заставлял бы человека гадать, что делать дальше."""
    if isinstance(e, textcount.TooBig):
        return HTTPException(413, str(e))
    if isinstance(e, textcount.NotAvailable):
        return HTTPException(503, str(e))
    return HTTPException(415, str(e))


class RateRow(BaseModel):
    src: str
    tgt: str
    price: float


class PricingBody(BaseModel):
    currency: Optional[str] = None
    default: Optional[float] = None       # цена страницы для пар без своей строки
    clearDefault: bool = False            # снять общую цену (None ≠ 0)
    rates: Optional[List[RateRow]] = None
    norms: Optional[dict] = None          # переопределение нормы: {"RU": 1800}
    minPages: Optional[float] = None
    roundTo: Optional[float] = None


@app.get("/api/pricing")
def get_pricing(request: Request):
    """Ценовая карточка СВОЕЙ организации плюс таблица норм страницы.

    Норму и цены отдаёт сервер, а не хранит браузер: второй прайс-лист в `.jsx`
    — это ровно та беда, ради которой модели и их цены живут в `OPENAI_MODELS`
    и уезжают через `/api/models`."""
    _current_user(request)
    t = textcount.norms()
    return {"ok": True, "tenant": _current_tenant(), "pricing": _pricing_of(),
            "norms": {"default": t["default"], "basis": t["basis"],
                      "rows": sorted(t["rows"].values(), key=lambda r: r["lang"])}}


@app.post("/api/pricing")
def save_pricing(req: PricingBody, request: Request):
    """Правит цены ВЛАДЕЛЕЦ организации (строка в `_OWNER_ONLY`), а не
    суперпользователь: `limitUsd` — наш расход на модели и наше дело, а цена
    страницы — то, что агентство продаёт своим клиентам, и в это мы не лезем.
    Переводчику остаётся GET: смету он посчитать вправе, прайс — нет."""
    me = _current_user(request)
    rec = _tenant_rec(_current_tenant())
    if not rec:
        raise HTTPException(404, "Записи вашей организации нет в базе — прайс сохранять некуда. "
                                 "Это сбой данных, а не права: позовите администратора сервиса.")
    card = {**PRICING_DEFAULTS, **(rec.get("pricing") or {})}
    if req.currency is not None:
        cur = req.currency.strip().upper()
        if not re.fullmatch(r"[A-Z]{3}", cur):
            raise HTTPException(400, "Валюта — три буквы кода ISO 4217: USD, EUR, UZS")
        card["currency"] = cur
    if req.clearDefault:
        card["default"] = None
    elif req.default is not None:
        if not (0 < req.default <= PRICE_MAX):
            # Ноль — это не цена, а «не задана»: пустое поле в браузере
            # приходит нулём, и без этой проверки клиенту показали бы
            # бесплатный перевод. Снимают цену полем clearDefault.
            raise HTTPException(400, "Цена страницы: больше нуля и не выше %g "
                                     "(снять цену — пустым полем)" % PRICE_MAX)
        card["default"] = float(req.default)
    if req.rates is not None:
        rows, seen = [], set()
        for r in req.rates:
            s, t = _check_lang_pair(r.src, r.tgt)   # пара проверяется каталогом языков
            if (s, t) in seen:
                raise HTTPException(400, "Пара %s→%s указана дважды" % (s, t))
            if not (0 < r.price <= PRICE_MAX):
                raise HTTPException(400, "Цена %s→%s: больше нуля и не выше %g "
                                         "(бесплатной строки в прайсе не бывает — уберите строку)"
                                    % (s, t, PRICE_MAX))
            seen.add((s, t))
            rows.append({"src": s, "tgt": t, "price": float(r.price)})
        card["rates"] = rows
    if req.norms is not None:
        norms = {}
        for k, v in req.norms.items():
            code = str(k).strip().upper()
            if not re.fullmatch(r"[A-Z]{2}", code):
                raise HTTPException(400, "Код языка нормы: две буквы, а не %r" % k)
            try:
                n = int(v)
            except (TypeError, ValueError):
                raise HTTPException(400, "Норма для %s — целое число знаков" % code)
            if not (NORM_MIN <= n <= NORM_MAX):
                raise HTTPException(400, "Норма для %s вне разумного: %d…%d"
                                    % (code, NORM_MIN, NORM_MAX))
            norms[code] = n
        card["norms"] = norms
    if req.minPages is not None:
        if not (0 <= req.minPages <= 100):
            raise HTTPException(400, "Минимальный заказ: 0…100 страниц")
        card["minPages"] = float(req.minPages)
    if req.roundTo is not None:
        if not (0 <= req.roundTo <= 1):
            raise HTTPException(400, "Шаг округления: от 0 (без округления) до 1 страницы")
        card["roundTo"] = float(req.roundTo)
    card["updated"] = datetime.now().strftime("%Y-%m-%d %H:%M")
    card["by"] = me.get("login")
    rec["pricing"] = card
    _audit("pricing.update", currency=card["currency"], default=card["default"],
           rates=len(card.get("rates") or []), norms=len(card.get("norms") or {}))
    save_state(STATE)
    return {"ok": True, "pricing": card}


@app.post("/api/quote")
async def quote_file(request: Request, file: UploadFile = File(...),
                     src: str = Form("RU"), tgt: str = Form("EN"),
                     save: bool = Form(True)):
    """Сколько знаков, страниц и денег в ЭТОМ файле — до всякого импорта.

    Файл никуда не сохраняется — писать можно только в `data/`, и незачем:
    это расчёт, а не проект. А вот САМА смета уходит в историю организации
    (`save=false` отключает): человек считает сегодня, платит через неделю,
    и к числам надо иметь возможность вернуться. Вызовов модели нет, поэтому
    команда бесплатна и работает на исчерпанном лимите. Разбор идёт
    в отдельном потоке: воркер ОДИН, и распаковка чужого пакета не должна
    держать всех."""
    _current_user(request)
    src, tgt = _check_lang_pair(src, tgt)
    # Потолок проверяется ДО чтения тела: воркер один, и полугигабайтная
    # загрузка не должна сначала лечь в память, а потом получить отказ.
    declared = request.headers.get("content-length")
    if declared and declared.isdigit() and int(declared) > textcount.MAX_BYTES * 1.1:
        raise HTTPException(413, "Файл больше %d МБ — разберите его по частям"
                            % (textcount.MAX_BYTES // 1024 // 1024))
    content = await file.read()
    card = _pricing_of()
    try:
        m = await run_in_threadpool(textcount.measure, file.filename or "", content,
                                    src, **_measure_kwargs(card))
    except (textcount.Unsupported, textcount.NotAvailable, textcount.TooBig) as e:
        raise _count_error(e)
    q = _quote_of(m["counts"], src, tgt, card, "file", m["notes"])
    saved = None
    if save:
        # Отпечаток содержимого, а не имя файла: «договор.docx» бывает разным,
        # и по имени повторный расчёт того же файла не опознать.
        sha = hashlib.sha256(content).hexdigest()[:16]
        saved = _quote_save(q, file.filename or "", m["kind"], sha)
    return {"ok": True, "file": m["file"], "kind": m["kind"], **q,
            "saved": saved, "supported": textcount.SUPPORTED_EXT}


@app.get("/api/projects/{pid}/quote")
def quote_project(pid: int, request: Request, withFile: bool = False):
    """Смета по УЖЕ разобранному проекту: считаются исходники его сегментов.

    Это ДРУГАЯ величина, чем `/api/quote` по тому же файлу, и ответ говорит
    об этом прямо: импорт выбрасывает чисто цифровые и совсем короткие абзацы,
    а соседние одинаковые склеивает в один сегмент. Сегменты, заведённые
    разбором картинок, сюда входят — их тоже переводят.
    `withFile=true` доcчитывает объём приложенного исходника для сравнения:
    он читает и разбирает весь .docx, поэтому по умолчанию выключен."""
    _current_user(request)
    project = get_project(pid)
    card = _pricing_of()
    src = project.get("src") or "RU"
    tgt = project.get("tgt") or "EN"
    segs = project.get("segments") or []
    counts = textcount.count_blocks([s.get("source") or "" for s in segs])
    notes = ["Посчитаны исходники %d сегментов проекта. Объём файла бывает больше: "
             "импорт не заводит сегменты на чисто цифровые и очень короткие абзацы, "
             "а соседние одинаковые склеивает в один." % len(segs)]
    imgs = sum(1 for s in segs if (s.get("origin") or {}).get("kind") == "image")
    if imgs:
        notes.append("Из них %d — надписи, распознанные на картинках." % imgs)
    q = _quote_of(counts, src, tgt, card, "segments", notes)
    out = {"ok": True, "project": pid, "segments": len(segs), **q}
    if withFile:
        # Сравнение с исходником: расхождение объёмов — законный вопрос
        # клиента, и отвечать на него молчанием нельзя.
        info = _load_source_map(pid)
        if not info:
            out["file"] = {"error": "Исходник к проекту не приложен"}
        else:
            try:
                # Через measure, а не своим путём: те же потолки размера и та же
                # обработка отказов, что у загруженного файла. Мимо них чужой
                # .docx на сотню мегабайт душил бы единственный воркер именно
                # отсюда — из безобидного на вид GET.
                fm = textcount.measure(info["path"].name, info["path"].read_bytes(),
                                       src, **_measure_kwargs(card))
                out["file"] = {"counts": fm["counts"], "pages": fm["pages"],
                               "notes": fm["notes"]}
            except Exception as e:
                out["file"] = {"error": "Исходник не разобрался: %s" % e}
    return out


# ── История смет: к чему возвращаться при оплате ────────────────────
# Смета — ОФЕРТА, а не расчёт «на сейчас»: клиент увидел сумму, ушёл думать,
# вернулся через неделю платить. За это время прайс мог смениться, файл —
# исчезнуть, норма — быть переопределена. Поэтому в историю кладутся ЧИСЛА,
# а не ссылка на пересчёт: знаки, страницы, норма, цена страницы и итог
# замораживаются в момент расчёта и потом не пересчитываются НИКОГДА.
# Пересчитанная задним числом смета — это другая сумма под тем же счётом.
#
# Файл не хранится (писать можно только в data/, и незачем — это расчёт,
# а не проект), поэтому запись несёт отпечаток содержимого (sha) и имя:
# по отпечатку видно, что принесли ТОТ ЖЕ файл, а не похожий.
QUOTE_HISTORY_MAX = max(50, int(os.environ.get("QUOTE_HISTORY_MAX", "500")))
QUOTE_STATUSES = ("new", "invoiced", "paid")


def _quotes() -> list:
    return STATE.setdefault("quotes", [])


def _tenant_quotes() -> list:
    t = _current_tenant()
    return [q for q in _quotes() if _tenant_of(q) == t]


def _quote_fingerprint(rec: dict) -> tuple:
    """Что делает смету ТОЙ ЖЕ: тот же файл, та же пара, та же норма и цена.
    Смени прайс — и повторный расчёт станет НОВОЙ записью, потому что это
    другая оферта. Без этого история заросла бы копиями одного нажатия."""
    return (rec.get("sha"), rec.get("src"), rec.get("tgt"), rec.get("normChars"),
            rec.get("pricePerPage"), rec.get("currency"), rec.get("pagesBilled"))


def _quote_save(q: dict, filename: str, kind: str, sha: str) -> dict:
    """Положить смету в историю. Повторный расчёт того же файла по той же цене
    НЕ плодит запись: обновляется время и счётчик обращений — иначе три нажатия
    подряд выглядели бы тремя заказами."""
    rec = {"id": max((x.get("id", 0) for x in _quotes()), default=0) + 1,
           "tenant": _current_tenant(),          # инвариант 11: каждая новая запись несёт организацию
           "at": datetime.now().strftime("%Y-%m-%d %H:%M"),
           "by": (_actor() or {}).get("login"),
           "file": filename, "kind": kind, "sha": sha,
           "src": q["src"], "tgt": q["tgt"],
           "chars": q["counts"]["chars"], "charsNoSpaces": q["counts"]["charsNoSpaces"],
           "words": q["counts"]["words"], "repeatChars": q["counts"]["repeatChars"],
           "normChars": q["norm"]["chars"], "normSource": q["norm"]["source"],
           "pagesExact": q["pages"]["exact"], "pagesBilled": q["pages"]["billed"],
           "minPages": q["pages"]["minPages"], "roundTo": q["pages"]["roundTo"],
           "pricePerPage": q["rate"]["price"], "rateSource": q["rate"]["source"],
           "currency": q["currency"], "total": q["total"], "formula": q["formula"],
           "status": "new", "count": 1}
    fp = _quote_fingerprint(rec)
    for old in _quotes():
        if _tenant_of(old) == rec["tenant"] and _quote_fingerprint(old) == fp:
            old["at"], old["count"] = rec["at"], (old.get("count") or 1) + 1
            save_state(STATE)
            return old
    _quotes().insert(0, rec)
    # Подрезка ПО ОРГАНИЗАЦИИ: общий потолок съедал бы историю тихого клиента
    # ради шумного соседа. Оплаченные и выставленные не выбрасываем — по ним
    # ещё придут деньги; выбрасываются только черновые расчёты.
    mine = [x for x in _quotes() if _tenant_of(x) == rec["tenant"]]
    if len(mine) > QUOTE_HISTORY_MAX:
        drop = {id(x) for x in [y for y in mine if y.get("status") == "new"][QUOTE_HISTORY_MAX:]}
        STATE["quotes"] = [x for x in _quotes() if id(x) not in drop]
    save_state(STATE)
    return rec


@app.get("/api/quotes")
def list_quotes(request: Request, limit: int = 200):
    """История смет своей организации, новые первыми. Числа отдаются такими,
    какими их посчитали тогда: пересчёт по нынешнему прайсу показал бы другую
    сумму под тем же счётом."""
    _current_user(request)
    rows = _tenant_quotes()[:max(1, min(limit, QUOTE_HISTORY_MAX))]
    return {"ok": True, "quotes": rows, "total": len(_tenant_quotes()),
            "statuses": list(QUOTE_STATUSES)}


class QuotePatch(BaseModel):
    status: Optional[str] = None      # new | invoiced | paid
    note: Optional[str] = None


@app.post("/api/quotes/{qid}")
def update_quote(qid: int, req: QuotePatch, request: Request):
    """Пометить смету выставленной или оплаченной. Право владельца: это
    коммерческое решение, а не работа переводчика. Числа сметы не меняются
    здесь НИКОГДА — меняется только её судьба."""
    _current_user(request)
    rec = next((q for q in _tenant_quotes() if q.get("id") == qid), None)
    if not rec:
        raise HTTPException(404, "Смета не найдена")
    if req.status is not None:
        if req.status not in QUOTE_STATUSES:
            raise HTTPException(400, "Состояние сметы: %s" % ", ".join(QUOTE_STATUSES))
        rec["status"] = req.status
        stamp = datetime.now().strftime("%Y-%m-%d %H:%M")
        # След решения человека остаётся и при возврате в «новую»: кто и когда
        # объявил смету оплаченной — вопрос денег, и стирать его нельзя.
        rec.setdefault("log", []).append({"at": stamp, "by": (_actor() or {}).get("login"),
                                          "status": req.status})
        if req.status == "paid":
            rec["paidAt"] = stamp
    if req.note is not None:
        rec["note"] = req.note.strip()[:500]
    _audit("quote.update", quote=qid, status=rec.get("status"))
    save_state(STATE)
    return {"ok": True, "quote": rec}


@app.delete("/api/quotes/{qid}")
def delete_quote(qid: int, request: Request):
    _current_user(request)
    rec = next((q for q in _tenant_quotes() if q.get("id") == qid), None)
    if not rec:
        raise HTTPException(404, "Смета не найдена")
    STATE["quotes"] = [q for q in _quotes() if q is not rec]
    _audit("quote.delete", quote=qid, total=rec.get("total"))
    save_state(STATE)
    return {"ok": True}


@app.post("/api/auth/logout")
def logout(request: Request):
    _drop_session(_token_from_request(request))
    return {"ok": True}


@app.get("/api/seed")
def get_seed():
    """Initial data dump — glossary capped at 150 terms for performance; full list via /api/glossary."""
    # Пользователи и организации в общую выдачу НЕ идут: там хеши паролей.
    t = _current_tenant()
    public = {k: v for k, v in STATE.items() if k not in ("users", "tenants")}
    public.pop("audit", None)
    public.pop("spend", None)
    # Приглашения — почты людей из ДРУГИХ организаций. Верхний ключ уехал бы
    # каждому вошедшему вместе с ними: `/api/seed` отдаёт всё, что не названо.
    public.pop("invites", None)
    for key in ("exportHistory", "termQueue", "autoBatches", "runCosts", "quotes"):
        public[key] = [e for e in (STATE.get(key) or []) if _tenant_of(e) == t]
    return {**public, "projects": [_project_for_client(p) for p in _tenant_projects()],
            "glossary": [g for g in STATE["glossary"] if _tenant_of(g) == t][:150]}


@app.get("/api/glossary")
def list_glossary(q: str = "", cat: str = "", limit: int = 200, offset: int = 0,
                  lang: str = "", domain: str = ""):
    """Full glossary with optional search and pagination.
    lang/domain сужают выдачу до области проекта — той самой, что уходит
    в промпт. Без них отдаём всё: вкладка «Глоссарий» листает общий список."""
    t_id = _current_tenant()
    items = [t for t in STATE["glossary"] if _tenant_of(t) == t_id]
    if lang or domain:
        items = [t for t in items
                 if (not lang or _scope_of(t)[0] == lang)
                 and (not domain or _scope_of(t)[1] == domain)]
    if cat and cat != "all":
        items = [t for t in items if t.get("cat") == cat]
    if q:
        ql = q.lower()
        items = [t for t in items if ql in t.get("src", "").lower() or ql in t.get("tgt", "").lower()]
    total = len(items)
    return {"total": total, "items": items[offset:offset + limit]}


@app.post("/api/glossary/import")
async def import_glossary(request: Request, file: UploadFile = File(...),
                          lang: str = Form(""), domain: str = Form(""),
                          tier: str = Form(GLOSSARY_TIER_SOFT), dry_run: bool = Form(True)):
    """Словарь клиента приходит ФАЙЛОМ, а не из кода: стартовый глоссарий пуст.

    TSV/CSV; колонки по заголовку (src/source/original, tgt/target/translation,
    cat, note) либо первые две. Область — пара и тематика, организация —
    текущей сессии. Уровень по умолчанию — ПОДСКАЗКА: приказом («use these
    exact translations») импорт становится только по явному выбору владельца,
    потому что приказ уходит модели мимо всех проверок. Повторы в пределах
    области пропускаются — импорт не переписывает то, что уже решено.
    `dry_run` по умолчанию: сначала числа и образец, потом запись."""
    import csv, io as _io
    me = _current_user(request)
    if tier not in (GLOSSARY_TIER_SOFT, GLOSSARY_TIER_HARD):
        raise HTTPException(400, "Уровень: %s | %s" % (GLOSSARY_TIER_SOFT, GLOSSARY_TIER_HARD))
    # Роль — из СЕССИИ (роль в активной команде), а не с записи пользователя:
    # владелец своей команды в чужой может быть переводчиком, и запись пустила
    # бы его импортировать приказы в чужой глоссарий.
    if tier == GLOSSARY_TIER_HARD and _actor_role() != "owner":
        raise HTTPException(403, "Импорт приказом — только владелец организации")
    pair = (lang or "").replace("->", "→").strip().upper()
    if "→" not in pair:
        raise HTTPException(400, "Языковая пара вида RU→EN")
    src_l, tgt_l = _check_lang_pair(*pair.split("→", 1))
    scope = _scope(f"{src_l}→{tgt_l}", _resolve_domain(domain)["id"])
    raw = (await file.read()).decode("utf-8-sig", errors="replace")
    first = raw.split("\n", 1)[0]
    delim = "\t" if "\t" in first else (";" if first.count(";") > first.count(",") else ",")
    rows = list(csv.reader(_io.StringIO(raw), delimiter=delim))
    if not rows:
        raise HTTPException(400, "Файл пуст")
    head = [h.strip().lower() for h in rows[0]]
    names = {"src": ("src", "source", "original", "term", "russian", "оригинал", "термин"),
             "tgt": ("tgt", "target", "translation", "english", "перевод"),
             "cat": ("cat", "category", "категория"), "note": ("note", "comment", "примечание")}
    col = {}
    for k, alts in names.items():
        for i, h in enumerate(head):
            if h in alts:
                col[k] = i
                break
    has_head = "src" in col and "tgt" in col
    if not has_head:
        col = {"src": 0, "tgt": 1}
    body = rows[1:] if has_head else rows
    existing = {_norm_key(g.get("src")) for g in STATE["glossary"] if _scope_of(g) == scope}
    today = datetime.now().strftime("%Y-%m-%d")
    added, seen, dup, bad = [], set(), 0, 0
    # Приказ подписывает владелец — на КАЖДОЙ записи, потому что дальше она
    # живёт и правится отдельно от файла. Подсказка подписи не несёт: за неё
    # никто не ручается, а десять тысяч записей × след — лишний вес документа.
    signed_import = _signed_field("import") if tier == GLOSSARY_TIER_HARD else {}
    for r in body:
        get = lambda k: (r[col[k]].strip() if k in col and col[k] < len(r) else "")
        s_, t_ = get("src"), get("tgt")
        if not s_ or not t_ or len(s_) > 200 or len(t_) > 200:
            bad += 1
            continue
        key = _norm_key(s_)
        if key in existing or key in seen:
            dup += 1
            continue
        seen.add(key)
        added.append({"src": s_, "tgt": t_, "cat": get("cat") or "Term", "freq": 1,
                      "conf": "high" if tier == GLOSSARY_TIER_HARD else "medium",
                      "note": get("note"), "tier": tier,
                      "origin": ("import:" + (file.filename or "file"))[:60],
                      "lang": scope[0], "domain": scope[1], "tenant": scope[2], "updated": today,
                      **signed_import})
    out = {"ok": True, "dryRun": dry_run, "rows": len(body), "added": len(added),
           "skippedDup": dup, "skippedBad": bad, "tier": tier,
           "lang": scope[0], "domain": scope[1], "header": has_head,
           "sample": [{"src": a["src"], "tgt": a["tgt"]} for a in added[:10]]}
    if not dry_run and added:
        STATE["glossary"][0:0] = added
        _audit("glossary.import", file=file.filename, added=len(added), tier=tier)
        _invalidate_gloss_index()
        save_state(STATE)
    return out


@app.get("/api/glossary/usage")
def glossary_usage(src: str, limit: int = 6, lang: str = "", domain: str = ""):
    """Где термин реально встречается. Совпадение ищем тем же _term_match, что и
    инъекция в промпт: иначе список «затронутых сегментов» разошёлся бы с тем,
    на что глоссарий действительно влияет.

    В `violating` — сегменты, где термин есть в оригинале, перевод уже готов,
    а утверждённого варианта в нём нет. Именно их имеет смысл переперевести."""
    # Без области сюда попадала запись чужой языковой пары, и тогда «нарушением
    # глоссария» помечался каждый сегмент: искомого перевода там и не могло быть.
    # Пустые lang/domain читаем как область по умолчанию — так же, как их читает
    # _scope_of; запасной поиск по имени срабатывает, только если запись одна.
    want = _scope(lang, domain)
    entry = _glossary_entry(src, want)
    if entry is None and not (lang or domain):
        same = [g for g in STATE["glossary"] if _norm_key(g.get("src")) == _norm_key(src)
                and _tenant_of(g) == _current_tenant()]
        entry = same[0] if len(same) == 1 else None
        if entry is not None:
            want = _scope_of(entry)
    tgt = (entry or {}).get("tgt", "")
    # Язык оригинала — из области записи: от него зависит, работает ли
    # стем-поиск и по какой таблице окончаний.
    want_lang = _src_lang({"lang": want[0]})
    projects, examples = [], []
    for p in _tenant_projects():
        # Проект чужой области этой записью не управляется: помечать его
        # сегменты нарушением глоссария — то же враньё, только на уровне отчёта.
        if entry is not None and _project_scope(p) != want:
            continue
        ids, violating = [], []
        for seg in p["segments"]:
            if not _term_match(src, seg.get("source", ""), want_lang):
                continue
            ids.append(seg["id"])
            target = (seg.get("target") or "").strip()
            if target and tgt and not _tgt_has_term(target, tgt):
                violating.append(seg["id"])
            if len(examples) < max(1, min(limit, 20)):
                examples.append({"project": p["id"], "projectTitle": p.get("title", ""),
                                 "id": seg["id"], "source": seg.get("source", "")[:400],
                                 "target": target[:400], "status": seg.get("status")})
        if ids:
            projects.append({"id": p["id"], "title": p.get("title", ""), "segments": ids,
                             "violating": violating})
    return {"ok": True, "term": src, "tgt": tgt, "prevTgt": (entry or {}).get("prevTgt", ""),
            "total": sum(len(x["segments"]) for x in projects),
            "violatingTotal": sum(len(x["violating"]) for x in projects),
            "projects": projects, "examples": examples}


_IMPACT_CACHE: dict = {}          # pid -> (отпечаток, отчёт)
_GLOSS_EPOCH = [0]                # растёт на каждой правке глоссария


def _impact_fingerprint(project: dict) -> str:
    """Отпечаток входных данных отчёта: переводы сегментов + поколение глоссария.

    Считаем по содержимому, а не по «версии состояния»: любой будущий код,
    поменявший перевод в обход save_state, иначе получал бы устаревший отчёт,
    а тихо устаревшая цифра «сколько переперевести» хуже её отсутствия."""
    parts = [str(s.get("id")) + "|" + (s.get("status") or "") + "|" + (s.get("target") or "")
             for s in project["segments"]]
    body = chr(10).join(parts)
    return str(_GLOSS_EPOCH[0]) + ":" + hashlib.sha1(body.encode("utf-8")).hexdigest()


class SegmentsFetchRequest(BaseModel):
    ids: List[int]


@app.post("/api/projects/{pid}/segments/fetch")
def fetch_segments(pid: int, req: SegmentsFetchRequest):
    """Отдать конкретные сегменты. Во время прогона клиенту нужны только те,
    что изменились: полный проект на 2670 сегментов весит 5 МБ, и тянуть его
    каждые несколько секунд — это мегабайты трафика и подвисающая таблица."""
    project = get_project(pid)
    wanted = set(req.ids[:1000])
    return {"ok": True, "segments": [_segment_for_client(s) for s in list(project["segments"])
                                     if s.get("id") in wanted]}


# Подбор приказных записей по исходнику — САМЫЙ дорогой расчёт бесплатной
# части: десять тысяч записей глоссария против текста сегмента, 13 мс на
# сегмент. Зовут его отчёт о соответствии, разбор прогона, /analysis, ремонт
# и промпты — каждый своим проходом, и на боевом проекте в 2711 сегментов
# один такой проход занимал 33 секунды ЕДИНСТВЕННОГО воркера: подтверждение
# одного сегмента меняло отпечаток, и экран «Анализ» покупал весь расчёт
# заново, а на это время сервис был недоступен всем.
# Зависит ответ РОВНО от трёх вещей: текста оригинала, области проекта
# и поколения глоссария. Ни перевод, ни статус, ни проверки на него не
# влияют — потому подтверждение сегмента (и любая правка перевода) кэш
# не роняет вовсе.
_HITS_CACHE: dict = {}
_HITS_CACHE_EPOCH = [-1]
# Потолок — про ПАМЯТЬ, а не про терпение: запись держит копии найденных
# записей глоссария. На проект в 2711 сегментов уходит столько же ключей.
_HITS_CACHE_MAX = int(os.environ.get("HITS_CACHE_MAX", "20000"))


def _verified_hits(source: str, project: Optional[dict]) -> list:
    """Записи глоссария уровня ПРИКАЗ, применимые к этому исходнику.

    Одно место на всех, кто спрашивает «чего требует глоссарий от сегмента»:
    расчёт нарушений (`_gloss_misses`), промпт ремонта и корзина спора
    с проверкой. Разойдись они в определении приказной записи — начали бы
    переписывать сегмент друг за другом по кругу (та же беда, от которой
    `_gloss_misses` и `/glossary-impact` считаются одним расчётом).
    Пустой перевод отсеивается здесь же: требовать «используйте ничего»
    нельзя, и записью-ответом такая запись тоже не является.

    Ответ КЭШИРУЕТСЯ по (область, исходник) при неизменном поколении
    глоссария (`_HITS_CACHE`). Наружу уходят КОПИИ записей — ровно так их
    отдаёт и `_get_context`: вызывающие дописывают в найденное свои поля
    (`_form`, пометки ремонта), и общая на всех запись однажды уехала бы
    в чужой промпт."""
    if project is None:
        return []
    # Поколение глоссария сторожит таблицу ЦЕЛИКОМ, а не лежит в ключе:
    # записи прошлого поколения не спросит уже никто, и держать их — значит
    # копить мусор до потолка.
    if _HITS_CACHE_EPOCH[0] != _GLOSS_EPOCH[0]:
        _HITS_CACHE.clear()
        _HITS_CACHE_EPOCH[0] = _GLOSS_EPOCH[0]
    key = (_project_scope(project), source or "")
    got = _HITS_CACHE.get(key)
    if got is None:
        got = [h for h in _get_context(source or "", with_tm=False, project=project)[0]
               if _hit_tier(h) == GLOSSARY_TIER_HARD and (h.get("tgt") or "").strip()]
        # Переполнение чистит таблицу целиком: выбирать, кого выселить,
        # дороже, чем посчитать заново.
        if len(_HITS_CACHE) >= _HITS_CACHE_MAX:
            _HITS_CACHE.clear()
        _HITS_CACHE[key] = got
    return [dict(h) for h in got]


@app.get("/api/projects/{pid}/glossary-impact")
def glossary_impact(pid: int, refresh: bool = False):
    """Сегменты проекта, чей перевод не соответствует ПРОВЕРЕННЫМ записям глоссария.

    После одобрения термина старые переводы сами не меняются — это и есть список
    «что переперевести». Считаем только по verified: автоимпорт модель вправе
    игнорировать, требовать соответствия ему нельзя."""
    project = get_project(pid)
    # Полный проход по проекту — секунды на тысячах сегментов. Пока состояние
    # не менялось, отдаём посчитанное: клиент дёргает отчёт при каждом открытии
    # редактора и после каждого прогона.
    fp = _impact_fingerprint(project)
    cached = _IMPACT_CACHE.get(pid)
    if cached and cached[0] == fp and not refresh:
        return cached[1]
    by_term: dict = {}
    seg_ids, confirmed_ids, case_ids, term_ids = set(), set(), set(), set()
    for seg in project["segments"]:
        target = (seg.get("target") or "").strip()
        if not target:
            continue
        hits = _verified_hits(seg.get("source", ""), project)
        # Регистр приказных терминов считаем ЗДЕСЬ же, по тем же самым hits:
        # `_verified_hits` — самый дорогой вызов в этом проходе (13 мс на
        # сегмент), и звать его второй раз ради того же ответа значит
        # удвоить полминуты работы единственного воркера.
        if _term_case_hits(seg, hits):
            case_ids.add(seg["id"])
        # И список «где вообще есть что сверять» — по тем же hits и тем же
        # проходом. Шагу сверки терминов нужен именно он, а второй проход
        # `_verified_hits` по проекту стоит 40 секунд единственного воркера:
        # разбор прогона пересчитывается на каждую смену модели и галочки.
        if hits:
            term_ids.add(seg["id"])
        for h in hits:
            if _tgt_has_term(target, h["tgt"]):
                continue
            key = h["src"]
            e = by_term.setdefault(key, {"src": key, "tgt": h["tgt"], "cat": h.get("cat", ""),
                                         "updated": h.get("updated", ""), "prevTgt": h.get("prevTgt", ""),
                                         "segments": [], "confirmed": []})
            e["segments"].append(seg["id"])
            seg_ids.add(seg["id"])
            if seg.get("status") == "confirmed":
                e["confirmed"].append(seg["id"])
                confirmed_ids.add(seg["id"])
    terms = sorted(by_term.values(), key=lambda t: len(t["segments"]), reverse=True)
    result = {"ok": True, "terms": terms,
              "segments": sorted(seg_ids), "confirmed": sorted(confirmed_ids),
              "pending": sorted(seg_ids - confirmed_ids),
              # Отдельным списком, а не вперемешку с `segments`: там «термина
              # в переводе нет вовсе», здесь «термин есть, но не в том
              # начертании». Смешай их — и карточка соответствия начала бы
              # звать переводить заново то, где надо поправить одну букву.
              "caseSegments": sorted(case_ids),
              # Сегменты, к которым применим хоть один приказной термин.
              # Не «расходятся», а «есть что сверить»: по нему считается состав
              # шага termaudit.
              "termSegments": sorted(term_ids)}
    # Кого ремонт уже не возьмёт: тот же текст с теми же претензиями он
    # проходил, и второй заход вернёт то же самое. Считаем ЗДЕСЬ, чтобы
    # карточка и состав прогона брали одно число из одного расчёта.
    by_id = {sg["id"]: sg for sg in project["segments"]}
    result["futile"] = sorted(i for i in result["segments"]
                              if i in by_id and _repair_futile(by_id[i], project))
    _IMPACT_CACHE[pid] = (fp, result)
    return result


_ANALYSIS_CACHE: dict = {}


# ─── Покрытие проверок для пары и области проекта ────────────────────
# Пара проекта может быть любой, и закон у детерминированных проверок один:
# нет правил для этой пары — МОЛЧИМ. Но молчание неотличимо от успеха, пока
# о нём не сказано: человек видит «находок нет» и считает текст проверенным.
# Поэтому здесь по ТЕМ ЖЕ таблицам, из которых проверки берут правила,
# считается три списка — что работает, что молчит и почему, что покупается
# вызовом модели (и потому работает на любой паре). Отдельный список
# «поддерживаемых языков» разошёлся бы с таблицами первой же правкой.
def _coverage(project: dict) -> dict:
    src = (project.get("src") or "").upper()
    tgt = (project.get("tgt") or "").upper()
    dom = _resolve_domain(project.get("domain"))
    works, silent, model = [], [], []
    # Языконезависимые — работают всегда.
    works += [
        {"key": "numbers", "label": "Числа и единицы измерения"},
        {"key": "glossary", "label": "Соответствие утверждённым терминам (точное совпадение)"},
        {"key": "case", "label": "Регистр букв"},
        {"key": "script", "label": "Буквы чужого письма в переводе"},
        {"key": "dup", "label": "Самоповтор текста"},
        {"key": "consist", "label": "Единство терминологии по документу"},
    ]
    # Морфологический подбор терминов — по таблице окончаний языка оригинала.
    if src in _LANG_ENDINGS:
        works.append({"key": "morph", "label": "Термины в косвенных формах (морфология %s)" % src})
    else:
        silent.append({"key": "morph", "label": "Термины в косвенных формах",
                       "why": "нет таблицы окончаний для %s — термин находится только в словарной форме" % src})
    # Правила области и пары (лево/право, стиль, скип-списки).
    rules = (checks_mod.rules_for(dom["id"], src, tgt) if checks_mod
             else {"pairs": [], "style": []})
    if rules.get("pairs") or rules.get("style"):
        works.append({"key": "domain_rules", "label": "Правила области «%s» для %s→%s (подмена стороны, стиль)" % (dom["label"], src, tgt)})
    else:
        silent.append({"key": "domain_rules", "label": "Правила области «%s»" % dom["label"],
                       "why": "для пары %s→%s правила не описаны — проверка стороны и стиля не срабатывает" % (src, tgt)})
    # Инверсия отрицания — по маркерам языка оригинала (сравниваются оригинал и обратный перевод).
    neg = checks_mod.negation_markers(src) if checks_mod else []
    if neg:
        works.append({"key": "negation", "label": "Инверсия отрицания"})
    else:
        silent.append({"key": "negation", "label": "Инверсия отрицания",
                       "why": "нет маркеров отрицания для %s" % src})
    # Балл back-check по словам — письмо без пробелов мерить нечем.
    info = _LANG_BY_CODE.get(src) or {}
    if info.get("script") in ("HAN", "THAI", "KHMER", "MYANMAR"):
        silent.append({"key": "recall", "label": "Доля слов оригинала, переживших обратный перевод",
                       "why": "письмо без пробелов между словами — балл не измеряется, решает судья"})
    else:
        works.append({"key": "recall", "label": "Доля слов оригинала, переживших обратный перевод"})
    # Платные — через модель, языка не требуют.
    model += [
        {"key": "backcheck", "label": "Обратный перевод (back-check)"},
        {"key": "termcheck", "label": "Проверка терминологии моделью"},
        {"key": "judge", "label": "Судья смысла"},
        {"key": "termaudit", "label": "Сверка терминов в контексте"},
        {"key": "meaning", "label": "Сверка смысла кандидатов глоссария"},
    ]
    return {"ok": True, "src": src, "tgt": tgt, "domain": dom["id"],
            "works": works, "silent": silent, "model": model}


@app.get("/api/projects/{pid}/coverage")
def project_coverage(pid: int):
    """Что проверяется на паре и области ЭТОГО проекта, что молчит и почему."""
    return _coverage(get_project(pid))


def _analysis_seg_fp(s: dict) -> str:
    """Отпечаток СОДЕРЖИМОГО сегмента для экрана «Анализ».

    Перечня «что влияет на экран» здесь намеренно НЕТ. Раньше он был и рос
    с каждой проверкой: сперва хеши текста, потом баллы, потом судья, потом
    Medical QA и арбитр, — а забытое поле означало доперегонные цифры
    на экране, то есть враньё без единого видимого признака. Считать
    содержимое целиком стоит 30 мс на 2711 сегментов, промахнуться им нельзя
    по построению, и тем же отпечатком помечена готовая строка разбора
    (`_ANALYSIS_ROWS`) — одно определение на кэш и на сверку свежести.

    Полей ровно столько, сколько читает `_analysis_row` со своими вызовами
    (`_repair_findings`, `_machine_clean`, `_repair_clamped`,
    `_confirm_override`, `_judge_pending`, `_repair_score_vetoed`). Начал
    читать новое поле — впиши его СЮДА, иначе строка не пересчитается."""
    return hashlib.sha1(repr([
        s.get("source"), s.get("target"), s.get("status"), s.get("backcheck"),
        s.get("termcheck"), s.get("repair"), s.get("qa_result"),
        s.get("qa_issues"), s.get("risk_color"), s.get("termContext"),
        s.get("confirmWithdrawn"), s.get("review")]).encode("utf-8")).hexdigest()


# Готовые строки разбора: {pid: {id сегмента: (ключ, строка)}}. Ключ — тройка
# «отпечаток сегмента + расходится ли он с глоссарием + порог балла»: ровно
# то, от чего строка зависит. Подтверждение одного сегмента меняет ОДИН ключ,
# значит и считается один сегмент, а не весь проект. Без этого пересчёт стоил
# 10 секунд единственного воркера на боевом проекте — и человек, нажавший
# «Подтвердить», ждал экран «Анализ» дольше, чем правил сам текст.
_ANALYSIS_ROWS: dict = {}


def _analysis_row(s: dict, gloss_bad: bool, min_score: int) -> dict:
    """Разбор ОДНОГО сегмента для экрана «Анализ»: в какие корзины он идёт.

    Вынесено из общего прохода ради кэша по содержимому (`_ANALYSIS_ROWS`),
    и раскладка по спискам осталась у вызывающего: строка описывает сегмент,
    а не его место среди других, — только поэтому её и можно переиспользовать.
    Всё, от чего строка зависит, обязано входить в ключ: содержимое сегмента
    (`_analysis_seg_fp`), расхождение с глоссарием и порог балла области."""
    row = {"untranslated": False, "withdrawn": False, "withdrawnOpen": False,
           "unjudgedBlind": False, "unverified": False, "judgeExt": False,
           "confirmed": False, "qaCritical": False, "repaired": False,
           "sourceSuspect": False, "reviewFlagged": False, "reviewVouched": False,
           "reverted": False, "scoreVetoed": False, "findings": False,
           "clamped": False, "confirmedFindings": False, "override": False,
           "bucket": None, "why": ""}
    target = (s.get("target") or "").strip()
    if not target:
        row["untranslated"] = True
        return row
    if s.get("confirmWithdrawn"):
        row["withdrawn"] = True
        if s.get("status") != "confirmed":
            row["withdrawnOpen"] = True
    _bc = s.get("backcheck") or {}
    # Свежесть считаем ОДИН раз: это sha1 перевода, а ниже её спрашивают
    # обе ветки — и «никто не проверял», и расширенная зона судьи.
    _bc_fresh = bool(_bc) and not _check_stale(_bc, target)
    if _bc_fresh and not _bc.get("judged"):
        _sc = _bc.get("score")
        if _lex_blind(s.get("source") or ""):
            row["unjudgedBlind"] = True
        elif _sc is not None and _sc > JUDGE_ZONE[1]:
            row["unverified"] = True
    # Расширенная зона (judge_all=True): прогон с разрешением спросит
    # судью и здесь. Свежесть проверяется отдельно — _judge_pending про
    # неё не знает, а вердикт по прежнему тексту недостачей не считается:
    # сам перевод сначала перепроверит back-check (корзина unchecked).
    if _bc_fresh and _judge_pending(s, above=True):
        row["judgeExt"] = True
    if s.get("status") == "confirmed":
        row["confirmed"] = True
        _qa = s.get("qa_result") or {}
        _hardqa = any((i.get("severity") or i.get("sev")) in
                      ("critical", "high", "major")
                      for i in (s.get("qa_issues") or ()))
        # Только свежий результат: устаревший пересчитает шаг Medical QA
        # ближайшего прогона, и находка либо подтвердится, либо снимется.
        if (_hardqa or s.get("risk_color") == "red") and _qa \
                and not _check_stale(_qa, target):
            row["qaCritical"] = True
    # Ревизор усомнился в САМОМ ОРИГИНАЛЕ (обрывок, ошибка распознавания,
    # бессвязная фраза). Машине тут делать нечего по построению: чинить
    # перевод догадкой по битому исходнику — значит сочинять. Свежесть
    # проверяем как у всех: вердикт мог описывать текст, которого уже нет.
    # Свежесть — тем же предикатом, что у прогона (`_review_stale`): он знает
    # и про версию вопросов, и про правку ОРИГИНАЛА. `_check_stale` смотрит
    # только на перевод, и корзина «повреждён оригинал» не осушалась бы даже
    # после того, как исходник выправили.
    _rv = s.get("review") or {}
    if _rv.get("sourceSuspect") and not _review_stale(s):
        row["sourceSuspect"] = True
    # Ревизия НАШЛА проблему, а текст остался прежним. Самый ценный сигнал
    # для человека из всего, что она даёт: модель прочитала пару целиком
    # и считает перевод дефектным, но машина чинить не стала.
    #
    # Читаем ЗАПИСАННОЕ решение (`code`), а не пересчитываем условия. Пересчёт
    # не знает, что решил прогон, и врал дважды: сегмент с готовым кандидатом
    # после СУХОГО прогона (сверку прошёл, ждёт `apply_saved`) показывался как
    # «сверка не пустила», а смена `REVIEW_APPLY_MAX` задним числом
    # переклассифицировала записи, у которых вето вообще не считалось.
    #
    # Что входит:
    #   • `veto` — кандидат не прошёл объективные сверки: чинить его машине
    #     нечем, следующий заход даст то же самое;
    #   • `ok` при НИЗКОЙ оценке — модель считает перевод плохим, но варианта
    #     не дала: машине тут делать нечего вовсе.
    # Чего нет: `suspect` — своя корзина (вопрос к исходнику, а не к переводу);
    # `above` — «можно улучшить», а не «сломано», звать человека к вкусовщине
    # нельзя; `undone` — он уже ответил; вердикт БЕЗ кода — это ждущий
    # применения кандидат (сухой прогон), и решает его `apply_saved`, машина.
    elif _rv and not _rv.get("applied") and not _rv.get("undone")             and not _review_stale(s):
        _code = _review_code(_rv)
        if _code == REVIEW_VETOED or (
                _code == REVIEW_OK and _rv.get("score") is not None
                and _rv["score"] <= REVIEW_FLAG_SCORE):
            row["reviewFlagged"] = True
    rp = s.get("repair") or {}
    # Настоящие находки держим отдельно от `open_findings`: у второго
    # бывает фолбэк-заглушка `[{"kind": "gloss"}]` без ключа `text`,
    # и отпечатку захода её отдавать нельзя. А сам список нужен ниже
    # `_repair_clamped`, чтобы тот не считал то же самое второй раз:
    # на боевом проекте это лишние полсекунды единственного воркера
    # при каждом открытии экрана.
    rp_findings = _repair_findings(s)
    open_findings = rp_findings or ([{"kind": "gloss"}] if gloss_bad else [])
    if rp.get("applied") and _repair_tried(s):
        row["repaired"] = True
    elif rp and not rp.get("applied") and open_findings and _repair_tried(s):
        # Модель пробовала починить и не смогла — дальше только человек.
        # _repair_tried обязателен: запись о неудачной правке могла остаться
        # от прежнего текста, а к нынешнему уже не относится.
        row["reverted"] = True
        # Отдельной строкой — те, где отмена держалась ТОЛЬКО на упавшем
        # балле, а термины правка почистила. Это не «модель не смогла»:
        # текст написан и оплачен, лежит в repair.candidate и ждёт одного
        # нажатия. В общей корзине они выглядят безнадёжными и потому
        # не разбираются никогда.
        if _repair_score_vetoed(s):
            row["scoreVetoed"] = True
    # Только неподтверждённые: подтверждённые с находками пакетный ремонт
    # без явного разрешения не трогает, и обещать «это починится само»
    # было бы неправдой. Они уходят в свою корзину — не потому, что с ними
    # нечего делать, а потому, что решение принимает человек.
    if open_findings and s.get("status") != "confirmed":
        row["findings"] = True
        # Тот же предикат, что у _plan_step с retry=False: совпавший
        # отпечаток захода прогон не берёт — это работа человека.
        if _repair_clamped(s, rp_findings):
            row["clamped"] = True
    elif open_findings:
        row["confirmedFindings"] = True
        # Объективная находка сильнее заверения: таких ремонт берёт сам.
        if _confirm_override(s):
            row["override"] = True
    # Корзины обязаны быть исчерпывающими: сегмент, не попавший ни в одну,
    # исчезает с экрана, и картина выглядит благополучнее, чем есть.
    why = _machine_clean(s, min_score)
    if why is None:
        row["bucket"] = "clean"
    elif why == CLEAN_TERMCHECK_SKIP:
        # НЕ «не проверено»: `_termcheck_trivial` сам сказал, что проверять
        # нечего — в переводе нет слов либо он совпадает с оригиналом
        # («40%», «kV 120.0 mA: 283»). Такой сегмент из корзины «не
        # проверено» не уйдёт НИКОГДА: текст не изменится, и проверять
        # в нём по-прежнему нечего. На боевом проекте так висели 70 из 138,
        # то есть половина корзины звала человека к работе, которой нет.
        row["bucket"] = "nothing"
    elif why in CLEAN_UNCHECKED:
        row["bucket"] = "unchecked"
    elif why == CLEAN_REPAIRED and row["repaired"]:
        # НЕ в «оценку ниже порога». У такого сегмента back-check прошёл
        # и termcheck чист (глоссарий _machine_clean не смотрит вовсе —
        # расхождение с ним видно отдельной строкой), а отказ он получил
        # только за то, что систему нельзя пускать заверять собственную
        # правку. На боевом проекте это 306 сегментов из 511: корзина на
        # 60% состояла из благополучных строк и звала человека разбираться
        # там, где разбираться не в чем. Своя строка у них есть —
        # «Исправила машина», и ждут они там ровно того, что им нужно:
        # подтверждения человеком. Условие `row["repaired"]` — не
        # перестраховка: без него выброс держался бы на том, что два
        # предиката в разных концах файла совпадают буква в букву.
        pass
    elif not row["findings"] and not row["confirmedFindings"]:
        # Балл ниже порога у КОРОТКОГО оригинала — это не «оценка ниже
        # порога», а «оценки нет»: мерить было нечем, и пока судья не
        # ответил, число ничего не значит. Своей причиной, чтобы человек
        # не шёл разбираться с верным переводом.
        if (why.startswith("back-check ниже")
                and not _bc.get("judged")
                and _lex_blind(s.get("source") or "")):
            why = CLEAN_LEX_BLIND
        if (why.startswith("back-check ниже")
                and (_bc.get("judge") or {}).get("severity") in ("major", "critical")
                and _review_wrote(s)):
            why = CLEAN_JUDGE_VS_REVIEW
        # Ручательство ревизии. Она единственная читает ПАРУ целиком, а балл
        # back-check меряет долю основ оригинала, вернувшихся через обратный
        # перевод, — вознаграждает кальку и роняет верный синоним. Свежая
        # оценка ≥ REVIEW_VOUCH_SCORE на том же тексте — прямое чтение против
        # косвенной меры, и прав ревизор: тот же закон, по которому вердикт
        # арбитра снимает `term_lost`. Чего ручательство не переживает,
        # решает сам `_review_vouches` (объективная находка, вердикт судьи);
        # детерминированные претензии (глоссарий, регистр, письмо, повтор)
        # сюда не доходят — с ними сегмент ушёл в findings выше. Донором
        # глоссария он не становится: `_machine_clean` ревизию не спрашивает,
        # «можно ли учить» и «нужен ли человек» — разные вопросы.
        if _review_vouches(s):
            row["reviewVouched"] = True
            row["bucket"] = "vouched"
        else:
            row["bucket"] = "weak"
            row["why"] = why
    return row


@app.get("/api/projects/{pid}/analysis")
def project_analysis(pid: int, refresh: bool = False):
    """Итог работы по проекту одним экраном: что чисто, что исправила машина,
    что она предлагает одобрить и что осталось человеку.

    Считается по состоянию, а не по последнему прогону: прогонов может быть
    несколько, страницу могли перезагрузить, а вопрос у пользователя один —
    «что сейчас с проектом». Ни одного вызова модели здесь нет — но проход
    по проекту не бесплатен: на 2670 сегментах это секунды CPU единственного
    воркера, а экран открывают часто. Поэтому результат кэшируется по
    отпечатку тех же данных, из которых считается."""
    project = get_project(pid)
    q = _term_queue()
    # Отпечаток считается ПО СОДЕРЖИМОМУ сегментов (`_analysis_seg_fp`),
    # тем же, каким помечены готовые строки разбора. Ручного перечня полей
    # здесь больше нет: он рос с каждой проверкой (хеши текста, баллы, судья,
    # Medical QA, арбитр), а забытое поле означало доперегонные цифры
    # на экране — враньё без единого видимого признака.
    seg_fp = {s["id"]: _analysis_seg_fp(s) for s in project["segments"]}
    fp = (_impact_fingerprint(project) + "|" + str(len(q)) + "|"
          + str(sum(1 for c in q if c.get("status", "pending") == "pending"))
          + "|" + hashlib.sha1("".join(
              "%s:%s;" % (i, f) for i, f in seg_fp.items()).encode("utf-8")).hexdigest())
    cached = _ANALYSIS_CACHE.get(pid)
    if cached and cached[0] == fp and not refresh:
        return cached[1]
    pol = _auto_policy(project.get("domain"))
    scope = _project_scope(project)
    impact = glossary_impact(pid)

    clean, repaired, reverted, untranslated, unchecked = [], [], [], [], []
    findings, weak = [], []
    # Претензии слепых измерителей снял свежий вердикт ревизии — «готово»,
    # но названное числом (`turnkey.reviewVouched`): снятое молча
    # неотличимо от потерянного.
    review_vouched: list = []
    # Подтверждённые с находками — своя корзина. Раньше они растворялись
    # в «оценка ниже порога» вперемешку с машинными сегментами, и по экрану
    # нельзя было понять, что это ручные подтверждения, до которых ни один
    # прогон не дотянется без явного разрешения.
    confirmed_findings: list = []
    conf_set: set = set()
    # Заверение, снятое МАШИНОЙ по объективной находке. Своя строка обязательна:
    # человек поставил отметку, машина её отменила — он должен об этом узнать
    # и увидеть доказательство, а не обнаружить пропажу случайно.
    withdrawn: list = []
    # Сегменты, которые НИКТО не проверял по существу. Две разные причины,
    # и обе нельзя показывать как «чисто»:
    #   • балл выше потолка зоны судьи — детерминированные проверки довольны,
    #     а спросить того, кто читает смысл, правило не даёт. Именно здесь
    #     прячется «беглое неверное слово»: monostable, cusps, actinoid —
    #     нормальные английские слова не из той области, и балл у них 100;
    #   • балл не измеряется (короткий оригинал), а судья не смотрел —
    #     система честно сказала «мерить нечем» и на этом остановилась.
    # Показываются ЧИСЛОМ и лечатся одной кнопкой (судья), а не разбором
    # руками: 845 сегментов человеку не пересмотреть, и не в этом смысл.
    unverified: list = []
    unjudged_blind: list = []
    # Сегменты, где проверять нечего по существу. Своя корзина, а не «не
    # проверено»: работой это не является ни сейчас, ни когда-либо.
    nothing_to_check: list = []
    # Подмножество `reverted`: правку отменил только балл back-check, а термины
    # она почистила. Подмножество, а не отдельная корзина, — иначе исчерпаемость
    # держалась бы на совпадении двух предикатов в разных концах файла.
    score_vetoed: list = []
    # Множество, а не повторение условия: ниже переписанные ремонтом сегменты
    # НЕ идут в «оценку ниже порога», и правомерно это ровно потому, что они
    # уже попали в `repaired`. Два одинаковых предиката в тысяче строк друг от
    # друга однажды разойдутся, и тогда сегмент исчезнет с экрана совсем —
    # худшая из здешних ошибок. Со множеством он в худшем случае окажется
    # в «оценке ниже порога»: видно и слегка неверно вместо не видно и совсем.
    repaired_set: set = set()
    # ── Сырьё для трёх корзин «под ключ» (см. сборку ниже, turnkey) ──────
    # Судья ещё не смотрел, считая по расширенной зоне (judge_all): эту
    # недостачу закрывает прогон с разрешением, а не человек.
    judge_ext: set = set()
    # Находки есть, но ремонт с retry=False сюда не пойдёт (совпавший
    # отпечаток захода) — дальше только человек или другая модель.
    clamped_ids: set = set()
    # Подтверждённые с ОБЪЕКТИВНОЙ находкой: ремонт возьмёт их и без галочки
    # (_confirm_override), то есть это работа машины, а не вопрос человеку.
    override_ids: set = set()
    # Снятое машиной заверение, которое человек ещё не пересмотрел
    # (статус так и не confirmed) — вопрос к человеку, пока он не закрыт.
    withdrawn_open: set = set()
    # Подтверждённые с критической находкой Medical QA. Единственное место,
    # где они были видны, — вкладка «Замечания»; корзины /analysis qa_issues
    # не читали вовсе, и удаление вкладки спрятало бы их отовсюду.
    qa_critical: list = []
    source_suspect: list = []
    review_flagged: list = []
    # Заверенные человеком — нужны корзинам: без явного разрешения прогон их
    # не переписывает, значит обещать «машина доделает» про них нельзя.
    confirmed_ids: set = set()
    # Расхождения с глоссарием берём из отчёта, а не считаем заново: там тот же
    # расчёт на весь проект и он кэширован. Вызов _repair_findings с project
    # гонял бы _get_context на каждый сегмент — 10 секунд CPU единственного
    # воркера на 2670 сегментах, и это при каждом открытии экрана.
    # Сюда же расхождения по НАЧЕРТАНИЮ приказного термина: они тоже считаются
    # по глоссарию, а `_repair_findings(s)` ниже зовётся без проекта.
    gloss_bad = set(impact["segments"]) | set(impact.get("caseSegments") or ())
    # Проход по сегментам: сам разбор лежит в `_analysis_row`, здесь — только
    # раскладка по спискам. Готовая строка переиспользуется, пока не менялись
    # ни сегмент, ни его расхождение с глоссарием, ни порог области: экран
    # открывают после каждой правки, и пересчитывать из-за одного
    # подтверждения весь проект — это те самые секунды единственного воркера,
    # на которые сервис недоступен всем.
    rows_prev = _ANALYSIS_ROWS.get(pid) or {}
    rows_new: dict = {}
    for s in project["segments"]:
        sid = s["id"]
        key = (seg_fp.get(sid), sid in gloss_bad, pol["backcheck_min"])
        got = rows_prev.get(sid)
        row = got[1] if got and got[0] == key else _analysis_row(
            s, sid in gloss_bad, pol["backcheck_min"])
        rows_new[sid] = (key, row)
        if row["untranslated"]:
            untranslated.append(sid)
            continue
        if row["withdrawn"]:
            withdrawn.append(sid)
            if row["withdrawnOpen"]:
                withdrawn_open.add(sid)
        if row["unjudgedBlind"]:
            unjudged_blind.append(sid)
        if row["unverified"]:
            unverified.append(sid)
        if row["judgeExt"]:
            judge_ext.add(sid)
        if row["confirmed"]:
            confirmed_ids.add(sid)
        if row["qaCritical"]:
            qa_critical.append(sid)
        if row["sourceSuspect"]:
            source_suspect.append(sid)
        if row["reviewFlagged"]:
            review_flagged.append(sid)
        if row["repaired"]:
            repaired.append(sid)
            repaired_set.add(sid)
        if row["reverted"]:
            reverted.append(sid)
        if row["scoreVetoed"]:
            score_vetoed.append(sid)
        if row["findings"]:
            findings.append(sid)
        if row["clamped"]:
            clamped_ids.add(sid)
        if row["confirmedFindings"]:
            confirmed_findings.append(sid)
            conf_set.add(sid)
        if row["override"]:
            override_ids.add(sid)
        if row["bucket"] == "clean":
            clean.append(sid)
        elif row["bucket"] == "nothing":
            nothing_to_check.append(sid)
        elif row["bucket"] == "unchecked":
            unchecked.append(sid)
        elif row["bucket"] == "weak":
            weak.append({"id": sid, "why": row["why"]})
        elif row["bucket"] == "vouched":
            review_vouched.append(sid)
    _ANALYSIS_ROWS[pid] = rows_new

    # ── Termcheck спорит с утверждённым термином ─────────────────────
    # Находка указывает на слово, которое И ЕСТЬ verified-перевод для этого
    # сегмента. Машине тут делать нечего: ремонт по такой находке всегда
    # откатится — _repair_scores считает нарушенные утверждённые термины, и
    # замена приказного термина поднимает счётчик. Решает человек, и решение
    # у него ровно одно из двух: неверна запись глоссария либо неверна
    # проверка. Без этой корзины спор не виден нигде: находка тонет в «с
    # замечаниями», а запись выглядит соблюдённой.
    # Считаем только по сегментам с находками — _get_context на каждом
    # сегменте проекта это секунды CPU единственного воркера.
    disputes: dict = {}
    for s in project["segments"]:
        tc = s.get("termcheck") or {}
        if not tc or _check_stale(tc, s.get("target") or ""):
            continue
        bad = [f for f in (tc.get("findings") or [])
               if f.get("severity") in TERMCHECK_DISPUTING and f.get("tgt_term")]
        if not bad:
            continue
        approved = {_norm_key(h["tgt"]): h for h in _verified_hits(s.get("source", ""), project)}
        for f in bad:
            h = approved.get(_norm_key(f.get("tgt_term")))
            if h is None:
                continue
            d = disputes.setdefault((_norm_key(h["src"]), _norm_key(h["tgt"])),
                                    {"src": h["src"], "tgt": h["tgt"],
                                     "suggests": [], "segments": []})
            sug = (f.get("suggestion") or "").strip()
            if sug and sug not in d["suggests"] and len(d["suggests"]) < 5:
                d["suggests"].append(sug)
            # Сегмент посещается ровно раз, поэтому повтором может быть только
            # предыдущая находка того же сегмента — линейный поиск по списку
            # на 2670 сегментов с одним спорным термином был бы O(n²).
            if not d["segments"] or d["segments"][-1] != s["id"]:
                d["segments"].append(s["id"])
    disputed = sorted(disputes.values(), key=lambda d: -len(d["segments"]))

    # ── Контекстный арбитр: кого ещё не спрашивали и что он ответил ──────
    # Спор виден и без него, но без довода человек читает только «проверка
    # против глоссария» и решить не может. Арбитр — единственный, кто смотрит
    # на сегмент в ряду соседей; его «передан верно» уже сняло претензию
    # с ремонта, а «передан неверно» здесь и показывается: вопрос про ЗАПИСЬ
    # глоссария, поэтому список — по записям, а не по сегментам.
    # Считаем по ВСЕМ вердиктам, а не только по спорным сегментам. Шаг сверки
    # спрашивает про все приказные термины: на боевом проекте это 713 сегментов
    # против 178 спорных, то есть при отборе по спору три четверти оплаченных
    # ответов не читал бы никто — в ремонт «передан неверно» не идёт намеренно,
    # а больше его показывать негде. Корзины обязаны быть исчерпывающими.
    ctx_pending, ctx_bad = 0, {}
    for s in project["segments"]:
        if not (s.get("target") or "").strip():
            continue
        seg_ctx = s.get("termContext") or {}
        fresh = bool(seg_ctx) and not _term_context_stale(s)
        # «Ждут арбитра» и «что он ответил» — РАЗНЫЕ вопросы, и считать их
        # одним условием нельзя. Вердикт разбора спора — настоящий ответ про
        # настоящий термин, и показывать его надо; но сегмент при этом всё
        # равно ждёт полной сверки, потому что неспорные термины в нём никто
        # не спрашивал.
        if not fresh or not seg_ctx.get("all_terms"):
            if _term_terms_of(s, project, disputes_only=False):
                ctx_pending += 1
        if not fresh:
            continue
        for t in _term_context_of(s):
            # Вердикты по забракованным СЛОВАМ — не спор про запись глоссария:
            # «негодно» уже стало находкой ремонта, «годно» сняло претензию.
            if t.get("stale"):
                continue
            if t.get("ok") is False:
                # Предъявить человеку нечего в двух случаях, и оба — шум:
                # совет ПУСТ (корпус его снял или арбитр не предложил) либо
                # совет СОВПАДАЕТ с приказной записью — тогда это не спор
                # про запись, а «здесь термин передан неверно», и это уже
                # находка ремонта (`kind: "term_ctx"`), а не вопрос к человеку.
                # На боевом проекте так набралось 8 карточек из 10, и строка
                # звала разбираться там, где разбираться не в чем.
                use = (t.get("use") or "").strip()
                if not use or _norm_key(use) == _norm_key(t.get("tgt") or ""):
                    continue
                k = (_norm_key(t.get("src")), _norm_key(t.get("tgt")))
                e = ctx_bad.setdefault(k, {"src": t.get("src"), "tgt": t.get("tgt"),
                                           "use": t.get("use"), "why": t.get("why"),
                                           "segments": []})
                if not e["segments"] or e["segments"][-1] != s["id"]:
                    e["segments"].append(s["id"])
    ctx_wrong = sorted(ctx_bad.values(), key=lambda d: -len(d["segments"]))

    # Забракованное termcheck слово, которое ВСЁ ЕЩЁ стоит в тексте.
    # Карточка очереди помнит `wasTgt` — формулировку, которую проверка
    # отвергла, — а сегмент об этом не помнит: свежий termcheck мог передумать
    # между прогонами, и тогда дефект остаётся в готовом на вид тексте.
    # На боевом проекте так вышло у 83 сегментов, 59 из которых числились
    # готовыми: «bioptate» вместо «biopsy specimen», «nodules» вместо
    # «papules», «hormone therapy» вместо «corticosteroid therapy».
    # Сверка бесплатна — подстрока по границам слов.
    #
    # Это НОМИНАЦИЯ, а не находка ремонта, и намеренно: одно суждение termcheck
    # о строке — не приказ переписывать (тот же закон, что у разнобоя, где
    # для применения по документу нужны два голоса). Тем более что проверка
    # с тех пор передумала, и кто из двух её мнений верен, решает человек.
    # Один расчёт с шагом сверки (_stale_words_of) — смета и работа не должны
    # видеть разные списки. Человеку остаются только слова БЕЗ свежего вердикта
    # арбитра: «годно» претензию сняло, «негодно» стало находкой ремонта
    # (сегмент уходит в findings, то есть в машинную корзину) — второй голос
    # получен, звать сюда человека больше не за чем.
    by_id = {sg["id"]: sg for sg in project["segments"]}
    stale_bad: dict = {}
    for sid, words in _stale_words_of(project).items():
        # Тем же _stale_unasked, что смета шага и эндпоинт: вторая формула
        # «какие слова уже отвечены» разошлась бы с ними первой же правкой.
        left = _stale_unasked(by_id[sid], words)
        if left:
            stale_bad[sid] = left
    stale_findings = sorted(stale_bad)

    # Разнобой по документу — один раз: он нужен и корзинам «под ключ»,
    # и списку todo.consistency, а считается кэшированным проходом.
    consist_pairs = _consistency_of(project)

    # ── Три корзины «под ключ»: готово / возьмёт прогон / нужен человек ──
    # Пользователю, которому нужен перевод под ключ, экран обязан отвечать
    # на один вопрос тремя числами, и числа обязаны сходиться с total:
    # сегмент, не попавший ни в одну корзину, исчез бы с экрана.
    #
    # Правила раздачи:
    #   • «нужен человек» — то, что прогон НЕ решает ПО ПОСТРОЕНИЮ: заверено
    #     человеком и без объективной находки (ремонт не тронет), спор проверки
    #     с приказной записью, вердикт арбитра о записи, забракованное слово
    #     в тексте, снятое заверение, откаченные и заклеймлённые ремонтом,
    #     «ремонт уже не берёт» (futile), слабый балл после судьи, критика
    #     Medical QA на подтверждённом. Приоритет у этой корзины: если сегмент
    #     прогон и возьмёт (освежить проверку), находку это не закроет,
    #     и обещание «машина доделает» было бы враньём.
    #   • «возьмёт прогон» — то, что берут шаги _plan_step с параметрами
    #     кнопки (use_judge=True, judge_all=True, retry=False): не переведено,
    #     не проверено, находки для ремонта, расхождения с глоссарием,
    #     разнобой, судья (в зоне, ниже зоны у коротких, выше зоны
    #     по judge_all) — плюс начертание, которое чинит бесплатная команда
    #     той же кнопки.
    #   • «готово» — остаток. Остатком, а не повторением предикатов clean/
    #     repaired/nothingToCheck: разойдись копии, сегмент исчез бы с экрана
    #     совсем, а так он в худшем случае виден в не той корзине.
    human_set: set = set()
    human_set.update(i for i in confirmed_findings if i not in override_ids)
    human_set.update(impact["confirmed"])
    human_set.update(qa_critical)
    human_set.update(source_suspect)
    human_set.update(review_flagged)
    human_set.update(i for d in disputed for i in d["segments"])
    human_set.update(stale_findings)
    human_set.update(withdrawn_open)
    human_set.update(i for d in ctx_wrong for i in d["segments"])
    human_set.update(reverted)
    human_set.update(impact.get("futile") or ())
    human_set.update(clamped_ids)
    # Слабый балл, который судья уже не поднимет (смотрел либо не позовут), —
    # читать глазами. Тот, куда судья ещё придёт, — работа прогона.
    human_set.update(w["id"] for w in weak if w["id"] not in judge_ext)

    # Разнобой у ЗАВЕРЕННОГО сегмента — работа человека, а не прогона.
    # Ловится это только здесь: `_repair_findings` в /analysis зовётся БЕЗ
    # проекта (ради скорости), а `_consist_misses` без проекта возвращает
    # пусто — то есть в `confirmed_findings` такой сегмент не попадает и сам
    # в human_set не уходит. Прогон же идёт с include_confirmed=False и его
    # пропускает: оставь его машине — и корзина «возьмёт прогон» держала бы
    # число, которое не осушится никогда. Исключение — объективная находка
    # (`override_ids`): такие ремонт берёт и без разрешения.
    _consist_ids = {i for p in consist_pairs for i in p["segments"]}
    _consist_human = (_consist_ids & confirmed_ids) - override_ids
    human_set.update(_consist_human)

    machine_set: set = set(untranslated)
    machine_set.update(unchecked)
    machine_set.update(findings)
    machine_set.update(override_ids)
    machine_set.update(impact["pending"])
    machine_set.update(_consist_ids)
    # caseSegments заверенных сюда попадают, но human_set их уже забрал:
    # фолбэк `[{"kind": "gloss"}]` по `gloss_bad` заводит их в
    # `confirmed_findings` выше. Вычитание ниже и есть страховка от того,
    # что это когда-нибудь перестанет быть правдой.
    machine_set.update(impact.get("caseSegments") or ())
    machine_set.update(judge_ext)
    machine_set -= human_set
    # Заверение человека — сильнейший сигнал системы (инвариант про
    # confirmedBy), и корзины обязаны его видеть: раньше подтверждение НЕ
    # меняло на этом экране ни одной цифры — подтверждённый сегмент без
    # находок так и стоял в «возьмёт прогон» из-за недостающего back-check
    # или несмотревшего судьи. Человек прочитал и заверил — открытых
    # вопросов нет, это «готово». Прогон такие сегменты по-прежнему ВОЗЬМЁТ
    # (проверки статус не фильтруют — это их право и защита), и как только
    # появится находка, сегмент сам уйдёт в human (confirmed_findings /
    # qaCritical): «готово» — снимок «сейчас вопросов нет», а не пожизненный
    # пропуск. Объективные находки (override_ids) остаются машине: их ремонт
    # берёт и без разрешения, обещание честное.
    machine_set -= (confirmed_ids - override_ids)

    _order = [s["id"] for s in project["segments"]]
    ready_set = set(_order) - human_set - machine_set

    # Что автоодобрение разложило бы прямо сейчас — тем же движком, что и кнопка,
    # иначе цифра на экране разошлась бы с тем, что произойдёт по нажатию.
    pending = [c for c in _term_queue() if c.get("status", "pending") == "pending"
               and _scope_of(c) == scope]
    ctx = _auto_context(pending, pol)
    ready, need_human, need_wait = 0, {}, {}
    for cand in pending:
        action, reason = _auto_verdict(cand, ctx)
        if action in (GLOSSARY_TIER_HARD, GLOSSARY_TIER_SOFT):
            ready += 1
        elif action == "wait":
            # Не хватает ДАННЫХ, а не решения: следующие чистые прогоны
            # приносят доноров сами. Считать это «ждёт человека» значило
            # звать его к работе, которой нет, — на боевом проекте так
            # пугали 412 карточек из 684.
            need_wait[reason] = need_wait.get(reason, 0) + 1
        elif action != "close":
            need_human[reason] = need_human.get(reason, 0) + 1

    result = {
        "ok": True,
        "total": len(project["segments"]),
        # Три корзины «под ключ» (см. сборку выше). Порядок — документа.
        # params — единственный источник параметров для кнопки «Перевести
        # и доделать»: состав machine считан ПОД НИХ, и запуск с другими
        # оставил бы число недостижимым (например, без judge_all корзину
        # держали бы сегменты выше зоны судьи, к которым судья не придёт).
        "turnkey": {
            "ready": [i for i in _order if i in ready_set],
            "machine": [i for i in _order if i in machine_set],
            "human": [i for i in _order if i in human_set],
            # Заверенные человеком (по статусу — тем же признаком, которым
            # раздаются корзины). Не корзина, а СРЕЗ поверх них: работа
            # человека обязана быть видна на экране числом, а не только
            # через рост «готово».
            "confirmed": [i for i in _order if i in confirmed_ids],
            # Претензии слепых измерителей (балл back-check, одиночное мнение
            # termcheck) снял свежий вердикт ревизии (`_review_vouches`).
            # Тоже срез поверх корзин: сегменты лежат в «готово».
            "reviewVouched": review_vouched,
            # Начертание приказных терминов: чинится бесплатной командой
            # /term-case, кнопка предлагает её отдельной галочкой.
            "case": [i for i in (impact.get("caseSegments") or ())
                     if i not in human_set],
            "params": {"steps": list(FULL_RUN_STEPS), "use_judge": True,
                       "judge_all": True, "retry": False,
                       "include_confirmed": False},
        },
        "clean": clean,
        "repaired": repaired,
        # Готовность БЕЗ двойного счёта. `clean` и `repaired` считаются
        # независимо от корзин остатка и пересекаются с ними: на боевом
        # проекте 181 сегмент числился и готовым, и требующим работы, отчего
        # готовность показывалась 93.3% вместо 86.6%. Считаем здесь, а не
        # в браузере: два расчёта одного числа однажды разойдутся.
        # Вычитаются КОРЗИНЫ, а не перечень списков: перечень отставал от
        # корзин (в нём не было judge_ext, споров, забракованных слов), и
        # у `clean` это было недостижимо, а у `vouched` — штатно: сегмент
        # с ручательством, которого ещё ждёт судья, числился бы и готовым,
        # и работой прогона разом.
        "readyIds": sorted(i for i in ready_set
                           if i in (set(clean) | set(repaired) | set(review_vouched))),
        "machine": {"repaired": len(repaired), "reverted": len(reverted)},
        "proposed": {"terms": ready},
        "human": {
            "terms": sorted(({"reason": k, "count": v} for k, v in need_human.items()),
                            key=lambda x: -x["count"]),
            "termsTotal": sum(need_human.values()),
            # Ждут не решения, а ДАННЫХ (доноров, чистых проверок) — дорешает
            # автоматика следующих прогонов. Отдельно, чтобы не пугать числом.
            "termsWaiting": sorted(({"reason": k, "count": v} for k, v in need_wait.items()),
                                   key=lambda x: -x["count"]),
            "termsWaitingTotal": sum(need_wait.values()),
            "reverted": reverted,
            # Слабый балл, который судья уже не поднимет (смотрел либо
            # не позовут) — та часть `todo.weak`, что лежит в human_set.
            # Экран группирует ручную работу по действию и обязан брать
            # состав у сервера, а не вычитать judge_ext сам.
            "weak": [w["id"] for w in weak if w["id"] not in judge_ext],
            # Из них: правка была верной, отменил её негодный измеритель.
            # Готовый текст лежит в сегменте, применяется без вызова модели
            # (POST /api/segments/{pid}/{sid}/repair/accept).
            "revertedByScore": score_vetoed,
            "glossaryConfirmed": impact["confirmed"],
            "confirmedFindings": confirmed_findings,
            # Машина сняла заверение человека — с доказательством на сегменте.
            "confirmWithdrawn": withdrawn,
            # Список терминов, а не сегментов: вопрос здесь про ЗАПИСЬ
            # глоссария («правильна ли она»), и отвечать на него посегментно
            # значит задать один и тот же вопрос десятки раз.
            "termcheckDisputes": disputed,
            "termcheckDisputesSegments": sorted({i for d in disputed for i in d["segments"]}),
            # Сколько спорных сегментов арбитр ещё не видел (платно, по кнопке)
            # и что он уже сказал про записи глоссария.
            "termContextPending": ctx_pending,
            "termContextWrong": ctx_wrong,
            # Забракованное проверкой слово всё ещё в тексте — с примерами.
            "staleFindings": stale_findings,
            "staleFindingWords": [{"id": k, "words": v[:3]}
                                  for k, v in sorted(stale_bad.items())][:20],
            # Подтверждённые с критической находкой Medical QA. Раньше их
            # показывала только вкладка «Замечания» — с её уходом это
            # единственное место, где такую находку видно.
            "qaCritical": qa_critical,
            # Ревизор говорит, что повреждён САМ ОРИГИНАЛ. Единственная
            # корзина, где машина бессильна по построению: перевод чинить
            # нечем, пока не выправлен исходник. До появления шага сказать
            # это было негде вовсе — termcheck прямо инструктирован
            # не трогать ничего в SOURCE.
            "sourceSuspect": source_suspect,
            # Ревизия нашла проблему, но текст не тронула: кандидат не прошёл
            # объективные сверки либо варианта не было. Совет — в карточке
            # сегмента, решение за человеком.
            "reviewFlagged": review_flagged,
        },
        "todo": {"untranslated": untranslated, "unchecked": unchecked,
                 "findings": findings, "glossaryPending": impact["pending"],
                 "weak": [w["id"] for w in weak],
                 # Не «плохо», а «никто не смотрел». Лечится судьёй.
                 "nothingToCheck": nothing_to_check,
                 "unverified": unverified,
                 "unjudgedBlind": unjudged_blind,
                 # Разнобой по документу: один оборот переведён по-разному.
                 # Список ПАР, а не сегментов: решение одно на пару, и в этом
                 # весь смысл — ручной работы одно нажатие вместо сотни.
                 "consistency": [
                     {"was": p["was"], "want": p["want"], "why": p["why"],
                      "segments": p["segments"], "already": p["already"]}
                     for p in consist_pairs],
                 "weakWhy": sorted(
                     ({"reason": r, "count": sum(1 for w in weak if w["why"] == r)}
                      for r in {w["why"] for w in weak}),
                     key=lambda x: -x["count"])},
    }
    _ANALYSIS_CACHE[pid] = (fp, result)
    return result


@app.get("/api/models")
def list_models():
    """Каталог GPT-моделей с ценами — для выпадающего списка и оценки стоимости пакета.
    Полосы back-check отдаются отсюда же, чтобы границы не дублировались на фронтенде."""
    return {
        # Ранг приклеивается здесь, а не хранится в OPENAI_MODELS: он живёт
        # в отдельном справочнике, который правят без деплоя (см. model_rank).
        "models": [dict(m, rank=model_rank(m["id"])) for m in OPENAI_MODELS],
        # Модели, которых нет в списке выбора, но которые стоят денег
        # (эмбеддинги back-check). Смета обязана брать их цену отсюда:
        # цифра в .jsx — это второй прайс-лист рядом с настоящим.
        "aux": AUX_MODEL_PRICES,
        "embedModel": EMBED_MODEL,
        "domains": ([{"id": d["id"], "label": d["label"]} for d in DOMAINS]
                    + [{"id": d["id"], "label": d["label"], "custom": True} for d in _tenant_domains()]),
        "domainDefault": DEFAULT_DOMAIN,
        "languages": LANGUAGES,
        "default": DEFAULT_OPENAI_MODEL,
        "backcheckDefault": BACKCHECK_DEFAULT_MODEL,
        "termcheckDefault": TERMCHECK_DEFAULT_MODEL,
        "repairDefault": REPAIR_DEFAULT_MODEL,
        "judgeDefault": JUDGE_DEFAULT_MODEL,
        # Модель сверки терминов. Отдаётся по той же причине, что и
        # четыре соседние: браузер не знает, чем пойдёт шаг с пустым
        # выбором, — а без цены шага прочерком становится ВСЯ смета
        # главной кнопки (шаг с работой и без цены обнуляет её намеренно).
        "termauditDefault": TERM_CONTEXT_DEFAULT_MODEL,
        # По той же причине, что и termauditDefault: браузер заполняет выбор
        # модели для КАЖДОГО ключа FULL_STEP_MODEL, а шаг с работой и без цены
        # обнуляет ВСЮ смету главной кнопки.
        "reviewDefault": REVIEW_DEFAULT_MODEL,
        # Порог, ниже которого (и при котором) ревизия правит текст. Числом
        # в .jsx он был бы вторым порогом рядом с настоящим — тот же закон,
        # что у judgeZone и backcheckBands.
        "reviewApplyMax": REVIEW_APPLY_MAX,
        "judgeZone": list(JUDGE_ZONE),
        # Уровни находок termcheck, по которым ремонт работает. Фронтенд по ним
        # считает состав кнопки «запустить только ремонт», и держать этот список
        # в .jsx литералом нельзя: он разошёлся с сервером ровно в тот день,
        # когда ремонту разрешили minor, — строка обещала 168 сегментов, а
        # кнопка под ней говорила «нечего запускать».
        "termcheckActionable": list(TERMCHECK_ACTIONABLE),
        # Порог «оригинал слишком короток, чтобы лексика что-то значила».
        # Отдаётся по той же причине, что judgeZone и backcheckBands: подсказка
        # у тумблера судьи называет это число словами, и вбитое в .jsx оно
        # разошлось бы с medical_qa молча.
        "backcheckMinStems": getattr(checks_mod, "BACKCHECK_MIN_STEMS", 3) if checks_mod else 3,
        "backcheckBands": getattr(checks_mod, "BACKCHECK_BANDS", []) if checks_mod else [],
        "available": bool(os.environ.get("OPENAI_API_KEY")),
        "brand": APP_BRAND,
        "pricesChecked": "2026-08-15",
    }


@app.get("/api/projects")
def list_projects():
    return [{k: v for k, v in p.items() if k != "segments"} | {"segmentCount": len(p["segments"])} for p in _tenant_projects()]


@app.get("/api/projects/{pid}")
def get_project_detail(pid: int):
    return _project_for_client(get_project(pid))


class CreateProjectRequest(BaseModel):
    title: str
    src: str = "RU"
    tgt: str = "EN"
    domain: str = DEFAULT_DOMAIN
    fileName: Optional[str] = None

@app.post("/api/projects")
def create_project(req: CreateProjectRequest):
    src, tgt = _check_lang_pair(req.src, req.tgt)
    # Номер — глобальный (один ряд на все организации: файлы исходников
    # и якоря картинок именуются им). Сегментов у нового проекта НЕТ: прежде
    # сюда копировались восемь сегментов первого проекта в списке — то есть
    # текст одного клиента оказывался в проекте другого.
    new_id = max((p["id"] for p in STATE["projects"]), default=0) + 1
    new_project = {
        "id": new_id,
        "title": req.title or "Новый проект",
        "titleEn": req.title or "New Project",
        "src": src, "tgt": tgt,
        "domain": _resolve_domain(req.domain)["id"],
        "tenant": _current_tenant(),
        "status": "in_progress",
        "created": datetime.now().strftime("%Y-%m-%d"),
        "deadline": "",
        "segments": [],
    }
    STATE["projects"].insert(0, new_project)
    save_state(STATE)
    return new_project


# ─── Исходный .docx: хранение и разметка ────────────────────────────
# Экспорт «как в оригинале» невозможен из одних сегментов: импорт забирает
# голый текст, а шрифты, картинки, таблицы и разметка живут в файле. Значит
# файл надо хранить и при выгрузке подменять в НЁМ текст, а не собирать
# документ заново — тогда всё, чего мы не трогали, остаётся ровно таким,
# каким его сделал автор.
#
# Якорь сегмента — НОМЕР абзаца в порядке XML. Файл лежит у нас и больше
# не меняется, поэтому номер стабилен; хранить w14:paraId незачем — его нет
# в документах старых версий Word, а разбирать два вида якорей значит однажды
# разойтись с самим собой.
SOURCE_DIR = DATA_DIR / "sources"   # внутри ReadWritePaths systemd-юнита

# Отбор абзацев в сегменты. Вынесен из upload_project, потому что теперь те же
# правила нужны второму месту — привязке исходника к УЖЕ существующему проекту.
# Разойдись они, привязка сажала бы перевод не на те абзацы.
_SKIP_PARA_RE = re.compile(r'[\d\s\-–—.,:;()\[\]/]+')


def _docx_clean(text: str) -> str:
    text = _html_mod.unescape(text)
    text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)
    return re.sub(r'\s+', ' ', text).strip()


def _docx_flat_paragraphs(doc) -> list:
    """Все абзацы документа одним списком: сначала тело, затем колонтитулы
    по имени части пакета.

    Колонтитулы лежат отдельными частями и в `body` не попадают — без этого
    бегущий заголовок остаётся на языке оригинала, а человек узнаёт об этом,
    только открыв готовый файл. Обход по ЧАСТЯМ, а не по секциям: секции делят
    колонтитулы между собой, и обход по секциям выдал бы один и тот же абзац
    несколько раз, а порядок зависел бы от того, какая секция первой его
    попросила.

    Тело идёт первым не случайно: номера абзацев тела от появления
    колонтитулов не сдвинулись, и карты, записанные раньше, остались верны."""
    return [p for p, _part in _docx_flat_parts(doc)]


def _docx_flat_parts(doc) -> list:
    """[(абзац, часть пакета, которой он принадлежит)] в том же порядке.

    Часть нужна ссылкам на картинки: r:embed указывает на связь ВНУТРИ своей
    части, и связь абзаца из колонтитула в связях документа не найдётся.
    Порядок задаётся здесь один раз: разойдись он с `_docx_flat_paragraphs`,
    номера абзацев в карте перестали бы значить то же самое, и переводы
    встали бы по всему документу не на свои места."""
    from docx.oxml.ns import qn as _qn
    main = doc.part
    out = [(p, main) for p in doc.element.body.findall(".//" + _qn("w:p"))]
    for part in sorted(doc.part.package.iter_parts(), key=lambda p: str(p.partname)):
        ct = part.content_type or ""
        el = getattr(part, "element", None)
        if el is not None and ("header+xml" in ct or "footer+xml" in ct):
            out.extend((p, part) for p in el.findall(".//" + _qn("w:p")))
    return out


def _docx_paragraph_texts(content: bytes) -> list:
    """[(текст, куда встанет перевод; полный текст абзаца)] по КАЖДОМУ абзацу
    в порядке разбора — включая те, что в сегменты не пойдут. Индекс в этом
    списке и есть якорь, поэтому список обязан быть полным: выбросишь пустые
    абзацы — и номера уедут.

    В сегмент идёт текст СЛОТОВ (`_para_slots`) — ровно то, куда экспорт 1в1
    пишет перевод: без результата вычисляемых полей, скрытого текста
    и вложенных w:p. Раньше склеивался весь `.//w:t`, и номер страницы из
    поля PAGEREF оглавления попадал в сегмент («ЛЕЧЕНИЯ.80»): его переводили
    как текст, арбитр и ремонт отрывали или удаляли, back-check ставил ложное
    жёсткое «расхождение чисел», а экспорт снимал хвост обратно (`trimmed`).
    Полный текст остаётся рядом ради СТАРЫХ проектов: их сегменты равны ему
    (это склейка прежнего импорта — весь .//w:t, с вложенными надписями),
    и по нему их узнают привязка исходника и экспорт. Экспорт сравнивает
    с полным текстом `_para_slots` (без вложенных w:p), поэтому абзац
    с надписью у старого проекта там по-прежнему `mismatch` — как и прежде."""
    from docx import Document
    from docx.oxml.ns import qn as _qn
    doc = Document(io.BytesIO(content))
    out = []
    for p in _docx_flat_paragraphs(doc):
        slots, _full, _dropped = _para_slots(p, _qn)
        # Полный текст — ровно так, как склеивал ПРЕЖНИЙ импорт: весь .//w:t,
        # включая вложенные надписи. Сегменты старых проектов равны ему буква
        # в букву, и привязка исходника узнаёт их без потерь.
        out.append((_docx_clean("".join((t.text or "") for t, _sig in slots)),
                    _docx_clean("".join(t.text for t in p.iter(_qn("w:t")) if t.text))))
    return out


def _docx_paragraphs(content: bytes) -> list:
    """Текст сегмента по каждому абзацу — см. `_docx_paragraph_texts`."""
    return [t for t, _full in _docx_paragraph_texts(content)]


def _docx_units(paras: list, full: Optional[list] = None) -> list:
    """(текст сегмента, [номера абзацев]) по правилам импорта: слишком короткие
    и чисто цифровые абзацы переводить нечего, соседние одинаковые строки —
    это один сегмент.

    Одинаковыми соседи считаются, только если равен и ПОЛНЫЙ текст (`full`,
    с результатом поля): две строки оглавления «Введение» с разными номерами
    страниц — разные абзацы. Склей их — и на старом проекте (там это два
    сегмента) второй абзац получил бы перевод первого рядом с живым полем.

    Соседний повтор раньше просто выбрасывался. Теперь он попадает в тот же
    сегмент вторым якорем: при выгрузке перевод должен встать в ОБА абзаца,
    иначе второй останется на языке оригинала."""
    units: list = []
    prev, prev_i = None, -1
    for i, t in enumerate(paras):
        if len(t) < 2 or _SKIP_PARA_RE.fullmatch(t):
            continue
        if t == prev and units and (full is None or full[i] == full[prev_i]):
            units[-1][1].append(i)
            prev_i = i
            continue
        units.append((t, [i]))
        prev, prev_i = t, i
    return units


def _source_paths(pid: int) -> tuple:
    return SOURCE_DIR / ("%d.docx" % pid), SOURCE_DIR / ("%d.json" % pid)


def _store_source_docx(project: dict, content: bytes, filename: str,
                       pairs: list, paras: int) -> dict:
    """Кладёт исходник и карту «абзац → сегмент» рядом с ним.

    Карта принадлежит файлу, а не состоянию проекта: state.json целиком лежит
    в памяти и переписывается ПРИ КАЖДОМ сохранении, а на 2670 сегментах карта
    весит десятки килобайт — это лишний мегабайт записи на каждую правку
    одного сегмента. В самом проекте остаётся только отметка о наличии."""
    SOURCE_DIR.mkdir(parents=True, exist_ok=True)
    docx_path, map_path = _source_paths(project["id"])
    tmp = docx_path.with_name(docx_path.name + ".tmp")
    tmp.write_bytes(content)
    os.replace(str(tmp), str(docx_path))
    payload = {"file": filename, "paras": paras, "pairs": pairs,
               "at": datetime.now().strftime("%Y-%m-%d %H:%M")}
    # Разбор картинок ПЕРЕЖИВАЕТ повторную привязку исходника. Карта абзацев
    # и карта картинок лежат в одном файле, но пишутся из разных мест: пока
    # эта запись их не сохраняла, нажатие «Заменить» в карточке исходника
    # молча уносило 199 секунд разбора, прочитанный за деньги текст и связь
    # с сегментами — а сегменты оставались в проекте, и их перевод исчезал
    # из выгрузки без единого счётчика. Приложили ДРУГОЙ файл — картинки
    # в нём другие, и это ловится отпечатком (sha) при выгрузке и при
    # следующем разборе, а не молчаливой потерей работы.
    if map_path.exists():
        try:
            was = json.loads(map_path.read_text(encoding="utf-8"))
            for k in ("images", "imagesAt", "imagesSkipped", "imagesTotal"):
                if k in was:
                    payload[k] = was[k]
        except Exception as e:
            print("[backend] прежняя карта проекта %s не прочиталась: %s"
                  % (project["id"], e), file=sys.stderr)
    tmp = map_path.with_name(map_path.name + ".tmp")
    tmp.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
    os.replace(str(tmp), str(map_path))
    mark = {"file": filename, "at": payload["at"], "paras": paras,
            "segments": len({p[1] for p in pairs})}
    project["sourceDocx"] = mark
    return mark


def _load_source_map(pid: int) -> Optional[dict]:
    """Карта и файл читаются только вместе: карта без файла и файл без карты
    одинаково бесполезны, а отметка в проекте могла пережить их обоих —
    восстановление state.json из бэкапа файлы не возвращает."""
    docx_path, map_path = _source_paths(pid)
    if not docx_path.exists() or not map_path.exists():
        return None
    try:
        data = json.loads(map_path.read_text(encoding="utf-8"))
    except Exception as e:
        print("[backend] карта исходника проекта %s не читается: %s" % (pid, e),
              file=sys.stderr)
        return None
    data["path"] = docx_path
    return data


@app.post("/api/projects/upload")
async def upload_project(
    file: UploadFile = File(...),
    title: str = Form(""),
    src: str = Form("RU"),
    tgt: str = Form("EN"),
    domain: str = Form(DEFAULT_DOMAIN),
):
    src, tgt = _check_lang_pair(src, tgt)
    try:
        import docx  # noqa: F401 — проверка наличия, разбор идёт в _docx_paragraphs
    except ImportError:
        raise HTTPException(500, "python-docx not installed")

    content = await file.read()
    # Разбор и отбор — общие с привязкой исходника к готовому проекту
    # (_docx_paragraphs / _docx_units): две копии правил однажды разошлись бы,
    # и перевод при выгрузке встал бы не в те абзацы.
    texts = _docx_paragraph_texts(content)
    paras = [t for t, _full in texts]
    units = _docx_units(paras, [f for _t, f in texts])
    deduped = [t for t, _ in units]

    new_id = max((p["id"] for p in STATE["projects"]), default=0) + 1
    proj_title = title or file.filename.rsplit(".", 1)[0]
    new_project = {
        "id": new_id,
        "title": proj_title,
        "titleEn": proj_title,
        "src": src, "tgt": tgt,
        "domain": _resolve_domain(domain)["id"],
        "tenant": _current_tenant(),
        "status": "in_progress",
        "created": datetime.now().strftime("%Y-%m-%d"),
        "deadline": "",
        "fileName": file.filename,
        "segments": [
            {
                "id": i + 1,
                "source": text,
                "target": "",
                "status": "new",
                "comments": [],
                "qa": [],
                "wordCount": len(text.split()),
                "risk": "high" if len(text.split()) > 30 else "medium" if len(text.split()) > 8 else "low",
                "route": "GPT_REQUIRED",
                "tm": None,
            }
            for i, text in enumerate(deduped)
        ],
    }
    # Исходник храним сразу: без него экспорт «как в оригинале» невозможен
    # в принципе, а второй раз тот же файл человек может и не найти.
    # Ошибка записи не роняет импорт: сегменты разобраны, переводить можно,
    # а исходник прикладывается отдельной командой.
    try:
        pairs = [[i, u + 1] for u, (_t, idxs) in enumerate(units) for i in idxs]
        _store_source_docx(new_project, content, file.filename, pairs, len(paras))
    except Exception as e:
        print("[backend] исходник проекта %s не сохранён: %s" % (new_id, e),
              file=sys.stderr)
    STATE["projects"].insert(0, new_project)
    save_state(STATE)
    return new_project


# ─── Привязка исходника к УЖЕ существующему проекту ─────────────────
# Проекты, импортированные до появления экспорта «как в оригинале», исходника
# не сохранили — а переводить их заново нельзя, там оплаченная работа. Разбор
# детерминированный, поэтому тот же файл даёт тот же список абзацев, и якоря
# садятся на существующие сегменты без единого изменения в переводах.
#
# Сопоставление идёт ПО ТЕКСТУ, а не по номерам: правила отбора абзацев
# со временем меняются (номера страниц из оглавления, табуляции), и жёсткая
# привязка «первый абзац = первый сегмент» после любой такой правки посадила
# бы перевод на чужие строки — молча и по всему документу.
_SOURCE_LOOKAHEAD = 50   # насколько далеко вперёд ищем сегмент под абзац


def _match_key(text: str) -> str:
    return re.sub(r'\s+', ' ', (text or "")).strip().lower()


def _map_source_to_segments(units: list, segments: list,
                            full: Optional[list] = None) -> tuple:
    """(пары «абзац → id сегмента», сколько сегментов нашлось).

    `full` — полный текст абзацев по индексу: сегмент СТАРОГО импорта равен
    ему (с номером страницы из поля), нового — тексту слотов. Свой абзац
    узнаётся по любому из двух.

    Идём двумя указателями вперёд: документ тот же и порядок тот же, поэтому
    окно поиска маленькое. Абзац, которому сегмента не нашлось, пропускаем —
    при выгрузке он останется на языке оригинала, и это честнее, чем сдвинуть
    на него чужой перевод."""
    # Сегменты из картинок в этом поиске не участвуют вовсе: абзацами они
    # не являются, а мешают дважды. Во-первых, два указателя идут вперёд
    # с окном в _SOURCE_LOOKAHEAD, и отсканированная страница на сорок надписей
    # рядом с такой же соседней сдвигает окно так, что след теряется.
    # Во-вторых, они раздували знаменатель порога «совпало меньше половины»,
    # и родной файл отклонялся как чужой.
    segments = [s for s in segments if (s.get("origin") or {}).get("kind") != "image"]
    pairs, matched = [], set()
    j = 0
    for text, idxs in units:
        keys = {_match_key(text)}
        if full:
            keys |= {_match_key(full[i]) for i in idxs if i < len(full)}
        hit = None
        for k in range(j, min(j + _SOURCE_LOOKAHEAD, len(segments))):
            if _match_key(segments[k].get("source")) in keys:
                hit = k
                break
        if hit is None:
            continue
        for i in idxs:
            pairs.append([i, segments[hit]["id"]])
        matched.add(segments[hit]["id"])
        j = hit + 1
    return pairs, len(matched)


@app.post("/api/projects/{pid}/source")
async def attach_source(pid: int, file: UploadFile = File(...), force: bool = Form(False)):
    """Приложить исходный .docx к готовому проекту — для экспорта 1в1.

    Переводы, проверки и статусы не трогаются вообще: пишется только файл
    и карта абзацев рядом с ним."""
    project = get_project(pid)
    if _job_busy(pid, "images"):
        # Разбор держит копию карты в памяти и переписывает файл после каждой
        # картинки: новая привязка легла бы под его следующее сохранение,
        # и экспорт пошёл бы по устаревшим номерам абзацев молча.
        raise HTTPException(409, "Идёт разбор картинок — дождитесь конца "
                                 "или остановите его")
    try:
        import docx  # noqa: F401
    except ImportError:
        raise HTTPException(500, "python-docx not installed")
    content = await file.read()
    try:
        texts = _docx_paragraph_texts(content)
    except Exception as e:
        return {"ok": False, "error": "Файл не читается как .docx: %s" % e}

    paras = [t for t, _full in texts]
    full = [f for _t, f in texts]
    units = _docx_units(paras, full)
    pairs, matched = _map_source_to_segments(units, project["segments"], full)
    total = len([s for s in project["segments"]
                 if (s.get("origin") or {}).get("kind") != "image"])
    stats = {"paras": len(paras), "units": len(units),
             "segments": total, "matched": matched, "unmatched": total - matched}

    # Не тот файл виден по числу совпадений, и молча положить его нельзя:
    # экспорт потом расставил бы переводы по чужим абзацам. Порог — половина:
    # правки отбора абзацев столько не съедают, а другой документ не наберёт.
    if not force and (matched == 0 or matched * 2 < total):
        return {"ok": False, "stats": stats,
                "error": ("Совпало %d сегментов из %d — похоже, это другой файл "
                          "или другая его редакция. Экспорт 1в1 расставил бы "
                          "переводы по чужим абзацам." % (matched, total))}

    mark = _store_source_docx(project, content, file.filename, pairs, len(paras))
    save_state(STATE)
    return {"ok": True, "stats": stats, "sourceDocx": mark}


# ─── Текст, впечатанный в картинки ───────────────────────────────────
# В учебнике фтизиатрии 158 растровых картинок, и часть текста живёт ТОЛЬКО
# в них: подписи под рисунками, схемы, куски отсканированных страниц.
# Абзацного якоря у такого текста нет, переводить его было нечем, и в выгрузке
# «1в1» он оставался на языке оригинала.
#
# Работа делится надвое, и это не косметика:
#   ГЕОМЕТРИЮ даёт локальный детектор (image_text.detect_lines): по рамке
#   придётся стирать и писать, а зрячая модель координаты выдумывает —
#   ошибка в рамке означает заплатку посреди снимка;
#   ТЕКСТ даёт зрячая модель: распознавалка, идущая с детектором, знает
#   китайский и английский, а русский вернула латинской абракадаброй
#   («пиопневмоторакс» → «WONHeBMOTOpaKC»).
#
# Второй вид якоря: «часть пакета + номер блока». Карта лежит в том же
# sidecar-файле рядом с исходником, что и карта абзацев, и по той же причине:
# state.json целиком в памяти и переписывается при КАЖДОМ сохранении.
try:
    import image_text
except Exception as _e:            # модуль в том же репозитории, но частичный
    image_text = None              # деплой не должен ронять весь бэкенд
    print("[backend] image_text не импортирован: %s" % _e, file=sys.stderr)

IMAGE_MEDIA_EXT = {".png", ".jpeg", ".jpg", ".gif", ".bmp", ".tif", ".tiff"}
# Сколько блоков уходит модели одним вызовом. Отсканированная страница даёт
# два-три десятка блоков, и складывать их в один запрос — это картинка
# на несколько мегабайт в теле и ответ, в котором модель теряет нумерацию.
IMAGE_READ_MAX_BLOCKS = 8
# Модель чтения: та же, что по умолчанию у перевода, либо названная в
# окружении. Выбора на экране пока нет — и обещать его комментарием нельзя;
# задача принимает `ocr_model`, так что появится он правкой одной строки
# в карточке, а не переделкой.
IMAGE_READ_MODEL = os.environ.get("IMAGE_READ_MODEL", "") or DEFAULT_OPENAI_MODEL


def _docx_media(content: bytes) -> tuple:
    """({имя части: байты} растровых картинок, [имена нерастровых]).

    EMF/WMF — векторные метафайлы Windows, растровые библиотеки их не читают.
    В учебнике таких четыре, и они честно считаются пропущенными: молча
    выброшенная картинка выглядела бы как «текста в ней нет»."""
    import zipfile
    raster, other = {}, []
    with zipfile.ZipFile(io.BytesIO(content)) as z:
        for name in z.namelist():
            if not name.startswith("word/media/"):
                continue
            ext = ("." + name.rsplit(".", 1)[-1]).lower() if "." in name else ""
            if ext in IMAGE_MEDIA_EXT:
                raster[name] = z.read(name)
            else:
                other.append(name)
    return raster, sorted(other)


def _docx_media_part(path, name: str) -> Optional[bytes]:
    """Одна картинка из пакета, не читая пакет целиком.

    Карточка сегмента спрашивает кроп при каждом открытии, а исходник учебника
    весит 22 МБ: `_docx_media` распаковал бы в память все 158 картинок ради
    одной, и на каждый клик по сегменту."""
    import zipfile
    try:
        with zipfile.ZipFile(str(path)) as z:
            return z.read(name)
    except Exception as e:
        print("[backend] картинка %s не достаётся: %s" % (name, e), file=sys.stderr)
        return None


def _docx_image_anchors(doc) -> dict:
    """{имя части: [номера абзацев, где она стоит]}.

    Нужен, чтобы посадить сегмент картинки рядом с тем местом документа,
    где человек её видит: сегмент, приехавший в конец списка, переводят
    без всякого контекста. Обход тот же, что у карты абзацев
    (`_docx_flat_parts`), иначе номера разъедутся."""
    R = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
    out: dict = {}
    for idx, (p, part) in enumerate(_docx_flat_parts(doc)):
        seen = set()
        for el in p.iter():
            for key, rid in el.attrib.items():
                if not key.startswith(R) or rid in seen:
                    continue
                seen.add(rid)
                rel = part.rels.get(rid)
                if rel is None or "image" not in (rel.reltype or "") or rel.is_external:
                    continue
                name = str(rel.target_part.partname).lstrip("/")
                out.setdefault(name, []).append(idx)
    return out


def _save_source_map(pid: int, data: dict) -> None:
    """Карта рядом с исходником, запись атомарная — тем же приёмом, что
    и save_state: разбор идёт минутами, и оборванная посередине запись
    оставила бы проект с картой, по которой экспорт сажает переводы
    неизвестно куда."""
    _docx_path, map_path = _source_paths(pid)
    payload = {k: v for k, v in data.items() if k != "path"}
    tmp = map_path.with_name(map_path.name + ".tmp")
    tmp.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
    os.replace(str(tmp), str(map_path))


def _image_tokens(w: int, h: int) -> int:
    """Грубая оценка входных токенов на картинку: плитки 512×512 плюс базовая.
    Нужна только для сметы ДО прогона — факт снимается с ответа модели
    (`_note_usage`), как и везде."""
    tiles = max(1, -(-int(w) // 512)) * max(1, -(-int(h) // 512))
    return 85 + 170 * tiles


def _image_read_system(dom: dict, src_lang: str) -> str:
    """Промпт чтения. Два вопроса, и второй не менее важен первого.

    `text` — что написано в рамке, дословно. `overlay` — надпись ли это
    САМОГО АППАРАТА, а не документа: на 38 снимках учебника из 51 с текстом
    в рамках лежат фамилии пациентов, даты исследования и настройки томографа
    («KARIMOV SH.», «DEPTH:14cm», «kV 120.0»). Переводить их незачем, а тащить
    в сегменты, память переводов и очередь терминов нельзя тем более —
    это персональные данные, и попав в TM они разъедутся по чужим проектам."""
    return (
        "You read text that is printed INSIDE an illustration of a "
        + dom["en"] + " document. The document language is " + src_lang + ".\n"
        "You are given the whole picture (for context) and then a series of "
        "close-up crops, one per numbered block.\n\n"
        "For EVERY block return:\n"
        '  "text"    — exactly what is written in that crop, verbatim, in the '
        "original language. Keep numbers, units and capitalisation. If the "
        "block is split into several lines, join them into one string; if a "
        "word is hyphenated at a line break, join it back into one word. "
        "If you cannot read it, return an empty string — never guess.\n"
        '  "overlay" — true if this is NOT text of the document but a label '
        "burnt in by equipment or software: patient name, examination date, "
        "scanner settings, ruler marks, menu items, watermark, file name. "
        "Otherwise false.\n\n"
        'Return ONLY JSON: {"blocks": [{"i": 0, "text": "...", "overlay": false}]}'
    )


def _image_read_parse(raw: str) -> list:
    """Блоки из ответа модели, даже если ответ оборван.

    Целиком JSON разбирается редко: на отсканированной странице ответ упирается
    в потолок и обрывается на середине строки. Тогда берём все ЗАКОНЧЕННЫЕ
    объекты — за них заплачено, и выбрасывать их значит платить за них снова
    следующим заходом, который оборвётся ровно там же."""
    import json as _json
    # Целый ответ: как есть и обрезанный по крайним скобкам — модель любит
    # обернуть JSON в ```json.
    lo, hi = raw.find("{"), raw.rfind("}")
    for probe in (raw, raw[lo:hi + 1] if 0 <= lo < hi else ""):
        if not probe:
            continue
        try:
            data = _json.loads(probe)
        except Exception:
            continue
        if isinstance(data, dict) and isinstance(data.get("blocks"), list):
            return [b for b in data["blocks"] if isinstance(b, dict)]
    out = []
    for m in re.finditer(r"\{[^{}]*\}", raw, re.S):
        try:
            obj = _json.loads(m.group(0))
        except Exception:
            continue
        if isinstance(obj, dict) and "i" in obj:
            out.append(obj)
    return out


def _openai_read_image(img_bytes: bytes, blocks: list, src_lang: str,
                       domain_id: Optional[str] = None,
                       model: str = None) -> Optional[list]:
    """Прочитать текст в найденных рамках. Список той же длины, что blocks,
    либо None — вызов не удался.

    None и пустой текст — разные вещи: None означает «не спросили», и разбор
    обязан оставить блок нерешённым, а не объявить картинку пустой."""
    import json as _json
    import base64
    import openai
    if image_text is None or not blocks:
        return None
    dom = _resolve_domain(domain_id)
    mdl = _resolve_model(model or IMAGE_READ_MODEL)
    client = openai.OpenAI(api_key=os.environ.get("OPENAI_API_KEY"), timeout=120,
                           max_retries=1)
    # Обзорный кадр — уменьшенный PNG, а не сырые байты части: в пакете лежат
    # и jpeg, и gif, и tiff, а уходили они объявленные как image/png.
    over = image_text.preview(img_bytes)
    content = [{"type": "text", "text": "Whole picture, for context only:"}]
    if over:
        content.append({"type": "image_url",
                        "image_url": {"url": "data:image/png;base64," +
                                      base64.b64encode(over).decode(),
                                      "detail": "low"}})
    asked = []
    for i, b in enumerate(blocks):
        # lineH/rows — чтобы поле кропа считалось от СТРОКИ: у подписи
        # в шесть строк «четверть блока» захватывает полторы чужие, и модель
        # добросовестно перепишет соседний абзац в этот сегмент.
        crop = image_text.crop(img_bytes, b["box"], line_h=b.get("lineH"),
                               rows=b.get("rows"))
        if not crop:
            continue
        asked.append(i)
        content.append({"type": "text", "text": "Block %d:" % i})
        content.append({"type": "image_url",
                        "image_url": {"url": "data:image/png;base64," +
                                      base64.b64encode(crop).decode(),
                                      "detail": "high"}})
    if not asked:
        return None
    extra = ({"max_completion_tokens": 4096} if mdl["api"] == "modern"
             else {"max_tokens": 4000, "temperature": 0})
    try:
        resp = client.chat.completions.create(
            model=mdl["id"],
            messages=[{"role": "system", "content": _image_read_system(dom, src_lang)},
                      {"role": "user", "content": content}],
            **extra)
        _note_usage("ocr", mdl["id"], resp)
        raw = (resp.choices[0].message.content or "").strip()
        why = getattr(resp.choices[0], "finish_reason", "") or ""
    except Exception as e:
        print("[backend] чтение картинки не удалось: %s" % e, file=sys.stderr)
        return None
    items = _image_read_parse(raw)
    if not items:
        print("[backend] чтение картинки: ответ не разобран (%s), %d симв."
              % (why or "?", len(raw)), file=sys.stderr)
        return None
    if why == "length":
        # Ответ оборвался на потолке. Разобранные блоки НЕ выбрасываем: за них
        # уже заплачено, а остальные останутся нерешёнными и уйдут в следующий
        # заход. Прежде выбрасывался весь ответ целиком, и следующий заход
        # с той же нарезкой обрывался ровно так же — оплачиваемая карусель.
        print("[backend] чтение картинки: ответ оборван, разобрано блоков %d из %d"
              % (len(items), len(blocks)), file=sys.stderr)
    out = [None] * len(blocks)
    for item in items:
        if not isinstance(item, dict):
            continue
        try:
            i = int(item.get("i"))
        except (TypeError, ValueError):
            continue
        if not (0 <= i < len(blocks)):
            continue
        out[i] = {"text": (item.get("text") or "").strip(),
                  "overlay": bool(item.get("overlay")),
                  "model": mdl["id"]}
    return out


def _image_seg_of(by_id: dict, block: dict, part: str, i: int) -> Optional[dict]:
    """Сегмент ЭТОГО блока — или None.

    Номера сегментов переиспользуются: `max(id) + 1` после сноса выдаёт те же
    числа заново. Значит одного номера мало — сегмент обязан подтвердить, что
    он и есть этот блок, своим якорем. Иначе блок цепляется к чужому сегменту,
    остаётся без перевода и молча выпадает из выгрузки."""
    seg = by_id.get(block.get("seg")) if block.get("seg") else None
    if seg is None:
        return None
    o = seg.get("origin") or {}
    return seg if (o.get("kind") == "image" and o.get("part") == part
                   and o.get("block") == i) else None


def _image_anchor_sid(data: dict, paras: list) -> Optional[int]:
    """Сегмент, после которого встанут блоки этой картинки: последний текст
    ПЕРЕД ней. Нет такого — первый после неё; нет и его — конец списка."""
    pairs = sorted((int(p[0]), int(p[1])) for p in (data.get("pairs") or []))
    if not pairs or not paras:
        return None
    first = min(paras)
    before = [sid for idx, sid in pairs if idx <= first]
    if before:
        return before[-1]
    return pairs[0][1]


def _image_new_segment(text: str, part: str, block: int, next_id: int) -> dict:
    """Сегмент из картинки. Поля те же, что у абзацного: он идёт через тот же
    перевод, те же проверки и те же корзины разбора — отдельная сущность
    означала бы второе место, меняющее текст, и второй набор проверок."""
    words = len(text.split())
    return {
        "id": next_id, "source": text, "target": "", "status": "new",
        "comments": [], "qa": [],
        "wordCount": words,
        "risk": "high" if words > 30 else "medium" if words > 8 else "low",
        "route": "GPT_REQUIRED", "tm": None,
        # Якорь: по нему экспорт находит, куда вписать перевод, а карточка
        # сегмента — какой кусок картинки показать человеку.
        "origin": {"kind": "image", "part": part, "block": block},
    }


def _image_place_segment(project: dict, seg: dict, anchor_sid: Optional[int],
                         after_id: Optional[int]) -> None:
    """Вставить сегмент сразу за его картинкой. Сегмент, уехавший в конец
    списка, переводят и подтверждают без всякого контекста — а контекст здесь
    и есть половина работы."""
    segs = project["segments"]
    want = after_id if after_id is not None else anchor_sid
    if want is not None:
        for i, s in enumerate(segs):
            if s["id"] == want:
                segs.insert(i + 1, seg)
                return
    segs.append(seg)


def _image_stats(images: list, total: int = 0) -> dict:
    """Что известно про картинки проекта. Один расчёт на карточку экрана
    и на отчёт задачи: разойдись они — под соседними числами стояли бы
    разные ответы на один вопрос."""
    st = {"images": total or len(images), "withText": 0, "blocks": 0, "text": 0,
          "overlay": 0, "noise": 0, "unread": 0, "segments": 0,
          "repaintable": 0, "captioned": 0, "scanned": 0, "unreadable": 0,
          # Сколько ещё сделает прогон: непрошенные блоки плюс прочитанные,
          # которым сегмента пока нет. По этому числу решает кнопка чтения —
          # раньше она смотрела на смету, и после сноса сегментов (текст
          # остался, значит платить не за что, значит смета ноль) кнопка
          # гасла навсегда, хотя работа была возможна.
          "pending": 0, "unasked": 0}
    for im in images:
        if im.get("unreadable"):
            # Посмотреть не удалось. Это не «текста нет» и прятать это нельзя:
            # иначе картинка навсегда числится разобранной и пустой.
            st["unreadable"] += 1
            continue
        blocks = im.get("blocks") or []
        st["blocks"] += len(blocks)
        if im.get("blocks") is not None:
            st["scanned"] += 1
        if blocks:
            st["withText"] += 1
        for b in blocks:
            if "text" not in b:
                st["unasked"] += 1
                st["pending"] += 1
                continue
            skip = b.get("skip")
            if skip == "overlay":
                st["overlay"] += 1
            elif skip == "noise":
                st["noise"] += 1
            elif b.get("text"):
                st["text"] += 1
                if b.get("seg"):
                    st["segments"] += 1
                else:
                    st["pending"] += 1
                # Порог плоскости фона — тот же, по которому откажет
                # перерисовка при экспорте. Разойдись они, карточка обещала бы
                # одно, а готовый файл показывал другое.
                floor = image_text.IMG_FLAT_MIN if image_text else 1.1
                if (b.get("flat") or 0) >= floor:
                    st["repaintable"] += 1
                else:
                    st["captioned"] += 1
            else:
                st["unread"] += 1
    return st


def _image_est_tokens(images: list) -> tuple:
    """(входных, выходных) токенов на то, что ещё не читали.

    Считается по КРОПАМ — модели уходят они, а не картинка целиком; обзорный
    кадр идёт с detail=low и стоит фиксированные 85 токенов."""
    tin = 0
    for im in images:
        blocks = [b for b in (im.get("blocks") or []) if "text" not in b]
        if not blocks:
            continue
        tin += 85
        for b in blocks:
            x0, y0, x1, y1 = b["box"]
            tin += _image_tokens(max(1, x1 - x0), max(1, y1 - y0))
    return tin, tin // 6


def _image_est_by_model(images: list) -> dict:
    """Смета по КАЖДОЙ модели каталога: {id: цена}.

    Считает сервер, а не браузер, и по той же таблице, по которой потом
    списывается факт. Цифра в .jsx была бы вторым прайс-листом рядом
    с настоящим — тем самым, который однажды разойдётся. Цена неизвестна —
    None, а не ноль: неизвестное, посчитанное нулём, показывает расход
    меньше настоящего."""
    tin, tout = _image_est_tokens(images)
    return {m["id"]: (0.0 if not tin else _usage_cost(m["id"], tin, tout))
            for m in OPENAI_MODELS}


@app.get("/api/projects/{pid}/images")
def images_report(pid: int):
    """Что известно про текст на картинках проекта. Ни одного вызова модели."""
    project = get_project(pid)
    ready, why = image_text.engine_ready() if image_text else (False, "модуль не собран")
    data = _load_source_map(pid) if project.get("sourceDocx") else None
    images = (data or {}).get("images") or []
    mdl = _resolve_model(IMAGE_READ_MODEL)
    tin, tout = _image_est_tokens(images)
    return {"ok": True, "engine": ready, "why": why,
            "hasSource": data is not None,
            # Модель по умолчанию — от неё пляшет выпадающий список; выбранную
            # человеком задача получает отдельным полем `ocr_model`.
            "model": mdl["id"],
            "stats": _image_stats(images, (data or {}).get("imagesTotal") or 0),
            "est": _image_est_by_model(images),
            "estTokens": {"in": tin, "out": tout},
            "at": (data or {}).get("imagesAt"),
            "skipped": (data or {}).get("imagesSkipped") or []}


@app.get("/api/projects/{pid}/images/crop")
def image_crop(pid: int, seg: int = 0, part: str = "", block: int = 0):
    """Кусок картинки с одной надписью — картинкой PNG.

    Без него распознанное проверить нечем: человек видит в карточке строку
    текста и не может знать, то ли это, что нарисовано. Спрашивать можно
    и по сегменту, и по паре «часть + блок»."""
    project = get_project(pid)
    if image_text is None:
        raise HTTPException(503, "Работа с картинками недоступна: модуль не собран")
    data = _load_source_map(pid) if project.get("sourceDocx") else None
    if data is None:
        raise HTTPException(404, "К проекту не приложен исходный .docx")
    if seg:
        s = next((x for x in project["segments"] if x["id"] == seg), None)
        origin = (s or {}).get("origin") or {}
        if origin.get("kind") != "image":
            raise HTTPException(404, "Сегмент #%d пришёл не из картинки" % seg)
        part, block = origin.get("part") or "", int(origin.get("block") or 0)
    rec = next((im for im in (data.get("images") or []) if im.get("part") == part), None)
    blocks = (rec or {}).get("blocks") or []
    if rec is None or not (0 <= block < len(blocks)):
        raise HTTPException(404, "Такой надписи в карте картинок нет")
    blob = _docx_media_part(data["path"], part)
    if blob is None:
        raise HTTPException(404, "Картинки %s в исходнике больше нет" % part)
    b = blocks[block]
    png = image_text.crop(blob, b["box"], line_h=b.get("lineH"), rows=b.get("rows"))
    if not png:
        raise HTTPException(404, "Кусок картинки не вырезался")
    return Response(content=png, media_type="image/png",
                    headers={"Cache-Control": "private, max-age=86400"})


@app.get("/api/projects/{pid}/images/blocks")
def images_blocks(pid: int, skip: str = "", limit: int = 400):
    """Найденные надписи списком: что отсеяно и почему.

    Без этого «Отсеяно: 230» — число, которое человеку нечем проверить.
    А проверять есть что: отсев делает модель, и ошибается она в обе стороны.
    `skip`: overlay | noise | none (то, что стало сегментами) | пусто — всё."""
    project = get_project(pid)
    data = _load_source_map(pid) if project.get("sourceDocx") else None
    if data is None:
        raise HTTPException(404, "К проекту не приложен исходный .docx")
    by_id = {s["id"]: s for s in project["segments"]}
    out = []
    for im in (data.get("images") or []):
        for i, b in enumerate(im.get("blocks") or []):
            if "text" not in b:
                continue
            mark = b.get("skip") or "none"
            if skip and mark != skip:
                continue
            seg = _image_seg_of(by_id, b, im.get("part"), i)
            out.append({"part": im.get("part"), "block": i, "skip": b.get("skip"),
                        "text": b.get("text") or "", "conf": b.get("conf"),
                        "flat": b.get("flat"), "seg": seg["id"] if seg else None,
                        # Кто решил. Метку ставят трое: модель при чтении,
                        # согласие между картинками и сам человек — и найти
                        # СВОИ пометки в списке «отсеяла модель» он должен
                        # уметь, иначе разберёт их заново.
                        "by": b.get("by") or "model"})
    return {"ok": True, "blocks": out[:max(1, min(limit, 2000))],
            "total": len(out)}


class ImageRestoreRequest(BaseModel):
    # «Это текст документа, а не надпись аппарата» — обратное решение к тому,
    # что принимает кнопка в карточке сегмента.
    part: str
    block: int


@app.post("/api/projects/{pid}/images/restore")
def image_restore_block(pid: int, req: ImageRestoreRequest):
    """Вернуть отсеянную надпись в работу: снять метку и завести сегмент.

    Отсев делает модель, и она ошибается в обе стороны. Обратный ход обязан
    существовать по тому же правилу, по которому у каждой пачки автоодобрения
    есть откат: решение машины, которое человек не может отменить, — это
    не помощь, а приговор."""
    project = get_project(pid)
    if _job_busy(pid, "images"):
        raise HTTPException(409, "Идёт разбор картинок — дождитесь конца")
    data = _load_source_map(pid) if project.get("sourceDocx") else None
    if data is None:
        raise HTTPException(404, "К проекту не приложен исходный .docx")
    rec = next((im for im in (data.get("images") or []) if im.get("part") == req.part), None)
    blocks = (rec or {}).get("blocks") or []
    if rec is None or not (0 <= req.block < len(blocks)):
        raise HTTPException(404, "Такой надписи в карте картинок нет")
    b = blocks[req.block]
    text = (b.get("text") or "").strip()
    if not text:
        raise HTTPException(400, "Надпись не прочитана — возвращать нечего")
    by_id = {s["id"]: s for s in project["segments"]}
    have = _image_seg_of(by_id, b, req.part, req.block)
    if have is not None:
        return {"ok": True, "segment": have["id"], "created": False}
    # Номер абзаца-якоря записан при разборе. Читать ради него весь .docx
    # (у учебника это 22 МБ и 2687 абзацев на единственном воркере) незачем;
    # старые карты его не знают — только для них и открываем документ.
    paras = rec.get("paras")
    if not paras:
        from docx import Document
        anchors = _docx_image_anchors(Document(io.BytesIO(Path(data["path"]).read_bytes())))
        paras = anchors.get(req.part) or []
    nid = max((s["id"] for s in project["segments"]), default=0) + 1
    seg = _image_new_segment(text, req.part, req.block, nid)
    # Встаём ЗА своими же соседями по картинке, а не перед ними: иначе
    # у возвращённого сегмента в промпте перевода окажутся чужие соседи,
    # а на одном якоре с двумя картинками он сядет в середину первой.
    after = None
    for i, other in enumerate(blocks):
        if i >= req.block:
            break
        sib = _image_seg_of(by_id, other, req.part, i)
        if sib is not None:
            after = sib["id"]
    _image_place_segment(project, seg, _image_anchor_sid(data, paras), after)
    b.pop("skip", None)
    b["by"] = "human"          # см. _image_harmonize: чужое решение не отменяем
    b["seg"] = nid
    _save_source_map(pid, data)
    save_state(STATE)
    return {"ok": True, "segment": nid, "created": True,
            "stats": _image_stats(data.get("images") or [],
                                  data.get("imagesTotal") or 0)}


class ImagesForgetRequest(BaseModel):
    # Снести заведённое: разбор оказался плохим, движок сменили, приложили
    # другой исходник.
    force: bool = False       # сносить и те сегменты, где уже есть перевод
    # Выбросить и ПРОЧИТАННЫЙ текст. По умолчанию он остаётся: геометрия —
    # это минуты работы детектора, а текст ещё и оплачен, и повторный заход
    # заведёт сегменты заново бесплатно. Отката у этой команды нет, поэтому
    # выбрасывание оплаченного должно быть отдельным решением человека,
    # а не побочным действием кнопки «забыть».
    wipe: bool = False


@app.post("/api/projects/{pid}/images/forget")
def images_forget(pid: int, req: ImagesForgetRequest):
    """Забыть сегменты, заведённые из картинок.

    Сегменты с готовым переводом по умолчанию НЕ трогаются и называются
    поимённо: это оплаченная работа, и снести её молча нельзя. Разбор картинок
    при этом остаётся (см. `wipe`) — повторный заход заведёт сегменты заново
    и ничего не заплатит."""
    project = get_project(pid)
    if _job_busy(pid, "images"):
        # Задача переписывает ту же карту и держит её копию в памяти: наша
        # правка легла бы под её следующее сохранение и пропала молча.
        raise HTTPException(409, "Идёт разбор картинок — дождитесь конца "
                                 "или остановите его")
    data = _load_source_map(pid) if project.get("sourceDocx") else None
    kept, removed = [], []
    for s in list(project["segments"]):
        if (s.get("origin") or {}).get("kind") != "image":
            continue
        if (s.get("target") or "").strip() and not req.force:
            kept.append(s["id"])
            continue
        project["segments"].remove(s)
        removed.append(s["id"])
    if data is not None:
        gone = set(removed)
        for im in (data.get("images") or []):
            for b in (im.get("blocks") or []):
                if b.get("seg") in gone:
                    b.pop("seg", None)
                if req.wipe:
                    # Выбрасываем ТОЛЬКО прочитанное. Геометрия остаётся:
                    # это минуты работы детектора, и она не зависит от того,
                    # верно ли модель прочитала текст.
                    b.pop("text", None)
                    b.pop("skip", None)
                    b.pop("reader", None)
        _save_source_map(pid, data)
    save_state(STATE)
    return {"ok": True, "removed": len(removed), "keptTranslated": kept,
            "wiped": bool(req.wipe),
            "stats": _image_stats((data or {}).get("images") or [],
                                  (data or {}).get("imagesTotal") or 0)}


def _image_harmonize(data: dict, project: dict) -> tuple:
    """Одинаковый текст на разных картинках размечается одинаково.

    Модель решает про каждую картинку отдельно и на границе ошибается: строка
    настроек томографа «kV 120.0 mA: 283 …» на боевом учебнике оказалась
    надпечаткой на двух снимках и «текстом документа» на третьем — и стала
    сегментом, за перевод которого человек заплатит. Спрашивать модель второй
    раз незачем: ответ у неё уже есть, просто разный. Берём большинство —
    это ноль вызовов и ноль догадок.

    Сегмент, у которого УЖЕ есть перевод, не трогаем: это оплаченная работа,
    и снимать её по счёту голосов нельзя."""
    groups: dict = {}
    for im in (data.get("images") or []):
        for b in (im.get("blocks") or []):
            txt = _norm_key(b.get("text") or "")
            if txt:
                groups.setdefault(txt, []).append(b)
    by_id = {sg["id"]: sg for sg in project["segments"]}
    changed = dropped = 0
    for blocks in groups.values():
        if len(blocks) < 2:
            continue
        over = sum(1 for b in blocks if b.get("skip") == "overlay")
        if over <= len(blocks) - over:
            continue
        for b in blocks:
            if b.get("skip") == "overlay":
                continue
            if b.get("by") == "human":
                # Человек уже сказал про эту надпись «это текст документа».
                # Отменить его решение счётом голосов — значит слушаться его
                # ровно в одну сторону: пометку «надпечатка» согласие не
                # трогает никогда, и обратное решение должно быть так же вечно.
                continue
            seg = by_id.get(b.get("seg"))
            if seg is not None and (seg.get("target") or "").strip():
                continue                       # переведённое не отменяем
            b["skip"] = "overlay"
            if seg is not None:
                project["segments"].remove(seg)
                b.pop("seg", None)
                dropped += 1
            changed += 1
    return changed, dropped


@app.post("/api/projects/{pid}/images/{sid}/overlay")
def image_mark_overlay(pid: int, sid: int):
    """«Это надпись аппарата, а не текст документа»: убрать сегмент и запомнить.

    Модель ошибается на границе, и ошибается в обе стороны. На боевом учебнике
    из 41 заведённого сегмента 15 оказались строками прибора — «LINEAR DISTANCE:
    004.43cm», «kV 120.0 mA: 283», «IPAP DOKTOR HAKIMOV M». Последнее — фамилия
    врача, и ей нечего делать в памяти переводов.

    Машиной это дальше не отсеять: правило «на картинке большинство надписей —
    надпечатка» на тех же данных убило бы законную подпись «Диссеминированный
    (милиарный) туберкулёз», а «нет букв языка оригинала» — латинское название
    вида. Значит решает человек, а система обязана СЛУШАТЬСЯ и ПОМНИТЬ: метка
    ложится на блок, и следующий разбор сегмент не заведёт заново."""
    project = get_project(pid)
    if _job_busy(pid, "images"):
        raise HTTPException(409, "Идёт разбор картинок — дождитесь конца")
    seg = next((x for x in project["segments"] if x["id"] == sid), None)
    origin = (seg or {}).get("origin") or {}
    if seg is None or origin.get("kind") != "image":
        raise HTTPException(404, "Сегмент #%d пришёл не из картинки" % sid)
    data = _load_source_map(pid) if project.get("sourceDocx") else None
    if data is None:
        raise HTTPException(404, "К проекту не приложен исходный .docx")
    had = (seg.get("target") or "").strip()
    for im in (data.get("images") or []):
        if im.get("part") != origin.get("part"):
            continue
        for i, b in enumerate(im.get("blocks") or []):
            if i == origin.get("block"):
                b["skip"] = "overlay"
                b["by"] = "human"
                b.pop("seg", None)
    project["segments"].remove(seg)
    _save_source_map(pid, data)
    save_state(STATE)
    # Про выброшенный перевод говорим прямо: человек решил, но знать, что
    # именно уходит вместе с сегментом, он обязан.
    return {"ok": True, "removed": sid, "hadTarget": had,
            "stats": _image_stats(data.get("images") or [],
                                  data.get("imagesTotal") or 0)}


def _job_images(job: dict) -> None:
    """Разбор картинок проекта: найти строки, собрать в блоки, прочитать
    моделью, завести сегменты.

    Идёт задачей, а не запросом: 158 картинок учебника — это две с половиной
    минуты только на поиск строк. Остановка мягкая, сделанное сохраняется:
    разбор кэшируется ПО СОДЕРЖИМОМУ картинки (sha), поэтому повторный заход
    не платит и не считает заново."""
    pid = job["project"]
    project = get_project(pid)
    dry = bool(job["params"].get("dry_run", True))
    if image_text is None:
        raise RuntimeError("Работа с картинками недоступна: модуль не собран")
    ready, why = image_text.engine_ready()
    if not ready:
        # Не «текста не найдено», а честный отказ с причиной: молчаливый ноль
        # означал бы «в книге нет ни одной подписи».
        raise RuntimeError("Поиск строк невозможен: " + why)
    data = _load_source_map(pid) if project.get("sourceDocx") else None
    if data is None:
        raise RuntimeError("К проекту не приложен исходный .docx — "
                           "искать текст не в чем")
    if not dry and not os.environ.get("OPENAI_API_KEY"):
        raise RuntimeError("Чтение текста требует ключ OpenAI")

    content = Path(data["path"]).read_bytes()
    raster, other = _docx_media(content)
    from docx import Document
    anchors = _docx_image_anchors(Document(io.BytesIO(content)))
    known = {im.get("part"): im for im in (data.get("images") or [])}
    # Порядок обхода — ПО ДОКУМЕНТУ, а не по имени части. Имена сортируются
    # как строки («image1, image10, image100»), и сегменты двух картинок,
    # стоящих между одними и теми же абзацами, вставали бы в списке задом
    # наперёд: читать такой проект человеку.
    names = sorted(raster, key=lambda n: (min(anchors.get(n) or [10 ** 9]), n))
    job["total"], job["done"] = len(names), 0
    job["counters"]["skippedFormat"] = len(other)
    model = job["params"].get("ocr_model") or IMAGE_READ_MODEL
    # Карту и состояние сохраняем ПО РАЗНЫМ поводам. Карта — файл на десятки
    # килобайт, её пишем после каждой картинки: разбор идёт минутами, и
    # остановка не должна стоить уже сделанной работы. А state.json на боевом
    # проекте это 5 МБ, и переписывать его на каждую из 158 картинок, когда
    # сегментов не прибавилось, — три четверти гигабайта записи впустую.
    out, map_dirty, segs_dirty = [], False, False
    # За каким сегментом встала последняя вставка у этого якоря. Без этого две
    # картинки между одними абзацами обе цепляются за один и тот же сегмент,
    # и вторая встаёт ПЕРЕД первой.
    placed_after: dict = {}

    def flush():
        """Карту на диск. Зовётся из двух мест, и оба обязаны писать её сами:
        нечитаемая картинка уходит `continue` мимо конца цикла."""
        data["images"] = out + [known[n] for n in names[len(out):] if n in known]
        data["imagesAt"] = datetime.now().strftime("%Y-%m-%d %H:%M")
        data["imagesSkipped"] = other
        data["imagesTotal"] = len(names)
        _save_source_map(pid, data)

    for name in names:
        if job["stop"]:
            job["status"] = "stopped"
            break
        # recent — это id сегментов, их ждёт /segments/fetch у редактора.
        # Имя части в этом поле давало 422 каждые три секунды.
        job["recent"] = []
        blob = raster[name]
        sha = hashlib.sha1(blob).hexdigest()
        prev = known.get(name) or {}
        rec = dict(prev) if prev.get("sha") == sha and prev.get("blocks") is not None \
            else {"part": name, "sha": sha}
        if rec.get("blocks") is None:
            lines = image_text.detect_lines(blob)
            if lines is None:
                # Битый файл, незнакомый формат, векторный метафайл. Это «не
                # знаю» про ОДНУ картинку: записать сюда пустой список значило
                # бы объявить, что надписей в ней нет. Разбор идёт дальше —
                # ронять из-за одной картинки полтораста остальных незачем.
                rec["unreadable"] = True
                rec.pop("blocks", None)
                out.append(rec)
                job["counters"]["unreadable"] = job["counters"].get("unreadable", 0) + 1
                job["done"] += 1
                flush()
                map_dirty = False
                continue
            rec.pop("unreadable", None)
            rec["blocks"] = image_text.group_blocks(lines)
            try:
                from PIL import Image as _PIL
                rec["w"], rec["h"] = _PIL.open(io.BytesIO(blob)).size
            except Exception:
                rec["w"] = rec["h"] = 0
            rec["paras"] = anchors.get(name) or []
            job["counters"]["detected"] = job["counters"].get("detected", 0) + 1
            map_dirty = True
        blocks = rec["blocks"]
        if blocks:
            job["counters"]["withText"] = job["counters"].get("withText", 0) + 1

        # ЧТЕНИЕ. Спрашиваем только про блоки, о которых ещё не спрашивали:
        # разбор кэширован по sha картинки, и повторный заход не платит
        # за то, что уже прочитано.
        todo = [i for i, b in enumerate(blocks) if "text" not in b]
        if todo and not dry:
            for start in range(0, len(todo), IMAGE_READ_MAX_BLOCKS):
                if job["stop"]:
                    break
                part_idx = todo[start:start + IMAGE_READ_MAX_BLOCKS]
                got = _openai_read_image(blob, [blocks[i] for i in part_idx],
                                         project.get("src") or "RU",
                                         project.get("domain"), model)
                if got is None:
                    # Не спросили — блок остаётся нерешённым. Записать сюда
                    # пустоту значило бы навсегда закрыть картинку от чтения.
                    job["counters"]["readFailed"] = job["counters"].get("readFailed", 0) + 1
                    continue
                for k, i in enumerate(part_idx):
                    ans = got[k]
                    if ans is None:
                        continue
                    b = blocks[i]
                    b["text"] = ans["text"]
                    b["reader"] = ans["model"]
                    if ans["overlay"]:
                        b["skip"] = "overlay"
                    elif image_text.is_noise(ans["text"]):
                        b["skip"] = "noise"
                    else:
                        b.pop("skip", None)
                    map_dirty = True

        # СЕГМЕНТЫ. Только то, что прочитано, не является надпечаткой аппарата
        # и не шум. Уже заведённые не задваиваются: карта помнит номер.
        if not dry:
            by_id = {s["id"]: s for s in project["segments"]}
            anchor = _image_anchor_sid(data, anchors.get(name) or [])
            after = placed_after.get(anchor)
            made = []
            for i, b in enumerate(blocks):
                if b.get("skip") or not (b.get("text") or "").strip():
                    continue
                have = _image_seg_of(by_id, b, name, i)
                if have is not None:
                    after = have["id"]
                    continue
                nid = max((s["id"] for s in project["segments"]), default=0) + 1
                seg = _image_new_segment(b["text"].strip(), name, i, nid)
                _image_place_segment(project, seg, anchor, after)
                b["seg"] = nid
                after = nid
                made.append(nid)
                map_dirty = segs_dirty = True
                job["counters"]["segments"] = job["counters"].get("segments", 0) + 1
            if after is not None:
                placed_after[anchor] = after
            if made:
                job["recent"] = made
        out.append(rec)
        job["done"] += 1
        if map_dirty:
            flush()
            map_dirty = False
        if segs_dirty:
            save_state(STATE)
            segs_dirty = False
    if not dry:
        # Согласие между картинками — последний бесплатный фильтр перед тем,
        # как человек увидит список.
        moved, dropped = _image_harmonize(data, project)
        if moved:
            job["counters"]["harmonized"] = moved
            # Вычитаем ровно снятые сегменты, а не все согласованные блоки:
            # часть из них сегментами не была.
            job["counters"]["segments"] = max(
                0, job["counters"].get("segments", 0) - dropped)
    flush()
    save_state(STATE)
    st = _image_stats(data["images"], len(names))
    for k in ("blocks", "text", "overlay", "noise", "unread"):
        job["counters"][k] = st[k]
    job["counters"]["repaintable"] = st["repaintable"]
    # Модели детектора держат сотни мегабайт, а разбор картинок бывает раз
    # в жизни проекта: воркер uvicorn один, и эта память отнимается у сервиса.
    image_text.release_engine()


# ─── Segment actions ────────────────────────────────────────────────
class TranslateRequest(BaseModel):
    # engine больше нет: движок один — выбранная модель. Поле осталось бы
    # враньём про выбор, которого не существует.
    force: bool = False  # True = пропустить TM-шорткат (ручной перевод)
    model: Optional[str] = None  # id из OPENAI_MODELS; неизвестный → DEFAULT_OPENAI_MODEL

@app.post("/api/segments/{pid}/{sid}/translate")
def translate_segment(pid: int, sid: int, req: TranslateRequest):
    _guard_project_write(pid)
    # обычный def, а не async: внутри блокирующий вызов модели (см. batch_translate)
    seg = get_segment(pid, sid)
    project = get_project(pid)
    src_text = seg["source"]

    # Глоссарий + TM контекст
    gloss_hits, tm_hit = _get_context(src_text, project=project)
    gloss_hits = gloss_hits + _doc_hits(src_text, project, gloss_hits)

    # TM точное совпадение → 0 токенов (только для авто/пакетного, не для ручного force-перевода)
    # и только для записей, которые завёл человек: см. _tm_trusted.
    if not req.force and _tm_trusted(tm_hit) and tm_hit.get("tgt"):
        # НЕ confirmed: одно подтверждение человека на одной строке не заверяет
        # все будущие строки с тем же исходником. Сегмент проходит проверки,
        # как всякий машинный перевод, и подтверждает его снова человек.
        _replace_target(seg, tm_hit["tgt"], PROVIDER_TM, "EXACT_TM")
        seg.pop("docTerms", None)      # промпта не было — следа быть не должно
        save_state(STATE)
        return {"ok": True, "segment": seg, "usedRealApi": False, "source": "TM"}

    if not os.environ.get("OPENAI_API_KEY"):
        raise HTTPException(503, "Перевод требует ключ OpenAI: бесплатного движка "
                                 "в системе больше нет")
    try:
        prev_src, next_src = _neighbours(project, seg)
        translation = _openai_translate(src_text, project["src"], project["tgt"],
                                        gloss_hits=gloss_hits, tm_context=tm_hit,
                                        style=_style_block(project),
                                        model=req.model, domain=project.get("domain"),
                                        prev_src=prev_src, next_src=next_src)
    except Exception as e:
        print(f"[backend] translate failed seg#{sid}: {e}", file=sys.stderr)
        translation = None
    if not translation:
        # Ни заглушек, ни «черновика бесплатным движком»: в медицинском переводе
        # подсунуть вместо ответа что попало хуже, чем не ответить. Сегмент
        # не трогаем, ошибку называем.
        raise HTTPException(502, "Модель не вернула перевод. Попробуйте ещё раз "
                                 "или выберите другую модель.")
    # След для замера (`_termlist_measure`) — только у состоявшегося вызова:
    # у ошибки или TM-совпадения промпта не было.
    _dt = [h["tgt"] for h in gloss_hits if h.get("tier") == "doc"]
    if _dt:
        seg["docTerms"] = _dt
    else:
        seg.pop("docTerms", None)

    _replace_target(seg, translation, _resolve_model(req.model)["id"], "GPT_REQUIRED")
    save_state(STATE)
    return {"ok": True, "segment": seg, "usedRealApi": True}


@app.post("/api/segments/{pid}/{sid}/qa")
def qa_segment(pid: int, sid: int):
    seg = get_segment(pid, sid)
    # Simple local QA checks
    qa_issues = []
    src, tgt = seg["source"], seg.get("target", "")
    # Check numbers preserved
    import re
    src_nums = re.findall(r"\d+(?:[.,]\d+)?", src)
    tgt_nums = re.findall(r"\d+(?:[.,]\d+)?", tgt)
    if sorted(src_nums) != sorted(tgt_nums) and src_nums:
        qa_issues.append({"sev": "high", "type": "numeric",
                          "msg": f"Числа не совпадают: source={src_nums}, target={tgt_nums}"})
    # Check length not crazy
    if tgt and len(tgt) > 3 * len(src):
        qa_issues.append({"sev": "medium", "type": "length",
                          "msg": "Перевод более чем в 3 раза длиннее оригинала."})

    seg["qa"] = qa_issues
    seg["status"] = "qa"
    save_state(STATE)
    return {"ok": True, "segment": seg, "issues": qa_issues}



# ─── Что система выучивает из подтверждённого сегмента ───────────────
# Подтверждение — единственная точка, где появляется достоверная пара
# «оригинал → перевод». Из неё берём три вещи: обновляем TM, складываем
# терминологические находки в очередь кандидатов и находим повторы того же
# исходника. В сам глоссарий автоматически не попадает НИЧЕГО: он инжектится
# в промпт как правило, и автопополнение закрепляло бы собственные ошибки.
# Потолок очереди. Он не про «столько вопросов человек осилит», а про размер
# state.json: файл целиком лежит в памяти и переписывается при КАЖДОМ
# сохранении. Ожидающая карточка тяжелее решённой — в ней лежат примеры
# сегментов (`sampleSrc`/`sampleTgt`, по 240 символов), которые `_mark_decided`
# снимает при решении. Восемьсот на документе в 2670 сегментов кончились
# в первый же день: очередь встала у потолка, и каждая новая находка стала
# вытеснять старую — то есть терялась работа, ради которой прогон и шёл.
# Переменной окружения, потому что зависит от размера документов, а не от кода.
TERM_QUEUE_MAX = max(200, int(os.environ.get("TERM_QUEUE_MAX", "2500")))
# Сколько чужих вариантов помнит решённая карточка. Это история спора, а не
# данные: без потолка один термин рос бы в state.json на каждом прогоне.
TERM_REASKED_MAX = 8
# Кандидаты, выброшенные потолком, с момента старта процесса. Растёт только
# под _TERM_QUEUE_LOCK (подрезка живёт внутри _queue_insert), поэтому свой лок
# не нужен. Прогон снимает разницу до и после — см. _job_loop.
_TERM_DROPPED = {"total": 0}
TERM_DROP_LOG_EVERY = 100


def _norm_key(text: str) -> str:
    return " ".join((text or "").lower().replace("ё", "е").split())


def _same_words(a: str, b: str) -> bool:
    """Тот же текст с точностью до пробелов, но НЕ до регистра.

    Отдельно от `_norm_key` намеренно: тот отвечает на вопрос «то же ли это
    слово» (ключ поиска, дедупликация) и потому опускает регистр. А здесь
    вопрос другой — «изменился ли текст», и правка, у которой всё отличие
    в заглавной букве, это изменение."""
    return " ".join((a or "").split()) == " ".join((b or "").split())


def _tm_upsert(source: str, target: str, project: dict = None) -> str:
    """Пара в TM: обновить существующую запись, а не пропустить её.
    Раньше дедуп находил старую пару и молча оставлял как есть — исправленный
    перевод в память не попадал, а прежний, неверный, продолжал автоматически
    подставляться в новые проекты как EXACT_TM."""
    key = _norm_key(source)
    today = datetime.now().strftime("%Y-%m-%d")
    lang = f"{(project or {}).get('src', 'RU')}→{(project or {}).get('tgt', 'EN')}"
    tenant = _tenant_of(project) if project else _current_tenant()
    for t in STATE["tm"]:
        if _norm_key(t.get("src")) != key or _tenant_of(t) != tenant:
            continue
        # Пара языков — часть ключа. Без неё подтверждение RU→DE переписывало
        # tgt у RU→EN записи, оставляя ей прежний lang: следующий RU→EN проект
        # получал немецкий текст как EXACT_TM, да ещё со статусом confirmed.
        if (t.get("lang") or DEFAULT_GLOSS_LANG) != lang:
            continue
        if (t.get("tgt") or "").strip() == (target or "").strip():
            return "kept"
        t["prevTgt"] = t.get("tgt", "")
        t["tgt"] = target
        t["quality"] = "verified"
        t["score"] = 100
        t["updated"] = today
        t["by"] = _actor_id()            # кто заверил этот перевод (id, «human» без сессии)
        return "updated"
    STATE["tm"].insert(0, {
        "src": source, "tgt": target, "lang": lang, "tenant": tenant,
        "score": 100, "quality": "verified", "used": 1, "created": today,
        "by": _actor_id(),
    })
    return "added"


def _term_queue() -> list:
    return STATE.setdefault("termQueue", [])


def _trim_term_queue():
    """Очередь не должна расти бесконечно: state.json целиком лежит в памяти
    и переписывается при каждом сохранении."""
    q = _term_queue()
    if len(q) <= TERM_QUEUE_MAX:
        return
    # Сначала обработанные: своё они уже отыграли. Кроме тех, что принадлежат
    # ещё откатываемой пачке: без кандидата откат вернёт глоссарий, но потеряет
    # само решение, и человек его больше не увидит.
    live = {b.get("id") for b in STATE.get("autoBatches", [])}
    done = [c for c in q if c.get("status") != "pending" and c.get("autoBatch") not in live]
    # Решения человека снимаем последними: решённый кандидат — это память
    # «про этот термин уже спрашивали». Выбросив её, сбор терминологии заведёт
    # кандидата заново, и одобренный термин вернётся в очередь неодобренным.
    # Решения человека сортируем по ДАТЕ РЕШЕНИЯ, а не по месту в очереди:
    # порядок в очереди — это порядок появления кандидатов, и старый кандидат,
    # одобренный сегодня, лежит в самом низу. По позиции его выбросило бы
    # первым — то есть терялась бы память именно о свежих решениях.
    humans = [c for c in done if _human_decision(c)]
    machines = [c for c in done if not _human_decision(c)]
    humans.sort(key=lambda c: c.get("decidedAt") or "", reverse=True)
    for c in (humans + machines)[TERM_QUEUE_MAX // 4:]:
        q.remove(c)
    if len(q) <= TERM_QUEUE_MAX:
        return
    # Дальше — машинный урожай. На документе в 2670 сегментов коротких пар
    # тысячи, а pending раньше не подрезался вовсе. Решения, которые может
    # принять только человек (conflict и всё, что пришло с подтверждённых
    # сегментов), не трогаем никогда — их больше никто не примет.
    droppable = [c for c in q if c.get("status", "pending") == "pending"
                 and c.get("via") == "auto" and c.get("kind") != "conflict"]
    # Сортируем по возрасту, а не по числу доноров. Кандидат с одним донором —
    # это ровно тот, кто ждёт второго: выбрасывая их первыми, мы бы вычищали
    # именно ту популяцию, ради которой существует порог auto_min_segments.
    # Одиночки уходят раньше многодонорных, но внутри группы — самые старые:
    # у них было больше всего шансов набрать доноров и они их не набрали.
    droppable.sort(key=lambda c: (1 if c.get("hits", 1) >= 2 else 0, c.get("id", 0)))
    over = len(q) - TERM_QUEUE_MAX
    for c in droppable[:over]:
        q.remove(c)
    if over > 0:
        # Молчаливых потолков не бывает. Но подрезка зовётся на КАЖДУЮ вставку,
        # и на полной очереди это один выброшенный кандидат за раз: прогон
        # 22.08 оставил в журнале 890 одинаковых строк «снято 1», из которых
        # не видно главного — что за прогон сбор терминологии потерял 890
        # находок. Поэтому счётчик: в журнал пишем изредка и с накопленным
        # итогом, а точное число за прогон уходит в его счётчики (terms_dropped).
        _TERM_DROPPED["total"] += min(over, len(droppable))
        n = _TERM_DROPPED["total"]
        if n == 1 or n % TERM_DROP_LOG_EVERY == 0:
            print(f"[backend] очередь кандидатов переполнена: выброшено {n} машинных "
                  f"кандидатов с момента старта (потолок {TERM_QUEUE_MAX})", file=sys.stderr)


_TERM_QUEUE_LOCK = threading.Lock()


# Виды находок, где пара ПРЕДЛАГАЕТСЯ из свободного текста, а не берётся
# из готового решения: извлечение терминов моделью, короткий сегмент целиком
# и пара, извлечённая из правки человека (`edit` — модель выравнивает диф,
# то есть формулировка предложена ею же). У `conflict` исходная сторона —
# уже существующая запись глоссария, у `audit` — слово из находки termcheck;
# сомневаться в самой формулировке там не наше дело, и проверять её значит
# выбрасывать чужой вопрос.
TERM_HARVEST_KINDS = ("extract", "segment", "edit")
# Сколько пар отсеяно как «это не словарная запись». Считаем, а не молчим:
# отсев без счётчика неотличим от «модель ничего не нашла».
_TERM_NOT_TERM = [0]


# Кириллическое слово целиком: по нему из общего набора служебных слов
# отбираются русские.
_CYR_WORD_RE = re.compile(r"[а-яё]+")


def _head_word(text: str) -> str:
    """Первое НАСТОЯЩЕЕ слово строки, в нижнем регистре.

    По пробелам, а не по буквам: «В-лимфоцит», «Т-хелперы», «С-реактивный
    белок» начинаются одной буквой, но она часть слова, а не предлог. Токены
    из одной пунктуации («— туберкулёз») пропускаются, иначе строка отклонялась
    бы целиком из-за тире."""
    for token in (text or "").split():
        head = re.sub(r"^\W+|\W+$", "", token, flags=re.UNICODE).lower()
        if head:
            return head
    return ""


def _looks_like_term(src: str, tgt: str) -> bool:
    """Похоже ли это на СЛОВАРНУЮ запись, а не на обрывок предложения.

    Правило одно: запись не начинается служебным словом С ОБЕИХ СТОРОН.
    «в лёгких → in the lungs», «у больного → The patient», «при кашле →
    in cough» — это куски предложения, а не термины; на боевом проекте модель
    извлечения выдавала такие пары вопреки прямой просьбе промпта («dictionary
    form», `_term_extract_system`, правило 2), и они потом держали балл
    back-check как «потерянный термин».

    Сторон именно две, и это главное в правиле. По одному только оригиналу
    отличить предлог от буквенной метки нельзя: «в лёгких» и «В-лимфоцит»,
    «с кашлем» и «С реактивный белок», «у больного» и «витамин У» начинаются
    одинаково. А вот перевод их разводит: у обрывка фразы служебное слово
    стоит и в переводе («in the lungs»), у термина — нет («B lymphocytes»,
    «C-reactive protein»). Конец записи не проверяется вообще: «Гепатит В»,
    «Гепатит С», «Витамин С» кончаются буквами, которые в любом наборе
    служебных слов стоят предлогами.

    Со стороны оригинала берутся только КИРИЛЛИЧЕСКИЕ служебные слова, поэтому
    «in situ» и «in vivo» остаются, а для источника на другом алфавите проверка
    просто МОЛЧИТ — тот же закон, что у `DOMAIN_RULES` без правил для пары.
    Со стороны перевода годится любое: там вопрос только «служебное ли это
    слово», а язык перевода заранее неизвестен.

    Чего правило НЕ ловит: «имеются у пациента» — служебное слово там
    в середине, а широкое «содержит предлог» убило бы «боль в груди». Такие
    пары останавливает сверка смысла (`_meaning_check` поставила этой паре
    `rule: false`): приказом они не становятся, а балл back-check считается
    только по приказным записям.

    Набор берётся из checks: там он уже есть и участвует в подсчёте балла,
    а второй список тех же слов однажды разошёлся бы с первым. Набора нет
    (модуль не подключён) — НЕ запрещаем: молчаливый ответ здесь «пропустить»,
    тот же закон, что у `attested()` и `_lex_blind`."""
    if not re.search(r"[^\W\d_]", src or ""):
        # Пара без единой буквы («1.3», «—») словарной записью не бывает.
        return False
    stop = set(getattr(checks_mod, "RU_STOPWORDS", ()) or ()) if checks_mod else set()
    if not stop:
        return True
    src_head, tgt_head = _head_word(src), _head_word(tgt)
    if not src_head or not tgt_head:
        return bool(src_head)
    return not (_CYR_WORD_RE.fullmatch(src_head) and src_head in stop
                and tgt_head in stop)


def _term_shape_reject(pol: dict, src: str, tgt: str) -> Optional[str]:
    """Форма, которую автоодобрение отвергнет при ЛЮБЫХ обстоятельствах:
    длиннее лимита слов политики либо кончается знаком предложения.

    ОДИН предикат на два места — `_auto_verdict` и ворота очереди в
    `_queue_term_locked`. Раньше ворота держались копией этих трёх условий,
    а обещание «отказ ровно тот же» — комментарием: первая же правка политики
    (новый знак, свой лимит области) легла бы в одно место, и очередь снова
    начала бы копить неодобряемое. Возвращает причину, а не bool: причина
    уходит человеку в вердикте автоодобрения.

    Про пустой tgt предикат молчит намеренно: пары без перевода
    (conflict-карточки) до проверки формы не доходят вовсе — их
    `_auto_verdict` отдаёт человеку раньше."""
    if len(src.split()) > pol["max_src_words"]:
        return "длинный термин — похоже на фразу"
    if tgt and len(tgt.split()) > pol["max_tgt_words"]:
        return "длинный перевод — похоже на фразу"
    if src.rstrip().endswith((".", "!", "?", ":", ";")):
        return "похоже на предложение, а не термин"
    return None


def _queue_term(kind: str, src: str, tgt: str, **extra) -> Optional[dict]:
    """Кандидат в глоссарий. Повтор той же пары не плодит запись, а поднимает
    hits: по нему видно, какая проблема встречается чаще всего. Отклонённое
    второй раз не всплывает.

    Под локом: сегменты порции считаются параллельно, а очередь одна — без него
    два потока могли бы завести две записи об одном и том же. Ворота формы
    (`_term_shape_reject`) живут ВНУТРИ лока и ПОСЛЕ дедупликации — см.
    `_queue_term_locked`: стоя здесь, на входе, они (а) молча глотали
    несогласие — у решённых терминов переставали расти `hits` и `reasked`,
    (б) резали conflict-карточки, которые по построению решает человек,
    и (в) крутили общий счётчик `_TERM_NOT_TERM` без лока из рабочих потоков."""
    src_n, tgt_n = _norm_key(src), _norm_key(tgt)
    if not src_n:
        return None
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    with _TERM_QUEUE_LOCK:
        return _queue_term_locked(kind, src, tgt, src_n, tgt_n, now, extra)


def _cand_pair(c: dict) -> tuple:
    """Пара, с которой кандидат ПОЯВИЛСЯ, — по ней и идёт дедупликация.
    Одобрение вписывает в src/tgt решение человека (у конфликта перевод вообще
    рождается пустым и заполняется только им), и без этой памяти следующий сбор
    терминологии не узнавал бы решённого кандидата: одобренный термин
    возвращался в очередь новым, будто его никто не подтверждал."""
    return (_norm_key(c["origSrc"] if "origSrc" in c else c.get("src")),
            _norm_key(c["origTgt"] if "origTgt" in c else c.get("tgt")))


def _answer_keys(c: dict) -> set:
    """Термины, на которые отвечает эта карточка, если она вообще ответ.

    Ключей два: нынешний термин и тот, с которым карточка родилась —
    одобрение вправе исправить сам термин. Пустое множество означает «это
    не ответ»: отклонение сюда не входит намеренно («эта пара неверна» не то
    же самое, что «с термином разобрались», см. reject_term_candidate), как
    и решения автоматики (`_human_decision`).

    Одно множество на всех, кто спрашивает «этот вопрос уже задавали»:
    дедупликация очереди (`_term_answered`) и разбор накопленного хвоста
    (`_migrate_term_queue`). Разойдись они в определении одного и того же
    термина — миграция снимала бы то, что дедупликация задаст заново."""
    if c.get("status") != "approved" or not _human_decision(c):
        return set()
    return {k for k in (_cand_pair(c)[0], _norm_key(c.get("src"))) if k}


def _hard_answer(entry: Optional[dict]) -> bool:
    """Запись глоссария — это ответ на вопрос «как переводить термин».

    Ответом делает уровень ПРИКАЗ (его ставит человек либо выверенный
    справочник) И непустой перевод: запись без перевода не отвечает ничего,
    и закрывать ею вопрос значило бы похоронить его молча."""
    return bool(entry) and _hit_tier(entry) == GLOSSARY_TIER_HARD         and bool((entry.get("tgt") or "").strip())


def _term_answered(c: dict, src_n: str) -> bool:
    """Карточка — это ОТВЕТ человека про тот же термин, а не другая пара."""
    return src_n in _answer_keys(c)


def _migrate_ui_lang() -> None:
    """Язык записи, доставшийся ей ПО УМОЛЧАНИЮ КОДА, — это не выбор человека.

    До появления узбекского каждая учётная запись заводилась с `uiLang: "ru"`
    литералом: выбора там не было и быть не могло, языка в системе был один.
    Оставить это значением значит показать русский интерфейс тем, кто ничего
    не выбирал, — то есть сделать перевод невидимым.

    Отличить одно от другого нечем, кроме СЛЕДА решения: `uiLangSet` пишется
    только когда язык поменял человек (`/api/profile`, `/api/admin/users`).
    Тот же приём, что `_human_touched` в глоссарии, и тот же закон: своё
    предположение машина вправе пересмотреть, чужое решение — нет.

    Идемпотентна: запись без следа каждый старт получает нынешний язык
    по умолчанию, запись со следом не трогается НИКОГДА.
    """
    changed = 0
    for u in STATE.get("users") or []:
        if u.get("uiLangSet"):
            continue
        if u.get("uiLang") != DEFAULT_UI_LANG:
            u["uiLang"] = DEFAULT_UI_LANG
            changed += 1
    if changed:
        print("[backend] язык интерфейса по умолчанию (%s) проставлен %d записям "
              "без следа выбора человека" % (DEFAULT_UI_LANG, changed), file=sys.stderr)
    return changed


def _migrate_term_queue(state: dict) -> int:
    """Снять карточки про термины, по которым решение УЖЕ есть.

    Пока дедупликация смотрела только на пару «термин + перевод» и вид находки,
    каждый прогон заводил новую карточку про уже утверждённый термин: свой
    вариант перевода или другой вид находки давали новый ключ. Заводить их
    перестал `_queue_term_locked`, но накопленное висит в очереди и просит
    решения, которое человек уже принял.

    Живёт отдельно от `_apply_migrations` намеренно: та зовётся при импорте
    раньше, чем определены `_norm_key`, `_cand_pair` и `_human_decision`, а
    дублировать их правила здесь — значит однажды разойтись с ними в том,
    что считать одним и тем же термином."""
    queue = state.get("termQueue") or []
    decided, hard = {}, {}
    for c in queue:
        for k in _answer_keys(c):
            decided.setdefault((_scope_of(c), k), c)
    for g in state.get("glossary", []):
        if _hard_answer(g):
            hard.setdefault((_scope_of(g), _norm_key(g.get("src"))), g)
    today, closed = datetime.now().strftime("%Y-%m-%d"), 0
    for c in queue:
        if c.get("status", "pending") != "pending":
            continue
        key = (_scope_of(c), _cand_pair(c)[0])
        answer = decided.get(key)
        if answer is None and key not in hard:
            continue
        # Тот же приём, что у _close_same_term: «approved» здесь означает
        # «вопрос закрыт», а autoWrote=False — что в глоссарий ничего не
        # писали. Решением человека такая карточка НЕ становится (см.
        # _human_decision), поэтому сама она вопросы не закрывает и уходит
        # при подрезке первой.
        c.update({"status": "approved", "autoWrote": False, "decidedAt": today,
                  "note": "вопрос уже решён: "
                          + ((answer or hard[key]).get("tgt") or "")})
        if answer is not None:
            c["decidedWith"] = answer.get("id")
        closed += 1
    # Обрывки фраз, накопленные ДО ворот формы (_term_shape_reject в
    # _queue_term_locked): политика отвергнет их при любых обстоятельствах,
    # а место у потолка TERM_QUEUE_MAX они занимают и вытесняют настоящие
    # находки — на боевом проекте так висело 115 карточек. Тем же предикатом,
    # что ворота и _auto_verdict, — копия условий разошлась бы первой же
    # правкой политики. conflict не трогаем: его _auto_verdict отдаёт человеку
    # раньше любых проверок формы, и ловить расхождение заверенного перевода
    # с записью длиннее лимита — его работа. Решением человека закрытая так
    # карточка не становится (autoWrote в ключах — см. _human_decision).
    for c in queue:
        if c.get("status", "pending") != "pending" or c.get("kind") == "conflict":
            continue
        src = (c.get("src") or "").strip()
        tgt = (c.get("tgt") or "").strip()
        if not src or not tgt:
            continue
        why = _term_shape_reject(_auto_policy(_scope_of(c)[1]), src, tgt)
        if why:
            c.update({"status": "rejected", "autoWrote": False, "decidedAt": today,
                      "note": "снято разбором очереди: " + why})
            for k in ("sampleSrc", "sampleTgt"):
                c.pop(k, None)
            closed += 1
    return closed


def _queue_term_locked(kind, src, tgt, src_n, tgt_n, now, extra):
    # Область входит в ключ дедупликации: «договор → contract» в RU→EN и
    # «договор → Vertrag» в RU→DE — разные кандидаты, а не спор двух вариантов.
    scope = _scope_of(extra)
    donor = (f"{extra.get('project')}:{extra['segment']}"
             if extra.get("segment") and extra.get("project") else None)
    answered = None
    for c in _term_queue():
        if _scope_of(c) != scope:
            continue
        if c.get("kind") == kind and _cand_pair(c) == (src_n, tgt_n):
            c["hits"] = c.get("hits", 1) + 1
            c["at"] = now
            # Список сегментов-доноров, а не только счётчик: автоодобрение
            # считает согласие НЕЗАВИСИМЫХ сегментов, а hits растёт и от
            # повторного подтверждения одного и того же — так одна строка
            # накрутила бы себе доказательства.
            segs = c.setdefault("segments", _donor_ids(c))
            if donor and donor not in segs:
                segs.append(donor)
            if extra.get("via") == "confirmed":
                c["via"] = "confirmed"      # подтверждение человеком не понижается
            return c if c.get("status") == "pending" else None
        if answered is None and _term_answered(c, src_n):
            answered = c

    # ── Вопрос про этот термин уже отвечен ───────────────────────────
    # Дедупликация выше сверяет ПАРУ «термин + перевод» и вид карточки, и этого
    # мало: следующий прогон предлагает СВОЙ перевод того же термина, а находка
    # termcheck приходит другим видом. Пара получается новая — и утверждённый
    # человеком термин возвращался в очередь как нерешённый, каждым прогоном
    # заново. Спрашивать второй раз нечего: ответ есть.
    if answered is not None:
        answered["hits"] = answered.get("hits", 1) + 1
        answered["at"] = now
        # Молча глотать несогласие нельзя: в решении остаётся список того, что
        # прогоны предлагали взамен. Потолок — чтобы история одного термина
        # не росла в state.json без края.
        other = _norm_key(tgt)
        if other and other != _norm_key(answered.get("tgt")):
            again = answered.setdefault("reasked", [])
            if other not in again and len(again) < TERM_REASKED_MAX:
                again.append(other)
        return None

    # Запись уровня «приказ» — тот же ответ, только данный не через очередь
    # (ручная правка глоссария, выверенный справочник). Расхождение перевода
    # с ней — находка ремонта и строка в «Соответствии глоссарию», а не вопрос
    # «как переводить этот термин». Записи уровня «подсказка» сюда не попадают
    # намеренно: массовый автоимпорт как раз и ловится conflict-кандидатами
    # («задний → rear»), ради которых они заведены.
    if _hard_answer(_glossary_entry(src, scope)):
        return None

    # Последней — и только для видов, где пара ПРЕДЛАГАЕТСЯ из свободного
    # текста. Именно последней: выше растут `hits` у существующей карточки
    # и пишется `reasked` — список того, что прогоны предлагали взамен
    # решённого. Поставь проверку раньше — и система начала бы молча глотать
    # несогласие, чего делать нельзя.
    if kind in TERM_HARVEST_KINDS and not _looks_like_term(src, tgt):
        # Считаем И называем: отсев без следа неотличим от «модель ничего
        # не нашла», а платили за вызов одинаково. Тот же закон, что у списка
        # отсеянных надписей на картинках.
        _TERM_NOT_TERM[0] += 1
        print("[backend] не словарная запись, в очередь не идёт: %r → %r"
              % ((src or "")[:60], (tgt or "")[:60]), file=sys.stderr)
        return None

    # Форму, которую автоодобрение отвергнет ВСЕГДА (`_term_shape_reject`),
    # заворачиваем ЗДЕСЬ ЖЕ — на заведении НОВОЙ карточки, а не на входе
    # в `_queue_term`. Смысл ворот прежний (115 неодобряемых карточек длиннее
    # лимита стояли у потолка и вытесняли настоящие находки), а место — нет,
    # и по трём причинам:
    #   1) выше растут `hits` и пишется `reasked` — на входе ворота молча
    #      глотали несогласие с решёнными терминами, ровно то, что запрещено
    #      у `_looks_like_term` абзацем выше;
    #   2) `conflict` — исключение ПО ПОСТРОЕНИЮ: `_auto_verdict` отдаёт его
    #      человеку раньше любых проверок формы («нет готового перевода»),
    #      то есть «отвергнет всегда» про него неправда. А ловить расхождение
    #      заверенного перевода с записью длиннее лимита — его работа:
    #      в медицине лимит три слова, и «фиброзно-кавернозный туберкулёз
    #      лёгких» иначе не всплывал бы никогда;
    #   3) общий счётчик крутится под локом, как ему и положено.
    if kind != "conflict":
        shape_why = _term_shape_reject(_auto_policy(extra.get("domain")), src, tgt)
        if shape_why:
            _TERM_NOT_TERM[0] += 1
            print("[backend] %s, в очередь не идёт: %r → %r"
                  % (shape_why, (src or "")[:60], (tgt or "")[:60]), file=sys.stderr)
            return None

    cand = {"id": max((c.get("id", 0) for c in _term_queue()), default=0) + 1,
            "kind": kind, "src": (src or "").strip(), "tgt": (tgt or "").strip(),
            "status": "pending", "hits": 1, "at": now,
            "segments": [donor] if donor else []}
    cand.update(extra)
    return _queue_insert(cand)


def _queue_insert(cand: dict) -> dict:
    _term_queue().insert(0, cand)
    _trim_term_queue()
    return cand


def _tgt_has_term(target: str, term: str) -> bool:
    """Есть ли глоссарный перевод в готовом переводе. Терпимо к числу и
    артиклям: сверяем по началам слов, иначе «lymph nodes» не найдёт
    «lymph node» и очередь захлебнётся ложными конфликтами."""
    tn, mn = _norm_key(target), _norm_key(term)
    if not mn:
        return True
    if mn in tn:
        return True
    words = [w for w in mn.split() if len(w) > 3]
    if not words:
        return False
    tgt_words = tn.replace("(", " ").replace(")", " ").replace(",", " ").split()
    return all(any(x.startswith(w[:max(4, len(w) - 2)]) for x in tgt_words) for w in words)


def _gloss_by_src() -> dict:
    """(область, нормализованный термин) → запись. Отдельный индекс от
    _gloss_index: тот ищет вхождения в тексте, этот — точную запись по термину.
    Без него каждый короткий чистый сегмент внутри параллельного прогона гнал
    линейный проход по 10 000 записей."""
    global _GLOSS_BY_SRC
    idx = _GLOSS_BY_SRC
    if idx is not None:
        return idx
    with _GLOSS_INDEX_LOCK:
        if _GLOSS_BY_SRC is None:
            built: dict = {}
            for g in STATE.get("glossary", []):
                built.setdefault((_scope_of(g), _norm_key(g.get("src"))), g)
            _GLOSS_BY_SRC = built
        # Возвращаем локальную ссылку: параллельный _invalidate_gloss_index()
        # обнуляет глобал, и вызывающий получил бы None вместо словаря.
        return _GLOSS_BY_SRC


def _glossary_entry(src: str, scope: tuple) -> Optional[dict]:
    """Запись глоссария по термину В ПРЕДЕЛАХ области. Раньше поиск шёл по
    всему списку, и запись из чужой языковой пары выглядела как «уже есть»."""
    if len(scope) == 2:                   # прежний вид (пара, тематика)
        scope = (scope[0], scope[1], _current_tenant())
    return _gloss_by_src().get((scope, _norm_key(src)))


def _harvest_terms(seg: dict, project: dict, via: str = "confirmed") -> list:
    """Терминологические находки сегмента:

    conflict — глоссарий предлагал перевод, а в тексте его нет. Значит либо
               запись глоссария врёт (наш случай «задний → rear»), либо
               переводчик отступил осознанно. Решает человек.
    segment  — короткий сегмент без финальной точки сам по себе является
               терминологической парой.

    via="confirmed" — сегмент подтвердил человек, это сильнейшее доказательство.
    via="auto"      — сегмент прошёл back-check и termcheck чисто. Отсюда берём
                      ТОЛЬКО готовые пары: расхождения с глоссарием всё равно
                      идут человеку, а на большом проекте их сотни, и очередь
                      (её pending никогда не подрезается) распухла бы зря.
    """
    out = []
    source = seg.get("source", "")
    target = (seg.get("target") or "").strip()
    if not target:
        return out
    scope = _project_scope(project)
    meta = {"project": project["id"], "segment": seg["id"],
            "lang": scope[0], "domain": scope[1], "tenant": scope[2], "via": via}
    if via == "confirmed":
        hits, _tm = _get_context(source, project=project)
        for h in hits:
            if _tgt_has_term(target, h["tgt"]):
                continue
            c = _queue_term("conflict", h["src"], "",
                            wasTgt=h["tgt"], tier=_hit_tier(h), cat=h.get("cat", ""),
                            sampleSrc=source[:240], sampleTgt=target[:240], **meta)
            if c:
                out.append(c)
    words = source.strip().split()
    if 1 <= len(words) <= 4 and not source.strip().endswith((".", "!", "?", ":")) \
            and len(target.split()) <= 8:
        term_src = source.strip().strip(" .,;:")
        term_tgt = target.strip().strip(" .,;:")
        # Числа и обозначения («38,5 °C», «IFN-γ») терминами не бывают: пара
        # без букв — это не словарная запись, а строка документа.
        if re.search(r"[A-Za-zА-Яа-яЁё]{3}", term_src) and re.search(r"[A-Za-zА-Яа-яЁё]{3}", term_tgt):
            known = _glossary_entry(term_src, scope)
            if not (known and _norm_key(known.get("tgt")) == _norm_key(term_tgt)):
                c = _queue_term("segment", term_src, term_tgt,
                                cat=(known or {}).get("cat", ""),
                                wasTgt=(known or {}).get("tgt", ""), **meta)
                if c:
                    out.append(c)
    return out


def _check_stale(check: Optional[dict], target: str) -> bool:
    """Та же производная, что уходит клиенту в _segment_for_client: проверка
    относится к другому тексту, значит её результат уже ничего не значит.

    Хеш ОБРЕЗАННОГО текста — ровно так его и записывают back-check, termcheck
    и Medical QA. Сравнивая с необрезанным, мы объявляли бы устаревшей свежую
    проверку любого перевода с висящим пробелом, и прогон платил бы за неё
    заново."""
    return (check or {}).get("target_hash") != _text_hash((target or "").strip())


# Уровни находок termcheck, по которым РЕМОНТ имеет право переписывать текст.
# Кортеж отдаётся браузеру через /api/models: состав кнопки «запустить только
# ремонт» считается по нему, и литерал в .jsx разошёлся бы с сервером в тот же
# день, когда список меняют.
# Осторожно: это НЕ список «кому верить вообще». Спор с приказом глоссария
# (`_note_term_disputes`, корзина в /analysis) по-прежнему считается только по
# critical/major — там вопрос стоит иначе: «отменить решение человека или
# признать проверку ошибочной», и задавать его от стилистической придирки
# значит звать человека к ложной дилемме, а запись — к платной пересверке.
# По составу совпадает с TERMCHECK_SEVERITY, но это разные вопросы: там —
# порядок тяжести для вычисления худшей находки, здесь — право менять текст.
TERMCHECK_ACTIONABLE = ("critical", "major", "minor")
# Находки, которых достаточно, чтобы усомниться в записи глоссария.
TERMCHECK_DISPUTING = ("critical", "major")

# Причины отказа _machine_clean — константами, а не литералами по месту.
# По ним разбирается корзина в /analysis, и сравнение подстрокой русской фразы
# («не делался» in why) ломается от любой правки формулировки молча: сегмент
# просто уезжает в чужую корзину, и на экране это выглядит как правда.
CLEAN_NO_TARGET = "нет перевода"
CLEAN_NO_BACKCHECK = "back-check не делался или устарел"
CLEAN_NO_TERMCHECK = "termcheck не делался или устарел"
CLEAN_TERMCHECK_SKIP = "termcheck пропущен — проверять было нечего"
CLEAN_TERMCHECK_FINDINGS = "termcheck нашёл замечания"
CLEAN_REPAIRED = "текст переписан автоматическим ремонтом"
# «Проверка не делалась» — это «неизвестно», а не «плохо»: такие сегменты
# идут в свою строку «переведено, но не проверено».
CLEAN_UNCHECKED = (CLEAN_NO_BACKCHECK, CLEAN_NO_TERMCHECK, CLEAN_TERMCHECK_SKIP)
# Причина ТОЛЬКО для корзины /analysis, не для _machine_clean: донором глоссария
# такой сегмент всё равно не станет (балл ниже порога — это правда), но человеку
# «оценка ниже порога» здесь говорит неправду. На коротком оригинале лексическая
# мера не мерит (см. _judge_zone), балл 0 значит «нечем измерить», и поднять его
# вправе только судья — а судья по умолчанию выключен. Без своей строки эти
# сегменты остаются ровно там, откуда только что вынесли 306 починенных, и по
# той же причине: система знает, что цифра ничего не значит, а на экране
# показывает её как приговор.
CLEAN_LEX_BLIND = "балл не измерен: оригинал короче трёх содержательных слов, нужен судья"
# Текст написала ревизия, а судья back-check ставит major/critical: два
# оплаченных мнения разошлись. В ремонт это не идёт (`_repair_findings`,
# `_review_wrote`), и «оценка ниже порога» здесь говорит не то — человек
# видел бы балл 45 без намёка, что спорят ревизор и судья.
CLEAN_JUDGE_VS_REVIEW = "текст написала ревизия, судья не согласен — решает человек"
# Пара стояла в промпте перевода из терм-листа документа: согласие такого
# сегмента с другими — одно решение, повторённое трижды, а не три подтверждения
# (тот же закон, что у DUPLICATE/propagatedFrom). Иначе терм-лист утекал бы
# в глоссарий приказом через автоодобрение, мимо инварианта 8.
CLEAN_TERMLIST = "перевод подсказан терм-листом документа"


def _machine_clean(seg: dict, min_score: int) -> Optional[str]:
    """None, если с сегмента можно собирать терминологию без человека.
    Иначе — причина отказа (её показываем в разборе автоодобрения).

    Смысл проверки: перевод оценивал не тот прогон, который его делал, и не
    та модель. Обратный перевод идёт в другую сторону, termcheck смотрит
    только на целевой текст. Совпали оба — пара заслуживает доверия.
    """
    target = (seg.get("target") or "").strip()
    if not target:
        return CLEAN_NO_TARGET
    bc, tc = seg.get("backcheck"), seg.get("termcheck")
    if not bc or _check_stale(bc, target):
        return CLEAN_NO_BACKCHECK
    if bc.get("score") is None or bc["score"] < min_score:
        return f"back-check ниже {min_score}%"
    if not tc or _check_stale(tc, target):
        return CLEAN_NO_TERMCHECK
    if tc.get("model") == "skip":
        # «Нечего проверять» — это не «проверено и чисто». Иначе сегмент, где
        # перевод совпал с оригиналом, дарил глоссарию пару вида «X → X».
        return CLEAN_TERMCHECK_SKIP
    if tc.get("findings"):
        return CLEAN_TERMCHECK_FINDINGS
    # Сверка хеша обязательна: сегмент, который однажды чинили, а потом
    # перевели заново, к ремонту уже не относится. Без неё он навсегда
    # оставался бы «переписанным» и никогда не стал бы донором для глоссария.
    if (seg.get("repair") or {}).get("applied") and _repair_tried(seg):
        return CLEAN_REPAIRED
    # Ревизия — тот же случай, что ремонт, и по той же причине: текст написан
    # МАШИНОЙ по её собственному вердикту. Пусти такой сегмент в доноры — и
    # система начнёт заверять в глоссарии собственную правку, то есть закреплять
    # свои ошибки перевода правилом на весь документ (инвариант 8). Сверка хеша
    # обязательна ровно как у ремонта: сегмент, переведённый заново после
    # ревизии, к ней уже не относится.
    rv = seg.get("review") or {}
    if (rv.get("applied") and not _review_stale(seg)) or seg.get("route") == "REVIEW":
        # Маршрут REVIEW остаётся и после ПЕРЕревизии (смена стайл-шита или
        # версии вопросов пишет свежий вердикт с applied=False): текст всё
        # так же написала машина.
        return CLEAN_REPAIRED
    return None


def _harvest_if_clean(seg: dict, project: dict) -> list:
    """Сбор терминологии с машинно-чистого сегмента. Вызывается в конце
    back-check и termcheck: какой из двух прогонов отработает вторым, тот и
    увидит обе оценки. Подтверждённые сегменты сюда не попадают — их
    терминологию собирает confirm_segment с пометкой «подтвердил человек»."""
    if seg.get("status") == "confirmed":
        return []
    pol = _auto_policy(project.get("domain"))
    if _machine_clean(seg, pol["backcheck_min"]) is not None:
        return []
    return _harvest_terms(seg, project, via="auto")


def _identical_source_segments(project: dict, seg: dict) -> dict:
    """Сегменты с тем же исходником и другим переводом. Ключ нормализован:
    лишний пробел или регистр не должны решать, повтор это или нет."""
    key = _norm_key(seg.get("source"))
    target = (seg.get("target") or "").strip()
    pending, confirmed = [], []
    for s in project["segments"]:
        if s["id"] == seg["id"] or _norm_key(s.get("source")) != key:
            continue
        if (s.get("target") or "").strip() == target:
            continue
        (confirmed if s.get("status") == "confirmed" else pending).append(s["id"])
    return {"pending": pending, "confirmed": confirmed}


@app.post("/api/segments/{pid}/{sid}/confirm")
def confirm_segment(pid: int, sid: int):
    """Подтверждение = момент обучения: пара уходит в TM, термины — в очередь
    кандидатов, повторы исходника возвращаются клиенту предложением.
    Ничего чужого сами не переписываем: распространение — отдельная команда."""
    _guard_project_write(pid)
    seg = get_segment(pid, sid)
    project = get_project(pid)
    seg["status"] = "confirmed"
    # Отметка о человеке нужна именно здесь: `confirmed` в проекте есть и на
    # старых сегментах, которые так пометило точное совпадение с TM (сейчас
    # оно ставит `translated`, см. _replace_target). Автоодобрение опирается
    # на «подтвердил человек» — без отметки такая подстановка сходила бы
    # за подтверждение и накручивала бы доказательства сама себе.
    seg["confirmedBy"] = _actor_id()
    seg["confirmedAt"] = datetime.now().strftime("%Y-%m-%d %H:%M")
    # Роль на момент заверения — ЧАСТЬ следа: роль человеку потом меняют,
    # а вопрос «кто и в каком качестве подписал» задают про прошлое.
    if _actor_role():
        seg["confirmedRole"] = _actor_role()
    else:
        seg.pop("confirmedRole", None)
    seg.pop("unconfirmed", None)
    _audit("segment.confirm", project=pid, segment=sid)
    tm_action, candidates = None, []
    edit_harvest = None
    if (seg.get("target") or "").strip():
        tm_action = _tm_upsert(seg["source"], seg["target"], project)
        candidates = _harvest_terms(seg, project)
        # Исправленное ЧЕЛОВЕКОМ внутри сегмента — тем же моментом обучения:
        # законы и коды исхода описаны у самой функции.
        edit_harvest = _harvest_edited_terms(seg, project)
        candidates += edit_harvest.pop("cards")
    same = _identical_source_segments(project, seg)
    save_state(STATE)
    return {"ok": True, "segment": seg, "tm": tm_action, "propagate": same,
            "editHarvest": edit_harvest,
            "termCandidates": [{"id": c["id"], "kind": c["kind"], "src": c["src"],
                                "tgt": c.get("tgt", ""), "wasTgt": c.get("wasTgt", "")}
                               for c in candidates]}


class PropagateRequest(BaseModel):
    ids: Optional[List[int]] = None
    include_confirmed: bool = False


@app.post("/api/segments/{pid}/{sid}/propagate")
def propagate_segment(pid: int, sid: int, req: PropagateRequest = PropagateRequest()):
    """Разослать подтверждённый перевод в сегменты с идентичным исходником.
    Только по явной команде и по умолчанию мимо подтверждённых: молча
    переписать чужую правку — худшее, что может сделать CAT-инструмент.
    Затронутые сегменты получают статус translated, а не confirmed: подтвердить
    перевод может только человек, иначе автоподстановка сама себя заверяет."""
    _guard_project_write(pid)
    seg = get_segment(pid, sid)
    project = get_project(pid)
    target = (seg.get("target") or "").strip()
    if not target:
        raise HTTPException(400, "У сегмента нет перевода")
    same = _identical_source_segments(project, seg)
    allowed = set(same["pending"])
    if req.include_confirmed:
        allowed |= set(same["confirmed"])
    if req.ids is not None:
        allowed &= set(req.ids)
    changed = []
    for s in project["segments"]:
        if s["id"] not in allowed:
            continue
        # Через _replace_target: у подтверждённого сегмента снимается отметка
        # «подтвердил человек» и статус становится «требует проверки». Иначе
        # рассылка оставляла бы чужую подпись на подставленном ею тексте.
        s["prevTarget"] = s.get("target", "")      # ручной откат остаётся возможен
        _replace_target(s, target, PROVIDER_TM, "EXACT_TM")
        s["propagatedFrom"] = seg["id"]
        changed.append(s["id"])
    if changed:
        save_state(STATE)
    return {"ok": True, "changed": changed,
            "skippedConfirmed": [] if req.include_confirmed else same["confirmed"]}


# ─── Очередь кандидатов в глоссарий ──────────────────────────────────
@app.get("/api/term-queue")
def list_term_queue(status: str = "pending", limit: int = 200, offset: int = 0,
                    project: Optional[int] = None):
    """Кандидаты, отсортированные по частоте: сверху то, что мешает чаще всего.

    Вместе с карточками отдаём РАЗБОР: почему автоматика не берёт каждую и
    сколько таких же. Без него очередь на четыреста штук — стена одинаковых
    карточек, и человек не видит, что треть из них вообще не термины."""
    t = _current_tenant()
    items = [c for c in _term_queue() if _tenant_of(c) == t]
    counts = {}
    for c in items:
        st = c.get("status", "pending")
        counts[st] = counts.get(st, 0) + 1
    if status and status != "all":
        items = [c for c in items if c.get("status", "pending") == status]
    items = sorted(items, key=lambda c: (-c.get("hits", 1), -c.get("id", 0)))

    groups, reason_of = [], {}
    if project:
        # Область применяем ко ВСЕМУ ответу, а не только к разбору: иначе
        # значок показывает одно число, сумма групп другое, а карточки чужой
        # языковой пары висят без причины и пропадают под любым фильтром.
        scope = _project_scope(get_project(project))
        items = [c for c in items if _scope_of(c) == scope]
    if status == "pending" and items and project:
        project_obj = get_project(project)
        pol = _auto_policy(project_obj.get("domain"))
        ctx = _auto_context(items, pol)
        ctx["corpus"] = {}          # без сети: это список, а не прогон
        buckets: dict = {}
        for c in items:
            action, why = _auto_verdict(c, ctx)
            if action in (GLOSSARY_TIER_HARD, GLOSSARY_TIER_SOFT):
                key = "ready"
            elif action == "close":
                key = "closed"      # уже в глоссарии — не «ждёт человека»
            else:
                key = why or "—"
            reason_of[c["id"]] = key
            b = buckets.setdefault(key, {"reason": key, "count": 0, "ids": [],
                                         # Отклонять пачкой можно не всё: см. bulk
                                         "bulk": key not in ("ready", "closed")})
            b["count"] += 1
            b["ids"].append(c["id"])
        # Конфликт нельзя отклонить пачкой: у него нет своего перевода, и
        # отказ по нему заблокировал бы вопрос навсегда — см. bulk_decide_terms.
        for b in buckets.values():
            if b["bulk"] and all(not (c.get("tgt") or "").strip()
                                 for c in items if c["id"] in set(b["ids"])):
                b["bulk"] = False
        groups = sorted(buckets.values(), key=lambda x: -x["count"])
    out = items[offset:offset + limit]
    return {"total": len(items), "counts": counts,
            "groups": groups,
            "items": [{**c, "why": reason_of.get(c["id"])} for c in out]}


class BulkDecision(BaseModel):
    ids: List[int]
    action: str = "reject"      # только отклонение: см. ниже


@app.post("/api/term-queue/bulk")
def bulk_decide_terms(req: BulkDecision):
    """Массовое ОТКЛОНЕНИЕ пачки кандидатов.

    Массового одобрения здесь нет намеренно. Одобрение пишет правило для всех
    будущих текстов, и «одобрить всё, что видно» — это подпись не глядя под
    четырьмястами правилами: ровно то, от чего защищает вся остальная система.
    Массовое одобрение уже есть в другом месте и с доказательствами —
    /term-queue/auto-approve, который берёт только то, за что может поручиться.

    Отклонение безопасно: оно ничего не пишет в глоссарий, а лишь убирает
    из очереди то, что термином не является. Ошиблись — кандидат вернётся,
    как только пара встретится снова с другим переводом."""
    if req.action != "reject":
        raise HTTPException(400, "Массовым может быть только отклонение. "
                                 "Одобрение пачкой — /term-queue/auto-approve, "
                                 "оно берёт лишь то, что подтверждено.")
    want = set(req.ids)
    done, kept = [], []
    # Под локом: очередь одна, а прогон на сервере пишет в неё из рабочего
    # потока — проход с мутацией без лока мог бы пропустить часть указанных.
    with _TERM_QUEUE_LOCK:
        for c in _term_queue():
            if c.get("id") not in want or c.get("status", "pending") != "pending":
                continue
            if not (c.get("tgt") or "").strip():
                # Кандидат без своего перевода (конфликт с глоссарием).
                # Дедупликация помнит его по паре «термин + пустой перевод»,
                # поэтому отказ закрыл бы этот вопрос НАВСЕГДА — а обещание
                # «встретится снова — спросим заново» тут выполнить нечем.
                # Такие решает человек по одному, вписав верный вариант.
                kept.append(c["id"])
                continue
            _mark_decided(c, "rejected", note="отклонён пачкой")
            done.append(c["id"])
    save_state(STATE)
    return {"ok": True, "rejected": done, "count": len(done),
            "kept": kept,
            "keptWhy": ("у этих кандидатов нет своего перевода — отказ закрыл бы "
                        "вопрос навсегда, их решают по одному") if kept else ""}


class TermDecision(BaseModel):
    src: Optional[str] = None
    tgt: Optional[str] = None
    cat: Optional[str] = None
    # «Знаю о замечании, всё равно одобряю». Без него одобрение пары, которую
    # судья считает негодной, возвращает предупреждение и НЕ пишет в глоссарий:
    # предупредить после записи — то же самое, что не предупредить.
    confirm: bool = False


def _mark_decided(cand: dict, status: str, src: str = None, tgt: str = None, **extra):
    """Решение по кандидату. Прежнюю пару запоминаем ДО правки: по ней работает
    дедупликация (см. _cand_pair). Примеры сегментов снимаем — решённый кандидат
    живёт дальше не как карточка, а как память «этот вопрос уже задавали»."""
    if src is not None and _norm_key(src) != _norm_key(cand.get("src")):
        cand.setdefault("origSrc", cand.get("src", ""))
        cand["src"] = src
    if tgt is not None and _norm_key(tgt) != _norm_key(cand.get("tgt")):
        cand.setdefault("origTgt", cand.get("tgt", ""))
        cand["tgt"] = tgt
    cand["status"] = status
    # Кто решил — идентификатор пользователя; прежнее "human" (без сессии)
    # законно навсегда, как у `confirmedBy`. Имя и роль — рядом: карточка
    # показывает их без похода за списком пользователей.
    cand["decidedBy"] = _actor_id()
    cand["decidedAt"] = datetime.now().strftime("%Y-%m-%d %H:%M")
    who = _actor()
    if who:
        cand["decidedName"] = who.get("name") or who["login"]
        cand["decidedRole"] = _actor_role() or who.get("role", "translator")
    else:
        cand.pop("decidedName", None)
        cand.pop("decidedRole", None)
    cand.update(extra)
    for k in ("sampleSrc", "sampleTgt", "wasTgtLeft"):
        cand.pop(k, None)
    return cand


def _human_decision(c: dict) -> bool:
    """Решение принял человек, а не автоодобрение. Такие живут в очереди дольше:
    именно они не дают спросить про уже решённый термин второй раз."""
    by = c.get("decidedBy")
    if by == "human" or isinstance(by, int):
        return True
    return not (c.get("autoBatch") or c.get("autoTier") or "autoWrote" in c)


def _close_same_term(decided: dict, scope: tuple) -> list:
    """Закрыть остальные ОЖИДАЮЩИЕ карточки про тот же термин в той же области.
    Человек ответил не на карточку, а на вопрос «как переводить этот термин»:
    оставлять рядом второй вопрос про него же — значит показать после
    обновления страницы уже решённое как нерешённое."""
    key = _cand_pair(decided)[0]
    tgt = _norm_key(decided.get("tgt"))
    closed = []
    # Под локом: очередь одна, а прогон на сервере пишет в неё из рабочего
    # потока. Проход с мутацией без лока мог бы разойтись с _queue_term.
    with _TERM_QUEUE_LOCK:
        return _close_same_term_locked(decided, scope, key, tgt, closed)


def _close_same_term_locked(decided, scope, key, tgt, closed):
    for c in _term_queue():
        if c is decided or c.get("status", "pending") != "pending":
            continue
        if _cand_pair(c)[0] != key or _scope_of(c) != scope:
            continue
        rival = _norm_key(c.get("tgt"))
        if rival and rival != tgt:
            # Карточка с ДРУГИМ переводом — не повтор вопроса, а проигравший
            # вариант: в глоссарий она не попала, и помечать её «одобрено»
            # значило бы соврать. Пишем «отклонён» и называем победителя —
            # если выбор окажется неверным, запись глоссария правится руками.
            _mark_decided(c, "rejected", decidedWith=decided["id"],
                          note="не выбран: для термина одобрен вариант «%s»"
                               % decided.get("tgt", ""))
        else:
            _mark_decided(c, "approved", decidedWith=decided["id"],
                          note="решено вместе с кандидатом #%d" % decided["id"])
        closed.append(c["id"])
    return closed


@app.post("/api/term-queue/{cid}/approve")
def approve_term_candidate(cid: int, req: TermDecision = TermDecision()):
    """Одобренный кандидат становится проверенной записью глоссария — только
    такие уходят в промпт жёстким правилом.

    Перед записью пара проходит ту же сверку судьёй, что и машинное
    автоодобрение. Раньше этот путь не проверялся ВООБЩЕ, и выходило наоборот:
    решения машины проверялись строже решений человека — который целевого языка
    может не знать и оценить пару не в состоянии. Замечание возвращается ДО
    записи (`written: false`), одобрить вопреки ему можно полем `confirm`."""
    _audit("term.approve", candidate=cid)
    cand = next((c for c in _term_queue() if c.get("id") == cid
                 and _tenant_of(c) == _current_tenant()), None)
    if not cand:
        raise HTTPException(404, "Кандидат не найден")
    src = (req.src or cand.get("src") or "").strip()
    tgt = (req.tgt or cand.get("tgt") or "").strip()
    if not src or not tgt:
        raise HTTPException(400, "Нужен и термин, и перевод. У кандидата-конфликта "
                                 "перевод пуст: впишите верный вариант.")
    # Категорию берём от кандидата: жёсткий медицинский дефолт в юридическом
    # проекте помечал договоры как «Disease».
    cat = req.cat or cand.get("cat") or "Term"
    today = datetime.now().strftime("%Y-%m-%d")
    scope = _scope_of(cand)

    # ── Тот же вопрос, что и машине ──────────────────────────────────
    # Одобрение пишет ПРАВИЛО на все будущие тексты, и человек, не знающий
    # целевого языка, оценить пару не может. Спрашиваем судью: то же ли это
    # понятие и годится ли пара правилом. «Не знаю» (нет ключа, сбой вызова)
    # не блокирует — тот же закон, что у корпуса.
    verdict = None
    if not req.confirm and os.environ.get("OPENAI_API_KEY"):
        got = _openai_meaning([(src, tgt)], scope)
        verdict = (got or {}).get((_norm_key(src), _norm_key(tgt)))
    if verdict and (verdict.get("same") is False or verdict.get("rule") is False):
        # В глоссарий НИЧЕГО не пишем и кандидата не решаем: человек ещё
        # не видел замечания, а значит и решения пока нет.
        return {"ok": True, "written": False, "warning": {
            "src": src, "tgt": tgt,
            "kind": "meaning" if verdict.get("same") is False else "rule",
            "back": verdict.get("back") or "",
            "why": verdict.get("why") or "",
            "text": ("перевод означает другое: " + (verdict.get("back") or "иное понятие"))
                    if verdict.get("same") is False
                    else ("правилом на весь документ не годится: "
                          + (verdict.get("why") or "зависит от контекста"))}}
    existing = _glossary_entry(src, scope)
    # Вердикт судьи кладём на запись: аудит глоссария читает его как свой
    # (`_meaning_stale` сверяет отпечаток пары) и не переспрашивает платно
    # то, что только что спросили здесь.
    mark = ({"meaning": {"same": bool(verdict.get("same")), "rule": verdict.get("rule"),
                         "back": verdict.get("back") or "", "why": verdict.get("why") or "",
                         "pair": _text_hash(_norm_key(src) + "||" + _norm_key(tgt)),
                         # Счётчик жалоб — НЫНЕШНИЙ, а не ноль: иначе вердикт
                         # рождается устаревшим, и аудит переспрашивает то,
                         # за что уже заплатили секунду назад.
                         "disputed": int((existing or {}).get("disputed") or 0),
                         "v": MEANING_VERSION,
                         "model": _resolve_model(JUDGE_DEFAULT_MODEL)["id"],
                         "at": today}}
            if verdict and verdict.get("same") is not None else {})
    if existing:
        existing.update({"tgt": tgt, "cat": cat, "conf": "high", "tier": GLOSSARY_TIER_HARD,
                         "note": "уточнено вручную " + today, "updated": today,
                         "prevTgt": existing.get("tgt", ""), **mark,
                         **_signed_field("approve")})
        _clear_auto_marks(existing)
    else:
        STATE["glossary"].insert(0, {"src": src, "tgt": tgt, "cat": cat, "freq": 1,
                                     "conf": "high", "note": "", "tier": GLOSSARY_TIER_HARD,
                                     "lang": scope[0], "domain": scope[1], "tenant": scope[2],
                                     "updated": today, **mark, **_signed_field("approve"),
                                     "origin": "confirmed:" + str(cand.get("segment", ""))})
    # Пара запоминается до правки, остальные карточки про этот же термин
    # закрываются: иначе одобренный термин всплывал бы снова — и как сосед
    # по очереди, и как заново созданный кандидат на следующем сборе.
    _mark_decided(cand, "approved", src=src, tgt=tgt)
    closed = _close_same_term(cand, scope)
    _invalidate_gloss_index()
    save_state(STATE)
    return {"ok": True, "written": True, "candidate": cand,
            "replaced": bool(existing), "closed": closed}


class ExplainRequest(BaseModel):
    variants: Optional[List[str]] = None   # что сравниваем; по умолчанию — все из очереди
    # Вариант, который обязан попасть в сравнение: то, что человек напечатал
    # в поле карточки. Не список, а добавка — иначе черновик вытеснял бы
    # остальные варианты, и сравнивать стало бы не с чем.
    include: Optional[str] = None
    model: Optional[str] = None


@app.post("/api/term-queue/{cid}/explain")
def explain_term_variants(cid: int, req: ExplainRequest = ExplainRequest()):
    """Объяснить варианты перевода НА ЯЗЫКЕ ОРИГИНАЛА.

    Ключевая мысль: пользователь может не знать целевого языка — и тогда
    вопрос «какой перевод верный» для него бессмыслен. Но вопрос «какое из
    двух значений вы имели в виду» он понимает, если оба значения написаны
    по-русски. Поэтому по каждому варианту спрашиваем у модели:
      обратный перевод — как этот вариант читается носителем целевого языка;
      определение     — что этот термин означает, одной строкой на языке оригинала;
      область         — где он употребляется.
    Человек сравнивает РУССКОЕ с РУССКИМ и выбирает смысл, а не строку.

    Вызов платный, поэтому только по кнопке и только на конкретную карточку."""
    if not os.environ.get("OPENAI_API_KEY"):
        raise HTTPException(503, "Разбор вариантов требует ключ OpenAI")
    cand = next((c for c in _term_queue() if c.get("id") == cid
                 and _tenant_of(c) == _current_tenant()), None)
    if not cand:
        raise HTTPException(404, "Кандидат не найден")
    scope = _scope_of(cand)
    src_lang, tgt_lang = (authorities_mod.source_lang(scope[0]),
                          authorities_mod.target_lang(scope[0])) if authorities_mod \
        else (scope[0].split("→")[0], scope[0].split("→")[-1])
    term = (cand.get("src") or "").strip()

    # Что сравниваем: присланные варианты, иначе всё, что предлагает очередь
    # и справочники по этому термину. Пустой перевод у конфликта не вариант.
    variants = [v.strip() for v in (req.variants or []) if v and v.strip()]
    if variants and (req.include or "").strip():
        # include — «этот вариант обязан участвовать». Игнорировать его при
        # заданном списке значило бы потерять именно то, что человек напечатал.
        if _norm_key(req.include) not in {_norm_key(v) for v in variants}:
            variants.insert(0, req.include.strip())
    if not variants:
        seen = set()
        # Первым идёт то, что человек напечатал (или, если не печатал, перевод
        # самой карточки): обрезая список до шести, нельзя выбросить именно тот
        # вариант, ради которого нажали кнопку.
        for first in ((req.include or "").strip(), (cand.get("tgt") or "").strip()):
            if first and _norm_key(first) not in seen:
                variants.append(first)
                seen.add(_norm_key(first))
        for c in _term_queue():
            if (c.get("status", "pending") == "pending" and _scope_of(c) == scope
                    and _norm_key(c.get("src")) == _norm_key(term) and (c.get("tgt") or "").strip()):
                if _norm_key(c["tgt"]) not in seen:
                    seen.add(_norm_key(c["tgt"]))
                    variants.append(c["tgt"].strip())
        # Запись глоссария — раньше очередных вариантов: её уже утверждали,
        # и вытеснять её потолком в шесть штук нельзя.
        known = _glossary_entry(term, scope)
        if known and known.get("tgt") and _norm_key(known["tgt"]) not in seen:
            seen.add(_norm_key(known["tgt"]))
            variants.insert(len(variants) and 1 or 0, known["tgt"])
        for v in _authority_suggests(term, scope):
            if _norm_key(v) not in seen:
                seen.add(_norm_key(v))
                variants.append(v)
    # Молчаливых потолков не бывает: сколько не влезло — скажем в ответе.
    dropped = max(0, len(variants) - 6)
    variants = variants[:6]
    if not variants:
        raise HTTPException(400, "Нечего сравнивать: у термина нет ни одного варианта перевода")

    dom = _resolve_domain(scope[1])
    system = (
        f"You are a {dom['expert']}. The user does NOT speak {tgt_lang} and must choose "
        f"between candidate {tgt_lang} translations of a {src_lang} term by MEANING.\n\n"
        f"For each candidate return, WRITTEN IN {src_lang} (this is essential — the user "
        f"reads only {src_lang}):\n"
        f"  back — how the candidate reads to a native {tgt_lang} speaker, translated back "
        f"literally into {src_lang};\n"
        f"  meaning — what the candidate actually denotes, one short sentence in {src_lang};\n"
        f"  usage — where this wording is used (register, field), a few words in {src_lang};\n"
        f"  same — true if it denotes exactly the same concept as the {src_lang} term.\n"
        + MEANING_TRAPS + "\n\n"
        'Return ONLY JSON: {"variants":[{"tgt":"...","back":"...","meaning":"...",'
        '"usage":"...","same":true}]}. No commentary.'
    )
    body = (f"{src_lang} term: {term}\n"
            + (f"Context sentence: {cand.get('sampleSrc')}\n" if cand.get("sampleSrc") else "")
            + f"Candidate {tgt_lang} translations:\n"
            + "\n".join("  - " + v for v in variants))
    try:
        import openai
        mdl = _resolve_model(req.model or JUDGE_DEFAULT_MODEL)
        client = openai.OpenAI(api_key=os.environ.get("OPENAI_API_KEY"), timeout=90, max_retries=1)
        extra = ({"max_completion_tokens": 2048} if mdl["api"] == "modern"
                 else {"max_tokens": 900, "temperature": 0})
        resp = client.chat.completions.create(
            model=mdl["id"], response_format={"type": "json_object"},
            messages=[{"role": "system", "content": system},
                      {"role": "user", "content": body}], **extra)
        _note_usage("terms", mdl["id"], resp)
        raw = (resp.choices[0].message.content or "").strip()
        if not raw:
            # Пустой ответ (модель израсходовала лимит на рассуждения) — это
            # отказ. Отдать его как «разобрано, но всё пусто» значит показать
            # человеку прочерки там, где он ждёт объяснения смысла.
            raise ValueError("модель вернула пустой ответ")
        data = json.loads(raw)
        by_tgt = {_norm_key(v.get("tgt")): v
                  for v in (data.get("variants") or []) if isinstance(v, dict)}
    except Exception as e:
        print(f"[backend] explain term #{cid}: {e}", file=sys.stderr)
        raise HTTPException(502, "Модель не разобрала варианты. Попробуйте ещё раз.")

    # Корпус по всем вариантам разом: шесть последовательных внешних запросов
    # с троттлингом под лимит источника — это ещё минута поверх уже оплаченного
    # вызова модели, и человек всё это время смотрит на «Разбираем…».
    att_by = dict(zip([_norm_key(v) for v in variants],
                      _run_parallel(variants, lambda v: _corpus_check(v, scope))))
    out = []
    for v in variants:
        got = by_tgt.get(_norm_key(v))
        # Модель про этот вариант ничего не сказала — так и передаём (None).
        # «Не знаю» нельзя показывать как «иное понятие»: человек, который по
        # условию не читает целевой язык, отвергнет верный перевод.
        answered = isinstance(got, dict)
        got = got if answered else {}
        # Корпус тут же: «сколько раз этот термин вообще встречается в языке» —
        # цифра, которую человек оценит без знания языка.
        att = att_by.get(_norm_key(v))
        out.append({
            "tgt": v,
            "back": got.get("back") or "",
            "meaning": got.get("meaning") or "",
            "usage": got.get("usage") or "",
            "same": (bool(got.get("same")) if answered and "same" in got else None),
            "corpus": ({"hits": att["hits"], "label": att["label"]} if att else None),
            "authority": _authority_match(term, v, scope),
        })
    return {"ok": True, "term": term, "variants": out, "dropped": dropped,
            "model": _resolve_model(req.model or JUDGE_DEFAULT_MODEL)["id"]}


@app.post("/api/term-queue/{cid}/reject")
def reject_term_candidate(cid: int):
    _audit("term.reject", candidate=cid)
    cand = next((c for c in _term_queue() if c.get("id") == cid
                 and _tenant_of(c) == _current_tenant()), None)
    if not cand:
        raise HTTPException(404, "Кандидат не найден")
    # Соседей не трогаем: «эта пара неверна» — не то же самое, что «с термином
    # разобрались». Другой вариант перевода того же термина остаётся вопросом.
    _mark_decided(cand, "rejected")
    save_state(STATE)
    return {"ok": True, "candidate": cand}


# ─── Автоодобрение однозначных кандидатов ────────────────────────────
# Смысл: человеку, далёкому от переводов, не нужно щёлкать сотни карточек.
# Правила опираются ТОЛЬКО на сигналы, не зависящие от языка и тематики:
#   1. единственность варианта — у термина ровно один перевод в очереди;
#   2. согласие независимых сегментов — одна и та же пара пришла из разных мест;
#   3. оценки ЧУЖИХ прогонов — back-check идёт в обратную сторону, termcheck
#      смотрит только на целевой текст. Мнение того вызова, который сделал
#      перевод, доказательством не считается.
#
# Ключевое отличие от ручного одобрения: автоматика пишет в tier "auto" —
# подсказку, которую модель вправе игнорировать. Приказом ("verified") запись
# становится от человека либо от трёх независимых подтверждений, а в медицине,
# фарме и юриспруденции — только от человека: там цена приказа выше.
AUTO_POLICY_VERSION = "v2"          # v2: внешние справочники и корпус
AUTO_BATCH_HISTORY = 20
# Потолок корпусных запросов на один разбор очереди. Каждый — внешний HTTP;
# 200 штук по шесть секунд таймаута в шесть потоков — это минуты, а не часы.
AUTO_CORPUS_MAX = int(os.environ.get("AUTO_CORPUS_MAX", "200"))

AUTO_APPROVE_DEFAULT = {
    "auto_min_segments": 2,        # независимых машинно-чистых сегментов на подсказку
    "verified_min_segments": 3,    # ... на приказ
    "backcheck_min": 90,           # ниже этого сегмент-донор не считается чистым
    "allow_verified": True,
    "max_src_words": 3,
    "max_tgt_words": 6,
}

AUTO_APPROVE_BY_DOMAIN = {
    "medical": {"allow_verified": False},
    "pharma": {"allow_verified": False},
    "legal": {"allow_verified": False},
}


# ─── Внешние источники приказов ──────────────────────────────────────
# Справочники лежат файлами в data/authorities/*.tsv и подхватываются при
# старте: добавить источник — значит положить файл, а не править код.
# Два каталога, и это не дублирование:
#   authority_data/ — источники, едущие в git вместе с кодом (общие для всех);
#   data/authorities/ — то, что добавил владелец сервера (в git не хранится,
#                       как и все данные в data/).
AUTHORITY_DIRS = [ROOT / "backend" / "authority_data", DATA_DIR / "authorities"]
_DICTIONARIES: list = []


def _load_authorities():
    global _DICTIONARIES
    if not authorities_mod:
        return
    loaded = []
    for d in AUTHORITY_DIRS:
        try:
            loaded.extend(authorities_mod.load_dictionaries(d))
        except Exception as e:                               # pragma: no cover
            print(f"[backend] справочники из {d} не загружены: {e}", file=sys.stderr)
    _DICTIONARIES = loaded


_load_authorities()      # ОБЯЗАТЕЛЬНО при импорте: без этого вызова весь путь
                         # «приказ от справочника» мёртв, и медицина с фармой
                         # навсегда остаются без приказов, кроме человеческих.


def _authority_match(src: str, tgt: str, scope: tuple) -> Optional[dict]:
    """Совпадение пары с отраслевым справочником для ЭТОЙ области и пары языков.

    Это и есть приказ без человека: норму зафиксировал тот, кто в предметной
    области разбирается. Справочник чужой пары языков или чужой области не
    подходит никогда — выдуманное подтверждение хуже отсутствующего."""
    best = None
    for d in _DICTIONARIES:
        if d.covers(scope[0], scope[1]) and d.match(src, tgt):
            hit = {"id": d.id, "label": d.label, "tier": getattr(d, "tier", "verified")}
            if hit["tier"] == GLOSSARY_TIER_HARD:
                return hit          # выверенный источник — дальше искать нечего
            best = best or hit      # краудсорсный запоминаем, вдруг найдётся лучше
    return best


def _authority_suggests(src: str, scope: tuple) -> list:
    """Что справочники предлагают для термина — показываем человеку рядом
    с кандидатом: «не совпало» без «а как правильно» бесполезно."""
    out = []
    for d in _DICTIONARIES:
        if d.covers(scope[0], scope[1]):
            out.extend(d.suggest(src))
    return sorted(set(out))


def _corpus_check(tgt: str, scope: tuple) -> Optional[dict]:
    """Живёт ли перевод в целевом языке. None — «проверить нечем», и это
    НЕ отрицательный ответ: молчащий источник не должен ни одобрять, ни
    блокировать."""
    if not authorities_mod:
        return None
    try:
        return authorities_mod.attested(tgt, authorities_mod.target_lang(scope[0]), scope[1])
    except Exception as e:                                   # pragma: no cover
        print(f"[backend] корпусная проверка не удалась: {e}", file=sys.stderr)
        return None


def _authority_sources(scope: tuple) -> dict:
    """Чем в этой области и паре языков вообще есть чем проверять. Нужно
    интерфейсу: разница в качестве между парами языков должна быть названа,
    а не обнаружена пользователем на своих текстах."""
    dicts = [{"id": d.id, "label": d.label, "terms": len(d.pairs),
              "tier": getattr(d, "tier", "verified")}
             for d in _DICTIONARIES if d.covers(scope[0], scope[1])]
    corpus = None
    if authorities_mod:
        c = authorities_mod.corpus_for(authorities_mod.target_lang(scope[0]), scope[1])
        if c and authorities_mod.corpus_available(authorities_mod.target_lang(scope[0]), scope[1]):
            corpus = {"id": c["id"], "label": c["label"]}
    return {"dictionaries": dicts, "corpus": corpus}


def _auto_policy(domain_id: Optional[str]) -> dict:
    pol = dict(AUTO_APPROVE_DEFAULT)
    dom = _resolve_domain(domain_id)
    pol.update(AUTO_APPROVE_BY_DOMAIN.get(dom["id"], {}))
    if dom.get("custom") and dom.get("strict", True):
        pol["allow_verified"] = False
    return pol


# Поздний проход по очереди — ПОСЛЕ определения _auto_policy: разбор формы
# внутри миграции читает лимиты слов той же политикой, что и _auto_verdict.
_TERM_QUEUE_MIGRATED = _migrate_term_queue(STATE)
_migrate_ui_lang()
# Занятость очереди — в журнал при старте. О потолке узнавали только из строки
# «выброшено N», то есть уже после потери находок.
try:
    _q = _term_queue()
    _q_pending = sum(1 for c in _q if c.get("status", "pending") == "pending")
    print(f"[backend] очередь кандидатов: {len(_q)} из {TERM_QUEUE_MAX} "
          f"(ожидают решения: {_q_pending})", file=sys.stderr)
except Exception as _e:                                      # pragma: no cover
    print(f"[backend] очередь кандидатов: размер не посчитан: {_e}", file=sys.stderr)
if _TERM_QUEUE_MIGRATED:
    # Молча укоротить очередь на сотню карточек нельзя: человек увидит другое
    # число и должен знать, что это разбор старого хвоста, а не потеря данных.
    print(f"[backend] очередь терминов: снято {_TERM_QUEUE_MIGRATED} карточек "
          f"про термины, по которым решение уже есть либо которые политика "
          f"не одобрит никогда", file=sys.stderr)
    save_state(STATE)


def _donor_ids(cand: dict) -> list:
    """Доноры кандидата как «проект:сегмент». Номера сегментов уникальны только
    внутри проекта — без префикса сегмент #12 из другого проекта сошёл бы за
    того же донора и накрутил бы доказательства."""
    ids = cand.get("segments")
    if ids:
        return list(dict.fromkeys(ids))
    if cand.get("segment") and cand.get("project"):
        return [f"{cand['project']}:{cand['segment']}"]
    return []


def _auto_context(pending: list, pol: dict) -> dict:
    """Индексы, собранные ОДИН раз на прогон. Без них разбор 2000 кандидатов
    гонял бы линейный поиск по 10 000 записей глоссария и по всем сегментам
    проекта на каждого — единственный воркер вставал бы на секунды."""
    # Только СВОЯ организация: иначе чужой сегмент становится донором
    # перевода и глоссария — утечка содержания, а не прав.
    segs = {(p["id"], s["id"]): s for p in _tenant_projects() for s in p["segments"]}
    gloss: dict = {}
    for g in STATE["glossary"]:
        gloss.setdefault((_scope_of(g), _norm_key(g.get("src"))), g)
    return {"variants": _auto_variants(pending), "segs": segs, "gloss": gloss, "pol": pol}


def _donor_quality(cand: dict, ctx: dict) -> tuple:
    """(годных доноров, был ли подтверждённый человеком, сколько среди них
    РАЗНЫХ исходников, причина отказа).

    Донор годится, если сегмент подтвердил человек либо он прошёл back-check и
    termcheck чисто. Копии чужого решения донорами не считаются вовсе, а число
    разных исходников отделяет «три независимых сегмента» от «один и тот же
    заголовок, переведённый трижды»: во втором случае согласие ничего не
    доказывает, это одна и та же строка."""
    good, confirmed, why = 0, False, None
    sources = set()
    for ref in _donor_ids(cand):
        try:
            pid, sid = (int(x) for x in str(ref).split(":", 1))
        except ValueError:
            continue
        seg = ctx["segs"].get((pid, sid))
        if seg is None:
            why = why or "сегмент-источник удалён"
            continue
        if seg.get("route") == "DUPLICATE" or seg.get("propagatedFrom"):
            why = why or "перевод скопирован с другого сегмента"
            continue
        if _norm_key(cand.get("tgt")) in {_norm_key(t) for t in (seg.get("docTerms") or [])}:
            why = why or CLEAN_TERMLIST
            continue
        if seg.get("status") == "confirmed" and _confirmed_by_human(seg):
            good += 1
            confirmed = True
            sources.add(_norm_key(seg.get("source")))
            continue
        # Подтверждён, но не человеком (подстановка из TM) — обычный кандидат:
        # решают проверки, а не статус.
        reason = _machine_clean(seg, ctx["pol"]["backcheck_min"])
        if reason is None:
            good += 1
            sources.add(_norm_key(seg.get("source")))
        else:
            why = why or reason
    return good, confirmed, len(sources), why


def _auto_variants(pending: list) -> dict:
    """(область, термин) → множество различных переводов среди кандидатов.
    Два варианта — это и есть определение неоднозначности: такой термин
    автоматика не трогает вообще."""
    out: dict = {}
    for c in pending:
        tgt = (c.get("tgt") or "").strip()
        if not tgt:
            continue
        out.setdefault((_scope_of(c), _norm_key(c.get("src"))), set()).add(_norm_key(tgt))
    return out


def _auto_verdict(cand: dict, ctx: dict) -> tuple:
    """(действие, причина). Действие: "verified" | "auto" | "close" | "wait" | None.
    "close" — пара уже есть в глоссарии слово в слово: кандидата закрываем,
    глоссарий не трогаем. "wait" — решать НЕЧЕГО, не хватает ДАННЫХ (доноров,
    чистых проверок): следующие прогоны могут дорешать сами, и звать человека
    сюда — звать его к работе, которой нет. Для одобрения оба не годятся
    одинаково (кандидат остаётся ждать), различие читает только разбор:
    на боевом проекте 412 карточек из 684 «ждущих человека» ждали на самом
    деле данных, и число пугало без причины."""
    pol = ctx["pol"]
    src = (cand.get("src") or "").strip()
    tgt = (cand.get("tgt") or "").strip()
    if cand.get("kind") == "conflict" or not tgt:
        return None, "нет готового перевода — решает человек"
    if not src:
        return None, "пустой термин"
    # Форма — общим предикатом с воротами очереди (_term_shape_reject):
    # копия этих условий уже расходилась и копила неодобряемое у потолка.
    shape_why = _term_shape_reject(pol, src, tgt)
    if shape_why:
        return None, shape_why

    scope = _scope_of(cand)
    known = ctx["gloss"].get((scope, _norm_key(src)))
    if known:
        if _norm_key(known.get("tgt")) == _norm_key(tgt):
            return "close", "уже в глоссарии"
        if _hit_tier(known) == GLOSSARY_TIER_HARD:
            return None, "спорит с проверенной записью глоссария"
        # Запись, занятую прошлой пачкой, писать поверх МОЖНО: _auto_write
        # сохраняет её цепочку отката (prev* не затираются, autoCreated
        # наследуется), поэтому откат НОВОЙ пачки возвращает состояние до
        # СТАРОЙ, а не чужой машинный вариант. Прежний отказ «сначала
        # откатите пачку #N» копил неразрешимые карточки: на боевом проекте
        # 36 кандидатов висели за шестью пачками, откатывать которые никто
        # не собирался. Старой пачке запись больше не принадлежит — её откат
        # такие называет числом superseded.
        # Запись уровня "auto" (массовый импорт) — как раз то, что автоодобрение
        # и должно чинить: у нас есть доказательства, у неё их не было.

    # ── Внешний справочник: приказ без человека ──────────────────────
    # Совпадение с отраслевым справочником — не мнение системы о собственном
    # переводе, а зафиксированная норма. Это единственное, что даёт verified
    # там, где allow_verified выключен (медицина, фарма, юриспруденция):
    # запрет касается САМООДОБРЕНИЯ, а справочник находится вне контура.
    # Проверяется РАНЬШЕ спора вариантов: когда у термина два перевода,
    # справочник как раз и говорит, какой из них норма.
    auth = _authority_match(src, tgt, scope)
    if auth:
        # Справочник может держать несколько норм на один термин. Если под них
        # подходит сразу два кандидата из очереди — он не разрешает спор, а
        # участвует в нём: решает человек. Иначе первый по порядку очереди
        # получал бы приказ, а второй молча закрывался как решённый.
        rivals = [v for v in ctx["variants"].get((scope, _norm_key(src)), set())
                  if _authority_match(src, v, scope)]
        if len(rivals) > 1:
            return None, "справочник допускает несколько вариантов — решает человек"
        # Тумблер «только подсказки» — это про ЭТОТ запуск, а не про политику
        # области: человек попросил ничего не поднимать до приказа, и справочник
        # не повод его переспрашивать.
        if auth["tier"] == GLOSSARY_TIER_HARD and not pol.get("cap_soft"):
            return GLOSSARY_TIER_HARD, "совпадает со справочником: " + auth["label"]
        if auth["tier"] == GLOSSARY_TIER_HARD:
            return GLOSSARY_TIER_SOFT, ("совпадает со справочником (%s), "
                                        "но выбран режим «только подсказки»" % auth["label"])
        # Краудсорсный источник приказывать в одиночку не вправе: выборочная
        # проверка находит в таких неверные нормы. Он идёт ГОЛОСОМ — дальше по
        # коду его учитывают вместе с согласием сегментов и корпусом.

    if len(ctx["variants"].get((scope, _norm_key(src)), set())) > 1:
        return None, "у термина несколько вариантов перевода"

    # ── Корпус целевого языка: вето на кальки ────────────────────────
    # Корпус не подтверждает перевод, он отвечает на другой вопрос — есть ли
    # такой термин в языке. Ноль вхождений («rear cyclitis») означает, что мы
    # изобрели слово, и никакое согласие доноров этого не искупает: они все
    # сделаны одной и той же моделью с одной и той же кальки.
    corpus = ctx.get("corpus", {}).get((scope, _norm_key(tgt)))
    if corpus and corpus.get("absent"):
        return None, "перевода нет в текстах целевого языка (%s)" % corpus["label"]

    good, confirmed, distinct, why = _donor_quality(cand, ctx)
    # Дальше — не вопросы к человеку, а нехватка ДАННЫХ: новые чистые прогоны
    # приносят доноров сами (сбор терминологии висит на back-check и termcheck),
    # и карточка дорешается без человека. "wait", а не None.
    if good == 0:
        return "wait", why or "сегмент-источник не проходил проверок"
    if cand.get("kind") == "audit" and good < 2:
        # Находка termcheck — это мнение модели о собственном переводе.
        # Одного такого мнения мало даже для подсказки.
        return "wait", "находка termcheck встретилась только раз"
    if not confirmed and good < pol["auto_min_segments"]:
        return "wait", "подтверждений: %d, нужно %d" % (good, pol["auto_min_segments"])

    # Краудсорсный справочник + корпус СНИЖАЮТ порог согласия сегментов на один,
    # но не отменяют запрет области. Почему только снижают: голоса независимы
    # не полностью — модель могла выучить те же ошибки справочника, а корпус
    # подтверждает лишь то, что строка в языке существует. «Анизакидоз →
    # Anisakis» (болезнь против рода паразита) прошёл бы все три проверки, и
    # в медицине цена такого приказа выше пользы от автоматизации: там приказ
    # по-прежнему даёт человек или ВЫВЕРЕННЫЙ справочник (см. выше по коду).
    corroborated = bool(auth) and bool(corpus) and corpus.get("ok")
    need = pol["verified_min_segments"]
    if corroborated:
        need = max(2, need - 1)
    if (pol["allow_verified"] and good >= need and distinct >= need
            and good == len(_donor_ids(cand))):
        if corroborated:
            return GLOSSARY_TIER_HARD, ("%d независимых сегмента + справочник %s + %s (%d)"
                                        % (distinct, auth["label"],
                                           corpus["label"], corpus["hits"]))
        # Корпус доступен, но термин в языке почти не встречается — согласия
        # доноров мало: приказ не даём, оставляем подсказкой. Обратное неверно —
        # частотность не делает перевод правильным, поэтому поднять до приказа
        # она не может, только удержать.
        if corpus and not corpus.get("ok") and corpus.get("vetoAllowed", True):
            return GLOSSARY_TIER_SOFT, ("%d независимых сегмента, но термин редок в %s (%d)"
                                        % (distinct, corpus["label"], corpus["hits"]))
        return GLOSSARY_TIER_HARD, ("%d независимых чистых сегмента" % distinct
                                    + (" · подтверждён в %s (%d)" % (corpus["label"], corpus["hits"])
                                       if corpus else ""))
    return GLOSSARY_TIER_SOFT, ("подтвердил человек" if confirmed
                                else "%d независимых чистых сегмента" % good)


def _forget_auto_batch(batch: int):
    """Пачка выпала из истории — откатить её уже нечем. Снимаем пометки, иначе
    записи навсегда застревали бы с отказом «сначала откатите пачку #N», а
    кнопки для этого в интерфейсе уже нет."""
    for g in STATE.get("glossary", []):
        if g.get("autoBatch") == batch:
            _clear_auto_marks(g)
            for k in ("prevTgt", "prevConf", "prevOrigin", "prevSignedBy"):
                g.pop(k, None)
    for c in _term_queue():
        if c.get("autoBatch") == batch:
            c.pop("autoBatch", None)


def _clear_auto_marks(entry: dict):
    """Человек тронул запись — она больше не принадлежит пачке автоодобрения.
    Иначе откат пачки снёс бы или откатил именно то, что человек исправил."""
    for k in ("autoBatch", "autoCreated", "prevTier", "prevNote", "prevSignedBy"):
        entry.pop(k, None)


def _auto_write(cand: dict, tier: str, batch: int, today: str,
                override: bool = False) -> bool:
    """Записать пару в глоссарий. True, если существующая запись заменена.
    Прежние значения сохраняем — на них держится откат пачки."""
    scope = _scope_of(cand)
    src, tgt = cand["src"].strip(), cand["tgt"].strip()
    cat = cand.get("cat") or "Term"
    existing = _glossary_entry(src, scope)
    if existing:
        upd = {
            "tgt": tgt, "tier": tier, "cat": existing.get("cat") or cat,
            "conf": "high" if tier == GLOSSARY_TIER_HARD else "medium",
            "note": ("автоодобрено " + today
                     + (" (приказ по разрешению человека)" if override else "")),
            "updated": today,
            "autoBatch": batch,
            "byOverride": bool(override),
            "origin": "auto:" + AUTO_POLICY_VERSION,
        }
        if existing.get("autoBatch"):
            # Запись ещё принадлежит ПРОШЛОЙ пачке. Цепочку отката НЕ трогаем:
            # prev* и autoCreated остаются как есть (в upd их нет) — они
            # указывают на состояние ДО той пачки, и откат ЭТОЙ вернёт именно
            # его (или уберёт запись, рождённую пачками). Прошлая пачка запись
            # теряет — её откат считает такие в superseded.
            pass
        else:
            upd.update({
                "prevTgt": existing.get("tgt", ""), "prevTier": _hit_tier(existing),
                "prevNote": existing.get("note", ""), "prevConf": existing.get("conf", ""),
                "prevOrigin": existing.get("origin", ""), "autoCreated": False,
                "prevSignedBy": existing.get("signedBy"),
            })
        existing.update(upd)
        # Подпись человека относилась к ПРЕЖНЕМУ переводу: на машинной записи
        # она врала бы про ответственного. Откат пачки её возвращает.
        existing.pop("signedBy", None)
        return True
    entry = {
        "src": src, "tgt": tgt, "cat": cat, "freq": 1,
        "conf": "high" if tier == GLOSSARY_TIER_HARD else "medium",
        "note": ("автоодобрено " + today
                 + (" (приказ по разрешению человека)" if override else "")),
        "tier": tier,
        "lang": scope[0], "domain": scope[1], "tenant": scope[2], "updated": today,
        "autoBatch": batch, "autoCreated": True, "byOverride": bool(override),
        "origin": "auto:" + AUTO_POLICY_VERSION,
    }
    STATE["glossary"].insert(0, entry)
    # Индекс по термину живёт всю пачку: без досыпки следующий кандидат той же
    # пары считал бы, что записи ещё нет. Полная инвалидация тут стоила бы
    # пересборки 10 000 записей на каждое одобрение.
    idx = _GLOSS_BY_SRC
    if idx is not None:
        idx.setdefault((scope, _norm_key(src)), entry)
    return False


# Классические подмены, на которых ломается любая проверка, кроме вопроса
# «то же ли это понятие». Один список на оба места, где этот вопрос задаётся
# (пачкой в автоодобрении и по кнопке в разборе вариантов): разойдись они —
# «Анизакидоз → Anisakis» отклонялся бы автоматикой и подтверждался вручную.
MEANING_TRAPS = ("Watch for classic traps: a disease vs its causative agent, "
                 "an organ vs a finding, a substance vs its class, "
                 "a procedure vs its result, a patient vs an animal.")

# Набор вопросов, на которые отвечает сверка. Растёт вместе с промптом.
# Вердикт, записанный ПРЕЖНИМ набором, отвечает не на те вопросы: когда
# к «то же ли это понятие» добавился вопрос «годится ли правилом», полторы
# тысячи записей с готовым вердиктом остались с ним навсегда. Отпечаток пары
# у них совпадал, устаревшими они не считались, и «Досверить новые» находило
# ноль — а «Переспросить всё» упиралось в потолок и переспрашивало одни и те же
# восемьсот, до хвоста не доходя никогда.
# 3 — с версии 2 вердикты считались сломанным разбором: явный JSON null
# («не знаю») читался как твёрдое «понятие другое». Часть находок была
# выдумана, и доверять записанному нельзя — версия поднята, чтобы всё
# спросить заново уже исправленным кодом.
MEANING_VERSION = 3

AUTO_MEANING_MAX = 400      # потолок пар за одно применение
AUTO_MEANING_CHUNK = 10     # пар в одном вызове судьи


def _openai_meaning(pairs: list, scope: tuple) -> Optional[dict]:
    """Один вызов судьи на пачку пар «термин → кандидат»: то же ли это ПОНЯТИЕ.

    Это единственный вопрос, который не задаёт ни одна другая проверка. Корпус
    подтверждает существование строки в языке, termcheck — что термин настоящий,
    согласие сегментов — что модель повторяет себя; «Анизакидоз → Anisakis»
    (болезнь против рода паразита-возбудителя) проходит всё это разом, потому
    что каждый отвечал не на тот вопрос.

    Вопросов два, и второй не менее важен. `same` — про понятие. `rule` — про
    то, годится ли пара ПРАВИЛОМ на весь документ: «Клинику → clinical practice»
    верен в «внедрено в клинику» и неверен в «направлен в клинику», а
    «частоте → frequency» — падежная форма вместо словарной. Понятия при этом
    совпадают, и вопрос `same` такое пропускает — а запись между тем принуждает
    модель во всех сегментах разом.

    Возвращает {(термин, перевод) нормализованные: {"same": bool|None,
    "back": чем кандидат является НА ЯЗЫКЕ ОРИГИНАЛА, "rule": bool|None,
    "why": почему правилом не годится, НА ЯЗЫКЕ ОРИГИНАЛА}} или None при сбое.
    None — это «не знаю», и по тому же закону, что attested() у корпуса,
    он не одобряет и не блокирует."""
    import openai
    src_lang, tgt_lang = ((authorities_mod.source_lang(scope[0]),
                           authorities_mod.target_lang(scope[0])) if authorities_mod
                          else (scope[0].split("→")[0], scope[0].split("→")[-1]))
    dom = _resolve_domain(scope[1])
    system = (
        f"You are a {dom['expert']}. For each pair below decide whether the {tgt_lang} "
        f"candidate denotes EXACTLY the same concept as the {src_lang} term, AND whether "
        "the pair is fit to become a glossary RULE applied to a whole document.\n"
        + MEANING_TRAPS + "\n"
        "For each pair return:\n"
        "  same — true ONLY if the concepts are identical;\n"
        f"  back — what the candidate actually denotes, a few words IN {_explain_lang_name()} "
        "(this line is read by a person, so it follows the interface language);\n"
        f"  rule — false if this pair must NOT become a document-wide rule: the "
        f"{src_lang} side is an inflected or otherwise non-dictionary form instead of "
        "a lemma, or it is an ordinary word rather than a term of the field, or the "
        "correct translation depends on the surrounding context. Otherwise true;\n"
        f"  why — when rule is false, the reason IN {_explain_lang_name()}, a few words.\n"
        'Return ONLY JSON: {"pairs":[{"src":"...","tgt":"...","same":true,"back":"...",'
        '"rule":true,"why":""}]}. '
        "No commentary."
    )
    body = "\n".join(f"  - {a} → {b}" for a, b in pairs)
    try:
        mdl = _resolve_model(JUDGE_DEFAULT_MODEL)
        client = openai.OpenAI(api_key=os.environ.get("OPENAI_API_KEY"), timeout=90, max_retries=1)
        extra = ({"max_completion_tokens": 2048} if mdl["api"] == "modern"
                 else {"max_tokens": 900, "temperature": 0})
        resp = client.chat.completions.create(
            model=mdl["id"], response_format={"type": "json_object"},
            messages=[{"role": "system", "content": system},
                      {"role": "user", "content": body}], **extra)
        _note_usage("terms", mdl["id"], resp)
        data = json.loads((resp.choices[0].message.content or "").strip() or "{}")
        out = {}
        for v in (data.get("pairs") or []):
            if isinstance(v, dict) and v.get("src") and v.get("tgt"):
                # ПУСТОТА — это «не знаю», а не «нет». Ключ мог отсутствовать
                # (вердикты прежней версии промпта) либо прийти с явным JSON
                # null — так модель и выражает неуверенность. Проверка «ключ
                # есть» ловила только первое, а `bool(None)` превращала второе
                # в твёрдое «понятие другое»: запись понижалась, и откат
                # переписывал сегменты — по замечанию, которого не было.
                out[(_norm_key(v["src"]), _norm_key(v["tgt"]))] = {
                    "same": (None if v.get("same") is None else bool(v["same"])),
                    "back": (v.get("back") or "").strip(),
                    "rule": (None if v.get("rule") is None else bool(v["rule"])),
                    "why": (v.get("why") or "").strip()}
        return out
    except Exception as e:
        print(f"[backend] смысловая сверка терминов: {e}", file=sys.stderr)
        return None


def _meaning_check(cands: list, cap: int = AUTO_MEANING_MAX) -> tuple:
    """Смысловая сверка кандидатов пачками по AUTO_MEANING_CHUNK.
    Возвращает ({(область, термин, перевод): вердикт}, отвечено, сверх потолка).
    Пара, на которую судья не ответил, в словарь не попадает — «не знаю»."""
    uniq, order = {}, []
    for c in cands:
        src, tgt = (c.get("src") or "").strip(), (c.get("tgt") or "").strip()
        if not src or not tgt:
            continue
        key = (_scope_of(c), _norm_key(src), _norm_key(tgt))
        if key not in uniq:
            uniq[key] = (src, tgt)
            order.append(key)
    capped = max(0, len(order) - cap)
    if capped:
        # Молчаливых потолков не бывает.
        print(f"[backend] смысловая сверка: {capped} пар сверх потолка не проверены",
              file=sys.stderr)
    order = order[:cap]
    # Один проход и порядок вставки: обход множества областей давал бы разный
    # порядок вызовов от запуска к запуску — такое не воспроизвести по журналу.
    by_scope: dict = {}
    for k in order:
        by_scope.setdefault(k[0], []).append(k)
    jobs = []
    for scope, keys in by_scope.items():
        for i in range(0, len(keys), AUTO_MEANING_CHUNK):
            jobs.append((scope, keys[i:i + AUTO_MEANING_CHUNK]))
    out = {}
    for (scope, chunk), got in _run_parallel(
            jobs, lambda j: (j, _openai_meaning([uniq[k] for k in j[1]], j[0]))):
        if not got:
            continue        # сбой вызова: эти пары остаются «не знаю»
        for k in chunk:
            v = got.get((k[1], k[2]))
            if v is not None:
                out[k] = v
    return out, len(out), capped


class AutoApproveRequest(BaseModel):
    dry_run: bool = True
    # Корпусная проверка — это внешние HTTP-запросы с лимитом источника
    # (у PubMed 3 в секунду). В разборе «показать», который дёргается при каждом
    # открытии проекта, она превращала бы страницу в минутное ожидание. Поэтому
    # по умолчанию она идёт только при ПРИМЕНЕНИИ, а разбор честно помечается
    # как «до корпусной проверки». None = решить по dry_run.
    corpus: Optional[bool] = None
    project: Optional[int] = None      # смотреть только область этого проекта
    max_tier: Optional[str] = None     # "auto" — не поднимать до приказа
    # Снять запрет области на приказ по согласию сегментов (медицина, фарма,
    # юриспруденция). Это РАЗРЕШЕНИЕ ЧЕЛОВЕКА на конкретный запуск, а не новая
    # политика: по умолчанию None — как решает область. Отдельным полем, а не
    # правкой AUTO_APPROVE_BY_DOMAIN, именно потому, что разрешение разовое,
    # видно в ответе (policy.humanOverride) и откатывается вместе с пачкой.
    allow_verified: Optional[bool] = None
    # Смысловая сверка судьёй (то же ли понятие) — платные вызовы, поэтому как
    # корпус: по умолчанию только при ПРИМЕНЕНИИ. None = решить по dry_run.
    meaning: Optional[bool] = None
    limit: int = 2000


@app.post("/api/term-queue/auto-approve")
def auto_approve_terms(req: AutoApproveRequest = AutoApproveRequest()):
    """Разложить однозначных кандидатов по глоссарию без участия человека.

    dry_run=True (по умолчанию) НИЧЕГО не меняет — возвращает, что попадёт и
    что отсеяно с причинами. Смотреть до применения, а не после."""
    project = get_project(req.project) if req.project else None
    scope = _project_scope(project) if project else None
    pol = _auto_policy(project.get("domain") if project else None)
    # Снимаем ДО любых override: интерфейсу нужно знать, есть ли запрет
    # в принципе, иначе тумблер исчезал бы ровно от того, что его включили.
    domain_banned = not pol.get("allow_verified")
    if req.allow_verified is True and not pol.get("allow_verified"):
        # Человек снял запрет области на этот запуск. Обмануться тут нечем:
        # порог сегментов, требование РАЗНЫХ исходников, единственность
        # варианта и вето корпуса остаются на месте — снимается ровно запрет
        # «в медицине приказ даёт только человек». Помечаем разрешение, чтобы
        # оно было видно в ответе и в записях пачки: молчаливое ослабление
        # защиты хуже её отсутствия.
        pol = {**pol, "allow_verified": True, "humanOverride": True}
    if req.max_tier == GLOSSARY_TIER_SOFT:
        # cap_soft отдельно от allow_verified: первое — просьба пользователя
        # на этот запуск, второе — политика области. Их нельзя смешивать,
        # иначе справочник (который политику области обходит законно) обходил
        # бы и явно выбранный режим «только подсказки».
        pol = {**pol, "allow_verified": False, "cap_soft": True}

    pending = [c for c in _term_queue() if c.get("status", "pending") == "pending"
               and _tenant_of(c) == _current_tenant()
               and (scope is None or _scope_of(c) == scope)]
    ctx = _auto_context(pending, pol)

    considered = pending[:max(1, min(req.limit, 5000))]

    # Корпус спрашиваем ТОЛЬКО про тех, кто иначе прошёл бы: это внешний
    # HTTP-запрос, и гонять его на всю очередь из двух тысяч кандидатов —
    # значит потратить полчаса на термины, которые всё равно отсеются
    # по другим причинам. Сначала предварительный вердикт без корпуса,
    # потом проверка прошедших, потом окончательный вердикт.
    ctx["corpus"] = {}
    prelim = [c for c in considered
              if _auto_verdict(c, ctx)[0] in (GLOSSARY_TIER_HARD, GLOSSARY_TIER_SOFT)]
    # Корпус может ПОДНЯТЬ вердикт (справочник + корпус + сегменты), а не только
    # отсеять кальку. Значит разбор без него врёт в обе стороны, и молчать об
    # этом нельзя — см. corpusPending в ответе.
    corpus_asked, corpus_capped = 0, 0
    use_corpus = (not req.dry_run) if req.corpus is None else bool(req.corpus)
    if prelim and authorities_mod and use_corpus:
        # Ключ включает ОБЛАСТЬ: очередь общая на сервис, и без неё английский
        # корпус спрашивали бы про немецкие термины (и наоборот).
        by_key = {}
        for c in prelim:
            if (c.get("tgt") or "").strip():
                by_key.setdefault((_scope_of(c), _norm_key(c.get("tgt"))),
                                  (c.get("tgt") or "").strip())
        want = list(by_key)
        corpus_capped = max(0, len(want) - AUTO_CORPUS_MAX)
        want = want[:AUTO_CORPUS_MAX]
        results = _run_parallel(want, lambda k: (k, _corpus_check(by_key[k], k[0])))
        ctx["corpus"] = {k: v for k, v in results if v}
        corpus_asked = len(ctx["corpus"])
        if corpus_capped:
            # Молчаливых потолков не бывает: сказали, скольких не проверили.
            print(f"[backend] корпусная проверка: {corpus_asked} терминов, "
                  f"{corpus_capped} сверх потолка не проверены", file=sys.stderr)

    # ── Смысловая сверка: то же ли это понятие ───────────────────────
    # Спрашивается, как и корпус, только про тех, кто иначе прошёл бы, и
    # по умолчанию только при применении. Явное «не то понятие» отклоняет
    # кандидата совсем — даже подсказка из ложного друга вредна, модель
    # вправе её взять. «Не знаю» (сбой, пара не разобрана) не одобряет и
    # не блокирует — тот же закон, что у attested().
    use_meaning = (not req.dry_run) if req.meaning is None else bool(req.meaning)
    meaning, meaning_asked, meaning_capped = {}, 0, 0
    if prelim and use_meaning and os.environ.get("OPENAI_API_KEY"):
        meaning, meaning_asked, meaning_capped = _meaning_check(prelim)

    picked, closed, mrejected, skipped = [], [], [], {}
    taken = set()
    for cand in considered:
        action, reason = _auto_verdict(cand, ctx)
        row = {"id": cand["id"], "kind": cand.get("kind"), "src": cand.get("src"),
               "tgt": cand.get("tgt"), "lang": _scope_of(cand)[0],
               "domain": _scope_of(cand)[1], "reason": reason,
               "hits": cand.get("hits", 1), "donors": len(_donor_ids(cand))}
        if action in (GLOSSARY_TIER_HARD, GLOSSARY_TIER_SOFT):
            m = meaning.get((_scope_of(cand), _norm_key(cand.get("src")),
                             _norm_key(cand.get("tgt"))))
            if m and (m.get("same") is False or m.get("rule") is False):
                # Ложный друг: строка в языке существует, понятие другое.
                # Причина хранит обратный смысл ПО-РУССКИ (языку оригинала) —
                # человек, не знающий целевого языка, должен видеть, ЧТО
                # именно машина у него отвела.
                reason = ("смысл расходится: " + (m.get("back") or "иное понятие")
                          if m.get("same") is False
                          else "правилом не годится: "
                               + (m.get("why") or "зависит от контекста"))
                mrejected.append((cand, {**row, "tier": None, "reason": reason}))
                continue
            if m and m.get("same") is True:
                row["reason"] = row["reason"] + " · смысл сверен судьёй"
            # Один термин — одна запись за пачку. Иначе два кандидата разных
            # kind с одной и той же парой писали её дважды, и второй проход
            # затирал prevTgt собственным значением: запись пережила бы откат.
            key = (_scope_of(cand), _norm_key(cand.get("src")))
            if key in taken:
                closed.append((cand, {**row, "tier": None,
                                      "reason": "та же пара уже одобрена в этой пачке"}))
                continue
            taken.add(key)
            picked.append((cand, action, {**row, "tier": action}))
        elif action == "close":
            closed.append((cand, {**row, "tier": None}))
        else:
            bucket = skipped.setdefault(reason, {"reason": reason, "count": 0, "samples": []})
            bucket["count"] += 1
            if len(bucket["samples"]) < 3:
                bucket["samples"].append({"src": row["src"], "tgt": row["tgt"]})

    counts = {
        "verified": sum(1 for _, t, _ in picked if t == GLOSSARY_TIER_HARD),
        "auto": sum(1 for _, t, _ in picked if t == GLOSSARY_TIER_SOFT),
        "closed": len(closed),
        "rejectedMeaning": len(mrejected),
        "skipped": sum(b["count"] for b in skipped.values()),
        # pending — по рассмотренному срезу: вердикты считаются только по limit,
        # иначе на большой очереди цифры в панели не сходились бы.
        "pending": len(considered),
        "queueTotal": len(pending),
    }
    result = {
        "ok": True, "dryRun": req.dry_run, "batch": None, "counts": counts,
        "policy": {**pol, "version": AUTO_POLICY_VERSION, "domainBanned": domain_banned},
        "scope": list(scope) if scope else None,
        # Чем проверялись термины. Разница в покрытии между парами языков
        # огромна, и назвать её честнее, чем дать пользователю обнаружить
        # её на своих текстах.
        "sources": _authority_sources(scope) if scope else {"dictionaries": [], "corpus": None},
        "corpusChecked": corpus_asked,
        "corpusSkipped": corpus_capped,
        # Разбор без корпуса — это верхняя оценка: при применении часть
        # кандидатов может отсеяться как отсутствующие в целевом языке.
        # Названо явно, чтобы цифра на кнопке не обещала лишнего.
        "corpusPending": (not use_corpus) and bool(prelim),
        "meaningChecked": meaning_asked,
        "meaningSkipped": meaning_capped,
        # Разбор без сверки — верхняя оценка: при применении часть кандидатов
        # будет отклонена как иное понятие. Названо явно, как corpusPending.
        "meaningPending": (not use_meaning) and bool(prelim),
        # Отклонённые идут в общий список, а не отдельным полем: у них та же
        # форма строки, что у закрытых, и вторая копия данных ради одного
        # экрана — это ещё одно место, которое однажды разойдётся.
        "items": ([row for _, _, row in picked] + [row for _, row in closed]
                  + [row for _, row in mrejected]),
        "skipped": sorted(skipped.values(), key=lambda b: -b["count"]),
    }
    if req.dry_run or not (picked or closed or mrejected):
        return result

    today = datetime.now().strftime("%Y-%m-%d")
    batch = _next_batch_seq()
    for cand, tier, row in picked:
        # Пометка только на приказах: подсказку разрешение не меняет, и писать
        # на ней «по разрешению человека» значит преувеличить его участие.
        by_override = bool(pol.get("humanOverride")) and tier == GLOSSARY_TIER_HARD
        row["byOverride"] = by_override
        row["replaced"] = _auto_write(cand, tier, batch, today, by_override)
        cand.update({"status": "approved", "autoBatch": batch, "autoTier": tier,
                     "autoWrote": True, "autoNote": row["reason"], "decidedAt": today})
    for cand, row in closed:
        cand.update({"status": "approved", "autoBatch": batch, "autoWrote": False,
                     "autoNote": row["reason"], "decidedAt": today})
    for cand, row in mrejected:
        # Отклонение машиной, не человеком: _human_decision это различает,
        # и решённый так кандидат вопрос не закрывает. Откат пачки вернёт его
        # в очередь — автоматика без отката недопустима.
        cand.update({"status": "rejected", "autoBatch": batch, "autoWrote": False,
                     "autoNote": row["reason"], "decidedAt": today})
    batches = STATE.setdefault("autoBatches", [])
    batches.insert(0, {"id": batch, "at": datetime.now().strftime("%Y-%m-%d %H:%M"),
                       "counts": counts, "scope": list(scope) if scope else None,
                       # Разрешение человека видно в истории пачек: через месяц
                       # «откуда в медицине приказы от машины» отвечается здесь,
                       # а не раскопками по записям глоссария.
                       "override": bool(pol.get("humanOverride"))})
    for gone in batches[AUTO_BATCH_HISTORY:]:
        _forget_auto_batch(gone["id"])
    del batches[AUTO_BATCH_HISTORY:]
    _invalidate_gloss_index()
    save_state(STATE)
    result["batch"] = batch
    return result


@app.post("/api/term-queue/auto-approve/{batch}/undo")
def undo_auto_approve(batch: int):
    """Откатить пачку целиком: созданные записи убрать, заменённые вернуть,
    кандидатов — обратно в очередь. Автоматике, которую нельзя отменить одной
    кнопкой, доверять нельзя."""
    # Откат ПОНИЖЕНИЯ — это решение человека «запись оставить приказом».
    # Без этой пометки следующий аудит увидел бы записанный вердикт «не то
    # понятие» и предложил бы понизить её снова, по кругу.
    # Номер пачки — общий счётчик на все организации, а откат переписывает
    # глоссарий: чужой номер → 404 (не 403: не подтверждаем, что пачка есть).
    rec = next((b for b in STATE.get("autoBatches", []) if b.get("id") == batch), None)
    if rec is not None and _tenant_of(rec) != _current_tenant():
        raise HTTPException(404, "Пачка не найдена")
    audit = bool(rec and rec.get("kind") == "audit")
    removed, restored = 0, 0
    keep = []
    for g in STATE["glossary"]:
        if g.get("autoBatch") != batch or _tenant_of(g) != _current_tenant():
            keep.append(g)
            continue
        if g.get("autoCreated"):
            removed += 1
            continue
        g["tgt"] = g.pop("prevTgt", g.get("tgt"))
        g["tier"] = g.pop("prevTier", GLOSSARY_TIER_SOFT)
        g["note"] = g.pop("prevNote", "")
        g["conf"] = g.pop("prevConf", "")
        prev_origin = g.pop("prevOrigin", "")
        if prev_origin:
            g["origin"] = prev_origin
        else:
            g.pop("origin", None)
        prev_signed = g.pop("prevSignedBy", None)
        if prev_signed:
            g["signedBy"] = prev_signed
        for k in ("autoBatch", "autoCreated"):
            g.pop(k, None)
        if audit:
            g["meaningKept"] = True
        restored += 1
        keep.append(g)
    STATE["glossary"] = keep
    back = 0
    superseded = 0
    for c in _term_queue():
        if c.get("autoBatch") == batch and _tenant_of(c) == _current_tenant():
            # Запись, которую эта пачка писала, могла быть ПЕРЕЗАПИСАНА более
            # поздней пачкой (см. _auto_write): она больше не наша, откат её
            # не трогает — трогать значило бы затереть более позднее решение.
            # Молчать нельзя: человек уверен, что откатил пачку целиком.
            # И В ОЧЕРЕДЬ такая карточка не возвращается: снова ставшая
            # pending, она следующим же «Одобрить и применить» вписала бы
            # СТАРЫЙ машинный вариант поверх более позднего решения — глоссарий
            # молча откатывался бы назад без единого выбора человека.
            if c.get("autoWrote"):
                g = _glossary_entry(c.get("src") or "", _scope_of(c))
                if g is not None and g.get("autoBatch") not in (None, batch):
                    superseded += 1
                    c["note"] = ("запись перехвачена пачкой #%s — вопрос "
                                 "решает она" % g["autoBatch"])
                    for k in ("autoBatch", "autoTier", "decidedAt"):
                        c.pop(k, None)
                    continue
            c["status"] = "pending"
            for k in ("autoBatch", "autoTier", "autoWrote", "autoNote", "decidedAt"):
                c.pop(k, None)
            back += 1
    STATE["autoBatches"] = [b for b in STATE.get("autoBatches", []) if b.get("id") != batch]
    _invalidate_gloss_index()
    save_state(STATE)
    return {"ok": True, "removed": removed, "restored": restored, "returned": back,
            "superseded": superseded}


@app.get("/api/term-queue/auto-batches")
def list_auto_batches():
    t = _current_tenant()
    return {"ok": True, "batches": [b for b in STATE.get("autoBatches", []) if _tenant_of(b) == t]}


_GLOSS_MARK_LOCK = threading.Lock()


def _meaning_pair(entry: dict) -> str:
    """Отпечаток пары, к которой относится вердикт. Правили перевод — вердикт
    о прежней паре больше ничего не значит; тот же приём, что у `target_hash`
    у back-check и termcheck."""
    return _text_hash(_norm_key(entry.get("src")) + "||" + _norm_key(entry.get("tgt")))


def _verdict_bad(v: Optional[dict]) -> bool:
    """Вердикт забраковал пару: понятие другое ЛИБО правилом не годится.
    Одно определение на всех — разойдись оно между записью вердикта и разбором
    находок, аудит показывал бы одно, а понижал другое."""
    if not v:
        return False
    return v.get("same") is False or v.get("rule") is False


def _meaning_stale(entry: dict) -> bool:
    """Записи нужен вопрос: её ещё не спрашивали, пара изменилась либо после
    прошлой сверки на неё пожаловалась проверка терминов.

    Без этого аудит переспрашивал ВСЕ записи каждым запуском, а судья на
    границе «то же понятие / не то» отвечает неустойчиво: список находок
    каждый раз получался новый, и понижать приходилось бесконечно."""
    m = entry.get("meaning") or {}
    if m.get("pair") != _meaning_pair(entry):
        return True
    # Вердикт отвечает не на тот набор вопросов — значит его нет.
    if int(m.get("v") or 1) != MEANING_VERSION:
        return True
    return int(entry.get("disputed") or 0) > int(m.get("disputed") or 0)


def _note_term_disputes(seg: dict, project: Optional[dict]) -> int:
    """Проверка терминов забраковала слово, которое И ЕСТЬ приказной перевод.

    Понизить запись прямо отсюда нельзя: termcheck — один вызов модели, он
    ошибается (на «infiltrate → induration» ошибается именно он), и отдать
    ему право переписывать глоссарий значит поставить шумную проверку выше
    решения. Поэтому он НОМИНИРУЕТ: счётчик `disputed` растёт, вердикт сверки
    смысла для этой записи считается устаревшим, и следующий аудит спросит
    про неё заново — а решает по-прежнему вопрос «то же ли это понятие».

    Под локом: termcheck идёт в рабочих потоках, а записи глоссария общие."""
    if project is None:
        return 0
    tc = seg.get("termcheck") or {}
    bad = [f for f in (tc.get("findings") or [])
           if f.get("severity") in TERMCHECK_DISPUTING and f.get("tgt_term")]
    every = [f for f in (tc.get("findings") or [])
             if f.get("severity") in TERMCHECK_ACTIONABLE and f.get("tgt_term")]
    if not every:
        return 0
    # Терм-лист документа: один голос любой действующей тяжести снимает пару
    # из промпта (`_termlist_dispute`) — той же, по которой ремонт заменил бы
    # её; решает дальше человек.
    doc = _doc_hits(seg.get("source", ""), project)
    for f in every:
        for h in doc:
            if _term_forms_overlap(f.get("tgt_term") or "", h["tgt"]):
                _termlist_dispute(project, h["src"], "termcheck: " + (f.get("issue") or f.get("why") or "")[:160])
    hits = {_norm_key(h["tgt"]): h for h in _verified_hits(seg.get("source", ""), project)}
    if not hits:
        return 0
    # Метку «эта находка воюет с приказной записью» ставим НА САМУ НАХОДКУ,
    # здесь и сейчас: проект тут есть, а `_repair_findings` его не имеет —
    # разбор состава зовёт её без проекта ради скорости. Считай мы это
    # по-разному, смета обещала бы одно, а прогон делал другое.
    # Ремонту такая находка не отдаётся: он подставит совет, `_repair_scores`
    # увидит нарушенный приказный термин и откатит правку. На боевом проекте
    # так сгорели 22 захода, и исход у них был известен заранее.
    for f in every:
        if _norm_key(f.get("tgt_term")) in hits:
            f["vsVerified"] = True
    if not bad:
        return 0
    scope = _project_scope(project)
    marked = 0
    with _GLOSS_MARK_LOCK:
        for f in bad:
            h = hits.get(_norm_key(f.get("tgt_term")))
            if h is None:
                continue
            # _get_context отдаёт КОПИИ записей (в них подмешан контекст поиска),
            # поэтому метку надо ставить настоящей записи глоссария — иначе она
            # уходит во временный словарь и пропадает вместе с ним.
            entry = _glossary_entry(h.get("src"), scope)
            if entry is None:
                continue
            entry["disputed"] = int(entry.get("disputed") or 0) + 1
            sug = (f.get("suggestion") or "").strip()
            if sug:
                entry["disputedSuggest"] = sug
            marked += 1
    return marked


def _human_touched(entry: dict) -> bool:
    """У записи есть СЛЕД решения человека, а не просто уровень «приказ».

    Разница решающая. Уровень мог достаться записи по умолчанию миграции
    («её нет в массовом импорте — значит добавлена руками»), и это
    ПРЕДПОЛОЖЕНИЕ, а не чьё-то решение: на боевых данных ровно так получили
    приказ 670 записей без единого следа. След оставляют только настоящие
    действия: одобрение кандидата (`origin: confirmed:*`), ручная правка
    (`уточнено вручную`) и автоодобрение (`autoBatch`/`byOverride`), которое
    само по себе шло с разрешения человека.

    Своё предположение машина вправе пересмотреть, чужое решение — нет."""
    origin = (entry.get("origin") or "").lower()
    note = (entry.get("note") or "").lower()
    return bool(origin.startswith("confirmed:") or "уточнено вручную" in note
                or entry.get("autoBatch") or entry.get("byOverride")
                # Подпись ответственного — и есть след, ради которого заведена.
                or entry.get("signedBy")
                # Человек откатил понижение — значит решил оставить запись
                # приказом, и предлагать понижение снова нельзя.
                or entry.get("meaningKept"))


class GlossaryAuditRequest(BaseModel):
    project: Optional[int] = None      # область: чей глоссарий проверяем
    dry_run: bool = True
    # Переспросить и то, что уже сверялось. По умолчанию НЕТ: вердикт лежит
    # на записи, и повторный проход стоит денег и даёт другой ответ на
    # пограничных парах — ради этого кэш и заведён.
    force: bool = False
    # Понижать и записи со следом решения человека. Это РАЗРЕШЕНИЕ на пачку,
    # а не новая политика: сама машина чужое решение не отменяет никогда.
    # Записи с `meaningKept` не трогаются и здесь — человек уже возвращал их
    # из понижения, и переспрашивать значит не считать его ответ ответом.
    include_human: bool = False
    limit: int = 800                   # потолок НОВЫХ пар за проверку


@app.post("/api/glossary/audit")
def audit_glossary(req: GlossaryAuditRequest = GlossaryAuditRequest()):
    """Смысловая сверка ЗАПИСЕЙ глоссария, уже стоящих приказом.

    Сверка при автоодобрении сторожит ВХОД: новый ложный друг в глоссарий
    не сядет. Но записи, попавшие туда раньше, её не проходили — а именно они
    приказывают модели и гонят ремонт. Этот проход задаёт им тот же вопрос:
    то же ли это понятие.

    Что делает с находкой: понижает до ПОДСКАЗКИ, а не удаляет. Запись
    остаётся в глоссарии и остаётся видна, но перестаёт приказывать модели
    и перестаёт быть основанием для ремонта — то есть перестаёт вредить.
    Записи со следом решения человека (`_human_touched`) не понижаются
    никогда: они только помечаются, решение по ним остаётся за человеком.

    Порядок проверки — по вреду: сначала те, что уже расходятся с переводом
    в этом проекте (их применяет ремонт), потом остальные. Потолок обрежет
    хвост, а не голову."""
    if not os.environ.get("OPENAI_API_KEY"):
        raise HTTPException(503, "Смысловая сверка требует ключ OpenAI")
    project = get_project(req.project) if req.project else None
    scope = _project_scope(project) if project else None
    entries = [g for g in STATE.get("glossary", [])
               if _hard_answer(g) and _tenant_of(g) == _current_tenant()
               and (scope is None or _scope_of(g) == scope)]
    if not entries:
        return {"ok": True, "dryRun": req.dry_run, "checked": 0, "capped": 0,
                "total": 0, "bad": [], "batch": None,
                "scope": list(scope) if scope else None}

    # Порядок по вреду: записи, которые прямо сейчас расходятся с переводом,
    # идут первыми — ремонт применит именно их.
    weight = {}
    if project is not None:
        try:
            for t in glossary_impact(project["id"])["terms"]:
                weight[_norm_key(t["src"])] = len(t.get("segments") or ())
        except Exception as e:                                   # pragma: no cover
            print(f"[backend] аудит глоссария: вес по импакту не посчитан: {e}",
                  file=sys.stderr)
    # Вторым ключом — сам термин: вес считается по расхождениям, а они меняются
    # после каждого понижения. Без устойчивого второго ключа порядок плавал бы
    # от захода к заходу, и потолок отрезал бы каждый раз ДРУГОЙ хвост.
    entries.sort(key=lambda g: (-weight.get(_norm_key(g.get("src")), 0),
                                _norm_key(g.get("src"))))

    # Спрашиваем только то, чего ещё не спрашивали. Вердикт живёт НА ЗАПИСИ,
    # поэтому список находок от запуска к запуску не пляшет: судья на границе
    # «то же понятие / не то» отвечает неустойчиво, и без памяти каждый проход
    # приносил новые записи на понижение — бесконечно.
    todo = [g for g in entries if req.force or _meaning_stale(g)]
    verdicts, asked, capped = _meaning_check(todo, cap=max(1, min(req.limit, 2000)))
    today = datetime.now().strftime("%Y-%m-%d")
    mdl_id = _resolve_model(JUDGE_DEFAULT_MODEL)["id"]
    def _write(g, v, flipped=False):
        g["meaning"] = {"same": bool(v["same"]), "back": v.get("back") or "",
                        "rule": v.get("rule"), "why": v.get("why") or "",
                        "pair": _meaning_pair(g), "disputed": int(g.get("disputed") or 0),
                        "v": MEANING_VERSION, "model": mdl_id, "at": today,
                        "flips": int((g.get("meaning") or {}).get("flips") or 0)
                                 + (1 if flipped else 0)}

    # Переспрос НЕ должен переворачивать готовый вердикт с одной попытки.
    # Судья на границе «годится / не годится» отвечает неустойчиво, и
    # «Переспросить всё» превращалось в переброс монеты: каждый заход находил
    # новые записи на понижение, хотя и пара, и вопрос те же. Отсюда ощущение,
    # что работа нескончаема. Разошёлся с прежним ответом — спрашиваем ТРЕТИЙ
    # раз и берём большинство; платит за это только сама спорная пара.
    flips = []
    for g in todo:
        v = verdicts.get((_scope_of(g), _norm_key(g.get("src")), _norm_key(g.get("tgt"))))
        # «Не знаю» (судья не ответил, вызов упал) не записываем: иначе молчание
        # закрыло бы запись от нормальной проверки навсегда. Тот же закон, что
        # у attested() и у отметки Medical QA.
        if not v or v.get("same") is None:
            continue
        prev = g.get("meaning") or {}
        # Сравнимо только с ответом на ТОТ ЖЕ вопрос о ТОЙ ЖЕ паре.
        comparable = (prev.get("pair") == _meaning_pair(g)
                      and int(prev.get("v") or 0) == MEANING_VERSION
                      and prev.get("same") is not None)
        if comparable and _verdict_bad(prev) != _verdict_bad(v):
            flips.append((g, prev, v))
            continue
        _write(g, v)

    reasked = 0
    if flips:
        tie, _n, _c = _meaning_check([g for g, _p, _v in flips], cap=len(flips))
        reasked = len(flips)
        for g, prev, v in flips:
            t = tie.get((_scope_of(g), _norm_key(g.get("src")), _norm_key(g.get("tgt"))))
            if t and t.get("same") is not None and _verdict_bad(t) == _verdict_bad(v):
                _write(g, v, flipped=True)      # третий согласен с новым
            else:
                # Большинство за прежним ответом (или третий смолчал). Клеймим
                # его свежим, иначе следующий заход снова потянет ту же пару.
                _write(g, prev, flipped=True)
    # Список строится по ЗАПИСЯННЫМ вердиктам — и свежим, и лежавшим с прошлого
    # раза: иначе разбор показывал бы только новое и врал бы про объём работы.
    bad = []
    for g in entries:
        m = g.get("meaning") or {}
        if m.get("pair") != _meaning_pair(g):
            continue
        # Две разные беды, и обе делают запись негодной ПРИКАЗОМ: перевод
        # означает другое либо пара не годится правилом на весь документ
        # (падежная форма, обычное слово, перевод по контексту).
        wrong = m.get("same") is False
        if not _verdict_bad(m):
            continue
        bad.append({"src": g["src"], "tgt": g["tgt"],
                    "kind": "meaning" if wrong else "rule",
                    "why": m.get("why") or "",
                    "back": m.get("back") or "иное понятие",
                    "segments": weight.get(_norm_key(g.get("src")), 0),
                    "disputed": int(g.get("disputed") or 0),
                    "humanTouched": _human_touched(g),
                    # Человек уже возвращал эту запись из понижения — его ответ
                    # окончателен, и разрешение на пачку его не отменяет.
                    "kept": bool(g.get("meaningKept")),
                    "lang": _scope_of(g)[0], "domain": _scope_of(g)[1],
                    "tenant": _scope_of(g)[2]})
    # Остаток — по СВЕЖЕСТИ вердикта, а не по арифметике «todo минус asked»:
    # судья мог не ответить про часть пар, и такие остаются неспрошенными.
    left = sum(1 for g in entries if _meaning_stale(g))
    result = {"ok": True, "dryRun": req.dry_run, "total": len(entries),
              "checked": asked, "cached": len(entries) - len(todo),
              "pending": left, "capped": capped, "batch": None,
              # Сколько пар разошлись с прежним ответом и решались третьим
              # голосом. Ноль здесь означает «переспрашивать больше нечего».
              "reasked": reasked,
              "scope": list(scope) if scope else None,
              "bad": sorted(bad, key=lambda b: (b["humanTouched"], -b["segments"])),
              "downgradable": sum(1 for b in bad if not b["humanTouched"]),
              # Сколько добавит разрешение — цифра на кнопке обязана относиться
              # к тому, что произойдёт, а не к тому, что можно было бы.
              "downgradableHuman": sum(1 for b in bad
                                       if b["humanTouched"] and not b.get("kept")),
              "keptByHuman": sum(1 for b in bad if b.get("kept")),
              "includeHuman": bool(req.include_human)}
    if req.dry_run or not bad:
        # Вердикты — это КЭШ, а не решение: платный ответ судьи выброшенным
        # быть не должен, иначе следующий разбор спросит то же самое и получит
        # другой ответ. Понижения тут по-прежнему нет.
        if asked:
            save_state(STATE)
        return result

    batch = _next_batch_seq()
    done = 0
    flagged = {(b["lang"], b["domain"], _norm_key(b["src"]), _norm_key(b["tgt"])): b
               for b in bad
               if (not b["humanTouched"]) or (req.include_human and not b.get("kept"))}
    for g in STATE.get("glossary", []):
        # Ключ — без организации, и без этой строки понижалась бы и одноимённая
        # запись ЧУЖОГО глоссария (`bad` собран по своей организации).
        if _tenant_of(g) != _current_tenant():
            continue
        b = flagged.get((_scope_of(g)[0], _scope_of(g)[1],
                         _norm_key(g.get("src")), _norm_key(g.get("tgt"))))
        if b is None:
            continue
        # Прежние значения — на них держится откат пачки (undo_auto_approve).
        # Перевод не трогаем: понижается только уровень доверия.
        # prevOrigin обязателен: откат при его отсутствии СНИМАЕТ origin
        # (см. undo_auto_approve) — запись потеряла бы своё происхождение.
        # prevTgt — НЫНЕШНИЙ перевод, а не оставшийся от чужой пачки. Понижение
        # перевода не трогает, но откат читает `prevTgt` безусловно
        # (undo_auto_approve), и чужое значение воскресило бы перевод,
        # от которого давно отказались.
        g.update({"prevTgt": g.get("tgt", ""),
                  "prevTier": _hit_tier(g), "prevNote": g.get("note", ""),
                  "prevConf": g.get("conf", ""), "prevOrigin": g.get("origin", ""),
                  "tier": GLOSSARY_TIER_SOFT, "conf": "medium",
                  # Причина в примечании — та, по которой запись и понижена:
                  # у «правилом не годится» она в why, а не в back.
                  "note": ("понижено сверкой смысла " + today + ": "
                           + (b.get("why") or "правилом не годится"
                              if b.get("kind") == "rule" else b.get("back") or "иное понятие")),
                  "updated": today, "autoBatch": batch, "autoCreated": False})
        done += 1
    batches = STATE.setdefault("autoBatches", [])
    # Понижение снимает повод чинить дальше, но уже переписанное так и осталось.
    # Возвращаем его тем же нажатием: руками это сотни сегментов, а откат
    # ничего не сочиняет — подставляет текст, что стоял до правки.
    # Область берём у самих записей: разбор без проекта охватывает ВСЕ пары
    # языков, и один общий откат по области по умолчанию трогал бы чужие
    # сегменты, а до остальных не доходил вовсе.
    back = {"reverted": 0, "requeued": 0, "skipped": 0}
    by_scope: dict = {}
    for b in flagged.values():
        by_scope.setdefault((b["lang"], b["domain"], b.get("tenant") or DEFAULT_TENANT), []).append(b)
    for sc, items in by_scope.items():
        part = _revert_repairs_bulk(items, sc)
        for k in back:
            back[k] += part[k]
    if back["reverted"] or back["requeued"]:
        _IMPACT_CACHE.clear()
        _ANALYSIS_CACHE.clear()
    result["reverted"] = back
    batches.insert(0, {"id": batch, "at": datetime.now().strftime("%Y-%m-%d %H:%M"),
                       "kind": "audit", "scope": list(scope) if scope else None,
                       "counts": {"verified": 0, "auto": 0, "closed": 0,
                                  "downgraded": done, "skipped": 0}})
    for gone in batches[AUTO_BATCH_HISTORY:]:
        _forget_auto_batch(gone["id"])
    del batches[AUTO_BATCH_HISTORY:]
    _invalidate_gloss_index()
    save_state(STATE)
    result["batch"] = batch
    result["downgraded"] = done
    return result


# Извлечение терминов из подтверждённых сегментов. Платный прогон: вызывается
# только по кнопке и только по подтверждённым парам.
def _term_extract_system(domain: dict) -> str:
    """Промпт извлечения терминов. Категории и примеры берутся из области:
    список «Anatomy|Cardiology|Dosage» в юридическом проекте бессмыслен."""
    return (
        "You extract bilingual " + domain["en"] + " terminology pairs from confirmed\n"
        "translation segments. Return ONLY a JSON array, no prose.\n\n"
        'Each item: {"src": <term in the source language>, "tgt": <its translation, copied from the\n'
        'target segment>, "cat": <one of ' + "|".join(domain["cats"]) + '>}\n\n'
        "RULES:\n"
        "1. " + domain["extract"] + "\n"
        "2. Give the source term in dictionary form (nominative singular).\n"
        "3. The target side MUST be copied from the segment as written, never invented.\n"
        "4. Skip general vocabulary, numbers, whole sentences, anything longer than 5 words.\n"
        "5. At most 5 pairs per segment. Return [] if the segment has no terminology.\n"
    )


def _extract_terms_call(pairs: list, model: Optional[str] = None,
                        domain_id: Optional[str] = None) -> list:
    """Один вызов модели на пачку сегментов. Возвращает список пар или []."""
    import json as _json
    import openai
    dom = _resolve_domain(domain_id)
    mdl = _resolve_model(model or DEFAULT_OPENAI_MODEL)
    client = openai.OpenAI(api_key=os.environ.get("OPENAI_API_KEY"), timeout=90, max_retries=1)
    body = "\n\n".join(f"[{i + 1}] SRC: {p[0]}\n    TGT: {p[1]}" for i, p in enumerate(pairs))
    extra = ({"max_completion_tokens": 4096} if mdl["api"] == "modern"
             else {"max_tokens": 1500, "temperature": 0})
    try:
        resp = client.chat.completions.create(
            model=mdl["id"],
            messages=[{"role": "system", "content": _term_extract_system(dom)},
                      {"role": "user", "content": body}],
            **extra,
        )
        _note_usage("terms", mdl["id"], resp)
        raw = (resp.choices[0].message.content or "").strip()
        lo, hi = raw.find("["), raw.rfind("]")
        if lo == -1 or hi <= lo:
            return []
        data = _json.loads(raw[lo:hi + 1])
        return [d for d in data if isinstance(d, dict) and d.get("src") and d.get("tgt")]
    except Exception as e:
        print(f"[backend] term extraction failed: {e}", file=sys.stderr)
        return []


# Потолок пар из одной правки: человек за один заход не исправляет больше —
# длиннее список бывает только у выдумки, и оплачивать её разбор незачем.
EDIT_TERMS_MAX = 10


def _edit_terms_prompt(src_lang: str, tgt_lang: str, domain: dict) -> str:
    """Промпт извлечения ИСПРАВЛЕННОЙ человеком терминологии.

    Отдельной функцией — тот же закон, что у `_translate_system`: собирается
    без сети и проверяется тестом с подменённым openai (см. раздел «Промпты
    проверяются НАСТОЯЩИМ кодом» в CLAUDE.md).

    Соседей здесь НЕТ намеренно: правка локальна, контекст ей даёт сам
    оригинал сегмента, а лишние сотни токенов на каждое подтверждение —
    расход без измеримой пользы. Свободного текста в ответе нет, поэтому
    язык объяснений (`_explain_lang_name`) сюда не протаскивается."""
    return (
        "You extract " + domain["en"] + " terminology that a HUMAN translator corrected.\n"
        "You are given the source text (" + src_lang + "), a DRAFT translation (" + tgt_lang + ")\n"
        "and the FINAL translation the human produced by editing that draft.\n\n"
        "Compare DRAFT and FINAL. Return ONLY a JSON array, no prose.\n"
        'Each item: {"src": <source-language term, dictionary form>,\n'
        '            "tgt": <the human\'s corrected translation, dictionary form>,\n'
        '            "was": <the draft\'s rejected wording for this term>}\n\n'
        "RULES:\n"
        "1. List ONLY terminology corrections: a domain term whose translation the human changed.\n"
        "2. Grammar, style, word-order or punctuation edits are NOT terminology — skip them.\n"
        "3. Do not invent: src must occur in the source text, tgt must occur in the FINAL text.\n"
        "4. At most " + str(EDIT_TERMS_MAX) + " items. Return [] if no terminology was corrected.\n"
    )


def _openai_edit_terms(source: str, before: str, after: str, project: dict):
    """Один вызов модели про одну правку. None — СБОЙ, а не «пар нет»:
    вызывающий обязан назвать причину, молчаливый пропуск запрещён."""
    import json as _json
    import openai
    dom = _resolve_domain(project.get("domain"))
    mdl = _resolve_model(DEFAULT_OPENAI_MODEL)
    src_l = (project.get("src") or "").upper() or "SRC"
    tgt_l = (project.get("tgt") or "").upper() or "TGT"
    # Таймаут 60 — младший из прецедентов проекта; без повторов: подтверждение
    # ждёт этот ответ, и вторая попытка удвоила бы паузу человеку.
    client = openai.OpenAI(api_key=os.environ.get("OPENAI_API_KEY"),
                           timeout=60, max_retries=0)
    cut = 2000     # сегмент столько не занимает; это страховка, не норма
    body = ("SOURCE (" + src_l + "): " + (source or "")[:cut]
            + "\n\nDRAFT (" + tgt_l + "): " + (before or "")[:cut]
            + "\n\nFINAL (" + tgt_l + "): " + (after or "")[:cut])
    extra = ({"max_completion_tokens": 2048} if mdl["api"] == "modern"
             else {"max_tokens": 1000, "temperature": 0})
    try:
        resp = client.chat.completions.create(
            model=mdl["id"],
            messages=[{"role": "system", "content": _edit_terms_prompt(src_l, tgt_l, dom)},
                      {"role": "user", "content": body}],
            **extra,
        )
        _note_usage("edit_terms", mdl["id"], resp)
        raw = (resp.choices[0].message.content or "").strip()
        # Разбор как в _extract_terms_call: JSON между первой «[» и последней
        # «]» — устойчиво к мусору вокруг.
        lo, hi = raw.find("["), raw.rfind("]")
        if lo == -1 or hi <= lo:
            return None
        data = _json.loads(raw[lo:hi + 1])
        return [d for d in data if isinstance(d, dict)]
    except Exception as e:
        print(f"[backend] разбор правки не удался: {e}", file=sys.stderr)
        return None


def _was_tgt_left(project: dict, sid: int, was: str) -> dict:
    """Сколько сегментов ЭТОГО проекта ещё содержат отвергнутый человеком
    вариант. Только свой проект и только счёт с горсткой id: по остальным
    после одобрения честно отвечает `glossary_impact`, а обход «по всем
    проектам» был бы здесь и лишним сканом на каждое подтверждение, и обходом
    изоляции организаций (инвариант 11)."""
    rx = _word_re(was)
    if not rx:
        return {"count": 0, "ids": []}
    ids = [s["id"] for s in project.get("segments") or ()
           if s.get("id") != sid and rx.search(s.get("target") or "")]
    return {"count": len(ids), "ids": ids[:20]}


def _harvest_edited_terms(seg: dict, project: dict) -> dict:
    """Термины, исправленные ЧЕЛОВЕКОМ внутри сегмента, — кандидатами в очередь.

    `_harvest_terms` видит только глоссарные конфликты и короткие сегменты
    целиком; правка термина ВНУТРИ длинного сегмента не извлекалась ничем —
    человек чинил «bioptate» на «biopsy specimen», а система не узнавала
    об этом никогда. Здесь диф «база правки → подтверждённый текст» уходит
    одним вызовом модели, и каждая возвращённая пара проверяется
    ДЕТЕРМИНИРОВАННО, прежде чем стать карточкой.

    Законы, которые нельзя ослаблять:
    • в глоссарий НИЧЕГО не пишется (инвариант 8: подтверждение сегмента
      приказа не даёт) — только кандидат; приказ появится, когда человек
      одобрит карточку, и одобрение пройдёт штатную сверку смысла;
    • fail-open: подтверждение — бесплатная работа (инвариант 15), и сбой
      здесь его не останавливает. Нет ключа, исчерпан лимит, модель молчит —
      пропуск с КОДОМ причины в ответе (`skipped`), а не молча и не 402;
      confirm в `_PAID` не входит намеренно — 402 запер бы подтверждение.
      Код, а не русская фраза: браузер переводит по коду (закон корзин
      `CLEAN_*`), подстрока сломалась бы от правки формулировки;
    • база одноразовая: `editedFrom` снимается ДО вызова, повторное
      подтверждение то же извлечение не покупает. Оборванный вызов пару
      теряет, и это названный закон, а не случайность: пара уже в TM,
      а сторожить неудачу значило бы копить в сегменте вечное поле;
    • ответ модели — только сырьё: src обязан НАЙТИСЬ в оригинале
      (`_term_match` — та же морфология, что у подбора терминов: модель
      отвечает словарной формой, а в тексте форма косвенная), tgt — в
      подтверждённом тексте (`_tgt_has_term`). Не нашлось — выдумка, отсев
      считается (`dropped`);
    • правка, расходящаяся с ПРИКАЗНОЙ записью, карточкой не становится
      (`_queue_term` её и не завёл бы — `_hard_answer` закрывает вопрос),
      но и не глотается: пара уходит в `disputed`, тост называет спор
      вслух. Решается он правкой самой записи (понижение/правка в
      глоссарии) — машина чужой приказ не трогает;
    • охват — только свой проект (`wasTgtLeft`), см. `_was_tgt_left`.
    """
    out: dict = {"pairs": [], "skipped": None, "dropped": 0,
                 "disputed": [], "cards": []}
    before = seg.pop("editedFrom", None)
    to_hash = seg.pop("editedToHash", None)
    target = (seg.get("target") or "").strip()
    if not before or not str(before).strip() or not target:
        return out
    # База жива, только если ПОСЛЕДНИМ target писал человек: цепочку хешей
    # ведёт _note_hand_edit. Разошлось — между правками писала машина
    # (ремонт, пакетный перевод, undo), и диф приписал бы человеку её слова.
    if to_hash != _text_hash(seg.get("target") or ""):
        return out
    if _norm_key(before) == _norm_key(target):
        return out
    if not os.environ.get("OPENAI_API_KEY"):
        out["skipped"] = "no_key"
        return out
    if _spend_status().get("over"):
        out["skipped"] = "limit"
        return out
    items = _openai_edit_terms(seg.get("source") or "", str(before), target, project)
    if items is None:
        out["skipped"] = "error"
        return out
    scope = _project_scope(project)
    src_lang = _src_lang({"lang": scope[0]})
    for it in items[:EDIT_TERMS_MAX]:
        t_src = (it.get("src") or "").strip()
        t_tgt = (it.get("tgt") or "").strip()
        t_was = (it.get("was") or "").strip()
        if not t_src or not t_tgt or _norm_key(t_tgt) == _norm_key(t_was):
            out["dropped"] += 1
            continue
        if not _term_match(t_src, seg.get("source") or "", src_lang):
            out["dropped"] += 1
            continue
        if not _tgt_has_term(target, t_tgt):
            out["dropped"] += 1
            continue
        entry = _glossary_entry(t_src, scope)
        if _hard_answer(entry) and _norm_key(entry.get("tgt")) != _norm_key(t_tgt):
            out["disputed"].append({"src": t_src, "tgt": t_tgt,
                                    "gloss": (entry.get("tgt") or "").strip()})
            continue
        meta = {"project": project["id"], "segment": seg["id"],
                "lang": scope[0], "domain": scope[1], "tenant": scope[2],
                "via": "confirmed",
                "sampleSrc": (seg.get("source") or "")[:240],
                "sampleTgt": target[:240]}
        if t_was:
            meta["wasTgt"] = t_was
            left = _was_tgt_left(project, seg["id"], t_was)
            if left["count"]:
                meta["wasTgtLeft"] = left
        c = _queue_term("edit", t_src, t_tgt, **meta)
        if c:
            out["cards"].append(c)
            out["pairs"].append({"src": t_src, "tgt": t_tgt, "wasTgt": t_was})
    return out


class ExtractTermsRequest(BaseModel):
    segment_ids: Optional[List[int]] = None
    limit: int = 30
    model: Optional[str] = None


@app.post("/api/projects/{pid}/extract-terms")
def extract_terms(pid: int, req: ExtractTermsRequest = ExtractTermsRequest()):
    """Достаёт терминологические пары из подтверждённых сегментов проекта.
    Кладёт их в очередь кандидатов, а не в глоссарий. Обычный def: внутри
    блокирующие вызовы модели (см. batch_translate)."""
    if not os.environ.get("OPENAI_API_KEY"):
        raise HTTPException(503, "Извлечение терминов требует ключ OpenAI")
    project = get_project(pid)
    segs = [s for s in project["segments"]
            if s.get("status") == "confirmed" and (s.get("target") or "").strip()]
    if req.segment_ids:
        ids = set(req.segment_ids)
        segs = [s for s in segs if s["id"] in ids]
    # Пара «оригинал+перевод» не менялась с прошлого извлечения — читать нечего
    fresh = []
    skipped_cached = 0
    for sg in segs:
        h = _text_hash((sg.get("source") or "") + "||" + (sg.get("target") or ""))
        if sg.get("extracted_hash") == h:
            skipped_cached += 1
            continue
        sg["_extract_hash"] = h
        fresh.append(sg)
    segs = fresh[:max(1, min(req.limit, 100))]
    if not segs:
        return {"ok": True, "scanned": 0, "skipped_cached": skipped_cached, "candidates": []}
    found = []
    not_terms = 0
    CHUNK = 10
    for i in range(0, len(segs), CHUNK):
        chunk = segs[i:i + CHUNK]
        for item in _extract_terms_call([(s["source"], s["target"]) for s in chunk],
                                        req.model, project.get("domain")):
            _sc = _project_scope(project)
            known = _glossary_entry(item.get("src", ""), _sc)
            if known and _norm_key(known.get("tgt")) == _norm_key(item.get("tgt")):
                continue      # уже знаем ровно эту пару
            if not _looks_like_term(item.get("src", ""), item.get("tgt", "")):
                # Своей проверкой, а не по дельте глобального счётчика: тот
                # растёт и из сбора терминологии в рабочих потоках, и число
                # в ответе включало бы чужие отсевы.
                not_terms += 1
                continue
            c = _queue_term("extract", item.get("src", ""), item.get("tgt", ""),
                            cat=item.get("cat", ""), wasTgt=(known or {}).get("tgt", ""),
                            lang=_sc[0], domain=_sc[1], tenant=_sc[2], via="auto",
                            project=pid, model=_resolve_model(req.model or DEFAULT_OPENAI_MODEL)["id"])
            if c:
                found.append(c)
        for sg in chunk:
            sg["extracted_hash"] = sg.pop("_extract_hash", None)
    save_state(STATE)
    # «Отсеяно как не словарная запись» — отдельным числом. Молчаливый отсев
    # неотличим от «модель ничего не нашла», а платили за вызов одинаково.
    return {"ok": True, "scanned": len(segs), "skipped_cached": skipped_cached,
            "skipped_not_terms": not_terms, "candidates": found}


# ─── Параллельные вызовы внутри порции ───────────────────────────────
# Сегменты порции независимы, а каждый вызов модели — это 3-6 секунд ожидания
# сети, а не работы процессора. Гнать их по одному значит держать паузу длиной
# в прогон: 2670 сегментов по 5 с — почти четыре часа. Считаем параллельно.
#
# Потолок скромный: слишком агрессивная параллельность упирается в rate limit
# провайдера, и вместо ускорения получаются 429 и повторы. RUN_WORKERS можно
# поднять переменной окружения, если лимиты аккаунта позволяют.
RUN_WORKERS = max(1, min(int(os.environ.get("RUN_WORKERS", "6")), 16))


def _run_parallel(items: list, fn):
    """Выполнить fn для каждого элемента, сохранив порядок результатов.
    fn обязана ловить свои исключения сама: одна упавшая пара не должна
    ронять всю порцию."""
    if RUN_WORKERS <= 1 or len(items) <= 1:
        return [fn(x) for x in items]
    from concurrent.futures import ThreadPoolExecutor
    # Организация — в рабочие потоки: они свои и thread-local не наследуют.
    # Язык объяснений — там же и по той же причине: без него модель в потоках
    # порции пишет `why` на языке по умолчанию, а в основном потоке — на языке
    # человека, и один прогон даёт объяснения на двух языках вперемешку.
    tid = _current_tenant()
    lang = _explain_lang()

    def run(x):
        _JOB_TENANT.id = tid
        _JOB_LANG.code = lang
        return fn(x)
    with ThreadPoolExecutor(max_workers=min(RUN_WORKERS, len(items)),
                            thread_name_prefix="mcat-run") as pool:
        return list(pool.map(run, items))


# Флаг «текущий прогон просят остановить». Пакетные циклы сверяются с ним
# после каждого сегмента: раньше остановка ждала конца порции, а это до минуты
# на переводе и несколько минут на ремонте — пользователь считал кнопку мёртвой.
# Прогон один (см. фоновые прогоны), поэтому хватает одной переменной.
_ACTIVE_JOB: dict = {}


_STOP_CHECK = {"t": 0.0}


def _job_should_stop() -> bool:
    job = _ACTIVE_JOB.get("job")
    if not job:
        return False
    if job.get("stop"):
        return True
    if IS_WORKER and time.time() - _STOP_CHECK["t"] > 2:
        _STOP_CHECK["t"] = time.time()
        try:
            fresh = STORE.get_job(job["id"])
            if fresh and fresh.get("stop"):
                job["stop"] = True
            else:
                _job_persist(job)      # заодно прогресс — API показывает его из базы
        except Exception as e:
            print(f"[backend] стоп-флаг из базы не прочитан: {e}", file=sys.stderr)
    return bool(job.get("stop"))


def _text_hash(text: str) -> str:
    return hashlib.sha1((text or "").encode("utf-8")).hexdigest()[:16]


def _hard_mark_trusted(bc: dict) -> bool:
    """Верна ли отметка `judge_skipped: "hard"`, стоящая в записи.

    Своей версии правил верим сразу. У записей постарше отметка могла быть
    вынесена за «потерян термин» — прежде он считался жёсткой находкой на
    длинном оригинале. Теперь он не жёсткий ни при какой длине (см.
    `checks._hard_issue`), и верить такой отметке значит навсегда закрыть
    от судьи ровно те сегменты, ради которых он и нужен. Настоящая объективная
    находка (числа, единицы, отрицание, подмена стороны) отметку держит."""
    if bc.get("v") == _bc_version():
        return True
    return not (bool(bc.get("terms_lost"))
                and not any(h in r for r in (bc.get("reasons") or [])
                            for h in BACKCHECK_OBJECTIVE_REASONS))


def _judge_pending(seg: dict, above: bool = False) -> bool:
    """Судья ещё не смотрел на ЭТОТ перевод, и его вердикт что-то изменит.

    Один предикат на всех, кто задаёт этот вопрос: состав прогона
    (`_backcheck_cached`) и признак `needs_judge`, который уезжает браузеру
    (`_segment_for_client`). Раньше это были две разные формулы — и первая
    после введения версии правил начала отвечать «беру» там, где вторая
    отвечала «не нужен». Получались противоречащие числа под соседними
    кнопками: составной прогон брал 303 сегмента, а «Запустить только этот
    шаг» их не видел и смета судьи их не считала. Ровно та беда, ради которой
    состав вынесен на сервер.

    «Пропущен по зоне» решением не считается: зона зависит от сегмента
    (`_judge_zone`), и записанный отказ мог быть вынесен по прежней, узкой
    шкале. Поэтому зону считаем заново, а не верим отметке. «Молчание судьи»
    (`failed`) законченной проверкой тоже не считается — спросить ещё раз
    правильно, и теперь это дёшево: обратный перевод берётся готовым.

    `above=True` — разовое разрешение звать судью и ВЫШЕ потолка зоны
    (`judge_all` у прогона): балл выше 97 означает лишь, что детерминированные
    проверки довольны, а смысл при этом не читал никто — именно там живёт
    «беглое неверное слово» (monostable, sanguiferous), у которого другой
    меры нет. Это параметр ЗАПУСКА, а не новая политика: по умолчанию верх
    зоны стоит где стоял, а решение с прямой ценой принимает человек кнопкой,
    на которой цена написана. Сам потолок разрешения живёт в `_judge_zone` —
    здесь только передаём флаг. НИЗ зоны разрешение не открывает: ниже низа
    детерминированная проверка уже вынесла решение, и спорить там не о чем.
    Жёсткую отметку (числа, единицы, отрицание) оно тоже не обходит — её
    судья не вправе отменить ни в какой зоне."""
    bc = seg.get("backcheck") or {}
    score = bc.get("score")
    if score is None or bc.get("judged"):
        return False
    if bc.get("judge_skipped") == "hard" and _hard_mark_trusted(bc):
        return False
    lo, hi = _judge_zone(seg.get("source") or "", above)
    return lo <= score <= hi


def _backcheck_cached(seg: dict, mdl_id: str, use_judge: bool,
                      judge_all: bool = False) -> bool:
    """Сегмент уже проверен ЭТИМ переводом — считать нечего.
    Судья вне своей зоны не вызывается, поэтому сегмент за её границами полный
    даже без judged: иначе включённый судья гнал бы весь проект заново.
    `judge_all` — разрешение этого прогона звать судью и выше зоны
    (см. `_judge_pending`): с ним сегмент с баллом 98 без вердикта — работа.

    Модель здесь НЕ сравнивается, в отличие от termcheck, и это не упущение.
    У обратного перевода нет шкалы «сильнее — лучше»: ему нужна максимально
    буквальная модель, а сильная чинит кривой английский на лету и прячет
    ровно ту ошибку, которую проверка ищет. Значит «проверено другой моделью»
    здесь не хуже и не лучше — это просто проверено, и платить второй раз
    незачем. Кому нужен другой взгляд, тот запускает back-check отдельной
    кнопкой: она идёт со skip_cached=False и проверяет что попросили."""
    bc = seg.get("backcheck") or {}
    # Обрезанный текст: ровно так хеш и записывается. Сравнивая с необрезанным,
    # мы перезапускали бы платную проверку на каждом переводе с висящим пробелом,
    # а интерфейс при этом показывал бы её свежей.
    if bc.get("target_hash") != _text_hash((seg.get("target") or "").strip()):
        return False
    # ...но проверка СВОЕЙ ЖЕ работы проверкой не является. Раз модель здесь
    # больше не сравнивается, остаётся ровно один случай, когда готовый
    # обратный перевод переиспользовать нельзя: его делала та самая модель,
    # которая писала этот текст. Она вернёт свой же замысел, а не то, что
    # в тексте написано, — и на такой оценке стоит автоодобрение глоссария.
    # Раньше это лечилось сменой модели в списке; теперь смена модели ничего
    # не перезапускает, поэтому случай назван явно.
    if bc.get("model") and bc.get("model") == seg.get("provider"):
        return False
    # Балла нет — проверки не было. Случай штатный: без medical_qa
    # `run_backcheck` не зовётся вовсе, а запись всё равно пишется. Раньше это
    # ловил хвост «score is not None and ...»; в _judge_pending переносить
    # нельзя — там «балла нет» правильно означает «судья не нужен».
    if bc.get("score") is None:
        return False
    if not use_judge or bc.get("judged"):
        return True
    return not _judge_pending(seg, judge_all)


def _segment_for_client(seg: dict) -> dict:
    """Сегмент с производными признаками stale/tried. Хеши считаются здесь:
    браузеру sha1 не пересчитать, а без них он не отличит устаревшую проверку
    от актуальной."""
    try:
        out = dict(seg)
    except RuntimeError:
        out = dict(seg)          # правка из фонового потока длится микросекунды
    # Имена ответственных — для карточки: в сегменте лежит идентификатор
    # (`confirmedBy`, `editedBy`), а списка пользователей браузеру не дают.
    for k, name_k in (("confirmedBy", "confirmedByName"), ("editedBy", "editedByName")):
        if isinstance(out.get(k), int):
            out[name_k] = _user_label(out[k])
    cur = _text_hash(out.get("target") or "")
    # back-check и termcheck кладут хеш ОБРЕЗАННОГО текста, ремонт — сырого.
    # Сравнивать надо тем же способом, каким писали: иначе перевод с висящим
    # пробелом показывался бы в UI устаревшим, а _machine_clean считал бы его
    # свежим и пускал бы в глоссарий.
    cur_trimmed = _text_hash((out.get("target") or "").strip())
    bc, tc, rp = out.get("backcheck"), out.get("termcheck"), out.get("repair")
    qa = out.get("qa_result")
    if bc:
        # needs_judge считает СЕРВЕР, а не браузер. Зона вызова судьи зависит
        # от длины оригинала (_judge_zone), а длина считается русской
        # морфологией из checks — повторять её в .jsx значит однажды
        # разойтись с сервером в том, кого судья ещё не смотрел, и получить
        # два разных состава под соседними кнопками. То же правило, что у
        # _backcheck_cached, и выведено оно здесь один раз.
        out["backcheck"] = {
            **bc, "stale": bc.get("target_hash") != cur_trimmed,
            "needs_judge": _judge_pending(out)}
    if tc:
        out["termcheck"] = {**tc, "stale": tc.get("target_hash") != cur_trimmed}
    rv = out.get("review")
    if rv:
        # Свежесть считает СЕРВЕР тем же `_review_stale`, что и прогон: он
        # знает и про версию вопросов, и про правку ОРИГИНАЛА, а браузеру
        # ни того, ни другого не вычислить. Повтори он «сравню хеш перевода» —
        # карточка показывала бы свежим вердикт, который прогон считает
        # протухшим.
        # `candidate` отдаём ТОЛЬКО у неприменённой правки. У применённой он
        # буква в букву равен переводу — это вторая копия текста на каждую
        # строку без единого читателя. А вот у отклонённой он и есть совет,
        # ради которого человека зовут в карточку: экран «Анализ» обещает
        # «совет в карточке сегмента», и обещание надо выполнять.
        out["review"] = {k: v for k, v in rv.items()
                         if k != "candidate" or not rv.get("applied")}
        out["review"]["stale"] = _review_stale(seg)
        # Ручательство считает СЕРВЕР: соло-состав ремонта в браузере считает
        # находки termcheck сам, а `_repairable` их под ручательством не видит
        # — без признака кнопка обещала бы N сегментов, а делала меньше.
        out["review"]["vouches"] = _review_vouches(seg)
        # «Текст написала ревизия» — мнение судьи в ремонт не идёт
        # (`_repair_findings`); соло-состав и кнопка «Починить» в браузере
        # обязаны видеть тот же список, иначе строка обещает N, а сервер
        # берёт меньше.
        out["review"]["wrote"] = _review_wrote(seg)
        # Подписи вето собирает СЕРВЕР: в `veto` лежат внутренние ключи
        # (`gloss`, `hard`), которых нет ни в одном словаре, — на экране
        # это была латиница посреди узбекской фразы. Таблица
        # REVIEW_VETO_LABELS теперь используется, а не лежит мёртвой.
        out["review"]["vetoLabels"] = [REVIEW_VETO_LABELS.get(k, k)
                                       for k in (rv.get("veto") or [])]
        # `flagged` — тот же признак, что корзина `human.reviewFlagged`: одно
        # правило, одно место. Иначе таблица красит по своему порогу и своему
        # чтению кода, а экран «Анализ» числит по серверному — на записях без
        # `code` (155 вердиктов первого прогона) они расходятся сразу.
        code = _review_code(rv)
        out["review"]["flagged"] = bool(
            not rv.get("applied") and not rv.get("undone")
            and not out["review"]["stale"]
            and (code == REVIEW_VETOED
                 or (code == REVIEW_OK and rv.get("score") is not None
                     and rv["score"] <= REVIEW_FLAG_SCORE)))
    if qa:
        # Тот же признак и у Medical QA: без него карточка прогона показывала
        # весь проект и после того, как всё уже проверено. Хеш sha1 браузеру
        # не посчитать, поэтому производную считаем здесь.
        out["qa_result"] = {**qa, "stale": qa.get("target_hash") != cur_trimmed}
    if rp:
        # `acceptable` считает СЕРВЕР по той же причине, что и `tried`: правило
        # «отмену держал только балл, а термины стали чище» разобрано в
        # `_repair_score_vetoed`, и повтори его браузер — кнопка «Принять»
        # предлагалась бы там, где эндпоинт отвечает 400. Тот же закон, что
        # у TERMCHECK_ACTIONABLE: одно правило — одно место.
        out["repair"] = {**rp, "tried": rp.get("source_hash") == cur,
                         # Несостоявшийся заход — только пока текст тот же:
                         # после ручной правки подпись «сбой перепроверки»
                         # врала бы, текст как раз менялся.
                         "retryable": bool(rp.get("retryable")
                                           and rp.get("attemptHash") == cur),
                         "acceptable": _repair_score_vetoed(seg)}
    # Советы арбитра, которые есть чем исполнить, — признак для кнопки
    # «Применить» в карточке сегмента. Считает сервер (см. `_ctx_advices`).
    if out.get("termContext"):
        adv = _ctx_advices(seg)
        if adv:
            out["ctxAdvice"] = adv
    return out


def _project_for_client(project: dict) -> dict:
    """Копия проекта с производными признаками у каждого сегмента."""
    return {**project, "segments": [_segment_for_client(s) for s in list(project["segments"])]}


# ── Бесплатный пересчёт сохранённых оценок back-check ────────────────
# Оценка живёт в сегменте и переживает всё: хеш перевода сторожит ТЕКСТ, а
# правила подсчёта он не сторожит никак. Когда 25.08 потерю термина ограничили
# приказными записями, а обрезку основ заменили сравнением форм слова, 798
# сегментов боевого проекта остались с оценками, посчитанными по-старому: текст
# не менялся, `_backcheck_cached` считал их свежими, и пересчитать их не мог
# уже НИКТО. У 303 из них судья был погашен как при жёсткой находке, то есть
# претензию, выведенную морфологией, стало нечем обжаловать.
#
# Пересчёт возможен без единого вызова модели, и это здесь главное: обратный
# перевод, оригинал и косинус лежат в самой записи, глоссарий берётся нынешний,
# а `run_backcheck` — чистая функция от них. Ничего не выдумывается: тот же
# текст, та же мера, нынешние правила. Поэтому косинус передаётся ГОТОВЫМ
# числом, а не `semantic_fn`: ленивый вариант полез бы считать эмбеддинги,
# то есть пересчёт стал бы платным.
#
# Вердикт судьи — оплаченная работа, и он применяется заново поверх свежего
# балла, а не выбрасывается.
BACKCHECK_RESCORE_DIR = DATA_DIR / "backups"
# Порог, ниже которого сегмент не считается машинно-чистым и донором глоссария
# не становится. Берётся из политики автоодобрения, а не пишется числом заново:
# два порога с одним смыслом однажды разойдутся, и отчёт начнёт обещать не то,
# что потом произойдёт.
_DONOR_MIN = AUTO_APPROVE_DEFAULT["backcheck_min"]


def _rescore_backcheck(seg: dict, project: dict,
                       hits_cache: Optional[dict] = None) -> Optional[dict]:
    """Новая запись back-check для сегмента или None, если считать нечем.

    Нечем — это либо проверки не было (`score is None`), либо не сохранён
    обратный перевод: выдумать его нельзя, а посчитать балл без него значит
    выдать за измерение то, чего не измеряли."""
    bc = seg.get("backcheck") or {}
    back = (bc.get("back") or "").strip()
    if not checks_mod or bc.get("score") is None or not back:
        return None
    source = seg.get("source") or ""
    # `_verified_hits` — тринадцать миллисекунд на сегмент, и это самое дорогое
    # здесь. У учебника заголовки повторяются, а требования глоссария зависят
    # только от оригинала — считаем их один раз на текст.
    if hits_cache is None:
        hits = _verified_hits(source, project)
    elif source in hits_cache:
        hits = hits_cache[source]
    else:
        hits = hits_cache[source] = _verified_hits(source, project)
    # Кэш общий на текст, а вердикт арбитра — свой у КАЖДОГО сегмента,
    # поэтому фильтр применяется после кэша, а не внутри него.
    res = checks_mod.run_backcheck(
        source, back, _hits_for_score(seg, hits),
        semantic=bc.get("semantic"),
        domain=project.get("domain"), src_lang=project.get("src", "RU"))
    if res.get("score") is None:                             # pragma: no cover
        return None
    judged = bool(bc.get("judged") and bc.get("judge"))
    if judged:
        res = checks_mod.apply_judge_verdict(res, bc["judge"])
    # Отметка о судье выводится заново — она и была главной потерей. «Молчание
    # судьи» пересчёту не подлежит: его звали, он не ответил, и этот факт
    # правилами подсчёта не отменяется.
    skipped = bc.get("judge_skipped")
    if judged:
        skipped = None
    elif skipped != "failed":
        lo, hi = _judge_zone(source)
        skipped = ("zone" if not (lo <= res["score"] <= hi)
                   else "hard" if res.get("hard") else None)
    out = dict(bc)
    out.update({
        "score": res["score"], "band": res.get("band"), "recall": res.get("recall"),
        "semantic": res.get("semantic"), "reasons": res.get("reasons", []),
        "terms_lost": res.get("terms_lost", []),
        "judge": res.get("judge") if judged else bc.get("judge"),
        "judged": judged, "judge_skipped": skipped,
        "v": _bc_version(),
    })
    return out


def _rescore_backchecks(state: dict, dry_run: bool = False,
                        pid: Optional[int] = None, force: bool = False) -> dict:
    """Пересчитать сохранённые оценки back-check по нынешним правилам.

    `force=False` — только записи, посчитанные ЧУЖОЙ версией правил; повторный
    запуск при этом не делает ничего.
    `force=True` нужен после правки глоссария: версия правил та же, а вот
    приказных записей стало меньше или больше, и претензия «потерян термин»
    в сегментах ссылается на записи, которые приказом уже не являются.

    Отчёт называет не только «сколько», но и КУДА сдвинулось: полосы до и
    после. «Пересчитано 789» без распределения — не отчёт, а отговорка."""
    before: dict = {}
    after: dict = {}
    scanned = rescored = changed = freed = 0
    clean_before = clean_after = no_semantic = 0
    hits_cache: dict = {}
    saved: dict = {}
    for project in state.get("projects", []):
        if pid is not None and project.get("id") != pid:
            continue
        for seg in project.get("segments", []):
            bc = seg.get("backcheck") or {}
            if bc.get("score") is None:
                continue
            scanned += 1
            if not force and bc.get("v") == _bc_version():
                continue
            new = _rescore_backcheck(seg, project, hits_cache)
            if new is None:
                continue
            rescored += 1
            # «Было» считаем только у тех, кого пересчитали: иначе пара
            # «было/стало» посчиталась бы на разных множествах.
            if _machine_clean(seg, _DONOR_MIN) is None:
                clean_before += 1
            b0, b1 = checks_mod.band_of(bc["score"]), checks_mod.band_of(new["score"])
            before[b0] = before.get(b0, 0) + 1
            after[b1] = after.get(b1, 0) + 1
            if new["score"] != bc["score"]:
                changed += 1
            if bc.get("judge_skipped") == "hard" and new.get("judge_skipped") != "hard":
                freed += 1
            if new.get("semantic") is None:
                # Косинус пересчёт не покупает — и не может: `semantic_fn`
                # звал бы эмбеддинги, то есть пересчёт стал бы платным.
                # У записей, где прежде стояла жёсткая находка, косинуса нет
                # вовсе (его тогда не считали), и теперь он БЫ повлиял. Значит
                # у стольких-то сегментов балл посчитан без него — и сказать
                # это надо, а не молчать: платный перезапуск дал бы им другое
                # число.
                no_semantic += 1
            # Сегмент, чей балл поднялся выше порога донора, начинает годиться
            # в источники для глоссария. Это правильное следствие — оценка была
            # занижена, — но молчать о нём нельзя: расширение донорской базы
            # человек обязан увидеть числом, а не обнаружить через месяц
            # по машинным записям в медицинском глоссарии. Само по себе
            # автоодобрение отсюда не запускается: у него своя команда,
            # свой dry_run и свои пороги согласия.
            # Запись подставляется временно — чтобы `_machine_clean` посчитал
            # «стало» по ней. При разборе она обязана вернуться на место при
            # ЛЮБОМ исходе: STATE живёт в памяти процесса, и «посчитали и ничего
            # не изменили» должно означать ровно это, а не «изменили и упали».
            seg["backcheck"] = new
            try:
                if _machine_clean(seg, _DONOR_MIN) is None:
                    clean_after += 1
            finally:
                if dry_run:
                    seg["backcheck"] = bc
                else:
                    saved.setdefault(str(project["id"]), {})[str(seg["id"])] = bc
    # Копия прежних записей — то, чем работает откат
    # (`POST /api/backcheck/rescore/{stamp}/undo`). Почасовой копии состояния
    # тут мало: она снимок ВСЕГО состояния, и откатить ею пересчёт значит
    # выбросить вместе с ним всю работу, сделанную после. Результат пересчёта
    # восстановить можно и так — входные данные на месте, — а прежние оценки
    # после перезаписи взять было бы неоткуда.
    backup = None
    if saved and not dry_run:
        try:
            BACKCHECK_RESCORE_DIR.mkdir(exist_ok=True)
            stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
            path = BACKCHECK_RESCORE_DIR / f"backcheck-{stamp}.json"
            path.write_text(json.dumps({"at": stamp, "version": _bc_version(),
                                        "projects": saved}, ensure_ascii=False),
                            encoding="utf-8")
            backup = str(path)
            # Подрезаем, как почасовые копии состояния: в файле лежат целые
            # записи вместе с обратными переводами, а `force` можно жать
            # сколько угодно.
            for stale in sorted(BACKCHECK_RESCORE_DIR.glob("backcheck-*.json"))[:-_BACKUP_KEEP]:
                stale.unlink(missing_ok=True)
        except Exception as e:                               # pragma: no cover
            print(f"[backend] пересчёт back-check: копия не сохранена: {e}", file=sys.stderr)
    return {"scanned": scanned, "rescored": rescored, "changed": changed,
            "freed_judge": freed, "before": before, "after": after,
            "machine_clean": {"before": clean_before, "after": clean_after,
                              "min_score": _DONOR_MIN},
            "no_semantic": no_semantic,
            "dry_run": dry_run, "backup": backup, "version": _bc_version()}


class RescoreRequest(BaseModel):
    dry_run: bool = True
    force: bool = False


@app.post("/api/projects/{pid}/backcheck/rescore")
def rescore_backchecks(pid: int, req: RescoreRequest = RescoreRequest()):
    """Пересчитать оценки back-check проекта по нынешним правилам. Бесплатно.

    Без `force` пересчитываются записи, посчитанные ПРЕЖНИМИ правилами. Сам
    по себе, проходом при старте, пересчёт не идёт намеренно: он зовёт
    `_verified_hits` на каждый сегмент, а это тринадцать миллисекунд —
    полминуты заблокированного единственного воркера на проекте в 2700 строк,
    то есть отвалившийся по таймауту деплой. И правило тут то же, что
    у выноса глоссария и пачек автоодобрения: посчитали, показали числа,
    человек нажал.

    `force` нужен для второго случая: поменялся ГЛОССАРИЙ. Понижение записи до подсказки
    убирает её из требований к переводу, но в сохранённых оценках претензия
    «потерян термин» по ней остаётся и продолжает держать балл — на это и
    нужен `force`. Вызов ничего не переводит и не проверяет: он пересчитывает
    уже измеренное."""
    _guard_project_write(pid)
    get_project(pid)
    # Прогон пишет `seg["backcheck"]` из рабочего потока, а этот обработчик —
    # обычный def, то есть живёт в пуле FastAPI и идёт параллельно. Между
    # чтением записи и её заменой свежая ОПЛАЧЕННАЯ проверка была бы затёрта
    # пересчитанной копией старой — и следующий прогон заплатил бы за неё
    # заново. Два писателя одного места: тот же отказ, что у `/source`
    # и `/images/forget`.
    if not req.dry_run and _job_live(pid):
        raise HTTPException(409, "Идёт прогон — пересчёт подождёт: он заменяет "
                                 "те же записи, которые прогон сейчас пишет")
    report = _rescore_backchecks(STATE, dry_run=req.dry_run, pid=pid, force=req.force)
    if not req.dry_run and report["rescored"]:
        # Оба кэша — руками. Экран «Анализ» протух бы и сам (в его отпечаток
        # входят балл и длина списка потерянных терминов), а вот отчёт
        # о соответствии глоссарию — нет: его отпечаток это тексты переводов
        # плюс поколение глоссария, и пересчёт не трогает ни то, ни другое.
        # Между тем строка «ремонт уже не берёт» (`impact["futile"]`) считается
        # через `_repair_findings`, то есть по back-check, — и без явной чистки
        # осталась бы доперегонной навсегда.
        _IMPACT_CACHE.pop(pid, None)
        _ANALYSIS_CACHE.pop(pid, None)
        save_state(STATE)
    return {"ok": True, **report}


@app.post("/api/backcheck/rescore/{stamp}/undo")
def rescore_backchecks_undo(stamp: str):
    """Вернуть оценки back-check такими, какими они были до пересчёта {stamp}.

    Массовая правка без отката недопустима — тот же закон, что у пачек
    автоодобрения и у выноса глоссария. Почасовой копии состояния тут мало:
    она снимок ВСЕГО состояния, и откатить ею пересчёт значит выбросить
    вместе с ним всю работу, сделанную после.

    Откат возвращает числа, а не решения: пересчёт ничего человеческого
    не трогает. Поэтому и повторный пересчёт после отката сделает то же самое —
    если пересчёт оказался неверным, откатывать надо и код."""
    if not re.fullmatch(r"\d{8}-\d{6}", stamp or ""):
        # Имя файла склеивается из этой строки, и подставлять в путь что
        # угодно из URL нельзя.
        raise HTTPException(400, "Неверная отметка времени")
    path = BACKCHECK_RESCORE_DIR / f"backcheck-{stamp}.json"
    if not path.exists():
        raise HTTPException(404, "Копия пересчёта не найдена")
    data = json.loads(path.read_text(encoding="utf-8"))
    restored, missing, skipped = 0, 0, 0
    for pid_s, saved in (data.get("projects") or {}).items():
        project = next((p for p in STATE.get("projects", [])
                        if str(p.get("id")) == str(pid_s)), None)
        if project is None:
            missing += len(saved)
            continue
        if _job_live(project["id"]):
            raise HTTPException(409, "Идёт прогон — откат подождёт")
        by_id = {str(s.get("id")): s for s in project.get("segments", [])}
        for sid, bc in saved.items():
            seg = by_id.get(str(sid))
            if seg is None:
                # Сегмент снесли после пересчёта — возвращать оценку некуда.
                # Молчать нельзя: «откачено 700 из 798» и «откачено 798» —
                # разные ответы.
                missing += 1
                continue
            cur = seg.get("backcheck") or {}
            # Возвращаем, только если в сегменте до сих пор лежит результат
            # ЭТОГО пересчёта. Пересчёт не трогает ни обратный перевод, ни его
            # модель, ни время проверки — значит настоящая проверка, прошедшая
            # после него, отличается хотя бы одним из этих полей. Подставить
            # поверх неё старую оценку значит молча выбросить оплаченную
            # работу; тот же закон, что у отката правок ремонта
            # (`_repair_tried`): прежний текст возвращают, только если в
            # сегменте стоит именно тот, что клали.
            if (cur.get("at") != bc.get("at") or cur.get("back") != bc.get("back")
                    or cur.get("target_hash") != bc.get("target_hash")):
                skipped += 1
                continue
            seg["backcheck"] = bc
            restored += 1
        _IMPACT_CACHE.pop(project["id"], None)
        _ANALYSIS_CACHE.pop(project["id"], None)
    if restored:
        save_state(STATE)
    return {"ok": True, "restored": restored, "missing": missing,
            "skipped": skipped, "stamp": stamp}


def _backcheck_model(seg: dict, requested: Optional[str]) -> str:
    """Модель обратного перевода для ЭТОГО сегмента.

    Обратный перевод моделью-автором текста — не проверка: она вернёт свой
    замысел, а не то, что в тексте написано, и _backcheck_cached такой
    результат действительным не признаёт. Раньше совпадение просто повторяло
    вызов той же моделью: прогон платил при каждом запуске, а зачётной
    проверки не получал никогда — вечная переплата за недействительный
    результат. Совпала — берём запасную, столь же буквальную и дешёвую."""
    mid = _resolve_model(requested or BACKCHECK_DEFAULT_MODEL)["id"]
    provider = seg.get("provider") or ""
    if mid != provider:
        return mid
    for alt in (BACKCHECK_DEFAULT_MODEL, BACKCHECK_FALLBACK_MODEL):
        if alt != provider:
            return alt
    return mid          # недостижимо: две запасных не совпадают между собой


def _run_segment_backcheck(seg: dict, project: dict, model: Optional[str] = None,
                           use_judge: bool = False, judge_model: Optional[str] = None,
                           harvest: bool = True, judge_all: bool = False) -> dict:
    """Обратный перевод сегмента + оценка соответствия оригиналу.
    Результат кладётся в seg['backcheck'] вместе с хешем перевода — по нему
    повторный прогон понимает, что пересчитывать нечего.
    `judge_all` — разрешение этого прогона звать судью и выше зоны."""
    target_text = (seg.get("target") or "").strip()
    if not target_text:
        return {"ok": False, "error": "Сегмент ещё не переведён"}

    mdl_id = _backcheck_model(seg, model)
    # Единственная недостача — судья? Тогда обратный перевод уже есть и лежит
    # в сегменте: платить за него второй раз незачем, а выбросить готовый
    # и заказать новый — значит ещё и потерять тот текст, по которому считали
    # прежний балл. Условие узкое намеренно: берём готовое ровно тогда, когда
    # БЕЗ судьи эта проверка считалась бы законченной (_backcheck_cached), а
    # С судьёй — нет. Прогон «другой моделью» (skip_cached=False) сюда не
    # попадает: там первое условие ложно, и обратный перевод делается заново.
    # Из-за чего написано: зона судьи открылась вниз для коротких сегментов,
    # и без этой ветки каждый такой сегмент оплачивал бы обратный перевод
    # заново — работу, результат которой лежал рядом.
    have = seg.get("backcheck") or {}
    reuse = bool(use_judge and (have.get("back") or "").strip()
                 and _backcheck_cached(seg, mdl_id, False)
                 and not _backcheck_cached(seg, mdl_id, True, judge_all))
    back = ""
    if reuse:
        back = have["back"]
        # Модель называем ту, что обратный перевод и делала: подписать чужой
        # работой ту, которую сейчас выбрали в списке, значит соврать о
        # происхождении оценки — и сломать правило «проверял тот, кто переводил».
        mdl_id = have.get("model") or mdl_id
    else:
        try:
            # Обратный перевод делает ДРУГАЯ модель, а не бесплатный движок: на нём
            # весь смысл проверки. Движок переводил дословно и одинаково хорошо
            # возвращал и верный термин, и кальку — балл получался ни о чём.
            back = _openai_translate(target_text, project["tgt"], project["src"],
                                     model=mdl_id, literal=True)
        except Exception as e:
            print(f"[backend] backcheck seg#{seg.get('id')}: {e}", file=sys.stderr)
            return {"ok": False, "error": str(e)}

    if not (back or "").strip():
        return {"ok": False, "error": "Обратный перевод не получен"}

    source_text = seg.get("source", "")
    # ТОЛЬКО приказные записи, тем же расчётом, что у _gloss_misses и
    # /glossary-impact. Раньше сюда уходил полный _get_context — и подсказки
    # автоимпорта тоже, — а back-check снижал балл за то, что обратный перевод
    # не сохранил ПОДСКАЗКУ. Модели её игнорировать разрешено прямо в промпте
    # («use these exact translations» только для verified), и наказывать за
    # принятое разрешение нельзя. На боевом проекте это давало 56 спорных
    # сегментов из 67: балл держали «лёгких», «высокой», «оценка», «метод» —
    # падежные формы обычных слов из массового импорта, которым сверка смысла
    # УЖЕ проставила rule: false. Ещё и судья к таким сегментам не приходил:
    # потерянный термин гасил его как жёсткую находку.
    gloss_hits = _hits_for_score(seg, _verified_hits(source_text, project))
    # semantic_fn, а не готовое число: косинус учитывается только при отсутствии
    # жёстких находок, и там, где он всё равно не повлияет, платить за эмбеддинги
    # незачем. Решение принимает medical_qa — он и знает состав находок.
    res = (checks_mod.run_backcheck(
        source_text, back, gloss_hits,
        semantic_fn=lambda: _semantic_similarity(source_text, back),
        domain=project.get("domain"), src_lang=project.get("src", "RU")) if checks_mod else {})

    # Судья — только для средней зоны: наверху и внизу шкалы вопрос уже решён.
    # `judge_all` открывает ВЕРХ (потолок считает сам `_judge_zone`): выше
    # потолка вопрос решён только детерминированно, а смысл там не читал
    # никто. Низ закрыт и с разрешением — там решение настоящее.
    judged, judge_skipped = False, None
    if use_judge and checks_mod and res.get("score") is not None:
        lo, hi = _judge_zone(source_text, judge_all)
        if not (lo <= res["score"] <= hi):
            judge_skipped = "zone"
        elif res.get("hard"):
            # Объективное расхождение (числа, единицы, отрицание, сторона) уже
            # найдено детерминированно. Судья такую находку отменить не может —
            # его вердикт ничего не изменит, а вызов стоит денег.
            judge_skipped = "hard"
        else:
            verdict = _openai_judge(source_text, back, judge_model,
                                    project.get("domain"), project.get("src", "RU"))
            if verdict:
                res = checks_mod.apply_judge_verdict(res, verdict)
                judged = True
            else:
                # Судью звали, он не ответил (сеть, лимит, неразобранный ответ).
                # Отметку ставим, чтобы это было видно в сегменте, но
                # ЗАКОНЧЕННОЙ проверкой она не считается (_backcheck_cached
                # признаёт безусловной только "hard"): молчание судьи — это
                # «не знаю», и спросить его ещё раз правильно. Дорого это
                # больше не стоит: обратный перевод берётся готовым.
                judge_skipped = "failed"

    seg["backtranslated_ru"] = back
    seg["backcheck"] = {
        "score": res.get("score"),
        "band": res.get("band"),
        "recall": res.get("recall"),
        "semantic": res.get("semantic"),
        "reasons": res.get("reasons", []),
        "terms_lost": res.get("terms_lost", []),
        # Объективная находка (числа, единицы, отрицание, подмена стороны —
        # BACKCHECK_HARD_TYPES). Кладём в запись, потому что её спрашивает
        # ремонт: падение балла он больше не считает отменой, если правка
        # почистила термины, — а вот жёсткая находка остаётся вето при любом
        # улучшении терминологии. Считать её по `reasons` подстрокой нельзя:
        # это ровно то сравнение с русской фразой, от которого заведены
        # CLEAN_* коды. Поле новое, у старых записей его нет; читает его
        # только тот, кто сам эту запись сейчас и написал (перепроверка
        # внутри ремонта), поэтому None ни во что не превращается.
        "hard": bool(res.get("hard")),
        "judge": res.get("judge"),
        "back": back,
        "model": mdl_id,
        "judged": judged,
        "judge_skipped": judge_skipped,
        "target_hash": _text_hash(target_text),
        # Правила, по которым посчитан балл. По ним идёт бесплатный пересчёт,
        # когда правила меняются: без клейма старая оценка неотличима от свежей.
        "v": _bc_version(),
        "at": datetime.now().strftime("%Y-%m-%d %H:%M"),
    }
    # Обе оценки на месте — можно собрать терминологию без участия человека.
    # Порядок прогонов пользователь выбирает сам, поэтому сбор висит на обоих.
    harvested = _harvest_if_clean(seg, project) if harvest else []
    return {"ok": True, "back": back, "backcheck": seg["backcheck"],
            "queued": [c["id"] for c in harvested]}


def _termcheck_cached(seg: dict, mdl_id: str) -> bool:
    """Тот же перевод уже разобран моделью НЕ СЛАБЕЕ запрошенной — платить
    второй раз незачем.

    Сравнение по рангу, а не по равенству id: у проверки терминов шкала есть,
    и вердикт сильной модели слабой не перезаписывается. Иначе достаточно было
    сменить модель в выпадающем списке, чтобы тысяча уже проверенных Sol
    сегментов ушла на перепроверку Terra — оплаченную и худшую по качеству.
    Обратное направление разрешено: усилить проверку можно всегда."""
    tc = seg.get("termcheck") or {}
    if tc.get("target_hash") != _text_hash((seg.get("target") or "").strip()):
        return False
    # Пропущенный как «нечего проверять» не зависит от модели — пересчитывать нечего
    if tc.get("model") == "skip":
        return True
    return _rank_not_weaker(tc.get("model") or "", mdl_id)


def _termcheck_trivial(source: str, target: str) -> Optional[str]:
    """Сегменты, где проверять нечего, — без вызова модели.
    Терминов не бывает там, где нет слов; а совпадение перевода с оригиналом
    (числа, латинские обозначения, «IFN-γ») — это не терминологическая ошибка."""
    if not re.search(r"[A-Za-zА-Яа-яЁё]{3}", target or ""):
        return "в переводе нет слов — только числа или обозначения"
    if _norm_key(source) == _norm_key(target):
        return "перевод совпадает с оригиналом — переводить нечего"
    return None


def _run_segment_termcheck(seg: dict, project: dict, model: Optional[str] = None,
                           harvest: bool = True) -> dict:
    """Проверка терминологии перевода + кандидаты в глоссарий из находок.

    Находка с предложенной заменой — это готовая пара «термин оригинала →
    правильный термин», то есть ровно то, что нужно глоссарию. Кладём её в ту
    же очередь кандидатов, что и расхождения при подтверждении: одно место,
    где человек принимает терминологические решения."""
    target = (seg.get("target") or "").strip()
    if not target:
        return {"ok": False, "error": "Сегмент ещё не переведён"}
    trivial = _termcheck_trivial(seg.get("source", ""), target)
    if trivial:
        seg["termcheck"] = {"findings": [], "severity": "none", "model": "skip",
                            "note": trivial,
                            "domain": _resolve_domain(project.get("domain"))["id"],
                            "target_hash": _text_hash(target),
                            "at": datetime.now().strftime("%Y-%m-%d %H:%M")}
        return {"ok": True, "termcheck": seg["termcheck"], "queued": [], "skipped": trivial}
    res = _openai_termcheck(seg.get("source", ""), target,
                            project.get("src", "RU"), project.get("tgt", "EN"),
                            project.get("domain"), model)
    if res is None:
        return {"ok": False, "error": "Модель не ответила"}
    findings = res["findings"]
    worst = next((sev for sev in TERMCHECK_SEVERITY
                  if any(f["severity"] == sev for f in findings)), "none")
    queued = []
    for f in findings:
        # В глоссарий просятся только уверенные находки с обеими сторонами пары
        if f["severity"] in ("critical", "major") and f["src_term"] and f["suggestion"]:
            _sc = _project_scope(project)
            c = _queue_term("audit", f["src_term"], f["suggestion"],
                            wasTgt=f["tgt_term"], project=project["id"], segment=seg["id"],
                            lang=_sc[0], domain=_sc[1], tenant=_sc[2], via="auto",
                            note=f["why"], model=res["model"],
                            sampleSrc=seg.get("source", "")[:240], sampleTgt=target[:240])
            if c:
                queued.append(c["id"])
    seg["termcheck"] = {
        "findings": findings,
        "severity": worst,
        "model": res["model"],
        "domain": _resolve_domain(project.get("domain"))["id"],
        "target_hash": _text_hash(target),
        "at": datetime.now().strftime("%Y-%m-%d %H:%M"),
    }
    if harvest:
        queued += [c["id"] for c in _harvest_if_clean(seg, project)]
    # Жалоба на приказной термин — не приговор записи, а повод переспросить
    # у сверки смысла (см. _note_term_disputes).
    disputed = _note_term_disputes(seg, project)
    return {"ok": True, "termcheck": seg["termcheck"], "queued": queued,
            "disputedTerms": disputed}


class TermcheckRequest(BaseModel):
    model: Optional[str] = None


@app.post("/api/segments/{pid}/{sid}/termcheck")
def termcheck_segment(pid: int, sid: int, req: TermcheckRequest = TermcheckRequest()):
    """Обычный def: внутри блокирующий вызов модели (см. batch_translate)."""
    if not os.environ.get("OPENAI_API_KEY"):
        raise HTTPException(503, "Проверка терминологии требует ключ OpenAI")
    seg = get_segment(pid, sid)
    project = get_project(pid)
    result = _run_segment_termcheck(seg, project, req.model)
    if result.get("ok"):
        save_state(STATE)
        return result
    raise HTTPException(502, result.get("error", "Проверка не удалась"))


class TermcheckBatchRequest(BaseModel):
    segment_ids: Optional[List[int]] = None
    limit: int = 10
    model: Optional[str] = None
    skip_cached: bool = True


@app.post("/api/projects/{pid}/termcheck/batch")
def termcheck_batch(pid: int, req: TermcheckBatchRequest):
    """Порционно, как back-check: клиент гоняет порции по 10, чтобы один
    запрос не жил дольше таймаута прокси."""
    if not os.environ.get("OPENAI_API_KEY"):
        raise HTTPException(503, "Проверка терминологии требует ключ OpenAI")
    project = get_project(pid)
    id_filter = set(req.segment_ids) if req.segment_ids is not None else None
    mdl_id = _resolve_model(req.model or TERMCHECK_DEFAULT_MODEL)["id"]

    candidates, skipped_cached = [], 0
    for seg in project["segments"]:
        if id_filter is not None and seg["id"] not in id_filter:
            continue
        if not (seg.get("target") or "").strip():
            continue
        if req.skip_cached and _termcheck_cached(seg, mdl_id):
            skipped_cached += 1
            continue
        candidates.append(seg)

    limit = max(1, min(req.limit, 100))
    remaining_after = max(0, len(candidates) - limit)
    targets = candidates[:limit]

    processed, errors, flagged = [], [], 0
    groups: dict = {}
    order: list = []
    for seg in targets:
        pair = (_norm_key(seg.get("source")), _norm_key(seg.get("target")))
        if pair not in groups:
            groups[pair] = []
            order.append(pair)
        groups[pair].append(seg)
    duplicates, skipped_trivial = 0, 0

    def _tc_one(pair):
        seg = groups[pair][0]
        if _job_should_stop():
            return {"pair": pair, "skip": True}
        try:
            r = _run_segment_termcheck(seg, project, req.model)
            return {"pair": pair, "ok": bool(r.get("ok")), "error": r.get("error"),
                    "trivial": bool(r.get("skipped"))}
        except Exception as e:
            print(f"[backend] termcheck batch seg#{seg['id']}: {e}", file=sys.stderr)
            return {"pair": pair, "ok": False, "error": str(e)}

    for res in _run_parallel(order, _tc_one):
        segs = groups[res["pair"]]
        if res.get("skip"):
            continue
        if not res.get("ok"):
            errors.append({"id": segs[0]["id"], "error": res.get("error", "unknown")})
            continue
        lead = segs[0]
        processed.append(lead["id"])
        if res.get("trivial"):
            skipped_trivial += 1
        if lead["termcheck"]["findings"]:
            flagged += 1
        for sg in segs[1:]:
            sg["termcheck"] = json.loads(json.dumps(lead["termcheck"]))
            processed.append(sg["id"])
            duplicates += 1
            if sg["termcheck"]["findings"]:
                flagged += 1
    save_state(STATE)
    return {"ok": True, "processed": processed, "count": len(processed),
            "flagged": flagged, "remaining": remaining_after,
            "skipped_cached": skipped_cached, "duplicates": duplicates,
            "skipped_trivial": skipped_trivial, "errors": errors, "model": mdl_id}


# ─── Автоматический ремонт сегмента ──────────────────────────────────
# Чинит по КОНКРЕТНЫМ находкам, а не «переведи получше»: список претензий
# приходит из back-check (числа, единицы, отрицания, потерянные термины,
# вердикт судьи) и из проверки терминологии (кальки с предложенной заменой).
#
# Три предохранителя, без которых цикл начинает угождать метрике:
#   1. чиним только проверяемые претензии — «часть текста не совпала дословно»
#      это шум, по нему ремонтировать нечего;
#   2. после ремонта перепроверяем ТЕМИ ЖЕ проверками и оставляем новый текст,
#      только если стало лучше, иначе откат;
#   3. статус после ремонта — review, не confirmed: заверяет человек.
# Причины back-check, за которыми стоит ОБЪЕКТИВНАЯ находка: числа, единицы,
# отрицание, подмена стороны и «обратный перевод про другое». По ним ремонт
# берёт претензию из строки причины, и по ним же видно, что «жёсткая» отметка
# в старой записи стоит за настоящей находкой, а не за сравнением основ.
# «Потерян термин» сюда НЕ входит: эта претензия выставляется отдельной строкой
# из `terms_lost` — с учётом вердикта арбитра и вердикта судьи, чего сравнение
# подстрокой русской фразы не знает.
# Находки, которые СИЛЬНЕЕ заверения человеком. Класс намеренно уже, чем
# BACKCHECK_OBJECTIVE_REASONS: сюда входит только то, чья ПОСЫЛКА не зависит
# ни от чьей записи, ни от мнения модели, ни от языка и области — расхождение
# чисел, единиц, инверсия отрицания и подмена стороны. Считаются они
# детерминированно из пары «оригинал → перевод».
#
# Чего здесь НЕТ и почему:
#   • «обратный перевод про другое» — это сравнение по косинусу, на коротком
#     сегменте оно врёт («Посев и ТЛЧ» → «Culture and drug susceptibility
#     testing» помечено как «про другое» при верном переводе);
#   • НАРУШЕННЫЙ ПРИКАЗНЫЙ ТЕРМИН — вычисление детерминированное, но посылка
#     («запись верна») недоказана: 20 приказных записей пришли импортом без
#     следа человека, а «Бактериовыделение → bacillary excretion» заверена
#     человеком и при этом калька. Дай глоссарию право снимать заверение —
#     и редактор, исправивший кальку руками, получит бесконечную борьбу:
#     он правит, ремонт вписывает запись обратно, статус слетает, и так по
#     кругу. Единственная дверь (правка самой записи) с экрана сегмента
#     не видна. Такой спор уходит в human.termcheckDisputes, где и стоит
#     вопрос «неверна запись либо неверна проверка»;
#   • регистр, чужое письмо и самоповтор — сигналы для НОМИНАЦИИ: у них есть
#     законные исключения (торговая марка, название закона на чужом письме;
#     повтор наименований сторон в договоре), и снимать по ним решение
#     человека нельзя.
CONFIRM_OVERRIDE_REASONS = ("расхождение чисел", "расхождение единиц",
                            "инверсия отрицания", "подмена на противоположное")


def _confirm_override(seg: dict) -> list:
    """Доказательства, которые сильнее заверения человеком.

    Пусто — заверение неприкосновенно, как и было. Непусто — ремонт вправе
    переписать текст без отдельного разрешения, но ОБЯЗАН оставить след:
    что нашли, каким правилом, какой текст был. Молча снятое заверение —
    это ровно та потеря доверия, ради которой правило и заведено узким."""
    bc = seg.get("backcheck") or {}
    if not bc or _check_stale(bc, seg.get("target") or ""):
        return []
    return [r for r in (bc.get("reasons") or [])
            if any(r.startswith(k) for k in CONFIRM_OVERRIDE_REASONS)]


BACKCHECK_OBJECTIVE_REASONS = ("расхождение чисел", "расхождение единиц",
                               "инверсия отрицания", "подмена на противоположное",
                               "обратный перевод про другое")


def _gloss_misses(seg: dict, project: Optional[dict]) -> list:
    """Расхождения с ПРОВЕРЕННЫМИ записями глоссария: термин есть в оригинале,
    утверждённого варианта в переводе нет.

    Считается ровно так же, как в /glossary-impact, и по тем же `verified`:
    если ремонт и отчёт о соответствии разойдутся в том, что считать
    нарушением, они начнут переписывать сегменты друг за другом по кругу.
    Устаревать тут нечему — сверяется текущий текст с текущим глоссарием."""
    if project is None:
        return []
    target = (seg.get("target") or "").strip()
    if not target:
        return []
    out = []
    for h in _verified_hits(seg.get("source", ""), project):
        if _tgt_has_term(target, h["tgt"]):
            continue
        out.append({"kind": "gloss", "use": h["tgt"], "src": h["src"],
                    "text": "утверждённый перевод термина «" + h["src"] + "» — «"
                            + h["tgt"] + "», в переводе его нет"})
    return out


# ── Регистр букв ─────────────────────────────────────────────────────
# Проверка детерминированная и БЕСПЛАТНАЯ — заведена там же и по той же
# причине, что `_gloss_misses`: испортить регистр умеет не модель сама
# по себе, а запись глоссария. «use these exact translations» модель
# понимает буквально и копирует НАЧЕРТАНИЕ записи, поэтому «Туберкулема →
# Tuberculoma» ставит заглавную посреди фразы, «ТУБЕРКУЛЕЗ ОРГАНОВ ДЫХАНИЯ
# → RESPIRATORY TUBERCULOSIS» — капс посреди фразы, а «Фиброзно-кавернозный
# туберкулёз → fibrocavitary tuberculosis» роняет строчную в начало
# заголовка. Промпт перевода это правило теперь называет (см.
# `_translate_system`), но промпт лечит только НОВЫЕ переводы: у 2711
# сегментов боевого учебника текст уже написан, и без находки его никто
# не пересмотрит.
#
# Правил ровно три, и все три проверены на боевом проекте: 38 находок
# на 2711 сегментах, ни одной ложной. Четвёртого («в оригинале слово
# со строчной, в переводе с заглавной») здесь нет намеренно: английский
# законно поднимает регистр там, где русский его не поднимает — названия
# народов, месяцев, дней недели, — и такое правило кричало бы на верный
# перевод.
CASE_CAPS_MIN = 8          # букв, с которых надпись считается КАПС-заголовком,
                           # а не аббревиатурой: «МБТ» и «ВИЧ» это не заголовки
_CASE_WORD = re.compile(r"[^\W\d_]+", re.UNICODE)


def _has_case(text: str) -> bool:
    """Различает ли письмо этого текста заглавные и строчные.

    Иероглифы, арабица, иврит регистра не имеют, и спрашивать с них
    заглавную букву — выдуманная находка. Тот же закон, что у DOMAIN_RULES:
    нечем проверить — молчим."""
    return text.upper() != text.lower()


def _first_alpha(text: str) -> str:
    """Первая БУКВА текста. Не первый символ: подпись начинается с номера
    («139-Рис.»), пункт списка — с цифры и скобки, и по первому символу
    регистр не виден вовсе."""
    for ch in text:
        if ch.isalpha():
            return ch
    return ""


def _all_caps(text: str) -> bool:
    """Все буквы текста — заглавные (цифры и знаки не мешают)."""
    letters = [c for c in text if c.isalpha()]
    return bool(letters) and all(c.isupper() for c in letters)


def _caps_runs(text: str) -> list:
    """Цепочки ПОДРЯД идущих слов целиком заглавными, длиннее одного слова.

    Одиночное слово капсом сюда не попадает намеренно: это аббревиатура
    («ELISA», «RSNPMCFP»), законная в любом месте фразы. А вот два слова
    подряд — уже кусок заголовка, приехавший из записи глоссария.

    Однобуквенное слово цепочку РАЗРЫВАЕТ: в «A16 — RESPIRATORY TUBERCULOSIS»
    буква из шифра болезни попадала в цепочку, и человеку показывалась
    «A RESPIRATORY TUBERCULOSIS» — строка, которой в переводе нет. Назвать
    находку тем, чего в тексте не написано, значит послать ремонт искать
    несуществующее."""
    out, run = [], []
    for w in _CASE_WORD.findall(text):
        if len(w) > 1 and _all_caps(w) and _has_case(w):
            run.append(w)
        else:
            if len(run) > 1:
                out.append(run)
            run = []
    if len(run) > 1:
        out.append(run)
    return out


def _case_misses(seg: dict) -> list:
    """Расхождения по регистру между оригиналом и переводом.

    Ничего не переписывает и не стоит ни одного вызова модели — как
    `_gloss_misses`, устаревать тут нечему: сверяются нынешние тексты."""
    src = (seg.get("source") or "").strip()
    tgt = (seg.get("target") or "").strip()
    if not src or not tgt or not _has_case(src) or not _has_case(tgt):
        return []
    out = []
    a, b = _first_alpha(src), _first_alpha(tgt)
    if a and b and a.isupper() and b.islower():
        out.append({"kind": "case",
                    "text": "перевод начинается со строчной буквы («" + tgt[:40]
                            + "…»), а оригинал — с заглавной"})
    src_caps = _all_caps(src) and len([c for c in src if c.isalpha()]) >= CASE_CAPS_MIN
    if src_caps and not _all_caps(tgt):
        out.append({"kind": "case",
                    "text": "оригинал набран ЗАГЛАВНЫМИ целиком, перевод — нет"})
    # Обратная сторона той же поломки: капс, которого в оригинале не было.
    # Спрашиваем только у сегментов, где оригинал капсом НЕ набран и своих
    # капс-цепочек не имеет, — иначе это законный перенос заголовка.
    if not src_caps and not _caps_runs(src):
        for run in _caps_runs(tgt):
            if (sum(len(w) for w in run) >= CASE_CAPS_MIN
                    and max(len(w) for w in run) >= 5):
                out.append({"kind": "case",
                            "text": "«" + " ".join(run) + "» набрано ЗАГЛАВНЫМИ, "
                                    "хотя в оригинале это обычный текст"})
    return out


def _case_fit(src: str, text: str) -> str:
    """Привести НАЧЕРТАНИЕ текста к оригиналу. Без вызова модели.

    Нужна ремонту. Модель приносит верный термин в неверном регистре
    («ИНСТРУМЕНТАЛЬНЫЕ ИССЛЕДОВАНИЯ» → «Diagnostic investigations»), и правка
    откатывается расхождением по регистру — на боевом проекте так сгорели
    18 заходов. Промпт про регистр просит дважды, и всё равно; а поправить
    это можно бесплатно и однозначно, потому что сочинять нечего: слова
    и их порядок те же, меняются только заглавные и строчные.

    Правила ровно два, и оба ПОДНИМАЮТ регистр — направление, безопасное
    всегда (см. `_case_like`):
      1) оригинал набран ЗАГЛАВНЫМИ целиком (от CASE_CAPS_MIN букв) —
         переводу тоже быть заглавными;
      2) оригинал начинается с заглавной — и перевод начинается с заглавной.
    Опускать регистр здесь нельзя: «в оригинале строчная» законно уживается
    с заглавной в переводе (народы, месяцы, названия), и правило кричало бы
    на верный текст. Письмо без регистра — молчим, тот же закон, что
    у DOMAIN_RULES."""
    src, text = (src or "").strip(), (text or "").strip()
    if not src or not text or not _has_case(src) or not _has_case(text):
        return text
    if (_all_caps(src) and len([c for c in src if c.isalpha()]) >= CASE_CAPS_MIN
            and not _all_caps(text)):
        return text.upper()
    a, b = _first_alpha(src), _first_alpha(text)
    if a and b and a.isupper() and b.islower():
        i = text.find(b)
        if i >= 0:
            return text[:i] + b.upper() + text[i + 1:]
    return text


# ── Регистр глоссарного термина: 1в1 по оригиналу ────────────────────# ── Регистр глоссарного термина: 1в1 по оригиналу ────────────────────
# Запись глоссария хранит ОДНО начертание, а мест, куда она встаёт, много:
# «Туберкулема» стоит и заголовком, и посреди фразы. Правильный ответ даёт
# не запись, а САМ ОРИГИНАЛ: как термин написан в этом сегменте, так он
# и должен быть написан в переводе. Форма из текста у нас есть — `_form`
# из `_get_context` (её же возвращает `_term_match`), поэтому подгонку
# делаем детерминированно, а не просьбой к модели.
#
# Три направления и у каждого своя оговорка:
#   ОРИГИНАЛ КАПСОМ  → перевод капсом, но только от CASE_CAPS_MIN букв:
#      «RW» — аббревиатура, и «WASSERMANN REACTION» из неё было бы криком;
#   Оригинал с заглавной → заглавная (поднять регистр безопасно всегда);
#   оригинал со строчной → строчная, КРОМЕ имён собственных и аббревиатур.
# Последнее и есть единственное опасное направление: «реакция Манту» →
# «Mantoux reaction», «очаг Гона» → «Ghon focus», «вакцина БЦЖ» → «BCG
# vaccine» — опустив первую букву, мы испортили бы имя. Признак берётся
# из САМОЙ записи: заглавная НЕ у первого слова оригинала означает имя
# внутри термина, а две заглавные в слове перевода — аббревиатуру.


def _name_bearing(term: str) -> bool:
    """В термине есть имя собственное: заглавная не у первого слова.

    «реакция Манту», «очаг Гона», «вакцина БЦЖ» — опускать регистр у таких
    переводов нельзя, там имя или аббревиатура."""
    words = [w for w in (term or "").split() if any(c.isalpha() for c in w)]
    for w in words[1:]:
        first = next((c for c in w if c.isalpha()), "")
        if first and first.isupper():
            return True
    return False


CASE_ACRONYM_MAX = 5     # букв, до которых слово капсом — аббревиатура, а не крик


def _acronymish(word: str) -> bool:
    """Слово, у которого заглавные — часть написания: BCG, XDR-TB, «M.», pH.

    Длинное слово капсом сюда НЕ попадает: «RESPIRATORY» — не аббревиатура,
    а крик, приехавший из записи глоссария, и опустить его как раз надо."""
    letters = [c for c in word if c.isalpha()]
    if not letters:
        return False
    if word.endswith(".") and len(word) <= 2:
        return True
    if all(c.isupper() for c in letters):
        return len(letters) <= CASE_ACRONYM_MAX
    # Заглавная НЕ в начале — McDonald, pH, DNase: регистр внутри значащий.
    return any(c.isupper() for c in letters[1:])


def _shout_keep(word: str) -> bool:
    """Оставить ли слово капсом при снятии КРИКА с записи.

    Внутри кричащей записи аббревиатуру от обычного слова отличить нечем —
    заглавные у всех, и правило «короткое слово капсом = аббревиатура»
    превращало «GENERALIZED FORMS OF TUBERCULOSIS» в «generalized FORMS OF
    tuberculosis». Поэтому признак только СТРУКТУРНЫЙ: цифра или дефис
    («MDR-TB», «2HRE») — у обычного слова их не бывает."""
    return any(c.isdigit() or c == "-" for c in word)


def _case_like(form: str, target: str, term: str = "") -> str:
    """Перевод термина, подогнанный под начертание НАЙДЕННОЙ формы оригинала.

    Сначала с записи снимается её собственный КРИК, если оригинал здесь так
    не написан: «ТУБЕРКУЛЕЗ ОРГАНОВ ДЫХАНИЯ → RESPIRATORY TUBERCULOSIS» посреди
    фразы должно стать «respiratory tuberculosis», а не «rESPIRATORY
    TUBERCULOSIS». Аббревиатуры внутри при этом не трогаются."""
    if not target or not form:
        return target
    letters = [c for c in form if c.isalpha()]
    if not letters or not _has_case(form) or not _has_case(target):
        return target
    all_caps = all(c.isupper() for c in letters)
    if all_caps and len(letters) < CASE_CAPS_MIN:
        # Оригинал — АББРЕВИАТУРА. Её заглавные это часть написания, а не
        # признак места в предложении: «ПТП» стоит капсом и посреди фразы,
        # но расшифровка «anti-tuberculosis drugs» — обычное словосочетание,
        # и поднимать ей регистр значило бы портить перевод в 17 сегментах.
        # Начертание неизвестно — не трогаем ничего.
        return target
    src_caps = all_caps and len(letters) >= CASE_CAPS_MIN
    base = target
    t_letters = [c for c in target if c.isalpha()]
    if (not src_caps and len(t_letters) >= CASE_CAPS_MIN
            and all(c.isupper() for c in t_letters)):
        base = _re.sub(r"\S+",
                       lambda m: m.group(0) if _shout_keep(m.group(0))
                       else m.group(0).lower(), target)
    i = next((k for k, c in enumerate(base) if c.isalpha()), None)
    if i is None:
        return base
    if src_caps:
        return base.upper()
    if letters[0].isupper():
        return base[:i] + base[i].upper() + base[i + 1:]
    head = base[i:].split()[0] if base[i:].split() else ""
    if _name_bearing(term) or _acronymish(head):
        return base
    return base[:i] + base[i].lower() + base[i + 1:]


def _case_class(form: str) -> str:
    """Класс начертания найденной формы: «caps», «abbr», «upper», «lower», «"»."""
    letters = [c for c in form if c.isalpha()]
    if not letters or not _has_case(form):
        return ""
    if all(c.isupper() for c in letters):
        return "caps" if len(letters) >= CASE_CAPS_MIN else "abbr"
    return "upper" if letters[0].isupper() else "lower"


def _agreed_form(term: str, text: str, lang: str = "") -> str:
    """Форма термина в оригинале — если ВСЕ его вхождения написаны одинаково.

    Иначе "": в одном сегменте термин может стоять и заголовком, и посреди
    фразы («Туберкулема плотна. …части туберкулемы…»), а требование у нас
    одно на сегмент. Навязать одно начертание всем вхождениям значит испортить
    то из них, которое и так верно. Не знаем — молчим, как везде."""
    got = _term_forms(term, text, lang)
    if not got:
        return ""
    classes = {_case_class(g) for g in got}
    return got[0] if len(classes) == 1 else ""


def _sentence_start(text: str, pos: int) -> bool:
    """Стоит ли позиция в начале предложения: перед ней либо ничего,
    либо знак конца. Заглавная там законна при любом оригинале."""
    before = text[:pos].rstrip()
    return not before or before[-1] in ".!?:;•–—-”)"


def _term_case_misses(seg: dict, project: Optional[dict]) -> list:
    """Приказные термины, чьё начертание в переводе разошлось с оригиналом.

    Считается по тем же `_verified_hits`, что `_gloss_misses`: разойдись
    определение — ремонт и отчёт правили бы сегмент по кругу."""
    if project is None:
        return []
    return _term_case_hits(seg, _verified_hits(seg.get("source", ""), project))


def _term_case_spans(seg: dict, hits: list) -> list:
    """[(начало, конец, как надо, как есть, термин)] — места, где начертание
    приказного термина разошлось с оригиналом.

    Один расчёт на проверку и на правку: разойдись они — отчёт показывал бы
    одно, а команда делала другое.

    Смотрим только БУКВАЛЬНЫЕ вхождения перевода термина ЦЕЛЫМ СЛОВОМ
    (сосед — не буква; дефис и апостроф границей считаются, как в
    `_bound_l`/`_bound_r`): нет вхождения вовсе — это забота `_gloss_misses`,
    и говорить об одном и том же дважды нельзя. Цена: форма с приписанным
    окончанием («Tuberculomas» у записи «tuberculoma») на начертание
    не проверяется — молчим, как везде, где не знаем. Термин ВНУТРИ другого термина своего начертания не диктует:
    «ревакцинация → revaccination» сидит внутри «Ревакцинация БЦЖ → BCG
    revaccination», и без этого правила две записи правили одно место по
    очереди, каждая ломая работу другой."""
    target = (seg.get("target") or "").strip()
    if not target or not _has_case(target):
        return []
    low = target.lower()
    claims = []
    for h in hits:
        tgt = (h.get("tgt") or "").strip()
        # Форма берётся СОГЛАСОВАННАЯ по всему оригиналу, а не первая
        # попавшаяся: требование у нас одно на сегмент, и разнобой в оригинале
        # означает «не знаем».
        form = _agreed_form(h.get("src") or "", seg.get("source") or "",
                            _src_lang(h))
        if not tgt or not form:
            continue
        if _case_class(form) == "abbr":
            if _case_class(seg.get("source") or "") != "caps":
                # Оригинал — АББРЕВИАТУРА (капс короче CASE_CAPS_MIN): начертание
                # неизвестно, и `_case_like` отдаёт написание ЗАПИСИ. Требовать
                # его от текста нельзя: «БОЛЬНЫХ → patient» превращало
                # «TB PATIENTS» в «TB patientS» (боевой #22), а ремонт, вернувший
                # капс, откатывался этим же счётчиком.
                continue
            # Короткий капс ВНУТРИ капс-заголовка — слово заголовка, а не
            # аббревиатура: «НАБЛЮДЕНИЕ БОЛЬНЫХ ТБ» → «MANAGEMENT OF TB PATIENTS».
            # Тот же закон, что у `_case_misses` на уровне сегмента.
            want = tgt.upper()
        else:
            want = _case_like(form, tgt, h.get("src") or "")
        need = tgt.lower()
        k = low.find(need)
        while k >= 0:
            e = k + len(tgt)
            # Только ЦЕЛОЕ слово: «patient» внутри «PATIENTS» — не то место,
            # и перекраска куска слова даёт «patientS». Нет целого вхождения —
            # молчим, как везде, где не знаем.
            whole = ((k == 0 or not low[k - 1].isalpha())
                     and (e >= len(low) or not low[e].isalpha()))
            if whole:
                claims.append((k, e, want, target[k:e], h.get("src") or ""))
            k = low.find(need, k + 1)
    # Длинный термин главнее: слева направо, длинный первым, а всё, что попало
    # внутрь уже занятого места, пропускаем.
    claims.sort(key=lambda x: (x[0], -(x[1] - x[0])))
    taken, out = [], []
    for c in claims:
        if any(c[0] >= t[0] and c[1] <= t[1] for t in taken):
            continue
        taken.append(c)
        # Заглавная в начале предложения законна при любом оригинале.
        if c[3] == c[2] or (c[2][:1].islower() and _sentence_start(target, c[0])
                            and c[3][:1].upper() + c[3][1:] == c[2][:1].upper() + c[2][1:]):
            continue
        out.append(c)
    return out


def _term_case_hits(seg: dict, hits: list) -> list:
    """То же по УЖЕ найденным приказным записям: `glossary_impact` ходит по
    всему проекту и второй проход по глоссарию себе позволить не может."""
    seen, out = set(), []
    for _a, _b, want, _got, src in _term_case_spans(seg, hits):
        if (src, want) in seen:
            continue
        seen.add((src, want))
        out.append({"kind": "term_case", "use": want, "src": src,
                    "text": "в оригинале термин «" + src + "» написан так, что "
                            "в переводе он пишется «" + want + "»"})
    return out


def _term_case_fix(seg: dict, project: Optional[dict]) -> tuple:
    """(новый текст, [(было, стало)]) — начертание приказных терминов по оригиналу.

    Меняются ТОЛЬКО заглавные и строчные: буквы, слова и порядок остаются те же,
    поэтому вызов модели здесь не нужен и вреден. Ремонт на такую претензию
    переписывает предложение целиком, оценка находит, что стало чуть хуже,
    и откатывает правку — на боевом учебнике так застряли 5 сегментов из 77,
    каждый ценой платного вызова. Это вторая команда в системе, меняющая текст
    без модели (первая — `/glossary/revert-repairs`), и по той же причине:
    она ничего не сочиняет."""
    target = (seg.get("target") or "").strip()
    if not target or project is None:
        return target, []
    spans = _term_case_spans(seg, _verified_hits(seg.get("source", ""), project))
    if not spans:
        return target, []
    out, moves, pos = [], [], 0
    for a, b, want, got, _src in spans:
        if a < pos:
            continue
        out.append(target[pos:a])
        out.append(want)
        moves.append((got, want))
        pos = b
    out.append(target[pos:])
    return "".join(out), moves


# ── Буквы чужого письма ──────────────────────────────────────────────
# Третья бесплатная детерминированная претензия, рядом с `_gloss_misses`
# и `_case_misses` и по той же причине. Промпт требует «Output must be 100%
# EN» с самого начала, но требование это про ТЕКСТ, а рвётся оно на кусках,
# которые текстом не выглядят: на боевом учебнике в английском переводе
# остались «(МБТ+)», «Р02»/«РС02», «38°С», «2HSЕ/10HE», «2-х», «3х» —
# кириллица внутри формул, единиц и аббревиатур. Найти это глазами нельзя:
# «РС02» и «PCO2» на экране неотличимы, а в поиске по документу — разные
# строки.
#
# Письмо считается ПО ТЕКСТАМ, а не по таблице «RU→EN»: хардкодить пару
# языков в новом коде нельзя (см. DOMAIN_RULES). Совпали письменности
# оригинала и перевода (RU→UK, EN→DE) — молчим: там буквы общие и
# претензия была бы выдумана.
_SCRIPT_CACHE: dict = {}


def _script_of(ch: str) -> str:
    """Письменность одной БУКВЫ («CYRILLIC», «LATIN», «HAN»…) или "".

    По имени символа в Юникоде, без внешних зависимостей: первое слово имени
    и есть письменность у всех алфавитов, которые нас касаются."""
    got = _SCRIPT_CACHE.get(ch)
    if got is None:
        got = unicodedata.name(ch, "").split(" ")[0] if ch.isalpha() else ""
        if got == "CJK":
            got = "HAN"
        _SCRIPT_CACHE[ch] = got
    return got


def _dominant_script(text: str) -> str:
    """Письменность, которой набрано БОЛЬШИНСТВО букв текста.

    Именно большинство, а не «любая встреченная»: латиница законно попадает
    в русский текст («Mycobacterium bovis»), и по факту её присутствия язык
    определять нельзя."""
    counts: dict = {}
    for ch in text:
        sc = _script_of(ch)
        if sc:
            counts[sc] = counts.get(sc, 0) + 1
    if not counts:
        return ""
    return max(counts.items(), key=lambda kv: (kv[1], kv[0]))[0]


def _script_misses(seg: dict) -> list:
    """Буквы письменности ОРИГИНАЛА, оставшиеся в переводе."""
    src = (seg.get("source") or "").strip()
    tgt = (seg.get("target") or "").strip()
    if not src or not tgt:
        return []
    s_script = _dominant_script(src)
    t_script = _dominant_script(tgt)
    if not s_script or not t_script or s_script == t_script:
        return []
    bad, seen = [], set()
    for word in tgt.split():
        if any(_script_of(c) == s_script for c in word) and word not in seen:
            seen.add(word)
            bad.append(word.strip("()[].,;:«»\"'"))
    if not bad:
        return []
    # Куски называются поимённо: «остались буквы оригинала» без списка —
    # претензия, которую нечем исполнить, а искать глазами кириллическую «С»
    # среди латинских бессмысленно.
    return [{"kind": "script",
             "text": "в переводе остались буквы исходного письма ("
                     + s_script.lower() + "): "
                     + ", ".join("«" + w + "»" for w in bad[:6])
                     + (" и ещё " + str(len(bad) - 6) if len(bad) > 6 else "")}]


# ── Контекстный арбитр спорного термина ──────────────────────────────
# Проверки смотрят на сегмент в одиночку, а термин живёт в ряду: «туберкулёз
# лёгких» обратный перевод возвращает как «лёгочный туберкулёз», и по словам
# это потеря, а по смыслу — то же самое. Ответить на такое можно только увидев,
# в каком месте документа сегмент стоит. Отсюда единственный вызов, которому
# дают соседей: сегмент ДО, этот, ПОСЛЕ — и вопрос «правильно ли передан
# термин ЗДЕСЬ, а если нет, то как правильно».
#
# Вердикт кэшируется НА СЕГМЕНТЕ по хешу перевода и версии вопросов — тем же
# способом, что back-check, termcheck и сверка смысла глоссария. Меняешь
# промпт — поднимай версию, иначе новый вопрос не задаётся никогда.
TERM_CONTEXT_VERSION = 2
TERM_CONTEXT_DEFAULT_MODEL = JUDGE_DEFAULT_MODEL


def _term_terms_of(seg: dict, project: Optional[dict],
                   disputes_only: bool = True) -> list:
    """Спорные ПРИКАЗНЫЕ термины сегмента: где проверка и глоссарий разошлись.

    Два разных спора и оба про одно и то же слово:
      • термин не пережил обратный перевод (`backcheck.terms_lost`);
      • termcheck забраковал слово, которое И ЕСТЬ утверждённый перевод.
    Считается только по verified: спорить с подсказкой автоимпорта не о чем —
    модели разрешено её игнорировать."""
    if project is None:
        return []
    target = (seg.get("target") or "").strip()
    if not target:
        return []
    hits = _verified_hits(seg.get("source", ""), project)
    if not hits:
        return []
    by_src = {_norm_key(h.get("src")): h for h in hits}
    by_tgt = {_norm_key(h.get("tgt")): h for h in hits}
    out, seen = [], set()

    def add(h, why, form=None):
        k = _norm_key(h.get("src"))
        if k in seen:
            # Форму дописываем и к уже заведённому спору: одна запись глоссария
            # может встретиться в сегменте несколькими формами, а вопрос к ней
            # один. По формам потом сверяется снятие претензии в ремонте.
            for d in out:
                if _norm_key(d["src"]) == k and form and form not in d["forms"]:
                    d["forms"].append(form)
            return
        seen.add(k)
        out.append({"src": h.get("src"), "tgt": h.get("tgt"), "why": why,
                    "forms": [form] if form else []})

    bc = seg.get("backcheck") or {}
    if bc.get("target_hash") == _text_hash(target):
        for t in (bc.get("terms_lost") or []):
            h = by_src.get(_norm_key(t))
            if h is None:
                # terms_lost хранит НАЙДЕННУЮ форму («туберкулёза лёгких»),
                # а запись глоссария — словарную. Ищем по вхождению формы
                # в исходник записи и обратно.
                h = next((x for x in hits
                          if _norm_key(t).startswith(_norm_key(x.get("src"))[:6])), None)
            if h:
                add(h, "не пережил обратный перевод: в обратном тексте его нет", t)
    tc = seg.get("termcheck") or {}
    if tc.get("target_hash") == _text_hash(target):
        for f in (tc.get("findings") or []):
            if f.get("severity") not in TERMCHECK_DISPUTING:
                continue
            h = by_tgt.get(_norm_key(f.get("tgt_term")))
            if h:
                add(h, "проверка терминов забраковала его и предлагает «%s»"
                    % (f.get("suggestion") or "другой вариант"))
    if not disputes_only:
        # Режим сверки: спрашиваем про ВСЕ приказные термины сегмента, а не
        # только про те, вокруг которых уже вышел спор. Спор находят
        # детерминированные проверки, а они знают морфологию одного языка
        # и знают её грубо: «туберкулёз лёгких» им возвращается как «лёгочный
        # туберкулёз» и считается потерей. Модель отвечает на тот же вопрос
        # без всякой морфологии и на любом языке — ради этого шаг и заведён.
        for h in hits:
            add(h, "", h.get("_form"))
    return out


def _term_disputes_of(seg: dict, project: Optional[dict]) -> list:
    """Только СПОРНЫЕ приказные термины: где детерминированная проверка
    разошлась с глоссарием. Отдельным именем, потому что вопрос другой:
    здесь платят за разбор конфликта, а в сверке — за проверку всего."""
    return _term_terms_of(seg, project, disputes_only=True)


def _term_context_stale(seg: dict) -> bool:
    tcx = seg.get("termContext") or {}
    return (tcx.get("target_hash") != _text_hash((seg.get("target") or "").strip())
            or tcx.get("version") != TERM_CONTEXT_VERSION)


def _stale_words_of(project: dict) -> dict:
    """{id сегмента: [слова]} — забракованные termcheck слова, всё ещё стоящие
    в тексте. Карточка очереди помнит `wasTgt` — формулировку, которую проверка
    отвергла, — а сегмент об этом не помнит: свежий termcheck мог передумать
    между прогонами, и дефект остаётся в готовом на вид тексте. Сверка
    бесплатна — подстрока по границам слов. ОДИН расчёт на три места
    (/analysis, _plan_step, сам шаг сверки): разойдись они, смета обещала бы
    одно, а работа делала другое.

    Слово, которое САМО есть утверждённый перевод, отсеивается ЗДЕСЬ, а не
    в шаге: это спор с ЗАПИСЬЮ (human.termcheckDisputes), второго голоса
    у него не будет никогда. Отсеянный только у шага, такой сегмент вечно
    числился бы «спросит арбитра», оплачивал вызов каждым прогоном — а слово
    молча выпадало бы из вопроса и из охвата: платная карусель."""
    by_id = {sg["id"]: sg for sg in project["segments"]}
    out: dict = {}
    for c in (STATE.get("termQueue") or ()):
        was = (c.get("wasTgt") or "").strip()
        sid = c.get("segment")
        if not was or c.get("project") != project["id"] or sid not in by_id:
            continue
        rx = _word_re(was)
        if rx and rx.search(by_id[sid].get("target") or ""):
            words = out.setdefault(sid, [])
            if was not in words:
                words.append(was)
    for sid in list(out):
        hits_tgt = {_norm_key(h.get("tgt") or "") for h in
                    _verified_hits(by_id[sid].get("source", ""), project)}
        out[sid] = [w for w in out[sid] if _norm_key(w) not in hits_tgt]
        if not out[sid]:
            del out[sid]
    return out


def _stale_unasked(seg: dict, words: list) -> list:
    """Забракованные слова, про которые арбитра ещё не спрашивали СВЕЖИМ
    вердиктом. Охват лежит на вердикте (`termContext.staleAsked`) — тот же
    приём, что `all_terms`: без отметки вердикт без этих вопросов закрывал бы
    сегмент от них навсегда. Записываются только ОТВЕЧЕННЫЕ слова: пропущенное
    моделью слово обязано спроситься снова."""
    if not words:
        return []
    if _term_context_stale(seg):
        return list(words)
    asked = {_norm_key(w) for w in
             ((seg.get("termContext") or {}).get("staleAsked") or ())}
    return [w for w in words if _norm_key(w) not in asked]


# Перевод строки для промптов, которые собираются склейкой. Отдельной
# константой — и она ОБЯЗАНА быть объявлена: без неё `_openai_term_context`
# падал с NameError на построении тела запроса, то есть ДО вызова модели.
# Наружу это выглядело как «арбитр не ответил», и шаг сверки терминов
# отчитывался отказом по всем 698 сегментам, ни разу не сходив в модель.
# Ловится только тестом, который гоняет НАСТОЯЩИЙ сборщик промпта:
# остальные тесты подменяют функцию целиком (см. tests/test_term_context.py).
NL = "\n"


def _openai_term_context(seg: dict, project: dict, disputes: list,
                         prev_src: str, next_src: str, model: Optional[str],
                         stale: list = ()) -> Optional[dict]:
    """Один вызов на сегмент: соседи + спорные термины → вердикт по каждому.

    `stale` — забракованные termcheck слова, всё ещё стоящие в переводе:
    их секция добавляется в промпт ТОЛЬКО когда такие слова есть, поэтому
    для остальных сегментов промпт байт в байт прежний и TERM_CONTEXT_VERSION
    не поднимается — подъём перекупил бы сотни готовых вердиктов ради
    вопросов, которых им никто не задаёт."""
    import openai
    mdl = _resolve_model(model or TERM_CONTEXT_DEFAULT_MODEL)
    dom = _resolve_domain(project.get("domain"))
    src_lang, tgt_lang = project.get("src", "RU"), project.get("tgt", "EN")
    system = (
        "Ты — редактор перевода, специализация: " + dom["label"].lower() + ". "
        "Тебе дают три подряд идущих сегмента документа (язык: " + src_lang + "), "
        "перевод СРЕДНЕГО из них на " + tgt_lang + " и список утверждённых терминов, "
        "применимых к этому сегменту. У некоторых в скобках сказано, чем "
        "недовольна автоматическая проверка; у остальных скобок нет — их надо "
        "просто сверить.\n\n"
        "Для КАЖДОГО термина ответь, передан ли он в переводе среднего сегмента "
        "правильно ИМЕННО ЗДЕСЬ, с учётом соседних сегментов.\n"
        "  ok — true, если смысл термина в переводе передан верно (пусть другими "
        "словами, другой частью речи или другим порядком слов);\n"
        "  ok — false, если термин передан неверно, потерян или подменён;\n"
        "  use — когда ok=false, ПРАВИЛЬНЫЙ вариант на " + tgt_lang + ", несколько слов;\n"
        "  why — причина на " + _explain_lang_name() + ", несколько слов "
        "(её читает человек, поэтому язык — его, а не документа).\n\n"
        "Соседние сегменты даны только как обстановка. Не оценивай их перевод.\n"
        'Верни ТОЛЬКО JSON: {"terms":[{"src":"...","ok":true,"use":"","why":""}]}. '
        "Без пояснений."
    )
    if stale:
        # Второй голос по забракованному слову: одно мнение termcheck — не
        # приказ переписывать (тот же закон, что у разнобоя), а проверка
        # с тех пор могла и передумать. Арбитр видит сегмент в ряду соседей
        # и решает: слово годно (претензия снимается) либо негодно (замена
        # уходит ремонту). Ответ идёт тем же форматом, слово кладётся в src.
        # ПРАВИШЬ ФОРМУЛИРОВКУ ЭТОЙ СЕКЦИИ — поднимай TERM_CONTEXT_VERSION:
        # охват staleAsked своей версии не несёт, и без подъёма уже отвеченные
        # слова навсегда останутся с вердиктами по старому вопросу.
        system += (
            "\n\nОтдельным списком даны СЛОВА ИЗ ПЕРЕВОДА (" + tgt_lang + "), "
            "которые автоматическая проверка терминологии раньше браковала. "
            "Для каждого ответь тем же форматом JSON, положив само слово в поле src:\n"
            "  ok — true, если слово здесь уместно и браковать его не за что;\n"
            "  ok — false, если слово в этом контексте действительно негодно; "
            "тогда в use дай верный вариант на " + tgt_lang + ".\n"
        )
    body = (
        "[сегмент ДО] " + (prev_src or "—") + NL +
        ">>> [этот сегмент] " + (seg.get("source") or "") + NL +
        "[сегмент ПОСЛЕ] " + (next_src or "—") + NL + NL +
        "Перевод этого сегмента (" + tgt_lang + "): " + (seg.get("target") or "")
    )
    # Пустой заголовок «Утверждённые термины:» без единой строки под ним
    # толкал бы модель выдумывать записи (сегмент из одних забракованных
    # слов): секция появляется только с содержимым. Для сегментов с
    # терминами тело байт в байт прежнее.
    if disputes:
        body += (NL + NL + "Утверждённые термины:" + NL
                 + NL.join("  - %s → %s%s" % (d["src"], d["tgt"],
                                              (" (%s)" % d["why"]) if d.get("why") else "")
                           for d in disputes))
    if stale:
        body += (NL + NL + "Забракованные проверкой слова перевода:" + NL
                 + NL.join("  - " + w for w in stale))
    client = openai.OpenAI(api_key=os.environ.get("OPENAI_API_KEY"), timeout=90, max_retries=2)
    extra = ({"max_completion_tokens": 2048} if mdl["api"] == "modern"
             else {"max_tokens": 700, "temperature": 0.0})
    try:
        resp = client.chat.completions.create(
            model=mdl["id"],
            messages=[{"role": "system", "content": system},
                      {"role": "user", "content": body}],
            **extra)
        _note_usage("term_context", mdl["id"], resp)
        raw = (resp.choices[0].message.content or "").strip()
        # Модель иногда оборачивает JSON в ```json ... ``` — вырезаем тело,
        # ровно как у судьи и у сверки смысла.
        m = re.search(r"\{.*\}", raw, re.S)
        data = json.loads(m.group(0)) if m else None
    except Exception as e:
        print("[backend] контекстный арбитр seg#%s: %s" % (seg.get("id"), e), file=sys.stderr)
        return None
    if not isinstance(data, dict) or not isinstance(data.get("terms"), list):
        return None
    return {"terms": data["terms"], "model": mdl["id"]}


def _ctx_corpus_gate(item: dict, scope: tuple) -> None:
    """Бесплатная проверка совета арбитра корпусом целевого языка — ОДНО
    правило на оба вида вердиктов (термин глоссария и забракованное слово):
    правь его здесь, а не в двух копиях. Запрещать вправе только корпус,
    СВОЙ для этой области: ноль в Википедии для узкого медицинского термина
    доказывает лишь отсутствие статьи с таким названием — пока PubMed был
    недоступен с сервера, так молча срезали готовый совет у пяти вердиктов.
    Молчание корпуса («не знаю») не одобряет и не блокирует."""
    use = item.get("use") or ""
    if item.get("ok") is not False or not use:
        return
    c = _corpus_check(use, scope)
    if c is None:
        return
    item["corpus"] = {"hits": c.get("hits"), "source": c.get("source")}
    if not c.get("ok") and c.get("vetoAllowed", True):
        # Корпус НАЛАГАЕТ ВЕТО, но приказа не даёт: совета, которого нет
        # в целевом языке, в ремонте быть не должно.
        item["use"] = ""
        item["why"] = ((item["why"] + "; ") if item.get("why") else "") + \
            "вариант не найден в корпусе целевого языка"


def _run_segment_term_context(seg: dict, project: dict,
                              model: Optional[str] = None,
                              disputes_only: bool = True,
                              stale_words: Optional[list] = None) -> dict:
    """Спросить арбитра про спорные термины сегмента и записать вердикт.

    Ответ проверяется БЕСПЛАТНО, прежде чем стать поводом что-то переписывать:
    предложенный вариант спрашивается у корпуса целевого языка (`_corpus_check`).
    Нулевое число вхождений — это калька, и такой совет в ремонт не пойдёт;
    молчание корпуса («не знаю») не одобряет и не блокирует, как везде в этой
    системе. Дальше правку всё равно принимает ремонт со своей перепроверкой
    и откатом — арбитр только предлагает.

    `stale_words` — забракованные termcheck слова, всё ещё стоящие в тексте:
    арбитр даёт по ним ВТОРОЙ голос. Слова приходят из `_stale_words_of` —
    слово-приказный-перевод отсеяно уже там (спор с ЗАПИСЬЮ решает человек).
    Свежие вердикты по УЖЕ отвеченным словам сохраняются, а не затираются:
    точечный разбор спора (stale_words=None) прежде переписывал termContext
    целиком и выбрасывал оплаченные ответы вместе с находкой ремонта — та же
    работа покупалась следующим прогоном второй раз."""
    disputes = _term_terms_of(seg, project, disputes_only=disputes_only)
    # Прежние stale-вердикты живы, пока жив сам вердикт (тот же текст и та же
    # версия вопросов): переносим их и спрашиваем только неспрошенное.
    prev = seg.get("termContext") or {}
    keep_stale, keep_asked = [], []
    if prev and not _term_context_stale(seg):
        keep_stale = [t for t in (prev.get("terms") or []) if t.get("stale")]
        keep_asked = list(prev.get("staleAsked") or [])
    asked_norm = {_norm_key(w) for w in keep_asked}
    stale = [w for w in (stale_words or []) if _norm_key(w) not in asked_norm]
    if not disputes and not stale:
        return {"ok": False, "error": "Утверждённых терминов в сегменте нет"}
    prev_src, next_src = _neighbours(project, seg)
    res = _openai_term_context(seg, project, disputes, prev_src, next_src, model,
                               stale=stale)
    if not res:
        return {"ok": False, "error": "Арбитр не ответил"}
    by_src = {_norm_key(d["src"]): d for d in disputes}
    by_stale = {_norm_key(w): w for w in stale}
    scope = _project_scope(project)
    terms, stale_asked = list(keep_stale), list(keep_asked)
    for t in res["terms"]:
        d = by_src.get(_norm_key(t.get("src")))
        if d is None:
            # Ответ про забракованное слово: модель кладёт его в src.
            w = by_stale.get(_norm_key(t.get("src")))
            if w is None:
                continue
            ok = t.get("ok")
            use = (t.get("use") or "").strip()
            item = {"stale": True, "src": "", "tgt": w, "forms": [],
                    "ok": bool(ok) if ok is not None else None,
                    "use": use, "why": (t.get("why") or "").strip()}
            _ctx_corpus_gate(item, scope)
            if item["ok"] is not None:
                # Охват пишется только по ОТВЕЧЕННЫМ словам: пропущенное
                # моделью обязано спроситься снова (см. _stale_unasked).
                stale_asked.append(w)
            terms.append(item)
            continue
        ok = t.get("ok")
        use = (t.get("use") or "").strip()
        item = {"src": d["src"], "tgt": d["tgt"], "forms": d["forms"],
                "ok": bool(ok) if ok is not None else None,
                "use": use, "why": (t.get("why") or "").strip()}
        _ctx_corpus_gate(item, scope)
        terms.append(item)
    seg["termContext"] = {
        "version": TERM_CONTEXT_VERSION,
        # Охват вопроса. Разбор СПОРА спрашивает только про спорные термины,
        # сверка — про все. Без этой отметки вердикт разбора закрывал сегмент
        # для сверки навсегда: на боевом проекте так осталось бы не сверено
        # 89 приказных терминов в 61 сегменте.
        "all_terms": not disputes_only,
        # Охват по забракованным словам — тот же закон (см. _stale_unasked).
        "staleAsked": stale_asked,
        "target_hash": _text_hash((seg.get("target") or "").strip()),
        "model": res["model"], "terms": terms,
        "at": datetime.now().strftime("%Y-%m-%d %H:%M"),
    }
    return {"ok": True, "termContext": seg["termContext"]}


def _arbiter_settled(seg: dict) -> set:
    """Термины, про которые арбитр СВЕЖИМ вердиктом сказал «передан верно».

    Нужна балльщику. `_terms_lost_open` уже снимает по этому вердикту претензию
    ремонта, а `run_backcheck` считает штраф за потерянный термин заново и
    вердикта не спрашивает — то есть балл держится на претензии, которую
    оплаченный арбитр официально отменил. На боевом учебнике так вышло у 59
    сегментов: у #445 «Туберкулёз органов дыхания» балл 8 при верном переводе,
    а восьмёрка ещё и гасит судью, потому что лежит ниже низа его зоны.
    Половина корзины «оценка ниже порога» держалась именно на этом."""
    out = set()
    for c in _term_context_of(seg):
        if c.get("ok") is not True:
            continue
        out.add(_norm_key(c.get("src") or ""))
        for f in (c.get("forms") or ()):
            out.add(_norm_key(f))
    out.discard("")
    return out


def _hits_for_score(seg: dict, hits: list) -> list:
    """Требования глоссария, по которым СЧИТАЕТСЯ БАЛЛ. Из них выброшены те,
    что арбитр уже признал переданными верно: наказывать за потерю термина,
    про который сказано «на месте», значит спорить с собственным оплаченным
    вердиктом."""
    settled = _arbiter_settled(seg)
    if not settled:
        return hits
    return [h for h in hits if _norm_key(h.get("src") or "") not in settled]


def _term_context_of(seg: dict) -> list:
    """Свежие вердикты арбитра для этого текста. Устаревшие не читаем: они
    описывают перевод, которого больше нет."""
    return [] if _term_context_stale(seg) else ((seg.get("termContext") or {}).get("terms") or [])


def _terms_lost_open(seg: dict) -> list:
    """Приказные термины, чью потерю НИКТО ещё не снял.

    Одно определение на двоих: по нему ремонт заводит претензию
    (`_repair_findings`) и по нему же считается его собственный счётчик
    (`_repair_scores["terms_lost"]`). Разойдись они — ремонт откатывал бы
    правку из-за потери, которую сам же претензией не считает.

    Снимают потерю двое. Устаревшая проверка: описывает другой текст, значит
    не описывает ничего. Вердикт контекстного арбитра «передан верно»: он
    единственный видел сегмент в ряду соседей, а «туберкулёз лёгких» обратный
    перевод законно возвращает как «лёгочный туберкулёз». (Судья снимает свою
    претензию иначе — вычищая `terms_lost` прямо в записи, см.
    `checks.apply_judge_verdict`.)"""
    bc = seg.get("backcheck") or {}
    if not bc or _check_stale(bc, seg.get("target") or ""):
        return []
    ctx = _term_context_of(seg)
    settled = {_norm_key(f) for c in ctx if c.get("ok") is True
               for f in (c.get("forms") or ())}
    settled |= {_norm_key(c["src"]) for c in ctx if c.get("ok") is True}
    return [t for t in (bc.get("terms_lost") or []) if _norm_key(t) not in settled]


def _repair_findings(seg: dict, project: Optional[dict] = None) -> list:
    """Претензии к ТЕКУЩЕМУ переводу. Устаревшие проверки (перевод правили
    после них) игнорируем — чинить по ним нечего.

    Соответствие глоссарию входит сюда наравне с проверками: цель ремонта —
    привести перевод в порядок целиком, а не по одной подсистеме. Без этого
    ремонт и кнопка «Соответствие глоссарию» чинили сегмент по очереди, каждый
    затирая работу другого."""
    # По ОБРЕЗАННОМУ тексту: back-check и termcheck пишут хеш именно так
    # (см. _check_stale и соседей). Сравнение с необрезанным означало, что
    # у перевода с висящим пробелом находок «нет» — сегмент молча выпадал
    # из ремонта, из корзины «с замечаниями» и из отчёта, а на экране
    # проверки числились свежими. Хеш ремонта (repair.source_hash) остаётся
    # сырым: его так же сырым и пишут.
    cur = _text_hash((seg.get("target") or "").strip())
    # Регистр — такая же бесплатная детерминированная претензия, как
    # соответствие глоссарию, и приходит она отсюда же: приказная запись
    # копируется в перевод вместе со своим начертанием. Устаревать ей нечему,
    # поэтому хеша у неё нет — сверяются нынешние тексты.
    # Начертания приказных терминов здесь НЕТ намеренно: оно чинится
    # детерминированно и бесплатно (`_term_case_fix`, POST …/term-case).
    # Отдать его модели значит платить за переписывание предложения ради
    # одной буквы — и получить откат, потому что «чуть хуже» по баллу
    # перевесит. В `_repair_scores` оно остаётся: ремонт, ЛОМАЮЩИЙ начертание,
    # обязан откатиться.
    items = (_gloss_misses(seg, project) + _consist_misses(seg, project)
             + _case_misses(seg) + _script_misses(seg) + _dup_misses(seg))
    # Потерянные приказные термины — через _terms_lost_open, тем же расчётом,
    # что и счётчик в `_repair_scores`. Он и снимает вердикты контекстного
    # арбитра: арбитр единственный видел сегмент в ряду соседей, и его
    # «передан верно» СНИМАЕТ претензию.
    #
    # «Передан неверно» поводом для ремонта НЕ становится, и это не забывчивость.
    # Арбитр в таком случае предлагает вариант, ОТЛИЧНЫЙ от утверждённого
    # перевода, а `_repair_scores["gloss"]` считает нарушенные приказные термины
    # и всегда: подставь совет — счётчик вырастет, правка откатится. Заход
    # с заранее известным исходом, ровно та оплачиваемая карусель, от которой
    # заведён `_repair_futile`.
    # Спор «проверка против приказа» машина не решает по построению (см. CLAUDE.md),
    # и решать его посегментно бессмысленно: вопрос про ЗАПИСЬ глоссария, а не
    # про строку. Поэтому вердикт уходит человеку в /analysis — с доводом
    # и готовым вариантом. Исправит он запись — и все затронутые сегменты
    # приведёт в порядок существующий расчёт соответствия глоссарию.
    for t in _terms_lost_open(seg):
        items.append({"kind": "term_lost", "must": t,
                      "text": "термин «" + t + "» не пережил обратный перевод"})
    # Вердикт арбитра «передан неверно» — ремонту, но ТОЛЬКО когда предложенный
    # вариант СОГЛАСЕН с приказной записью (содержит утверждённый перевод).
    #
    # Прежде сюда не шло ничего, и довод был такой: арбитр предлагает вариант,
    # отличный от утверждённого, `_repair_scores["gloss"]` считает нарушенные
    # приказные термины всегда, и подстановка совета откатилась бы сама —
    # оплаченная карусель. Довод верен ровно для СПОРА С ЗАПИСЬЮ и неверен для
    # второго случая: когда арбитр говорит «здесь утверждённый термин передан
    # неверно» и предлагает сам этот термин, подстановка счётчик не поднимает,
    # а ОПУСКАЕТ. Карусели тут нет по построению.
    # А смысла в проверке, которая только подтверждает ошибку и оставляет её
    # в тексте, нет: на боевом проекте так стояли «an infected animal» вместо
    # «patient» — арбитр сказал бы «неверно», и всё осталось бы как есть.
    # Спор с записью по-прежнему уходит человеку (`human.termContextWrong`):
    # там вопрос про ЗАПИСЬ, и машина его не решает.
    for c in _term_context_of(seg):
        if c.get("stale"):
            # Забракованное слово, чью негодность подтвердил арбитр, — находка
            # ремонта. Один голос termcheck переписывать не приказывает (тот же
            # закон, что у разнобоя: для применения нужны ДВА голоса), поэтому
            # без вердикта такие слова ждут человека в human.staleFindings.
            # Арбитр и есть второй голос: «негодно» — чиним, «годно» — претензия
            # снята. Слово-приказный-перевод сюда не попадает по построению:
            # _run_segment_term_context такие не спрашивает (спор с записью).
            w = (c.get("tgt") or "").strip()
            if c.get("ok") is False and w:
                rx = _word_re(w)
                if rx and rx.search(seg.get("target") or ""):
                    use = (c.get("use") or "").strip()
                    items.append({
                        "kind": "term_ctx",
                        "replace": [w, use] if use else None,
                        "text": "слово «" + w + "» забраковано проверкой, "
                                "арбитр подтвердил"
                                + (" — " + c["why"] if c.get("why") else "")})
            continue
        use = (c.get("use") or "").strip()
        if c.get("ok") is not False or not use:
            continue
        # Тем же предикатом, каким `_gloss_misses` решает, стоит ли
        # утверждённый перевод в тексте: разойдись они — ремонт подставлял бы
        # совет, который проверка соответствия тут же объявит нарушением.
        if not _tgt_has_term(use, c.get("tgt") or ""):
            continue
        items.append({"kind": "term_ctx", "use": use, "must": c.get("tgt") or "",
                      "text": "термин «" + (c.get("src") or "") + "» передан неверно"
                              + (" — " + c["why"] if c.get("why") else "")})
    bc = seg.get("backcheck") or {}
    if bc and bc.get("target_hash") == cur:
        # По причинам берём только ОБЪЕКТИВНЫЕ (числа, единицы, отрицание,
        # подмена стороны, «обратный перевод про другое»). «Потерян термин»
        # отсюда исключён намеренно: эта претензия уже выставлена выше, строкой
        # `kind: "term_lost"`, и выставлена ПРАВИЛЬНО — с учётом вердикта
        # арбитра (`settled`). Сравнение подстрокой русской фразы такого учёта
        # не знает, поэтому одна и та же жалоба попадала в список дважды,
        # а после того, как судья её снял и написал об этом в reasons, его
        # отречение само становилось поводом для ремонта: строка «„потерян
        # термин: X“ снято судьёй» содержит искомую подстроку. Заход
        # с заранее известным исходом — ровно та карусель, от которой заведён
        # `_repair_futile`.
        for r in (bc.get("reasons") or []):
            if any(h in r for h in BACKCHECK_OBJECTIVE_REASONS):
                items.append({"kind": "backcheck", "text": r})
        j = bc.get("judge") or {}
        # Мнение судьи НЕ гонит ремонт на текст, который только что написала
        # РЕВИЗИЯ (`review.applied` на нынешнем тексте): судья видел оригинал
        # и ОБРАТНЫЙ перевод, ревизор — саму пару. Два оплаченных мнения
        # разошлись — решает человек (по баллу под вердиктом major сегмент
        # уходит в «оценку ниже порога»). Иначе карусель: на боевом #4 ревизия
        # поставила «Phthisiology», судья на обратном «Фтизиология» объявил
        # «другую дисциплину», ремонт вернул кальку «Phthisiatry» с баллом
        # 100 (калька возвращается дословно), ревизия снова устарела — и
        # следующий прогон повторил бы всё за те же деньги. Объективные
        # причины выше остаются: мнение снимает только мнение.
        if j.get("severity") in ("major", "critical") and not _review_wrote(seg):
            for d in (j.get("divergences") or []):
                items.append({"kind": "judge", "text": str(d)})
            if j.get("comment"):
                items.append({"kind": "judge", "text": j["comment"]})
    tc = seg.get("termcheck") or {}
    # Свежее ручательство ревизии (`_review_vouches`) снимает мнение termcheck
    # о строке: та смотрит только на перевод, а ревизор прочитал пару целиком
    # и поставил ≥ REVIEW_VOUCH_SCORE. Одно суждение termcheck переписывать
    # не приказывает (закон двух голосов), а против него здесь стоит более
    # осведомлённое мнение. Снимается ТОЛЬКО мнение модели: детерминированные
    # претензии выше остаются, объективную находку ручательство не переживает
    # по построению. Считается ЗДЕСЬ, а не в разборе экрана: список претензий
    # один на состав прогона, ремонт, отпечаток захода и корзины — снятое
    # в одном месте вернулось бы в другом платным заходом.
    if tc and tc.get("target_hash") == cur and not _review_vouches(seg):
        # minor входит наравне с critical/major. Раньше он не входил никуда:
        # ремонт его не брал, а _machine_clean всё равно объявлял сегмент
        # нечистым — на боевом проекте 168 сегментов висели между двумя
        # политиками и не двигались ни в какую сторону. При этом сами находки
        # содержательные: «EPT → EPTB», «Amount → Number», «extensive drug
        # resistance → extensively drug-resistant». Их исправление — ровно та
        # автоматизация, ради которой ремонт и заведён.
        # Тяжесть едет вместе с находкой: по ней _repair_scores решает, стало
        # ли лучше, и без неё размен major на minor был бы неотличим от
        # размена minor на major.
        for f in (tc.get("findings") or []):
            # Находка ПРОТИВ приказной записи ремонту не отдаётся: подстановка
            # совета поднимет счётчик нарушенных приказных терминов, и правка
            # откатится сама. Спор про ЗАПИСЬ решает человек — он видит его
            # в `human.termcheckDisputes`. Метку ставит `_note_term_disputes`
            # там, где есть проект (здесь его может не быть).
            if f.get("severity") in TERMCHECK_ACTIONABLE and not f.get("vsVerified"):
                items.append({"kind": "term", "sev": f.get("severity"),
                              "replace": [f.get("tgt_term", ""), f.get("suggestion", "")],
                              "text": "«" + f.get("tgt_term", "") + "»"
                                      + (" → «" + f["suggestion"] + "»" if f.get("suggestion") else "")
                                      + (" — " + f["why"] if f.get("why") else "")})
    seen, out = set(), []
    for i in items:
        if i["text"] not in seen:
            seen.add(i["text"])
            out.append(i)
    return out


# Версия ПРАВИЛ, по которым ремонт решает «принять или откатить». Поднимается
# при каждой правке этих правил — как BACKCHECK_VERSION у формулы балла.
# Вердикт, вынесенный прежними правилами, нынешние правила не описывает,
# и держать по нему сегмент закрытым значит отказывать в починке по решению,
# которого больше нет. Из переменной окружения — чтобы поднять её можно было
# и без выката кода (например, после правки политики в данных).
REPAIR_RULES_VERSION = os.environ.get("REPAIR_RULES_VERSION", "4")


def _repair_attempt_key(seg: dict, findings: Optional[list] = None) -> str:
    """Отпечаток ЗАХОДА: что ремонт увидит, если пойдёт сюда сейчас.

    Три составляющие, и каждая меняет исход:
      • сам текст — иначе чиним другое;
      • СПИСОК претензий — от него зависит промпт, а значит и ответ модели;
      • версия правил — от неё зависит, примут правку или откатят.

    Отпечаток и есть ответ на вопрос «даст ли второй заход что-то новое».
    Совпал — заход вернёт то же самое за те же деньги, и его не делаем;
    разошёлся — работа новая, и запрещать её незачем. Так второй заход
    разрешён там, где он осмыслен, и запрещён там, где это перерасход.

    Глоссарий в отпечаток НЕ входит, и это не упущение. Считается он всюду
    одинаково — `_repair_findings(seg, None)`, — потому что иначе отпечаток
    зависел бы от того, кто его считает: разбор состава ходит без проекта ради
    скорости, `_segment_for_client` проекта не имеет вовсе, а прогон имеет.
    Разойдись они — сегмент чинился бы по кругу либо не чинился никогда.
    Смену глоссария по-прежнему открывает `retry=True` у «Применить термины»,
    а бессмысленный повтор там останавливает `_repair_futile`.

    `findings` — уже посчитанный `_repair_findings(seg, None)`: он не бесплатен
    (письменность, начертание, триграммы — около миллисекунды на сегмент),
    а зовущие его обычно только что посчитали то же самое. Передавать сюда
    можно РОВНО этот список: посчитанный с проектом (то есть с глоссарием)
    сделал бы отпечаток зависящим от того, кто его считает, — а это ровно то,
    от чего предостерегает абзац выше."""
    items = sorted(f["text"] for f in
                   (findings if findings is not None else _repair_findings(seg, None)))
    return _text_hash("|".join([_text_hash(seg.get("target") or ""),
                                REPAIR_RULES_VERSION] + items))


def _repair_tried(seg: dict) -> bool:
    """Такой же заход уже делали — второй даст то же самое за те же деньги.

    «Такой же» — это совпавший отпечаток (текст + претензии + версия правил),
    а не просто тот же текст. Прежняя проверка смотрела только на текст, и
    поэтому держала сегмент закрытым даже после того, как появились НОВЫЕ
    находки или изменились правила отмены: на боевом проекте так простаивали
    93 сегмента из 115 найденных дефектов — не «не смогли», а «не пустили».

    Отпечатка нет (записи прежних версий) — откатываемся на старое сравнение
    по тексту, иначе разом открылся бы весь проект."""
    r = seg.get("repair") or {}
    key = r.get("attemptKey")
    if key:
        return key == _repair_attempt_key(seg)
    return r.get("source_hash") == _text_hash(seg.get("target") or "")


def _repair_clamped(seg: dict, findings: Optional[list] = None,
                    model: Optional[str] = None) -> bool:
    """Закрыт ли сегмент от ремонта. Клеймит ТОЛЬКО совпавший отпечаток.

    `findings` — готовый `_repair_findings(seg, None)`, если он уже посчитан
    у зовущего (см. `_repair_attempt_key`): экран «Анализ» считает его на
    каждый сегмент, и второй проход стоил бы полсекунды единственного воркера
    на боевом проекте — при том, что ответ лежит строкой выше.

    Отличие от `_repair_tried` существенное и вот в чём. `_repair_tried`
    отвечает на вопрос «проходил ли ЭТОТ текст через ремонт» — на нём стоят
    экраны и разбор наследства, и он должен работать для записей любого
    возраста. А «не пускать сюда снова» — вопрос другой: вердикт, вынесенный
    ПРЕЖНИМИ правилами, нынешние правила не описывает.

    У записей без отпечатка (сделаны до появления `REPAIR_RULES_VERSION`)
    правила неизвестны, а менялись они с тех пор не раз: вето по баллу,
    подгонка начертания, зачёт частичного успеха, отсев споров с приказной
    записью. Держать по такому вердикту сегмент закрытым — отказывать
    в починке по решению, которого больше нет. На боевом проекте так стояли
    79 сегментов с открытыми находками при пустом составе ремонта.

    Открывается при этом НЕ весь проект: в состав попадают только сегменты
    с находками, а их считает разбор и называет числом до запуска.

    `model` — РАЗРЕШЁННЫЙ id модели, которой заход пойдёт сейчас. Другая
    модель — другой заход: тот же промпт она ответит по-своему, и держать
    сегмент закрытым вердиктом чужой модели значит отказывать во втором
    мнении, которое ремонту прямо обещано («дальше только человек или другая
    модель»). Записи помнят, кто уже ходил на этот отпечаток (`triedModels`;
    у старых записей — их `model`, она пишется с самого начала). `model=None`
    означает «какой моделью пойдут, неизвестно» (экран «Анализ»,
    `_segment_for_client`) — клеймо держит, как и раньше: обещать «возьмёт
    прогон» про сегмент, который возьмёт только смена модели, нельзя."""
    rp = seg.get("repair") or {}
    key = rp.get("attemptKey")
    if not key or key != _repair_attempt_key(seg, findings):
        return False
    if model is None:
        return True
    # Правило «кто уже ходил» — одно, в _models_tried: вторая копия вывода
    # списка из старой записи разошлась бы с ним первой же правкой.
    tried = _models_tried(seg, key)
    # Запись без модели (пустой список) — чей заход, неизвестно: клеймо держит.
    return not tried or model in tried


def _models_tried(seg: dict, attempt_key: str) -> list:
    """Модели, уже ходившие на ЭТОТ отпечаток захода.

    У старых записей списка нет — берём модель самой записи: она пишется
    в запись ремонта с первого дня. Запись о ДРУГОМ отпечатке не в счёт:
    сменился текст или претензии — прежние попытки к новому заходу
    отношения не имеют."""
    rp = seg.get("repair") or {}
    if rp.get("attemptKey") != attempt_key:
        return []
    tried = rp.get("triedModels")
    if tried is None:
        tried = [rp["model"]] if rp.get("model") else []
    return list(tried)


def _repair_futile(seg: dict, project: Optional[dict] = None) -> bool:
    """Второй заход по ЭТОМУ тексту с ЭТИМИ претензиями даст то же самое.

    `_repair_tried` смотрит только на текст, и этого мало: глоссарий меняется,
    и по НОВОЙ претензии зайти на тот же текст осмысленно. А вот когда и текст
    тот же, и список претензий тот же, повторять нечего — модель получит тот
    же промпт и вернёт тот же ответ, а оценка отвергнет его так же.

    Из-за чего написано: `apply_terms` ходит со `retry=True` (глоссарий с
    прошлого раза мог измениться) и потому забирал сегменты, где ремонт уже
    провалился на том же тексте. На боевых данных это все 57 оставшихся
    расхождений: каждое нажатие «Применить» стоило денег и не меняло ничего,
    а список оставался тем же — работа выглядела нескончаемой."""
    rp = seg.get("repair") or {}
    if not _repair_tried(seg):
        return False
    # Записи, сделанные до того, как претензия про потерянный термин перестала
    # дублироваться строкой причины, несут её ДВАЖДЫ: «термин «X» не пережил
    # обратный перевод» из terms_lost и «потерян термин: X» из reasons. Второй
    # строки больше не бывает, и без этой чистки сравнение «тот же список»
    # никогда бы не сошлось — то есть сегмент, который ремонт уже пробовал
    # и откатил, снова пошёл бы в платный заход с заранее известным исходом.
    was = {i for i in (rp.get("issues") or ()) if not i.startswith("потерян термин")}
    if not was:
        return False
    return was == {f["text"] for f in _repair_findings(seg, project)}


def _repair_score_vetoed(seg: dict) -> bool:
    """Правку отменил ТОЛЬКО упавший балл back-check, а термины она почистила.

    Это разбор НАСЛЕДСТВА, а не действующая ветка: нынешний `_run_segment_repair`
    такую правку больше не откатывает (см. там про кальку и `_content_recall`).
    Но записи, сделанные прежним правилом, никуда не делись — на боевом проекте
    это 111 сегментов, где перевод и сегодня несёт «medicine physicians» вместо
    «pulmonologists», а верный вариант лежит рядом в `repair.candidate`,
    написанный и оплаченный.

    Предикат намеренно узкий, три условия:
      1) причина отмены РОВНО одна и это балл — «;» в строке означает, что были
         и другие претензии (термины, глоссарий, регистр), а они законны;
      2) термины при этом стали ЧИЩЕ по бесплатным объективным счётчикам;
      3) кандидат относится к НЫНЕШНЕМУ тексту (`_repair_tried`) — иначе перевод
         правили после отмены, и подставлять старого кандидата значит выбросить
         чужую работу, ровно как в `_revert_repairs`.

    Сравнение по строке причины здесь законно (в отличие от корзин `/analysis`,
    где заведены CLEAN_*): строку писал этот же файл, она лежит в данных
    прошлых прогонов и переформулировать её задним числом нельзя."""
    rp = seg.get("repair") or {}
    if rp.get("applied") or not (rp.get("candidate") or "").strip():
        return False
    if not _repair_tried(seg):
        return False
    why = rp.get("reason") or ""
    if not why.startswith("балл back-check упал") or ";" in why:
        return False
    # ЖЁСТКАЯ находка на кандидате — отказ по существу, а не «упал балл».
    # Строка причины у обоих исходов почти одна, поэтому решает поле записи.
    if rp.get("hardAfter"):
        return False
    b, a = rp.get("before") or {}, rp.get("after") or {}
    cleaner = bool(
        (b.get("terms") is not None and a.get("terms") is not None
         and a["terms"] < b["terms"])
        or (b.get("gloss") is not None and a.get("gloss") is not None
            and a["gloss"] < b["gloss"]))
    if not cleaner:
        return False
    # Второй рубеж, и он нужен из-за НАСЛЕДСТВА: у записей прежнего кода поля
    # `hardAfter` нет вовсе, а отменить их могла ровно та же жёсткая находка —
    # строка причины тогда писалась одна на оба исхода. Поэтому кандидата
    # сверяем с оригиналом ЗАНОВО, детерминированными правилами: числа,
    # единицы и отрицание считаются по паре «оригинал → перевод» без единого
    # вызова модели и без знания языка.
    # Стоит ПОСЛЕДНИМ: предикат зовётся на каждый сегмент при выдаче проекта
    # (2711 штук), а до этой строки доходят единицы.
    # Чего этот рубеж не ловит у старых записей — подмену стороны из
    # DOMAIN_RULES, если правил для пары языков нет: тот же закон, молчим.
    if checks_mod:
        try:
            bad = checks_mod.deterministic_issues(
                seg.get("source") or "", rp.get("candidate") or "")
            if any(i.get("type") in checks_mod.BACKCHECK_HARD_TYPES
                   or i.get("type") in ("number_unit_dosage_mismatch", "negation_shift")
                   for i in (bad or [])):
                return False
        except Exception as e:                                  # pragma: no cover
            print(f"[backend] сверка кандидата seg#{seg.get('id')}: {e}", file=sys.stderr)
            return False        # не смогли проверить — не предлагаем
    return True


def _repairable(seg: dict, allow_tried: bool = False, project: Optional[dict] = None,
                model: Optional[str] = None) -> bool:
    """allow_tried — человек сам отметил галочкой уже чинившиеся сегменты.
    По умолчанию второй заход по тому же тексту не делаем: те же претензии
    дадут тот же результат за те же деньги. `model` — разрешённый id модели
    ЭТОГО прогона: другая модель — другой заход (см. _repair_clamped)."""
    if not (seg.get("target") or "").strip() or not _repair_findings(seg, project):
        return False
    return allow_tried or not _repair_clamped(seg, model=model)


def _repair_system(dom: dict, src_lang: str, tgt_lang: str, style: str = "") -> str:
    return (
        "You are a senior " + dom["expert"] + ". You are given a SOURCE text in " + src_lang
        + ", its TRANSLATION into " + tgt_lang + ", and a list of ISSUES found by quality control.\n\n"
        "Fix ONLY the listed issues. This is a repair, not a retranslation.\n"
        "RULES:\n"
        "1. Return ONLY the corrected translation — no explanations, no comments, no quotes.\n"
        "2. Change as little as possible: keep every wording that is not part of an issue.\n"
        "3. Required replacements must be applied exactly as given.\n"
        "4. Terms reported as lost MUST be present in the corrected translation.\n"
        "5. Keep all numbers, units, negations and abbreviations exactly as in the SOURCE.\n"
        "6. Use " + dom["terminology"] + ". Output must be 100% " + tgt_lang + " — not a single\n"
        "   letter of the source script, inside formulas and abbreviations included.\n"
        "7. If an issue looks wrong to you, leave that part unchanged rather than inventing something new.\n"
        # То же правило, что в промпте перевода, и по той же причине: ремонт
        # переписывает текст последним, и промолчи он про регистр — вернул бы
        # строчную в начало заголовка следом за находкой про термин.
        "8. Follow the capitalisation of the SOURCE: a sentence, heading or caption that starts\n"
        "   with a capital letter starts with one in the translation, an ALL-CAPS heading stays\n"
        "   ALL-CAPS, and no word is shouted that the SOURCE does not shout. Glossary terms are\n"
        "   listed with their case already matched to the source — copy them as printed.\n"
        + ("\n" + style if style else "")
    )


def _openai_repair(seg: dict, project: dict, findings: list, model: Optional[str]) -> Optional[str]:
    import openai
    dom = _resolve_domain(project.get("domain"))
    mdl = _resolve_model(model or REPAIR_DEFAULT_MODEL)
    client = openai.OpenAI(api_key=os.environ.get("OPENAI_API_KEY"), timeout=120, max_retries=1)
    lines = []
    for i, f in enumerate(findings, 1):
        lines.append(str(i) + ". " + f["text"])
        if f.get("replace") and f["replace"][1]:
            lines.append('   REQUIRED: replace "' + f["replace"][0] + '" with "' + f["replace"][1] + '"')
        if f.get("must"):
            lines.append('   REQUIRED: the translation must convey "' + f["must"] + '"')
        if f.get("use"):
            lines.append('   REQUIRED: use exactly this approved term: "' + f["use"] + '"')
    body = ("SOURCE:\n" + seg.get("source", "") + "\n\nTRANSLATION:\n" + (seg.get("target") or ""))
    # Проверенные записи глоссария кладём в промпт ЦЕЛИКОМ, а не только те, что
    # нарушены: чиня одно, модель свободно переписывала соседний утверждённый
    # термин — и сегмент возвращался в отчёт о соответствии уже по другой строке.
    approved = _verified_hits(seg.get("source", ""), project)
    if approved:
        body += ("\n\nAPPROVED GLOSSARY for this segment — the right-hand side must be present "
                 "in the corrected text letter for letter (its case is already matched to the "
                 "source fragment on the left):\n"
                 + "\n".join("  " + (h.get("_form") or h["src"]) + " → "
                              + _case_like(h.get("_form") or h["src"], h["tgt"], h["src"])
                              for h in approved))
    doc = _doc_hits(seg.get("source", ""), project, approved)
    if doc:
        body += ("\n\nDOCUMENT TERM SHEET (agreed for this document, not verified by a human — "
                 "the approved glossary above always wins):\n"
                 + "\n".join("  " + (h.get("_form") or h["src"]) + " → " + h["tgt"] for h in doc))
    back = ((seg.get("backcheck") or {}).get("back") or "").strip()
    if back:
        body += ("\n\nBACK-TRANSLATION of the current translation (for reference — it shows how the "
                 "translation reads to a reviewer):\n" + back)
    body += "\n\nISSUES:\n" + "\n".join(lines)
    extra = ({"max_completion_tokens": 4096} if mdl["api"] == "modern"
             else {"max_tokens": 1500, "temperature": 0})
    try:
        resp = client.chat.completions.create(
            model=mdl["id"],
            messages=[{"role": "system", "content": _repair_system(dom, project.get("src", "RU"), project.get("tgt", "EN"),
                                                                  _style_block(project))},
                      {"role": "user", "content": body}],
            **extra,
        )
        _note_usage("repair", mdl["id"], resp)
        return (resp.choices[0].message.content or "").strip()
    except Exception as e:
        print(f"[backend] repair failed seg#{seg.get('id')}: {e}", file=sys.stderr)
        return None


_DUP_WORD_RE = re.compile(r"[^\W\d_]+(?:-[^\W\d_]+)*", re.UNICODE)


def _dup_count(text: str) -> int:
    """Сколько ЛИШНИХ повторов трёхсловных сочетаний в тексте.

    Мера грубая и в одиночку ничего не значит: в длинном абзаце «number of
    patients» законно встречается дважды. Поэтому она и не показывается
    человеку — её сравнивают ТОЛЬКО «до и после» одной и той же правки, где
    шум сокращается: выросло — значит повтор добавила именно правка.

    Триграммы, а не слова: повтор одного слова в тексте норма, повтор трёх
    подряд — почти всегда дописанный кусок."""
    w = _DUP_WORD_RE.findall(text or "")
    if len(w) < 3:
        return 0
    seen: dict = {}
    for i in range(len(w) - 2):
        k = " ".join(w[i:i + 3]).lower()
        seen[k] = seen.get(k, 0) + 1
    return sum(c - 1 for c in seen.values() if c >= 2)


# Слово ИЛИ ФРАЗА повторены сами собой в скобках: «Prevalence (prevalence)»,
# «Eales disease (Eales disease)». Четыре буквы и больше в ПЕРВОМ слове —
# короткие («ТБ (TB)») бывают законной расшифровкой.
#
# Фраза, а не одно слово, потому что между двумя правилами была дыра ровно
# в один шаг: скобку ловило только повторённое ОДНО слово, а подряд идущий
# повтор — только от ТРЁХ слов (_adjacent_repeat). Двусловное «Eales disease
# (Eales disease)» не ловилось ни тем, ни другим и уходило в готовый перевод
# при балле back-check 4/10, где претензия была совсем о другом. Расшифровкой
# такое не бывает: аббревиатуру раскрывают ДРУГИМИ словами, а не теми же.
_SELF_GLOSS_RE = re.compile(
    r"\b([^\W\d_]{4,}(?:[ \t\-][^\W\d_]+){0,3})\s*\(\s*\1\s*\)",
    re.IGNORECASE | re.UNICODE)


def _adjacent_repeat(text: str) -> list:
    """Подряд идущие повторы: хвост из N слов равен следующим N словам.

    Один расчёт на ОРИГИНАЛ и на ПЕРЕВОД: по нему же решается, не повторяется
    ли автор сам, — а два расчёта одного и того же однажды разойдутся."""
    w = _DUP_WORD_RE.findall(text or "")
    low = [x.lower() for x in w]
    out, seen = [], set()
    for n in (6, 5, 4, 3):
        for i in range(len(low) - 2 * n + 1):
            if i in seen:
                continue
            if low[i:i + n] == low[i + n:i + 2 * n]:
                seen.update(range(i, i + 2 * n))
                out.append((" ".join(w[i:i + n]), n))
    return out


def _dup_misses(seg: dict) -> list:
    """Текст сам себя повторяет. Бесплатно, детерминированно, без знания языка.

    Две находки, и обе были СОВЕРШЕННО невидимы всем проверкам:
      1. подряд идущий повтор из трёх и более слов — «Infiltrative pulmonary
         tuberculosis Infiltrative pulmonary tuberculosis» (балл back-check 100,
         termcheck молчит);
      2. слово, поясняющее само себя скобкой, — «Prevalence (prevalence)»
         (тоже 100). Так схлопываются два РАЗНЫХ термина оригинала
         («Распространённость (болезненность)») в одно английское слово,
         напечатанное дважды.

    Почему этого не видел никто: back-check считает долю слов ОРИГИНАЛА,
    вернувшихся через обратный перевод, и от повтора она не падает; termcheck
    спрашивает, нормальный ли это термин целевого языка, — а повторённый термин
    нормален; глоссарию повтор угодил. Проверка ничего не стоит и потому идёт
    наравне с регистром и чужим письмом.

    Скобка ловится только когда в ней стоит ТО ЖЕ САМОЕ слово. «ТБ (tuberculosis)»
    — законная расшифровка, и таких мы не трогаем; «Prevalence (prevalence)»
    законным не бывает ни при каком оригинале. Наличие скобки в оригинале
    признаком не является: у #128 она там есть, и именно в неё схлопнулись
    два разных русских термина."""
    tgt = (seg.get("target") or "").strip()
    if not tgt:
        return []
    src = seg.get("source") or ""
    out = []
    # Повтор ЕСТЬ И В ОРИГИНАЛЕ — значит так написал автор, и отличить его
    # повтор от нашего нечем: языки разные, n-граммы не сопоставимы. Молчим,
    # тот же закон, что у DOMAIN_RULES. Правило не косметическое: на боевом
    # проекте так отсеиваются 3 находки из 12 («туберкулёз семенных пузырьков»
    # повторён в самом оригинале), а в договорах и законах повтор наименований
    # сторон и определённых терминов — норма оформления, и там доля ложных
    # была бы много выше. Система не только медицинская.
    #
    # Гасит оно РОВНО СВОЮ находку, а не обе. Прежде любой авторский повтор
    # подряд закрывал заодно и скобку — а это разные улики: то, что автор
    # дважды назвал стороны договора, ничего не говорит о том, законно ли
    # в переводе стоит «Prevalence (prevalence)». Скобку гасит только такая же
    # скобка в оригинале.
    if not _adjacent_repeat(src):
        for a, b in _adjacent_repeat(tgt):
            out.append({"kind": "dup", "text": "кусок повторён подряд: «" + a + "»"})
    if not _SELF_GLOSS_RE.search(src):
        for m in _SELF_GLOSS_RE.finditer(tgt):
            out.append({"kind": "dup",
                        "text": "слово поясняет само себя: «" + m.group(0) + "»"})
    return out


# Сколько раз пару «было → надо» должен назвать termcheck, чтобы считать её
# мнением о ДОКУМЕНТЕ, а не придиркой к одной строке.
#
# Двойка, а не единица, и это принципиально. termcheck — ОДИН вызов модели,
# и он ошибается: на боевом учебнике из 47 пар 45 держались на единственном
# голосе, и среди них «Bronchoscopy → sputum smear microscopy» (совсем другая
# процедура, 16 мест), «bacillary excretion → bacteriological conversion»
# (противоположный смысл, 27 мест) и «infiltrative pulmonary tuberculosis →
# inflammatory tuberculosis» (другая болезнь, 10 мест). Применить их по всему
# документу было бы много хуже исходного разнобоя. На двух голосах выжили
# ровно две пары, и обе верные: «caverns → cavities» и «fibrocavernous →
# fibrocavitary». Мнение о ДОКУМЕНТЕ должно быть высказано дважды —
# иначе это мнение о строке. Из окружения, чтобы порог можно было поднять
# ещё выше на шумном проекте.
CONSIST_MIN_VOTES = int(os.environ.get("CONSIST_MIN_VOTES", "2"))
# Потолок пар в отчёте: список ведёт человек, и бесконечным он быть не должен.
CONSIST_MAX_PAIRS = int(os.environ.get("CONSIST_MAX_PAIRS", "40"))


def _word_re(term: str):
    """Поиск варианта в переводе по границам слова. Границей цифра не считается
    только там, где сам вариант ею не кончается, — иначе «CD4» ловил бы «CD40»
    (то же правило, что у подбора терминов)."""
    t = (term or "").strip()
    if not t:
        return None
    key = "consist::" + t
    rx = _PATTERN_CACHE.get(key)
    if rx is None:
        rx = re.compile(r"(?<![^\W\d_])" + re.escape(t) + r"(?![^\W\d_])",
                        re.IGNORECASE | re.UNICODE)
        _PATTERN_CACHE[key] = rx
    return rx


def _consistency_pairs(project: dict) -> list:
    """Разнобой по ДОКУМЕНТУ: один и тот же оборот переведён по-разному.

    Ни back-check, ни termcheck, ни глоссарий не смотрят дальше своего
    сегмента, поэтому книга спокойно пишет «MBT» в 37 местах и «MTB» в 49,
    а титул говорит «Phthisiatry» при «phthisiology» в тексте — и всё это
    «чисто». Читателю же нужен единый перевод, а не среднее по сегментам.

    Откуда берутся пары «было → надо»: из САМИХ НАХОДОК termcheck. Он уже
    сказал про какой-то сегмент «здесь X, а надо Y» — это и есть суждение
    о терминологии, и его незачем спрашивать второй раз про каждое из
    оставшихся мест. Так одна оплаченная находка распространяется на весь
    документ бесплатно, а ручной работы становится одно решение на пару
    вместо одного на сегмент.

    Никаких зашитых списков вариантов тут нет и быть не может: система
    работает с любой областью и любой парой языков, а словарь вариантов
    существует только в данных проекта."""
    segs = project.get("segments") or []
    votes: dict = {}
    for s in segs:
        tc = s.get("termcheck") or {}
        if _check_stale(tc, s.get("target") or ""):
            continue
        for f in (tc.get("findings") or []):
            if f.get("severity") not in TERMCHECK_DISPUTING:
                continue
            was, want = (f.get("tgt_term") or "").strip(), (f.get("suggestion") or "").strip()
            if not was or not want or was.lower() == want.lower():
                continue
            v = votes.setdefault((was.lower(), want.lower()),
                                 {"was": was, "want": want, "votes": 0, "why": f.get("why") or ""})
            v["votes"] += 1
    out = []
    for v in votes.values():
        if v["votes"] < CONSIST_MIN_VOTES:
            continue
        rx_was, rx_want = _word_re(v["was"]), _word_re(v["want"])
        if not rx_was or not rx_want:
            continue
        with_was = [s["id"] for s in segs if rx_was.search(s.get("target") or "")]
        with_want = [s["id"] for s in segs if rx_want.search(s.get("target") or "")]
        if not with_was:
            continue
        out.append({"was": v["was"], "want": v["want"], "why": v["why"],
                    "votes": v["votes"], "segments": with_was,
                    # Сколько мест уже пишут правильно. Ноль означает, что
                    # разнобоя нет — есть один сплошной вариант, и вопрос
                    # к нему тот же, но человек должен видеть разницу.
                    "already": len(with_want)})
    out.sort(key=lambda x: (-len(x["segments"]), x["was"]))
    return out[:CONSIST_MAX_PAIRS]


_CONSIST_CACHE: dict = {}


def _consistency_of(project: dict) -> list:
    """Пары разнобоя проекта, посчитанные один раз на отпечаток.

    Считать их на каждый сегмент нельзя: это проход по всему проекту на каждую
    пару. Отпечаток тот же, что у отчёта о соответствии, — тексты и находки."""
    pid = project.get("id")
    fp = _text_hash("|".join(
        (s.get("target") or "") + str(len(((s.get("termcheck") or {}).get("findings") or [])))
        for s in (project.get("segments") or [])))
    hit = _CONSIST_CACHE.get(pid)
    if hit and hit[0] == fp:
        return hit[1]
    pairs = _consistency_pairs(project)
    _CONSIST_CACHE[pid] = (fp, pairs)
    return pairs


def _consist_misses(seg: dict, project: Optional[dict]) -> list:
    """Сегмент пишет то, что termcheck УЖЕ забраковал в другом месте документа.

    Это и есть «проверка смотрит дальше своего сегмента»: находку оплатили
    один раз, а применяется она везде, где встретился тот же оборот.
    Совет, спорящий с приказной записью глоссария, сюда не идёт — тот же
    закон, что у вердикта арбитра: подстановка подняла бы счётчик нарушенных
    приказных терминов, и правка откатилась бы сама."""
    if not project:
        return []
    tgt = seg.get("target") or ""
    if not tgt.strip():
        return []
    # Приказные переводы ЭТОГО сегмента: если разнобой предлагает заменить
    # то, что требует глоссарий, — молчим. Подстановка подняла бы счётчик
    # нарушенных приказных терминов, правка откатилась бы сама, и вышел бы
    # платный вызов с заранее известным исходом. Спор с записью решает
    # человек, а не документ (тот же закон, что у вердикта арбитра).
    ordered = {(h.get("tgt") or "").lower() for h in _verified_hits(seg.get("source") or "", project)}
    # Про что termcheck уже сказал В ЭТОМ сегменте. Об одном и том же дважды
    # не говорим: смысл проверки — донести находку до ОСТАЛЬНЫХ мест
    # документа, а здесь она и так есть, со своей тяжестью и своим доводом.
    # Под ручательством ревизии (`_review_vouches`) находки termcheck из
    # списка претензий сняты — значит «здесь она и так есть» неправда,
    # и разнобой обязан сказать сам: у него ДВА голоса по документу.
    tc = seg.get("termcheck") or {}
    own = ({(f.get("tgt_term") or "").lower() for f in (tc.get("findings") or [])}
           if not _check_stale(tc, tgt) and not _review_vouches(seg) else set())
    out = []
    for pr in _consistency_of(project):
        rx = _word_re(pr["was"])
        if not rx or not rx.search(tgt):
            continue
        if pr["was"].lower() in ordered or pr["was"].lower() in own:
            continue
        out.append({"kind": "consist", "replace": [pr["was"], pr["want"]],
                    "text": "по документу принято «" + pr["want"] + "», здесь «"
                            + pr["was"] + "»"
                            + (" — " + pr["why"] if pr.get("why") else "")})
    return out


def _repair_scores(seg: dict, project: Optional[dict] = None,
                   doc_skip: Optional[set] = None) -> dict:
    """Снимок качества сегмента: балл back-check, число серьёзных замечаний
    по терминам и число нарушенных утверждённых терминов. По нему решаем,
    стало ли лучше. Глоссарий считается всегда: это единственная из трёх
    оценок, которая не стоит ни одного вызова модели."""
    bc = seg.get("backcheck") or {}
    tc = seg.get("termcheck") or {}
    # terms = None, если проверка этого текста не видела: ноль замечаний
    # у непроверенного текста — не «чисто», а «неизвестно». Сравнение с таким
    # нулём откатывало бы верную правку из-за унаследованной проблемы.
    fresh_tc = bool(tc) and not _check_stale(tc, seg.get("target") or "")
    findings = (tc.get("findings") or []) if fresh_tc else ()
    return {
        "score": bc.get("score"),
        "terms": (len([f for f in findings
                       if f.get("severity") in ("critical", "major")])
                  if fresh_tc else None),
        # Отдельным числом, а не в общей сумме: размен серьёзного замечания
        # на мелкое — выигрыш, и сложи мы их вместе, он выглядел бы ничьёй.
        # Считается всегда, а сверяется только когда чинили ИЗ-ЗА мелких
        # (см. _run_segment_repair): иначе правка ради глоссария откатывалась
        # бы из-за побочного minor, то есть п.3 отнимал бы автоматизацию
        # ровно там, где должен её добавить.
        "terms_minor": (len([f for f in findings
                             if f.get("severity") == "minor"])
                        if fresh_tc else None),
        "gloss": len(_gloss_misses(seg, project)),
        # Потерянные ПРИКАЗНЫЕ термины — ещё один бесплатный счётчик. Нужен
        # потому, что после перехода на долевой штраф потеря термина почти
        # не двигает балл на длинном сегменте (одно слово из сорока — два-три
        # пункта), а `gloss` её не видит вовсе: он смотрит, стоит ли
        # утверждённый вариант в ПЕРЕВОДЕ, и перефразировку, из-за которой
        # термин перестал возвращаться через обратный перевод, не замечает.
        # Без этого счётчика ремонт мог принять правку, потерявшую термин,
        # а следующий прогон завёл бы по ней новую находку — второй платный
        # заход. None по тому же правилу, что у `terms`: back-check не видел
        # этого текста, значит ноль означает «неизвестно», а не «чисто».
        # None и тогда, когда претензию СНЯЛ арбитр: после правки его вердикт
        # устареет вместе с текстом, `_terms_lost_open` вернёт сырой список
        # нового back-check, и правка откатится по претензии, которую только
        # что отменил тот, кто читал оба текста, — платный заход с заранее
        # известным исходом.
        "terms_lost": (len(_terms_lost_open(seg))
                       if bc and not _check_stale(bc, seg.get("target") or "")
                       and not any(x.get("ok") is True
                                   for x in ((seg.get("termContext") or {}).get("terms") or ()))
                       else None),
        # Как и глоссарий, считается ВСЕГДА и бесплатно: правка ради термина
        # не должна попутно ронять заглавную в начале заголовка.
        "case": len(_case_misses(seg)),
        # И буквы чужого письма — тоже бесплатно и тоже всегда.
        "script": len(_script_misses(seg)),
        # Терм-лист документа: бесплатный счётчик, не находка — ремонт за ним
        # не ходит, но ломать согласованную пару правкой по другой претензии
        # не вправе. Без включённого терм-листа всегда ноль.
        "doc": len(_doc_misses(seg, project, doc_skip)),
        # Повторы. Ремонту говорят «термин потерян», и он ДОПИСЫВАЕТ вариант
        # вместо замены: «areas of increased bone density, areas of increased
        # bone density», «pulmonary tuberculosis (lung tuberculosis)» через
        # весь абзац, пять имён одного диагноза подряд. Не видит этого никто:
        # back-check считает долю слов ОРИГИНАЛА, вернувшихся через обратный
        # перевод, и от лишних слов только растёт; termcheck смотрит на термины,
        # а не на повторы; глоссарию дописанный термин угодил. На боевом
        # проекте так испорчено 12 правок из 1176 — мало, но чинит это только
        # человек, а стоит бесплатно.
        "dup": _dup_count(seg.get("target") or ""),
        # Сами НАХОДКИ самоповтора — скобка «Prevalence (prevalence)» и повтор
        # подряд. `dup` выше их не считает (там триграммы), и заход ТОЛЬКО по
        # такой находке мерился нулём «0 → 0»: верная правка «Prevalence
        # (morbidity)» на боевом #128 откатывалась и клеймилась `tried`.
        # Бесплатно, как всё в этой группе.
        "self_dup": len(_dup_misses(seg)),
        # Регистр приказных терминов: бесплатно, но глоссарий нужен — с
        # project=None считается нулём, как и `gloss` рядом.
        "term_case": len(_term_case_misses(seg, project)),
    }


def _repair_desync(seg: dict, want: str) -> bool:
    """Текст сегмента разошёлся с решением, которое только что записано.

    Такого быть не должно: и откат, и применение кладут текст сами, следующей
    строкой. Но в данных прогона от 22.08 нашлись 34 сегмента, где запись
    о ремонте описывает один текст, а в сегменте лежит другой — отвергнутый
    вариант на месте отката и дореморный на месте применённой правки.
    Воспроизвести на нынешнем коде не удалось, а молча жить с расхождением
    нельзя: по этим записям считаются «откачено N» в отчёте прогона, и по ним
    же ремонт решает, заходить ли на сегмент второй раз.

    Поэтому условие проверяется на выходе. Не чинит — называет: в журнале
    видно сегмент и оба текста, в счётчиках прогона видно, что это случилось.
    Дальше уже есть с чем идти разбираться, а не выяснять через два дня
    по остывшему state.json."""
    if (seg.get("target") or "") == want:
        return False
    print("[backend] ремонт seg#%s: текст разошёлся с записью о решении — "
          "в сегменте %r, ожидалось %r"
          % (seg.get("id"), (seg.get("target") or "")[:120], (want or "")[:120]),
          file=sys.stderr)
    return True


def _run_segment_repair(seg: dict, project: dict, model: Optional[str] = None,
                        bc_model: Optional[str] = None, tc_model: Optional[str] = None,
                        use_judge: bool = False, judge_model: Optional[str] = None,
                        judge_all: bool = False) -> dict:
    """Один заход ремонта с обязательной перепроверкой и откатом.

    `judge_all` доезжает до перепроверки: судья обязан смотреть на текст ДО
    и ПОСЛЕ по одному правилу, иначе сравниваются числа разной природы."""
    findings = _repair_findings(seg, project)
    if not findings:
        return {"ok": False, "error": "Нет находок, по которым можно чинить"}
    old_target = seg.get("target") or ""
    old_hash = _text_hash(old_target)
    had_bc = any(f["kind"] in ("term_lost", "backcheck", "judge") for f in findings)
    had_gloss = any(f["kind"] == "gloss" for f in findings)
    # Правка ради глоссария меняет формулировку, и проверить её обязан кто-то,
    # кроме самой правки. Termcheck смотрит только на целевой текст и стоит один
    # вызов — без него подстановка термина принималась бы на веру: сравнивать
    # было бы нечего, обе оценки остались бы от прежнего текста.
    had_tc = any(f["kind"] in ("term", "term_ctx") for f in findings) or had_gloss
    had_minor = any(f["kind"] == "term" and f.get("sev") == "minor" for f in findings)
    # Заход ТОЛЬКО из-за мелких замечаний — единственный случай, где у правки
    # нет ни одной измеримой цели, кроме самих этих замечаний: балл back-check
    # не пересчитывается (had_bc False), нарушать глоссарию нечего (had_gloss
    # False), а счётчик серьёзных находок и до правки был нулём. «Не стало
    # хуже» тут засчитало бы успехом любой переписанный текст.
    # Поэтому спрашиваем по существу: ушла ли хоть одна из тех находок, ради
    # которых зашли, — и сверяем их ПОИМЕННО, а не по количеству. Счёт врал бы
    # в обе стороны: termcheck на переписанном тексте почти всегда добавляет
    # свою придирку, и верная правка «EPT → EPTB» откатывалась бы по «1 → 1»,
    # а сегмент клеймился бы `tried`, то есть чинить его больше не пришли бы
    # никогда. Рост общего числа мелочи при этом всё равно откат: полторы
    # новые придирки взамен одной старой — не работа.
    only_minor = had_minor and not had_bc and not had_gloss and all(
        f.get("sev") == "minor" for f in findings if f["kind"] == "term")
    minor_targets = {_norm_key(f["replace"][0]) for f in findings
                     if f["kind"] == "term" and f.get("sev") == "minor"
                     and (f.get("replace") or [""])[0]}
    # То же самое для СЕРЬЁЗНЫХ замечаний. Раньше их сверяли по количеству,
    # и это врало ровно так же, как врал бы счёт у мелких: termcheck на
    # переписанном тексте почти всегда добавляет свою придирку, а придирка
    # эта часто про текст, который правка не трогала. На боевом проекте так
    # откачены верные правки «an infected animal» → «a patient» и
    # «urethral narrowing» → «narrowing of the small ureter» — обе по счёту
    # «1 → 2», обе с клеймом `tried`, то есть чинить их больше не пришли бы.
    major_targets = {_norm_key(f["replace"][0]) for f in findings
                     if f["kind"] == "term" and f.get("sev") in ("critical", "major")
                     and (f.get("replace") or [""])[0]}
    # Все заказанные замечания любой тяжести. По ним меряется «сделала ли
    # правка хоть что-то из того, ради чего заходили»; `major_targets` остаётся
    # мерой для вето по баллу, где речь именно о серьёзном улучшении.
    all_targets = {_norm_key(f["replace"][0]) for f in findings
                   if f["kind"] == "term" and (f.get("replace") or [""])[0]}
    # Заход ТОЛЬКО из-за регистра — ровно тот же случай, что only_minor: ни балл,
    # ни termcheck не пересчитываются, глоссарию нарушать нечего, и «не стало
    # хуже» засчитало бы успехом любой переписанный текст. Мерить тут есть чем
    # и это бесплатно — сама находка детерминированная, поэтому спрашиваем
    # прямо: убавилось ли расхождений по регистру.
    had_free = any(f["kind"] in ("case", "script", "term_case", "dup") for f in findings)
    only_free = had_free and not had_bc and not had_tc
    # Доказательство, сильнее ли находка заверения, снимаем ДО правки: после
    # перепроверки back-check описывает уже НОВЫЙ текст, и находки на старом
    # там нет — след получился бы пустым ровно там, где он нужнее всего.
    attempt_key = _repair_attempt_key(seg)
    # Кто уже ходил на этот отпечаток. Второй заход другой моделью ДОПИСЫВАЕТ
    # её к списку, а не затирает чужую: затёртая модель снова считалась бы
    # «не пробовала», и смена моделей туда-сюда открывала бы сегмент заново.
    mdl_planned = _resolve_model(model or REPAIR_DEFAULT_MODEL)["id"]
    prev_tried = _models_tried(seg, attempt_key)
    tried_now = prev_tried if mdl_planned in prev_tried else prev_tried + [mdl_planned]
    was_confirmed = seg.get("status") == "confirmed"
    override_ev = _confirm_override(seg) if was_confirmed else []
    doc_skip = _doc_flagged(seg, project)
    before = _repair_scores(seg, project, doc_skip)
    # Проверки старого текста сохраняем целиком: при откате их надо вернуть,
    # иначе сегмент окажется «непроверенным» и человек заплатит за прогон снова.
    bc_before = json.loads(json.dumps(seg.get("backcheck"))) if seg.get("backcheck") else None
    tc_before = json.loads(json.dumps(seg.get("termcheck"))) if seg.get("termcheck") else None

    new_target = _openai_repair(seg, project, findings, model)
    if not new_target:
        return {"ok": False, "error": "Модель не вернула исправленный текст"}
    # Начертание правим САМИ и бесплатно, до всякой оценки. Модель приносит
    # верный термин не в том регистре, и правка откатывалась расхождением
    # по регистру — 18 заходов на боевом проекте, каждый оплаченный. Сочинять
    # тут нечего: слова и порядок те же.
    new_target = _case_fit(seg.get("source") or "", new_target)
    # Пробелы схлопываем, регистр — НЕТ. `_norm_key` приводит текст к нижнему
    # регистру, и правка, у которой всё отличие в заглавной букве, читалась бы
    # как «модель не нашла, что менять»: единственная находка, ради которой
    # заходили, оставалась бы неисправленной, а сегмент получал бы клеймо
    # `tried` и больше не чинился НИКОГДА.
    if _same_words(new_target, old_target):
        seg["repair"] = {"applied": False, "reason": "Модель не нашла, что менять",
                         "source_hash": old_hash, "attemptKey": attempt_key,
                         "model": mdl_planned, "triedModels": tried_now,
                         "issues": [f["text"] for f in findings],
                         "at": datetime.now().strftime("%Y-%m-%d %H:%M")}
        return {"ok": True, "applied": False, "repair": seg["repair"]}

    # Перепроверяем ровно теми проверками, которые ругались. Лишних не гоняем.
    seg["target"] = new_target
    # harvest=False: вердикт «лучше/хуже» ещё не вынесен, и repair.applied не
    # проставлен. Собирать терминологию с текста, который через две строки
    # может быть откачен, — значит закрепить в глоссарии отменённую правку.
    if had_bc:
        # Судья с ОБЕИХ сторон либо ни с одной. Прежний балл мог сложиться
        # с участием судьи (JUDGE_CAP опускает, JUDGE_FLOOR_NONE поднимает),
        # а перепроверка идёт с `use_judge` из прогона — тумблер «Судья»
        # выключен по умолчанию. Тогда сравниваются два числа РАЗНОЙ природы:
        # вердикт против сырого измерения. На боевом проекте так вышло у 66
        # отмен из 176, и падение у них медианно 25 пунктов против 12
        # у остальных — то есть отменяла правку сама несимметричность.
        # Зовём судью там, где он участвовал в прежней оценке: лишний вызов
        # только на таких сегментах, зато сравнение честное.
        judge_after = use_judge or bool(bc_before and bc_before.get("judged"))
        # `judge_all` обязан доехать и сюда — по той же причине, по которой
        # заведён `judge_after`. Прежний балл в прогоне с разрешением мог
        # сложиться с участием судьи ВЫШЕ обычной зоны (JUDGE_CAP опускает
        # балл 98 до 70), а перепроверка без разрешения на сыром балле 98
        # судью не позовёт — и мы снова сравним вердикт с сырым измерением,
        # ровно ту асимметрию, из-за которой откатывались верные правки.
        _run_segment_backcheck(seg, project, bc_model, judge_after, judge_model,
                               harvest=False, judge_all=judge_all)
    if had_tc:
        _run_segment_termcheck(seg, project, tc_model, harvest=False)
    after = _repair_scores(seg, project, doc_skip)
    # Что говорит СВЕЖИЙ termcheck и ушли ли заказанные замечания. Считаем
    # один раз: этим пользуются и вето по баллу, и сверка по числу замечаний,
    # а два расчёта одного и того же однажды разойдутся.
    tc_left = {_norm_key(f.get("tgt_term")) for f in
               ((seg.get("termcheck") or {}).get("findings") or [])
               if f.get("severity") in ("critical", "major")}
    # Свежий termcheck обязателен: без него `tc_left` пуст не потому, что
    # замечания сняты, а потому, что их никто не смотрел. «Не знаю» выдать
    # за «снято» — значит принять правку, которую подтвердить нечем.
    ordered_fixed = (after["terms"] is not None and bool(major_targets)
                     and bool(major_targets - tc_left))
    # Снимок находок ОТВЕРГНУТОГО текста берём ЗДЕСЬ: при откате проверки
    # восстанавливаются раньше, чем пишется запись, и оттуда пришли бы находки
    # прежнего текста — то есть запись врала бы о том, за что откатили.
    after_findings = [{"tgt_term": f.get("tgt_term"), "severity": f.get("severity"),
                       "suggestion": f.get("suggestion")}
                      for f in ((seg.get("termcheck") or {}).get("findings") or [])
                      if f.get("severity") in TERMCHECK_ACTIONABLE][:8]

    better = True
    why = []
    # Отмена по причине, к КАЧЕСТВУ правки отношения не имеющей (перепроверка
    # не выполнилась). Такой заход не засчитывается: см. ниже про source_hash.
    infra_fail = False
    # Жёсткая находка на КАНДИДАТЕ (числа, единицы, отрицание, подмена стороны).
    # Уезжает в запись отката: по ней решают, можно ли предлагать кандидата
    # человеку к принятию. Считается, только когда back-check перепроверялся.
    hard_after = False
    # Решения, принятые ВОПРЕКИ падению балла, — в запись: без них человек
    # видит принятую правку с упавшим баллом и не понимает, почему её оставили.
    notes = []
    if had_bc and before["score"] is not None and after["score"] is not None:
        if after["score"] < before["score"]:
            # Балл back-check — доля основ ОРИГИНАЛА, вернувшихся через
            # обратный перевод (_content_recall). Значит он вознаграждает
            # КАЛЬКУ: «Erect solar rays» возвращается как «прямые солнечные
            # лучи» слово в слово и набирает почти единицу, а верное «direct
            # sunlight» возвращается синонимом и балл роняет. Termcheck заведён
            # ровно против кальки — и его правку отменяла проверка, для которой
            # калька выглядит образцовой. На боевом проекте так выброшены 111
            # верных правок: «medicine physicians» → «pulmonologists»,
            # «sanguiferous bed» → «bloodstream», «an infected animal» →
            # «the patient». Медиана падения 24 пункта, и это не эффект
            # коротких сегментов — медиана длины 23 содержательных слова.
            #
            # Поэтому балл больше не отменяет правку, которая ПОЧИСТИЛА
            # ТЕРМИНЫ. Обоснование то же, каким обоснован only_minor: у захода
            # должна быть измеримая цель, а на правке кальки балл меряет не то.
            # Ослабление узкое, и границы у него две:
            #   1) улучшение считается по БЕСПЛАТНЫМ и объективным счётчикам
            #      (серьёзные находки termcheck, нарушенные приказные термины),
            #      а не по мнению модели о собственной работе;
            #   2) ЖЁСТКАЯ находка на новом тексте (числа, единицы, отрицание,
            #      подмена стороны) остаётся вето при любом улучшении
            #      терминологии: она не зависит ни от морфологии, ни от
            #      буквализма, и размен числа на термин — не работа.
            # «Стало чище» — это и убывший счётчик, и СНЯТОЕ ЗАКАЗАННОЕ
            # замечание. Без второго правка, снявшая критичную находку
            # и получившая взамен другую («1 → 1»), не проходила ни сюда,
            # ни в сверку числом — и откатывалась баллом, который на правке
            # кальки меряет не то. Именно так стоял «an infected animal».
            terms_cleaner = (
                (before["terms"] is not None and after["terms"] is not None
                 and after["terms"] < before["terms"])
                or after["gloss"] < before["gloss"]
                or ordered_fixed)
            # Флаг читаем у записи, которую только что написала перепроверка
            # ВЫШЕ, — то есть он всегда свежий и всегда есть. У записей
            # прежних версий его нет, но сюда такие не попадают.
            hard_now = bool((seg.get("backcheck") or {}).get("hard"))
            hard_after = hard_now
            if terms_cleaner and not hard_now:
                notes.append("балл back-check упал " + str(before["score"]) + " → "
                             + str(after["score"]) + ", но правка почистила термины "
                             + "и жёстких находок нет — отменой не считаем")
            else:
                better = False
                why.append("балл back-check упал " + str(before["score"]) + " → " + str(after["score"])
                           + ("" if not hard_now else ", жёсткая находка на новом тексте"))
    if had_tc and after["terms"] is None:
        # Перепроверка не состоялась (вызов упал) — подтвердить правку нечем.
        # Откат: автоправка не заверяет сама себя, и «проверка не ответила»
        # это не «проверка сказала, что всё хорошо».
        better = False
        infra_fail = True
        why.append(REPAIR_RECHECK_FAILED)
    elif had_tc and before["terms"] is not None and after["terms"] > before["terms"]:
        # Заказ ОТ TERMCHECK меряется поимённо, а не счётом — тот же закон, что
        # у `only_minor`, и по той же причине. Спрашиваем по существу: ушли ли
        # те замечания, ради которых заходили, и не пришло ли новое на слово,
        # которое подставили МЫ. За своё отвечаем целиком; чужая, уже стоявшая
        # в тексте проблема не повод выбрасывать верную правку вместе
        # с оплаченной проверкой — её человек увидит на экране итогов.
        left = tc_left
        left_all = {_norm_key(f.get("tgt_term")) for f in
                    ((seg.get("termcheck") or {}).get("findings") or [])
                    if f.get("severity") in TERMCHECK_ACTIONABLE}
        # Подставленное разбирает отдельная проверка ниже — она работает
        # при любом движении счётчиков, а не только когда их стало больше.
        if not major_targets:
            # Заказа от termcheck не было (чинили по back-check или глоссарию) —
            # сравнивать по именам нечего, остаётся счёт.
            better = False
            why.append("замечаний по терминам стало больше "
                       + str(before["terms"]) + " → " + str(after["terms"]))
        elif not (all_targets - left_all):
            # Сверяем ВСЕ заказанные замечания, а не только серьёзные: правка,
            # снявшая одно из двух, — это движение вперёд, и выбрасывать её
            # целиком значит терять оплаченную работу. Оставшееся замечание
            # никуда не девается, оно так и висит находкой, а текст изменился —
            # значит отпечаток захода другой и следующий заход разрешён.
            better = False
            why.append("ни одно из заказанных замечаний не снято: "
                       + ", ".join(sorted(all_targets)[:3]))
        else:
            notes.append("замечаний по терминам стало больше "
                         + str(before["terms"]) + " → " + str(after["terms"])
                         + ", но заказанные сняты, а новые — не про подставленное")
    # Забракованный ПОДСТАВЛЕННЫЙ термин откатывает ВСЕГДА, при любой тяжести
    # и независимо от того, как двинулись счётчики. За слово, которое вписали
    # МЫ, отвечаем мы целиком; чужая, уже стоявшая в тексте проблема — не повод
    # выбрасывать верную правку вместе с оплаченной проверкой, её человек
    # увидит на экране итогов.
    # Проверка стоит ОТДЕЛЬНО, а не внутри ветки «сравнивать не с чем»: когда
    # число замечаний не изменилось (сняли одно, пришло другое), ни одна ветка
    # выше не срабатывает — и забракованная подстановка проходила молча.
    # Оставь тут только critical/major — и мелкое замечание на свежеподставленном
    # термине правку не отменит, зато следующий прогон сделает из него повод для
    # ремонта, ремонт полезет менять утверждённый термин, `gloss` вырастет
    # и правку откатят: платный вызов с заранее известным исходом.
    if better and had_tc and after["terms"] is not None:
        # Подставленное — это и совет из `use`, и замена из `replace`
        # (находка по забракованному слову несёт её именно там): за слово,
        # которое вписали мы, отвечаем мы, каким бы полем оно ни приехало.
        inserted = ({_norm_key(f["use"]) for f in findings if f.get("use")}
                    | {_norm_key(f["replace"][1]) for f in findings
                       if f["kind"] == "term_ctx" and f.get("replace")
                       and f["replace"][1]})
        hit = next((f for f in ((seg.get("termcheck") or {}).get("findings") or [])
                    if f.get("severity") in TERMCHECK_ACTIONABLE
                    and _norm_key(f.get("tgt_term")) in inserted), None) if inserted else None
        if hit:
            better = False
            why.append("подставленный термин «%s» забракован проверкой"
                       % hit.get("tgt_term", ""))
    # Мелкие замечания сверяем ТОЛЬКО у захода, где кроме них ничего не было.
    # Смешанный заход (мелочь плюс глоссарий или серьёзная находка) уже измерен
    # своими счётчиками; откатывать верную подстановку термина из-за побочной
    # придирки значит отнять автоматизацию там, где она работала.
    # Отдельной ветки на «перепроверка не ответила» нет намеренно: terms и
    # terms_minor становятся None вместе (один признак свежести termcheck),
    # и этот случай уже разобран выше — вторая ветка добавила бы к откату
    # вторую причину о том же самом.
    if only_minor and after["terms_minor"] is not None and minor_targets:
        left = {_norm_key(f.get("tgt_term")) for f in
                ((seg.get("termcheck") or {}).get("findings") or [])
                if f.get("severity") == "minor"}
        if not (minor_targets - left):
            better = False
            why.append("ни одно из мелких замечаний не снято: "
                       + ", ".join(sorted(minor_targets)[:3]))
        elif (before["terms_minor"] is not None
                and after["terms_minor"] > before["terms_minor"]):
            better = False
            why.append("мелких замечаний по терминам стало больше "
                       + str(before["terms_minor"]) + " → " + str(after["terms_minor"]))

    # Глоссарий сверяем независимо от того, из-за него ли чинили: правка одного
    # места не должна выбивать утверждённый термин в другом. Проверка бесплатна,
    # поэтому идёт всегда.
    if after["gloss"] > before["gloss"]:
        better = False
        why.append("нарушено утверждённых терминов больше "
                   + str(before["gloss"]) + " → " + str(after["gloss"]))
    # И по той же причине — потерянные приказные термины. Сверяется только
    # когда известны обе стороны: сравнивать с «неизвестно» значит откатывать
    # правку из-за того, что проверки прежнего текста не было.
    if (before["terms_lost"] is not None and after["terms_lost"] is not None
            and after["terms_lost"] > before["terms_lost"]):
        better = False
        why.append("приказных терминов не пережило обратный перевод больше "
                   + str(before["terms_lost"]) + " → " + str(after["terms_lost"]))
    # Регистр — по тем же двум причинам, что глоссарий: проверка бесплатна,
    # а правка ради термина не должна попутно ронять заглавную в заголовке.
    if after["case"] > before["case"]:
        better = False
        why.append("расхождений по регистру стало больше "
                   + str(before["case"]) + " → " + str(after["case"]))
    if after["script"] > before["script"]:
        better = False
        why.append("букв чужого письма стало больше "
                   + str(before["script"]) + " → " + str(after["script"]))
    if after.get("doc", 0) > before.get("doc", 0):
        better = False
        why.append("пар терм-листа документа нарушено больше "
                   + str(before.get("doc", 0)) + " → " + str(after.get("doc", 0)))
    # Повторы — та же бесплатная и безусловная сверка. Ремонт обязан ЗАМЕНЯТЬ
    # неверный термин, а не дописывать верный рядом с ним.
    if after["dup"] > before["dup"]:
        better = False
        why.append("повторов в тексте стало больше "
                   + str(before["dup"]) + " → " + str(after["dup"])
                   + " — термин дописан вместо замены")
    if after["term_case"] > before["term_case"]:
        better = False
        why.append("приказных терминов не в начертании оригинала стало больше "
                   + str(before["term_case"]) + " → " + str(after["term_case"]))
    # Заход ТОЛЬКО по бесплатным находкам меряется ими же и по существу:
    # ни балл, ни termcheck не пересчитывались, глоссарию нарушать нечего,
    # и «не стало хуже» засчитало бы успехом любой переписанный текст.
    # Считаем их вместе: размен регистра на кириллицу — не работа.
    _free = lambda d: d["case"] + d["script"] + d["term_case"] + d["dup"] + d["self_dup"]
    if only_free and _free(after) >= _free(before):
        better = False
        why.append("правка не сняла ни регистра, ни чужого письма, ни самоповтора: "
                   + str(_free(before)) + " → " + str(_free(after)))

    mdl_id = _resolve_model(model or REPAIR_DEFAULT_MODEL)["id"]
    if not better:
        # Откат вместе с проверками: они уже пересчитаны под отвергнутый текст,
        # возвращаем те, что относились к прежнему переводу
        seg["target"] = old_target
        # Проверки возвращаем те, что относились к ПРЕЖНЕМУ тексту. Прежней
        # не было — снимаем совсем, а не оставляем свежую: `seg[k] = seg.pop(k)`
        # выковыривал ключ и клал обратно, то есть на восстановленном тексте
        # оставалась проверка ОТВЕРГНУТОГО варианта. На экране это находка про
        # слова, которых в сегменте нет; ремонт получал по ней повод чинить
        # сегмент ещё раз, а Medical QA — обратный перевод выброшенного текста.
        # Случай не редкий: расхождение с глоссарием (`had_gloss`) заказывает
        # termcheck сегменту, который termcheck до этого ни разу не видел.
        if had_bc:
            if bc_before:
                seg["backcheck"] = bc_before
            else:
                seg.pop("backcheck", None)
        if had_tc:
            if tc_before:
                seg["termcheck"] = tc_before
            else:
                seg.pop("termcheck", None)
        # Заход ЗАСЧИТЫВАЕТСЯ (source_hash), только если правку отвергла ОЦЕНКА:
        # тот же текст с теми же претензиями даст тот же ответ модели, и второй
        # заход — это платный вызов с заранее известным исходом (_repair_tried,
        # _repair_futile). А вот сбой перепроверки о качестве правки не говорит
        # НИЧЕГО: причина в оборванном вызове, и клеймить за неё сегмент значит
        # закрыть его от ремонта навсегда из-за чужой сетевой ошибки. Откат при
        # этом правильный — система не заверяет сама себя, — но попытка не
        # записывается, и сегмент остаётся доступен. На боевом проекте так
        # потеряны 5 верных правок (#645: «accidental» → «casual», балл
        # не падал вовсе); при плохой сети их были бы десятки.
        seg["repair"] = {"applied": False, "reason": "; ".join(why) or "не стало лучше",
                         "model": mdl_id, "candidate": new_target,
                         "issues": [f["text"] for f in findings], "before": before, "after": after,
                         # Была ли на КАНДИДАТЕ жёсткая находка (числа, единицы,
                         # отрицание, подмена стороны). Отдельным полем, а не
                         # разбором строки причины: строка у обоих исходов почти
                         # одна, и по ней их не различить — а различать надо,
                         # потому что кандидата с расхождением чисел нельзя
                         # предлагать человеку к принятию.
                         "hardAfter": bool(hard_after),
                         # Находки перепроверки ОТВЕРГНУТОГО текста. В `after`
                         # лежат только числа, и по ним нельзя отличить шум
                         # termcheck от настоящей регрессии — а именно этот
                         # вопрос и встаёт, когда разбираются, за что откатили.
                         "afterFindings": after_findings,
                         "at": datetime.now().strftime("%Y-%m-%d %H:%M")}
        # Заход НЕ засчитывается, когда правку отвергли сбоем перепроверки и
        # больше ничем ПО СУЩЕСТВУ. Упавший балл существом не считается, и это
        # не поблажка: балл — та самая мера, которую этот же код объявил
        # негодной для правки термина, а взвесить её было нечем именно потому,
        # что termcheck не ответил (`after["terms"] is None`, значит и
        # `terms_cleaner` посчитать не из чего). Считать такой заход
        # состоявшимся значит закрыть сегмент навсегда по показанию прибора,
        # которому в этом вопросе не верят, — а это самая частая форма отказа:
        # заход почти всегда смешанный.
        # Жёсткая находка на кандидате существом является всегда.
        substantive = [w for w in why
                       if w != REPAIR_RECHECK_FAILED
                       and not w.startswith("балл back-check упал")]
        if infra_fail and not substantive and not hard_after:
            # Признак для экрана: заход не состоялся, сегмент ждёт повтора.
            seg["repair"]["retryable"] = True
            # Хеш текста, на котором заход не состоялся. НЕ `source_hash`:
            # тот клеймит сегмент, а этот только уточняет подпись на экране —
            # правь текст руками, и «заход не состоялся» станет неправдой,
            # потому что текст как раз менялся.
            seg["repair"]["attemptHash"] = old_hash
        else:
            # Отпечаток ЗАХОДА пишем вместе с клеймящим хешем и по той же
            # причине: заход состоялся, вердикт вынесен. У несостоявшегося
            # (оборванная перепроверка) не пишем ни того, ни другого —
            # иначе сегмент закрыт из-за чужой сетевой ошибки.
            seg["repair"]["source_hash"] = old_hash
            seg["repair"]["attemptKey"] = attempt_key
            # Вместе с отпечатком, и только с ним: несостоявшийся заход
            # (оборванная перепроверка) не в счёт — модель не «пробовала».
            seg["repair"]["triedModels"] = tried_now
        return {"ok": True, "applied": False, "repair": seg["repair"],
                "desync": _repair_desync(seg, old_target)}

    # Заверение снимаем ГРОМКО. Человек его поставил, и если машина его
    # отменяет, она обязана предъявить, за что: какая находка, каким правилом,
    # какой текст стоял. Без этого следа отмена решения человека неотличима
    # от обычной правки, и доверять системе больше нельзя.
    if was_confirmed and override_ev:
        seg["confirmWithdrawn"] = {
            "by": seg.get("confirmedBy"), "at": seg.get("confirmedAt"),
            "withdrawnAt": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "evidence": override_ev, "was": old_target}
    seg["prevTarget"] = old_target
    seg["status"] = "review"          # заверяет человек, автоправка себя не подтверждает
    # Чинили заверенный перевод — отметка «подтвердил человек» относилась к
    # прежнему тексту. Оставить её на новом значит соврать и себе (автоодобрение
    # считает такие сегменты доказательством), и пользователю.
    seg.pop("confirmedBy", None)
    seg.pop("confirmedAt", None)
    seg.pop("confirmedRole", None)
    seg["provider"] = mdl_id
    # Принятая правка не выходит из прогона с устаревшей проверкой. Ремонт —
    # последний платный шаг конвейера, back-check и termcheck идут ДО него,
    # и правка по терминам (had_tc без had_bc) оставляла back-check от
    # прежнего текста: сегмент тут же возвращался в «переведено, но не
    # проверено», а следующий прогон покупал обратный перевод заново — после
    # каждого прогона оставался хвост (боевые #1349, #1741, #1925, #2577).
    # Освежаем только НЕДОСТАЮЩУЮ проверку и только у ПРИНЯТОЙ правки: у
    # откачённой восстановлены прежние проверки, и платить за выброшенного
    # кандидата незачем. Денег это не прибавляет — тот же вызов всё равно
    # купил бы следующий прогон, просто теперь сегмент закрывается этим.
    # Судья — по правилу judge_after (см. had_bc выше): прежний балл мог
    # сложиться с его участием. Сбой освежения правку НЕ откатывает: вердикт
    # уже вынесен теми проверками, которые ругались, а проверка просто
    # остаётся устаревшей — ровно как было до этой правки. Идёт ДО записи
    # repair: attemptKey обязан видеть находки свежих проверок.
    # Каждая проверка в СВОЁМ try: вызовы независимы, и сбой back-check
    # не повод не делать дешёвый termcheck (заход только по бесплатным
    # находкам освежает обе) — иначе один обрыв сети возвращал бы тот самый
    # платный хвост, ради которого освежение заведено.
    if not had_bc:
        try:
            _run_segment_backcheck(seg, project, bc_model,
                                   use_judge or bool(bc_before and bc_before.get("judged")),
                                   judge_model, harvest=False, judge_all=judge_all)
        except Exception as e:
            print(f"[backend] repair: освежение back-check после принятой правки не удалось: {e}",
                  file=sys.stderr)
    if not had_tc:
        try:
            _run_segment_termcheck(seg, project, tc_model, harvest=False)
        except Exception as e:
            print(f"[backend] repair: освежение termcheck после принятой правки не удалось: {e}",
                  file=sys.stderr)
    seg["repair"] = {"applied": True, "from": old_target, "source_hash": _text_hash(new_target),
                     # По НОВОМУ тексту: находки пересчитаны, и если их не
                     # осталось, повторять нечего; появились новые — заход
                     # будет другим и разрешён.
                     "attemptKey": _repair_attempt_key(seg),
                     # Список заново: отпечаток описывает НОВЫЙ текст, и на
                     # него ходила только модель, которая его написала.
                     "model": mdl_id, "triedModels": [mdl_id],
                     "issues": [f["text"] for f in findings],
                     "before": before, "after": after,
                     # Чем правка обязана решению, принятому вопреки падению
                     # балла. Без этого на экране стоит принятая правка
                     # с упавшим баллом и ничем не объяснённая.
                     "notes": notes,
                     "at": datetime.now().strftime("%Y-%m-%d %H:%M")}
    return {"ok": True, "applied": True, "repair": seg["repair"], "target": new_target,
            "desync": _repair_desync(seg, new_target)}


class TermCaseRequest(BaseModel):
    dry_run: bool = True                      # по умолчанию только считаем
    segment_ids: Optional[List[int]] = None   # None — весь проект
    include_confirmed: bool = False            # трогать заверенное человеком


@app.post("/api/projects/{pid}/term-case")
def term_case(pid: int, req: TermCaseRequest = TermCaseRequest()):
    """Приводит начертание приказных терминов к оригиналу. БЕЗ вызова модели.

    Запись глоссария хранит одно начертание, а мест, куда она встаёт, много:
    «Туберкулема» стоит и заголовком, и посреди фразы. Как термин написан
    в ЭТОМ сегменте оригинала — так он пишется и в переводе (`_case_like`).
    Меняются только заглавные и строчные: слова, их порядок и знаки те же,
    поэтому сочинять тут нечего и модель не нужна.

    `dry_run` по умолчанию: сначала показываем, что изменится, — то же правило,
    что у выноса глоссария и пересчёта баллов. Заверенное человеком не трогаем
    без явного разрешения: правка снимет с него отметку (`_replace_target`)."""
    _guard_project_write(pid)
    project = get_project(pid)
    ids = set(req.segment_ids) if req.segment_ids is not None else None
    changed, skipped_confirmed, samples = [], [], []
    for seg in project["segments"]:
        if ids is not None and seg["id"] not in ids:
            continue
        new_text, moves = _term_case_fix(seg, project)
        if not moves:
            continue
        if seg.get("status") == "confirmed" and not req.include_confirmed:
            skipped_confirmed.append(seg["id"])
            continue
        changed.append(seg["id"])
        if len(samples) < 20:
            samples.append({"id": seg["id"],
                            "fixed": [{"was": a, "now": b} for a, b in moves]})
        if not req.dry_run:
            _replace_target(seg, new_text, seg.get("provider") or "", "TERM_CASE")
    if changed and not req.dry_run:
        save_state(STATE)
    # Список id нужен браузеру: подтянуть ПРАВЛЕННЫЕ сегменты, а не весь проект
    # на пять мегабайт ради десятка изменившихся строк.
    return {"ok": True, "dryRun": req.dry_run, "segments": len(changed),
            "ids": changed, "skippedConfirmed": skipped_confirmed,
            "samples": samples}


# ─── Ревизия: единственная проверка, которая читает ПАРУ целиком ─────────────
# Все прежние проверки задают узкий вопрос и потому слепы к целому классу
# дефектов. Back-check спрашивает «пережил ли смысл круг» и меряет долю основ
# ОРИГИНАЛА, вернувшихся через обратный перевод, — то есть вознаграждает кальку
# и не видит английского вовсе (судье уходят оригинал и обратный перевод,
# а не сам перевод). Termcheck смотрит ТОЛЬКО на терминологию: в его промпте
# прямо стоит «DO NOT flag sentence structure» и «Be conservative», а формат
# ответа требует `tgt_term` — кусок текста, — поэтому ошибку, которую нельзя
# ткнуть пальцем в одно слово, ему некуда записать. Арбитр видит пару, но
# спрашивают его про СПИСОК приказных терминов, а не про перевод.
#
# Отсюда дефекты, доезжавшие до готового текста при чистых проверках:
# «conduct differential diagnosis and tests in addition to tuberculosis»
# (синтаксис, а не термин), «Artificial pneumothorax treatment is closed»
# (фраза целиком), «Анамнез жизни → Social history» (другой раздел, при этом
# совершенно нормальный английский термин). Балл у таких сегментов высокий
# ровно потому, что они буквальны.
#
# Ревизор задаёт недостающий вопрос: прочитай оригинал и перевод рядом
# и скажи, годится ли это. И возвращает он не претензию, а ГОТОВЫЙ текст —
# в этом всё устройство шага:
#   • не нужен отдельный вызов ремонта (модель уже написала исправленный
#     вариант), а значит правку не отменяет балл back-check: вето живёт
#     в `_run_segment_repair`, куда мы просто не заходим;
#   • подстановка становится ПЯТОЙ командой, меняющей текст без вызова
#     модели, — рядом с `/glossary/revert-repairs`, `/term-case`,
#     `/repair/accept` и `/term-context/apply`, и по той же причине: она
#     ничего не сочиняет.
#
# Что вместо балла. Кандидат проверяется тем, что БЕСПЛАТНО и ОБЪЕКТИВНО:
# числа, единицы, отрицание и подмена стороны (`deterministic_issues` —
# считаются из пары без модели и без знания языка), нарушенные приказные
# термины, регистр, чужое письмо, самоповтор. Мера сменилась, а не исчезла:
# негодный измеритель заменён на набор годных.
REVIEW_VERSION = os.environ.get("REVIEW_VERSION", "1")
REVIEW_DEFAULT_MODEL = os.environ.get("REVIEW_MODEL", "gpt-5.6-terra")
# Оценка, при которой И НИЖЕ правка применяется. Семёрка — не круглое число
# наугад: на боевой выборке 9-10 стоит у переводов, которые править нечего,
# 7.5 — у спора о приказном термине (его решает человек, а не мы), а 4-7 —
# у настоящих дефектов. Из окружения, потому что это цена решения: порог выше
# означает больше автоматических правок и больше риска.
try:
    REVIEW_APPLY_MAX = float(os.environ.get("REVIEW_APPLY_MAX", "7"))
except ValueError:
    # Мусор в переменной окружения не должен ронять ИМПОРТ модуля: сервис
    # тогда не стартует вовсе, и починить это можно только с консоли сервера.
    print("[backend] REVIEW_APPLY_MAX: не число, беру 7", file=sys.stderr)
    REVIEW_APPLY_MAX = 7.0
# «7», а не «7.0»: порог уходит в промпт текстом, и дробный хвост там читается
# как требование точности, которой у оценки нет.
REVIEW_APPLY_LABEL = ("%g" % REVIEW_APPLY_MAX)
# Что сверяем у кандидата. Ровно бесплатные ключи `_repair_scores` — считает
# их ОДНА функция на всех, потому что второй расчёт того же однажды разойдётся
# с первым. Балла (`score`) и платных проверок (`terms`, `terms_lost`) здесь
# нет намеренно: первый меряет не то, вторые потребовали бы вызовов модели,
# ради отказа от которых шаг и заведён.
REVIEW_FREE_KEYS = ("gloss", "case", "script", "dup", "self_dup", "term_case", "doc")
# Оценка, ниже которой сегмент зовёт человека, даже если правки не было.
# Отдельно от REVIEW_APPLY_MAX намеренно: тот отвечает на вопрос «когда машина
# правит сама», а этот — «когда звать человека», и двигают их по разным
# причинам. Связав их, мы бы переклассифицировали корзину при каждой правке
# политики применения.
REVIEW_FLAG_SCORE = float(os.environ.get("REVIEW_FLAG_SCORE", "5"))
# Оценка, от которой свежий вердикт ревизии РУЧАЕТСЯ за перевод и снимает
# претензии слепых измерителей (`_review_vouches`). Девятка — из той же
# боевой выборки, что и семёрка выше: 9–10 стоит у переводов, которые править
# нечего. Отдельно от REVIEW_FLAG_SCORE: «звать человека» и «снять с человека»
# — разные решения с разной ценой ошибки.
REVIEW_VOUCH_SCORE = float(os.environ.get("REVIEW_VOUCH_SCORE", "9"))
# Коды решений ревизии. Русская фраза рядом с ними — для человека, а разбор
# идёт ПО КОДУ (закон корзин CLEAN_*): подстрока ломается от правки
# формулировки, а пересчёт условий на показе врёт ещё хуже — он не знает,
# что решил прогон. Так и вышло у первой версии корзины: сегмент с готовым
# кандидатом, прошедшим все сверки, показывался как «сверка не пустила».
REVIEW_OK = "ok"            # перевод годится, варианта модель не дала
REVIEW_ABOVE = "above"      # вариант есть, но оценка выше порога — не дефект
REVIEW_SUSPECT = "suspect"  # повреждён сам оригинал, чинить догадкой нельзя
REVIEW_VETOED = "veto"      # кандидат не прошёл объективные сверки
# Вердикты, записанные ДО появления кода, читаются по литералу причины.
# Строки точные: их писал наш код, они лежат в боевых данных, и
# переформулировать их задним числом нельзя — тот же закон, что
# у `REPAIR_RECHECK_FAILED`. Без этого 155 вердиктов первого прогона стали бы
# для корзины невидимы навсегда: они свежие, значит ни один прогон их больше
# не перезапишет и кода им не проставит.
_REVIEW_CODE_LEGACY = {
    "нет варианта": REVIEW_OK,                      # прежняя формулировка
    "перевод годится, править нечего": REVIEW_OK,
    "оценка выше порога": REVIEW_ABOVE,
    "оригинал под подозрением": REVIEW_SUSPECT,
    "не прошёл сверку": REVIEW_VETOED,
}


def _review_code(rv: dict) -> Optional[str]:
    """Код решения ревизии. Одно место на всех, кто его спрашивает: разойдись
    чтение старых записей по копиям — корзины показывали бы разное."""
    return rv.get("code") or _REVIEW_CODE_LEGACY.get(rv.get("skipped") or "")


REVIEW_VETO_LABELS = {
    "gloss": "нарушено приказных терминов больше",
    "case": "расхождений по регистру больше",
    "script": "букв чужого письма больше",
    "dup": "повторов больше",
    "self_dup": "самоповторов больше",
    "term_case": "приказных терминов не в начертании оригинала больше",
    "hard": "расхождение чисел, единиц или отрицания",
}


def _review_system(domain: dict, src_lang: str, tgt_lang: str, style: str = "") -> str:
    """Промпт ревизора. Отдельно от вызова — чтобы его гонял тест настоящим
    кодом: от формулировки зависит, что попадёт в текст клиента."""
    return (
        "Ты — редактор перевода, специализация: " + domain["label"].lower() + ". "
        "Тебе дают сегмент документа (язык: " + src_lang + "), его перевод "
        "на " + tgt_lang + " и соседние сегменты как обстановку.\n\n"
        "Оцени перевод по шкале 0-10. Если оценка "
        + REVIEW_APPLY_LABEL + " или ниже — верни ИСПРАВЛЕННЫЙ ПЕРЕВОД ЦЕЛИКОМ.\n\n"
        "Смотри на перевод как целое, а не на отдельные термины: смысл фразы, "
        "синтаксис, естественность " + tgt_lang + ", повторы, обрывки, "
        "нечитаемые места.\n"
        "НЕ снижай оценку за: синонимы, порядок слов, разное написание одного "
        "слова, а также за профессиональную нормализацию — когда вместо кальки "
        "оригинала стоит принятый в " + tgt_lang + " оборот. Это верный перевод.\n"
        "Утверждённые термины, если они даны списком, менять НЕЛЬЗЯ: они "
        "согласованы с заказчиком. Считаешь такой термин неверным — скажи это "
        "в issues и оставь его в fixed как есть.\n"
        "Числа, единицы, даты и отрицания переноси из оригинала точно.\n"
        "Соседние сегменты — только обстановка, их перевод не оценивай "
        "и в fixed не включай.\n\n"
        "Если ПОВРЕЖДЁН САМ ОРИГИНАЛ (обрывок, ошибка распознавания, "
        "бессвязная фраза) — поставь source_suspect: true и не чини перевод "
        "догадкой: пусть это увидит человек.\n\n"
        + (style + "\n" if style else "") +
        "Верни ТОЛЬКО JSON, без пояснений:\n"
        '{"score": 0-10, "source_suspect": false, '
        '"issues": ["короткая фраза на ' + _explain_lang_name() + '"], '
        '"fixed": "исправленный перевод целиком на ' + tgt_lang + '"}\n'
        'Перевод годится — верни {"score": N} и НИЧЕГО больше: ни issues, '
        "ни fixed."
    )


def _review_stale(seg: dict) -> bool:
    """Вердикт относится к другому тексту или вынесен по другим вопросам.
    Та же производная, что у back-check и termcheck, плюс версия промпта:
    меняешь вопрос — поднимай REVIEW_VERSION, иначе новый вопрос не задаётся
    (ровно та беда, что была у MEANING_VERSION)."""
    rv = seg.get("review") or {}
    if not rv or str(rv.get("v")) != str(REVIEW_VERSION):
        return True
    # Источник — вторая сторона пары, и вердикт описывает ИМЕННО ПАРУ.
    # У записей прежней версии поля нет: считаем их устаревшими по одному
    # только тексту перевода, как и раньше.
    # Стайл-шит сменили ПОСЛЕ вердикта (`set_project_style` ставит метку
    # на месте: `_review_stale` проекта не имеет, а отпечаток стиля живёт
    # на проекте). Ревизор перечитает под новые правила — это платно,
    # и число названо в ответе на смену.
    if rv.get("styleStale"):
        return True
    src_h = rv.get("source_hash")
    if src_h is not None and src_h != _text_hash((seg.get("source") or "").strip()):
        return True
    return _check_stale(rv, seg.get("target") or "")


def _review_wrote(seg: dict) -> bool:
    """Нынешний текст сегмента написала РЕВИЗИЯ: вердикт применён, и его хеш
    — хеш этого текста.

    Не `_review_stale`: та отвечает «свеж ли ВЕРДИКТ» и включает
    REVIEW_VERSION и source_hash. Подъём версии промпта разом «устарел» бы все
    применённые ревизии, и судейские находки вернулись бы в ремонт на всех
    переписанных сегментах — массовый платный откат от одной правки промпта."""
    rv = seg.get("review") or {}
    if not rv.get("applied") or not rv.get("target_hash"):
        return False
    return rv["target_hash"] == _text_hash((seg.get("target") or "").strip())


def _review_vouches(seg: dict) -> bool:
    """Свежий вердикт ревизии ручается за ЭТОТ перевод.

    Ревизия — единственная проверка, читающая пару целиком; back-check меряет
    долю основ оригинала, вернувшихся через обратный перевод (вознаграждает
    кальку, роняет верный синоним), termcheck смотрит только на перевод.
    Оценка ≥ REVIEW_VOUCH_SCORE на том же тексте и по тем же вопросам — прямое
    чтение против косвенной меры, и оно сильнее: тот же закон, по которому
    вердикт арбитра снимает `term_lost` (`_arbiter_settled`). На боевом
    проекте так стояли 30 сегментов «оценка ниже порога» с ревизией 9–10.

    Чего ручательство не переживает — не список исключений, а само правило
    «мнение снимает только мнение»:
      • объективная находка на паре (`_confirm_override`: числа, единицы,
        отрицание, сторона) — тот класс, что сильнее заверения человека,
        и мнение модели его не отменяет ни в какую сторону;
      • вердикт судьи major/critical — два оплаченных мнения расходятся,
        решает человек;
      • подозрение на оригинал, применённая или откачённая правка — там
        вердикт не о годности нынешнего текста.
    Донором глоссария сегмент от этого не становится (`_machine_clean`
    ревизию не спрашивает), заверением человека — тем более: `confirmedBy`
    машина не пишет никогда."""
    rv = seg.get("review") or {}
    if not rv or _review_stale(seg):
        return False
    if rv.get("applied") or rv.get("undone") or rv.get("sourceSuspect"):
        return False
    if _review_code(rv) in (REVIEW_VETOED, REVIEW_SUSPECT):
        return False
    sc = rv.get("score")
    if sc is None or sc < REVIEW_VOUCH_SCORE:
        return False
    if _confirm_override(seg):
        return False
    target = seg.get("target") or ""
    bc = seg.get("backcheck") or {}
    if bc and not _check_stale(bc, target) \
            and (bc.get("judge") or {}).get("severity") in ("major", "critical"):
        return False
    # Серьёзная находка termcheck (уровни TERMCHECK_DISPUTING — те же, что
    # вправе усомниться в решении человека) держит и здесь: промпт ревизора
    # прямо уводит его от терминов («смотри на перевод как целое»), и девятка
    # про подмену понятия не говорит ничего — «rear cyclitis» читается гладко.
    # Снимается только minor: там termcheck шумит больше всего.
    tc = seg.get("termcheck") or {}
    if tc and not _check_stale(tc, target) and any(
            f.get("severity") in TERMCHECK_DISPUTING for f in (tc.get("findings") or [])):
        return False
    return True


def _openai_review(seg: dict, project: dict, prev_src: str, next_src: str,
                   terms: list, model: Optional[str] = None) -> Optional[dict]:
    """Один вызов на сегмент. None — вызов не удался (сегмент не трогаем)."""
    import openai
    mdl = _resolve_model(model or REVIEW_DEFAULT_MODEL)
    dom = _resolve_domain(project.get("domain"))
    src_lang, tgt_lang = project.get("src", "RU"), project.get("tgt", "EN")
    body = ("[сегмент ДО] " + (prev_src or "—") + NL +
            ">>> [этот сегмент] " + (seg.get("source") or "") + NL +
            "[сегмент ПОСЛЕ] " + (next_src or "—") + NL + NL +
            "Перевод этого сегмента (" + tgt_lang + "): " + (seg.get("target") or ""))
    if terms:
        body += (NL + NL + "Утверждённые термины: "
                 + "; ".join((t.get("src") or "") + " → " + (t.get("tgt") or "")
                             for t in terms[:20]))
    doc = _doc_hits(seg.get("source") or "", project, terms)
    if doc:
        body += (NL + "Терм-лист документа (согласован машиной, не приказ; при споре "
                 "сильнее утверждённые термины): "
                 + "; ".join(h["src"] + " → " + h["tgt"] for h in doc))
    client = openai.OpenAI(api_key=os.environ.get("OPENAI_API_KEY"), timeout=90, max_retries=1)
    extra = ({"max_completion_tokens": 2048} if mdl["api"] == "modern"
             else {"max_tokens": 900, "temperature": 0})
    try:
        resp = client.chat.completions.create(
            model=mdl["id"],
            messages=[{"role": "system", "content": _review_system(dom, src_lang, tgt_lang,
                                                                  _style_block(project))},
                      {"role": "user", "content": body}],
            **extra,
        )
        _note_usage("review", mdl["id"], resp)
        # Цена ЭТОГО вызова — для отчёта запроса. Разностью процессного
        # счётчика её брать нельзя: он общий на процесс, и фоновый прогон
        # добавил бы в отчёт ревизии свой расход.
        u = getattr(resp, "usage", None)
        cost = _usage_cost(mdl["id"], _usage_field(u, "prompt_tokens"),
                           _usage_field(u, "completion_tokens")) or 0.0
        raw = (resp.choices[0].message.content or "").strip()
        m = re.search(r"\{.*\}", raw, re.S)
        if not m:
            return None
        data = json.loads(m.group(0))
        score = float(data.get("score"))
    except Exception as e:
        print(f"[backend] ревизия seg#{seg.get('id')}: {e}", file=sys.stderr)
        return None
    return {"score": max(0.0, min(10.0, score)),
            "issues": [str(x).strip() for x in (data.get("issues") or [])
                       if str(x).strip()][:6],
            "fixed": (data.get("fixed") or "").strip(),
            "source_suspect": bool(data.get("source_suspect")),
            "model": mdl["id"], "cost": cost}


def _review_veto(seg: dict, project: Optional[dict], candidate: str) -> list:
    """Почему кандидата НЕЛЬЗЯ поставить. Пусто — можно.

    Всё здесь бесплатно и объективно, и сравнивается «до и после», а не
    «кандидат чист». Абсолютная проверка отвергала бы верную правку термина
    в сегменте, где числа разошлись ещё в оригинале, — а такие есть, ради них
    и заведён `source_suspect`. Ухудшил — не ставим; унаследовал чужую
    проблему — не наше дело.

    Балл back-check сюда НЕ входит, и это главное свойство шага: он меряет
    долю основ оригинала, вернувшихся через обратный перевод, то есть
    вознаграждает кальку — ровно то, что ревизор и убирает. На боевом проекте
    этот балл выбросил 111 верных правок termcheck."""
    doc_skip = _doc_flagged(seg, project)
    before = _repair_scores(seg, project, doc_skip)
    probe = dict(seg)
    probe["target"] = candidate
    after = _repair_scores(probe, project, doc_skip)
    bad = [k for k in REVIEW_FREE_KEYS if (after.get(k) or 0) > (before.get(k) or 0)]
    if checks_mod:
        # Область и пара языков — из ПРОЕКТА. Без них `rules_for` берёт
        # medical RU→EN по умолчанию, то есть на немецком проекте маркеры
        # отрицания искались бы русские, а правила направлений — чужие.
        dom = (project or {}).get("domain")
        src_lang, tgt_lang = ((project or {}).get("src") or "RU",
                              (project or {}).get("tgt") or "EN")
        try:
            issues = checks_mod.deterministic_issues(
                seg.get("source") or "", candidate, domain=dom,
                src_lang=src_lang, tgt_lang=tgt_lang) or []
            # ЛЮБАЯ объективная находка на кандидате — вето, а не «стало
            # больше, чем было». Здесь та же асимметрия, что у ремонта
            # (`hard_now` в `_run_segment_repair`): числа, единицы, отрицание
            # и сторона неразменны — сняв одно расхождение и внеся другое,
            # правка даёт счёт «1 → 1» и прошла бы сравнением. Цена
            # строгости названа честно: сегмент, где числа разошлись ещё
            # в оригинале, ревизия не починит — его чинит человек, и он же
            # видит его в `sourceSuspect`.
            if any(i.get("type") in checks_mod.OBJECTIVE_ISSUE_TYPES for i in issues):
                bad.append("hard")
        except Exception as e:                                  # pragma: no cover
            print(f"[backend] сверка кандидата ревизии seg#{seg.get('id')}: {e}",
                  file=sys.stderr)
            bad.append("hard")      # не смогли проверить — не ставим
    return bad


def _review_ask(seg: dict, project: dict, model: Optional[str] = None) -> Optional[dict]:
    """Только вызов модели — то, что можно отдать в рабочий поток.

    Мутации STATE применяет основной поток (контракт `_run_parallel`):
    рабочая функция ВОЗВРАЩАЕТ результат, а не пишет в общие структуры.
    Своё исключение ловит сама — одна упавшая пара не должна ронять порцию."""
    try:
        prev_src, next_src = _neighbours(project, seg)
        return _openai_review(seg, project, prev_src, next_src,
                              _verified_hits(seg.get("source") or "", project), model)
    except Exception as e:                                      # pragma: no cover
        print(f"[backend] ревизия seg#{seg.get('id')}: {e}", file=sys.stderr)
        return None


def _run_segment_review(seg: dict, project: dict, model: Optional[str] = None,
                        apply: bool = True, include_confirmed: bool = False,
                        res: Optional[dict] = None) -> dict:
    """Один заход ревизии: вердикт, бесплатные сверки кандидата и подстановка.

    Вердикт кладётся на сегмент по хешу текста, который в нём СЕЙЧАС стоит —
    то есть после подстановки хеш описывает уже новый текст. Так сегмент
    закрыт от повторного платного захода: спрашивать модель о тексте, который
    она сама только что написала, — это оплата второго мнения у того же
    мнения. Что оценка относилась к ПРЕЖНЕМУ тексту, видно по `applied`
    и `from`, и врать этим нельзя."""
    old = (seg.get("target") or "").strip()
    if not old:
        return {"ok": False, "error": "Нечего ревизовать: перевода нет"}
    # `res` уже посчитан рабочим потоком (`_review_ask`) — тогда сеть не
    # трогаем. Одна ветка на оба пути: разойдись они, порционный прогон
    # и одиночный заход оставляли бы сегмент в разных состояниях.
    if res is None:
        res = _review_ask(seg, project, model)
    if res is None:
        return {"ok": False, "error": "Ревизор не ответил"}
    cand = res["fixed"]
    # «Модель вернула то же самое» — не правка. Сравниваем по словам, а не
    # по _norm_key: тот приводит к нижнему регистру, и правка, у которой всё
    # отличие в заглавной букве, читалась бы как «менять нечего».
    if cand and _same_words(cand, old):
        cand = ""
    rec = {"score": res["score"], "issues": res["issues"],
           "sourceSuspect": res["source_suspect"], "model": res["model"],
           "v": REVIEW_VERSION, "applied": False, "veto": [],
           "source_hash": _text_hash((seg.get("source") or "").strip()),
           "at": datetime.now().strftime("%Y-%m-%d %H:%M")}
    veto, why = [], None
    code = None
    if not cand:
        # Самая частая причина из всех: у хорошего перевода модель не
        # возвращает `fixed` вовсе. «Нет варианта» читалось как сбой,
        # хотя это ответ «перевод годится».
        why, code = "перевод годится, править нечего", REVIEW_OK
    elif res["score"] > REVIEW_APPLY_MAX:
        # Вариант есть, а оценка выше порога: модель предлагает улучшение,
        # а не чинит дефект. Переписывать по такому поводу — тратить деньги
        # клиента на вкусовщину.
        why, code = "оценка выше порога", REVIEW_ABOVE
    elif res["source_suspect"]:
        # Оригинал под подозрением: чинить перевод догадкой нельзя, это
        # работа человека. Единственный класс, которого в системе не было
        # вообще ни в каком виде.
        why, code = "оригинал под подозрением", REVIEW_SUSPECT
    else:
        veto = _review_veto(seg, project, cand)
        if veto:
            why, code = "не прошёл сверку", REVIEW_VETOED
    if cand:
        rec["candidate"] = cand
    rec["veto"] = veto
    if why:
        rec["skipped"] = why
        rec["code"] = code
    rec["target_hash"] = _text_hash(old)
    seg["review"] = rec
    if cand and not why and apply:
        _apply_review(seg, include_confirmed)
    return {"ok": True, "applied": rec["applied"], "review": rec,
            "cost": res.get("cost") or 0.0,
            "ready": bool(cand and not why)}


def _apply_review(seg: dict, include_confirmed: bool = False) -> bool:
    """Подставить кандидата ревизии. Вызова модели тут нет: ставится то, что
    она уже написала и что прошло бесплатные объективные сверки.

    Одна ветка на одиночную и пакетную команду — разойдись они, «применить»
    и «применить все» оставляли бы сегмент в разных состояниях (тот же закон,
    что у `_apply_repair_candidate`).

    Право переписать ЗАВЕРЕННОЕ человеком проверяется ЗДЕСЬ, а не только
    в эндпоинте (инвариант 9). Стоять этажом выше оно не может: `apply`
    у `_run_segment_review` по умолчанию True, и первый же будущий вызов —
    подключение шага к прогону, кнопка на сегменте — молча снял бы
    `confirmedBy`. Правило должно жить рядом с `_replace_target`."""
    rv = seg.get("review") or {}
    if not rv.get("candidate") or rv.get("skipped") or rv.get("applied"):
        return False
    if seg.get("status") == "confirmed" and not include_confirmed:
        return False
    rv["from"] = seg.get("target") or ""
    # Провайдер — модель РЕВИЗОРА, а не прежнего переводчика: текст целиком
    # написала она. Соврав здесь, мы ломаем защиту «обратный перевод делает
    # не тот, кто писал» (`_backcheck_model` сравнивает bc["model"]
    # с seg["provider"]) — и back-check пошёл бы к автору текста за отзывом
    # о собственной работе. Так же поступает `_apply_repair_candidate`.
    _replace_target(seg, rv["candidate"], rv.get("model") or "", "REVIEW")
    # Машина не заверяет сама себя: текст переписан, значит его смотрит человек.
    seg["status"] = "review"
    rv["applied"] = True
    # Кандидата из ЗАПИСИ убираем: после применения он буква в букву равен
    # переводу, а документ проекта переписывается при КАЖДОМ сохранении —
    # хранить текст сегмента трижды (target + from + candidate) значит
    # платить весом за копию. Так же поступает `_apply_repair_candidate`.
    rv.pop("candidate", None)
    # Хеш теперь описывает НОВЫЙ текст — сегмент закрыт от повторного платного
    # захода. Спрашивать модель о тексте, который она сама только что
    # написала, значит покупать второе мнение у того же мнения. Что оценка
    # относилась к прежнему тексту, видно по `applied` и `from`.
    rv["target_hash"] = _text_hash((seg.get("target") or "").strip())
    return True


# Потолок на один запрос. Воркер uvicorn ОДИН, а шаг идёт последовательно:
# каждый сегмент — вызов модели в теле HTTP-запроса. На боевом замере вышло
# 2.5 с на сегмент, то есть 150 штук держат сервис 6 минут, а весь проект
# в 2711 сегментов держал бы его почти два часа — и всё равно оборвался бы
# по таймауту клиента, потратив деньги впустую. Ограничение честнее: остаток
# называется числом (`capped`), и запрос повторяют.
REVIEW_LIMIT_MAX = int(os.environ.get("REVIEW_LIMIT_MAX", "300"))
REVIEW_SAMPLES = ("mixed", "all")


class ReviewRequest(BaseModel):
    segment_ids: Optional[List[int]] = None   # None — выбрать самим
    model: Optional[str] = None
    limit: int = 150
    dry_run: bool = True                      # считаем и показываем, текст не трогаем
    refresh: bool = False                     # переспросить уже отвеченные
    # Применить УЖЕ ПОЛУЧЕННЫЕ вердикты, не спрашивая модель. Без этого пути
    # штатный порядок массовой команды («посмотрел сухим прогоном → применил»)
    # не работал вовсе: сухой прогон пишет вердикт на сегмент, тот перестаёт
    # быть `_review_stale`, и боевой запуск не находил НИ ОДНОГО сегмента —
    # применять было нечего, а обойти можно было только `refresh`, то есть
    # заплатив за те же вердикты второй раз. Шестая команда в системе,
    # меняющая текст без вызова модели: подставляется то, что уже написано
    # и уже прошло объективные сверки.
    apply_saved: bool = False
    # Переписывать заверенное человеком. Как у ремонта — только по явному
    # разрешению на ЭТОТ запуск.
    include_confirmed: bool = False
    # Метка отката. Задаёт её ПРОГОН — одну на всю задачу, чтобы сотни порций
    # писали в одну копию и весь прогон откатывался одной командой.
    stamp: Optional[str] = None
    # Как выбирать сегменты, когда список не задан:
    #   mixed — половина из тех, что система считает ГОТОВЫМИ (без находок,
    #           балл высокий). Иначе замер отвечает не на тот вопрос: дефекты,
    #           ради которых шаг заведён, живут именно там, где все проверки
    #           довольны, и выборка «где и так есть находки» их не увидит;
    #   all   — подряд, как идут в документе.
    sample: str = "mixed"


def _review_pick(project: dict, req: ReviewRequest) -> tuple:
    """Кого спрашиваем и сколько всего таких. Порядок важен: `mixed` чередует
    «готовые» и остальные, чтобы потолок `limit` не срезал целиком одну
    из половин.

    Второе число — весь список ДО потолка: без него усечение по `limit` было
    бы молчаливым, а «спросили 150» неотличимо от «работа кончилась». Считать
    его вторым вызовом нельзя — при `mixed` отбор зовёт `_repair_findings`
    на каждый сегмент с переводом, то есть проходит книгу целиком."""
    segs = [s for s in project["segments"] if (s.get("target") or "").strip()]
    if req.segment_ids is not None:
        ids = set(req.segment_ids)
        segs = [s for s in segs if s["id"] in ids]
    if req.apply_saved:
        # Модель не спрашиваем вовсе: берём готовые вердикты с кандидатом,
        # который ещё не поставлен. `skipped` (оценка выше порога, вето,
        # подозрительный оригинал) сюда не попадает — это решения, а не
        # очередь.
        got = [s for s in segs
               if (s.get("review") or {}).get("candidate")
               and not (s.get("review") or {}).get("applied")
               and not (s.get("review") or {}).get("skipped")
               # Откачено человеком — не предлагаем снова: он уже ответил.
               and not (s.get("review") or {}).get("undone")
               and not _review_stale(s)]
        return got, len(got)
    if not req.include_confirmed:
        # Заверенное человеком не спрашиваем ВОВСЕ, а не спрашиваем и потом
        # отбрасываем правку. Применить вердикт к такому сегменту нечем
        # (`_apply_review` откажет), а показать его человеку пока негде —
        # значит это платный совет в никуда, тот же перерасход, от которого
        # заведён `_repair_futile`. Нужен вердикт — разрешите и правку: это
        # один тумблер. Разбор состава (`_plan_step`) читает то же правило
        # и называет причину вслух, иначе план обещал бы работу, которой
        # не будет.
        segs = [s for s in segs if s.get("status") != "confirmed"]
    if not req.refresh:
        # `undone` читают ОБА: и разбор состава, и этот отбор. Разойдись
        # они — план говорит «пропустим», а шаг идёт в модель и платит.
        segs = [s for s in segs
                if _review_stale(s) and not (s.get("review") or {}).get("undone")]
    if req.sample != "mixed" or req.segment_ids is not None:
        return segs[:max(0, req.limit)], len(segs)
    ready, rest = [], []
    for s in segs:
        bc = s.get("backcheck") or {}
        sc = bc.get("score")
        # «Готов» по мнению нынешних проверок: находок нет, балл высокий.
        # Без проекта — как везде в разборе состава, ради скорости.
        (ready if (not _repair_findings(s, None) and sc is not None and sc >= 90)
         else rest).append(s)
    out = []
    for i in range(max(len(ready), len(rest))):
        for pool in (ready, rest):
            if i < len(pool) and len(out) < max(0, req.limit):
                out.append(pool[i])
    return out, len(ready) + len(rest)


@app.post("/api/projects/{pid}/review")
def review_project(pid: int, req: ReviewRequest = ReviewRequest()):
    """Ревизия: прочитать пару «оригинал — перевод» и починить, что найдётся.

    `dry_run` по умолчанию — команда меняет текст, и правила у неё общие
    с остальными массовыми: заверенное человеком не трогается без явного
    разрешения, прежнее состояние уходит копией в `data/backups/`
    и возвращается `/review/{stamp}/undo`.

    `apply_saved` — поставить уже полученные вердикты, не спрашивая модель:
    без него сухой прогон закрывал сегменты от повторного отбора, и применить
    оплаченное было нечем.

    Шаг встроен в составной прогон (`FULL_RUN_STEPS`, вторым — сразу после
    перевода), в разбор состава и в браузер; этот эндпоинт остаётся точкой
    входа для точечного запуска и для `apply_saved`.

    Платный: строка в `_PAID` обязательна, иначе шаг бесплатен для клиента."""
    if not req.apply_saved and not os.environ.get("OPENAI_API_KEY"):
        raise HTTPException(503, "Ревизия требует ключ OpenAI")
    if req.sample not in REVIEW_SAMPLES:
        # Молча считать опечатку за «all» нельзя: это ровно та выборка, против
        # которой заведён `mixed`, и человек не узнал бы, что мерил не то.
        raise HTTPException(400, "sample: " + " | ".join(REVIEW_SAMPLES))
    project = get_project(pid)
    # Охранник нужен и в СУХОМ прогоне: он не трогает текст, но пишет вердикт
    # в документ проекта (`seg["review"]` — кэш, как вердикт сверки смысла при
    # dry_run), а во время прогона проект принадлежит воркеру. Без этой строки
    # сухой запуск затирал бы работу идущего прогона молча. Стоит ПОСЛЕ
    # `get_project`: иначе чужой проект с идущим прогоном ответил бы 409
    # вместо 404, то есть подтвердил бы своё существование (инвариант 11).
    _guard_project_write(pid)
    if req.limit > REVIEW_LIMIT_MAX:
        raise HTTPException(400, f"limit больше потолка {REVIEW_LIMIT_MAX}: "
                                 "воркер один, и шаг держит его на всё время запроса")
    todo, total = _review_pick(project, req)
    # Расход считаем ПО СВОИМ ответам, а не разностью процессного счётчика:
    # `_USAGE_TOTAL` один на процесс, и когда прогон идёт фоновым потоком
    # здесь же, его расход приписался бы ревизии.
    spent = [0.0]
    answered = failed = 0
    proposed, suspect, vetoed, samples, skipped_confirmed = [], [], {}, [], []
    ready, buckets = [], {"9-10": 0, "8": 0, "5-7": 0, "0-4": 0}
    # Вызовы модели идут ПАРАЛЛЕЛЬНО (вызов — это ожидание сети, а не
    # процессора), а пишет результаты основной поток — контракт `_run_parallel`.
    # Последовательно шаг держал единственный воркер 2.5 с на сегмент: книга
    # в 2711 строк заняла бы почти два часа, и всё это время сервис недоступен
    # всем. Порядок ответов сохраняется, поэтому zip с `todo` верен.
    answers = ([None] * len(todo) if req.apply_saved
               else _run_parallel(todo, lambda s: _review_ask(s, project, req.model)))
    # ПЕРВЫЙ проход — только вердикты, текст не трогаем. Разделение не
    # косметическое: копия для отката обязана лечь на диск ДО первой правки,
    # иначе неудачная запись оставила бы в памяти изменения, откатить которые
    # уже нечем.
    for seg, ans in zip(todo, answers):
        if req.apply_saved:
            # Вердикт уже оплачен и лежит на сегменте — модель не трогаем.
            rv = seg["review"]
            answered += 1
            s = rv["score"]
            buckets["9-10" if s >= 9 else "8" if s >= 8 else "5-7" if s >= 5 else "0-4"] += 1
            proposed.append(seg["id"])
            # Сверки ПЕРЕСЧИТЫВАЮТСЯ, а не берутся из вердикта. Они бесплатны,
            # а вердикт мог быть вынесен когда угодно раньше — в том числе
            # прежними правилами (первая версия не ловила подмену стороны
            # и потерю единиц) или до правки глоссария, после которой кандидат
            # стал нарушать приказный термин. Подставлять по устаревшему
            # разрешению значит писать в документ клиента по решению,
            # которого больше нет.
            veto = _review_veto(seg, project, rv["candidate"])
            if veto:
                rv["veto"] = veto
                rv["skipped"] = "не прошёл сверку"
                rv["code"] = REVIEW_VETOED
                for k in veto:
                    vetoed[k] = vetoed.get(k, 0) + 1
            elif seg.get("status") == "confirmed" and not req.include_confirmed:
                skipped_confirmed.append(seg["id"])
            else:
                ready.append(seg)
            if len(samples) < 25:
                samples.append({"id": seg["id"], "score": s,
                                "issues": rv.get("issues") or [],
                                "willApply": seg in ready, "skipped": None, "veto": [],
                                "was": (seg.get("target") or "")[:220],
                                "now": (rv.get("candidate") or "")[:220]})
            continue
        out = _run_segment_review(seg, project, req.model, apply=False, res=ans)
        spent[0] += out.get("cost") or 0.0
        if not out.get("ok"):
            failed += 1
            continue
        answered += 1
        rv = out["review"]
        s = rv["score"]
        buckets["9-10" if s >= 9 else "8" if s >= 8 else "5-7" if s >= 5 else "0-4"] += 1
        if rv.get("sourceSuspect"):
            suspect.append(seg["id"])
        if rv.get("candidate"):
            proposed.append(seg["id"])
        for k in rv.get("veto") or []:
            vetoed[k] = vetoed.get(k, 0) + 1
        if out.get("ready"):
            # Второй рубеж у самого `_replace_target`: без разрешения
            # заверенное не переписываем. В обычном заходе сюда не доходят —
            # `_review_pick` отсекает их раньше и денег на них не тратит;
            # ветка работает для `apply_saved`, где вердикты уже лежат.
            if seg.get("status") == "confirmed" and not req.include_confirmed:
                skipped_confirmed.append(seg["id"])
            else:
                ready.append(seg)
        if len(samples) < 25 and (rv.get("candidate") or rv.get("issues")):
            samples.append({"id": seg["id"], "score": s,
                            "issues": rv.get("issues") or [],
                            "willApply": bool(out.get("ready")),
                            "skipped": rv.get("skipped"),
                            "veto": rv.get("veto") or [],
                            "was": (seg.get("target") or "")[:220],
                            "now": (rv.get("candidate") or "")[:220]})
    stamp, applied = None, 0
    if ready and not req.dry_run:
        stamp = _backup_segments("review", pid,
                                 [_repair_accept_snapshot(s) for s in ready], req.stamp)
        applied = len([s for s in ready if _apply_review(s, req.include_confirmed)])
    if answered:
        # Журнал — ДО сохранения: `_audit` пишет в STATE["audit"], и рестарт
        # между ним и следующей записью стёр бы запись о массовой перезаписи
        # текстов клиента, то есть ровно то событие, ради которого журнал есть.
        if applied:
            _audit("review.apply", project=pid, count=applied, stamp=stamp)
        _IMPACT_CACHE.pop(pid, None)
        _ANALYSIS_CACHE.pop(pid, None)
        save_state(STATE)
    return {"ok": True, "dryRun": req.dry_run, "asked": len(todo),
            "answered": answered, "failed": failed,
            # Сколько ещё ждёт своей очереди: молчаливое усечение по `limit`
            # неотличимо от «работа кончилась».
            "capped": max(0, total - len(todo)),
            "proposed": len(proposed), "wouldApply": len(ready), "applied": applied,
            "skippedConfirmed": skipped_confirmed,
            "sourceSuspect": suspect, "vetoed": vetoed, "scores": buckets,
            "cost": round(spent[0], 4),
            "stamp": stamp, "samples": samples}


@app.post("/api/projects/{pid}/review/{stamp}/undo")
def undo_review(pid: int, stamp: str):
    """Вернуть тексты до ревизии — только там, где стоит именно её текст:
    правили после, значит откат затёр бы чужую работу."""
    project = get_project(pid)
    data = _read_backup("review", pid, stamp)
    by_id = {sg["id"]: sg for sg in project["segments"]}
    restored, changed_since = [], []
    for snap in data.get("segments") or []:
        seg = by_id.get(snap["id"])
        if seg is None:
            continue
        rv = seg.get("review") or {}
        # «Стоит ли сейчас именно наш текст» спрашиваем у ХЕША, а не у копии
        # кандидата: после применения `target_hash` описывает как раз
        # поставленный текст, а сам кандидат из записи убран — он был бы
        # третьей копией перевода в документе проекта. Правили после нас —
        # хеш не сойдётся, и чужую работу мы не затрём.
        if (not rv.get("applied")
                or _text_hash((seg.get("target") or "").strip()) != rv.get("target_hash")):
            changed_since.append(seg["id"])
            continue
        seg["target"] = snap["target"]
        seg["status"] = snap["status"]
        for k in ("provider", "route", "confirmedBy", "confirmedAt", "confirmedRole", "prevTarget"):
            if snap.get(k) is None:
                seg.pop(k, None)
            else:
                seg[k] = snap[k]
        seg["repair"] = snap["repair"]
        # Вердикт НЕ стираем, а помечаем отменённым. Стереть его значит
        # сделать сегмент снова «неспрошенным»: следующий прогон заплатил бы
        # ещё раз и поставил бы ровно тот же текст обратно — то есть машина
        # переиграла бы решение человека. Тот же закон, что у `meaningKept`
        # и `_human_touched`: своё предположение машина вправе пересмотреть,
        # чужое решение — нет. Хеш описывает восстановленный текст, поэтому
        # вердикт считается свежим и сегмент в отбор не попадает.
        rv = dict(seg.get("review") or {})
        rv["applied"] = False
        rv["undone"] = {"by": _actor_id(), "at": datetime.now().strftime("%Y-%m-%d %H:%M")}
        rv["target_hash"] = _text_hash((snap["target"] or "").strip())
        rv.pop("from", None)
        seg["review"] = rv
        restored.append(seg["id"])
    if restored:
        _IMPACT_CACHE.pop(pid, None)
        _ANALYSIS_CACHE.pop(pid, None)
        save_state(STATE)
    return {"ok": True, "restored": restored, "changedSince": changed_since}


class TermContextRequest(BaseModel):
    segment_ids: Optional[List[int]] = None   # None — весь проект
    model: Optional[str] = None
    limit: int = 40
    refresh: bool = False                     # переспросить уже отвеченные
    # False — разбирать только СПОР (дорого и точечно), True — сверять ВСЕ
    # приказные термины сегмента. Второе и есть штатный шаг конвейера:
    # детерминированная проверка знает морфологию одного языка и грубо,
    # а модель отвечает на тот же вопрос на любом языке.
    all_terms: bool = False


@app.post("/api/projects/{pid}/term-context")
def term_context(pid: int, req: TermContextRequest = TermContextRequest()):
    """Арбитр спорных терминов: сегмент ДО, этот, ПОСЛЕ — и вердикт по каждому
    утверждённому термину, вокруг которого проверки разошлись с глоссарием.

    Платный, поэтому с потолком и без переспроса уже отвеченного: вердикт лежит
    на сегменте и устаревает вместе с текстом (`_term_context_stale`), как
    back-check и termcheck. Считается ОДИН вызов на сегмент, сколько бы спорных
    терминов в нём ни было."""
    if not os.environ.get("OPENAI_API_KEY"):
        raise HTTPException(503, "Арбитр требует ключ OpenAI")
    project = get_project(pid)
    ids = set(req.segment_ids) if req.segment_ids is not None else None
    # Забракованные слова сверяются в режиме полной сверки (штатный шаг):
    # точечный разбор спора — про конкретную запись глоссария, и подмешивать
    # к нему чужие вопросы значило бы платить за неспрошенное.
    stale_map = _stale_words_of(project) if req.all_terms else {}
    todo, skipped, nothing = [], 0, 0
    for sg in project["segments"]:
        if ids is not None and sg["id"] not in ids:
            continue
        stale_pending = _stale_unasked(sg, stale_map.get(sg["id"]) or [])
        if (not _term_terms_of(sg, project, disputes_only=not req.all_terms)
                and not stale_pending):
            nothing += 1
            continue
        # Свежий вердикт закрывает сегмент, только если покрывает ОХВАТ
        # запроса: сверке (all_terms) вердикт разбора спора не ответ — иначе
        # смета шага обещает сегмент, а шаг его молча пропускает, — и вопросы
        # про забракованные слова должны быть заданы все.
        tcx = sg.get("termContext") or {}
        covered = (not _term_context_stale(sg) and not stale_pending
                   and (tcx.get("all_terms") or not req.all_terms))
        if not req.refresh and covered:
            skipped += 1
            continue
        todo.append(sg)
    capped = len(todo) > max(0, req.limit)
    todo = todo[:max(0, req.limit)]
    settled, wrong, failed = [], [], []
    # Дедупликации по паре «оригинал+перевод» здесь НЕТ намеренно, и это надо
    # сказать вслух: контракт прогонов её требует. Ответ арбитра зависит
    # от СОСЕДЕЙ, а не только от пары, поэтому копировать вердикт близнецу
    # значило бы приписать ему чужую обстановку — та же уступка, что у перевода
    # («соседи берутся у первого»), но здесь она бессмысленна: на боевом
    # проекте среди 713 берущихся сегментов 711 уникальных пар, то есть вся
    # экономия — два вызова из семисот.
    #
    # А вот параллельность нужна: вызов модели это ожидание сети, и порция
    # из десяти сегментов подряд держала бы прогон вдесятеро дольше.
    def _ask(sg):
        # Остановку проверяем ЗДЕСЬ, а не в цикле: порция уходит в потоки
        # разом, и сегменты, до которых очередь ещё не дошла, должны остаться
        # неспрошенными, а не «провалившимися».
        if _job_should_stop():
            return sg, None
        try:
            return sg, _run_segment_term_context(sg, project, req.model,
                                                 disputes_only=not req.all_terms,
                                                 stale_words=stale_map.get(sg["id"]))
        except Exception as e:                                   # pragma: no cover
            # _run_parallel требует, чтобы fn ловила своё сама: одна упавшая
            # пара не должна ронять всю порцию.
            print("[backend] сверка терминов seg#%s: %s" % (sg.get("id"), e),
                  file=sys.stderr)
            return sg, {"ok": False, "error": str(e)}

    asked = 0
    for sg, r in _run_parallel(todo, _ask):
        if r is None:
            continue
        asked += 1
        if not r.get("ok"):
            failed.append(sg["id"])
            continue
        for t in r["termContext"]["terms"]:
            (settled if t.get("ok") is True else wrong).append(
                {"segment": sg["id"], "src": t.get("src"), "tgt": t.get("tgt"),
                 "use": t.get("use"), "why": t.get("why")})
    save_state(STATE)
    _ANALYSIS_CACHE.pop(pid, None)
    return {"ok": True, "asked": asked, "cachedSkipped": skipped,
            # Сколько сегментов сверять было нечем. Без этого числа отдельный
            # прогон на весь проект молча терял бы три четверти списка,
            # а полоса дошла бы до «выполнено» на четверти.
            "nothingToCheck": nothing,
            "capped": capped, "failed": failed,
            # «Передан верно» — снятая претензия: ремонт по ней больше не пойдёт.
            "settled": settled,
            # «Передан неверно» — вопрос к ЗАПИСИ глоссария, а не к строке.
            # Ремонту он не отдаётся: подстановка чужого варианта нарушила бы
            # приказ и была бы откачена, см. _repair_findings.
            "wrong": wrong}


class RepairRequest(BaseModel):
    model: Optional[str] = None
    bc_model: Optional[str] = None
    tc_model: Optional[str] = None
    use_judge: bool = False
    judge_model: Optional[str] = None
    # Судья и выше зоны — то же разрешение, что у back-check: перепроверка
    # внутри ремонта обязана идти тем же правилом, каким считали прежний балл.
    judge_all: bool = False


@app.post("/api/segments/{pid}/{sid}/repair")
def repair_segment(pid: int, sid: int, req: RepairRequest = RepairRequest()):
    _guard_project_write(pid)
    if not os.environ.get("OPENAI_API_KEY"):
        raise HTTPException(503, "Ремонт требует ключ OpenAI")
    seg = get_segment(pid, sid)
    project = get_project(pid)
    result = _run_segment_repair(seg, project, req.model, req.bc_model, req.tc_model,
                                 req.use_judge, req.judge_model, req.judge_all)
    if result.get("ok"):
        save_state(STATE)
        return result
    raise HTTPException(502, result.get("error", "Ремонт не удался"))


# Копию пачки читают-и-переписывают порции одного прогона, а «читать, слить,
# записать» без лока — это потерянные записи и общий .tmp на двоих.
# RLock, а не Lock: `_backup_segments` берёт его и внутри может позвать
# `_backup_stamp`, который берёт тот же лок, — с обычным Lock это дедлок,
# то есть намертво вставший прогон.
_BACKUP_LOCK = threading.RLock()
# Метка ходит через публичный API (`ReviewRequest.stamp`) и через params
# задачи, поэтому проверяется ТАМ ЖЕ, где пишется, а не только при чтении:
# иначе `{"stamp": "abc"}` создаёт копию, которую `_read_backup` откажется
# открывать, — то есть массовая правка текста без действующего отката.
_STAMP_RE = re.compile(r"[0-9-]{8,24}$")


def _backup_stamp(kind: str) -> str:
    """Свободная метка для НОВОЙ пачки, СРАЗУ занятая пустым файлом.

    Отдельно от записи, потому что метка нужна раньше первой правки: прогон
    идёт порциями по пять сегментов, и без общей метки каждая порция заводила
    бы свою — на книге это ~250 копий по одному-два сегмента, то есть откат,
    которым нельзя воспользоваться.

    Файл создаётся сразу, а не «когда понадобится»: между выдачей метки
    и первой правкой проходят минуты, и ручная пачка по другому проекту,
    начатая в ту же секунду, заняла бы имя — прогон умер бы на первой же
    правке с 500. Пустую копию убирает `_backup_drop_empty`."""
    with _BACKUP_LOCK:
        PURGE_DIR.mkdir(parents=True, exist_ok=True)
        base = datetime.now().strftime("%Y%m%d-%H%M%S")
        stamp, n = base, 1
        while (PURGE_DIR / (kind + "-" + stamp + ".json")).exists():
            stamp = base + "-" + str(n)
            n += 1
        (PURGE_DIR / (kind + "-" + stamp + ".json")).write_text(
            json.dumps({"project": None, "segments": []}), encoding="utf-8")
        return stamp


def _backup_drop_empty(kind: str, stamp: Optional[str]) -> None:
    """Убрать зарезервированную копию, в которую так ничего и не записали."""
    if not stamp or not _STAMP_RE.fullmatch(stamp):
        return
    try:
        path = PURGE_DIR / (kind + "-" + stamp + ".json")
        if path.exists() and not (json.loads(path.read_text(encoding="utf-8"))
                                  .get("segments") or []):
            path.unlink()
    except Exception as e:                                      # pragma: no cover
        print(f"[backend] {kind}: пустая копия {stamp} не убрана: {e}", file=sys.stderr)


def _backup_segments(kind: str, pid: int, snapshot: list,
                     stamp: Optional[str] = None) -> str:
    """Копия состояния сегментов ПЕРЕД массовой правкой текста. Массовая
    команда без отката недопустима — тот же закон, что у пачек автоодобрения
    и выноса глоссария.

    Метка разводится суффиксом при занятости: две пачки подряд укладываются
    в одну секунду, и вторая молча затирала бы копию первой — а бэкап,
    который можно затереть, не бэкап. Не записалась копия — команда
    отменяется целиком: правка без отката хуже несделанной правки."""
    if stamp is not None and not _STAMP_RE.fullmatch(stamp):
        # Метка приходит из публичного API и из params задачи. Проверять её
        # надо ЗДЕСЬ, а не только в `_read_backup`: иначе команда перепишет
        # тексты и сложит копию под именем, которое чтение не примет, —
        # массовая правка без действующего отката.
        raise HTTPException(400, "Неверная метка отката")
    try:
        with _BACKUP_LOCK:
            PURGE_DIR.mkdir(parents=True, exist_ok=True)
            if stamp is None:
                stamp = _backup_stamp(kind)
            path = PURGE_DIR / (kind + "-" + stamp + ".json")
            segs = list(snapshot)
            if path.exists():
                # ДОПИСЫВАЕМ: метку задаёт прогон, а порций у него сотни. Своя
                # копия на каждую означала бы, что отменить прогон целиком нечем.
                old = json.loads(path.read_text(encoding="utf-8"))
                old_segs = old.get("segments") or []
                if old_segs and old.get("project") != pid:
                    raise RuntimeError("метка занята копией другого проекта")
                # Побеждает ПЕРВЫЙ снимок, а не последний, и это несущее
                # свойство отката. Сегмент попадает сюда дважды, когда порцию
                # повторяют (`JOB_CHUNK_RETRIES`) или когда прогон начинают
                # заново после рестарта: во второй раз в нём стоит уже НАШ
                # текст, и сохрани мы его — откат вернул бы машинную правку,
                # а перевод человека пропал бы навсегда (`prevTarget` пишется
                # только у заверенных, а `review.from` перезаписывается).
                have = {s["id"] for s in old_segs}
                segs = old_segs + [s for s in snapshot if s["id"] not in have]
            # Атомарно: оборванная запись растущего файла — это потерянный откат
            # у всех сегментов пачки, а не только у последней порции.
            tmp = path.with_suffix(".tmp")
            tmp.write_text(json.dumps({"project": pid, "segments": segs}, ensure_ascii=False),
                           encoding="utf-8")
            os.replace(tmp, path)
        return stamp
    except Exception as e:
        print(f"[backend] {kind}: копия для отката не записана: {e}", file=sys.stderr)
        raise HTTPException(500, "Не удалось сохранить копию для отката — команда отменена")


def _read_backup(kind: str, pid: int, stamp: str) -> dict:
    """Копия для отката, с проверкой метки и принадлежности проекту."""
    if not re.fullmatch(r"[0-9-]{8,24}", stamp or ""):
        raise HTTPException(400, "Неверная метка отката")
    path = PURGE_DIR / (kind + "-" + stamp + ".json")
    if not path.exists():
        raise HTTPException(404, "Копия для отката не найдена")
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except Exception as e:
        raise HTTPException(500, "Копия для отката не читается: " + str(e))
    if data.get("project") != pid:
        raise HTTPException(400, "Эта копия относится к другому проекту")
    return data


def _repair_accept_snapshot(seg: dict) -> dict:
    """Что надо вернуть, чтобы отменить принятие. Копия, а не ссылка: запись
    ремонта через строку будет переписана."""
    return {"id": seg["id"], "target": seg.get("target") or "",
            "status": seg.get("status"), "provider": seg.get("provider"),
            "route": seg.get("route"),
            "confirmedBy": seg.get("confirmedBy"), "confirmedAt": seg.get("confirmedAt"),
            "confirmedRole": seg.get("confirmedRole"),
            "prevTarget": seg.get("prevTarget"),
            "repair": json.loads(json.dumps(seg.get("repair") or {}))}


def _apply_repair_candidate(seg: dict) -> str:
    """Подставить `repair.candidate` в перевод. Возвращает ПРЕЖНИЙ текст.

    Вызова модели тут нет: подставляется то, что система уже написала.
    Одна ветка на одиночную и пакетную команду — разойдись они, «принять»
    и «принять все» оставляли бы сегмент в разных состояниях."""
    rp = seg["repair"]
    cand = rp["candidate"]
    was = seg.get("target") or ""
    # Тот же закон, что у любой машинной правки: заверенный человеком текст
    # уходит в prevTarget, статус становится review, отметка снимается.
    # Автоправка не заверяет сама себя — принял её человек, а не проверка.
    _replace_target(seg, cand, rp.get("model") or "", "REPAIR_ACCEPTED")
    # `_replace_target` ставит review только заверенному, остальным translated.
    # Здесь этого мало: текст НИКЕМ не проверен (проверки описывали отвергнутый
    # вариант и устарели вместе с ним), а сегмент мог стоять в review после
    # прошлой правки — и потерял бы отметку «посмотри человек». Ставим её сами,
    # ровно как это делает сам ремонт после применения правки.
    seg["status"] = "review"
    # Проверки описывают ОТВЕРГНУТЫЙ текст: их хеш от нового не совпадёт, и
    # `stale` встанет сам. Снимать их руками нельзя — на них держится история.
    seg["repair"] = {**rp, "applied": True, "from": was,
                     "source_hash": _text_hash(cand),
                     # След решения человека: без него следующий прогон
                     # не отличит принятого кандидата от машинной правки.
                     "acceptedBy": "human",
                     "acceptedAt": datetime.now().strftime("%Y-%m-%d %H:%M")}
    seg["repair"].pop("candidate", None)
    return was


@app.post("/api/segments/{pid}/{sid}/repair/accept")
def accept_repair_candidate(pid: int, sid: int):
    """Принять текст, который ремонт написал и отменил падением балла.

    Третья команда в системе, меняющая перевод БЕЗ вызова модели, — после
    `/glossary/revert-repairs` и `/term-case`, и по той же причине: она ничего
    не сочиняет. Подставляется `repair.candidate` — текст, который эта же
    система написала по конкретным находкам и который отвергла мерой,
    оказавшейся негодной для правки термина (балл back-check вознаграждает
    кальку, см. `_run_segment_repair`).

    Решение принимает ЧЕЛОВЕК. Пачкой то же самое делает
    `/api/projects/{pid}/repair/accept-batch` — с разбором, копией для отката
    и пощадой заверенному человеком. Оговорка одна на обе двери: проверки
    прежнего текста после подстановки устаревают сами (их хеш описывает
    отвергнутый вариант), а нового текста ещё нет ни у кого, — то есть сегмент
    остаётся непроверенным до ближайшего прогона. Это не запрет, а условие,
    и пачка называет его числом до применения."""
    _guard_project_write(pid)
    seg = get_segment(pid, sid)
    if not _repair_score_vetoed(seg):
        raise HTTPException(400, "У сегмента нет отменённой баллом правки, "
                                 "которую можно принять")
    was = _apply_repair_candidate(seg)
    cand = seg["target"]
    _IMPACT_CACHE.pop(pid, None)
    _ANALYSIS_CACHE.pop(pid, None)
    save_state(STATE)
    return {"ok": True, "target": cand, "prev": was,
            "segment": _segment_for_client(seg)}


class RepairAcceptBatchRequest(BaseModel):
    dry_run: bool = True                       # сначала показать, потом делать
    segment_ids: Optional[List[int]] = None    # None — все подходящие в проекте
    # Заверенное человеком пачкой не трогаем: подстановка снимет отметку
    # «подтвердил человек», а он заверял ДРУГОЙ текст. То же правило, что
    # у пакетного ремонта и переперевода по глоссарию.
    include_confirmed: bool = False


@app.post("/api/projects/{pid}/repair/accept-batch")
def accept_repair_candidates(pid: int, req: RepairAcceptBatchRequest = RepairAcceptBatchRequest()):
    """Принять пачкой все тексты, которые ремонт написал и отменил баллом.

    Одиночная кнопка есть в карточке сегмента, но сегментов таких сотня, и
    выбирать их в таблице по одному — та же работа руками, от которой заведены
    `/glossary/purge` и `/backcheck/rescore`. Поэтому команда есть, и живёт она
    по их правилам:

      1. `dry_run=True` по умолчанию — считает и показывает, ничего не меняя;
      2. заверенное человеком не трогается без явного `include_confirmed`,
         и сколько таких — сказано в ответе;
      3. прежнее состояние сегментов целиком уходит файлом в `data/backups/`
         и возвращается `/repair/accept/{stamp}/undo`. Массовая правка текста
         без отката недопустима.

    Оговорка, которую надо понимать. Проверки прежнего текста после подстановки
    устаревают САМИ (их хеш описывает отвергнутый вариант), а нового текста
    ещё нет ни у кого: пачка переводит сегменты в «никем не проверено».
    Это не порча — кандидат написан по конкретным находкам и отвергнут мерой,
    негодной для правки термина (см. `_run_segment_repair`), — но проверить его
    обязан ближайший прогон, и до тех пор сегмент стоит с непроверенным
    переводом. Поэтому число сказано до применения, а не после."""
    _guard_project_write(pid)
    project = get_project(pid)
    ids = set(req.segment_ids) if req.segment_ids is not None else None
    # Тот же отбор, что у корзины `human.revertedByScore` в /analysis, и это
    # обязательство: кнопка «Принять все» стоит в строке с числом N и обязана
    # применить ровно N. Корзина берёт сегмент только при ОТКРЫТЫХ находках —
    # если претензии с тех пор сняли (арбитр отменил `term_lost`, запись
    # глоссария понизили, откат снял termcheck), кандидат чинит то, чего
    # больше нет, и ставить его в перевод незачем.
    # Глоссарий берём из готового отчёта, а `_repair_findings` зовём БЕЗ
    # проекта — ровно как в /analysis: с проектом это `_get_context` на каждый
    # сегмент, десятки секунд единственного воркера.
    impact = glossary_impact(pid)
    gloss_bad = set(impact["segments"]) | set(impact.get("caseSegments") or ())
    matched, skipped_confirmed = [], []
    for seg in project["segments"]:
        if ids is not None and seg["id"] not in ids:
            continue
        if not _repair_score_vetoed(seg):
            continue
        if not (_repair_findings(seg) or seg["id"] in gloss_bad):
            continue
        if seg.get("status") == "confirmed" and not req.include_confirmed:
            skipped_confirmed.append(seg["id"])
            continue
        matched.append(seg)

    result = {
        "ok": True, "dryRun": req.dry_run,
        "matched": len(matched),
        "ids": [sg["id"] for sg in matched],
        # Молчаливой пощады не бывает: сказано, скольких не тронули и почему.
        "skippedConfirmed": skipped_confirmed,
        "samples": [{"id": sg["id"],
                     "was": (sg.get("target") or "")[:200],
                     "now": (sg["repair"].get("candidate") or "")[:200]}
                    for sg in matched[:12]],
        "accepted": 0, "stamp": None,
    }
    if req.dry_run or not matched:
        return result

    snapshot = [_repair_accept_snapshot(sg) for sg in matched]
    try:
        PURGE_DIR.mkdir(parents=True, exist_ok=True)
        # Метка с точностью до СЕКУНДЫ, а две пачки подряд укладываются
        # в одну: вторая молча затирала бы копию первой, и откат первой
        # возвращал бы чужие сегменты. Бэкап, который можно затереть, —
        # не бэкап, поэтому занятую метку разводим суффиксом.
        base = datetime.now().strftime("%Y%m%d-%H%M%S")
        stamp, n = base, 1
        while (PURGE_DIR / ("repair-accept-" + stamp + ".json")).exists():
            stamp = base + "-" + str(n)
            n += 1
        path = PURGE_DIR / ("repair-accept-" + stamp + ".json")
        path.write_text(json.dumps({"project": pid, "segments": snapshot},
                                   ensure_ascii=False), encoding="utf-8")
    except Exception as e:
        # Без бэкапа не применяем: откат — часть операции, а не украшение.
        print(f"[backend] принятие правок: бэкап не записан: {e}", file=sys.stderr)
        raise HTTPException(500, "Не удалось сохранить копию для отката — применение отменено")
    for sg in matched:
        _apply_repair_candidate(sg)
    _IMPACT_CACHE.pop(pid, None)
    _ANALYSIS_CACHE.pop(pid, None)
    save_state(STATE)
    print(f"[backend] принято отменённых правок: {len(matched)} "
          f"(проект {pid}), копия: {path.name}", file=sys.stderr)
    result["accepted"] = len(matched)
    result["stamp"] = stamp
    return result


@app.post("/api/projects/{pid}/repair/accept/{stamp}/undo")
def undo_accept_repair_candidates(pid: int, stamp: str):
    """Вернуть тексты, принятые пачкой.

    Возвращаем ТОЛЬКО те сегменты, где сейчас стоит именно принятый кандидат:
    правили после — значит поверх нашей подстановки легла чужая работа, и
    затирать её откатом нельзя. Тот же закон, что у `_repair_tried` в
    `_revert_repairs`. Неподходящие названы поимённо."""
    project = get_project(pid)
    # Метку из URL подставлять в путь как есть нельзя — то же правило и та же
    # проверка, что у откатов пересчёта баллов и выноса глоссария.
    if not re.fullmatch(r"[0-9-]{8,24}", stamp or ""):
        raise HTTPException(400, "Неверная метка отката")
    path = PURGE_DIR / ("repair-accept-" + stamp + ".json")
    if not path.exists():
        raise HTTPException(404, "Копия для отката не найдена")
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except Exception as e:
        raise HTTPException(500, "Копия для отката не читается: " + str(e))
    # Копия помнит, чей это проект. Без сверки метку от одного проекта можно
    # подать в откат другого: от порчи спасало бы только совпадение текстов,
    # то есть случайность, а не правило.
    if data.get("project") != pid:
        raise HTTPException(400, "Эта копия относится к другому проекту")
    by_id = {sg["id"]: sg for sg in project["segments"]}
    restored, changed_since = [], []
    for snap in data.get("segments") or []:
        seg = by_id.get(snap["id"])
        if seg is None:
            continue
        # Принятый кандидат лежит в снимке записи ремонта как `candidate`;
        # после применения он же стоит в переводе. Разошлись — текст правили.
        want = (snap.get("repair") or {}).get("candidate")
        if want is not None and (seg.get("target") or "") != want:
            changed_since.append(seg["id"])
            continue
        seg["target"] = snap["target"]
        seg["status"] = snap["status"]
        for k in ("provider", "route", "confirmedBy", "confirmedAt", "confirmedRole", "prevTarget"):
            if snap.get(k) is None:
                seg.pop(k, None)
            else:
                seg[k] = snap[k]
        seg["repair"] = snap["repair"]
        restored.append(seg["id"])
    if restored:
        _IMPACT_CACHE.pop(pid, None)
        _ANALYSIS_CACHE.pop(pid, None)
        save_state(STATE)
    return {"ok": True, "restored": restored, "changedSince": changed_since}



# ─── Совет арбитра одним нажатием ────────────────────────────────────
def _ctx_advices(seg: dict) -> list:
    """Вердикты арбитра «здесь термин передан неверно», которые есть чем
    исполнить: совет непустой, отличен от утверждённого перевода и БУКВАЛЬНО
    стоит в тексте сегмента — подставлять есть куда. Тот же отбор, что у
    корзины `human.termContextWrong` в /analysis, плюс проверка вхождения:
    кнопка обязана нажиматься там, где эндпоинт сработает, поэтому признак
    считает СЕРВЕР (`ctxAdvice` в `_segment_for_client`), а не браузер."""
    out = []
    target = seg.get("target") or ""
    for t in _term_context_of(seg):
        # Вердикт по забракованному СЛОВУ — не спор с записью глоссария:
        # он уже находка ремонта (kind term_ctx), и карточка «Арбитр: термин
        # передан неверно» с пустым src была бы про несуществующую запись.
        # Тот же отбор, что у human.termContextWrong в /analysis.
        if t.get("stale") or t.get("ok") is not False:
            continue
        use, tgt = (t.get("use") or "").strip(), (t.get("tgt") or "").strip()
        if not use or not tgt or _norm_key(use) == _norm_key(tgt):
            continue
        if not _ctx_pattern(tgt).search(target):
            continue
        out.append({"src": t.get("src"), "tgt": tgt, "use": use, "why": t.get("why")})
    return out


def _ctx_pattern(tgt: str):
    """Буквальное вхождение по границам слова, без учёта регистра. Пробелы
    в записи терпимы к переносам, остальное — как есть."""
    body = r"\s+".join(re.escape(w) for w in tgt.split())
    return re.compile(r"(?<![\w-])" + body + r"(?![\w-])", re.IGNORECASE)


def _ctx_substitute(target: str, tgt: str, use: str) -> str:
    """Заменить все вхождения `tgt` на `use`, унаследовав начертание найденного
    места: заглавная в начале и капс берутся у ТЕКСТА, а не у совета — тот же
    закон, что у `_case_like`. Ничего не сочиняется: слова совета те же."""
    def fit(m):
        found = m.group(0)
        if found.isupper() and len(found) >= CASE_CAPS_MIN:
            return use.upper()
        if found[:1].isupper() and not use[:1].isupper():
            return use[:1].upper() + use[1:]
        return use
    return _ctx_pattern(tgt).sub(fit, target)


class TermContextApplyRequest(BaseModel):
    src: str
    tgt: str
    use: str
    dry_run: bool = True
    segment_ids: Optional[List[int]] = None
    include_confirmed: bool = False



# ── Терм-лист документа (фаза 0) ──────────────────────────────────────
# Терминология решается ДО перевода, а не выкапывается из готовых сегментов:
# спор «Phthisiatry/Phthisiology» на боевой книге решился после того, как 2711
# сегментов были переведены и оплачены, и стоил ремонта. Задача `termsheet`
# читает ОРИГИНАЛ порциями, модель называет термины и стандартный перевод,
# дальше — ворота: приказ глоссария (тогда пара и не нужна: `shadowed`),
# форма (`_looks_like_term`, `_term_shape_reject`), корпус (вето только там,
# где у него есть право, `vetoAllowed`), сверка смысла (`same`/`rule` — тот же
# вопрос, что у автоодобрения). «Не знаю» остаётся `pending` и в промпт
# НЕ идёт.
# В глоссарий терм-лист НЕ пишет ничего (инвариант 8): это КОНТРАКТ
# КОНСИСТЕНТНОСТИ ОДНОГО ДОКУМЕНТА, `tier: "doc"`, и в промпт он попадает
# только по явному включению (`use`) — до A/B-замера на той же книге
# (`_termlist_measure`: вред termcheck на 10 тыс. вставок против прежних
# подсказок 15/11414). Приказ глоссария на том же куске сильнее всегда;
# в строгих областях (медицина, фарма, право) блок сформулирован просьбой,
# а не приказом. Один голос termcheck ЛЮБОЙ действующей тяжести
# (TERMCHECK_ACTIONABLE) снимает пару (`_termlist_dispute`): умолчание здесь
# «приказывать по всему документу», и цена ошибки другая, чем у разнобоя
# (там нужны два голоса). Той же тяжестью считает вред и замер.
TERMSHEET_CHUNK = 10
TERMSHEET_CORPUS_MAX = int(os.environ.get("TERMSHEET_CORPUS_MAX", "300"))
TERMSHEET_MEANING_MAX = int(os.environ.get("TERMSHEET_MEANING_MAX", "1500"))
TERMSHEET_VERSION = "1"
TERMLIST_PROMPT_MAX = 15


def _termsheet_system(domain: dict, src_lang: str, tgt_lang: str) -> str:
    """Промпт терм-листа. Отдельно от вызова — проверяется тестом настоящим кодом."""
    return (
        "You prepare a bilingual TERM SHEET for translating a " + domain["en"] + " document\n"
        "from " + src_lang + " to " + tgt_lang + ". You are given SOURCE segments only.\n"
        "Return ONLY a JSON array, no prose.\n\n"
        'Each item: {"src": <term in the source language, dictionary form (nominative singular)>,\n'
        '            "tgt": <the standard ' + tgt_lang + " term used in " + domain["en"] + ' publications>,\n'
        '            "cat": <one of ' + "|".join(domain["cats"]) + '>}\n\n'
        "RULES:\n"
        "1. " + domain["extract"] + "\n"
        "2. Include abbreviations with their standard " + tgt_lang + " equivalent — the target-language\n"
        "   abbreviation, never a transliteration of the source one.\n"
        "3. The target side is the accepted term of the field, never a word-by-word calque and never\n"
        "   a transliteration. If the field uses several names, give the most common one.\n"
        "4. Skip general vocabulary, numbers, whole sentences, anything longer than 5 words.\n"
        "5. At most 5 terms per segment. Return [] if the segments have no terminology.\n"
    )


def _termsheet_call(sources: list, model: Optional[str], domain_id: Optional[str],
                    src_lang: str, tgt_lang: str) -> list:
    """Один вызов на порцию оригиналов. Возвращает список пар или []."""
    import json as _json
    import openai
    dom = _resolve_domain(domain_id)
    mdl = _resolve_model(model or REVIEW_DEFAULT_MODEL)
    client = openai.OpenAI(api_key=os.environ.get("OPENAI_API_KEY"), timeout=90, max_retries=1)
    body = "\n\n".join(f"[{i + 1}] {t}" for i, t in enumerate(sources))
    extra = ({"max_completion_tokens": 4096} if mdl["api"] == "modern"
             else {"max_tokens": 1500, "temperature": 0})
    try:
        resp = client.chat.completions.create(
            model=mdl["id"],
            messages=[{"role": "system", "content": _termsheet_system(dom, src_lang, tgt_lang)},
                      {"role": "user", "content": body}],
            **extra,
        )
        _note_usage("terms", mdl["id"], resp)
        raw = (resp.choices[0].message.content or "").strip()
        lo, hi = raw.find("["), raw.rfind("]")
        if lo == -1 or hi <= lo:
            return None
        data = _json.loads(raw[lo:hi + 1])
        return [d for d in data if isinstance(d, dict) and d.get("src") and d.get("tgt")]
    except Exception as e:
        # None — «вызов не состоялся», и это НЕ то же, что «терминов нет»:
        # мёртвый ключ иначе выглядел бы как выполненный сбор с пустым списком.
        print(f"[backend] терм-лист: вызов не удался: {e}", file=sys.stderr)
        return None


def _termlist_entries(project: Optional[dict]) -> list:
    """Пары, которые вправе идти в промпт. В строгой области (медицина, фарма,
    право) — ТОЛЬКО принятые человеком: пары предлагает модель, и судит сверка
    той же моделью, — это самоодобрение, которое там запрещено даже глоссарию
    (инвариант 8). «Согласовано машиной» в строгой области — список на решение
    человека, не подсказка в промпт."""
    ents = ((project or {}).get("termlist") or {}).get("entries") or []
    strict = _termlist_strict(project) if project else True
    return [e for e in ents if e.get("status") == "agreed"
            and (not strict or e.get("by") == "human")]


def _termlist_active(project: Optional[dict]) -> bool:
    tl = (project or {}).get("termlist") or {}
    return bool(tl.get("use")) and bool(_termlist_entries(project))


def _termlist_prompt_fp(project: dict) -> str:
    """Отпечаток того, что из терм-листа видит промпт: смена — та же смена
    промпта ревизии, что у стайл-шита."""
    if not _termlist_active(project):
        return ""
    return _text_hash(json.dumps(sorted((e["src"], e["tgt"]) for e in _termlist_entries(project)),
                                 ensure_ascii=False))


def _reviews_mark_stale(project: dict) -> int:
    """Пометить свежие вердикты ревизии устаревшими (смена промпта: стайл-шит
    или терм-лист). Число — цена, её называют в ответе."""
    n = 0
    for seg in project.get("segments") or []:
        rv = seg.get("review")
        if rv and not rv.get("styleStale") and not _review_stale(seg):
            rv["styleStale"] = True
            n += 1
    return n


_TERMLIST_INDEX: dict = {}


def _termlist_index(project: dict) -> dict:
    """Индекс терм-листа по ключам-началам слов — как `_gloss_index`, но свой
    у проекта и по отпечатку согласованных пар: перебирать сотни регулярок
    на каждый сегмент незачем."""
    ents = _termlist_entries(project)
    fp = _text_hash(json.dumps([(e["src"], e["tgt"]) for e in ents], ensure_ascii=False))
    cached = _TERMLIST_INDEX.get(project["id"])
    if cached and cached[0] == fp:
        return cached[1]
    idx: dict = {}
    for e in ents:
        for k in _entry_keys(e.get("src") or ""):
            idx.setdefault(k, []).append(e)
    _TERMLIST_INDEX[project["id"]] = (fp, idx)
    return idx


def _doc_hits(source: str, project: Optional[dict], other: Optional[list] = None) -> list:
    """Согласованные пары терм-листа, чей термин стоит в ЭТОМ оригинале —
    хиты `tier: "doc"`. Пара, чей термин совпадает с приказной записью или
    вложен в неё (`other` — хиты глоссария), не идёт: приказ сильнее."""
    if not project or not source or not _termlist_active(project):
        return []
    idx = _termlist_index(project)
    seen, hits = set(), []
    for k in _text_keys(source):
        for e in idx.get(k, ()):
            if id(e) in seen:
                continue
            seen.add(id(e))
            form = _term_match(e["src"], source, _src_lang(e))
            if form:
                hits.append({"src": e["src"], "tgt": e["tgt"], "tier": "doc",
                             "lang": e.get("lang"), "_form": form})
    if not hits:
        return []
    hard = [_norm_key(h.get("src") or "") for h in (other or [])
            if h.get("tier") == GLOSSARY_TIER_HARD]

    def _inside(d: str) -> bool:
        # Целым словом: «рак» не вложен в «характер», «тест» — в «тестостерон».
        rx = re.compile(r"(?<!\w)" + re.escape(d) + r"(?!\w)")
        return any(rx.search(v) for v in hard if v)
    hits = [h for h in hits if not _inside(_norm_key(h["src"]))]
    best: dict = {}
    for h in hits:
        key = h["_form"].lower()
        if key not in best or len(h["src"]) > len(best[key]["src"]):
            best[key] = h
    return sorted(best.values(), key=lambda x: len(x["src"]), reverse=True)[:TERMLIST_PROMPT_MAX]


def _doc_misses(seg: dict, project: Optional[dict], skip: Optional[set] = None) -> list:
    """Пары терм-листа, чей термин есть в оригинале, а перевод — нет.
    Счётчик ремонта, НЕ находка: ремонт за терм-листом не ходит, но и ломать
    его правкой по другой претензии не вправе."""
    if not project or not _termlist_active(project):
        return []
    target = seg.get("target") or ""
    hits = _doc_hits(seg.get("source") or "", project, _verified_hits(seg.get("source") or "", project))
    skip = skip or set()
    return [h for h in hits if not _tgt_has_term(target, h["tgt"]) and _norm_key(h["tgt"]) not in skip]


def _doc_flagged(seg: dict, project: Optional[dict] = None) -> set:
    """Пары терм-листа, забракованные СВЕЖИМ termcheck в этом сегменте.
    Считается ОДИН раз, до правки, и передаётся в обе оценки: ремонт,
    заменивший такую пару по находке, иначе откатывался бы счётчиком «после»
    (у нового текста termcheck про исчезнувшее слово молчит) и клеймился."""
    tc = seg.get("termcheck") or {}
    if tc.get("target_hash") != _text_hash((seg.get("target") or "").strip()):
        return set()
    words = [f.get("tgt_term") for f in (tc.get("findings") or [])
             if f.get("severity") in TERMCHECK_ACTIONABLE and f.get("tgt_term")]
    if not words:
        return set()
    hits = _doc_hits(seg.get("source") or "", project)
    return {_norm_key(h["tgt"]) for h in hits if any(_term_forms_overlap(w, h["tgt"]) for w in words)}


def _term_forms_overlap(found: str, term: str) -> bool:
    """Находка termcheck цитирует фрагмент ПЕРЕВОДА («Bioptates»), а пара
    хранит словарную форму («bioptate»): сравнение буквальное их не сведёт,
    и ни диспут, ни исключение из счётчика не сработали бы ровно там, где
    нужны. Вхождение в любую сторону — тем же `_tgt_has_term`, каким
    `_gloss_misses` решает, стоит ли термин в тексте."""
    if not found or not term:
        return False
    return _norm_key(found) == _norm_key(term) or _tgt_has_term(found, term) or _tgt_has_term(term, found)


def _termlist_strict(project: dict) -> bool:
    return not _auto_policy((project or {}).get("domain")).get("allow_verified", True)


def _termsheet_block(terms: str, strict: bool) -> str:
    if strict:
        return ("\nDocument term sheet (agreed by the model for THIS document, not verified by a human):\n"
                + terms + "\nPrefer these translations for consistency across the document unless one is\n"
                "clearly wrong in this context. The approved glossary above always takes precedence.\n")
    return ("\nDocument term sheet — use these translations consistently across the document:\n"
            + terms + "\nThe approved glossary above always takes precedence.\n")


def _termlist_dispute(project: Optional[dict], src: str, why: str) -> bool:
    """Один голос снимает пару: согласованная машиной запись уходит из промпта
    и ждёт человека. Решение человека (`by: human`) машина не отменяет."""
    if not project:
        return False
    key = _norm_key(src)
    for e in ((project.get("termlist") or {}).get("entries") or []):
        if _norm_key(e.get("src")) == key and e.get("status") == "agreed" and e.get("by") != "human":
            e["status"] = "disputed"
            e["why"] = why
            e["votes"] = int(e.get("votes") or 0) + 1
            _TERMLIST_INDEX.pop(project["id"], None)
            return True
    return False


def _termlist_measure(project: dict) -> dict:
    """Замер вреда: среди сегментов, переведённых с терм-листом в промпте
    (`docTerms`) и проверенных termcheck, сколько вставленных и СТОЯЩИХ
    в переводе пар termcheck забраковал (любой действующей тяжести). Тот же
    счёт, каким похоронили подсказки автоимпорта: 15 подмен на 11 414 вставок.
    Меряется только вред: пользу можно измерить лишь двумя прогонами одних
    сегментов, и в карточке это сказано."""
    segs = ins = harm = 0
    samples = []
    for sg in project.get("segments") or []:
        dt = sg.get("docTerms") or []
        target = (sg.get("target") or "").strip()
        if not dt or not target:
            continue
        tc = sg.get("termcheck") or {}
        if tc.get("target_hash") != _text_hash(target):
            continue
        # Текст, переписанный ремонтом или ревизией, описывает не тот промпт.
        rp = sg.get("repair") or {}
        if _review_wrote(sg) or (rp.get("applied") and rp.get("source_hash") == _text_hash(target)):
            continue
        # Считаются только пары, которые СТОЯТ в переводе (модель послушалась):
        # тот же счёт, что у прежнего замера подсказок.
        dt = [t for t in dt if _tgt_has_term(target, t)]
        if not dt:
            continue
        segs += 1
        ins += len(dt)
        dts = {_norm_key(t) for t in dt}
        for f in tc.get("findings") or []:
            if f.get("severity") in TERMCHECK_ACTIONABLE and _norm_key(f.get("tgt_term")) in dts:
                harm += 1
                if len(samples) < 10:
                    samples.append({"id": sg["id"], "term": f.get("tgt_term"), "why": f.get("issue") or f.get("why") or ""})
    return {"segments": segs, "insertions": ins, "harm": harm,
            "per10k": (round(harm * 10000.0 / ins, 1) if ins else None),
            "baseline": {"insertions": 11414, "harm": 15, "per10k": 13.1}, "samples": samples}


def _termlist_counts(entries: list) -> dict:
    out = {k: 0 for k in ("agreed", "disputed", "rejected", "pending", "shadowed")}
    for e in entries:
        out[e.get("status")] = out.get(e.get("status"), 0) + 1
    return out


def _job_termsheet(job: dict) -> None:
    """Терм-лист документа: одна задача на проект, оригиналы порциями,
    вызовы параллельно, ворота после сбора. Прежние решения человека и уже
    оплаченные вердикты сверки/корпуса по той же паре переживают пересбор."""
    pid = job["project"]
    project = next((p for p in STATE["projects"] if p["id"] == pid), None)
    if project is None:
        job["status"], job["error"] = "error", "проект не найден"
        return
    src_lang, tgt_lang = project.get("src", "RU"), project.get("tgt", "EN")
    model = job["params"].get("model") or REVIEW_DEFAULT_MODEL
    segs = [sg for sg in project["segments"] if (sg.get("source") or "").strip()]
    chunks = [segs[i:i + TERMSHEET_CHUNK] for i in range(0, len(segs), TERMSHEET_CHUNK)]
    job["total"], job["done"] = len(segs), 0
    found: dict = {}

    def work(chunk):
        return _termsheet_call([sg["source"] for sg in chunk], model, project.get("domain"),
                               src_lang, tgt_lang)

    calls = failed = 0
    step = max(1, RUN_WORKERS * 2)
    for i in range(0, len(chunks), step):
        # Стоп-флаг читается ТЕМ ЖЕ помощником, что у порционных прогонов:
        # во внешнем воркере он перечитывает таблицу jobs и заодно пишет
        # прогресс; локальная копия `job["stop"]` там не обновляется.
        if _job_should_stop():
            job["status"] = "stopped"
            job["counters"].update({"calls": calls, "failed": failed, "stoppedAt": job["done"]})
            return    # список не пишется: половина книги под листом — разнобой по построению
        batch = chunks[i:i + step]
        for chunk, items in zip(batch, _run_parallel(batch, work)):
            calls += 1
            if items is None:
                failed += 1
                continue
            for it in items:
                s_, t_ = (it.get("src") or "").strip(), (it.get("tgt") or "").strip()
                if not s_ or not t_:
                    continue
                rec = found.setdefault(_norm_key(s_), {"src": s_, "tgts": {}, "cat": it.get("cat") or ""})
                rec["tgts"][t_] = rec["tgts"].get(t_, 0) + 1
            job["done"] += len(chunk)
    if calls and failed >= calls:
        # Ни один вызов не прошёл — это отозванный ключ, лимит или сеть,
        # а не пустой документ. Прежний список не трогается.
        job["status"], job["error"] = "error", "ни один вызов терм-листа не прошёл: ключ, лимит или сеть"
        job["counters"].update({"calls": calls, "failed": failed})
        return

    scope = _project_scope(project)
    pol = _auto_policy(project.get("domain"))
    old = ((project.get("termlist") or {}).get("entries") or [])
    prev = {(_norm_key(e.get("src")), _norm_key(e.get("tgt"))): e for e in old}
    entries = []
    for rec in found.values():
        tgts = sorted(rec["tgts"].items(), key=lambda x: -x[1])
        tgt = tgts[0][0]
        e = {"src": rec["src"], "tgt": tgt, "cat": rec["cat"], "lang": scope[0],
             "hits": sum(n for _t, n in tgts), "variants": [t for t, _n in tgts[1:4]],
             "status": "pending", "by": "model", "gates": {}, "why": ""}
        was = prev.get((_norm_key(e["src"]), _norm_key(tgt)))
        if was and was.get("by") == "human":
            # Решение человека переживает пересбор: своё предположение машина
            # вправе пересмотреть, чужое решение — нет.
            e.update({"status": was["status"], "by": "human", "why": was.get("why") or "",
                      "gates": dict(was.get("gates") or {})})
            entries.append(e)
            continue
        known = _glossary_entry(rec["src"], scope)
        if known and known.get("tier") == GLOSSARY_TIER_HARD:
            e["status"], e["gates"]["shadowedBy"] = "shadowed", known.get("tgt")
            entries.append(e)
            continue
        if not _looks_like_term(rec["src"], tgt):
            e["status"], e["why"] = "rejected", "не термин: обрывок фразы"
            entries.append(e)
            continue
        shape = _term_shape_reject(pol, rec["src"], tgt)
        if shape:
            e["status"], e["why"] = "rejected", shape
            entries.append(e)
            continue
        if len(tgts) > 1 and tgts[1][1] >= tgts[0][1]:
            e["status"] = "disputed"
            e["why"] = "модель предложила разные переводы: " + ", ".join(t for t, _n in tgts[:3])
        if was and was.get("gates"):
            e["gates"] = {k: v for k, v in was["gates"].items() if k in ("corpus", "meaning")}
            c = e["gates"].get("corpus") or {}
            if c and c.get("ok") is False and c.get("veto", True):
                e["status"], e["why"] = "rejected", "корпус целевого языка не знает такого термина — калька"
        entries.append(e)

    # Корпус: только про то, что не решено, в пределах потолка; вето — только
    # у корпуса своей области (`vetoAllowed`). Счётчики называют честно, что
    # корпус сделал: молчание и «нет права вето» — не защита.
    corpus_n = {"corpusChecked": 0, "corpusVeto": 0, "corpusSilent": 0, "corpusNoVeto": 0}
    for e in entries:
        if e["status"] not in ("pending", "disputed") or "corpus" in e["gates"]:
            continue
        if corpus_n["corpusChecked"] >= TERMSHEET_CORPUS_MAX:
            break
        c = _corpus_check(e["tgt"], scope)
        corpus_n["corpusChecked"] += 1
        if c is None:
            corpus_n["corpusSilent"] += 1
            continue
        e["gates"]["corpus"] = {"ok": c.get("ok"), "hits": c.get("hits"), "source": c.get("source"),
                                "veto": bool(c.get("vetoAllowed", True))}
        if not c.get("ok"):
            if c.get("vetoAllowed", True):
                corpus_n["corpusVeto"] += 1
                e["status"], e["why"] = "rejected", "корпус целевого языка не знает такого термина — калька"
            else:
                corpus_n["corpusNoVeto"] += 1
    # Готовый вердикт той же пары на записи глоссария (аудит уже платил за него)
    # не переспрашивается.
    for e in entries:
        if e["status"] not in ("pending", "disputed") or "meaning" in e["gates"]:
            continue
        g = _glossary_entry(e["src"], scope) or {}
        mv = g.get("meaning") or {}
        if mv and mv.get("pair") == _meaning_pair({"src": e["src"], "tgt": e["tgt"]}) \
                and str(mv.get("v")) == str(MEANING_VERSION) and mv.get("same") is not None:
            e["gates"]["meaning"] = {"same": mv.get("same"), "rule": mv.get("rule"), "why": mv.get("why") or ""}

    # Сверка смысла — те же два вопроса, что у автоодобрения. Готовый
    # вердикт по той же паре не переспрашивается.
    cands = [{"src": e["src"], "tgt": e["tgt"], "lang": scope[0], "domain": scope[1], "tenant": scope[2]}
             for e in entries if e["status"] in ("pending", "disputed") and "meaning" not in e["gates"]]
    verdicts, _answered, capped = _meaning_check(cands, cap=TERMSHEET_MEANING_MAX) if cands else ({}, 0, 0)
    for e in entries:
        if e["status"] not in ("pending", "disputed"):
            continue
        v = e["gates"].get("meaning")
        if v is None:
            raw = verdicts.get((scope, _norm_key(e["src"]), _norm_key(e["tgt"])))
            if not raw:
                continue
            v = {"same": raw.get("same"), "rule": raw.get("rule"), "why": raw.get("why") or ""}
            e["gates"]["meaning"] = v
        if v.get("same") is False:
            e["status"], e["why"] = "rejected", "не то понятие: " + (v.get("why") or "")
        elif v.get("rule") is False:
            e["status"], e["why"] = "disputed", "верно в контексте, но не годится правилом: " + (v.get("why") or "")
        elif v.get("same") and v.get("rule") and e["status"] == "pending":
            e["status"] = "agreed"
    entries.sort(key=lambda x: (-x["hits"], x["src"].lower()))
    # Слияние под локом: решение, принятое человеком ПОКА шёл сбор (в одном
    # процессе охранник записи не держит), не теряется, а принятая человеком
    # пара, которую модель в этот раз не назвала, остаётся в списке.
    with _SAVE_LOCK:
        cur = project.get("termlist") or {}
        human = {_norm_key(e.get("src")): e for e in (cur.get("entries") or []) if e.get("by") == "human"}
        seen = set()
        for e in entries:
            k = _norm_key(e["src"])
            h = human.get(k)
            if not h:
                continue
            same_pair = _norm_key(h.get("tgt")) == _norm_key(e["tgt"])
            # «Эта пара неверна» не равно «с термином разобрались»: отклонение
            # относится к ПАРЕ, и новый перевод того же термина проходит ворота
            # сам. Согласие человека — решение о термине, оно сильнее.
            if h.get("status") == "agreed" or same_pair:
                seen.add(k)
                e.update({"status": h["status"], "tgt": h.get("tgt") or e["tgt"], "by": "human",
                          "why": h.get("why") or ""})
                if same_pair:
                    e["gates"] = dict(h.get("gates") or e.get("gates") or {})
        for k, h in human.items():
            if k not in seen:
                entries.append(h)     # решение человека не исчезает, если модель термин не назвала
        if failed:
            # Часть порций не прошла (429, сеть): прежние машинные пары из них
            # неоткуда взять заново — оставляем как были, а число упавших
            # порций уходит в счётчики и на карточку.
            have = {(_norm_key(e["src"]), _norm_key(e["tgt"])) for e in entries}
            for e in cur.get("entries") or []:
                if (_norm_key(e.get("src")), _norm_key(e.get("tgt"))) not in have:
                    entries.append(e)
        project["termlist"] = {"v": TERMSHEET_VERSION, "at": datetime.now().strftime("%Y-%m-%d %H:%M"),
                               "model": _resolve_model(model)["id"], "strict": _termlist_strict(project),
                               "use": bool(cur.get("use", False)),
                               "acceptedBy": cur.get("acceptedBy"), "acceptedAt": cur.get("acceptedAt"),
                               "entries": entries}
        _TERMLIST_INDEX.pop(pid, None)
    job["counters"].update({"calls": calls, "failed": failed, "terms": len(entries),
                            "meaningCapped": capped, **corpus_n, **_termlist_counts(entries)})
    if failed:
        project["termlist"]["partial"] = failed
    _ANALYSIS_CACHE.pop(pid, None)
    save_state(STATE)


class TermlistBody(BaseModel):
    use: Optional[bool] = None
    decisions: Optional[list] = None     # [{"src", "status": agreed|rejected, "tgt"?}]
    # «Принять все согласованные машиной» — одно решение человека списком,
    # со следом (`acceptedBy/At`): в строгой области без него в промпт
    # не идёт ничего.
    accept_all: bool = False


def _termlist_view(project: dict) -> dict:
    tl = project.get("termlist") or {}
    ents = tl.get("entries") or []
    return {"ok": True, "built": bool(tl), "use": bool(tl.get("use")), "at": tl.get("at"),
            "model": tl.get("model"), "strict": _termlist_strict(project),
            "acceptedAt": tl.get("acceptedAt"), "active": len(_termlist_entries(project)),
            "partial": tl.get("partial") or 0,
            "pendingHuman": sum(1 for e in ents if e.get("status") == "agreed" and e.get("by") != "human"),
            "counts": _termlist_counts(ents),
            # Карточке нужны спорные (все) и согласованные (образец): тысячи
            # отклонённых пар с причинами — сотни килобайт на каждый клик.
            "entries": [{k: e.get(k) for k in ("src", "tgt", "status", "by", "why", "hits", "variants", "cat", "votes")}
                        for e in ents if e.get("status") == "disputed"]
                       + [{k: e.get(k) for k in ("src", "tgt", "status", "by", "hits")}
                          for e in ents if e.get("status") == "agreed"][:300],
            "measure": _termlist_measure(project)}


@app.get("/api/projects/{pid}/termlist")
def get_termlist(pid: int):
    return _termlist_view(get_project(pid))


@app.post("/api/projects/{pid}/termlist")
def set_termlist(pid: int, req: TermlistBody):
    """Включить терм-лист в промпты и записать решения человека по спорным
    парам. Бесплатно; в глоссарий не пишет ничего."""
    _guard_project_write(pid)
    if _job_busy(pid, "termsheet"):
        raise HTTPException(409, "Идёт сбор терм-листа — дождитесь конца или остановите его")
    project = get_project(pid)
    tl = project.get("termlist")
    if not tl:
        raise HTTPException(400, "Терм-лист ещё не собран")
    before = _termlist_prompt_fp(project)
    if req.use is not None:
        tl["use"] = bool(req.use)
    ents = tl.get("entries") or []
    decided = 0
    if req.accept_all:
        for e in tl.get("entries") or []:
            if e.get("status") == "agreed" and e.get("by") != "human":
                e["by"], e["why"] = "human", "принято списком"
                decided += 1
        tl["acceptedBy"] = _actor_id()
        tl["acceptedAt"] = datetime.now().strftime("%Y-%m-%d %H:%M")
    for d in req.decisions or []:
        d = d or {}
        st = d.get("status")
        ks, kt = _norm_key(d.get("src") or ""), _norm_key(d.get("tgt") or "")
        # Решение адресует ПАРУ: у одного термина в списке может стоять и
        # отклонённый человеком перевод, и новый от модели.
        e = next((x for x in ents if _norm_key(x.get("src")) == ks
                  and (not kt or _norm_key(x.get("tgt")) == kt)
                  and (kt or x.get("status") != "rejected")), None)
        if e is None or st not in ("agreed", "rejected"):
            continue
        e["status"], e["by"] = st, "human"
        e["why"] = "решение человека" if st == "agreed" else "отклонено человеком"
        decided += 1
    _TERMLIST_INDEX.pop(pid, None)
    # Состав промпта изменился — вердикты ревизии устарели, как при смене
    # стайл-шита; число названо.
    stale = _reviews_mark_stale(project) if _termlist_prompt_fp(project) != before else 0
    _ANALYSIS_CACHE.pop(pid, None)
    save_state(STATE)
    _audit("termlist.set", project=pid, use=req.use, decided=decided, accept_all=req.accept_all)
    return {**_termlist_view(project), "decided": decided, "reviewsStale": stale}


# ── Стайл-шит документа ────────────────────────────────────────────────
# Единая стилистика — свойство ДОКУМЕНТА, а не соседей: два соседних сегмента
# не скажут, как в статье пишутся заголовки, аббревиатуры при первом
# упоминании или американская ли орфография. Поэтому правила лежат таблицей
# из ВЫБОРОВ (человек, не знающий целевого языка, выбирает журнал и вариант
# орфографии, а не слова) и уходят ОДНИМ блоком в промпты перевода, ревизии
# и ремонта (`_style_block`), а детерминированная проверка (`_style_findings`)
# читает ТЕ ЖЕ поля — промпт и проверка разойтись не могут.
# Три слоя, как у областей: встроенный пресет по области → пресет организации
# (запись tenant, как pricing) → проект. Стайл-шит ВКЛЮЧАЕТСЯ на проекте
# (`project["style"]` — dict): без него блок пуст, промпты байт в байт прежние,
# и ни один оплаченный вердикт не устаревает.
# Чего в первой версии НЕТ намеренно: формат чисел и единиц (сравнение чисел
# идёт по строкам, и «1,3 → 1.3» стало бы объективной находкой, снимающей
# заверение человека), регистр заголовков кроме «как в оригинале» (так
# считают `_case_misses` и `_term_case_spans`, второе правило рядом дало бы
# две правки одного места по очереди), курсив (в сегменте нет оформления).
STYLE_FIELDS = {
    "preset":        ("", "ama", "vancouver", "apa", "nature"),
    "spelling":      ("", "US", "UK"),
    "register":      ("", "academic", "clinical", "textbook", "plain"),
    "abbreviations": ("", "expand_first", "as_source"),
    "quotes":        ("", "double", "single", "guillemets"),
    "headings":      ("", "as_source"),
}
STYLE_PRESETS = {
    "ama":       {"spelling": "US", "register": "academic", "abbreviations": "expand_first", "quotes": "double"},
    "vancouver": {"spelling": "US", "register": "academic", "abbreviations": "expand_first", "quotes": "double"},
    "apa":       {"spelling": "US", "register": "academic", "abbreviations": "expand_first", "quotes": "double"},
    "nature":    {"spelling": "UK", "register": "academic", "abbreviations": "expand_first", "quotes": "single"},
}
STYLE_PRESET_LABEL = {"ama": "AMA Manual of Style", "vancouver": "Vancouver (ICMJE)",
                      "apa": "APA", "nature": "Nature journals"}
# Встроенный пресет по области (по `base` своей области организации).
STYLE_DOMAIN_PRESET = {"medical": "ama", "pharma": "ama"}
STYLE_TEXT = {
    "spelling": {"US": "Spelling: American English (US).",
                 "UK": "Spelling: British English (UK, -ise forms)."},
    "register": {"academic": "Register: academic prose for a peer-reviewed publication — precise, impersonal, no colloquialisms.",
                 "clinical": "Register: clinical documentation — concise, standard clinical phrasing.",
                 "textbook": "Register: textbook prose — clear, explanatory, consistent terminology.",
                 "plain":    "Register: plain language for a general reader."},
    "abbreviations": {"expand_first": "Abbreviations: spell out at first use in the document as “Full form (ABBR)”, then use the abbreviation consistently.",
                      "as_source":    "Abbreviations: keep them exactly as the source uses them."},
    "quotes": {"double": "Quotation marks: “double”.", "single": "Quotation marks: ‘single’.",
               "guillemets": "Quotation marks: «guillemets»."},
    "headings": {"as_source": "Headings: keep the casing of the source."},
}


def _style_layer(fields: dict) -> dict:
    """Один слой: пресет заполняет поля, явные значения слоя сильнее пресета."""
    preset = (fields or {}).get("preset") or ""
    out = dict(STYLE_PRESETS.get(preset, {}))
    for k, v in (fields or {}).items():
        if k != "preset" and v:
            out[k] = v
    out["preset"] = preset
    return out


def _style_effective(project: dict) -> Optional[dict]:
    """Действующие поля: встроенный пресет области ← организация ← проект.
    None — стайл-шит на проекте не включён."""
    if not isinstance((project or {}).get("style"), dict):
        return None
    dom = _resolve_domain(project.get("domain"))
    rec = _tenant_rec(project.get("tenant") or _current_tenant()) or {}
    eff = _style_layer({"preset": STYLE_DOMAIN_PRESET.get(dom.get("base") or dom.get("id") or "", "")})
    for layer in (rec.get("style") or {}, project.get("style") or {}):
        for k, v in _style_layer(layer).items():
            if v:
                eff[k] = v
    if (project.get("tgt") or "").upper() != "EN":
        # Орфографический вариант есть только у английского; на другой паре
        # поле молчит, как и правила, которых для языка нет.
        eff["spelling"] = ""
    eff.setdefault("headings", "as_source")
    return {k: eff.get(k, "") for k in STYLE_FIELDS}


def _style_fp(project: dict) -> str:
    eff = _style_effective(project)
    return _text_hash(json.dumps(eff, sort_keys=True)) if eff else ""


def _style_block(project: Optional[dict]) -> str:
    """Блок STYLE SHEET для промптов. Пусто, пока стайл-шит не включён."""
    eff = _style_effective(project or {}) if project else None
    if not eff:
        return ""
    lines = []
    if eff.get("preset"):
        lines.append("Journal style: " + STYLE_PRESET_LABEL.get(eff["preset"], eff["preset"]) + ".")
    for k in ("spelling", "register", "abbreviations", "quotes", "headings"):
        t = STYLE_TEXT.get(k, {}).get(eff.get(k) or "")
        if t:
            lines.append(t)
    if not lines:
        return ""
    # Приказ глоссария сильнее стиля: иначе «Spelling: US» рядом с приказом
    # «haemoglobin» даёт модели два требования, ремонт подставит одно,
    # `_repair_scores["gloss"]` откатит — платный вызов с известным исходом.
    lines.append("Approved glossary terms take precedence over these conventions.")
    return ("STYLE SHEET — document-wide conventions, apply consistently:\n"
            + "\n".join("- " + l for l in lines) + "\n")


def _style_state(project: dict) -> dict:
    enabled = isinstance(project.get("style"), dict)
    dom = _resolve_domain(project.get("domain"))
    rec = _tenant_rec(project.get("tenant") or _current_tenant()) or {}
    return {"enabled": enabled,
            "project": dict(project.get("style") or {}) if enabled else {},
            "org": dict(rec.get("style") or {}),
            "builtinPreset": STYLE_DOMAIN_PRESET.get(dom.get("base") or dom.get("id") or "", ""),
            "effective": _style_effective(project) or {},
            "block": _style_block(project),
            "spellingApplies": (project.get("tgt") or "").upper() == "EN"}


# Орфография: (британская основа, американская основа, допустимый хвост).
# Хвост — ЗАКРЫТЫЙ список, а не «любые буквы»: «organis» + любые буквы ловил бы
# «organism», «liter» без границы — «literature», «analys» + «is» — «analysis»,
# который в обоих вариантах один. Слова, у которых вариант зависит от части речи
# или спорен (practise/practice, licence/license, programme, sulphur, foetus,
# metre/meter-прибор), в списке нет намеренно: пропуск дешевле подмены.
_SPELL_VERB = r"(?:e|es|ed|ing|er|ers|ation|ations|ational|ationally)"
_SPELLING = [
    ("tumour", "tumor", r"s?"), ("colour", "color", r"(?:s|ed|ing|less|ful)?"),
    ("behaviour", "behavior", r"(?:s|al)?"), ("labour", "labor", r"(?:s|ed|ing)?"),
    ("haemoglobin", "hemoglobin", r"s?"), ("oedema", "edema", r"s?"), ("oedemat", "edemat", r"[a-z]{1,4}"),
    ("oesophag", "esophag", r"[a-z]{1,8}"), ("anaemi", "anemi", r"[a-z]{1,3}"),
    ("leukaemi", "leukemi", r"[a-z]{1,3}"), ("ischaemi", "ischemi", r"[a-z]{1,3}"),
    ("haemorrhag", "hemorrhag", r"[a-z]{1,4}"), ("haematolog", "hematolog", r"[a-z]{1,5}"),
    ("anaesthe", "anesthe", r"[a-z]{1,6}"), ("paediatr", "pediatr", r"[a-z]{1,6}"),
    ("orthopaed", "orthoped", r"[a-z]{1,4}"), ("gynaecolog", "gynecolog", r"[a-z]{1,5}"),
    ("aetiolog", "etiolog", r"[a-z]{1,5}"), ("oestrogen", "estrogen", r"[a-z]{0,2}"),
    ("diarrhoea", "diarrhea", r"l?"), ("faeces", "feces", r""), ("faecal", "fecal", r""),
    ("caesarean", "cesarean", r"s?"), ("catalogue", "catalog", r"(?:s|d)?"),
    ("centre", "center", r"s?"), ("litre", "liter", r"s?"), ("fibre", "fiber", r"s?"),
    ("defence", "defense", r"s?"),
    ("analys", "analyz", _SPELL_VERB), ("paralys", "paralyz", _SPELL_VERB),
    ("organis", "organiz", _SPELL_VERB), ("randomis", "randomiz", _SPELL_VERB),
    ("immunis", "immuniz", _SPELL_VERB), ("hospitalis", "hospitaliz", _SPELL_VERB),
    ("sterilis", "steriliz", _SPELL_VERB), ("characteris", "characteriz", _SPELL_VERB),
    ("minimis", "minimiz", _SPELL_VERB), ("maximis", "maximiz", _SPELL_VERB),
    ("optimis", "optimiz", _SPELL_VERB), ("recognis", "recogniz", _SPELL_VERB),
    ("stabilis", "stabiliz", _SPELL_VERB), ("utilis", "utiliz", _SPELL_VERB),
    ("standardis", "standardiz", _SPELL_VERB), ("normalis", "normaliz", _SPELL_VERB),
    ("localis", "localiz", _SPELL_VERB), ("generalis", "generaliz", _SPELL_VERB),
    ("visualis", "visualiz", _SPELL_VERB), ("categoris", "categoriz", _SPELL_VERB),
]
_SPELL_RULES: dict = {}


def _spelling_rules(want: str) -> list:
    if want not in _SPELL_RULES:
        rules = []
        for uk, us, tail in _SPELLING:
            frm, to = (uk, us) if want == "US" else (us, uk)
            rules.append((re.compile(r"\b(" + frm + r")(" + tail + r")\b", re.I), to))
        _SPELL_RULES[want] = rules
    return _SPELL_RULES[want]


def _match_case(sample: str, word: str) -> str:
    if len(sample) > 1 and sample.isupper():
        return word.upper()
    if sample[:1].isupper():
        return word[:1].upper() + word[1:]
    return word


def _spelling_fix(text: str, want: str, protect: frozenset = frozenset()) -> tuple:
    """(текст в нужном варианте, [(было, стало)]). Начертание слова сохраняется.

    Не трогаем: слово из `protect` (приказные переводы глоссария — иначе
    правка и ремонт ходили бы по кругу) и слово с заглавной НЕ в начале
    предложения — имя собственное, название, латинский род: «Centers for
    Disease Control», «Department of Defense», «Oesophagostomum», заглавие
    цитируемой статьи. Пропуск дешевле подмены в библиографии."""
    if want not in ("US", "UK") or not text:
        return text, []
    changes = []
    for rx, to in _spelling_rules(want):
        cur = text

        def _r(m, to=to, cur=cur):
            whole, head = m.group(0), m.group(1)
            if whole.lower() in protect:
                return whole
            if head[:1].isupper() and not head.isupper() and not _sentence_start(cur, m.start()):
                return whole
            new = _match_case(head, to) + m.group(2)
            changes.append((whole, new))
            return new
        text = rx.sub(_r, text)
    return text, changes


def _style_protected_words(project: dict) -> frozenset:
    """Слова приказных переводов области проекта: орфография их не правит."""
    scope = _project_scope(project)
    words = set()
    for g in STATE.get("glossary") or []:
        if g.get("tier") != GLOSSARY_TIER_HARD or _scope_of(g) != scope:
            continue
        words.update(w.lower() for w in re.findall(r"[A-Za-z]{3,}", g.get("tgt") or ""))
    return frozenset(words)


# Аббревиатуры, которые журналы не расшифровывают, и то, что аббревиатурой
# лишь выглядит (римские числа, единицы).
_ABBR_KNOWN = frozenset((
    "DNA", "RNA", "HIV", "AIDS", "WHO", "PCR", "CT", "MRI", "US", "UK", "USA", "EU",
    "SI", "ISO", "AM", "PM", "OK", "TB", "BCG", "COVID", "SARS", "CD", "IG", "ID",
    "II", "III", "IV", "VI", "VII", "VIII", "IX", "XI", "XII", "XIII", "XIV", "XV",
    "MG", "ML", "KG", "CM", "MM", "MHZ", "KHZ", "GHZ", "KM", "NM", "MCG", "IU",
))
# Дефисное сложение («MDR-TB») — одна аббревиатура: расшифровка «(MDR-TB)»
# покрывает и её части.
_ABBR_RE = re.compile(r"\b[A-Z][A-Z0-9]{1,5}(?:-[A-Z][A-Z0-9]{1,5})*s?\b")


def _abbr_report(project: dict) -> list:
    """Аббревиатуры без расшифровки при ПЕРВОМ упоминании — в порядке
    документа (список сегментов и есть его порядок: сегменты с картинок
    стоят у своей картинки). Только отчёт: расшифровку сочинять нельзя."""
    seen, out = set(), []
    for seg in project.get("segments") or []:
        t = seg.get("target") or ""
        if not t.strip():
            continue
        words = re.findall(r"[A-Za-z]{2,}", t)
        if words and all(w.isupper() for w in words):
            continue    # капс-заголовок («MAIN FORMS», «NOTE:»): там каждое слово похоже на аббревиатуру
        for m in _ABBR_RE.finditer(t):
            a = m.group(0)
            if a.endswith("s") and len(a) > 2:
                a = a[:-1]
            if a in seen or a in _ABBR_KNOWN or a.isdigit():
                continue
            seen.add(a)
            expanded = bool(re.search(r"\(" + re.escape(a) + r"s?\)", t)
                            or re.search(r"\b" + re.escape(a) + r"s?\s*\(", t))
            if expanded:
                seen.update(a.split("-"))
            elif "-" in a and all(x in seen or x in _ABBR_KNOWN for x in a.split("-")):
                continue    # части уже расшифрованы порознь
            else:
                out.append({"abbr": a, "id": seg["id"]})
    return out


def _style_findings(project: dict, ids: Optional[set] = None) -> dict:
    """Детерминированные находки по действующему стайл-шиту: орфография
    (с готовой заменой) и аббревиатуры (отчёт). Бесплатно."""
    eff = _style_effective(project) or {}
    want = eff.get("spelling") or ""
    protect = _style_protected_words(project) if want else frozenset()
    spelling = []
    for seg in project.get("segments") or []:
        if ids is not None and seg["id"] not in ids:
            continue
        t = seg.get("target") or ""
        if not t.strip() or not want:
            continue
        new, ch = _spelling_fix(t, want, protect)
        if ch:
            spelling.append({"id": seg["id"], "status": seg.get("status"),
                             "changes": [list(c) for c in ch[:20]], "now": new})
    abbr = _abbr_report(project) if eff.get("abbreviations") == "expand_first" else []
    return {"spelling": spelling, "abbreviations": abbr, "want": want}


class StyleBody(BaseModel):
    fields: dict = {}
    enable: Optional[bool] = None


class StyleCheckRequest(BaseModel):
    dry_run: bool = True
    include_confirmed: bool = False
    segment_ids: Optional[list] = None


def _style_validate(fields: dict) -> dict:
    out = {}
    for k, v in (fields or {}).items():
        if k not in STYLE_FIELDS:
            raise HTTPException(400, "Неизвестное поле стайл-шита: " + str(k))
        v = v or ""
        if not isinstance(v, str) or v not in STYLE_FIELDS[k]:
            raise HTTPException(400, "Недопустимое значение поля «" + k + "»: " + str(v))
        out[k] = v
    return out


@app.get("/api/style")
def get_org_style(request: Request):
    """Стайл-шит СВОЕЙ организации плюс каталог полей. Наружу его не отдаёт
    никто другой: `/api/auth/me` перечисляет поля организации белым списком."""
    _current_user(request)
    rec = _tenant_rec(_current_tenant()) or {}
    return {"ok": True, "tenant": _current_tenant(), "style": dict(rec.get("style") or {}),
            "catalog": STYLE_FIELDS, "presets": STYLE_PRESETS,
            "domainDefaults": STYLE_DOMAIN_PRESET}


@app.post("/api/style")
def save_org_style(req: StyleBody, request: Request):
    """Правит ВЛАДЕЛЕЦ (строка в `_OWNER_ONLY`): это продукт агентства."""
    _current_user(request)
    rec = _tenant_rec(_current_tenant())
    if not rec:
        raise HTTPException(404, "Записи вашей организации нет в базе — стайл-шит сохранять некуда.")
    cur = dict(rec.get("style") or {})
    cur.update(_style_validate(req.fields))
    rec["style"] = {k: v for k, v in cur.items() if v}
    save_state(STATE)
    # Запись организации — документ, а эпоху документов поднимает только
    # воркер для проектов. Без этого `medcat-worker` до рестарта переводил бы
    # под прежним слоем организации, а API показывал бы новый.
    try:
        if hasattr(STORE, "bump_epoch"):
            STORE.bump_epoch("doc:tenants")
    except Exception as e:
        print(f"[backend] стайл-шит организации: эпоха не поднята: {e}", file=sys.stderr)
    _audit("style.org", fields=sorted(req.fields or {}))
    return {"ok": True, "style": rec["style"]}


@app.get("/api/projects/{pid}/style")
def get_project_style(pid: int):
    project = get_project(pid)
    return {"ok": True, **_style_state(project), "catalog": STYLE_FIELDS, "presets": STYLE_PRESETS}


@app.post("/api/projects/{pid}/style")
def set_project_style(pid: int, req: StyleBody):
    """Включает стайл-шит на проекте и правит его поля. Пустое поле снимает
    значение проекта — действует слой ниже. Смена ДЕЙСТВУЮЩИХ полей делает
    вердикты ревизии устаревшими (метка на месте, число в ответе): ревизор
    перечитает под новые правила, а это платно. Готовые переводы и заходы
    ремонта не трогаются — переперевод под новый стиль отдельное решение."""
    _guard_project_write(pid)
    project = get_project(pid)
    before = _style_fp(project)
    if req.enable is False:
        project.pop("style", None)
    else:
        cur = dict(project.get("style") or {}) if isinstance(project.get("style"), dict) else {}
        cur.update(_style_validate(req.fields))
        project["style"] = {k: v for k, v in cur.items() if v}
    after = _style_fp(project)
    stale = 0
    if after != before:
        stale = _reviews_mark_stale(project)
        _ANALYSIS_CACHE.pop(pid, None)
    save_state(STATE)
    _audit("style.project", project=pid, fields=sorted(req.fields or {}), enable=req.enable)
    return {"ok": True, **_style_state(project), "reviewsStale": stale, "changed": after != before}


def _style_backup(pid: int, snapshot: list) -> tuple:
    PURGE_DIR.mkdir(parents=True, exist_ok=True)
    base = datetime.now().strftime("%Y%m%d-%H%M%S")
    stamp, n = base, 1
    while (PURGE_DIR / ("style-" + stamp + ".json")).exists():
        stamp = base + "-" + str(n)
        n += 1
    path = PURGE_DIR / ("style-" + stamp + ".json")
    path.write_text(json.dumps({"project": pid, "segments": snapshot}, ensure_ascii=False),
                    encoding="utf-8")
    return stamp, path


@app.post("/api/projects/{pid}/style-check")
def style_check(pid: int, req: StyleCheckRequest):
    """Проверка стиля по действующему стайл-шиту: орфографический вариант
    (с готовой заменой) и аббревиатуры без расшифровки (только отчёт).
    Седьмая команда, меняющая текст БЕЗ вызова модели, и по той же причине:
    она ничего не сочиняет — меняются буквы внутри слов из закрытого списка.
    Правила массовых команд: `dry_run` по умолчанию, заверенное не трогается
    без `include_confirmed` и названо числом, копия в `data/backups/`,
    откат `/style-check/{stamp}/undo`. Бесплатно."""
    _guard_project_write(pid)
    project = get_project(pid)
    if not isinstance(project.get("style"), dict):
        raise HTTPException(400, "Стайл-шит проекта не включён")
    ids = set(req.segment_ids) if req.segment_ids is not None else None
    rep_ = _style_findings(project, ids)
    by_id = {sg["id"]: sg for sg in project["segments"]}
    todo, skipped_confirmed = [], []
    for f in rep_["spelling"]:
        sg = by_id[f["id"]]
        if sg.get("status") == "confirmed" and not req.include_confirmed:
            skipped_confirmed.append(sg["id"])
            continue
        todo.append((sg, f))
    # Правка меняет текст, и хеши back-check, termcheck и ревизии перестают
    # его описывать: ближайший прогон купит проверки заново. Число названо
    # ДО применения — как у пачки `/repair/accept-batch`.
    stale_checks = 0
    for sg, _f in todo:
        cur_h = _text_hash((sg.get("target") or "").strip())
        if any((sg.get(k) or {}).get("target_hash") == cur_h for k in ("backcheck", "termcheck")) \
                or not _review_stale(sg):
            stale_checks += 1
    result = {"ok": True, "dryRun": req.dry_run, "spelling": rep_["want"],
              "staleChecks": stale_checks,
              "spellingSegments": len(rep_["spelling"]),
              "spellingChanges": sum(len(f["changes"]) for f in rep_["spelling"]),
              "ids": [sg["id"] for sg, _ in todo], "skippedConfirmed": skipped_confirmed,
              "samples": [{"id": f["id"], "changes": f["changes"][:6]} for _sg, f in todo[:12]],
              "abbreviations": rep_["abbreviations"][:60],
              "abbreviationsTotal": len(rep_["abbreviations"]),
              "applied": 0, "stamp": None}
    if req.dry_run or not todo:
        return result
    snapshot = [{**_repair_accept_snapshot(sg), "now": f["now"]} for sg, f in todo]
    try:
        stamp, path = _style_backup(pid, snapshot)
    except Exception as e:
        print(f"[backend] стайл-шит: бэкап не записан: {e}", file=sys.stderr)
        raise HTTPException(500, "Не удалось сохранить копию для отката — применение отменено")
    at = datetime.now().strftime("%Y-%m-%d %H:%M")
    for sg, f in todo:
        was = sg.get("target") or ""
        _replace_target(sg, f["now"], sg.get("provider") or "", "STYLE_SPELLING")
        sg["status"] = "review"
        sg["styleApplied"] = {"from": was, "changes": f["changes"], "by": "human", "at": at}
    _IMPACT_CACHE.pop(pid, None)
    _ANALYSIS_CACHE.pop(pid, None)
    save_state(STATE)
    _audit("style.apply", project=pid, count=len(todo), spelling=rep_["want"], stamp=stamp)
    print(f"[backend] стайл-шит: орфография исправлена в {len(todo)} сегм. (проект {pid}), "
          f"копия: {path.name}", file=sys.stderr)
    result["applied"] = len(todo)
    result["stamp"] = stamp
    return result


@app.post("/api/projects/{pid}/style-check/{stamp}/undo")
def undo_style_check(pid: int, stamp: str):
    """Возвращает ТОЛЬКО те сегменты, где сейчас стоит именно наш текст:
    правили после — чужую работу откатом не затираем."""
    _guard_project_write(pid)
    project = get_project(pid)
    if not re.fullmatch(r"[0-9-]{8,24}", stamp or ""):
        raise HTTPException(400, "Неверная метка отката")
    path = PURGE_DIR / ("style-" + stamp + ".json")
    if not path.exists():
        raise HTTPException(404, "Копия для отката не найдена")
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except Exception as e:
        raise HTTPException(500, "Копия для отката не читается: " + str(e))
    if data.get("project") != pid:
        raise HTTPException(400, "Эта копия относится к другому проекту")
    by_id = {sg["id"]: sg for sg in project["segments"]}
    restored, changed_since = [], []
    for snap in data.get("segments") or []:
        sg = by_id.get(snap["id"])
        if sg is None:
            continue
        if (sg.get("target") or "") != (snap.get("now") or ""):
            changed_since.append(sg["id"])
            continue
        sg["target"] = snap["target"]
        sg["status"] = snap["status"]
        for k in ("provider", "route", "confirmedBy", "confirmedAt", "prevTarget"):
            if snap.get(k) is None:
                sg.pop(k, None)
            else:
                sg[k] = snap[k]
        sg.pop("styleApplied", None)
        restored.append(sg["id"])
    _IMPACT_CACHE.pop(pid, None)
    _ANALYSIS_CACHE.pop(pid, None)
    save_state(STATE)
    _audit("style.undo", project=pid, stamp=stamp, restored=len(restored))
    return {"ok": True, "restored": len(restored), "ids": restored, "changedSince": changed_since}


@app.post("/api/projects/{pid}/term-context/apply")
def apply_term_context(pid: int, req: TermContextApplyRequest):
    """Подставить вариант арбитра вместо утверждённого перевода — по строке,
    а не по записи глоссария.

    Четвёртая команда в системе, меняющая текст БЕЗ вызова модели, и по той же
    причине, что три прежние: она ничего не сочиняет. Вариант предложил
    оплаченный арбитр, единственный, кто видел сегмент в ряду соседей, а корпус
    целевого языка уже проверил его бесплатно (`_corpus_check`) — кальку он
    снял бы до показа. Запись глоссария при этом НЕ трогается: она может быть
    верна в остальных местах документа, а спор про запись остаётся человеку.

    Живёт по правилам массовых команд: `dry_run` по умолчанию, заверенное
    не трогается без `include_confirmed` и названо числом, прежнее состояние
    уходит копией в `data/backups/` и возвращается
    `/term-context/apply/{stamp}/undo`. После подстановки сегмент — `review`:
    проверки описывали прежний текст и устаревают сами, вердикт арбитра тоже
    (его хеш от нового текста не совпадёт) — это честно, текст другой.
    Бесплатно: в `_PAID` не входит."""
    _guard_project_write(pid)
    project = get_project(pid)
    ids = set(req.segment_ids) if req.segment_ids is not None else None
    want = (_norm_key(req.src), _norm_key(req.tgt), _norm_key(req.use))
    matched, skipped_confirmed = [], []
    for seg in project["segments"]:
        if ids is not None and seg["id"] not in ids:
            continue
        adv = next((a for a in _ctx_advices(seg)
                    if (_norm_key(a["src"]), _norm_key(a["tgt"]), _norm_key(a["use"])) == want), None)
        if not adv:
            continue
        if seg.get("status") == "confirmed" and not req.include_confirmed:
            skipped_confirmed.append(seg["id"])
            continue
        matched.append((seg, adv))
    result = {"ok": True, "dryRun": req.dry_run, "matched": len(matched),
              "ids": [sg["id"] for sg, _ in matched],
              "skippedConfirmed": skipped_confirmed,
              "samples": [{"id": sg["id"], "was": (sg.get("target") or "")[:200],
                           "now": _ctx_substitute(sg.get("target") or "", a["tgt"], a["use"])[:200]}
                          for sg, a in matched[:12]],
              "applied": 0, "stamp": None}
    if req.dry_run or not matched:
        return result
    snapshot = [{**_repair_accept_snapshot(sg),
                 "termContext": json.loads(json.dumps(sg.get("termContext") or {}))}
                for sg, _ in matched]
    try:
        PURGE_DIR.mkdir(parents=True, exist_ok=True)
        base = datetime.now().strftime("%Y%m%d-%H%M%S")
        stamp, n = base, 1
        while (PURGE_DIR / ("term-ctx-" + stamp + ".json")).exists():
            stamp = base + "-" + str(n)
            n += 1
        path = PURGE_DIR / ("term-ctx-" + stamp + ".json")
        path.write_text(json.dumps({"project": pid, "segments": snapshot}, ensure_ascii=False),
                        encoding="utf-8")
    except Exception as e:
        print(f"[backend] совет арбитра: бэкап не записан: {e}", file=sys.stderr)
        raise HTTPException(500, "Не удалось сохранить копию для отката — применение отменено")
    for sg, a in matched:
        was = sg.get("target") or ""
        now = _ctx_substitute(was, a["tgt"], a["use"])
        _replace_target(sg, now, sg.get("provider") or "", "TERM_CTX_ACCEPTED")
        sg["status"] = "review"
        # След решения: без него следующий прогон не отличит подстановку
        # по совету от правки руками, а человек — своё решение от машинного.
        sg["termCtxApplied"] = {"src": a["src"], "tgt": a["tgt"], "use": a["use"],
                                "from": was, "by": "human",
                                "at": datetime.now().strftime("%Y-%m-%d %H:%M")}
    _IMPACT_CACHE.pop(pid, None)
    _ANALYSIS_CACHE.pop(pid, None)
    save_state(STATE)
    _audit("term_context.apply", project=pid, count=len(matched), src=req.src, use=req.use, stamp=stamp)
    print(f"[backend] совет арбитра применён: {len(matched)} сегм. (проект {pid}), "
          f"копия: {path.name}", file=sys.stderr)
    result["applied"] = len(matched)
    result["stamp"] = stamp
    return result


@app.post("/api/projects/{pid}/term-context/apply/{stamp}/undo")
def undo_apply_term_context(pid: int, stamp: str):
    """Вернуть тексты до подстановки совета. Только там, где стоит именно
    наш текст: правили после — чужую работу откатом не затираем."""
    project = get_project(pid)
    if not re.fullmatch(r"[0-9-]{8,24}", stamp or ""):
        raise HTTPException(400, "Неверная метка отката")
    path = PURGE_DIR / ("term-ctx-" + stamp + ".json")
    if not path.exists():
        raise HTTPException(404, "Копия для отката не найдена")
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except Exception as e:
        raise HTTPException(500, "Копия для отката не читается: " + str(e))
    if data.get("project") != pid:
        raise HTTPException(400, "Эта копия относится к другому проекту")
    by_id = {sg["id"]: sg for sg in project["segments"]}
    restored, changed_since = [], []
    for snap in data.get("segments") or []:
        seg = by_id.get(snap["id"])
        if seg is None:
            continue
        rec = seg.get("termCtxApplied") or {}
        expect = _ctx_substitute(snap["target"], rec.get("tgt") or "", rec.get("use") or "")
        if rec.get("from") != snap["target"] or (seg.get("target") or "") != expect:
            changed_since.append(seg["id"])
            continue
        seg["target"] = snap["target"]
        seg["status"] = snap["status"]
        for k in ("provider", "route", "confirmedBy", "confirmedAt", "confirmedRole", "prevTarget"):
            if snap.get(k) is None:
                seg.pop(k, None)
            else:
                seg[k] = snap[k]
        seg["repair"] = snap["repair"]
        seg["termContext"] = snap["termContext"]
        seg.pop("termCtxApplied", None)
        restored.append(seg["id"])
    if restored:
        _IMPACT_CACHE.pop(pid, None)
        _ANALYSIS_CACHE.pop(pid, None)
        save_state(STATE)
    return {"ok": True, "restored": restored, "changedSince": changed_since}


class RepairBatchRequest(BaseModel):
    segment_ids: Optional[List[int]] = None
    limit: int = 5
    retry: bool = False          # чинить и то, что уже пытались чинить
    # Подтверждённые сегменты пакетный ремонт не трогает: он переписывает текст,
    # а молча переписать заверенный человеком перевод нельзя — то же правило,
    # что у переперевода по глоссарию. Одиночный ремонт по кнопке на сегменте
    # остаётся явным выбором и разрешён.
    include_confirmed: bool = False
    model: Optional[str] = None
    bc_model: Optional[str] = None
    tc_model: Optional[str] = None
    use_judge: bool = False
    judge_model: Optional[str] = None
    # См. RepairRequest: перепроверка судится тем же правилом, что и прежний балл.
    judge_all: bool = False


@app.post("/api/projects/{pid}/repair/batch")
def repair_batch(pid: int, req: RepairBatchRequest):
    """Порция маленькая (5): на сегмент уходит вызов ремонта плюс перепроверка,
    это самый дорогой прогон в системе."""
    if not os.environ.get("OPENAI_API_KEY"):
        raise HTTPException(503, "Ремонт требует ключ OpenAI")
    project = get_project(pid)
    id_filter = set(req.segment_ids) if req.segment_ids is not None else None
    # Разрешаем модель ТЕМ ЖЕ выражением, что _plan_step и сам заход: клеймо
    # «уже чинилось» смотрит, ходила ли на отпечаток ИМЕННО эта модель, и две
    # формулы разрешения дали бы смете один состав, а прогону другой.
    rp_mdl = _resolve_model(req.model or REPAIR_DEFAULT_MODEL)["id"]
    candidates = [s for s in project["segments"]
                  if (id_filter is None or s["id"] in id_filter)
                  and (req.include_confirmed or s.get("status") != "confirmed"
                       or _confirm_override(s))
                  and _repairable(s, req.retry, project, rp_mdl)]
    # Пощада названа поимённо и БЕЗ тех, кого забрала объективная находка:
    # иначе отчёт говорил бы «не тронули», а текст был бы переписан.
    skipped_confirmed = ([s["id"] for s in project["segments"]
                          if (id_filter is None or s["id"] in id_filter)
                          and s.get("status") == "confirmed"
                          and not _confirm_override(s)
                          and _repairable(s, req.retry, project, rp_mdl)]
                         if not req.include_confirmed else [])
    limit = max(1, min(req.limit, 30))
    remaining_after = max(0, len(candidates) - limit)
    targets = candidates[:limit]

    applied, skipped, errors, desync = [], [], [], []

    def _repair_one(seg):
        if _job_should_stop():
            return {"seg": seg, "skip": True}
        try:
            r = _run_segment_repair(seg, project, req.model, req.bc_model, req.tc_model,
                                    req.use_judge, req.judge_model, req.judge_all)
            return {"seg": seg, "res": r}
        except Exception as e:
            print(f"[backend] repair batch seg#{seg['id']}: {e}", file=sys.stderr)
            return {"seg": seg, "error": str(e)}

    for out in _run_parallel(targets, _repair_one):
        seg = out["seg"]
        if out.get("skip"):
            continue
        if out.get("error"):
            errors.append({"id": seg["id"], "error": out["error"]})
            continue
        r = out["res"]
        if r.get("desync"):
            desync.append(seg["id"])
        if not r.get("ok"):
            errors.append({"id": seg["id"], "error": r.get("error", "unknown")})
        elif r.get("applied"):
            applied.append(seg["id"])
        else:
            skipped.append({"id": seg["id"], "reason": (r.get("repair") or {}).get("reason", "")})
    save_state(STATE)
    return {"ok": True, "applied": applied, "count": len(applied), "skipped": skipped,
            "remaining": remaining_after, "errors": errors,
            # Молчаливых потолков не бывает: подтверждённые, которые есть что чинить,
            # называем поимённо, а не выбрасываем как будто их не было.
            "skipped_confirmed": skipped_confirmed,
            # Сегменты, где текст разошёлся с записью о решении. Ноль — норма,
            # не ноль — повод смотреть журнал, а не гадать по остывшим данным.
            "desync": desync,
            "model": _resolve_model(req.model or REPAIR_DEFAULT_MODEL)["id"]}


class BackcheckRequest(BaseModel):
    model: Optional[str] = None
    use_judge: bool = False
    judge_model: Optional[str] = None


@app.post("/api/segments/{pid}/{sid}/backcheck")
def backcheck_segment(pid: int, sid: int, req: BackcheckRequest = BackcheckRequest()):
    """Обратный перевод target → язык оригинала + оценка соответствия.
    Обычный def: внутри блокирующий вызов модели (см. batch_translate)."""
    seg = get_segment(pid, sid)
    project = get_project(pid)
    result = _run_segment_backcheck(seg, project, req.model, req.use_judge, req.judge_model)
    if result.get("ok"):
        save_state(STATE)
        result["segment"] = seg
    return result


class BackcheckBatchRequest(BaseModel):
    segment_ids: Optional[list] = None
    limit: int = 10
    model: Optional[str] = None
    skip_cached: bool = True     # не пересчитывать сегменты с неизменившимся переводом
    use_judge: bool = False
    judge_model: Optional[str] = None
    # Разовое разрешение звать судью и выше потолка зоны (см. _judge_pending).
    # Поле запуска, как include_confirmed: не настройка и не новая политика.
    judge_all: bool = False


@app.post("/api/projects/{pid}/backcheck/batch")
def backcheck_batch(pid: int, req: BackcheckBatchRequest):
    """Пакетный back-check. Порционный, как и пакетный перевод: клиент гоняет
    порции по 10, поэтому один запрос не живёт дольше таймаута прокси."""
    if not os.environ.get("OPENAI_API_KEY"):
        # 503 отдаём до работы, как это делают termcheck и ремонт: иначе
        # обратный перевод вырождается в посегментные ошибки, и составной
        # прогон принимает их за «порция целиком провалилась».
        raise HTTPException(503, "Back-check требует ключ OpenAI")
    project = get_project(pid)
    id_filter = set(req.segment_ids) if req.segment_ids is not None else None
    mdl_id = _resolve_model(req.model or BACKCHECK_DEFAULT_MODEL)["id"]

    candidates = []
    skipped_cached = 0
    for s in project["segments"]:
        if id_filter is not None and s["id"] not in id_filter:
            continue
        if not (s.get("target") or "").strip():
            continue
        if req.skip_cached and _backcheck_cached(s, mdl_id, req.use_judge,
                                                 req.judge_all):
            skipped_cached += 1
            continue
        candidates.append(s)

    limit = max(1, min(req.limit, 100))
    remaining_after = max(0, len(candidates) - limit)
    targets = candidates[:limit]

    processed, errors = [], []
    # Одинаковая пара «оригинал + перевод» даёт одинаковый результат: считаем раз,
    # а сами уникальные пары проверяем параллельно.
    groups: dict = {}
    order: list = []
    for seg in targets:
        pair = (_norm_key(seg.get("source")), _norm_key(seg.get("target")))
        if pair not in groups:
            groups[pair] = []
            order.append(pair)
        groups[pair].append(seg)
    duplicates = 0

    def _bc_one(pair):
        seg = groups[pair][0]
        if _job_should_stop():
            return {"pair": pair, "skip": True}
        try:
            r = _run_segment_backcheck(seg, project, req.model, req.use_judge,
                                       req.judge_model, judge_all=req.judge_all)
            return {"pair": pair, "ok": bool(r.get("ok")), "error": r.get("error")}
        except Exception as e:
            print(f"[backend] backcheck batch seg#{seg['id']}: {e}", file=sys.stderr)
            return {"pair": pair, "ok": False, "error": str(e)}

    for res in _run_parallel(order, _bc_one):
        segs = groups[res["pair"]]
        if res.get("skip"):
            continue
        if not res.get("ok"):
            errors.append({"id": segs[0]["id"], "error": res.get("error", "unknown")})
            continue
        processed.append(segs[0]["id"])
        for sg in segs[1:]:
            sg["backtranslated_ru"] = segs[0].get("backtranslated_ru", "")
            sg["backcheck"] = json.loads(json.dumps(segs[0]["backcheck"]))
            processed.append(sg["id"])
            duplicates += 1
    save_state(STATE)
    return {
        "ok": True,
        "processed": processed,
        "count": len(processed),
        "remaining": remaining_after,
        "skipped_cached": skipped_cached,
        "duplicates": duplicates,
        "errors": errors,
        "model": mdl_id,
    }


class ChecksRequest(BaseModel):
    run_backcheck: bool = True
    bc_model: Optional[str] = None


def _segment_checks(pid: int, sid: int, run_backcheck: bool = True,
                        bc_model: Optional[str] = None) -> dict:
    if not checks_mod:
        raise HTTPException(500, "Модуль проверок недоступен")

    seg = get_segment(pid, sid)
    project = get_project(pid)
    source_text = seg.get("source", "")
    target_text = seg.get("target", "").strip()
    if not target_text:
        return {"ok": False, "error": "Segment is not translated yet", "segment": seg}

    gloss_hits, tm_hit = _get_context(source_text, project=project)
    back = seg.get("backtranslated_ru", "")

    if run_backcheck and checks_enabled():
        fresh = seg.get("backcheck") or {}
        if fresh.get("back") and fresh.get("target_hash") == _text_hash(target_text):
            # Обратный перевод для этого же текста уже есть — второй раз не платим
            back = fresh["back"]
        else:
            try:
                # literal=True обязателен: обычный промпт «чинит» кривой
                # английский на лету и прячет ошибку, которую QA ищет.
                #
                # Модель — та же, что у back-check, и это не косметика. Без
                # явного указания сюда подставлялась модель перевода по
                # умолчанию: тот же самый вызов стоил впятеро дороже нужного
                # и делался моделью, которую для обратного перевода никто
                # не выбирал. У этой работы одна правильная модель — самая
                # буквальная и дешёвая, та же, что у back-check. Совпадение
                # с автором текста разрешается тем же _backcheck_model.
                back = _openai_translate(target_text, project["tgt"], project["src"],
                                         model=_backcheck_model(seg, bc_model),
                                         literal=True, step="medical_qa")
            except Exception as e:
                print(f"[backend] medical QA backcheck skipped: {e}", file=sys.stderr)

    qa_result = checks_mod.run_checks(
        source_text,
        target_text,
        backtranslated_ru=back,
        glossary_matches=gloss_hits,
        tm_match=tm_hit,
        engine_qa="medical_qa_mvp",
        domain=project.get("domain"),
        src_lang=project.get("src", "RU"),
        tgt_lang=project.get("tgt", "EN"),
    )

    seg["backtranslated_ru"] = qa_result["literal_backcheck"]["backtranslated_ru"]
    # Хеш текста — как у back-check и termcheck: по нему повторный прогон
    # понимает, что пересчитывать нечего. Без него Medical QA брала В КАЖДЫЙ
    # прогон все переведённые сегменты, и составная кнопка всегда показывала
    # полный проект, сколько бы раз её ни нажимали.
    #
    # Но ставим отметку ТОЛЬКО если проверка была полной: без обратного перевода
    # часть находок не считается вовсе, и закэшировать такой результат значит
    # закрыть сегмент от нормальной проверки навсегда.
    if (back or "").strip():
        qa_result["target_hash"] = _text_hash(target_text)
    seg["qa_result"] = qa_result
    seg["qa_issues"] = qa_result["qa_issues"]
    seg["qa"] = qa_result["ui_issues"]
    seg["term_candidates"] = qa_result["term_candidates"]
    seg["risk_score"] = qa_result["risk_score"]
    seg["risk_color"] = qa_result["risk_color"]
    seg["engine_qa"] = qa_result["engine_qa"]
    seg["medical_qa_enabled"] = checks_enabled()

    # Статус проверка может только ПОВЫСИТЬ до review, но не понизить.
    #
    # Подтверждённый не трогаем: она не изменила ни буквы текста, а разжаловать
    # решение человека молча — то же самое, что переписать его перевод.
    # Находка видна по risk_color и qa_issues, подтверждённые с замечаниями
    # показывает вкладка QA.
    #
    # review не трогаем по той же причине с другой стороны: этот статус ставят
    # те, кто ПЕРЕПИСАЛ текст (ремонт, переперевод, правка поверх заверенного),
    # и значит он «на текст смотрела машина, посмотри человек». Medical QA
    # проверяет числа и отрицания, а не то, читал ли кто-то результат правки;
    # её «замечаний нет» — не ответ на этот вопрос. Понижая review до qa, она
    # уводила починенные сегменты из зоны внимания на доске и в фильтрах.
    if seg.get("status") in ("translated", "qa"):
        seg["status"] = "review" if qa_result["risk_color"] == "red" else "qa"
    # seg["risk"] не трогаем. Это длина сегмента, по ней выбирается движок
    # перевода (low → Google); медицинскую опасность несут risk_score и
    # risk_color. Раньше проверка перетирала одно другим, и сегмент, чисто
    # прошедший QA, при следующем переводе уезжал в Google вместо GPT.

    return {"ok": True, "segment": seg, "qa_result": qa_result, "issues": qa_result["qa_issues"]}


# Путь назывался medical-qa; новое имя нейтрально, старое оставлено ещё
# релиз — по нему ходит вкладка, открытая до выката.
@app.post("/api/segments/{pid}/{sid}/checks")
@app.post("/api/segments/{pid}/{sid}/medical-qa")
def checks_segment(pid: int, sid: int, req: ChecksRequest = ChecksRequest()):
    _guard_project_write(pid)
    result = _segment_checks(pid, sid, run_backcheck=req.run_backcheck,
                                 bc_model=req.bc_model)
    save_state(STATE)
    return result


class ChecksBatchRequest(BaseModel):
    limit: int = 50
    segment_ids: Optional[list] = None
    run_backcheck: bool = True
    # Модель обратного перевода — та же, что у back-check. Своей модели у
    # Medical QA нет: правила детерминированные, вызов нужен только чтобы
    # получить обратный перевод, когда готового от back-check не осталось.
    bc_model: Optional[str] = None
    # Пропускать сегменты, чей результат относится к нынешнему тексту.
    skip_cached: bool = True


@app.post("/api/projects/{pid}/checks/batch")
@app.post("/api/projects/{pid}/medical-qa/batch")
def batch_checks(pid: int, req: ChecksBatchRequest = ChecksBatchRequest()):
    project = get_project(pid)
    id_filter = set(req.segment_ids) if req.segment_ids else None
    candidates = [
        s for s in project["segments"]
        if s.get("target", "").strip()
        and s.get("status") in {"translated", "qa", "review", "confirmed"}
        and (id_filter is None or s["id"] in id_filter)
    ]
    # Свежую проверку второй раз не считаем. Экономит не только процессор:
    # Medical QA заказывает обратный перевод, если back-check по этому тексту
    # устарел, — то есть повторный прогон по всему проекту мог стоить денег.
    skipped_cached = 0
    if req.skip_cached:
        before = len(candidates)
        candidates = [s for s in candidates
                      if _check_stale(s.get("qa_result"), s.get("target"))]
        skipped_cached = before - len(candidates)
    targets = candidates[:req.limit]
    processed = []
    errors = []
    def _qa_one(seg):
        if _job_should_stop():
            return {"seg": seg, "skip": True}
        try:
            return {"seg": seg, "res": _segment_checks(pid, seg["id"], run_backcheck=req.run_backcheck,
                                                           bc_model=req.bc_model)}
        except Exception as e:
            print(f"[backend] medical QA batch error seg#{seg['id']}: {e}", file=sys.stderr)
            return {"seg": seg, "error": str(e)}

    for out in _run_parallel(targets, _qa_one):
        seg = out["seg"]
        if out.get("skip"):
            continue
        if out.get("error"):
            errors.append({"id": seg["id"], "error": out["error"]})
        elif out["res"].get("ok"):
            processed.append(seg["id"])
        else:
            errors.append({"id": seg["id"], "error": out["res"].get("error", "unknown")})

    save_state(STATE)
    return {
        "ok": True,
        "processed": processed,
        "count": len(processed),
        "remaining": max(0, len(candidates) - len(targets)),
        "errors": errors,
        "skipped_cached": skipped_cached,
        "featureEnabled": checks_enabled(),
    }


@app.post("/api/segments/{pid}/{sid}/revert")
def revert_segment(pid: int, sid: int):
    _guard_project_write(pid)
    seg = get_segment(pid, sid)
    if seg["status"] == "confirmed":
        _withdraw_confirmation(seg, "revert")
        _audit("segment.unconfirm", project=pid, segment=sid)
    elif seg["status"] == "failed":
        seg["status"] = "new"
        seg["target"] = ""
    save_state(STATE)
    return {"ok": True, "segment": seg}


def _note_hand_edit(seg: dict, new_target: str) -> None:
    """Запомнить БАЗУ правки человека — текст, что стоял до его руки.

    Диф «база → подтверждённый текст» — сырьё для `_harvest_edited_terms`.
    Жива ли база, решает ЦЕПОЧКА ХЕШЕЙ, а не чистка по всем машинным путям
    записи: target пишут и ремонт, и пакетный перевод, и undo-восстановления —
    мимо `_replace_target`, и забытая чистка в любом из них приписала бы
    человеку правки машины. `editedToHash` — «последний текст, который писал
    человек»: совпал с тем, что лежит в сегменте, — цепочка правок
    не прерывалась, база жива; разошёлся — между правками писала машина,
    и базой становится ЕЁ свежий текст. Перевод, набранный с нуля
    (прежний target пуст), исправлением не считается: диффать его не с чем."""
    prev = seg.get("target") or ""
    if seg.get("editedFrom") and seg.get("editedToHash") == _text_hash(prev):
        seg["editedToHash"] = _text_hash(new_target or "")
    elif prev.strip():
        seg["editedFrom"] = prev
        seg["editedToHash"] = _text_hash(new_target or "")
    else:
        seg.pop("editedFrom", None)
        seg.pop("editedToHash", None)


class UpdateSegmentRequest(BaseModel):
    target: Optional[str] = None
    status: Optional[str] = None
    comment: Optional[str] = None
    commentAuthor: Optional[dict] = None

@app.post("/api/segments/{pid}/{sid}/update")
def update_segment(pid: int, sid: int, req: UpdateSegmentRequest):
    _guard_project_write(pid)
    seg = get_segment(pid, sid)
    want_confirm = req.status == "confirmed"
    if want_confirm and seg.get("status") != "confirmed":
        # Заверение — только командой /confirm: она ставит подпись
        # ответственного и проверяется по роли. Статус без подписи — это
        # «подтвердил кто-то», след, за который никто не отвечает.
        # На УЖЕ заверенном сегменте тот же статус в теле — не просьба
        # заверить, а копия браузера (store шлёт статус вместе с текстом).
        raise HTTPException(400, "Заверение ставится командой «Подтвердить» — она оставляет подпись ответственного")
    changing = req.target is not None and req.target != seg.get("target")
    if seg.get("status") == "confirmed" and (changing or (req.status and not want_confirm)):
        # Правка заверенного текста и уход из статуса снимают подпись: она
        # относилась к тексту, которого больше нет. Прежний текст — в
        # prevTarget, как у `_replace_target`: заверенное не пропадает молча.
        if changing:
            seg["prevTarget"] = seg.get("target", "")
        _withdraw_confirmation(seg, "edit" if changing else "revert")
        _audit("segment.unconfirm", project=pid, segment=sid)
    if req.target is not None:
        if changing:
            _audit("segment.edit", project=pid, segment=sid)
            _note_hand_edit(seg, req.target)
            # Кто правил руками — след на сегменте, не только в журнале:
            # журнал кольцевой, а вопрос «чья это формулировка» задают
            # про сегмент.
            seg["editedBy"] = _actor_id()
            seg["editedAt"] = datetime.now().strftime("%Y-%m-%d %H:%M")
        seg["target"] = req.target
        if seg["status"] == "new" and req.target.strip():
            seg["status"] = "translated"
    if req.status and not want_confirm:
        seg["status"] = req.status
    if req.comment:
        seg.setdefault("comments", []).append({
            "author": req.commentAuthor or {"name": "Вы", "initials": "ВЫ", "color": "#2c7be5"},
            "when": "только что",
            "text": req.comment,
        })
    save_state(STATE)
    return {"ok": True, "segment": seg}


def _needs_translation(seg: dict) -> bool:
    """Сегмент ждёт перевода: его возьмёт пакет БЕЗ force.

    Один предикат на три места — сам пакет (`batch_translate`), разбор шага
    перевода (`_plan_step`) и множество «переведём в этом же прогоне»
    (`will_translate` в `run_plan`, от которого зависит, попадут ли эти
    сегменты в состав ПРОВЕРОК). Три копии условия разошлись бы, и разойтись
    им хватило бы одного дня: смета обещала бы одно, а прогон делал другое —
    ровно та беда, ради которой состав вынесен на сервер.

    `failed` с ПУСТЫМ переводом — это «не переведён»: ошибка перевода
    сегмент не трогает (инвариант №4), и без повтора он остался бы в корзине
    «возьмёт прогон» навсегда, а прогон честно отвечал бы «нечего делать».
    `failed` с НЕПУСТЫМ переводом не берём: там лежит прежний текст, и его
    судьба — вопрос проверок, а не молчаливого переперевода."""
    st = seg.get("status")
    return st == "new" or (st == "failed" and not (seg.get("target") or "").strip())


class BatchRequest(BaseModel):
    # Движок один — выбранная модель, поэтому engine и low_engine убраны.
    # Заодно исчез отбор по risk: он существовал только чтобы делить сегменты
    # между Google и моделью, и оставлял половину проекта непереведённой,
    # если запустить не ту кнопку.
    limit: int = 50                      # максимум за один вызов
    segment_ids: Optional[list] = None  # если передан — обрабатывать только эти сегменты
    force: bool = False                  # True = явный выбор пользователя, пропустить фильтры статуса и риска
    model: Optional[str] = None          # id из OPENAI_MODELS; неизвестный → DEFAULT_OPENAI_MODEL
    include_confirmed: bool = False      # переписывать и подтверждённые (только по явной галочке)

@app.post("/api/projects/{pid}/batch")
def batch_translate(pid: int, req: BatchRequest):
    # ВАЖНО: обычный def, а не async def. Внутри блокирующие вызовы OpenAI/Google;
    # в async def они вешали единственный event loop uvicorn на всё время пакета —
    # сервер не отвечал даже на GET /api/projects/{pid} сразу после батча.
    project = get_project(pid)
    # is not None, а не truthy: пустой список — это «не выбрано ничего», и переводить
    # в этом случае надо ноль сегментов, а не весь проект.
    id_filter = set(req.segment_ids) if req.segment_ids is not None else None
    if req.force and id_filter is not None:
        # Явный выбор — переводим только указанные сегменты. Подтверждённые
        # мимо, пока человек не попросил отдельной галочкой: молча переписать
        # заверенный перевод нельзя. Попросил — переписываем, но с сохранением
        # прежнего текста и со статусом «требует проверки» (см. ниже).
        all_targets = [s for s in project["segments"] if s["id"] in id_filter
                       and (req.include_confirmed or s["status"] != "confirmed")]
        skipped_confirmed = ([s["id"] for s in project["segments"]
                              if s["id"] in id_filter and s["status"] == "confirmed"]
                             if not req.include_confirmed else [])
    else:
        # Предикат общий с разбором прогона — см. _needs_translation.
        all_targets = [s for s in project["segments"]
                       if _needs_translation(s)
                       and (id_filter is None or s["id"] in id_filter)]
        skipped_confirmed = []
    done_by_src: dict = {}
    if not req.force:
        for s0 in project["segments"]:
            t0 = (s0.get("target") or "").strip()
            # Тот же белый список, что и у остальных потребителей перевода:
            # сегмент со статусом failed не прошёл проверку, и копировать его
            # текст в близнецов — значит размножить брак.
            if not t0 or s0.get("status") not in ("translated", "qa", "review", "confirmed"):
                continue
            key0 = _norm_key(s0.get("source"))
            prev = done_by_src.get(key0)
            # Донором становится лучший, а не первый по порядку: заверенный
            # человеком перевод сильнее машинного, иначе старая ошибка
            # переехала бы в новые сегменты поверх уже исправленного близнеца.
            if prev is None or (_confirmed_by_human(s0) and not prev[1]):
                done_by_src[key0] = (t0, _confirmed_by_human(s0),
                                     s0.get("provider"))

    # Ключ нужен, только если хоть один сегмент придётся ПЕРЕВОДИТЬ. Порция из
    # одних совпадений с памятью или из повторов уже переведённого обходится
    # без вызовов модели — отказывать ей 503 значило бы запретить работу,
    # которая ничего не стоит. Проверка стоит здесь: до отбора состав неизвестен.
    def _free(sg):
        if _tm_trusted(_get_context(sg["source"], project=project)[1]):
            return True
        return _norm_key(sg["source"]) in done_by_src
    if (all_targets and not os.environ.get("OPENAI_API_KEY")
            and not all(_free(s) for s in all_targets)):
        # 503 до работы, как у termcheck, back-check и ремонта: иначе отсутствие
        # ключа вырождается в посегментные ошибки, и составной прогон принимает
        # их за «порция целиком провалилась» вместо «шаг недоступен».
        raise HTTPException(503, "Перевод требует ключ OpenAI")
    # Потолок на порцию: один HTTP-запрос не должен жить дольше proxy_read_timeout (1800s)
    # в nginx. При ~5-6 с на сегмент 100 штук — это ~10 минут, с большим запасом.
    limit = max(1, min(req.limit, 100))
    remaining_after = max(0, len(all_targets) - limit)
    targets = all_targets[:limit]
    translated = []
    tm_hits_count = 0
    dup_hits_count = 0
    errors = []
    # Повторы внутри порции переводим один раз: в документах один и тот же
    # заголовок встречается десятки раз, и каждый стоил отдельного вызова.
    # Группируем ДО вызовов, чтобы уникальные тексты можно было гнать разом.
    groups: dict = {}
    order: list = []
    for seg in targets:
        key = _norm_key(seg["source"])
        if key not in groups:
            groups[key] = []
            order.append(key)
        groups[key].append(seg)

    # Готовые переводы одинаковых исходников ПО ВСЕМУ проекту. Дедупликация
    # внутри порции ловит только соседей: один и тот же заголовок, встретившийся
    # в сегментах 12 и 500, попадает в разные порции и оплачивается дважды.
    # Проект уже переведён наполовину — брать оттуда бесплатно.

    def _translate_group(key):
        """Один вызов на группу одинаковых исходников. Возвращает результат,
        применяет его основной поток — так мутации STATE остаются в одном месте."""
        seg = groups[key][0]
        if _job_should_stop():
            return {"key": key, "skip": True}
        gloss_hits, tm_hit = _get_context(seg["source"], project=project)
        gloss_hits = gloss_hits + _doc_hits(seg["source"], project, gloss_hits)

        # TM точное совпадение → пропускаем API вызов.
        # При force (явный выбор пользователя: галочки или «перевести заново») шорткат
        # не применяем — иначе «перевести заново выбранной моделью» молча подставляло бы
        # старый текст из памяти. Так же ведёт себя одиночный перевод сегмента.
        # Память переводов идёт ПЕРЕД повтором внутри проекта: её записи человек
        # подтверждал, а близнец в проекте мог быть переведён машиной и с тех пор
        # исправлен только в памяти.
        if not req.force and _tm_trusted(tm_hit) and tm_hit.get("tgt"):
            return {"key": key, "tm": True, "text": tm_hit["tgt"]}

        # Такой текст в проекте уже переведён — платить за него второй раз
        # не за что. force не трогаем: «перевести заново» должно переводить.
        ready = done_by_src.get(key)
        if ready:
            return {"key": key, "reuse": True, "text": ready[0], "provider": ready[2]}

        # Соседи берутся у ПЕРВОГО сегмента группы. Группа — это сегменты
        # с одинаковым исходником, и перевод у них будет один на всех
        # (см. дедупликацию повторов): выбрать «правильных» соседей для
        # одинакового заголовка, стоящего в трёх местах документа, нельзя
        # в принципе. Берём обстановку донора — это честнее, чем не брать
        # никакой, и ровно так же ведёт себя перенос перевода на близнецов.
        prev_src, next_src = _neighbours(project, seg)
        try:
            translation = _openai_translate(seg["source"], project["src"], project["tgt"],
                                            domain=project.get("domain"),
                                            style=_style_block(project),
                                            gloss_hits=gloss_hits, tm_context=tm_hit,
                                            model=req.model,
                                            prev_src=prev_src, next_src=next_src)
        except Exception as e:
            print(f"[backend] batch error seg#{seg['id']}: {e}", file=sys.stderr)
            return {"key": key, "error": str(e)}
        if not translation:
            return {"key": key, "error": "модель не вернула перевод"}
        # След для замера: какие пары терм-листа ушли в ЭТОТ промпт.
        return {"key": key, "text": translation, "provider": _resolve_model(req.model)["id"],
                "docTerms": [h["tgt"] for h in gloss_hits if h.get("tier") == "doc"]}

    for res in _run_parallel(order, _translate_group):
        segs = groups[res["key"]]
        if res.get("skip"):
            continue
        if res.get("error"):
            errors.extend(sg["id"] for sg in segs)
            continue
        if res.get("tm") or res.get("reuse"):
            for sg in segs:
                # translated, а не confirmed — см. translate_segment
                sg["target"] = res["text"]
                sg["status"] = "translated"
                sg.pop("docTerms", None)   # текст из памяти или копия: промпта с терм-листом не было
                sg["route"] = "EXACT_TM" if res.get("tm") else "DUPLICATE"
                # Копия наследует провайдера донора: писать «tm» на тексте,
                # взятом у соседнего сегмента, значит соврать о происхождении.
                sg["provider"] = PROVIDER_TM if res.get("tm") else (res.get("provider") or "")
                translated.append(sg["id"])
            if res.get("tm"):
                tm_hits_count += len(segs)
            else:
                dup_hits_count += len(segs)
            continue
        for i, sg in enumerate(segs):
            if res.get("docTerms"):
                sg["docTerms"] = list(res["docTerms"])
            else:
                sg.pop("docTerms", None)
            # Переписываем заверенный человеком перевод — сохраняем прежний текст
            # и снимаем отметку о подтверждении: статус «требует проверки», а не
            # «подтверждено». Машина не заверяет сама себя, и «подтвердил человек»
            # не должно оставаться на строке, которой человек не видел.
            was_confirmed = sg.get("status") == "confirmed"
            if was_confirmed:
                sg["prevTarget"] = sg.get("target", "")
                sg.pop("confirmedBy", None)
                sg.pop("confirmedAt", None)
                sg.pop("confirmedRole", None)
            sg["target"] = res["text"]
            sg["status"] = "review" if was_confirmed else "translated"
            sg["provider"] = res["provider"]
            sg["route"] = "GPT_REQUIRED" if i == 0 else "DUPLICATE"
            translated.append(sg["id"])
        dup_hits_count += len(segs) - 1
    save_state(STATE)
    return {
        "ok": True,
        "translated": translated,
        "count": len(translated),
        "remaining": remaining_after,
        "errors": errors,
        "tm_hits": tm_hits_count,
        "duplicates": dup_hits_count,
        # Молчаливых потолков не бывает: сегменты, отсеянные как подтверждённые,
        # называем поимённо — иначе прогон рапортует «готово» и ничего не делает.
        "skipped_confirmed": skipped_confirmed,
        "model": _resolve_model(req.model)["id"],
    }


@app.post("/api/projects/{pid}/preflight")
def run_preflight(pid: int):
    project = get_project(pid)
    segs = project["segments"]
    total = len(segs)

    # Assign risk + route to every segment that lacks them
    for s in segs:
        if not s.get("risk"):
            words = len(s["source"].split())
            s["risk"] = "high" if words > 30 else "medium" if words > 8 else "low"
        if not s.get("route"):
            s["route"] = "GPT_REQUIRED"
        if s.get("tm") is None:
            s["tm"] = None

    save_state(STATE)

    routes: dict = {}
    for s in segs:
        r = s.get("route", "GPT_REQUIRED")
        routes[r] = routes.get(r, 0) + 1

    risks: dict = {}
    for s in segs:
        r = s.get("risk", "medium")
        risks[r] = risks.get(r, 0) + 1

    return {
        "ok": True,
        "totalSegments": total,
        "routes": routes,
        "risks": risks,
        "analysisTime": round(total * 0.045 + 0.6, 1),
    }


class ExportRequest(BaseModel):
    format: str          # "docx" | "xlsx" ("pdf" пока не поддерживается)
    source: bool = True  # включить колонку с оригиналом

EXPORT_DIR = DATA_DIR / "exports"   # внутри ReadWritePaths systemd-юнита

def _safe_filename(name: str) -> str:
    return _re.sub(r'[\\/:*?"<>|\x00-\x1f]+', "_", name).strip() or "project"

# ─── Экспорт «как в оригинале» ──────────────────────────────────────
# Документ не собирается заново, а открывается исходный и в нём подменяется
# текст. Всё, чего мы не тронули, остаётся байт в байт: шрифты, кегли,
# картинки, таблицы, колонтитулы, нумерация, разметка страницы. Собрать это
# из сегментов невозможно — в сегментах нет ни шрифта, ни картинок.
_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# Что видно глазом. Всё, чего здесь нет, — служебное: w:lang (каким языком
# проверять орфографию), w:noProof, w:bCs/w:iCs/w:szCs (сложные письменности),
# w:rFonts cs/eastAsia. Word ставит их сам, кусками, по ходу правки: в учебнике
# фтизиатрии 255 абзацев из 678 «разного оформления» различались ТОЛЬКО этим.
# Считать их разными значит резать перевод там, где резать нечего.
#
# w:spacing, w:kern, w:position и w:w тоже не в счёт, и это не небрежность:
# это микроподгонка межбуквенного расстояния под ширину строки (значения
# вроде -5 двадцатых пункта). Строка после перевода всё равно другой длины,
# и переносить такую подгонку не только незачем, но и вредно — ещё 96 абзацев.
_VIS_RPR = {"b", "i", "u", "strike", "dstrike", "color", "sz", "highlight",
            "vertAlign", "caps", "smallCaps", "rStyle", "shd", "em",
            "outline", "shadow", "imprint", "emboss"}


def _vis_sig(rpr) -> tuple:
    """Отпечаток ВИДИМОГО оформления прогона.

    Своими руками, а не сериализацией XML: lxml пришлось бы тащить в модуль,
    который умеет стартовать и без python-docx (тогда просто нет экспорта,
    а не падения при импорте)."""
    if rpr is None:
        return ()
    out = []
    for ch in rpr:
        name = ch.tag[len(_W_NS):] if ch.tag.startswith(_W_NS) else ch.tag
        if name == "rFonts":
            fonts = (ch.get(_W_NS + "ascii"), ch.get(_W_NS + "hAnsi"))
            if any(fonts):
                out.append(("rFonts", fonts))
            continue
        if name not in _VIS_RPR:
            continue
        val = ch.get(_W_NS + "val")
        if val in ("0", "false", "none"):
            continue          # <w:b w:val="0"/> — это «не жирный», а не отметка
        out.append((name, val))
    return tuple(sorted(out, key=repr))


# Поля, чей результат считает САМ движок вёрстки: номер страницы, дата,
# счётчики, зеркала чужого текста. Перевод, вписанный туда, либо ломает
# оглавление и нумерацию, либо стирается при первом же обновлении поля.
# Всё ОСТАЛЬНОЕ переводится: результат поля TOC — обычный видимый текст,
# и строка оглавления, несущая метку поля, оставалась единственной русской
# строкой в английском документе.
FIELD_COMPUTED = {
    "PAGE", "PAGEREF", "NUMPAGES", "SECTIONPAGES", "SECTION", "SEQ", "REF",
    "STYLEREF", "NOTEREF", "AUTONUM", "AUTONUMLGL", "AUTONUMOUT", "LISTNUM",
    "DATE", "TIME", "CREATEDATE", "SAVEDATE", "PRINTDATE", "EDITTIME",
    "REVNUM", "FILENAME", "FILESIZE", "NUMWORDS", "NUMCHARS", "AUTHOR",
    "USERNAME", "LASTSAVEDBY", "DOCPROPERTY", "TEMPLATE", "INFO", "TITLE",
    "SUBJECT", "KEYWORDS", "COMMENTS", "DOCVARIABLE",
}
# Их же (плюс сами таблицы-указатели) просим пересчитать при открытии.
FIELD_REFRESH = FIELD_COMPUTED | {"TOC", "TOA", "INDEX", "RD", "XE", "TC"}


_FIELD_SWITCH = chr(92)          # обратная косая: с неё начинаются ключи поля


def _field_key(instr: str) -> str:
    """Имя поля из инструкции: « PAGEREF _Toc219883320 [ключ]h » → «PAGEREF».

    Ключи-переключатели начинаются с обратной косой и именем поля не бывают."""
    for w in (instr or "").split():
        if w and w[0] != _FIELD_SWITCH:
            return w.upper()
    return ""


def _para_slots(p_elem, qn) -> tuple:
    """(куда можно писать, весь текст абзаца, пропущенный текст).

    Первое — список пар «узел w:t, отпечаток видимого оформления его прогона»
    в порядке документа. Пропускаем три вещи, и каждая не мелочь:

    * ВЫЧИСЛЯЕМЫЕ ПОЛЯ (`FIELD_COMPUTED`). В строке оглавления там лежит номер
      страницы, а в колонтитуле — номер текущей. Считает их движок вёрстки,
      и записать туда перевод значит сломать и оглавление, и нумерацию.
      А вот результат ОСТАЛЬНЫХ полей — обычный видимый текст, и он
      переводится: у поля TOC результат растянут на десятки абзацев, метку
      несёт только первый, и без этого он один оставался на языке оригинала.
    * СКРЫТЫЙ текст (w:webHidden, w:vanish) — его в документе не видно,
      и перевод, попавший туда, не увидит никто.
    * ВЛОЖЕННЫЕ абзацы: надпись или таблица внутри абзаца — это свои w:p
      со своими якорями. Захватить их значит написать один перевод дважды
      и стереть другой.

    Второе — текст абзаца ЦЕЛИКОМ, ровно так, как его склеивал импорт: подряд
    все w:t и без табуляций. Нужен, чтобы узнать сегмент, в который номер
    страницы попал ещё при импорте, — и снять этот номер с перевода, а не
    написать его вторым рядом с настоящим полем.
    """
    slots, full, dropped = [], [], []
    # Стек открытых полей: имя или None, пока инструкция не прочитана.
    # Инструкция всегда идёт ДО результата, поэтому к тексту имя уже известно;
    # непрочитанное поле до тех пор считается вычисляемым — молча писать
    # в неизвестное поле опаснее, чем пропустить его.
    fields: list = []

    def blocked():
        return any(k is None or k in FIELD_COMPUTED for k in fields)

    def walk(el, hidden, sig):
        for ch in el:
            tag = ch.tag
            if tag == qn("w:p"):
                continue
            if tag == qn("w:fldSimple"):
                if _field_key(ch.get(qn("w:instr")) or "") in FIELD_COMPUTED:
                    txt = "".join(t.text or "" for t in ch.iter(qn("w:t")))
                    full.append(txt)
                    dropped.append(txt)
                    continue
                walk(ch, hidden, sig)
                continue
            if tag == qn("w:fldChar"):
                kind = ch.get(qn("w:fldCharType"))
                if kind == "begin":
                    fields.append(None)
                elif kind == "end" and fields:
                    fields.pop()
                continue
            if tag == qn("w:instrText"):
                # Инструкция приезжает кусками; имя даёт первый непустой.
                if fields and fields[-1] is None:
                    key = _field_key(ch.text or "")
                    if key:
                        fields[-1] = key
                continue
            if tag == qn("w:t"):
                txt = ch.text or ""
                full.append(txt)
                if hidden or blocked():
                    dropped.append(txt)
                else:
                    slots.append((ch, sig))
                continue
            if tag == qn("w:r"):
                rpr = ch.find(qn("w:rPr"))
                walk(ch,
                     hidden or (rpr is not None
                                and (rpr.find(qn("w:webHidden")) is not None
                                     or rpr.find(qn("w:vanish")) is not None)),
                     _vis_sig(rpr))
                continue
            walk(ch, hidden, sig)

    walk(p_elem, False, ())
    return slots, "".join(full), "".join(dropped)


# ── Перенос выделений внутри абзаца ─────────────────────────────────
# Сегмент — это строка текста; границ жирного куска в ней нет. Значит перенести
# выделение можно только сопоставив перевод с исходником. Два способа, и первый
# точный: кусок, который перевод обязан сохранить ДОСЛОВНО (латинское название
# вида, аббревиатура, число), ищется в переводе как есть. Остальные границы
# ставятся по доле длины и подтягиваются к ближайшему пробелу — приблизительно,
# зато абзац сохраняет свой вид, а не уходит жирным целиком.
_CYRILLIC = _re.compile(r'[Ѐ-ӿ]')


def _verbatim_token(text: str) -> str:
    """Кусок, по которому выделение можно поставить точно: без кириллицы,
    но со значащим символом. «Mycobacterium tuberculosis», «MDR/RR-TB», «38»."""
    t = text.strip()
    if not t or not any(c.isalnum() for c in t) or _CYRILLIC.search(t):
        return ""
    return t


# Число вместе с разделителями внутри: «1,3», «450–500», «100 000», «1/4».
_NUM_RUN = _re.compile(r'\d[\d\s.,/–—-]*\d|\d')


def _find_number(target: str, want: str, start: int, guess: int) -> int:
    """Где в переводе стоит это же число. Сравниваются ЦИФРЫ, а не запись:
    «1,3» по-русски и «1.3» по-английски — одно число, а разделитель разрядов
    и дробной части перевод меняет. Без этого выделенная статистика («около
    1,3 млн человек») теряла опору на пустом месте."""
    digits = _re.sub(r'\D', '', want)
    if not digits:
        return -1
    best = -1
    for m in _NUM_RUN.finditer(target, start):
        if _re.sub(r'\D', '', m.group(0)) != digits:
            continue
        if best < 0 or abs(m.start() - guess) < abs(best - guess):
            best = m.start()
    return best


# Знаки, по которым граница выделения видна в обоих языках. Тире и кавычки
# в переводе почти всегда другого начертания, поэтому сравниваются классы,
# а не сами символы: «Алиментарный — заражение» и «Alimentary - infection»
# делятся в одном и том же месте.
_DASHES = set("—–—-−‒")


def _delim_class(ch: str) -> Optional[str]:
    if not ch:
        return None
    if ch in _DASHES:
        return "dash"
    if ch in ".!?…":
        return "stop"
    if ch in ":;":
        return ch
    if ch in ")]}»":
        return "close"
    if ch in "([{«":
        return "open"
    return None


def _word_starts(target: str, lo: int, hi: int) -> list:
    return [j for j in range(lo + 1, hi)
            if target[j - 1].isspace() and not target[j].isspace()]


def _snap_word(target: str, guess: int, lo: int, hi: int,
               left: str = "", right: str = "") -> tuple:
    """Ближайшая к guess осмысленная граница строго внутри (lo, hi).

    Сначала знак препинания: в исходнике выделения кончаются на нём в подавляющем
    большинстве случаев — «13 ГЛАВА.», «Клиника: », «ПЦР – полимеразная…».
    Знак есть и в переводе, стоит на своём месте, и граница по нему —
    не догадка, а совпадение. Догадка (доля длины) остаётся запасным ходом.

    Резать посреди слова нельзя в любом случае: половина слова жирной — это
    не «примерно так же», это брак.

    Возвращает (граница, точно ли). «Точно» решается по ПЕРЕВОДУ, а не по
    исходнику: знак в исходнике есть почти всегда, а найтись в переводе рядом
    с расчётным местом он может и не найтись. Считать такое совпадением значит
    отчитаться цифрой, которой не было."""
    if hi <= lo:
        return lo, False
    guess = max(lo, min(hi, guess))
    window = max(15, (hi - lo) // 3)

    # 1) граница у знака препинания
    lc = _delim_class((left.rstrip() or " ")[-1])
    rstr = right.lstrip()
    rc = _delim_class(rstr[0]) if rstr else None
    marks = []
    for i, ch in enumerate(target):
        cls = _delim_class(ch)
        if lc and cls == lc:
            j = i + 1
            while j < len(target) and target[j].isspace():
                j += 1          # пробелы за знаком уходят налево, как в исходнике
            marks.append(j)
            marks.append(i + 1)
        if rc and cls == rc:
            j = i
            while j > 0 and target[j - 1].isspace():
                j -= 1
            marks.append(j)
            marks.append(i)
    marks = [j for j in marks if lo < j < hi and abs(j - guess) <= window]
    if marks:
        return min(marks, key=lambda j: (abs(j - guess), j)), True

    # 2) запасной ход — ближайшая граница слова
    cands = _word_starts(target, lo, hi)
    if not cands:
        return guess, False
    return min(cands, key=lambda j: abs(j - guess)), False


def _split_target(target: str, sources: list) -> tuple:
    """Режет перевод на куски по группам оформления исходного абзаца.

    Возвращает (куски, было ли хоть одно приблизительное деление). Сумма кусков
    ВСЕГДА равна переводу целиком: границы — это индексы одной строки, идущие
    по возрастанию от 0 до конца. Ни потерять, ни задвоить текст нельзя
    по построению."""
    n = len(sources)
    out = [""] * n
    idx = [i for i, s in enumerate(sources) if s.strip()]
    if not idx:
        return out, False
    if len(idx) == 1 or len(target) < 2 * len(idx):
        # Делить нечего или перевод слишком короток, чтобы делить осмысленно.
        out[idx[0]] = target
        return out, False

    total = sum(len(sources[i]) for i in idx) or 1
    bounds = [None] * (len(idx) + 1)

    # 1) точные якоря
    cur, run = 0, 0
    for k, i in enumerate(idx):
        piece = sources[i]
        guess = int(run / total * len(target))
        run += len(piece)
        window = max(12, len(target) // 4)
        tok = _verbatim_token(piece)
        if tok:
            pos = target.find(tok, cur)
            if pos >= 0 and (len(tok) >= 3 or abs(pos - guess) <= window):
                # Короткий кусок вроде «2» встречается в переводе много раз:
                # нашёлся далеко от своего места — это другое вхождение.
                bounds[k] = pos
                bounds[k + 1] = pos + len(tok)
                cur = pos + len(tok)
                continue
        # Кусок с кириллицей целиком не найти, но выделение сплошь и рядом
        # НАЧИНАЕТСЯ с числа: «1,3 млн человек», «43,6 на 100 000 населения».
        # Число перевод сохраняет, поэтому начало ставится точно, а конец
        # достаётся общему правилу.
        head = _NUM_RUN.match(piece.strip())
        if head:
            pos = _find_number(target, head.group(0), cur, guess)
            if pos >= 0 and abs(pos - guess) <= window:
                bounds[k] = pos
                cur = pos

    # Края принадлежат крайним группам целиком: якорь в середине не имеет права
    # отрезать начало или хвост перевода.
    bounds[0], bounds[len(idx)] = 0, len(target)

    # 2) промежутки — по доле длины исходных кусков
    approx = False
    known = [j for j in range(len(bounds)) if bounds[j] is not None]
    for a, b in zip(known, known[1:]):
        if b - a <= 1:
            continue
        lo, hi = bounds[a], bounds[b]
        span = sum(len(sources[idx[j]]) for j in range(a, b)) or 1
        acc = 0
        for j in range(a, b - 1):
            acc += len(sources[idx[j]])
            guess = lo + int(round((hi - lo) * acc / span))
            left, right = sources[idx[j]], sources[idx[j + 1]]
            pos, exact = _snap_word(target, guess, bounds[j], hi, left, right)
            bounds[j + 1] = pos
            # Граница у знака препинания — не догадка: тот же знак стоит
            # и в переводе. Приблизительными считаем только те, где опереться
            # оказалось не на что.
            approx = approx or not exact

    for k, i in enumerate(idx):
        out[i] = target[bounds[k]:bounds[k + 1]]
    return out, approx


def _write_para(slots: list, target: str) -> tuple:
    """Пишет перевод в абзац, сохраняя выделения. (разделён ли, приблизительно ли).

    Прогоны с одинаковым видимым оформлением сливаются в группу — Word режет
    текст на прогоны сам по себе (орфография, метки правок), и такое деление
    к оформлению отношения не имеет.

    Пробелы по КРАЯМ абзаца возвращаются на место. В этом учебнике подписи
    к рисункам сдвинуты вправо не отступом абзаца, а тридцатью тремя пробелами
    в самом тексте; импорт их обрезал (`_docx_clean`), перевод их не содержит,
    и подпись уезжала к левому краю — а дальше Word заново обтекал ею плавающую
    картинку, и строка рассыпалась на куски по обе стороны от рисунка. Ровно то,
    что видно глазом как «текст съезжает». Внутренние пробелы не трогаем: они
    часть предложения, и у перевода свои."""
    keep = "".join(n.text or "" for n, _s in slots)
    if keep.strip():
        lead = keep[:len(keep) - len(keep.lstrip())]
        trail = keep[len(keep.rstrip()):]
        target = lead + target + trail
    groups: list = []
    for node, sig in slots:
        if groups and groups[-1][1] == sig:
            groups[-1][0] += node.text or ""
            groups[-1][2].append(node)
        else:
            groups.append([node.text or "", sig, [node]])

    nonblank = {g[1] for g in groups if g[0].strip()}
    if len(nonblank) <= 1:
        # Оформление одно на весь абзац: пишем в самую длинную группу, чтобы
        # не отдать текст группе из одних пробелов.
        main = max(range(len(groups)), key=lambda i: len(groups[i][0]))
        pieces = ["" if i != main else target for i in range(len(groups))]
        approx = False
    else:
        pieces, approx = _split_target(target, [g[0] for g in groups])

    for piece, g in zip(pieces, groups):
        nodes = g[2]
        nodes[0].text = piece
        nodes[0].set(_XML_SPACE, "preserve")
        for extra in nodes[1:]:
            extra.text = ""
    return len(nonblank) > 1, approx


def _image_caption(p_elem, text: str, qn):
    """Дописать перевод отдельным абзацем сразу за картинкой.

    Так уходит текст, который перерисовать нельзя: надпись на рентгенограмме
    или на фотографии стирается заплаткой, а заплатка посреди снимка — порча
    документа. Оформление берём у абзаца картинки (выключка, отступы), чтобы
    подпись встала под ней, а не у левого поля.

    Никакой приписки вроде «текст на рисунке» здесь нет намеренно: она была бы
    выдуманным текстом, которого в документе не было, да ещё и на неизвестно
    каком языке — целевых языков у системы много."""
    from docx.oxml import OxmlElement
    from copy import deepcopy
    new_p = OxmlElement("w:p")
    ppr = p_elem.find(qn("w:pPr"))
    if ppr is not None:
        new_p.append(deepcopy(ppr))
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    new_p.append(run)
    p_elem.addnext(new_p)
    return new_p


def _export_images(doc, data: dict, by_id: dict, qn, stats: dict) -> None:
    """Перевод, вписанный обратно в картинки.

    Байты картинки подменяются в КОПИИ документа: исходник на диске не
    трогается никогда, иначе следующий экспорт брал бы уже переведённую
    картинку и писал перевод поверх перевода.

    Решение по каждому блоку принимает `image_text.render_target` — одно
    место на всю систему: разойдись оно с тем, что показывает карточка
    «Текст на картинках», человек читал бы одно число, а получал другое."""
    images = (data.get("images") or []) if image_text is not None else []
    if not images:
        return
    anchors = _docx_image_anchors(doc)
    all_p = _docx_flat_paragraphs(doc)
    # Тело идёт первым (инвариант `_docx_flat_parts`), поэтому номер меньше
    # этого — абзац тела. Подпись в колонтитул вставлять нельзя: она
    # напечаталась бы на КАЖДОЙ странице и раздвинула вёрстку.
    body_n = len(doc.element.body.findall(".//" + qn("w:p")))
    parts = {str(p.partname).lstrip("/"): p for p in doc.part.package.iter_parts()}
    for im in images:
        name = im.get("part")
        items = []
        for i, b in enumerate(im.get("blocks") or []):
            if b.get("skip"):
                continue
            # Сегмент обязан подтвердить якорем, что он и есть этот блок:
            # номера переиспользуются, и по одному номеру перевод встал бы
            # в чужую рамку.
            seg = _image_seg_of(by_id, b, name, i)
            if seg is None:
                # Прочитано, но своего сегмента нет: разбор остановили
                # на полпути либо сегменты снесли. Надпись останется на языке
                # оригинала, и молчать об этом нельзя.
                if (b.get("text") or "").strip():
                    stats["img_noseg"] += 1
                continue
            target = (seg.get("target") or "").strip()
            if not target:
                # Не переведён — надпись остаётся на языке оригинала. Стереть
                # её «пока что» значило бы выбросить текст совсем.
                stats["img_untranslated"] += 1
                continue
            items.append({"box": b["box"], "text": target, "align": b.get("align"),
                          # Высота строки и число строк: по ним считаются поля
                          # рамки и решается, нужно ли увеличивать картинку.
                          "lineH": b.get("lineH"), "rows": b.get("rows")})
        if not items:
            continue
        part = parts.get(name)
        if part is None:
            stats["img_lost"] += len(items)
            continue
        # Отпечаток. Исходник могли приложить заново — карта картинок это
        # переживает намеренно, но рамки описывают ТУ картинку. Совпадения нет
        # — не трогаем ничего и говорим числом: перевод, вписанный по чужим
        # координатам, это заплатка посреди снимка.
        if im.get("sha") and hashlib.sha1(part.blob).hexdigest() != im["sha"]:
            stats["img_stale"] += len(items)
            continue
        stats["img_parts"] += 1
        new_bytes, report = image_text.render_target(part.blob, items)
        captions = []
        for r in report:
            if r.get("ok"):
                stats["img_repainted"] += 1
                continue
            why = r.get("why") or ""
            if why == "flat":
                stats["img_flat"] += 1
            elif why == "tiny_font":
                stats["img_font"] += 1
            else:
                stats["img_failed"] += 1
            captions.append(items[r["i"]]["text"])
        if new_bytes:
            part._blob = new_bytes
        idxs = [i for i in (anchors.get(name) or []) if i < min(len(all_p), body_n)]
        if not captions:
            continue
        if not idxs:
            # Картинка в теле документа не стоит (только в колонтитуле)
            # либо не стоит нигде: дописать перевод некуда. Считаем
            # потерянным, а не «сделанным».
            stats["img_lost"] += len(captions)
            continue
        # Каждая следующая подпись цепляется за предыдущую, а не за картинку:
        # addnext вставляет СРАЗУ после, и три подписи, посаженные на один
        # абзац, встали бы в обратном порядке.
        after = all_p[min(idxs)]
        for text in captions:
            after = _image_caption(after, text, qn)
            stats["img_captioned"] += 1


# Порядок элементов в settings.xml задан схемой, и Word разборчив: чужое
# место — «содержимое нечитаемо». `w:updateFields` стоит перед этой группой,
# поэтому вставляем ПЕРЕД первым найденным из неё, а не в конец.
_SETTINGS_AFTER = ("hdrShapeDefaults", "footnotePr", "endnotePr", "compat",
                   "docVars", "rsids")


def _refresh_fields(doc, qn) -> int:
    """Просит пересчитать поля при открытии файла. Возвращает число помеченных.

    Иначе в оглавлении остаются номера страниц ИСХОДНИКА: английский текст
    короче русского, книга похудела с 324 страниц до 302, а «311» в оглавлении
    осталось. Сами эти числа посчитать нельзя — их даёт вёрстка, которой у нас
    нет; всё, что можно, — попросить Word и LibreOffice пересчитать.
    Просим ДВУМЯ способами, потому что клиенты слушают разное: `w:dirty`
    на самом поле и `w:updateFields` в настройках документа.

    Пересчёт заодно перебирает оглавление по ЗАГОЛОВКАМ, а они переведены, —
    то есть чинит и текст строк, и номера разом."""
    n = 0
    roots = [doc.element]
    for part in sorted(doc.part.package.iter_parts(), key=lambda p: str(p.partname)):
        ct = part.content_type or ""
        el = getattr(part, "element", None)
        if el is not None and ("header+xml" in ct or "footer+xml" in ct):
            roots.append(el)
    for root in roots:
        stack: list = []
        for el in root.iter():
            if el.tag == qn("w:fldChar"):
                kind = el.get(qn("w:fldCharType"))
                if kind == "begin":
                    stack.append(el)
                elif kind == "end" and stack:
                    stack.pop()
            elif el.tag == qn("w:instrText") and stack:
                if (_field_key(el.text or "") in FIELD_REFRESH
                        and stack[-1].get(qn("w:dirty")) != "true"):
                    stack[-1].set(qn("w:dirty"), "true")
                    n += 1
            elif el.tag == qn("w:fldSimple"):
                if _field_key(el.get(qn("w:instr")) or "") in FIELD_REFRESH:
                    el.set(qn("w:dirty"), "true")
                    n += 1
    try:
        settings = doc.settings.element
        if settings.find(qn("w:updateFields")) is None:
            upd = settings.makeelement(qn("w:updateFields"), {qn("w:val"): "true"})
            anchor = next((settings.find(qn("w:" + name)) for name in _SETTINGS_AFTER
                           if settings.find(qn("w:" + name)) is not None), None)
            if anchor is not None:
                anchor.addprevious(upd)
            else:
                settings.append(upd)
    except Exception as e:
        # Не нашли настроек — поля всё равно помечены `w:dirty`. Молча
        # ронять экспорт из-за необязательной части пакета нельзя.
        print("[backend] settings.xml: updateFields не выставлен: %s" % e, file=sys.stderr)
    return n


def _export_docx_layout(project: dict, out: Path) -> dict:
    """Подставляет переводы в сохранённый исходник и сохраняет копию.

    Всё, чего мы не тронули, остаётся байт в байт: картинки, стили, нумерация,
    разметка страницы, поля. Трогаем только текстовые узлы тех абзацев, что
    названы в карте, — и внутри абзаца сохраняем выделения (см. `_write_para`)."""
    # Отметка в проекте — часть опознания, а не украшение: номера проектов
    # переиспользуются (id = max + 1), и файл удалённого проекта мог бы
    # достаться новому с тем же номером. Проект без отметки исходника
    # не имеет, чей бы файл ни лежал на диске.
    data = _load_source_map(project["id"]) if project.get("sourceDocx") else None
    if data is None:
        raise HTTPException(400, "К проекту не приложен исходный .docx — "
                                 "экспорт 1в1 собрать не из чего")
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        raise HTTPException(500, "python-docx not installed")

    doc = Document(str(data["path"]))
    all_p = _docx_flat_paragraphs(doc)
    body_n = len(doc.element.body.findall(".//" + qn("w:p")))
    if data.get("paras") not in (len(all_p), body_n):
        # Файл под картой подменили. Молча продолжать нельзя: номера абзацев
        # уехали, и перевод встал бы по всему документу не на свои места.
        # body_n принимается тоже: карты, записанные до того, как в разбор
        # вошли колонтитулы, считали только тело — а номера абзацев тела
        # от этого не сдвинулись, они идут первыми.
        raise HTTPException(400, "Исходник разошёлся с картой абзацев "
                                 "(%s против %s) — приложите файл заново"
                                 % (len(all_p), data.get("paras")))

    by_id = {s["id"]: s for s in project["segments"]}
    stats = {"paragraphs": len(all_p), "written": 0, "untranslated": 0,
             "noslot": 0, "mismatch": 0, "inline": 0, "approx": 0,
             "trimmed": 0, "lost": 0,
             # Текст, впечатанный в картинки: что перерисовали, что ушло
             # подписью и почему. Молча пропущенная надпись остаётся
             # на языке оригинала, и человек узнает об этом, только открыв
             # готовый файл.
             "img_parts": 0, "img_repainted": 0, "img_captioned": 0,
             "img_untranslated": 0, "img_flat": 0, "img_font": 0,
             "img_failed": 0, "img_lost": 0, "img_stale": 0, "img_noseg": 0,
             # Полей, помеченных к пересчёту: по ним Word пересоберёт
             # оглавление и номера страниц уже под перевод.
             "fields_refreshed": 0}
    shown = 0
    for idx, sid in data.get("pairs") or []:
        seg = by_id.get(sid)
        if seg is None or idx >= len(all_p):
            # Сегмент удалили после привязки — абзац остаётся на языке оригинала.
            stats["lost"] += 1
            continue
        target = (seg.get("target") or "").strip()
        if not target:
            stats["untranslated"] += 1
            continue
        slots, full, dropped = _para_slots(all_p[idx], qn)
        if not slots:
            # Весь текст абзаца лежит в поле или скрыт — вписывать некуда.
            stats["noslot"] += 1
            continue
        source = seg.get("source") or ""
        skey = _match_key(source)
        # Сегмент СТАРОГО импорта равен полному тексту абзаца (с номером
        # страницы из поля), нового — тексту слотов. Хвост-номер снимается
        # только у старого: у нового его в сегменте нет по построению.
        same_full = _match_key(full) == skey
        same = same_full or _match_key(
            "".join((t.text or "") for t, _sig in slots)) == skey
        dropped = dropped.strip()
        if same_full and dropped and target.rstrip().endswith(dropped):
            # Импорт склеивал весь текст абзаца подряд, поэтому номер страницы
            # из оглавления попал и в сегмент, и в его перевод. Снимаем ровно
            # этот хвост и ровно тогда, когда он подтверждён и текстом абзаца,
            # и текстом сегмента: иначе в оглавлении встанет «…trachea85»
            # рядом с настоящим полем, показывающим те же «85».
            target = target.rstrip()[:-len(dropped)].rstrip()
            stats["trimmed"] += 1
        elif not same:
            # Текст сегмента разошёлся с текстом абзаца. Пишем всё равно
            # (перевод абзаца лучше оригинала), но считаем и называем:
            # большое число здесь означает, что карта села не так.
            stats["mismatch"] += 1
            if shown < 5:
                shown += 1
                print("[backend] экспорт 1в1: абзац %d не совпал с сегментом #%d\n"
                      "  в файле:    %r\n  в сегменте: %r"
                      % (idx, sid, full[:120], source[:120]), file=sys.stderr)
        split, approx = _write_para(slots, target)
        stats["inline"] += 1 if split else 0
        stats["approx"] += 1 if approx else 0
        stats["written"] += 1

    # Картинки идут ПОСЛЕ абзацев и только теперь: подписи, которые здесь
    # добавляются, сдвигают абзацы, а `all_p` уже разобран по номерам.
    _export_images(doc, data, by_id, qn, stats)
    # ПОСЛЕ всех правок: пересчитывать надо готовый текст, а не промежуточный.
    stats["fields_refreshed"] = _refresh_fields(doc, qn)
    doc.save(str(out))
    return stats


# Расширение файла и приписка к имени по формату. Формат docx_layout тоже
# отдаёт .docx, но имя обязано отличаться: два экспорта одного проекта иначе
# перезаписывали бы друг друга, и человек скачивал бы не то, что просил.
EXPORT_EXT = {"docx": "docx", "xlsx": "xlsx", "docx_layout": "docx"}
EXPORT_SUFFIX = {"docx_layout": " 1в1"}


def _generate_export(project: dict, fmt: str, include_source: bool = True) -> tuple:
    """Собирает реальный файл экспорта и отчёт о том, что в него попало.
    Раньше экспорт был фиктивным — файл не создавался вовсе, только запись
    в историю."""
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    ext = EXPORT_EXT.get(fmt)
    if not ext:
        raise HTTPException(400, f"Формат {fmt} не поддерживается")
    out = EXPORT_DIR / (_safe_filename(project["title"])
                        + EXPORT_SUFFIX.get(fmt, "") + "." + ext)
    # Пишем во временный файл и подменяем готовый одним os.replace — тем же
    # приёмом, что и save_state. Две причины, и обе не теоретические:
    #   1) экспорт учебника это 21 МБ и несколько секунд. Записанный прямо
    #      в итоговый файл, он в это время доступен на скачивание НЕДОПИСАННЫМ;
    #   2) запись поверх существующего файла требует прав НА ФАЙЛ, а os.replace —
    #      только на каталог. Один файл, случайно оставленный в exports/ от
    #      другого владельца, ронял весь экспорт этого проекта навсегда:
    #      PermissionError на боевом сервере, «Сервер недоступен» на экране.
    tmp = out.with_name(out.name + ".tmp")
    stats: dict = {}
    segs = project["segments"]
    if fmt == "docx_layout":
        stats = _export_docx_layout(project, tmp)
    elif fmt == "docx":
        from docx import Document
        doc = Document()
        doc.add_heading(project["title"], level=1)
        doc.add_paragraph(f"{project.get('src','RU')} → {project.get('tgt','EN')} · "
                          f"сегментов: {len(segs)} · экспорт: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        if include_source:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = "#", "Источник", "Перевод"
            for s in segs:
                row = table.add_row().cells
                row[0].text = str(s["id"])
                row[1].text = s.get("source", "")
                row[2].text = s.get("target", "")
        else:
            for s in segs:
                if s.get("target"):
                    doc.add_paragraph(s["target"])
        doc.save(str(tmp))
    elif fmt == "xlsx":
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Segments"
        ws.append(["#", "Источник", "Перевод", "Статус", "Маршрут", "Риск"])
        for s in segs:
            ws.append([s["id"], s.get("source", ""), s.get("target", ""),
                       s.get("status", ""), s.get("route", ""), s.get("risk", "")])
        wb.save(str(tmp))
    os.replace(str(tmp), str(out))
    return out, stats

@app.post("/api/projects/{pid}/export")
def export_project(pid: int, req: ExportRequest):
    _audit("project.export", project=pid, format=getattr(req, "format", None))
    project = get_project(pid)
    fmt = req.format.lower()
    if fmt not in EXPORT_EXT:
        return {"ok": False,
                "error": f"Формат {fmt.upper()} пока не поддерживается — выберите DOCX или Excel."}
    try:
        path, stats = _generate_export(project, fmt, include_source=req.source)
    except ImportError as e:
        return {"ok": False, "error": f"На сервере нет библиотеки для {fmt.upper()}: {e}"}
    except HTTPException as e:
        # Нет исходника, файл разошёлся с картой — это не сбой сервера, а
        # разговор с человеком: он приложит файл и повторит.
        return {"ok": False, "error": e.detail}
    except Exception as e:
        # Сборка падает не только на нашей логике: не хватило места, файл занят,
        # исходник побился. Пятисотка превращается на экране в «сервер недоступен»
        # — сообщение, по которому нельзя ни понять причину, ни решить, что
        # делать. Разбор с настоящей причиной остаётся в журнале.
        import traceback
        traceback.print_exc(file=sys.stderr)
        return {"ok": False,
                "error": "Не удалось собрать файл (%s: %s). Подробности — в журнале сервера."
                         % (type(e).__name__, e)}
    size_kb = max(1, path.stat().st_size // 1024)
    STATE["exportHistory"].insert(0, {
        "tenant": _current_tenant(),
        "file": path.name,
        "when": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "size": f"{size_kb} КБ",
    })
    STATE["exportHistory"] = STATE["exportHistory"][:50]
    save_state(STATE)
    return {"ok": True, "file": path.name, "size": f"{size_kb} КБ", "stats": stats,
            "url": f"/api/projects/{pid}/export/download?format={fmt}&source={1 if req.source else 0}"}

@app.get("/api/projects/{pid}/export/download")
def download_export(pid: int, format: str = "docx", source: bool = True):
    project = get_project(pid)
    fmt = format.lower()
    if fmt not in EXPORT_EXT:
        raise HTTPException(400, "Поддерживаются только docx, docx_layout и xlsx")
    path, _stats = _generate_export(project, fmt, include_source=source)
    media = ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
             if fmt == "xlsx"
             else "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    return FileResponse(str(path), media_type=media, filename=path.name)


# ─── Glossary ───────────────────────────────────────────────────────
class TermRequest(BaseModel):
    src: str
    tgt: str
    cat: str = ""
    freq: int = 0
    conf: str = "Medium"
    isNew: bool = False
    lang: Optional[str] = None      # языковая пара записи; пусто — дефолтная область
    domain: Optional[str] = None

@app.post("/api/glossary")
def save_term(req: TermRequest):
    scope = _scope(req.lang, req.domain)
    existing = _glossary_entry(req.src, scope)
    if existing is None and not (req.lang or req.domain):
        # Клиент не прислал область (правка записи из общего списка) — правим ту
        # запись, что есть, а не заводим рядом дубль в области по умолчанию.
        existing = next((g for g in STATE["glossary"]
                         if _norm_key(g.get("src")) == _norm_key(req.src)
                         and _tenant_of(g) == _current_tenant()), None)
    # Правка руками = проверенная запись: только такие идут в промпт приказом.
    if existing and not req.isNew:
        # След решения человека обязателен. Без него `_human_touched` правку
        # руками не отличает от записи массового импорта, и аудит понижает её
        # без разрешения — то есть машина молча отменяет то, что человек
        # только что вписал сам.
        existing.update({"tgt": req.tgt, "cat": req.cat, "freq": req.freq, "conf": req.conf,
                         "tier": GLOSSARY_TIER_HARD,
                         "note": "уточнено вручную "
                                 + datetime.now().strftime("%Y-%m-%d"),
                         **_signed_field("edit")})
        _clear_auto_marks(existing)
    else:
        STATE["glossary"].insert(0, {**req.dict(exclude={"isNew"}), "tier": GLOSSARY_TIER_HARD,
                                     "lang": scope[0], "domain": scope[1], "tenant": scope[2],
                                     **_signed_field("add")})
    _audit("glossary.save", src=req.src, tgt=req.tgt)
    _invalidate_gloss_index()
    save_state(STATE)
    return {"ok": True}


@app.delete("/api/projects/{pid}")
def delete_project(pid: int):
    _audit("project.delete", project=pid)
    get_project(pid)                      # чужой проект — 404, удалять нечего
    STATE["projects"] = [p for p in STATE["projects"] if p["id"] != pid]
    # Исходник уходит вместе с проектом. Учебник весит 21 МБ, и оставлять его
    # на диске после удаления значит копить мусор, который никто уже не найдёт:
    # имя файла — номер проекта, а проекта больше нет.
    for path in _source_paths(pid):
        try:
            path.unlink(missing_ok=True)
        except OSError as e:
            print("[backend] исходник проекта %s не удалён: %s" % (pid, e),
                  file=sys.stderr)
    save_state(STATE)
    return {"ok": True}


def _revert_repairs_bulk(entries: list, scope: tuple) -> dict:
    """Вернуть правки, сделанные по КАЖДОЙ из записей, одним проходом.

    Поштучный `revert_repairs_by_term` на трёхстах записях означал бы триста
    проходов по всем сегментам проекта. Здесь проход один, а записи разложены
    по метке претензии — той самой строке, которую формулирует `_gloss_misses`.

    Правила те же, что у поштучного отката: единственная причина — возвращаем
    прежний текст; причин несколько — не трогаем (откат унёс бы верные правки),
    а снимаем `source_hash`, и ремонт вернётся к сегменту сам."""
    marks = {}
    for e in entries:
        src, tgt = (e.get("src") or "").strip(), (e.get("tgt") or "").strip()
        if src and tgt:
            marks["«" + src + "» — «" + tgt + "»"] = e
    if not marks:
        return {"reverted": 0, "requeued": 0, "skipped": 0}
    rev = req = skip = 0
    for p in STATE["projects"]:
        if _project_scope(p) != scope:
            continue
        if _active_job_for(p["id"]):
            # Проект сейчас пишет воркер прогона — пропускаем поимённо,
            # иначе два процесса затёрли бы документ друг друга.
            print(f"[backend] revert-repairs: проект {p['id']} пропущен — идёт прогон",
                  file=sys.stderr)
            skip += 1
            continue
        for seg in p["segments"]:
            rp = seg.get("repair") or {}
            if not rp.get("applied"):
                continue
            issues = list(rp.get("issues") or ())
            if not any(m in (t or "") for t in issues for m in marks):
                continue
            # В сегменте должен стоять ИМЕННО тот текст, что написал ремонт.
            # Иначе перевод правили после него — руками или другим прогоном —
            # и подстановка `repair.from` выбросила бы чужую работу молча.
            # Проверка идёт ПОСЛЕ сверки претензии: до неё она отсеивала бы
            # и сегменты, к этой записи отношения не имеющие.
            if not _repair_tried(seg):
                skip += 1
                continue
            old_text = rp.get("from")
            if not (old_text or "").strip():
                skip += 1
                continue
            if len(issues) > 1:
                rp.pop("source_hash", None)
                req += 1
                continue
            _replace_target(seg, old_text, rp.get("model") or seg.get("provider") or "",
                            "REPAIR_REVERT")
            seg["status"] = "review"
            seg["prevTarget"] = rp.get("candidate") or ""
            seg["repair"] = {"applied": False, "reason": "откачено: правило понижено",
                             "from": old_text, "model": rp.get("model", ""),
                             "at": datetime.now().strftime("%Y-%m-%d %H:%M")}
            rev += 1
    return {"reverted": rev, "requeued": req, "skipped": skip}


def _repaired_by_term(entry: dict, scope: tuple) -> list:
    """Сегменты, которые ремонт переписал ИЗ-ЗА этой записи глоссария.

    Ищем по тексту претензии, с которой ремонт работал (`repair.issues`):
    её формулирует `_gloss_misses` и в неё входят обе стороны пары. Хранить
    отдельную ссылку на запись было бы честнее, но записи правят и удаляют,
    а претензия остаётся фактом о том, что и почему переписали."""
    src, tgt = (entry.get("src") or "").strip(), (entry.get("tgt") or "").strip()
    if not src or not tgt:
        return []
    mark = "«" + src + "» — «" + tgt + "»"
    out = []
    for p in STATE["projects"]:
        if _project_scope(p) != scope:
            continue
        for seg in p["segments"]:
            rp = seg.get("repair") or {}
            if not rp.get("applied"):
                continue
            if any(mark in (t or "") for t in (rp.get("issues") or ())):
                out.append({"project": p["id"], "id": seg["id"]})
    return out


class TermScopeRequest(BaseModel):
    src: str
    lang: str = ""
    domain: str = ""


@app.post("/api/glossary/demote")
def demote_term(req: TermScopeRequest):
    """Понизить приказ до подсказки по решению ЧЕЛОВЕКА.

    Отдельно от `save_term` намеренно. Там действует правило «правка руками =
    приказ»: меняя перевод, человек за него ручается. Здесь намерение обратное
    и его нельзя выразить через ту же дверь — запись остаётся в глоссарии
    и остаётся видна, но перестаёт принуждать модель и быть основанием
    для ремонта.

    Зачем это нужно: сверка смысла находит неверные приказы, но записи со следом
    решения человека не понижает — чужое решение машина не отменяет. Значит
    человеку нужен способ согласиться с ней в один клик, иначе находка
    показывается и ничем не заканчивается.

    Удаления тут нет: перевод может быть верным В КОНТЕКСТЕ и негодным как
    правило на весь документ. Подсказка — ровно это и означает."""
    _audit("glossary.demote", src=req.src)
    scope = _scope(req.lang, req.domain)
    entry = _glossary_entry(req.src, scope)
    if entry is None and not (req.lang or req.domain):
        # Только СВОЯ организация: по одному имени термина эта дверь понижала
        # бы приказ в ЧУЖОМ глоссарии.
        entry = next((g for g in STATE["glossary"]
                      if _norm_key(g.get("src")) == _norm_key(req.src)
                      and _tenant_of(g) == _current_tenant()), None)
    if entry is None:
        raise HTTPException(404, "Запись не найдена в этой области")
    if _hit_tier(entry) != GLOSSARY_TIER_HARD:
        return {"ok": True, "already": True, "tier": _hit_tier(entry)}
    today = datetime.now().strftime("%Y-%m-%d")
    entry.update({"prevTier": GLOSSARY_TIER_HARD, "prevNote": entry.get("note", ""),
                  "prevConf": entry.get("conf", ""), "tier": GLOSSARY_TIER_SOFT,
                  "conf": "medium", "note": "понижено вручную " + today,
                  "updated": today, **_signed_field("demote")})
    # Пометка «человек решил оставить приказ» с записи снимается: он решил
    # обратное, и держать её значит соврать следующему аудиту.
    entry.pop("meaningKept", None)
    _invalidate_gloss_index()
    save_state(STATE)
    # Понижение снимает ПОВОД чинить, но не отменяет уже сделанного: сегменты,
    # куда ремонт этот термин вписал, остались переписанными. Молчать об этом
    # нельзя — человек уверен, что отменил правило целиком.
    touched = _repaired_by_term(entry, scope)
    return {"ok": True, "already": False, "tier": GLOSSARY_TIER_SOFT,
            "repaired": touched[:200], "repairedCount": len(touched)}


class RevertRepairsRequest(BaseModel):
    src: str
    lang: str = ""
    domain: str = ""
    project: Optional[int] = None


@app.post("/api/glossary/revert-repairs")
def revert_repairs_by_term(req: RevertRepairsRequest):
    """Вернуть сегменты, которые ремонт переписал ИЗ-ЗА этой записи глоссария.

    Понижение снимает повод чинить дальше, но уже переписанное так и остаётся.
    Здесь оно возвращается — и это единственная операция во всей системе,
    которая меняет текст БЕЗ вызова модели. Так можно ровно потому, что она
    ничего не сочиняет: подставляется `repair.from` — тот самый текст, который
    в сегменте стоял до правки и у которого были свои проверки.

    Два случая, и смешивать их нельзя:
      * претензия по этой записи была ЕДИНСТВЕННОЙ причиной правки — возвращаем
        прежний текст целиком;
      * причин было несколько — не трогаем: откат унёс бы и верные исправления
        (потерянные числа, кальки). Вместо этого снимаем `source_hash`, и
        сегмент снова становится доступен ремонту: он перечинит его по
        оставшимся находкам, уже без понижённого требования.

    Проверки после отката сами становятся устаревшими: они писали хеш
    ОТВЕРГНУТОГО текста, а в сегменте теперь другой (см. `_check_stale`).
    Статус — `review`: текст менял не человек."""
    scope = _scope(req.lang, req.domain)
    entry = _glossary_entry(req.src, scope) or {"src": req.src, "tgt": ""}
    # Записи может уже не быть (её удалили) — тогда пару берём из запроса.
    if not (entry.get("tgt") or "").strip():
        raise HTTPException(400, "Нужен перевод записи, по которой чинили: "
                                 "у этой записи его нет")
    mark = "«" + (entry.get("src") or "").strip() + "» — «" + entry["tgt"].strip() + "»"
    reverted, requeued, skipped = [], [], []
    for p in STATE["projects"]:
        if _project_scope(p) != scope:
            continue
        if req.project and p["id"] != req.project:
            continue
        for seg in p["segments"]:
            rp = seg.get("repair") or {}
            if not rp.get("applied"):
                continue
            issues = list(rp.get("issues") or ())
            if not any(mark in (t or "") for t in issues):
                continue
            if not _repair_tried(seg):
                # Текст правили после ремонта — возвращать нечего и нельзя.
                skipped.append({"project": p["id"], "id": seg["id"]})
                continue
            old = rp.get("from")
            if not (old or "").strip():
                # Ремонт не сохранил прежний текст — вернуть нечем, врать нельзя.
                skipped.append({"project": p["id"], "id": seg["id"]})
                continue
            if len(issues) > 1:
                # Правка чинила не только это. Откат унёс бы верные исправления —
                # отдаём сегмент ремонту заново, без понижённого требования.
                rp.pop("source_hash", None)
                requeued.append({"project": p["id"], "id": seg["id"]})
                continue
            _replace_target(seg, old, rp.get("model") or seg.get("provider") or "",
                            "REPAIR_REVERT")
            seg["status"] = "review"          # текст менял не человек
            seg["prevTarget"] = rp.get("candidate") or rp.get("from") or ""
            seg["repair"] = {"applied": False, "reason": "откачено: правило понижено",
                             "from": old, "model": rp.get("model", ""),
                             "at": datetime.now().strftime("%Y-%m-%d %H:%M")}
            reverted.append({"project": p["id"], "id": seg["id"]})
    if reverted or requeued:
        save_state(STATE)
        _IMPACT_CACHE.clear()
        _ANALYSIS_CACHE.clear()
    return {"ok": True, "reverted": reverted, "revertedCount": len(reverted),
            "requeued": requeued, "requeuedCount": len(requeued),
            "skipped": skipped, "skippedCount": len(skipped),
            "requeuedWhy": ("правка чинила не только это — прежний текст унёс бы "
                            "и верные исправления; сегменты отданы ремонту заново"
                            if requeued else ""),
            "skippedWhy": ("ремонт не сохранил прежний текст — вернуть нечем"
                           if skipped else "")}


@app.delete("/api/glossary")
def delete_term(src: str, lang: str = "", domain: str = ""):
    """Удаление — операция без отмены, поэтому область здесь трактуется строго.
    Пустые lang/domain означают область по умолчанию (так же читаются записи
    без полей), а НЕ «любую»: иначе удаление RU→EN термина уносило бы и его
    RU→DE тёзку из чужого проекта."""
    want = _scope(lang, domain)
    victims = [t for t in STATE["glossary"] if t.get("src") == src and _scope_of(t) == want]
    if not victims and not (lang or domain):
        # Область не назвали и в области по умолчанию записи нет. Удаляем по
        # одному имени, только если претендент ровно один — иначе непонятно,
        # какой именно, а угадывать в необратимой операции нельзя.
        # Только своя организация: удаление по одному имени термина иначе
        # уносило бы запись из чужого глоссария.
        same = [t for t in STATE["glossary"] if t.get("src") == src
                and _tenant_of(t) == _current_tenant()]
        if len(same) == 1:
            victims = same
    if not victims:
        raise HTTPException(404, "Запись не найдена в этой области")
    doomed = {id(t) for t in victims}
    STATE["glossary"] = [t for t in STATE["glossary"] if id(t) not in doomed]
    _invalidate_gloss_index()
    save_state(STATE)
    return {"ok": True}


PURGE_DIR = DATA_DIR / "backups"


def _used_term_ids() -> set:
    """id записей глоссария, применимых хоть к одному сегменту хоть одного
    проекта своей области.

    Обход СЕГМЕНТАМИ, а не записями, и это не оптимизация ради красоты.
    Поштучный `_term_used` на 9502 подсказках и 2670 сегментах — двадцать пять
    миллионов проверок регулярками, минуты работы ЕДИНСТВЕННОГО воркера: запрос
    отваливается по таймауту, а сервис на это время недоступен всем. Через
    индекс глоссария тот же ответ считается за секунды.
    Кандидаты берутся индексом, но НЕ через `_get_context`: он оставляет
    пятнадцать лучших записей на сегмент, и всё, что не влезло, оказалось бы
    «неиспользуемым» — то есть вынесенным вместе с работающими записями."""
    used = set()
    idx = _gloss_index()
    for p in STATE.get("projects", []):
        scope = _project_scope(p)
        for seg in p.get("segments", ()):
            text = seg.get("source", "")
            if not text:
                continue
            seen = set()
            for k in _text_keys(text):
                for g in idx.get(k, ()):
                    if id(g) in seen or id(g) in used:
                        continue
                    seen.add(id(g))
                    if (_scope_of(g) == scope and g.get("src")
                            and _term_match(g["src"], text, _src_lang(g))):
                        used.add(id(g))
    return used


# Разделитель ВАРИАНТОВ перевода в записи массового импорта. Ровно точка
# с запятой и ничего больше: на боевом глоссарии она стоит в 157 записях
# и НИ В ОДНОЙ приказной, то есть двусмысленна только у подсказок.
# Косая черта сюда НЕ добавлена, и это размен, а не факт: из семи записей
# со слэшем в трёх он часть термина («HIV/TB coinfection», «Eto/Pto»,
# «helper/suppressor T-cell balance»), в четырёх — разделитель вариантов.
# Ложный вынос верной записи дороже пропущенной кривой, поэтому берём только
# однозначный разделитель. Запятая не годится по той же причине, но с обратным
# перевесом: из двенадцати записей в десяти это дозировки и однородные члены
# («Ciprofloxacin (250, 500, 750 mg)»). Фильтр берёт 157 записей из ~163
# многовариантных — 96%.
_VARIANT_SEP = ";"


def _multi_variant(entry: dict) -> bool:
    """Запись хранит НЕСКОЛЬКО вариантов перевода, а не перевод.

    «аппендицит → appendicitis; ecphyaditis», «биоптат → biopsied; material»,
    «бледность → paleness; pallor» — это строка из словарной статьи, а не
    ответ на вопрос «как переводить термин». Модель получает её одной
    подсказкой целиком, и на боевом проекте 267 раз в перевод садился ПЕРВЫЙ
    вариант — а первый там сплошь и рядом хуже второго («aspergillomycosis»
    вместо «aspergillosis», «biopsied» вместо «material»).

    Чинить такую запись выбором варианта нельзя: какой из двух верен, знает
    предметная область, а не разделитель. Запись, которая не может сказать,
    что она означает, вреднее отсутствующей — поэтому её выносят, а не правят.

    Уровень записи предикат НЕ смотрит — отсев по уровню делает `_hit_tier`
    в самом отборе. На боевом глоссарии 26.08 «;» стояла в 157 подсказках
    и ни в одной приказной записи, но это измерение, а не свойство кода:
    выверенный справочник даёт приказы без следа человека, и запись с «;»
    оттуда режимом `tier=verified` вынеслась бы.

    Считаются непустые ЧАСТИ, а не наличие символа: висячая точка с запятой
    в конце («appendicitis;») — обычная грязь импорта, а не второй вариант."""
    parts = [p for p in (entry.get("tgt") or "").split(_VARIANT_SEP) if p.strip()]
    return len(parts) > 1


class GlossaryPurgeRequest(BaseModel):
    # Что выносим. По умолчанию подсказки: массовый импорт лежит именно там,
    # а приказы — это решения человека и выверенных справочников.
    tier: str = GLOSSARY_TIER_SOFT
    project: Optional[int] = None      # только область этого проекта
    unused_only: bool = False          # только те, что не встречаются ни в одном проекте
    multi_variant: bool = False        # только записи с НЕСКОЛЬКИМИ вариантами перевода
    dry_run: bool = True


@app.post("/api/glossary/purge")
def purge_glossary(req: GlossaryPurgeRequest = GlossaryPurgeRequest()):
    """Вынести массовый импорт из глоссария целиком.

    Зачем отдельной командой: `DELETE /api/glossary` работает по ОДНОЙ записи,
    а импорт — десять тысяч. Кликать их поштучно невозможно, поэтому «удалить
    всё» либо есть как осознанная операция с предпросмотром и откатом, либо
    делается руками по файлу состояния — что хуже во всех отношениях.

    Три предохранителя:
      1. `dry_run=True` по умолчанию — считает и показывает, ничего не трогая;
      2. записи со СЛЕДОМ РЕШЕНИЯ ЧЕЛОВЕКА (`_human_touched`) не выносятся
         никогда, даже если подходят под фильтр: человек их трогал, и стирать
         его работу пачкой нельзя. Сколько таких — сказано в ответе;
      3. вынесенное целиком уходит файлом в `data/backups/`, и его возвращает
         `/purge/{stamp}/undo`. Массовое удаление без отката недопустимо —
         тот же закон, что у пачек автоодобрения.

    На уже переведённый текст не влияет НИЧЕМ: расхождения и ремонт считаются
    только по приказам (`_verified_hits`), а подсказки в этот расчёт не входят.
    Меняется лишь то, что уйдёт в промпт при СЛЕДУЮЩЕМ переводе."""
    if not req.dry_run:
        _audit("glossary.purge", mode=getattr(req, "mode", None))
    tier = req.tier if req.tier in (GLOSSARY_TIER_SOFT, GLOSSARY_TIER_HARD) else GLOSSARY_TIER_SOFT
    scope = _project_scope(get_project(req.project)) if req.project else None

    # Считаем ОДИН раз на всю пачку, а не по записи: см. _used_term_ids.
    used_ids = _used_term_ids() if req.unused_only else None
    matched, kept_human = [], 0
    for g in STATE.get("glossary", []):
        if _tenant_of(g) != _current_tenant():
            continue
        if _hit_tier(g) != tier:
            continue
        if scope is not None and _scope_of(g) != scope:
            continue
        if req.multi_variant and not _multi_variant(g):
            continue
        if req.unused_only and id(g) in used_ids:
            continue
        # Пощада считается ПОСЛЕДНЕЙ, уже после всех сужений. Стой она первой,
        # `keptHuman` отвечал бы не на тот вопрос: при `multi_variant` он
        # называл бы все пощажённые подсказки области (690 против 157 взятых),
        # то есть тех, кого и не собирались трогать. Читалось бы это как
        # «у выноса огромный защищённый хвост», а на деле многовариантных
        # со следом человека нет ни одной.
        if _human_touched(g):
            kept_human += 1
            continue
        matched.append(g)

    result = {
        "ok": True, "dryRun": req.dry_run, "tier": tier,
        "scope": list(scope) if scope else None,
        "total": len(STATE.get("glossary", [])),
        "matched": len(matched),
        # Молчаливой пощады не бывает: сказано, скольких не тронули и почему.
        "keptHuman": kept_human,
        "unusedOnly": req.unused_only,
        "multiVariant": bool(req.multi_variant),
        "samples": [{"src": g.get("src"), "tgt": g.get("tgt")} for g in matched[:12]],
        # Всегда, а не только на удавшемся применении: пустой отбор при
        # dry_run=False уходил ранним return без этого ключа.
        "removed": 0,
        "stamp": None,
    }
    if req.dry_run or not matched:
        return result

    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    doomed = {id(g) for g in matched}
    try:
        PURGE_DIR.mkdir(parents=True, exist_ok=True)
        path = PURGE_DIR / ("glossary-purge-" + stamp + ".json")
        path.write_text(json.dumps(matched, ensure_ascii=False), encoding="utf-8")
    except Exception as e:
        # Без бэкапа не удаляем: откат — часть операции, а не украшение.
        print(f"[backend] вынос глоссария: бэкап не записан: {e}", file=sys.stderr)
        raise HTTPException(500, "Не удалось сохранить копию для отката — удаление отменено")
    STATE["glossary"] = [g for g in STATE["glossary"] if id(g) not in doomed]
    _invalidate_gloss_index()
    save_state(STATE)
    print(f"[backend] вынесено записей глоссария: {len(matched)} "
          f"(уровень {tier}), копия: {path.name}", file=sys.stderr)
    result["stamp"] = stamp
    result["removed"] = len(matched)
    return result


@app.get("/api/glossary/purge/list")
def list_glossary_purges():
    """Что можно вернуть. Без списка откат существует только на словах."""
    out = []
    try:
        for f in sorted(PURGE_DIR.glob("glossary-purge-*.json"), reverse=True):
            try:
                rows = json.loads(f.read_text(encoding="utf-8"))
            except Exception:
                continue
            # Копия — одной организации (вынос идёт по своей): чужая в списке
            # не показывается, как и её число.
            if any(_tenant_of(g) != _current_tenant() for g in rows[:1]):
                continue
            n = len(rows)
            out.append({"stamp": f.stem.replace("glossary-purge-", ""),
                        "count": n, "at": f.stem.replace("glossary-purge-", "")})
    except Exception as e:                                   # pragma: no cover
        print(f"[backend] список выносов: {e}", file=sys.stderr)
    return {"ok": True, "purges": out}


@app.post("/api/glossary/purge/{stamp}/undo")
def undo_glossary_purge(stamp: str):
    """Вернуть вынесенное. Возвращаем только те записи, которых сейчас нет:
    иначе откат создал бы дубль поверх записи, добавленной после выноса."""
    if not re.fullmatch(r"[0-9-]{8,20}", stamp or ""):
        raise HTTPException(400, "Неверная метка")
    path = PURGE_DIR / ("glossary-purge-" + stamp + ".json")
    if not path.exists():
        raise HTTPException(404, "Копия не найдена — вернуть нечем")
    try:
        saved = json.loads(path.read_text(encoding="utf-8"))
    except Exception as e:
        raise HTTPException(500, "Копия не читается: " + str(e))
    # Возвращается только СВОЁ: копия чужой организации по метке не читается.
    saved = [g for g in saved if _tenant_of(g) == _current_tenant()]
    if not saved:
        raise HTTPException(404, "Копия не найдена — вернуть нечем")
    have = {(_scope_of(g), _norm_key(g.get("src"))) for g in STATE.get("glossary", [])}
    back, skipped = 0, 0
    for g in saved:
        key = (_scope_of(g), _norm_key(g.get("src")))
        if key in have:
            skipped += 1
            continue
        STATE["glossary"].append(g)
        have.add(key)
        back += 1
    _invalidate_gloss_index()
    save_state(STATE)
    return {"ok": True, "restored": back, "skipped": skipped,
            "skippedWhy": ("запись с таким термином уже есть — её не трогали"
                           if skipped else "")}


# ─── TM ─────────────────────────────────────────────────────────────
@app.delete("/api/tm")
def delete_tm(src: str, lang: str = ""):
    """Как и у глоссария: пустой lang — это пара по умолчанию, а не «любая».
    _tm_upsert теперь держит по записи на языковую пару, и удаление RU→EN
    иначе уносило бы RU→DE запись того же исходника."""
    _audit("tm.delete", src=src)
    want = lang or DEFAULT_GLOSS_LANG
    victims = [t for t in STATE["tm"]
               if t.get("src") == src and (t.get("lang") or DEFAULT_GLOSS_LANG) == want
               and _tenant_of(t) == _current_tenant()]
    if not victims and not lang:
        same = [t for t in STATE["tm"] if t.get("src") == src
                and _tenant_of(t) == _current_tenant()]
        if len(same) == 1:
            victims = same
    if not victims:
        raise HTTPException(404, "Запись памяти переводов не найдена в этой паре языков")
    doomed = {id(t) for t in victims}
    STATE["tm"] = [t for t in STATE["tm"] if id(t) not in doomed]
    save_state(STATE)
    return {"ok": True}


@app.get("/api/health")
def health(request: Request):
    # Эндпоинт публичный — на него опирается смоук-проверка деплоя.
    # Пути на диске и список модулей отдаём только вошедшим.
    info = {
        "ok": True,
        "version": "5.6.0",
        "checksEnabled": checks_enabled(),
        # Прежний ключ — ещё релиз: на него смотрит смоук деплоя.
        "medicalQaEnabled": checks_enabled(),
        "projects": len(STATE["projects"]),
    }
    if _session_valid(_token_from_request(request)):
        info["backendModules"] = list(_BACKEND_MODULES.keys())
        info["stateFile"] = str(STATE_FILE)
    return info



# ─── Фоновые прогоны ─────────────────────────────────────────────────
# Раньше пакет гонял браузер: закрыл вкладку — прогон оборвался на середине,
# а оплаченные вызовы для оставшихся сегментов просто не случились. Теперь
# клиент только ставит задачу, а порции крутит сервер, и страница нужна лишь
# чтобы смотреть прогресс.
#
# Рабочий поток ОДИН и очередь одна: воркер uvicorn тоже один, STATE — общий
# словарь в памяти. Два параллельных прогона дрались бы за один state.json.
import queue as _queue

JOB_CHUNKS = {"translate": 10, "backcheck": 10, "termcheck": 10, "medical_qa": 10,
              # Сверка терминов моделью: один вызов на сегмент, порция как
              # у остальных проверок.
              "termaudit": 10,
              # Терм-лист документа: одна задача на проект, оригиналы порциями.
              "termsheet": 10,
              # Ревизия: один вызов на сегмент, но вызовы внутри порции идут
              # параллельно (`_run_parallel`), поэтому порция как у проверок.
              # Без записи ЗДЕСЬ отдельный запуск шага отвечает 400
              # («неизвестный тип прогона»), а `_job_run` падает KeyError.
              "review": 10,
              "repair": 5, "full": 5, "apply_terms": 5,
              # Разбор картинок идёт СВОИМ циклом (порция — картинка, а не
              # сегмент), но очередь и воркер те же: два тяжёлых прогона
              # на одном воркере uvicorn дрались бы за state.json.
              "images": 1}
JOB_KINDS = set(JOB_CHUNKS)

# Составной прогон: порция проходит все шаги подряд, порция за порцией. Порядок
# НЕ косметический и менять его нельзя:
#   перевод    — иначе проверять нечего;
#   back-check — Medical QA берёт из него готовый обратный перевод и не платит
#                за него второй раз;
#   термины    — вторая независимая проверка; та из двух, что отработала второй,
#                и собирает терминологию с чистых сегментов;
#   ремонт     — после обеих проверок: он чинит по ИХ находкам;
#   Medical QA — последней, и это её место, а не предпоследнее.
#
# Medical QA стояла перед ремонтом, пока считалось, что ремонту нужны её
# находки. Это неправда: _repair_findings читает back-check, termcheck и
# глоссарий, а qa_result не читает никто — числа, единицы и отрицания,
# которыми чинит ремонт, считает сам back-check через medical_qa.run_backcheck.
# Стоя перед ремонтом, проверка описывала текст, который через шаг переписывали,
# и оставалась устаревшей: следующий прогон забирал те же сегменты снова.
# Стоя после, она описывает окончательный текст и в следующий прогон
# не попадает.
# Ревизия стоит ВТОРОЙ, сразу после перевода, и это несущее свойство порядка:
# всё, что ПЕРЕПИСЫВАЕТ текст, обязано идти раньше всего, что его ОПИСЫВАЕТ.
# Тот же урок, из-за которого Medical QA уехала в конец — стоя перед ремонтом,
# она описывала текст, который через шаг заменяли, и её результат устаревал
# сразу. С ревизией цена ошибки выше: поставь её после back-check, и каждая
# правка обесценит оплаченный обратный перевод, а следующий прогон купит его
# заново. Ремонт при этом остаётся последним платным шагом и чинит остаток —
# расхождения с глоссарием, регистр, чужое письмо.
FULL_RUN_STEPS = ["translate", "review", "backcheck", "termcheck", "termaudit",
                  "repair", "medical_qa"]
FULL_STEP_LABELS = {"translate": "перевод", "backcheck": "back-check",
                    "termcheck": "проверка терминов", "medical_qa": "Medical QA",
                    "termaudit": "сверка терминов", "repair": "ремонт",
                    "review": "ревизия"}
# Откуда шаг берёт свою модель. Подшаги читают её из params["model"], а моделей
# в составном прогоне несколько: смысл в том, что переводит одна, а проверяют
# другие — иначе проверка перестаёт быть независимой.
FULL_STEP_MODEL = {"translate": "model", "backcheck": "bc_model",
                   "termcheck": "tc_model", "termaudit": "tcx_model",
                   "repair": "rp_model", "review": "rv_model"}


# ── Разбор прогона: что он сделает и чего делать не станет ───────────────────
# Состав и смету раньше считал браузер своими предикатами, а работу отбирал
# сервер своими — и разойтись они были обязаны. Список сегментов у составного
# прогона ОДИН, объединение целей всех шагов, и каждый шаг брал оттуда всё,
# что проходило ЕГО серверную проверку, а не то, что человек отметил галочками
# в карточке соседнего шага. Снятая галочка уменьшала смету, но не работу:
# показывали одно, списывалось другое.
#
# Теперь состав считает тот же код, который потом и работает, — предикаты
# _backcheck_cached / _termcheck_cached / _check_stale / _repairable. И разбор
# отвечает не только «сколько», но и «почему»: «пропущено 970» без причины —
# это не отчёт, а отговорка, после которой человек идёт кликать галочки
# наугад.
def _model_label(mid: str) -> str:
    m = _MODELS_BY_ID.get(mid or "")
    return m["label"] if m else (mid or "модель неизвестна")


def _plan_step(project: dict, step: str, params: dict, scope: list,
               will_translate: set, gloss_ids: set,
               term_ids: Optional[set] = None,
               consist_ids: Optional[set] = None,
               # Полный состав прогона: шаг обязан знать, кто работает ДО него.
               # Без этого разбор молчит о работе, которую сам же и создаст —
               # ревизия перепишет текст, и проверки этих сегментов протухнут.
               # Сколько сегментов возьмёт РЕВИЗИЯ в этом же прогоне. Не состав
               # шагов, а именно число: она идёт раньше всех, кто описывает
               # текст, и переписанные сегменты обесценят их проверки. Ноль —
               # значит и говорить не о чем (повторный прогон).
               review_takes: int = 0) -> dict:
    """Разбор одного шага: кого возьмёт, кого не возьмёт и почему.

    will_translate — сегменты, которые переведёт этот же прогон. Сейчас у них
    нет перевода и ни в одну проверку они не попадают, но к своему шагу уже
    будут переведены. Без них разбор непереведённого проекта показывал бы цену
    одного перевода, хотя платить придётся и за все проверки поверх него."""
    ids, runs, skips = [], {}, {}

    def run(reason, seg):
        ids.append(seg["id"])
        runs[reason] = runs.get(reason, 0) + 1

    def skip(reason):
        skips[reason] = skips.get(reason, 0) + 1

    mdl_id = None
    if step == "backcheck":
        mdl_id = _resolve_model(params.get("bc_model") or BACKCHECK_DEFAULT_MODEL)["id"]
    elif step == "termcheck":
        mdl_id = _resolve_model(params.get("tc_model") or TERMCHECK_DEFAULT_MODEL)["id"]
    elif step == "repair":
        mdl_id = _resolve_model(params.get("rp_model") or REPAIR_DEFAULT_MODEL)["id"]
    elif step == "translate":
        mdl_id = _resolve_model(params.get("model"))["id"]
    elif step == "termaudit":
        mdl_id = _resolve_model(params.get("tcx_model") or TERM_CONTEXT_DEFAULT_MODEL)["id"]
    elif step == "review":
        mdl_id = _resolve_model(params.get("rv_model") or REVIEW_DEFAULT_MODEL)["id"]
    elif step == "medical_qa":
        # Своей модели у неё нет: правила детерминированные. Но обратный
        # перевод, если готового не осталось, она закажет — моделью back-check.
        # Показываем именно её: «без вызова модели» было полуправдой, из-за
        # которой шаг молча платил моделью перевода по умолчанию.
        mdl_id = _resolve_model(params.get("bc_model") or BACKCHECK_DEFAULT_MODEL)["id"]

    # Забракованные слова — тем же расчётом, что шаг сверки и /analysis:
    # один раз на разбор, а не на сегмент (это проход по очереди кандидатов).
    stale_map = _stale_words_of(project) if step == "termaudit" else {}

    use_judge = bool(params.get("use_judge"))
    # Разрешение этого прогона звать судью и выше зоны. Разбор обязан его
    # читать по той же причине, что include_confirmed: иначе смета обещает
    # одно, а прогон делает другое.
    judge_all = bool(params.get("judge_all"))
    retry = bool(params.get("retry"))
    # Разрешение чинить заверенное человеком. Разбор ОБЯЗАН его читать: раньше
    # он отбрасывал подтверждённые безусловно, а прогон брал их по флагу —
    # смета обещала одно, работа делала другое. Флаг относится только к ремонту
    # (см. _job_chunk_full): он правит по конкретным находкам и точечно,
    # а перевод заново перегнал бы сегмент целиком и за полную цену.
    fix_confirmed = bool(params.get("include_confirmed"))
    # Разрешение ревизии трогать заверенное — своё. Разбор ОБЯЗАН читать
    # именно его, иначе строка обещает одно, а шаг делает другое.
    rv_confirmed = bool(params.get("rv_confirmed"))
    note = None

    for seg in scope:
        target = (seg.get("target") or "").strip()
        pending = seg["id"] in will_translate and not target

        if step == "translate":
            # Предикат общий с самим пакетом — см. _needs_translation.
            # Причины при этом РАЗНЫЕ: «ещё не переведён» и «прошлый перевод
            # не удался» — это не одно и то же для человека, читающего отчёт.
            if not _needs_translation(seg):
                skip("уже переведён")
            elif seg.get("status") == "failed":
                run("прошлый перевод не удался — попробуем снова", seg)
            else:
                run("ещё не переведён", seg)
            continue

        # Переводится в этом же прогоне — к своему шагу текст появится.
        if pending:
            # У сверки терминов есть о чём спрашивать не везде, и знать это
            # можно ЗАРАНЕЕ: приказные записи применимы к ОРИГИНАЛУ, а он
            # от перевода не меняется. Без этой оговорки разбор свежей книги
            # обещал бы сверку всего проекта, а спросил бы четверть — смета
            # вчетверо больше работы, а полоса прогона до галочки «шаг взял
            # всё» не дошла бы никогда.
            if step == "termaudit" and not _verified_hits(seg.get("source", ""), project):
                skip("приказных терминов в сегменте нет")
            else:
                run("появится после перевода", seg)
            continue
        if not target:
            skip("нет перевода")
            continue

        if step == "backcheck":
            if _backcheck_cached(seg, mdl_id, use_judge, judge_all):
                # «Уже проверен» — полуправда, когда судья ещё не смотрел:
                # прогон без тумблера «Судья» такие сегменты не осушает
                # НИКОГДА (тумблер выключен по умолчанию и не переживает
                # перезагрузку), и на боевом проекте 17 сегментов с баллом
                # 77–95 пережили все прогоны из редактора. Причина обязана
                # называть лекарство, которое работает.
                if not use_judge and _judge_pending(seg):
                    skip("судья не смотрел — его позовёт тумблер «Судья» "
                         "или кнопка «Перевести и доделать»")
                elif not judge_all and _judge_pending(seg, above=True):
                    skip("балл выше зоны судьи — смысл прочтёт только "
                         "«Перевести и доделать»")
                else:
                    bm = (seg.get("backcheck") or {}).get("model")
                    skip("уже проверен этим переводом: " + _model_label(bm))
            elif (seg.get("backcheck") or {}).get("score") is None:
                # is None, а не truthy: балл 0 — это проверенный сегмент
                # с провальной оценкой, а не непроверенный.
                run("ещё не проверялся", seg)
            elif _check_stale(seg.get("backcheck"), target):
                run("перевод изменился после проверки", seg)
            elif (seg.get("backcheck") or {}).get("model") == seg.get("provider"):
                run("проверял тот, кто переводил — это не проверка", seg)
            else:
                run("проверен без судьи, а судья включён", seg)

        elif step == "termcheck":
            tc = seg.get("termcheck") or {}
            if _termcheck_cached(seg, mdl_id):
                if tc.get("model") == "skip":
                    skip("нечего проверять — в переводе нет слов")
                elif tc.get("model") == mdl_id:
                    skip("уже проверен этой моделью")
                else:
                    skip("уже проверен моделью не слабее: " + _model_label(tc.get("model")))
            elif not tc:
                run("ещё не проверялся", seg)
            elif _check_stale(tc, target):
                run("перевод изменился после проверки", seg)
            elif model_rank(tc.get("model") or "") is None:
                # Ранга нет — утверждать «этого достаточно» не о чем. Называем
                # модель поимённо: строка в разборе и есть подсказка, что её
                # пора дописать в backend/model_ranks.json.
                run("прошлая проверка моделью неизвестной силы: "
                    + _model_label(tc.get("model")), seg)
            else:
                run("прошлая проверка слабее выбранной: " + _model_label(tc.get("model")), seg)

        elif step == "review":
            rv = seg.get("review") or {}
            # Ветки «нет перевода» и «переведём в этом прогоне» сюда не нужны:
            # общий блок выше перехватывает и то и другое раньше, и вторая
            # копия правила просто никогда не сработала бы.
            # По СВОЕМУ флагу, а не по ремонтному: у того «правь по
            # конкретным находкам», здесь «перечитай и перепиши целиком».
            # Прочитай мы `fix_confirmed` — галочка в строке РЕМОНТА
            # заставила бы строку ревизии обещать все заверенные сегменты
            # проекта, а шаг их не взял бы.
            if seg.get("status") == "confirmed" and not rv_confirmed:
                # Вердикт получить можно, но ПРАВКИ не будет: без разрешения
                # заверенное не переписывают. Платить за совет, который некуда
                # применить, — тот же перерасход, от которого заведён
                # `_repair_futile`. Разбор обязан сказать это вслух.
                skip("заверено человеком — правка не применится")
            elif rv.get("undone"):
                # Человек откатил правку: спрашивать заново значит предлагать
                # ему то же самое второй раз за его же деньги.
                skip("правка откачена человеком")
            elif not _review_stale(seg):
                skip("уже ревизован этим переводом")
            elif not rv:
                run("ещё не ревизовался", seg)
            elif str(rv.get("v")) != str(REVIEW_VERSION):
                # Текст не менялся — изменился набор вопросов. Сказать «перевод
                # изменился» значит соврать в отчёте, который для того и заведён.
                run("вопросы ревизии изменились", seg)
            elif rv.get("styleStale"):
                run("стайл-шит изменился после ревизии", seg)
            else:
                run("перевод изменился после ревизии", seg)

        elif step == "termaudit":
            tcx = seg.get("termContext") or {}
            # Спрашиваем только там, где есть о чём: сегмент без приказных
            # терминов сверять нечем, и платить за него незачем.
            # Готовый список — ускорение, а не источник правды: его отсутствие
            # обязано менять только скорость, но не ОТВЕТ. Прямой вызов
            # (тесты, будущий код) считает сам.
            has_terms = (seg["id"] in term_ids if term_ids is not None
                         else bool(_verified_hits(seg.get("source", ""), project)))
            # Забракованные слова — второй повод для шага: арбитр даёт по ним
            # второй голос (снять претензию либо отдать ремонту). Тот же
            # расчёт, что у самого шага, — иначе смета и работа расходятся.
            stale_pending = _stale_unasked(seg, stale_map.get(seg["id"]) or [])
            if not has_terms and not stale_pending:
                skip("приказных терминов в сегменте нет")
            elif (not _term_context_stale(seg) and tcx.get("all_terms")
                    and not stale_pending):
                skip("уже сверен этим переводом")
            elif not tcx:
                run("ещё не сверялся", seg)
            elif (stale_pending and not _term_context_stale(seg)
                    and tcx.get("all_terms")):
                run("проверка браковала слово — спросит арбитра", seg)
            elif not tcx.get("all_terms"):
                # Вердикт получен разбором СПОРА: там спрашивали только про
                # спорные термины, остальные не сверяли ни разу.
                run("сверялись только спорные термины", seg)
            elif tcx.get("version") != TERM_CONTEXT_VERSION:
                # Текст не менялся — изменился набор вопросов. Сказать «перевод
                # изменился» значит соврать в отчёте, который для того и заведён.
                run("вопросы сверки изменились", seg)
            else:
                run("перевод изменился после сверки", seg)

        elif step == "medical_qa":
            if seg.get("status") not in ("translated", "qa", "review", "confirmed"):
                skip("статус вне работы проверки")
            elif not _check_stale(seg.get("qa_result"), target):
                skip("результат относится к нынешнему тексту")
            else:
                run("нет свежего результата", seg)

        elif step == "repair":
            # Глоссарий берётся из готового отчёта о соответствии, а не считается
            # заново на каждый сегмент: расчёт тот же самый (это записано в
            # _gloss_misses как обязательство), но 13 мс на сегмент превращали
            # разбор проекта на 2670 строк в 34 секунды с заблокированным
            # воркером — при том, что смета пересчитывается на каждую смену
            # модели в списке. project=None тут и означает «без глоссария»:
            # остальные находки читаются из самого сегмента и бесплатны.
            # Находки считаются ОДИН раз и отдаются клейму готовыми:
            # _repair_clamped без списка пересчитывает их внутри отпечатка,
            # а разбор пересчитывается на каждую смену модели в панели —
            # лишние секунды единственного воркера на проекте в 2700 строк.
            rp_f = _repair_findings(seg, None)
            if (seg["id"] not in gloss_ids and seg["id"] not in (consist_ids or ())
                    and not rp_f):
                skip("чинить нечего — находок нет")
            elif (seg.get("status") == "confirmed" and not fix_confirmed
                    and not _confirm_override(seg)):
                skip("заверено человеком — включите «чинить подтверждённые»")
            elif seg.get("status") == "confirmed" and not fix_confirmed:
                # Объективная находка сильнее заверения. Причина названа
                # отдельно: у этих сегментов последствие особое — с них
                # снимется отметка человека, и он должен видеть, за что.
                run("расхождение чисел, единиц или отрицания — сильнее заверения", seg)
            elif not retry and _repair_clamped(seg, rp_f, mdl_id):
                skip("такой же заход уже делали")
            elif not retry and _repair_clamped(seg, rp_f):
                # Заход был, но ДРУГОЙ моделью — а другая модель на тот же
                # промпт отвечает по-своему, и это второе мнение человек
                # заказал сам, выбрав её в панели. Прежнюю называем поимённо:
                # без неё строка читается как повтор уже сделанного.
                run("уже чинила %s — выбранная модель зайдёт со вторым мнением"
                    % _model_label((seg.get("repair") or {}).get("model")), seg)
            elif seg.get("status") == "confirmed":
                # Причина названа отдельно намеренно: цена у этих сегментов та же,
                # а последствие другое — отметка «подтвердил человек» с них снимется.
                run("есть находки, подтверждение будет снято", seg)
            else:
                run("есть находки", seg)

    # Ремонт и Medical QA зависят от того, что найдут предыдущие шаги ЭТОГО же
    # прогона, а этого до запуска не знает никто. Молчать об этом нельзя:
    # смета, посчитанная по нынешним находкам, — нижняя граница, а не цена.
    if step == "repair":
        note = ("Считано по нынешним находкам. Проверки в этом же прогоне могут "
                "добавить ещё — такие сегменты будут починены, но в смету не вошли.")
    elif step == "medical_qa":
        note = ("Считано по нынешнему тексту. Сегменты, которые перепишет ремонт, "
                "проверка возьмёт тоже — они идут следом за ним.")
    # Ревизия идёт ПЕРЕД всеми, кто описывает текст, и переписывает часть
    # сегментов. Их проверки после этого протухают (`_check_stale`) и будут
    # куплены заново — `skip_cached` тут не спасает, потому что текст ДРУГОЙ.
    # Молчать об этом нельзя: смета таких сегментов не видит, а её число
    # уходит в `est_cost` и калибрует поправку по всей системе.
    if (step in ("backcheck", "termcheck", "termaudit", "repair")
            and review_takes):
        # Ремонт тоже в списке: переписанный текст — другие находки и другой
        # отпечаток захода, то есть другой состав.
        add = ("Ревизия идёт раньше и перепишет часть из %d сегментов — "
               "их проверки придётся сделать заново, и в этот состав они "
               "не вошли." % review_takes)
        # У medical_qa своя фраза про то же самое; двух формулировок об одном
        # человеку не нужно.
        note = (note + " " + add) if note else add

    fmt = lambda d: [{"reason": k, "count": v} for k, v in
                     sorted(d.items(), key=lambda kv: -kv[1])]
    return {"step": step, "label": FULL_STEP_LABELS[step], "model": mdl_id,
            "modelLabel": _model_label(mdl_id) if mdl_id else None,
            "ids": ids, "count": len(ids), "note": note,
            "runs": fmt(runs), "skips": fmt(skips)}


def _status_counts(project: dict) -> dict:
    """Сколько сегментов проекта в каждом статусе.

    Нужно ОДНОМУ потребителю — браузеру, чтобы заметить, что его копия проекта
    устарела. Числа сегментов при этом совпадают: прогон ничего не добавлял,
    он менял статусы. Поэтому сверки по длине списка мало, и без этой разбивки
    вкладка честно показывала «Новые 25» там, где на сервере переводить давно
    нечего, — а строка «Перевод» в том же окне говорила «—»."""
    out: dict = {}
    for s in project["segments"]:
        # or "new", а не голое значение: нормализация обязана совпадать
        # с браузерной буква в букву. Разойдись они — сверка находила бы
        # расхождение там, где его нет, и тянула бы проект целиком на каждый
        # разбор состава.
        k = s.get("status") or "new"
        out[k] = out.get(k, 0) + 1
    return out


class RunPlanRequest(BaseModel):
    steps: Optional[list] = None
    segment_ids: Optional[list] = None   # None — весь проект; [] — ничего
    model: Optional[str] = None
    bc_model: Optional[str] = None
    tc_model: Optional[str] = None
    tcx_model: Optional[str] = None
    rp_model: Optional[str] = None
    # Поле обязано быть у КАЖДОГО ключа FULL_STEP_MODEL: лишнее pydantic
    # выбрасывает молча, и разбор посчитал бы смету под модель по умолчанию,
    # а прогон пошёл бы под выбранную. Сторожит tests/test_full_run.py.
    rv_model: Optional[str] = None
    # Разрешение ревизии читать и переписывать заверенное человеком. СВОЁ,
    # а не общее с ремонтом: у того флаг значит «правь по конкретным
    # находкам», здесь — «перечитай и перепиши целиком». Одна галочка на два
    # разных решения означала бы, что человек, разрешивший точечный ремонт,
    # молча разрешил и переписывание заверенного текста.
    rv_confirmed: bool = False
    use_judge: bool = False
    # Судья и выше потолка зоны — разовое разрешение прогона, как
    # include_confirmed. Осушает корзину «смысл не читал никто».
    judge_all: bool = False
    retry: bool = False
    include_confirmed: bool = False   # чинить и заверенное человеком (только ремонт)


@app.post("/api/projects/{pid}/run-plan")
def run_plan(pid: int, req: RunPlanRequest):
    """Что сделает составной прогон с этими настройками — до его запуска.

    Клиент не считает состав сам и не подбирает его галочками: и то и другое
    он делал по своим правилам, а деньги тратились по серверным."""
    project = get_project(pid)
    # is not None, а не truthy: пустой список — это «не выбрано ни одного шага»,
    # и разбирать надо ноль шагов, а не весь конвейер.
    want = set(req.steps if req.steps is not None else FULL_RUN_STEPS)
    steps = [s for s in FULL_RUN_STEPS if s in want]
    id_filter = set(req.segment_ids) if req.segment_ids is not None else None
    scope = [s for s in project["segments"]
             if id_filter is None or s["id"] in id_filter]
    params = req.dict()
    # Тем же предикатом, что и сам перевод (_needs_translation): от этого
    # множества зависит, попадут ли сегменты в состав ПРОВЕРОК — к своему
    # шагу они уже будут переведены. Считай его по «status == new», и
    # сегмент, перевод которого сорвался, прогон переведёт и проверит,
    # а разбор скажет про него «нет перевода» и занизит смету.
    will_translate = ({s["id"] for s in scope if _needs_translation(s)}
                      if "translate" in steps else set())
    # Один расчёт соответствия глоссарию на весь разбор, из общего кэша по
    # отпечатку проекта — тот же, которым живёт отчёт «Соответствие глоссарию».
    # Только «утверждённого термина в переводе НЕТ»: это работа ремонта, и
    # `_repair_findings(seg, None)` её не видит — она зовётся БЕЗ проекта ради
    # скорости. Расхождения по НАЧЕРТАНИЮ (`caseSegments`) сюда не входят:
    # их чинит бесплатная детерминированная команда, и звать ради них модель
    # значит платить за то, что делается точно и даром.
    impact = (glossary_impact(pid)
              if ("repair" in steps or "termaudit" in steps) else None)
    gloss_ids = set(impact["segments"]) if (impact and "repair" in steps) else set()
    # Разнобой по документу считается тем же готовым проходом, что и глоссарий:
    # разбор зовёт `_repair_findings` БЕЗ проекта ради скорости, и без этого
    # списка обещал бы меньше работы, чем сделает прогон.
    consist_ids = set()
    if "repair" in steps:
        for _pr in _consistency_of(project):
            consist_ids.update(_pr["segments"])
    # Сегменты, где есть что сверять, берём из того же кэшированного отчёта.
    # Считать `_verified_hits` заново — 13 мс на сегмент, то есть сорок секунд
    # заблокированного воркера на каждый разбор; ровно та беда, из-за которой
    # ремонт уже ходит сюда, а не считает глоссарий сам.
    term_ids = set(impact["termSegments"]) if (impact and "termaudit" in steps) else set()
    # Ревизию разбираем ПЕРВОЙ: её число нужно остальным шагам, чтобы честно
    # сказать, сколько работы она им создаст. Второй раз её не считаем.
    rv_plan = (_plan_step(project, "review", params, scope, will_translate,
                          gloss_ids, term_ids, consist_ids)
               if "review" in steps else None)
    rv_takes = rv_plan["count"] if rv_plan else 0
    plans = [(rv_plan if st == "review" else
              _plan_step(project, st, params, scope, will_translate, gloss_ids,
                         term_ids, consist_ids, rv_takes))
             for st in steps]
    # Объединение — в порядке ДОКУМЕНТА, а не в порядке шагов: порции идут по
    # этому списку, и прогон должен двигаться по тексту сверху вниз, а не
    # прыгать по проекту в зависимости от того, какому шагу сегмент достался.
    seen: set = set()
    for p in plans:
        seen.update(p["ids"])
    ids = [s["id"] for s in scope if s["id"] in seen]
    return {"steps": plans, "ids": ids, "total": len(ids), "scope": len(scope),
            # Сколько сегментов у проекта на сервере. Браузер держит проект
            # с момента загрузки страницы, а прибавиться они могут где угодно
            # — например, разбором картинок с соседнего экрана. Тогда состав
            # прогона (его считает сервер) говорит про 41 непереведённый
            # сегмент, а в таблице их нет и выбрать их нечем. Число дешёвое,
            # ответ и так приходит на каждую правку настроек.
            "projectSegments": len(project["segments"]),
            # То же самое про ИЗМЕНИВШИЕСЯ сегменты, а не про добавленные.
            # Прогон идёт на сервере, а вкладка забирает только те сегменты,
            # что сервер назвал в job.recent; пропустила порцию — и держит
            # допрогонные статусы, а число сегментов сходится и молчит.
            "projectStatus": _status_counts(project)}


JOB_HISTORY = 30                # сколько завершённых прогонов помним
JOB_CHUNK_RETRIES = 2           # повтор порции при сбое: сеть моргнула — не всё потеряно
JOB_RETRY_PAUSE = 5             # секунд между попытками

_JOBS: dict = {}                # id -> job
_SERVER_STARTED = time.time()
_JOB_QUEUE = _queue.Queue()
_JOBS_LOCK = threading.Lock()
_JOB_WORKER = None


def _job_live(pid: int) -> bool:
    """Есть ли у проекта ЛЮБАЯ живая задача.

    Рядом с `_job_busy`, но вопрос другой: тому важен вид задачи (разбор
    картинок пишет свой файл), а пересчёт оценок сталкивается с любым прогоном
    — все они пишут `seg["backcheck"]` или сам перевод."""
    return any(j["project"] == pid and j["status"] in ("queued", "running")
               for j in list(_JOBS.values()))


def _job_busy(pid: int, kind: str) -> bool:
    """Есть ли у проекта живая задача такого рода. Нужен тем командам, что
    правят те же данные, что и задача: два писателя одного файла — это
    молчаливая потеря работы одного из них."""
    return any(j["project"] == pid and j["kind"] == kind
               and j["status"] in ("queued", "running")
               for j in list(_JOBS.values()))


def _job_public(job: dict) -> dict:
    """Наружу отдаём без внутренних полей (список id и флаг остановки)."""
    return {k: v for k, v in job.items() if k not in ("ids", "stop")}


def _first_error(result: dict) -> str:
    """Текст первой ошибки порции. Пакетные эндпоинты глотают ошибку каждого
    сегмента по отдельности, и наружу шло только их число — а «ошибок: 10»
    не отличить от «кончились деньги на счёте» и «отозван ключ»."""
    for e in (result.get("errors") or []):
        txt = e.get("error") if isinstance(e, dict) else str(e)
        if txt:
            return str(txt)[:180]
    return ""


def _job_chunk_full(pid: int, chunk: list, params: dict) -> dict:
    """Порция составного прогона: все шаги подряд, в порядке FULL_RUN_STEPS.

    Шаг, для которого нет ключа или модуля, не роняет прогон — он записывается
    в blocked и работа идёт дальше: отсутствие ключа OpenAI не повод терять
    уже сделанный перевод. Но если НИ ОДИН шаг не отработал, порция считается
    провалившейся: молча рапортовать «выполнено», не сделав ничего, нельзя."""
    # Порядок берём из FULL_RUN_STEPS, а не из присланного списка: клиент выбирает
    # СОСТАВ шагов, но не их очерёдность — она несущая (см. комментарий там же).
    want = set(params.get("steps") or FULL_RUN_STEPS)
    steps = [s for s in FULL_RUN_STEPS if s in want]
    out = {"done": len(chunk)}
    ran, blocked = [], []
    for st in steps:
        if _job_should_stop():
            break
        # Medical QA сообщает о своей недоступности пятисоткой посегментно, а не
        # 503 на пакет: без этой проверки отсутствие модуля выглядело бы как
        # «порция целиком провалилась» и роняло весь прогон.
        if st == "medical_qa" and not (checks_mod and checks_enabled()):
            blocked.append("Medical QA: модуль недоступен")
            continue
        sub = dict(params)
        # У каждого шага своя модель: одна переводит, другие проверяют. Подшаги
        # читают её из params["model"], поэтому подставляем нужную перед вызовом —
        # иначе back-check пошёл бы той же моделью, что делала перевод, и перестал
        # быть независимой проверкой.
        mkey = FULL_STEP_MODEL.get(st)
        if mkey:
            # or None, а не «если задано»: незаполненная модель шага означает
            # «возьми свою по умолчанию», а не «возьми ту, что переводила».
            # Иначе back-check молча шёл бы моделью переводчика и переставал
            # быть независимой проверкой — а на ней стоит автоодобрение.
            sub["model"] = params.get(mkey) or None
        # Разрешение трогать заверенное человеком относится ТОЛЬКО к ремонту:
        # он меняет минимум слов по конкретным находкам и откатывается, если
        # стало хуже. Отдай тот же флаг переводу — и одна галочка «починить
        # подтверждённые» перегнала бы их заново целиком, по полной цене
        # и без единой находки в основании. Это и есть точечность: правим
        # найденное, а не переводим сегмент сначала.
        sub["include_confirmed"] = bool(params.get("include_confirmed")) if st == "repair" else False
        # То же и для разрешения ревизии: страховка от будущего читателя,
        # который прочтёт чужой флаг и молча расширит себе права.
        sub["rv_confirmed"] = bool(params.get("rv_confirmed")) if st == "review" else False
        if st == "translate":
            # Уже переведённое не переводим заново: составной прогон гоняют
            # по всему проекту, и force затирал бы готовые переводы.
            sub["force"] = False
        if st in ("backcheck", "termcheck", "medical_qa"):
            # Свежую проверку второй раз не оплачиваем — за состав отвечает
            # отбор сегментов, а не повторный вызов модели.
            sub["skip_cached"] = True
        try:
            r = _job_chunk(st, pid, chunk, sub)
        except HTTPException as e:
            if e.status_code == 503:
                blocked.append("%s: %s" % (FULL_STEP_LABELS[st], e.detail))
                continue
            raise
        if r.get("done", 0) == 0 and r.get("errors", 0) >= len(chunk):
            # Называем ПРИЧИНУ, а не только симптом. «Порция завершилась
            # ошибкой» человек прочтёт как поломку программы и полезет
            # в настройки, тогда как на деле у аккаунта кончились деньги или
            # отозван ключ — и это видно только в журнале сервера.
            why = r.get("why") or ""
            raise RuntimeError("порция целиком завершилась ошибкой на шаге «%s»%s"
                               % (FULL_STEP_LABELS[st], (": " + why) if why else ""))
        ran.append(st)
        for k, v in r.items():
            key = st if k == "done" else k
            if isinstance(v, str):
                # Текстовые поля (причина ошибки) не складываются: берём первое
                # непустое. Складывать их с числом — TypeError посреди прогона.
                if v and not out.get(key):
                    out[key] = v
                continue
            out[key] = out.get(key, 0) + v
    if blocked and not ran:
        raise RuntimeError("ни один шаг не выполнен — " + "; ".join(blocked))
    if blocked:
        # Счётчики прогона суммируются по всем порциям, поэтому это число —
        # «сколько раз шаг пропускался», а не «сколько шагов». Называем его так
        # и на экране: три недоступных шага на 534 порциях дают 1602, и подпись
        # «пропущено шагов: 1602» была бы враньём.
        out["step_skips"] = len(blocked)
        print("[backend] составной прогон: пропущены шаги — " + "; ".join(blocked),
              file=sys.stderr)
    return out


def _job_chunk(kind: str, pid: int, chunk: list, params: dict) -> dict:
    """Одна порция. Возвращает счётчики, которые нарастают в прогрессе."""
    n = len(chunk)
    if kind == "full":
        return _job_chunk_full(pid, chunk, params)
    if kind == "apply_terms":
        # Ремонт по свежеодобренным терминам. Само одобрение пачки делает
        # первая порция (см. _job_run: ставится флаг), дальше идёт обычный
        # ремонт — расхождение с глоссарием у него такая же находка, как
        # потерянный термин. Отдельного «переперевода» тут нет намеренно:
        # ремонт меняет минимум слов и откатывается, если стало хуже.
        r = repair_batch(pid, RepairBatchRequest(
            segment_ids=chunk, limit=n, model=params.get("rp_model") or params.get("model"),
            bc_model=params.get("bc_model"), tc_model=params.get("tc_model"),
            use_judge=bool(params.get("use_judge")), judge_model=params.get("judge_model"),
            judge_all=bool(params.get("judge_all")),
            include_confirmed=bool(params.get("include_confirmed")), retry=True))
        # done = сделанное, а не размер порции: иначе проверка «порция целиком
        # завершилась ошибкой» никогда не сработает, и мёртвый ключ выглядел бы
        # как «выполнено, исправлено 0».
        return {"done": len(r.get("applied", [])) + len(r.get("skipped", [])),
                "applied": len(r.get("applied", [])), "why": _first_error(r),
                "reverted": len(r.get("skipped", [])), "errors": len(r.get("errors", [])),
                "desync": len(r.get("desync", [])),
                "skipped_confirmed": len(r.get("skipped_confirmed", []))}
    if kind == "translate":
        r = batch_translate(pid, BatchRequest(
            segment_ids=chunk, limit=n,
            force=bool(params.get("force", True)), model=params.get("model"),
            include_confirmed=bool(params.get("include_confirmed"))))
        return {"done": r["count"], "tm_hits": r.get("tm_hits", 0),
                "duplicates": r.get("duplicates", 0), "errors": len(r.get("errors", [])),
                "why": _first_error(r), "skipped_confirmed": len(r.get("skipped_confirmed", []))}
    if kind == "backcheck":
        r = backcheck_batch(pid, BackcheckBatchRequest(
            segment_ids=chunk, limit=n, model=params.get("model"),
            use_judge=bool(params.get("use_judge")), judge_model=params.get("judge_model"),
            judge_all=bool(params.get("judge_all")),
            skip_cached=bool(params.get("skip_cached", False))))
        return {"done": r["count"], "duplicates": r.get("duplicates", 0),
                "skipped_cached": r.get("skipped_cached", 0),
                "errors": len(r.get("errors", [])), "why": _first_error(r)}
    if kind == "termaudit":
        # Сверка ВСЕХ приказных терминов, а не разбор спора: сюда шаг и заведён.
        # `refresh=False` — свежий вердикт второй раз не оплачиваем, за состав
        # отвечает отбор сегментов (см. _plan_step).
        r = term_context(pid, TermContextRequest(
            segment_ids=chunk, limit=n, model=params.get("model"),
            all_terms=True, refresh=False))
        # done — вся порция, а не только спрошенное: сегменты без приказных
        # терминов пройдены, просто спрашивать в них нечего. Иначе полоса
        # прогона доходила бы до четверти и там останавливалась.
        return {"done": (r.get("asked", 0) + r.get("nothingToCheck", 0)
                         + r.get("cachedSkipped", 0)),
                "asked": r.get("asked", 0),
                "settled": len(r.get("settled") or []),
                "wrong": len(r.get("wrong") or []),
                "nothing": r.get("nothingToCheck", 0),
                "skipped_cached": r.get("cachedSkipped", 0),
                "errors": len(r.get("failed") or []),
                # Строкой, а не None: счётчики порции складываются, и None
                # среди них — TypeError посреди прогона.
                "why": ("арбитр не ответил" if r.get("failed") else "")}
    if kind == "review":
        # Ревизия ПРАВИТ текст, поэтому идёт с dry_run=False. Заверенное
        # человеком она читает и переписывает ТОЛЬКО по своему разрешению
        # (`rv_confirmed`); без него `_review_pick` их не спрашивает вовсе —
        # платить за совет, который некуда применить, незачем.
        r = review_project(pid, ReviewRequest(
            segment_ids=chunk, limit=n, model=params.get("model"),
            dry_run=False,
            # Именно `rv_confirmed`: общий `include_confirmed` доезжает только
            # до ремонта (`_job_chunk_full` гасит его остальным), и разрешение
            # на точечную починку не должно означать разрешения на
            # переписывание заверенного целиком.
            include_confirmed=bool(params.get("rv_confirmed")),
            stamp=params.get("review_stamp")))
        # done — ОТВЕЧЕННЫЕ, без провалов. Складывать провалы сюда нельзя:
        # `_job_chunk_full` роняет прогон по условию «done == 0 и ошибок
        # не меньше порции», и с провалами внутри done оно недостижимо —
        # отозванный ключ или кончившиеся деньги выглядели бы как
        # «выполнено, исправлено 0», то есть ровно тем враньём, ради которого
        # это условие и заведено.
        # Имена счётчиков СВОИ, а не общие с ремонтом: счётчики порций
        # складываются в прогресс задачи, и общий `applied` показал бы сумму
        # двух разных работ одним числом — «исправлено 40» без ответа на
        # вопрос, чем именно и что теперь смотреть человеку.
        return {"done": r.get("answered", 0),
                "revised": r.get("applied", 0),
                "review_proposed": r.get("proposed", 0),
                "suspect": len(r.get("sourceSuspect") or []),
                "errors": r.get("failed", 0),
                "skipped_confirmed": len(r.get("skippedConfirmed") or []),
                # Метка — СТРОКА, и счётчики её не складывают (первое непустое
                # значение побеждает): без неё человек не знает, чем отменить
                # прогон, переписавший сотни сегментов.
                "reviewStamp": r.get("stamp") or "",
                "why": ("ревизор не ответил" if r.get("failed") else "")}
    if kind == "termcheck":
        r = termcheck_batch(pid, TermcheckBatchRequest(
            segment_ids=chunk, limit=n, model=params.get("model"),
            skip_cached=bool(params.get("skip_cached", False))))
        return {"done": r["count"], "flagged": r.get("flagged", 0),
                "duplicates": r.get("duplicates", 0), "why": _first_error(r),
                "skipped_trivial": r.get("skipped_trivial", 0), "errors": len(r.get("errors", []))}
    if kind == "repair":
        r = repair_batch(pid, RepairBatchRequest(
            segment_ids=chunk, limit=n, model=params.get("model"),
            bc_model=params.get("bc_model"), tc_model=params.get("tc_model"),
            use_judge=bool(params.get("use_judge")), judge_model=params.get("judge_model"),
            judge_all=bool(params.get("judge_all")),
            include_confirmed=bool(params.get("include_confirmed")),
            retry=bool(params.get("retry"))))
        return {"done": len(r.get("applied", [])) + len(r.get("skipped", [])),
                "applied": len(r.get("applied", [])), "reverted": len(r.get("skipped", [])),
                "errors": len(r.get("errors", [])), "why": _first_error(r),
                "desync": len(r.get("desync", [])),
                "skipped_confirmed": len(r.get("skipped_confirmed", []))}
    if kind == "medical_qa":
        r = batch_checks(pid, ChecksBatchRequest(
            segment_ids=chunk, limit=n, bc_model=params.get("bc_model"),
            skip_cached=bool(params.get("skip_cached", False))))
        return {"done": r.get("count", 0), "errors": len(r.get("errors", [])),
                "why": _first_error(r), "skipped_cached": r.get("skipped_cached", 0)}
    raise ValueError("unknown job kind: " + kind)


def _job_run(job: dict):
    kind, pid = job["kind"], job["project"]
    # Организация прогона — в его поток: ContextVar сессии сюда не доезжает,
    # а get_project и области считаются по ней.
    _JOB_TENANT.id = job.get("tenant") or DEFAULT_TENANT
    _JOB_LANG.code = job.get("lang") or DEFAULT_UI_LANG
    chunk_size = JOB_CHUNKS[kind]
    if kind == "images":
        # Разбор картинок не идёт по сегментам: их ещё нет — они из него
        # и рождаются. Цикл свой, но остановка, счётчики и сохранение общие.
        _job_images(job)
        if job["status"] == "running":
            job["status"] = "done"
        job["finished"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        return
    if kind == "termsheet":
        _job_termsheet(job)
        if job["status"] == "running":
            job["status"] = "done"
        job["finished"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        return
    if kind == "apply_terms":
        # Одобрение пачки — один шаг, а не порция: оно про глоссарий, а не про
        # сегменты. Состав сегментов пересчитываем ПОСЛЕ него: до одобрения
        # неизвестно, какие сегменты разойдутся с новыми терминами, и клиент
        # физически не может прислать правильный список.
        if job["stop"]:
            # Остановили до записи — значит ничего и не записываем. Проверка
            # обязана стоять ДО одобрения: глоссарий меняется одним куском,
            # и «остановлено» после него было бы неправдой.
            job["status"] = "stopped"
            job["finished"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            return
        try:
            res = auto_approve_terms(AutoApproveRequest(
                dry_run=False, project=pid,
                max_tier=job["params"].get("max_tier"),
                allow_verified=job["params"].get("allow_verified"),
                limit=int(job["params"].get("term_limit", 2000))))
        except HTTPException as e:
            job["status"], job["error"] = "error", f"{e.status_code}: {e.detail}"
            job["finished"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            return
        job["counters"]["termsApproved"] = res["counts"]["verified"] + res["counts"]["auto"]
        job["counters"]["termsVerified"] = res["counts"]["verified"]
        job["counters"]["termsClosed"] = res["counts"]["closed"]
        job["counters"]["termsRejected"] = res["counts"].get("rejectedMeaning", 0)
        job["autoBatch"] = res.get("batch")
        # Чиним ТОЛЬКО то, что разошлось с утверждёнными терминами, а не всё,
        # где вообще есть находки: человек нажал «применить термины», а не
        # «перебрать весь проект». Список тот же, что показывает карточка
        # «Соответствие глоссарию», — расчёт один, чтобы цифры не расходились.
        imp = glossary_impact(pid, refresh=True)
        allow_conf = bool(job["params"].get("include_confirmed"))
        want = list(imp["segments"] if allow_conf else imp["pending"])
        # Сегменты, где тот же текст с теми же претензиями уже проходил ремонт,
        # не берём: заход вернёт тот же результат за те же деньги. Их видно
        # в отчёте о соответствии как «ремонт уже не берёт» — прятать нельзя,
        # расходиться с глоссарием они не перестали.
        project_obj = get_project(pid)
        by_id = {sg["id"]: sg for sg in project_obj["segments"]}
        futile = [i for i in want
                  if i in by_id and _repair_futile(by_id[i], project_obj)]
        job["ids"] = [i for i in want if i not in set(futile)]
        job["counters"]["futile"] = len(futile)
        job["total"] = len(job["ids"])
        save_state(STATE)
    ids = job["ids"]
    # Одна метка отката на ВЕСЬ прогон: ревизия идёт порциями по пять
    # сегментов, и без общей метки книга дала бы ~250 копий по одному-два
    # сегмента — отменить прогон целиком было бы нечем. Кладём в params,
    # потому что до `_job_chunk` доезжают именно они; `setdefault` —
    # чтобы повтор задачи после рестарта не завёл вторую метку.
    if kind in ("full", "review"):
        if not job["params"].get("review_stamp"):
            job["params"]["review_stamp"] = _backup_stamp("review")
            # Сохраняем СРАЗУ: иначе рестарт поднимет задачу с params без
            # метки, она заведёт вторую, а первая копия останется сиротой —
            # её имени не будет ни в счётчиках, ни в отчёте.
            _job_persist(job)
    for i in range(0, len(ids), chunk_size):
        if job["stop"]:
            job["status"] = "stopped"
            break
        chunk = ids[i:i + chunk_size]
        job["recent"] = chunk          # клиент подтянет только эти сегменты
        last_err = None
        for attempt in range(JOB_CHUNK_RETRIES + 1):
            try:
                counters = _job_chunk(kind, pid, chunk, job["params"])
                # Порция, где не прошло НИЧЕГО, — это не «выполнено»: так выглядит
                # мёртвый ключ или упавший провайдер. Пакетные эндпоинты глотают
                # ошибку каждого сегмента по отдельности, и без этой проверки
                # прогон рапортовал бы «готово» с нулевым результатом.
                if counters.get("done", 0) == 0 and counters.get("errors", 0) >= len(chunk):
                    # С причиной от провайдера: «проверьте ключи и связь» звучит
                    # одинаково и для кончившихся денег, и для отозванного ключа,
                    # и для оборванной сети — а действия у них разные.
                    raise RuntimeError("порция целиком завершилась ошибкой"
                                       + (": " + counters["why"] if counters.get("why")
                                          else " — проверьте ключи и связь"))
                for k, v in counters.items():
                    if k == "done":
                        job["done"] += v
                    elif isinstance(v, str):
                        # Причина ошибки — текст, а не счётчик. Складывать её
                        # с числом значит уронить TypeError посреди прогона,
                        # причём молча: общий обработчик повторит порцию
                        # трижды, оплатив её вызовы заново.
                        if v and not job["counters"].get(k):
                            job["counters"][k] = v
                    else:
                        job["counters"][k] = job["counters"].get(k, 0) + v
                last_err = None
                break
            except HTTPException as e:
                # 503 (нет ключа) или 404 — повторять бессмысленно
                last_err = f"{e.status_code}: {e.detail}"
                break
            except Exception as e:
                last_err = str(e)
                print(f"[backend] job#{job['id']} chunk failed ({attempt + 1}): {e}", file=sys.stderr)
                if attempt < JOB_CHUNK_RETRIES:
                    time.sleep(JOB_RETRY_PAUSE)
        if last_err:
            job["status"] = "error"
            job["error"] = last_err
            break
    if job["status"] == "running":
        job["status"] = "done"
    # Метка резервируется файлом ДО первой правки — если правок так и не
    # случилось, убираем пустышку, чтобы каталог копий не зарастал.
    _backup_drop_empty("review", job["params"].get("review_stamp"))
    job["finished"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def _job_loop():
    while True:
        job = _JOB_QUEUE.get()
        _job_execute(job)
        _JOB_QUEUE.task_done()


def _job_execute(job: dict):
    """Одна задача от начала до конца: статусы, расход, сохранение. Вынесено
    из цикла, потому что исполнителей два: поток в процессе API (файловое
    хранилище) и отдельный процесс medcat-worker (base + claim_job)."""
    # Перед прогоном — свежие глоссарий и очередь: их могла править
    # другая сторона (после разделения на процессы — обязательно).
    _sync_shared(force=True)
    if True:
        dropped_at_start = _TERM_DROPPED["total"]
        try:
            if job["stop"]:
                job["status"] = "stopped"
                job["finished"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                return
            job["status"] = "running"
            job["started"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            _job_persist(job)
            _ACTIVE_JOB["job"] = job
            _usage_begin(job)
            _job_run(job)
        except Exception as e:
            job["status"] = "error"
            job["error"] = str(e)
            job["finished"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            print(f"[backend] job#{job.get('id')} crashed: {e}", file=sys.stderr)
        finally:
            _ACTIVE_JOB.pop("job", None)
            # Сколько находок сбор терминологии потерял на потолке очереди.
            # Ноль — норма; не ноль означает, что часть платной работы прогона
            # ушла в никуда, и человек об этом узнает из отчёта, а не из
            # журнала, куда никто не смотрит.
            lost = _TERM_DROPPED["total"] - dropped_at_start
            if lost:
                job["counters"]["terms_dropped"] = lost
            # До save_state: запись о расходе должна уехать на диск вместе
            # с остальным результатом прогона, а не ждать следующего сохранения.
            _usage_end(job)
            try:
                save_state(STATE)
            except Exception as e:
                print(f"[backend] job#{job.get('id')} save failed: {e}", file=sys.stderr)
            _job_persist(job)


def _ensure_job_worker():
    global _JOB_WORKER
    if _JOB_WORKER is None or not _JOB_WORKER.is_alive():
        _JOB_WORKER = threading.Thread(target=_job_loop, name="mcat-jobs", daemon=True)
        _JOB_WORKER.start()


def _next_job_id() -> int:
    if STORE.kind == "pg":
        try:
            return STORE.next_counter("jobId", max(_JOBS, default=0))
        except Exception as e:
            print(f"[backend] счётчик задач из базы не взялся: {e}", file=sys.stderr)
    return max(_JOBS, default=0) + 1


def _refresh_jobs_from_db(force: bool = False) -> None:
    """Зеркало задач для API при внешнем воркере: прогресс и статусы пишет
    воркер в базу, здесь их только показывают. Не чаще раза в 2 секунды."""
    if not EXTERNAL_WORKER or IS_WORKER:
        return
    now = time.time()
    if not force and now - _JOBS_REFRESH["t"] < 2:
        return
    _JOBS_REFRESH["t"] = now
    try:
        for j in STORE.load_jobs():
            j.setdefault("stop", False)
            j.setdefault("recent", [])
            _JOBS[j["id"]] = j
    except Exception as e:
        print(f"[backend] зеркало задач не обновлено: {e}", file=sys.stderr)


_JOBS_REFRESH = {"t": 0.0}


def _job_persist(job: dict) -> None:
    try:
        STORE.save_job(job)
    except Exception as e:
        print(f"[backend] job#{job.get('id')} не сохранён в базу: {e}", file=sys.stderr)


def _trim_jobs():
    finished = [j for j in _JOBS.values() if j["status"] in ("done", "stopped", "error")]
    for j in sorted(finished, key=lambda x: x["id"])[:-JOB_HISTORY]:
        _JOBS.pop(j["id"], None)
        try:
            STORE.delete_job(j["id"])
        except Exception as e:
            print(f"[backend] job#{j['id']} не удалён из базы: {e}", file=sys.stderr)


class JobRequest(BaseModel):
    kind: str
    segment_ids: List[int]
    params: dict = {}


@app.post("/api/projects/{pid}/jobs")
def create_job(pid: int, req: JobRequest):
    """Поставить прогон в очередь. Клиенту достаточно отдать список сегментов —
    дальше страница может быть закрыта, сервер доведёт работу до конца."""
    _audit("job.create", project=pid, kind=req.kind)
    get_project(pid)                        # 404, если проекта нет
    if req.kind not in JOB_KINDS:
        raise HTTPException(400, "Неизвестный тип прогона: " + req.kind)
    ids = list(dict.fromkeys(req.segment_ids))
    # apply_terms сам считает состав после одобрения терминов — до него список
    # сегментов ещё неизвестен, и требовать его от клиента бессмысленно.
    if not ids and req.kind not in ("apply_terms", "images", "termsheet"):
        # apply_terms считает состав после одобрения терминов, а разбор
        # картинок — сам себе состав: сегментов из картинок ещё не существует.
        raise HTTPException(400, "Пустой список сегментов")
    with _JOBS_LOCK:
        job = {
            "id": _next_job_id(),
            "kind": req.kind, "project": pid, "status": "queued",
            "tenant": _current_tenant(),
            "user": (CURRENT_SESSION.get() or {}).get("user"),
            # Язык объяснений — того, кто запустил. В поток прогона ContextVar
            # не доезжает, поэтому задача несёт его в себе (как организацию).
            "lang": _explain_lang(),
            "total": len(ids), "done": 0, "counters": {}, "error": None,
            "params": dict(req.params or {}),
            "created": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "started": None, "finished": None,
            "ids": ids, "stop": False, "recent": [],
        }
        _JOBS[job["id"]] = job
        _trim_jobs()
    _job_persist(job)
    if EXTERNAL_WORKER:
        # Задачу заберёт medcat-worker из таблицы jobs (claim_job).
        pass
    else:
        _ensure_job_worker()
        _JOB_QUEUE.put(job)
    return {"ok": True, "job": _job_public(job)}


@app.get("/api/usage")
def usage_report(limit: int = 20):
    """Сколько прогоны стоили на самом деле — рядом с тем, во что их оценили.

    Ни одного вызова модели здесь нет: всё уже посчитано провайдером и снято
    с ответов. `process` — расход с момента старта сервиса, включая одиночные
    вызовы по кнопке, которые ни одному прогону не принадлежат."""
    t = _current_tenant()
    runs = [r for r in reversed(STATE.get("runCosts") or [])
            if _tenant_of(r) == t][:max(1, min(limit, 100))]
    priced = [r for r in runs if r.get("est") and r.get("cost")]
    return {
        "process": _USAGE_TOTAL,
        "runs": runs,
        # Во сколько раз смета в среднем расходится с фактом. Это и есть та
        # поправка, которой в системе не было: без неё «ориентировочно $15»
        # не с чем сравнить, и остаётся гадать, врёт она или прогон не отработал.
        "estRatio": (round(sum(r["est"] for r in priced) / sum(r["cost"] for r in priced), 2)
                     if priced and sum(r["cost"] for r in priced) else None),
        "estRuns": len(priced),
    }


@app.get("/api/jobs")
def list_jobs(project: Optional[int] = None, limit: int = 20):
    _refresh_jobs_from_db()
    t = _current_tenant()
    jobs = [j for j in _JOBS.values() if _tenant_of(j) == t
            and (project is None or j["project"] == project)]
    jobs.sort(key=lambda j: j["id"], reverse=True)
    active = [_job_public(j) for j in jobs if j["status"] in ("queued", "running")]
    return {"active": active, "jobs": [_job_public(j) for j in jobs[:max(1, min(limit, 100))]]}


@app.get("/api/jobs/{jid}")
def get_job(jid: int):
    _refresh_jobs_from_db()
    job = _JOBS.get(jid)
    if not job or (_tenant_of(job) != _current_tenant()
                   and not (CURRENT_SESSION.get() or {}).get("super")):
        raise HTTPException(404, "Прогон не найден")
    return {"ok": True, "job": _job_public(job)}


@app.post("/api/jobs/{jid}/stop")
def stop_job(jid: int):
    """Остановка мягкая: текущая порция досчитывается и сохраняется, следующая
    не начинается. Обрывать порцию на середине значило бы заплатить за вызовы
    и выбросить их результат."""
    _refresh_jobs_from_db(force=True)
    job = _JOBS.get(jid)
    if not job or (_tenant_of(job) != _current_tenant()
                   and not (CURRENT_SESSION.get() or {}).get("super")):
        raise HTTPException(404, "Прогон не найден")
    job["stop"] = True
    if EXTERNAL_WORKER:
        # Стоп-флаг доезжает до воркера через базу: он читает его между
        # сегментами (_job_should_stop) и мягко останавливается.
        _job_persist(job)
    if job["status"] == "queued":
        job["status"] = "stopped"
        job["finished"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    elif job["status"] == "running":
        # Промежуточный статус: обработка текущего сегмента доигрывается, но
        # пользователю сразу видно, что кнопка сработала.
        job["stopping"] = True
    return {"ok": True, "job": _job_public(job)}


# ─────────────────────────────────────────────────────────────────────
# Static file serving (the React design)
# Mounted last so /api/* takes precedence.
# ─────────────────────────────────────────────────────────────────────
if FRONTEND_DIR.exists():
    app.mount("/css", StaticFiles(directory=str(FRONTEND_DIR / "css")), name="css")
    app.mount("/js", StaticFiles(directory=str(FRONTEND_DIR / "js")), name="js")
    app.mount("/screens", StaticFiles(directory=str(FRONTEND_DIR / "screens")), name="screens")

    @app.get("/", response_class=HTMLResponse)
    def index():
        return FileResponse(str(FRONTEND_DIR / "index.html"))

    # Вход в админку — по нестандартному адресу, а не /admin. Адрес — из
    # ADMIN_PATH в окружении, иначе выводится из APP_PASSWORD (стабилен для
    # установки, не угадывается) и печатается в журнал при старте. Это
    # обфускация входа, а не защита: право на /api/admin/* даёт роль super.
    @app.get("/terms", response_class=HTMLResponse)
    def legal_terms():
        if not legal_mod:
            raise HTTPException(503, "Документы недоступны")
        return HTMLResponse(legal_mod.page("terms"))

    @app.get("/privacy", response_class=HTMLResponse)
    def legal_privacy():
        if not legal_mod:
            raise HTTPException(503, "Документы недоступны")
        return HTMLResponse(legal_mod.page("privacy"))

    @app.get("/" + ADMIN_PATH, response_class=HTMLResponse)
    def admin_console():
        html = (FRONTEND_DIR / "index.html").read_text(encoding="utf-8")
        marker = "<script>window.ADMIN_ENTRY=true;</script>"
        html = html.replace("<head>", "<head>" + marker, 1) if "<head>" in html else marker + html
        return HTMLResponse(html)
else:
    @app.get("/", response_class=HTMLResponse)
    def index_fallback():
        return "<h1>Frontend directory not found</h1><p>Expected at: " + str(FRONTEND_DIR) + "</p>"


if __name__ == "__main__":
    import uvicorn
    print(f"[backend] Starting Medical CAT Translator API on http://localhost:8000")
    print(f"[backend] Frontend dir: {FRONTEND_DIR}")
    print(f"[backend] Loaded modules: {list(_BACKEND_MODULES.keys())}")
    uvicorn.run(app, host="0.0.0.0", port=8000)


def _restore_jobs() -> None:
    """Очередь прогонов из базы после рестарта. Незаконченные (queued/running)
    ставятся в очередь заново: порция, оборванная рестартом, повторится, но
    готовые проверки у сегментов кэшированы — второй раз они не оплачиваются.
    Файловое хранилище очередь не хранит: там, как и раньше, рестарт её теряет."""
    try:
        jobs = STORE.load_jobs()
    except Exception as e:
        print(f"[backend] очередь прогонов из базы не прочитана: {e}", file=sys.stderr)
        return
    requeued = 0
    for j in jobs:
        j.setdefault("stop", False)
        j.setdefault("recent", [])
        j.setdefault("ids", [])
        j.setdefault("counters", {})
        _JOBS[j["id"]] = j
        # При внешнем воркере очередь ведёт он сам (claim_job + сброс running
        # на своём старте); API держит только зеркало для показа.
        if not EXTERNAL_WORKER and j.get("status") in ("queued", "running"):
            j["status"], j["stop"], j["started"] = "queued", False, None
            _JOB_QUEUE.put(j)
            requeued += 1
    if jobs:
        print(f"[backend] прогонов из базы: {len(jobs)}, поставлено в очередь заново: {requeued}",
              file=sys.stderr)
    if requeued:
        _ensure_job_worker()


_restore_jobs()
