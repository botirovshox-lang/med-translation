"""Разбор картинок проекта: найти надписи, прочитать, завести сегменты.

Что здесь сторожится:

  1. якорь второго вида — «часть пакета + номер блока», и сегмент из картинки
     встаёт РЯДОМ с ней, а не в конце списка: сегмент без соседей переводят
     без контекста;
  2. надпечатка аппарата (фамилия пациента, дата, настройки томографа)
     и шум сегментами НЕ становятся — это не текст документа, а в случае
     фамилий ещё и персональные данные, которым нечего делать в памяти
     переводов;
  3. разбор кэшируется по СОДЕРЖИМОМУ картинки: повторный заход не платит
     за уже прочитанное и не заводит те же сегменты второй раз;
  4. «не спросили» (None от модели) и «пусто» — разные вещи: первое обязано
     оставить блок нерешённым, иначе картинка навсегда считается пустой;
  5. разбор без исходника, без движка и без ключа отказывает СЛОВАМИ,
     а не отвечает «текста нет»;
  6. снос распознанного не уносит переводы: сегмент с готовым переводом —
     оплаченная работа, и без явного разрешения он остаётся.

Ни одного вызова модели: чтение подменяется заглушкой.
"""
import io
import os
import sys
import shutil
import tempfile
from pathlib import Path

os.environ.setdefault("APP_PASSWORD", "test")
os.environ["AUTHORITY_CORPUS"] = "0"
sys.path.insert(0, "backend")
import main                                                   # noqa: E402
import image_text                                             # noqa: E402

main.save_state = lambda *a, **k: None

fail = []


def check(cond, label):
    print(("  OK   " if cond else "  FAIL ") + label)
    if not cond:
        fail.append(label)


TMP = Path(tempfile.mkdtemp(prefix="medcat-images-"))
main.SOURCE_DIR = TMP / "sources"

from docx import Document                                     # noqa: E402
from docx.shared import Inches                                # noqa: E402
from PIL import Image, ImageDraw, ImageFont                   # noqa: E402


def picture(text: str, noise: bool = False) -> bytes:
    """Картинка из учебника как она есть: подпись документа плюс надпечатка
    аппарата (фамилия и дата) в стороне от неё.

    Надписи разнесены по вертикали больше чем на строку — иначе сборка в блоки
    справедливо склеит их в один текст. Кегль настоящий: штатный растровый
    шрифт Pillow даёт строку в восемь пикселей, которую детектор не видит."""
    if noise:
        import numpy as np
        px = (np.random.RandomState(3).rand(220, 520, 3) * 255).astype("uint8")
        im = Image.fromarray(px)
    else:
        im = Image.new("RGB", (520, 220), "white")
        d = ImageDraw.Draw(im)
        big = ImageFont.truetype(image_text.font_path(), 30)
        small = ImageFont.truetype(image_text.font_path(), 22)
        d.text((24, 20), text, fill=(10, 10, 10), font=big)              # блок 0
        d.text((330, 100), "KARIMOV SH.", fill=(10, 10, 10), font=small)  # блок 1
        d.text((24, 165), "25.03.2008", fill=(10, 10, 10), font=small)    # блок 2
    out = io.BytesIO()
    im.save(out, "PNG")
    return out.getvalue()


def build_source() -> bytes:
    doc = Document()
    doc.add_paragraph("Текст перед рисунком.")               # абзац 0
    doc.add_picture(io.BytesIO(picture("Рис. 1. Схема лёгких")),
                    width=Inches(3))                          # абзац 1 — картинка
    doc.add_paragraph("Текст после рисунка.")                # абзац 2
    doc.add_picture(io.BytesIO(picture("", noise=True)),
                    width=Inches(3))                          # абзац 3 — фотошум
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


content = build_source()
paras = main._docx_paragraphs(content)
units = main._docx_units(paras)
project = {
    "id": 1, "title": "Тест", "src": "RU", "tgt": "EN", "domain": "medical",
    "status": "in_progress", "segments": [
        {"id": i + 1, "source": t, "target": "", "status": "new", "comments": [],
         "qa": [], "wordCount": len(t.split()), "risk": "low",
         "route": "GPT_REQUIRED", "tm": None}
        for i, (t, _idx) in enumerate(units)],
}
main.STATE["projects"] = [project]
pairs = [[i, u + 1] for u, (_t, idxs) in enumerate(units) for i in idxs]
main._store_source_docx(project, content, "t.docx", pairs, len(paras))

print("\n── картинки в пакете ──")
raster, other = main._docx_media(content)
check(len(raster) == 2, "обе картинки найдены как растровые: %d" % len(raster))
check(other == [], "нерастровых частей нет")
anchors = main._docx_image_anchors(Document(io.BytesIO(content)))
check(len(anchors) == 2, "у обеих картинок есть абзац-якорь")
first = sorted(anchors.items())[0]
check(min(first[1]) in (1, 3), "якорь картинки — номер её абзаца: %s" % (first[1],))


def new_job(**params):
    return {"id": 1, "kind": "images", "project": 1, "status": "running",
            "total": 0, "done": 0, "counters": {}, "error": None,
            "params": params, "stop": False, "recent": [], "ids": []}


print("\n── отказы называются словами ──")
no_src = dict(project)
no_src.pop("sourceDocx")
main.STATE["projects"] = [no_src]
try:
    main._job_images(new_job(dry_run=True))
    check(False, "проект без исходника обязан отказать")
except RuntimeError as e:
    check("не приложен" in str(e), "без исходника: " + str(e))
main.STATE["projects"] = [project]

ready, why = image_text.engine_ready()
if not ready:
    try:
        main._job_images(new_job(dry_run=True))
        check(False, "без движка разбор обязан отказать, а не найти ноль")
    except RuntimeError as e:
        check("Поиск строк невозможен" in str(e), "без движка: " + str(e))
    print("\nДВИЖОК НЕ УСТАНОВЛЕН (%s) — дальше проверять нечем" % why)
    shutil.rmtree(TMP, ignore_errors=True)
    sys.exit(1 if fail else 0)

key = os.environ.pop("OPENAI_API_KEY", None)
try:
    main._job_images(new_job(dry_run=False))
    check(False, "чтение без ключа обязано отказать")
except RuntimeError as e:
    check("ключ" in str(e), "без ключа: " + str(e))
os.environ["OPENAI_API_KEY"] = key or "test-key"

print("\n── разбор без чтения (dry_run) ──")
job = new_job(dry_run=True)
main._job_images(job)
data = main._load_source_map(1)
images = data.get("images") or []
withtext = [im for im in images if im.get("blocks")]
check(len(images) == 2, "обе картинки в карте")
check(len(withtext) == 1, "надпись нашлась ровно на одной картинке")
check(all(im.get("sha") for im in images), "у каждой картинки записан отпечаток")
check(all("text" not in b for im in images for b in im["blocks"]),
      "разбор без чтения не выдумывает текст")
check(len(project["segments"]) == len(units), "разбор без чтения не заводит сегментов")

print("\n── чтение и сегменты ──")
calls = {"n": 0}
TEXT = "Рис. 1. Схема лёгких"


def fake_read(blob, blocks, src_lang, domain_id=None, model=None):
    calls["n"] += 1
    check(src_lang == "RU", "языку картинки не место в коде: пришёл из проекта")
    out = []
    for i, _b in enumerate(blocks):
        if i == 0:
            out.append({"text": TEXT, "overlay": False, "model": "test"})
        elif i == 1:
            out.append({"text": "KARIMOV SH.", "overlay": True, "model": "test"})
        else:
            out.append({"text": "25.03.2008", "overlay": False, "model": "test"})
    return out


main._openai_read_image = fake_read
job = new_job(dry_run=False)
main._job_images(job)
img_segs = [s for s in project["segments"] if (s.get("origin") or {}).get("kind") == "image"]
check(len(img_segs) == 1, "заведён ровно один сегмент — на настоящую подпись")
check(img_segs[0]["source"] == TEXT, "текст сегмента — распознанная надпись")
o = img_segs[0]["origin"]
check(o.get("part", "").startswith("word/media/") and "block" in o,
      "якорь сегмента — часть пакета и номер блока: %s" % o)

pos = [s["id"] for s in project["segments"]].index(img_segs[0]["id"])
check(project["segments"][pos - 1]["source"] == "Текст перед рисунком.",
      "сегмент картинки встал сразу за абзацем перед ней")

data = main._load_source_map(1)
blocks = [b for im in data["images"] for b in (im.get("blocks") or [])]
check(any(b.get("skip") == "overlay" for b in blocks),
      "надпечатка аппарата помечена и сегментом не стала")
check(any(b.get("skip") == "noise" for b in blocks),
      "дата помечена шумом и сегментом не стала")
check(job["counters"].get("segments") == 1, "счётчик задачи назвал заведённое")

print("\n── повторный заход ──")
was = calls["n"]
main._job_images(new_job(dry_run=False))
img_segs2 = [s for s in project["segments"] if (s.get("origin") or {}).get("kind") == "image"]
check(len(img_segs2) == 1, "повтор не задваивает сегменты")
check(calls["n"] == was, "повтор не платит за уже прочитанное")

print("")
print("── отсеянное видно и возвращается ──")
# «Отсеяно: 230» — число, которое человеку нечем проверить, а отсев делает
# модель и ошибается в обе стороны. Значит список обязан быть, и обратный ход
# тоже: решение машины, которое нельзя отменить, — это не помощь, а приговор.
blocks = main.images_blocks(1)
check(blocks["total"] >= 3 and all("text" in b for b in blocks["blocks"]),
      "надписи отдаются списком: %d" % blocks["total"])
over = main.images_blocks(1, skip="overlay")
check(over["total"] >= 1 and all(b["skip"] == "overlay" for b in over["blocks"]),
      "отсеянное как надпечатка отбирается отдельно")
kept = main.images_blocks(1, skip="none")
check(all(b["skip"] is None for b in kept["blocks"]) and kept["total"] >= 1,
      "ставшее сегментами тоже видно")

o = over["blocks"][0]
was = len(project["segments"])
res = main.image_restore_block(1, main.ImageRestoreRequest(part=o["part"], block=o["block"]))
check(res["ok"] and res["created"] and len(project["segments"]) == was + 1,
      "отсеянная надпись возвращена в работу сегментом #%s" % res.get("segment"))
data = main._load_source_map(1)
blk = (next(im for im in data["images"] if im["part"] == o["part"])["blocks"])[o["block"]]
check("skip" not in blk and blk.get("seg") == res["segment"], "метка снята, сегмент привязан")
again = main.image_restore_block(1, main.ImageRestoreRequest(part=o["part"], block=o["block"]))
check(again["ok"] and not again["created"] and again["segment"] == res["segment"],
      "повторный возврат не заводит второй сегмент")
main._job_images(new_job(dry_run=False))
check(len([s for s in project["segments"] if s["id"] == res["segment"]]) == 1,
      "разбор после возврата не задваивает сегмент")
main.image_mark_overlay(1, res["segment"])          # возвращаем состав как был
try:
    main.image_restore_block(1, main.ImageRestoreRequest(part="word/media/нет.png", block=0))
    check(False, "несуществующий блок обязан отказать")
except main.HTTPException as e:
    check(e.status_code == 404, "несуществующий блок: " + str(e.detail))


print("\n── «это надпись аппарата» ──")
# Модель ошибается в обе стороны, и дальше машиной это не отсеять. Решает
# человек — а система обязана слушаться и ПОМНИТЬ: следующий разбор не имеет
# права завести тот же сегмент заново.
mark = [s for s in project["segments"] if (s.get("origin") or {}).get("kind") == "image"][0]
res = main.image_mark_overlay(1, mark["id"])
check(res["ok"] and not [s for s in project["segments"] if s["id"] == mark["id"]],
      "сегмент убран по решению человека")
data = main._load_source_map(1)
o = mark["origin"]
blk = [b for i, b in enumerate((next(im for im in data["images"] if im["part"] == o["part"])
                                .get("blocks") or [])) if i == o["block"]][0]
check(blk.get("skip") == "overlay" and "seg" not in blk, "метка легла на блок")
main._job_images(new_job(dry_run=False))
check(not [s for s in project["segments"] if s["id"] == mark["id"]
           or (s.get("origin") or {}).get("block") == o["block"]
           and (s.get("origin") or {}).get("part") == o["part"]],
      "следующий разбор не заводит помеченное заново")
try:
    main.image_mark_overlay(1, 1)
    check(False, "абзацный сегмент так пометить нельзя")
except main.HTTPException as e:
    check("не из картинки" in str(e.detail), "абзацный сегмент отклонён: " + str(e.detail))
# Возвращаем метку, чтобы дальше проверять на полном составе
blk.pop("skip", None)
main._save_source_map(1, data)
main._job_images(new_job(dry_run=False))

print("\n── одинаковый текст размечается одинаково ──")
# Модель решает про каждую картинку отдельно и на границе ошибается: на боевом
# учебнике строка настроек томографа оказалась надпечаткой на двух снимках
# и «текстом документа» на третьем — и стала сегментом, за перевод которого
# человек заплатит. Спрашивать второй раз незачем: ответ уже есть, просто
# разный. Берём большинство.
SETTINGS = "kV 120.0 mA: 283 2.5 mm Tilt: 0.0 degrees"
data = main._load_source_map(1)
blocks = [b for im in data["images"] for b in (im.get("blocks") or [])]
blocks[0]["text"], blocks[0]["skip"] = SETTINGS, "overlay"
blocks[1]["text"], blocks[1]["skip"] = SETTINGS, "overlay"
odd = blocks[2]
odd["text"] = SETTINGS
odd.pop("skip", None)
odd_id = max(s["id"] for s in project["segments"]) + 1
project["segments"].append(main._image_new_segment(SETTINGS, "word/media/x.png", 0, odd_id))
odd["seg"] = odd_id
moved, dropped = main._image_harmonize(data, project)
check(moved == 1 and dropped == 1 and odd.get("skip") == "overlay",
      "строка, отсеянная на двух картинках, отсеяна и на третьей")
check(not [s for s in project["segments"] if s["id"] == odd_id],
      "сегмент, заведённый по ошибке, снят")

# А переведённое большинством голосов не отменяется: это оплаченная работа.
odd.pop("skip", None)
odd["seg"] = odd_id
project["segments"].append(main._image_new_segment(SETTINGS, "word/media/x.png", 0, odd_id))
project["segments"][-1]["target"] = "Scanner settings"
moved2, dropped2 = main._image_harmonize(data, project)
check(moved2 == 0 and dropped2 == 0 and [s for s in project["segments"] if s["id"] == odd_id],
      "переведённый сегмент не снимается по счёту голосов")
project["segments"] = [s for s in project["segments"] if s["id"] != odd_id]

# Единственное вхождение никем не переголосовывается.
alone = {"box": [0, 0, 10, 10], "text": "Рис. 5. Каверна", "rows": 1}
data["images"][0]["blocks"].append(alone)
main._image_harmonize(data, project)
check("skip" not in alone, "одиночная надпись остаётся как решила модель")
data["images"][0]["blocks"].remove(alone)
main.images_forget(1, main.ImagesForgetRequest(force=True, wipe=True))
main._job_images(new_job(dry_run=False))

print("\n── номер сегмента сам по себе ничего не доказывает ──")
# Номера переиспользуются: max(id)+1 после сноса выдаёт те же числа заново.
# Блок, чей номер достался ЧУЖОМУ сегменту, обязан завести свой, а не цепляться
# к соседу — иначе он остаётся без перевода и молча выпадает из выгрузки.
data = main._load_source_map(1)
blk = [b for im in data["images"] for b in (im.get("blocks") or []) if b.get("seg")][0]
blk["seg"] = 1                       # сегмент #1 — абзацный, чужой
main._save_source_map(1, data)
before_n = len([s for s in project["segments"] if (s.get("origin") or {}).get("kind") == "image"])
main._job_images(new_job(dry_run=False))
after_n = len([s for s in project["segments"] if (s.get("origin") or {}).get("kind") == "image"])
check(after_n == before_n + 1,
      "блок с чужим номером завёл свой сегмент, а не принял соседа: было %d, стало %d"
      % (before_n, after_n))
main.images_forget(1, main.ImagesForgetRequest(force=True))
main._job_images(new_job(dry_run=False))

print("\n── «не спросили» ≠ «пусто» ──")
data = main._load_source_map(1)
for im in data["images"]:
    for b in (im.get("blocks") or []):
        b.pop("text", None)
        b.pop("skip", None)
main._save_source_map(1, data)
main._openai_read_image = lambda *a, **k: None
job = new_job(dry_run=False)
main._job_images(job)
data = main._load_source_map(1)
blocks = [b for im in data["images"] for b in (im.get("blocks") or [])]
check(all("text" not in b for b in blocks),
      "несостоявшийся вызов не записывает пустоту")
check(job["counters"].get("readFailed", 0) >= 1, "несостоявшийся вызов посчитан")
main._openai_read_image = fake_read
main._job_images(new_job(dry_run=False))

print("\n── снос распознанного ──")
seg = [s for s in project["segments"] if (s.get("origin") or {}).get("kind") == "image"][0]
seg["target"] = "Fig. 1. Diagram of the lungs"
res = main.images_forget(1, main.ImagesForgetRequest(force=False))
check(res["removed"] == 0 and res["keptTranslated"] == [seg["id"]],
      "переведённый сегмент из картинки без разрешения не сносится")
res = main.images_forget(1, main.ImagesForgetRequest(force=True))
check(res["removed"] == 1, "с разрешением сносится")
check(not [s for s in project["segments"] if (s.get("origin") or {}).get("kind") == "image"],
      "сегментов из картинок не осталось")

print("\n── кусок картинки для человека ──")
main._job_images(new_job(dry_run=False))
seg = [s for s in project["segments"] if (s.get("origin") or {}).get("kind") == "image"][0]
resp = main.image_crop(1, seg=seg["id"])
check(resp.media_type == "image/png" and len(resp.body) > 100,
      "кроп отдаётся картинкой: %d байт" % len(resp.body))
try:
    main.image_crop(1, seg=1)
    check(False, "кроп у абзацного сегмента обязан отказать")
except main.HTTPException as e:
    check("не из картинки" in str(e.detail), "кроп абзацного сегмента: " + str(e.detail))

print("\n── отчёт для экрана ──")
rep = main.images_report(1)
st = rep["stats"]
check(st["images"] == 2 and st["withText"] == 1 and st["segments"] == 1,
      "отчёт совпадает с картой: %s" % {k: st[k] for k in ("images", "withText", "segments")})
check(st["overlay"] >= 1 and st["noise"] >= 1, "отсеянное показано, а не спрятано")
# Смету считает СЕРВЕР и по каждой модели каталога: цифра в .jsx была бы
# вторым прайс-листом рядом с настоящим.
check(isinstance(rep["est"], dict) and set(rep["est"]) == {m["id"] for m in main.OPENAI_MODELS},
      "смета названа по каждой модели каталога")
check(all(v == 0.0 for v in rep["est"].values()) and rep["estTokens"]["in"] == 0,
      "всё прочитано — смета нулевая, а не выдуманная")

# А когда читать есть что, смета обязана различать модели: иначе выбор
# не влияет ни на что, и человек не может решить, чем платить.
data = main._load_source_map(1)
for im in data["images"]:
    for b in (im.get("blocks") or []):
        b.pop("text", None)
main._save_source_map(1, data)
rep2 = main.images_report(1)
costs = {k: v for k, v in rep2["est"].items() if v}
check(rep2["estTokens"]["in"] > 0 and len(set(costs.values())) > 1,
      "у разных моделей разная цена разбора: %s" % sorted(costs.values())[:3])
check(rep2["est"][main.DEFAULT_OPENAI_MODEL] is not None,
      "у модели по умолчанию цена известна")
main._job_images(new_job(dry_run=False))

print("\n── разбор переживает повторную привязку исходника ──")
# Человек жмёт «Заменить» в карточке исходника — например, чтобы поправить
# расхождение по абзацам. Разбор картинок при этом стирался целиком: карточка
# говорила «разбор не делался», кроп отвечал 404, а перевод надписей молча
# исчезал из выгрузки, хотя сегменты оставались в проекте.
was = main._load_source_map(1)
was_blocks = sum(len(im.get("blocks") or []) for im in was["images"])


class _Upload:
    """UploadFile ровно в том объёме, в каком его читает attach_source."""

    def __init__(self, b, name="t.docx"):
        self._b, self.filename = b, name

    async def read(self): return self._b


import asyncio                                                # noqa: E402

# force=False намеренно: порог «совпало меньше половины» обязан считаться
# по АБЗАЦНЫМ сегментам. Сегменты из картинок абзацами не являются, и в общем
# знаменателе они отклоняли бы родной файл как чужой.
res = asyncio.run(main.attach_source(1, _Upload(content), False))
now = main._load_source_map(1)
check(res.get("ok") and sum(len(im.get("blocks") or []) for im in now.get("images") or []) == was_blocks,
      "повторная привязка исходника не стирает разбор картинок")
check(now.get("imagesAt") == was.get("imagesAt"), "отметка о разборе на месте")

print("\n── правки во время разбора ──")
job_live = {"id": 9, "kind": "images", "project": 1, "status": "running",
            "total": 1, "done": 0, "counters": {}, "params": {}, "stop": False,
            "ids": [], "recent": []}
main._JOBS[9] = job_live
try:
    asyncio.run(main.attach_source(1, _Upload(content), True))
    check(False, "привязка во время разбора обязана отказать")
except main.HTTPException as e:
    check(e.status_code == 409, "исходник не приложить во время разбора: " + str(e.detail))
try:
    main.images_forget(1, main.ImagesForgetRequest())
    check(False, "снос во время разбора обязан отказать")
except main.HTTPException as e:
    check(e.status_code == 409, "распознанное не снести во время разбора")
main._JOBS.pop(9, None)

print("\n── оборванный ответ модели ──")
whole = main._image_read_parse('{"blocks":[{"i":0,"text":"A","overlay":false}]}')
fenced = main._image_read_parse('```json\n{"blocks":[{"i":0,"text":"A"},{"i":1,"text":"B"}]}\n```')
cut = main._image_read_parse('{"blocks":[{"i":0,"text":"A","overlay":false},{"i":1,"text":"Bcd')
check(len(whole) == 1 and len(fenced) == 2, "целый ответ разбирается, обёрнутый тоже")
check(len(cut) == 1 and cut[0]["text"] == "A",
      "у оборванного ответа берутся законченные блоки: за них уже заплачено")
check(main._image_read_parse("не могу прочитать") == [],
      "ответ без JSON — пусто, а не выдуманные блоки")

print("\n── снос не выбрасывает оплаченное ──")
seg_img = [s for s in project["segments"] if (s.get("origin") or {}).get("kind") == "image"][0]
res = main.images_forget(1, main.ImagesForgetRequest(force=True))
after = main._load_source_map(1)
texts = [b for im in after["images"] for b in (im.get("blocks") or []) if b.get("text")]
check(res["removed"] >= 1 and texts, "снос сегментов оставляет прочитанный текст")
check(after["images"] and any(im.get("blocks") for im in after["images"]),
      "и геометрию тоже — она стоила минут работы детектора")
st = res["stats"]
check(st["pending"] >= 1,
      "кнопка чтения знает, что работа осталась: pending=%s" % st["pending"])
main._job_images(new_job(dry_run=False))
check(len([s for s in project["segments"] if (s.get("origin") or {}).get("kind") == "image"]) == 1,
      "повторный заход заводит сегмент заново и бесплатно")
res = main.images_forget(1, main.ImagesForgetRequest(force=True, wipe=True))
after = main._load_source_map(1)
check(res["wiped"] and not [b for im in after["images"]
                            for b in (im.get("blocks") or []) if b.get("text")],
      "с явным разрешением прочитанное выбрасывается")
check(any(im.get("blocks") for im in after["images"]),
      "геометрия остаётся даже при wipe: она не зависит от того, верно ли прочитано")
main._openai_read_image = fake_read
main._job_images(new_job(dry_run=False))

print("\n── экспорт 1в1: перевод возвращается в картинку ──")
import hashlib                                                # noqa: E402

main.EXPORT_DIR = TMP / "exports"
main.EXPORT_DIR.mkdir(parents=True, exist_ok=True)
docx_path, map_path = main._source_paths(1)
before = hashlib.sha1(docx_path.read_bytes()).hexdigest()
src_raster, _o = main._docx_media(docx_path.read_bytes())

# Надпись на белом фоне — её перевод впишется в саму картинку.
seg = [s for s in project["segments"] if (s.get("origin") or {}).get("kind") == "image"][0]
seg["target"] = "Fig. 1. Diagram of the lungs"

# А эту рамку кладём на фотошум руками: перерисовать там нельзя, и перевод
# обязан уйти подписью под картинкой, а не заплаткой поверх снимка.
data = main._load_source_map(1)
noisy = [im for im in data["images"] if not im.get("blocks")][0]
nid = max(s["id"] for s in project["segments"]) + 1
project["segments"].append(main._image_new_segment("Подпись на снимке", noisy["part"], 0, nid))
project["segments"][-1]["target"] = "Caption on the photograph"
noisy["blocks"] = [{"box": [20, 20, 480, 60], "rows": 1, "conf": 0.9, "flat": 0.1,
                    "align": "left", "text": "Подпись на снимке", "seg": nid}]
main._save_source_map(1, data)

path, stats = main._generate_export(project, "docx_layout")
check(stats["img_repainted"] == 1, "надпись на белом фоне перерисована: %s" % stats["img_repainted"])
check(stats["img_captioned"] == 1 and stats["img_flat"] == 1,
      "надпись на снимке ушла подписью, причина названа: %s"
      % {k: stats[k] for k in ("img_captioned", "img_flat")})

out_raster, _o2 = main._docx_media(path.read_bytes())
changed = [n for n in src_raster
           if hashlib.sha1(src_raster[n]).hexdigest()
           != hashlib.sha1(out_raster.get(n, b"")).hexdigest()]
check(len(changed) == 1, "заменена ровно одна картинка, остальные байт в байт: %s" % changed)
check(hashlib.sha1(docx_path.read_bytes()).hexdigest() == before,
      "исходник на диске не тронут — перерисовка живёт только в копии")

texts = [p.text for p in Document(str(path)).paragraphs]
check("Caption on the photograph" in texts, "подпись дописана в документ")
check(texts.index("Caption on the photograph") > texts.index("Текст перед рисунком."),
      "подпись стоит после своей картинки, а не в начале документа")

# Подменённая картинка не перерисовывается. Карта переживает повторную
# привязку исходника намеренно, но рамки описывают ТУ картинку: перевод,
# вписанный по чужим координатам, — заплатка посреди чужого снимка.
data = main._load_source_map(1)
mine = [im for im in data["images"] if any(b.get("seg") == seg["id"]
                                           for b in (im.get("blocks") or []))][0]
keep_sha = mine["sha"]
mine["sha"] = "0" * 40
main._save_source_map(1, data)
_p3, stats3 = main._generate_export(project, "docx_layout")
check(stats3["img_stale"] >= 1 and stats3["img_repainted"] == 0,
      "картинка разошлась с картой — не трогаем и говорим числом: %s"
      % {k: stats3[k] for k in ("img_stale", "img_repainted")})
mine["sha"] = keep_sha
main._save_source_map(1, data)

# Прочитанная надпись без сегмента (разбор остановили, сегменты снесли) —
# это не «сделано»: она останется на языке оригинала.
data = main._load_source_map(1)
for im in data["images"]:
    for b in (im.get("blocks") or []):
        if b.get("seg") == seg["id"]:
            b.pop("seg")
main._save_source_map(1, data)
_p4, stats4 = main._generate_export(project, "docx_layout")
check(stats4["img_noseg"] >= 1,
      "прочитанная надпись без сегмента посчитана, а не забыта: %s" % stats4["img_noseg"])
main._job_images(new_job(dry_run=False))

# Непереведённая надпись НЕ стирается: пустой перевод выбросил бы текст совсем.
seg = [s for s in project["segments"] if (s.get("origin") or {}).get("kind") == "image"][0]
seg["target"] = ""
_p2, stats2 = main._generate_export(project, "docx_layout")
check(stats2["img_untranslated"] == 1 and stats2["img_repainted"] == 0,
      "непереведённая надпись остаётся на языке оригинала и посчитана: %s"
      % {k: stats2[k] for k in ("img_untranslated", "img_repainted")})

shutil.rmtree(TMP, ignore_errors=True)
print("\n" + ("ВСЁ ПРОШЛО" if not fail else "ПРОВАЛЕНО: " + "; ".join(fail)))
sys.exit(1 if fail else 0)
