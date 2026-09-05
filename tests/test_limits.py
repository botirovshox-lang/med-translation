"""Лимит расхода по организации: 402 на платное, бесплатное работает.

Факт расхода складывается по организации и месяцу (`_spend_add` из
`_note_usage`), лимит ставит суперпользователь. На исчерпанном лимите
платные команды отвечают 402 с остатком, а правка начертания, откаты,
пересчёт back-check, принятие кандидатов, разбор состава и экспорт —
работают: лимит режет деньги, а не доступ к оплаченному. Ни одного
вызова модели, файл состояния не пишется.
"""
import os, sys
os.environ["APP_PASSWORD"] = "boot-password-1"
os.environ["AUTHORITY_CORPUS"] = "0"
os.environ["OPENAI_API_KEY"] = "test-key"
sys.path.insert(0, "backend")
import main
from starlette.testclient import TestClient

main.save_state = lambda *a, **k: None
main.STATE["users"], main.STATE["tenants"], main.STATE["spend"] = [], [], {}
main._SESSIONS.clear(); main._LOGIN_FAILS.clear()
fail = []


def check(cond, label):
    print(("  OK   " if cond else "  FAIL ") + label)
    if not cond:
        fail.append(label)


c = TestClient(main.app)
H = lambda t: {"Authorization": "Bearer " + t}
A = c.post("/api/auth/login", json={"login": "admin", "password": "boot-password-1"}).json()["token"]
c.post("/api/admin/tenants", headers=H(A),
       json={"id": "acme", "name": "ACME", "ownerLogin": "acme", "ownerPassword": "acme-pass-123"})
B = c.post("/api/auth/login", json={"login": "acme", "password": "acme-pass-123"}).json()["token"]
pid = c.post("/api/projects", headers=H(B), json={"title": "p", "src": "RU", "tgt": "EN"}).json()["id"]
proj = next(p for p in main.STATE["projects"] if p["id"] == pid)
proj["segments"].append({"id": 1, "source": "Тест.", "target": "", "status": "new",
                         "route": "GPT_REQUIRED", "risk": "low", "comments": [], "qa": []})


class _Resp:
    usage = {"prompt_tokens": 1000, "completion_tokens": 500}


print("=== 1. Расход складывается по организации и месяцу ===")
tok = main.CURRENT_SESSION.set({"tenant": "acme", "user": 2, "role": "owner"})
try:
    main._note_usage("translate", "gpt-4o", _Resp())
    main._note_usage("translate", "no-such-model", _Resp())
finally:
    main.CURRENT_SESSION.reset(tok)
st = main._spend_status("acme")
check(st["calls"] == 2 and st["unpriced"] == 1 and st["spentUsd"] > 0, "две записи, одна без цены, сумма > 0: %s" % st)
check(main._spend_status("default")["calls"] == 0, "у другой организации — ноль")
me = c.get("/api/auth/me", headers=H(B)).json()
check(me["spend"]["tenant"] == "acme" and me["spend"]["limitUsd"] is None and not me["spend"]["over"],
      "/auth/me показывает расход, лимита нет")

print("=== 2. Лимит ставит только super ===")
r = c.post("/api/admin/tenants/acme", headers=H(B), json={"limitUsd": 100})
check(r.status_code == 403, "владелец сам себе лимит не ставит")
r = c.post("/api/admin/tenants/acme", headers=H(A), json={"limitUsd": 0.001})
check(r.status_code == 200 and r.json()["spend"]["over"], "super поставил лимит ниже расхода — over")
r = c.get("/api/admin/tenants", headers=H(A))
check(r.status_code == 200 and any(t["id"] == "acme" and t["spend"]["over"] for t in r.json()["tenants"]),
      "список организаций с расходом")

print("=== 3. На исчерпанном лимите: платное — 402, бесплатное — работает ===")
paid = [("POST", "/api/projects/%d/jobs" % pid, {"kind": "full", "ids": [1]}),
        ("POST", "/api/segments/%d/1/translate" % pid, {}),
        ("POST", "/api/segments/%d/1/backcheck" % pid, {}),
        ("POST", "/api/projects/%d/batch" % pid, {}),
        ("POST", "/api/glossary/audit", {})]
for m, path, body in paid:
    r = c.request(m, path, headers=H(B), json=body)
    check(r.status_code == 402 and "spend" in r.json(), "%s → 402" % path)
free = [("POST", "/api/projects/%d/run-plan" % pid, {"steps": ["translate"]}),
        ("POST", "/api/projects/%d/term-case" % pid, {}),
        ("POST", "/api/projects/%d/backcheck/rescore" % pid, {}),
        ("POST", "/api/projects/%d/repair/accept-batch" % pid, {}),
        ("POST", "/api/glossary/revert-repairs", {"src": "x", "tgt": "y"}),
        ("GET", "/api/projects/%d/analysis" % pid, None),
        ("GET", "/api/projects/%d/coverage" % pid, None),
        ("POST", "/api/projects/%d/export" % pid, {"format": "xlsx"})]
for m, path, body in free:
    r = c.request(m, path, headers=H(B), **({"json": body} if body is not None else {}))
    check(r.status_code != 402, "%s → не 402 (%d)" % (path, r.status_code))

print("=== 4. Снятие лимита ===")
r = c.post("/api/admin/tenants/acme", headers=H(A), json={"clearLimit": True})
check(r.status_code == 200 and not r.json()["spend"]["over"], "лимит снят")
r = c.post("/api/segments/%d/1/backcheck" % pid, headers=H(B), json={})
check(r.status_code != 402, "платное снова доступно")
seed = c.get("/api/seed", headers=H(B)).json()
check("spend" not in seed, "/api/seed расход по организациям не отдаёт")

print("=== 5. Смета больше остатка — 402 на старте (число клиентское, рубеж от случайности) ===")
r = c.post("/api/admin/tenants/acme", headers=H(A), json={"limitUsd": 1.0})
check(r.status_code == 200 and not r.json()["spend"]["over"], "лимит $1, расход меньше")
body = {"kind": "backcheck", "segment_ids": [1], "params": {"est_cost": 5.0}}
r = c.post("/api/projects/%d/jobs" % pid, headers=H(B), json=body)
check(r.status_code == 402 and "Смета прогона" in r.json().get("detail", ""),
      "смета $5 больше остатка → 402: %s" % r.text[:120])
ew = main.EXTERNAL_WORKER
main.EXTERNAL_WORKER = True                      # задачу никто не подхватит — без сети
body["params"]["est_cost"] = 0.01
r = c.post("/api/projects/%d/jobs" % pid, headers=H(B), json=body)
main.EXTERNAL_WORKER = ew
check(r.status_code == 200, "смета в остаток → задача принята: %s" % r.text[:120])
if r.status_code == 200:
    main._JOBS.pop(r.json()["job"]["id"], None)

print("=== 6. Лимит исчерпан — задача из очереди останавливается ДО первой порции ===")
r = c.post("/api/admin/tenants/acme", headers=H(A), json={"limitUsd": 0.001})
calls = []
orig_chunk = main._job_chunk
main._job_chunk = lambda *a, **k: calls.append(a) or {"done": len(a[2])}
mk = lambda ids: {"id": 999, "kind": "backcheck", "project": pid, "status": "queued", "tenant": "acme",
                  "total": len(ids), "done": 0, "counters": {}, "error": None, "params": {}, "ids": ids,
                  "stop": False, "recent": [], "created": "", "started": None, "finished": None}
job = mk([1])
try:
    main._job_execute(job)
finally:
    main._job_chunk = orig_chunk
check(job["status"] == "stopped" and job.get("stopReason") == "limit" and job.get("finished"),
      "stopped с кодом limit до старта: %s/%s" % (job["status"], job.get("stopReason")))
check(job["error"] == main.JOB_STOP_LIMIT and job["counters"].get("limitStop") == 1, "причина и счётчик записаны")
check(not calls, "ни одна порция не вызвана")

bumped = []
orig_bump = getattr(main.STORE, "bump_epoch", None)
main.STORE.bump_epoch = lambda name: bumped.append(name) or 0
try:
    c.post("/api/admin/tenants/acme", headers=H(A), json={"limitUsd": 0.5})
finally:
    if orig_bump is not None:
        main.STORE.bump_epoch = orig_bump
check("doc:tenants" in bumped, "правка лимита поднимает эпоху doc:tenants — внешний воркер увидит новый потолок")

print("=== 6a. Лимит кончился ПОСЛЕ первой порции — вторая не идёт, сделанное сохранено ===")
r = c.post("/api/admin/tenants/acme", headers=H(A), json={"limitUsd": 1.0})
orig_ss, seen = main._spend_status, {"n": 0}
def _ss(tenant=None):
    st = orig_ss(tenant); st["over"] = seen["n"] >= 1; return st
def _chunk(*a, **k):
    seen["n"] += 1; calls.append(a); return {"done": len(a[2])}
main._spend_status, main._job_chunk = _ss, _chunk
job = mk(list(range(1, main.JOB_CHUNKS["backcheck"] + 2)))     # две порции
try:
    main._job_execute(job)
finally:
    main._spend_status, main._job_chunk = orig_ss, orig_chunk
check(len(calls) == 1 and job["status"] == "stopped" and job.get("stopReason") == "limit",
      "одна порция прошла, вторая остановлена лимитом: calls=%d %s" % (len(calls), job["status"]))
check(job["done"] == main.JOB_CHUNKS["backcheck"], "сделанное сохранено в done: %s" % job["done"])
r = c.post("/api/admin/tenants/acme", headers=H(A), json={"clearLimit": True})

print("=== 7. Потолки импорта: страницы на файл (413), проекты и страницы организации (402) ===")
try:
    from docx import Document
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False
if HAVE_DOCX:
    import io as _io
    d = Document()
    d.add_paragraph("Первый абзац про туберкулёз лёгких и его лечение в стационаре.")
    d.add_paragraph("Второй абзац про профилактику.")
    b = _io.BytesIO(); d.save(b); raw = b.getvalue()
    MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    up = lambda: c.post("/api/projects/upload", headers=H(B), files={"file": ("t.docx", raw, MIME)},
                        data={"src": "RU", "tgt": "EN"})
    caps = (main.IMPORT_MAX_PAGES, main.TENANT_MAX_PAGES, main.TENANT_MAX_PROJECTS)
    try:
        main.IMPORT_MAX_PAGES, main.TENANT_MAX_PAGES, main.TENANT_MAX_PROJECTS = 0.01, 0, 0
        r = up()
        check(r.status_code == 413 and "Файл на " in r.json().get("detail", ""),
              "потолок страниц на файл → 413: %s" % r.text[:110])
        main.IMPORT_MAX_PAGES, main.TENANT_MAX_PROJECTS = 0, 1          # у acme уже есть проект
        r = up()
        check(r.status_code == 402 and "проектов" in r.json().get("detail", ""), "потолок проектов → 402: %s" % r.text[:110])
        main.TENANT_MAX_PROJECTS, main.TENANT_MAX_PAGES = 0, 0.01
        r = up()
        check(r.status_code == 402 and " стр." in r.json().get("detail", ""),
              "потолок страниц организации → 402, старый проект без pages посчитан по сегментам: %s" % r.text[:110])
        main.TENANT_MAX_PAGES = 0
        r = up()
        ok = r.status_code == 200 and (r.json().get("pages") or 0) > 0
        check(ok, "без потолков импорт проходит, pages записан: %s" % (r.json().get("pages") if r.status_code == 200 else r.text[:110]))
        if r.status_code == 200:
            main.STATE["projects"] = [p for p in main.STATE["projects"] if p["id"] != r.json()["id"]]
    finally:
        main.IMPORT_MAX_PAGES, main.TENANT_MAX_PAGES, main.TENANT_MAX_PROJECTS = caps
else:
    print("python-docx нет — раздел 7 пропущен")

main.STATE["projects"] = [p for p in main.STATE["projects"] if p["id"] != pid]
print("\n" + ("ВСЁ ПРОШЛО" if not fail else "ПРОВАЛЕНО: " + "; ".join(fail)))
