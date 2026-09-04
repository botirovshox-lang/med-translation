"""Экспорт «как в оригинале»: перевод подставляется в исходный файл.

Собрать документ, похожий на исходник, из одних сегментов нельзя — в них нет
ни шрифта, ни картинок. Поэтому исходник хранится, а выгрузка подменяет текст
в НЁМ. Отсюда всё, что здесь проверяется:

  1. якорь сегмента — номер абзаца, и список абзацев обязан быть ПОЛНЫМ:
     выброси из него пустые строки, и переводы поедут по всему документу;
  2. соседние одинаковые абзацы — один сегмент, но ДВА якоря: иначе второй
     останется на языке оригинала;
  3. в ВЫЧИСЛЯЕМЫЕ поля (номер страницы в оглавлении) и в скрытый текст писать
     нельзя — номер считает Word, а скрытого не видит никто; зато результат
     остальных полей (строка оглавления внутри TOC) — обычный текст, и он
     переводится, иначе одна строка остаётся на языке оригинала;
  4. непереведённый сегмент оставляет абзац как есть — подставить туда пусто
     значит стереть оригинал и выдать это за перевод;
  5. привязка исходника к готовому проекту идёт ПО ТЕКСТУ и не имеет права
     сажать перевод на чужой абзац: чужой файл отклоняется по числу совпадений;
  6. подменённый под картой файл отклоняется — номера абзацев уехали;
  7. выделение ВНУТРИ абзаца переносится, а не схлопывается: сумма кусков
     всегда равна переводу целиком, резать посреди слова нельзя, а кусок,
     который перевод сохраняет дословно (латынь, число), ставится точно;
  8. служебное оформление (язык проверки орфографии, микрокернинг) деления
     не вызывает — иначе перевод режется там, где резать нечего;
  9. колонтитулы входят в разбор, но тело идёт ПЕРВЫМ: иначе номера абзацев
     тела уехали бы и все прежние карты стали бы врать.

Ни одного вызова модели и ни одного обращения к сети: собирается настоящий
.docx (python-docx) в отдельном временном каталоге.
"""
import os, sys, shutil, tempfile
from pathlib import Path

os.environ.setdefault("APP_PASSWORD", "test")
os.environ["AUTHORITY_CORPUS"] = "0"
sys.path.insert(0, "backend")
import main

main.save_state = lambda *a, **k: None

fail = []


def check(cond, label):
    print(("  OK   " if cond else "  FAIL ") + label)
    if not cond:
        fail.append(label)


TMP = Path(tempfile.mkdtemp(prefix="medcat-layout-"))
main.SOURCE_DIR = TMP / "sources"
main.EXPORT_DIR = TMP / "exports"
main.EXPORT_DIR.mkdir(parents=True)

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _run(p, text, bold=False, italic=False):
    r = p.add_run(text)
    if bold:
        r.bold = True          # False писало бы <w:b w:val="0"/> — тоже отметку
    if italic:
        r.italic = True
    return r


def build_source() -> bytes:
    """Документ со всеми случаями сразу: обычный абзац, абзац из двух прогонов
    с разным оформлением, соседний повтор, чисто цифровая строка и строка
    оглавления с полем PAGEREF."""
    doc = Document()
    doc.add_paragraph("Первый абзац документа.")           # 0
    p = doc.add_paragraph()                                 # 1: два прогона
    _run(p, "Туберкулёз", bold=True)
    _run(p, " — инфекционное заболевание.")
    doc.add_paragraph("Повтор строки.")                     # 2
    doc.add_paragraph("Повтор строки.")                     # 3: тот же сегмент
    doc.add_paragraph("14")                                 # 4: только цифры
    doc.add_paragraph("Абзац без перевода.")                # 5

    p = doc.add_paragraph()                                 # 6: латынь курсивом
    _run(p, "Возбудитель — ")
    _run(p, "Mycobacterium tuberculosis", italic=True)
    _run(p, ", открыт в 1882 году.")

    p = doc.add_paragraph()                                 # 7: служебное
    r1 = _run(p, "Первая половина строки ")
    r2 = _run(p, "и вторая половина строки.")
    # w:lang — каким языком проверять орфографию. Word ставит его сам, кусками;
    # видимого оформления это не меняет и делить абзац не повод.
    for r, lang in ((r1, "ru-RU"), (r2, "en-US")):
        rpr = r._r.get_or_add_rPr()
        el = OxmlElement("w:lang")
        el.set(qn("w:val"), lang)
        rpr.append(el)

    toc = doc.add_paragraph()                               # 6: оглавление
    _run(toc, "ГЛАВА ПЕРВАЯ")
    tab = OxmlElement("w:r")
    tab.append(OxmlElement("w:tab"))
    toc._p.append(tab)
    for kind, payload in (("begin", None), (None, " PAGEREF _Toc1 \\h "),
                          ("separate", None), (None, "85"), ("end", None)):
        r = OxmlElement("w:r")
        if kind:
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), kind)
            r.append(fld)
        elif payload.startswith(" PAGEREF"):
            instr = OxmlElement("w:instrText")
            instr.text = payload
            r.append(instr)
        else:
            t = OxmlElement("w:t")
            t.text = payload
            r.append(t)
        toc._p.append(r)

    # Подпись к рисунку, сдвинутая вправо ПРОБЕЛАМИ, а не отступом абзаца.
    # Так свёрстан весь учебник; импорт эти пробелы обрезает, и без их
    # возврата подпись уезжает к левому краю, а Word заново обтекает ею
    # плавающую картинку — строка рассыпается по обе стороны рисунка.
    doc.add_paragraph(" " * 33 + "Рис. 60. Казеозная пневмония.")   # 9

    # Колонтитул с настоящим текстом: он живёт отдельной частью пакета,
    # в body не попадает и без отдельного обхода остался бы на русском.
    doc.sections[0].header.paragraphs[0].text = "Фтизиатрия"

    path = TMP / "source.docx"
    doc.save(str(path))
    return path.read_bytes()


CONTENT = build_source()
TEXTS = main._docx_paragraph_texts(CONTENT)
paras = [t for t, _f in TEXTS]
units = main._docx_units(paras, [f for _t, f in TEXTS])   # тем же путём, что upload_project

BODY = 10     # абзацев в теле; дальше идут колонтитулы
check(len(paras) > BODY, "разбор дошёл до колонтитулов (%d абзацев)" % len(paras))
check([t for t, _ in units] == [
    "Первый абзац документа.",
    "Туберкулёз — инфекционное заболевание.",
    "Повтор строки.",
    "Абзац без перевода.",
    "Возбудитель — Mycobacterium tuberculosis, открыт в 1882 году.",
    "Первая половина строки и вторая половина строки.",
    "ГЛАВА ПЕРВАЯ",
    "Рис. 60. Казеозная пневмония.",
    "Фтизиатрия",
], "в сегменты идут не все абзацы, зато текст колонтитула идёт: %s"
   % [t for t, _ in units])
check(units[2][1] == [2, 3], "соседний повтор — один сегмент и ДВА якоря")
check(units[-1][1][0] >= BODY,
      "колонтитул стоит ПОСЛЕ тела — номера абзацев тела не сдвинулись")

# ── проект как после импорта ────────────────────────────────────────
project = {"id": 1, "title": "Тест", "src": "RU", "tgt": "EN", "domain": "medical",
           "segments": [{"id": i + 1, "source": t, "target": "", "status": "new"}
                        for i, (t, _idx) in enumerate(units)]}
pairs = [[i, u + 1] for u, (_t, idxs) in enumerate(units) for i in idxs]
main._store_source_docx(project, CONTENT, "source.docx", pairs, len(paras))
check(project["sourceDocx"]["segments"] == len(units)
      and project["sourceDocx"]["paras"] == len(paras),
      "отметка в проекте называет и абзацы, и сегменты")

docx_path, map_path = main._source_paths(1)
check(docx_path.exists() and map_path.exists(), "исходник и карта легли рядом")
check("pairs" not in project["sourceDocx"],
      "карта НЕ попала в проект: state.json переписывается на каждую правку")

TR = {1: "First paragraph of the document.",
      2: "Tuberculosis is an infectious disease.",
      3: "Repeated line.",
      5: "The causative agent is Mycobacterium tuberculosis, discovered in 1882.",
      6: "First half of the line and second half of the line.",
      7: "CHAPTER ONE",           # номера в сегменте больше нет: поле — не текст
      8: "Fig. 60. Caseous pneumonia.",
      9: "Phthisiology"}
# Статусы нарочно вперемешку. В файл идёт всё, у чего есть перевод: экспорт
# не судит о качестве, он выгружает то, что есть, а «подтвердил человек» —
# про доверие к переводу, а не про попадание в документ. Фильтр по статусу
# здесь означал бы, что готовый документ молча теряет часть работы.
STATUSES = ["new", "translated", "review", "qa", "confirmed", "review",
            "translated", "new", "confirmed"]
for s, st in zip(project["segments"], STATUSES):
    s["target"] = TR.get(s["id"], "")
    s["status"] = st

out, stats = main._generate_export(project, "docx_layout")
check(out.name.endswith(" 1в1.docx"),
      "имя файла отличает 1в1 от обычного docx: " + out.name)

res = Document(str(out))
all_p = main._docx_flat_paragraphs(res)
text = [main._docx_clean("".join(t.text for t in p.iter(qn("w:t")) if t.text))
        for p in all_p]

check(len(all_p) == len(paras), "число абзацев не изменилось")
check(text[0] == "First paragraph of the document.", "обычный абзац переведён")
check(text[1] == "Tuberculosis is an infectious disease.",
      "абзац из двух прогонов переведён целиком: " + text[1])
check(text[2] == "Repeated line." and text[3] == "Repeated line.",
      "перевод встал в ОБА абзаца повтора")
check(text[4] == "14", "цифровая строка не тронута")
check(text[5] == "Абзац без перевода.",
      "непереведённый сегмент оставил оригинал, а не пустоту")
check(text[8] == "CHAPTER ONE85",
      "оглавление: переведён заголовок, номер страницы остался полем: " + text[8])
check(stats["trimmed"] == 0 and stats["mismatch"] == 0,
      "новый импорт: номера в сегменте нет, абзац узнан по тексту слотов: "
      "trimmed=%s mismatch=%s" % (stats["trimmed"], stats["mismatch"]))

# Проект СТАРОГО импорта: номер из поля попал и в сегмент, и в перевод —
# абзац узнаётся по ПОЛНОМУ тексту, хвост-номер снимается с перевода.
_s7 = project["segments"][6]
_s7["source"], _s7["target"] = "ГЛАВА ПЕРВАЯ85", "CHAPTER ONE85"
_out2, _st2 = main._generate_export(project, "docx_layout")
_res2 = Document(str(_out2))
_t2 = [main._docx_clean("".join(t.text for t in p.iter(qn("w:t")) if t.text))
       for p in main._docx_flat_paragraphs(_res2)]
check(_t2[8] == "CHAPTER ONE85" and _st2["trimmed"] == 1 and _st2["mismatch"] == 0,
      "старый импорт: хвост-номер снят с перевода, абзац узнан по полному тексту: "
      "%r trimmed=%s mismatch=%s" % (_t2[8], _st2["trimmed"], _st2["mismatch"]))
_s7["source"], _s7["target"] = "ГЛАВА ПЕРВАЯ", "CHAPTER ONE"
check(stats["written"] == 9 and stats["untranslated"] == 1,
      "отчёт называет и написанное, и непереведённое: %s" % stats)
# Ровно столько, сколько якорей у сегментов с переводом, — и ни один статус
# не отнял себе ни одного абзаца.
with_target = {s["id"] for s in project["segments"] if (s.get("target") or "").strip()}
check(stats["written"] == sum(1 for _i, sid in pairs if sid in with_target),
      "статус не влияет на попадание в файл: пишутся все абзацы с переводом")
check(text[0] == TR[1] and text[2] == TR[3],
      "непроверенный сегмент («новый») выгружен наравне с подтверждённым")

# ── выделения внутри абзаца ─────────────────────────────────────────
def marked(i, tag):
    """Текст абзаца, попавший в прогоны с этой отметкой оформления."""
    out = []
    for r in all_p[i].iter(qn("w:r")):
        t = r.find(qn("w:t"))
        rpr = r.find(qn("w:rPr"))
        if t is None or not (t.text or ""):
            continue
        if rpr is not None and rpr.find(qn(tag)) is not None:
            out.append(t.text)
    return "".join(out)


def whole(i):
    return "".join(t.text or "" for t in all_p[i].iter(qn("w:t")))


# 1: жирное «Туберкулёз» + обычный хвост. Жирным обязана остаться ЧАСТЬ,
# а не абзац целиком и не пустота.
bold1 = marked(1, "w:b")
check(whole(1) == TR[2], "текст абзаца не потерян и не задвоен: " + whole(1))
check(bold1.strip() and bold1 != TR[2],
      "жирным осталась часть перевода, а не весь абзац: %r" % bold1)
check(bold1.strip() == "Tuberculosis",
      "граница жирного подтянута к слову: %r" % bold1)

# 5: латинское название вида курсивом — перевод сохраняет его дословно,
# значит курсив ставится ТОЧНО, а не на глаз.
check(whole(6) == TR[5], "текст абзаца с латынью цел: " + whole(6))
check(marked(6, "w:i").strip() == "Mycobacterium tuberculosis",
      "курсив встал ровно на латинское название: %r" % marked(6, "w:i"))

# 6: прогоны различаются только языком проверки орфографии — делить нечего.
check(whole(7) == TR[6], "текст абзаца со служебным оформлением цел")
holders6 = [r for r in all_p[7].iter(qn("w:r"))
            if r.find(qn("w:t")) is not None and (r.find(qn("w:t")).text or "")]
check(len(holders6) == 1,
      "служебное оформление не делит перевод: кусков %d" % len(holders6))

# колонтитул переведён
hdr = [p for p in all_p[BODY:]
       if "".join(t.text or "" for t in p.iter(qn("w:t"))).strip() == "Phthisiology"]
check(len(hdr) == 1, "текст колонтитула переведён")

check(stats["inline"] == 2,
      "выделения перенесены ровно там, где они есть: inline=%s" % stats["inline"])

# ── куда встают границы выделений ───────────────────────────────────
# Настоящие пары из учебника: слева куски исходного абзаца по оформлению,
# справа перевод. Все обязаны делиться ТОЧНО — по знаку препинания либо
# по куску, который перевод сохраняет дословно. Ни одной догадки: там, где
# опереться не на что, деление честно помечается приблизительным, и такие
# случаи в этом списке были бы видны.
SPLITS = [
    (["13 ГЛАВА.", " ПРОФИЛАКТИКА ТУБЕРКУЛЁЗА"],
     "CHAPTER 13. TUBERCULOSIS PREVENTION",
     ["CHAPTER 13.", " TUBERCULOSIS PREVENTION"]),
    (["Клиника: ", "зависит от объема поражения лёгких"],
     "Clinical picture: depends on the extent of lung involvement",
     ["Clinical picture:", " depends on the extent of lung involvement"]),
    (["ПЦР", " – полимеразная цепная реакция"],
     "PCR – polymerase chain reaction",
     ["PCR", " – polymerase chain reaction"]),
    (["Термин ", "tuberculosis", " происходит от латинского слова ", "tuberculum"],
     "The term tuberculosis comes from the Latin word tuberculum",
     ["The term ", "tuberculosis", " comes from the Latin word ", "tuberculum"]),
    (["Алиментарный", " — заражение ", "M. bovis", " при употреблении сырого молока"],
     "Alimentary — infection with M. bovis when drinking raw milk",
     ["Alimentary ", "— infection with ", "M. bovis", " when drinking raw milk"]),
    (["(", "Учебник для студентов медицинских институтов", ")"],
     "(Textbook for students of medical institutes)",
     ["(", "Textbook for students of medical institutes", ")"]),
    (["HBsAg", " - антиген вируса гепатита В"],
     "HBsAg - hepatitis B virus antigen",
     ["HBsAg", " - hepatitis B virus antigen"]),
]
for src, tgt, want in SPLITS:
    got, approx = main._split_target(tgt, src)
    check(got == want and not approx,
          "деление %r → %r%s" % (src[0][:22], got, " (ДОГАДКА)" if approx else ""))
    check("".join(got) == tgt, "сумма кусков равна переводу целиком")

# Резать посреди слова нельзя даже когда опереться не на что.
blind, approx = main._split_target("aaa bbb ccc ddd eee", ["раз ", "два три четыре"])
check(approx and "".join(blind) == "aaa bbb ccc ddd eee",
      "без опоры деление честно помечено догадкой")
check(all(p == "" or p[0] != " " or True for p in blind)
      and all(not p.strip() or p.strip() in "aaa bbb ccc ddd eee".split()
              or " " in p.strip() for p in blind),
      "догадка всё равно режет по словам: %r" % blind)

# ── пробелы по краям абзаца возвращаются на место ───────────────────
# Подпись сдвинута вправо тридцатью тремя пробелами В ТЕКСТЕ, а не отступом
# абзаца. Импорт их обрезал, перевод их не содержит — и без возврата подпись
# встаёт к левому краю. На странице с плавающей картинкой это видно сразу:
# Word заново обтекает ею подпись, и строка рассыпается по обе стороны рисунка.
cap = "".join(t.text or "" for t in all_p[9].iter(qn("w:t")))
check(cap.startswith(" " * 33),
      "пробелы, которыми подпись сдвинута вправо, сохранены: %r" % cap[:44])
check(cap.strip() == TR[8],
      "и сам перевод при этом не тронут: %r" % cap.strip())
# Внутренние пробелы — часть предложения, у перевода свои: их не добавляем.
check("  " not in cap.strip(),
      "внутрь перевода лишние пробелы не попали: %r" % cap.strip())

# ── запись атомарна ─────────────────────────────────────────────────
# Экспорт учебника это 21 МБ и несколько секунд. Записанный прямо в итоговый
# файл, он в это время доступен на скачивание НЕДОПИСАННЫМ, а перезапись
# требует прав на сам файл: один файл, случайно оставленный в exports/ от
# другого владельца, ронял экспорт этого проекта навсегда (PermissionError
# на боевом сервере, «Сервер недоступен» на экране). os.replace требует прав
# только на каталог.
stale = out.with_name(out.name + ".tmp")
stale.write_bytes("мусор от упавшей сборки".encode("utf-8"))
out2, _st2 = main._generate_export(project, "docx_layout")
check(out2 == out and not stale.exists(),
      "временный файл убран за собой, а не оставлен рядом с готовым")
check(Document(str(out2)) is not None, "итоговый файл читается как .docx")

if os.name == "posix" and hasattr(os, "geteuid") and os.geteuid() != 0:
    # Только не под root: ему права на файл не помеха, и проверка была бы
    # зелёной ни о чём.
    out.chmod(0o444)
    try:
        main._generate_export(project, "docx_layout")
        check(True, "экспорт переживает недоступный на запись прежний файл")
    except OSError as e:
        check(False, "экспорт упал на прежнем файле: %s" % e)
    finally:
        out.chmod(0o644)

# ── соседи склеиваются только при равном ПОЛНОМ тексте ─────────────
_u = main._docx_units(["Введение", "Введение"], ["Введение.3", "Введение.5"])
check(len(_u) == 2, "две строки оглавления с разными номерами — два сегмента: %s" % _u)
_u = main._docx_units(["Введение", "Введение"], ["Введение.3", "Введение.3"])
check(len(_u) == 1 and _u[0][1] == [0, 1], "одинаковые целиком — один сегмент, два якоря")
check(len(main._docx_units(["Введение", "Введение"])) == 1, "без полного текста — как прежде")

# ── привязка исходника к готовому проекту ───────────────────────────
old = {"id": 2, "title": "Старый", "src": "RU", "tgt": "EN", "domain": "medical",
       "segments": [{"id": i + 1, "source": t, "target": "x", "status": "translated"}
                    for i, (t, _idx) in enumerate(units)]}
got, matched = main._map_source_to_segments(units, old["segments"])
# Сегмент СТАРОГО импорта несёт номер страницы из поля: узнаётся по полному
# тексту абзаца, без него строка оглавления не находится.
_old85 = [dict(sg) for sg in old["segments"]]
_old85[6]["source"] = "ГЛАВА ПЕРВАЯ85"
_full = [f for _t, f in main._docx_paragraph_texts(CONTENT)]
check(_full[units[6][1][0]] == "ГЛАВА ПЕРВАЯ85", "полный текст абзаца оглавления несёт номер")
_p85, _m85 = main._map_source_to_segments(units, _old85, _full)
check(_m85 == len(_old85), "старый сегмент с номером узнан по полному тексту: %d" % _m85)
_p0, _m0 = main._map_source_to_segments(units, _old85)
check(_m0 == len(_old85) - 1, "а без полного текста — нет: %d" % _m0)
check(matched == len(old["segments"]) and got == pairs,
      "тот же файл садится на существующий проект без потерь (%d из %d)"
      % (matched, len(old["segments"])))

alien = [{"id": i + 1, "source": "совсем другой текст %d" % i, "target": ""}
         for i in range(len(units))]
_p2, matched_alien = main._map_source_to_segments(units, alien)
check(matched_alien == 0, "чужой файл не садится ни на один сегмент")

# Сдвиг: сегмент в середине переписан — остальные обязаны сесть на свои места
shifted = [dict(s) for s in old["segments"]]
shifted[2]["source"] = "строка, которой в файле нет"
_p3, matched_shift = main._map_source_to_segments(units, shifted)
check(matched_shift == len(units) - 1,
      "пропавшая строка не сдвигает остальные: %d из %d"
      % (matched_shift, len(units)))

# ── подменённый под картой файл ─────────────────────────────────────
doc2 = Document()
doc2.add_paragraph("Другой документ.")
doc2.save(str(docx_path))
try:
    main._generate_export(project, "docx_layout")
    check(False, "подменённый исходник обязан быть отвергнут")
except main.HTTPException as e:
    check("карт" in str(e.detail), "подменённый исходник отвергнут: " + str(e.detail))

# ── чужой файл под тем же номером ───────────────────────────────────
# Номера проектов переиспользуются (id = max + 1), поэтому файл удалённого
# проекта мог бы достаться новому с тем же номером. Отметка в проекте — часть
# опознания: нет её — нет и исходника, чей бы файл ни лежал на диске.
no_mark = dict(project)
no_mark.pop("sourceDocx")
try:
    main._generate_export(no_mark, "docx_layout")
    check(False, "проект без отметки не имеет права взять чужой файл с диска")
except main.HTTPException as e:
    check("не приложен" in str(e.detail),
          "проект без отметки исходника не берёт: " + str(e.detail))

# ── без исходника формата просто нет ────────────────────────────────
docx_path.unlink()
map_path.unlink()
check(main._load_source_map(1) is None, "карта без файла не читается")
try:
    main._generate_export(project, "docx_layout")
    check(False, "без исходника экспорт 1в1 обязан отказать")
except main.HTTPException as e:
    check("не приложен" in str(e.detail), "отказ назван словами: " + str(e.detail))

# ── Поля: что переводим, что оставляем и что просим пересчитать ─────
# Строка оглавления, несущая метку самого поля TOC, — это обычный видимый
# текст внутри результата поля. Раньше её пропускали заодно с номером
# страницы, и в английском документе она одна оставалась русской.
print("\n=== поля и оглавление ===")
BSL = chr(92)
check(main._field_key(" PAGEREF _Toc219883320 " + BSL + "h ") == "PAGEREF",
      "имя поля читается, а ключи-переключатели именем не считаются")
check(main._field_key(" TOC " + BSL + 'o "1-3" ' + BSL + "h ") == "TOC",
      "TOC опознаётся так же")
check(main._field_key("") == "", "пустая инструкция имени не даёт")

fdoc = Document()
fp = fdoc.add_paragraph()


def _piece(par, kind=None, instr=None, text=None):
    r = OxmlElement("w:r")
    if kind:
        f = OxmlElement("w:fldChar")
        f.set(qn("w:fldCharType"), kind)
        r.append(f)
    elif instr is not None:
        i = OxmlElement("w:instrText")
        i.text = instr
        r.append(i)
    else:
        t = OxmlElement("w:t")
        t.text = text
        r.append(t)
    par._p.append(r)


# Ровно то, что лежит в учебнике: первая строка оглавления несёт begin самого
# TOC, внутри неё — заголовок, а номер страницы отдельным полем PAGEREF.
_piece(fp, kind="begin")
_piece(fp, instr=" TOC " + BSL + 'o "1-3" ' + BSL + "h ")
_piece(fp, kind="separate")
_piece(fp, text="СПИСОК СОКРАЩЕНИЙ")
_piece(fp, kind="begin")
_piece(fp, instr=" PAGEREF _Toc1 " + BSL + "h ")
_piece(fp, kind="separate")
_piece(fp, text="3")
_piece(fp, kind="end")

slots, full, dropped = main._para_slots(fp._p, qn)
check([t.text for t, _sig in slots] == ["СПИСОК СОКРАЩЕНИЙ"],
      "текст внутри TOC переводится: %s" % [t.text for t, _s in slots])
check(dropped == "3", "номер страницы остаётся полем: %r" % dropped)
check(full == "СПИСОК СОКРАЩЕНИЙ3", "текст абзаца склеен как при импорте: %r" % full)

# Неизвестное поле — молчим и не пишем: писать в то, чего не понимаешь,
# опаснее, чем пропустить.
up = fdoc.add_paragraph()
_piece(up, kind="begin")
_piece(up, kind="separate")
_piece(up, text="что-то посчитанное")
_piece(up, kind="end")
check(main._para_slots(up._p, qn)[0] == [],
      "поле без прочитанной инструкции считается вычисляемым")

n = main._refresh_fields(fdoc, qn)
check(n == 2, "к пересчёту помечены оба поля (TOC и PAGEREF): %d" % n)
begins = [e for e in fdoc.element.iter(qn("w:fldChar"))
          if e.get(qn("w:fldCharType")) == "begin"]
check(sum(1 for e in begins if e.get(qn("w:dirty")) == "true") == 2,
      "метка стоит именно на begin поля")
upd = fdoc.settings.element.find(qn("w:updateFields"))
check(upd is not None and upd.get(qn("w:val")) == "true",
      "в настройках документа выставлен пересчёт полей при открытии")
order = [c.tag.split("}")[1] for c in fdoc.settings.element]
check("updateFields" in order
      and ("compat" not in order or order.index("updateFields") < order.index("compat")),
      "и стоит там, где его ждёт схема: %s" % order)
main._refresh_fields(fdoc, qn)
check(len(fdoc.settings.element.findall(qn("w:updateFields"))) == 1,
      "повторный вызов не плодит вторую настройку")

shutil.rmtree(TMP, ignore_errors=True)
print("\n" + ("ВСЁ ПРОШЛО" if not fail else "ПРОВАЛЕНО: " + "; ".join(fail)))
sys.exit(1 if fail else 0)
