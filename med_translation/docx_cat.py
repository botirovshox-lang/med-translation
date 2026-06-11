"""
docx_cat.py — Импорт/экспорт DOCX для CAT-рабочего процесса
"""
from pathlib import Path
from docx import Document
from db import create_project, add_segment, get_project, get_segments
import html
import re


def clean_text(text):
    """Очищает текст от HTML-сущностей и лишних пробелов."""
    # Декодируем HTML-сущности (&quot; → ", &amp; → &, и т.д.)
    text = html.unescape(text)
    # Удаляем управляющие символы
    text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)
    # Нормализуем пробелы
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def iter_docx_text_blocks(doc):
    """Итератор по всем текстовым блокам в DOCX (параграфы и ячейки таблиц)."""
    idx = 0
    # Параграфы
    for p in doc.paragraphs:
        text = clean_text(p.text)
        if text:
            yield ('paragraph', idx, p, text)
            idx += 1
    # Таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = clean_text(p.text)
                    if text:
                        yield ('table_cell', idx, p, text)
                        idx += 1


def import_docx_project(uploaded_file, title, source_language='Russian', target_language='English'):
    """
    Импортирует DOCX в проект CAT.
    Разбивает текст на сегменты (параграфы и таблицы).
    Возвращает (project_id, segment_count).
    """
    Path('data/source_docs').mkdir(parents=True, exist_ok=True)
    path = Path('data/source_docs') / uploaded_file.name.replace('/', '_').replace('\\', '_')
    with open(path, 'wb') as f:
        f.write(uploaded_file.getbuffer())

    pid = create_project(title or uploaded_file.name, str(path), source_language, target_language)
    doc = Document(str(path))
    order = 1

    for block_type, block_index, paragraph_obj, text in iter_docx_text_blocks(doc):
        add_segment(pid, order, block_type, block_index, text)
        order += 1

    return pid, order - 1


def replace_paragraph_text(paragraph, new_text):
    """Заменить текст параграфа (сохраняя форматирование первого run)."""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for r in paragraph.runs[1:]:
            r.text = ''
    else:
        paragraph.add_run(new_text)


def export_translated_docx(pid):
    """
    Экспортирует переведённый проект обратно в DOCX.
    Заменяет исходный текст на переводы.
    """
    project = get_project(pid)
    doc = Document(project['source_docx_path'])
    blocks = list(iter_docx_text_blocks(doc))
    segs = get_segments(pid)

    # Заменяем текст блоков на переводы
    for seg, block in zip(segs, blocks):
        target = seg.get('target_text') or seg.get('source_text') or ''
        replace_paragraph_text(block[2], target)

    Path('exports').mkdir(exist_ok=True)
    out = Path('exports') / (project['title'] + '_translated.docx')
    doc.save(str(out))
    return str(out)
